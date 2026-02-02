# -*- coding: utf-8 -*-
"""
Insurance Calculator Module
保险计算器模块 — 主险计算 + 附加险计算
"""

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QLabel,
    QPushButton, QComboBox, QDoubleSpinBox, QSpinBox,
    QScrollArea, QFrame, QSlider, QSplitter, QFileDialog, QLineEdit,
    QListWidget, QListWidgetItem, QDateEdit, QTextEdit,
    QGraphicsDropShadowEffect, QMessageBox, QGroupBox,
    QSizePolicy, QAbstractItemView, QDialog
)
from PyQt5.QtCore import Qt, pyqtSignal, QDate
from PyQt5.QtGui import QFont, QColor
import json
import os
import math
import re
from datetime import date

# 从主脚本导入设计系统
try:
    from Clause_Comparison_Assistant import AnthropicColors, AnthropicFonts, GlassCard
except ImportError:
    # 备用颜色定义
    class AnthropicColors:
        BG_PRIMARY = "#faf9f5"
        BG_CARD = "#f0eee6"
        BG_MINT = "#bcd1ca"
        BG_LAVENDER = "#cbcadb"
        BG_DARK = "#141413"
        ACCENT = "#d97757"
        ACCENT_DARK = "#c6613f"
        ACCENT_LIGHT = "#e8956f"
        TEXT_PRIMARY = "#141413"
        TEXT_SECONDARY = "#6b6960"
        TEXT_TERTIARY = "#57554e"
        TEXT_LIGHT = "#ffffff"
        BORDER = "#d8d6ce"
        BORDER_DARK = "#c0beb6"

    class AnthropicFonts:
        TITLE_LARGE = ("Anthropic Sans", 28)
        TITLE = ("Anthropic Sans", 22)
        TITLE_SMALL = ("Anthropic Sans", 16)
        BODY = ("Anthropic Serif", 14)
        BODY_SMALL = ("Anthropic Serif", 12)

    class GlassCard(QFrame):
        def __init__(self, parent=None, variant="default"):
            super().__init__(parent)
            bg = {"mint": "#bcd1ca", "lavender": "#cbcadb"}.get(variant, "#f0eee6")
            self.setStyleSheet(f"QFrame {{ background: {bg}; border: 1px solid #d8d6ce; border-radius: 12px; }}")


# =============================================
# 数据常量
# =============================================

MC_PRODUCTS = {
    "employerLiability": {
        "productName": "雇主责任险",
        "productType": "liability",
        "amountUnit": "万元",
        "amountLabel": "每人限额",
        "premiumCap": 0.70,
        "formulaText": "固定限额：年保费＝每人限额×基准费率×各项系数乘积×承保人数；工资总额：年保费＝年度工资总额×基准费率×各项系数乘积",
        "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算；短期承保保险费＝年保费×保险期间天数÷365",
        "versions": {
            "original": {
                "label": "雇主责任险费率",
                "baseRates": {
                    "fixed": {"class1": 0.0011, "class2": 0.0017, "class3": 0.0029},
                    "salary": {"class1": 0.0033, "class2": 0.0051, "class3": 0.0085}
                },
                "coefficients": [
                    {
                        "id": "perPersonLimit", "name": "每人赔偿限额调整系数", "applicableTo": ["fixed"],
                        "note": "未列明限额可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤10万元", "min": 1.2, "max": 1.3, "type": "range"},
                            {"parameter": "30万元", "value": 1.1, "type": "fixed"},
                            {"parameter": "50万元", "value": 1.0, "type": "fixed"},
                            {"parameter": "80万元", "value": 0.9, "type": "fixed"},
                            {"parameter": "≥100万元", "min": 0.8, "max": 0.85, "type": "range"}
                        ]
                    },
                    {
                        "id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"],
                        "note": "未列明人数可按线性插值法计算",
                        "rows": [
                            {"parameter": "＜100人", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "[100, 500)人", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "[500, 1000)人", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                            {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                        ]
                    },
                    {
                        "id": "deathDisabilityMonths", "name": "死亡/伤残每人赔偿限额调整系数", "applicableTo": ["salary"],
                        "rows": [
                            {"parameter": "36/48个月", "value": 1.0, "type": "fixed"},
                            {"parameter": "48/60个月", "value": 1.25, "type": "fixed"},
                            {"parameter": "60/72个月", "value": 1.4, "type": "fixed"},
                            {"parameter": "72/84个月", "value": 1.5, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "medicalLimit", "name": "医疗费用每人赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "医疗费用每人赔偿限额÷每人赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "10%", "value": 1.0, "type": "fixed"},
                            {"parameter": "15%", "value": 1.05, "type": "fixed"},
                            {"parameter": "20%", "value": 1.1, "type": "fixed"},
                            {"parameter": "≥25%", "min": 1.15, "max": 1.3, "type": "range"}
                        ]
                    },
                    {
                        "id": "lostWorkLimit", "name": "误工费用每人赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "误工费用每人赔偿限额÷每人赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "10%", "value": 1.0, "type": "fixed"},
                            {"parameter": "15%", "value": 1.05, "type": "fixed"},
                            {"parameter": "≥20%", "min": 1.1, "max": 1.2, "type": "range"}
                        ]
                    },
                    {
                        "id": "perAccidentRatio", "name": "每次事故赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每次事故赔偿限额÷每人赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤3倍", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                            {"parameter": "10倍", "value": 1.05, "type": "fixed"},
                            {"parameter": "≥15倍", "min": 1.1, "max": 1.2, "type": "range"}
                        ]
                    },
                    {
                        "id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "累计赔偿限额÷每次事故赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "1倍", "value": 0.95, "type": "fixed"},
                            {"parameter": "2倍", "value": 1.0, "type": "fixed"},
                            {"parameter": "3倍", "value": 1.05, "type": "fixed"},
                            {"parameter": "≥4倍", "min": 1.1, "max": 1.2, "type": "range"}
                        ]
                    },
                    {
                        "id": "deductible", "name": "免赔额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每次事故医疗费用每人免赔额；未列明免赔额可按线性插值法计算",
                        "rows": [
                            {"parameter": "0元", "value": 1.0, "type": "fixed"},
                            {"parameter": "500元", "value": 0.97, "type": "fixed"},
                            {"parameter": "1000元", "value": 0.95, "type": "fixed"},
                            {"parameter": "≥2000元", "min": 0.8, "max": 0.9, "type": "range"}
                        ]
                    },
                    {
                        "id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                            {"parameter": "后勤人员", "min": 0.9, "max": 1.0, "type": "range"},
                            {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                        ]
                    },
                    {
                        "id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                            {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "managementLevel", "name": "管理水平调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "制度完善，无明显缺陷", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "较完善，存在个别缺陷", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "不完善或存在较多缺陷", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                            {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                            {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                            {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                            {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                            {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "新保", "value": 1.0, "type": "fixed"},
                            {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                            {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                        ]
                    }
                ]
            },
            "v2026": {
                "label": "雇主责任险（2026版）费率",
                "baseRates": {
                    "fixed": {"class1": 0.0012, "class2": 0.0017, "class3": 0.0029},
                    "salary": {"class1": 0.0035, "class2": 0.0051, "class3": 0.0085}
                },
                "coefficients": [
                    {
                        "id": "perPersonLimit", "name": "每人赔偿限额调整系数", "applicableTo": ["fixed"],
                        "note": "每人赔偿限额按死亡/伤残赔偿限额高者取值；未列明限额可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤10万元", "min": 1.2, "max": 1.3, "type": "range"},
                            {"parameter": "30万元", "value": 1.1, "type": "fixed"},
                            {"parameter": "50万元", "value": 1.0, "type": "fixed"},
                            {"parameter": "80万元", "value": 0.9, "type": "fixed"},
                            {"parameter": "≥100万元", "min": 0.8, "max": 0.85, "type": "range"}
                        ]
                    },
                    {
                        "id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"],
                        "note": "未列明人数可按线性插值法计算",
                        "rows": [
                            {"parameter": "＜100人", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "[100, 500)人", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "[500, 1000)人", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                            {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                        ]
                    },
                    {
                        "id": "deathDisabilityMonths", "name": "每人赔偿限额调整系数（工资月数）", "applicableTo": ["salary"],
                        "rows": [
                            {"parameter": "36/48个月", "value": 1.0, "type": "fixed"},
                            {"parameter": "48/60个月", "value": 1.25, "type": "fixed"},
                            {"parameter": "60/72个月", "value": 1.4, "type": "fixed"},
                            {"parameter": "72/84个月", "value": 1.5, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "medicalLimit", "name": "每人医疗费用赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每人医疗费用赔偿限额÷每人赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "10%", "value": 1.0, "type": "fixed"},
                            {"parameter": "15%", "value": 1.05, "type": "fixed"},
                            {"parameter": "20%", "value": 1.1, "type": "fixed"},
                            {"parameter": "≥25%", "min": 1.15, "max": 1.3, "type": "range"}
                        ]
                    },
                    {
                        "id": "lostWorkDaily", "name": "每人每天误工费用金额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "基于雇员月平均工资；未列明金额可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤月工资÷40", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "月工资÷30", "value": 1.0, "type": "fixed"},
                            {"parameter": "≥月工资÷20", "min": 1.05, "max": 1.1, "type": "range"}
                        ]
                    },
                    {
                        "id": "lostWorkDays", "name": "单次及累计赔偿误工费用天数调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "未列明天数可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤120天", "min": 0.95, "max": 0.97, "type": "range"},
                            {"parameter": "180天", "value": 1.0, "type": "fixed"},
                            {"parameter": "240天", "value": 1.03, "type": "fixed"},
                            {"parameter": "≥300天", "min": 1.06, "max": 1.1, "type": "range"}
                        ]
                    },
                    {
                        "id": "lostWorkLimitRatio", "name": "每人误工费用赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每人误工费用赔偿限额÷（每人每天误工费用金额×天数）；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤50%", "min": 0.95, "max": 0.96, "type": "range"},
                            {"parameter": "75%", "value": 0.98, "type": "fixed"},
                            {"parameter": "100%", "value": 1.0, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "累计赔偿限额÷每人赔偿限额；未列明比例可按线性插值法计算",
                        "rows": [
                            {"parameter": "≤5倍", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "10倍", "value": 1.0, "type": "fixed"},
                            {"parameter": "20倍", "value": 1.05, "type": "fixed"},
                            {"parameter": "≥30倍", "min": 1.1, "max": 1.2, "type": "range"}
                        ]
                    },
                    {
                        "id": "deductibleRate", "name": "医疗费用免赔率调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每次事故每人医疗费用免赔率；若同时约定免赔率与免赔额，以两者系数的低者取值",
                        "rows": [
                            {"parameter": "0", "value": 1.0, "type": "fixed"},
                            {"parameter": "10%", "value": 0.97, "type": "fixed"},
                            {"parameter": "20%", "value": 0.94, "type": "fixed"},
                            {"parameter": "30%", "value": 0.91, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "deductibleAmount", "name": "医疗费用免赔额调整系数", "applicableTo": ["fixed", "salary"],
                        "note": "每次事故每人医疗费用免赔额；未列明免赔额可按线性插值法计算",
                        "rows": [
                            {"parameter": "0元", "value": 1.0, "type": "fixed"},
                            {"parameter": "500元", "value": 0.97, "type": "fixed"},
                            {"parameter": "1000元", "value": 0.94, "type": "fixed"},
                            {"parameter": "≥1500元", "min": 0.85, "max": 0.9, "type": "range"}
                        ]
                    },
                    {
                        "id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                            {"parameter": "后勤人员", "min": 0.9, "max": 1.0, "type": "range"},
                            {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                        ]
                    },
                    {
                        "id": "managementLevel", "name": "管理水平调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "制度完善，无明显缺陷", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "较完善，存在个别缺陷", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "不完善或存在较多缺陷", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                            {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                            {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                            {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                            {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                            {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "新保", "value": 1.0, "type": "fixed"},
                            {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                            {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                        ]
                    },
                    {
                        "id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed", "salary"],
                        "rows": [
                            {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                            {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                        ]
                    }
                ]
            },
            "v2016": {
                "label": "雇主责任险（2016版）费率",
                "baseRates": {
                    "fixed": {"class1": 0.00148, "class2": 0.00220, "class3": 0.00366},
                    "salary": {"class1": 0.00444, "class2": 0.00663, "class3": 0.01101}
                },
                "coefficients": [
                    {"id": "perPersonLimit", "name": "每人赔偿限额调整系数", "applicableTo": ["fixed"], "note": "未列明限额可按线性插值法计算", "rows": [
                        {"parameter": "≤10万元", "min": 1.2, "max": 1.3, "type": "range"},
                        {"parameter": "30万元", "value": 1.1, "type": "fixed"},
                        {"parameter": "50万元", "value": 1.0, "type": "fixed"},
                        {"parameter": "80万元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥100万元", "min": 0.8, "max": 0.85, "type": "range"}
                    ]},
                    {"id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"], "note": "未列明人数可按线性插值法计算", "rows": [
                        {"parameter": "＜100人", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "[100, 500)人", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "[500, 1000)人", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                        {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "deathDisabilityMonths", "name": "死亡/伤残每人赔偿限额调整系数", "applicableTo": ["salary"], "rows": [
                        {"parameter": "36/48个月", "value": 1.0, "type": "fixed"},
                        {"parameter": "48/60个月", "value": 1.25, "type": "fixed"},
                        {"parameter": "60/72个月", "value": 1.4, "type": "fixed"},
                        {"parameter": "72/84个月", "value": 1.5, "type": "fixed"}
                    ]},
                    {"id": "medicalLimit", "name": "医疗费用每人赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "note": "医疗费用每人赔偿限额÷每人赔偿限额", "rows": [
                        {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "10%", "value": 1.0, "type": "fixed"},
                        {"parameter": "15%", "value": 1.05, "type": "fixed"},
                        {"parameter": "20%", "value": 1.1, "type": "fixed"},
                        {"parameter": "≥25%", "min": 1.15, "max": 1.3, "type": "range"}
                    ]},
                    {"id": "lostWorkLimit", "name": "误工费用每人赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "10%", "value": 1.0, "type": "fixed"},
                        {"parameter": "15%", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥20%", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "perAccidentRatio", "name": "每次事故赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "≤3倍", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "10倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥15倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "1倍", "value": 0.95, "type": "fixed"},
                        {"parameter": "2倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "3倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥4倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0", "value": 1.0, "type": "fixed"},
                        {"parameter": "10%", "value": 0.9, "type": "fixed"},
                        {"parameter": "20%", "value": 0.8, "type": "fixed"},
                        {"parameter": "30%", "value": 0.7, "type": "fixed"}
                    ]},
                    {"id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0元", "value": 1.0, "type": "fixed"},
                        {"parameter": "2000元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥4000元", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                        {"parameter": "后勤人员", "min": 0.9, "max": 1.0, "type": "range"},
                        {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                    ]},
                    {"id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                        {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                    ]},
                    {"id": "safetySystem", "name": "安全管理制度情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全管理规章制度健全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全管理规章制度较健全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全管理规章制度不健全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "workExperience", "name": "员工工作经验调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "员工工作经验整体较多", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "员工工作经验整体一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "员工工作经验整体较少", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                        {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                        {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                        {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                        {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                        {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "新保", "value": 1.0, "type": "fixed"},
                        {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                        {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                    ]},
                    {"id": "govInspection", "name": "政府安全检查情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "定期对企业进行安全生产检查", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对企业进行安全生产检查", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对企业进行安全生产检查", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "overtime", "name": "员工长时间加班情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "基本没有", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "偶尔有", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "经常有", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "operationCompliance", "name": "员工操作情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "严格按照安全生产制度操作", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "存在个别违反安全生产制度的情况", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在较多违反安全生产制度的情况", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "educationLevel", "name": "员工平均学历情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "automationLevel", "name": "机器设备自动化程度调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "机器设备自动化程度较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "机器设备自动化程度一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "机器设备自动化程度较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]}
                ]
            },
            "vB": {
                "label": "雇主责任险（B款）费率",
                "baseRates": {
                    "fixed": {"class1": 0.0008, "class2": 0.0012, "class3": 0.0020},
                    "salary": {"class1": 0.0023, "class2": 0.0035, "class3": 0.0058}
                },
                "coefficients": [
                    {"id": "compensationStandard", "name": "赔偿标准调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "A表", "value": 1.0, "type": "fixed"},
                        {"parameter": "B表", "value": 1.15, "type": "fixed"}
                    ]},
                    {"id": "perAccidentRatio", "name": "每次事故赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "≤3倍", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "10倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥15倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "1倍", "value": 0.95, "type": "fixed"},
                        {"parameter": "2倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "3倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥4倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0", "value": 1.0, "type": "fixed"},
                        {"parameter": "10%", "value": 0.9, "type": "fixed"},
                        {"parameter": "20%", "value": 0.8, "type": "fixed"},
                        {"parameter": "30%", "value": 0.7, "type": "fixed"}
                    ]},
                    {"id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0元", "value": 1.0, "type": "fixed"},
                        {"parameter": "2000元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥4000元", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                        {"parameter": "后勤人员", "min": 0.9, "max": 1.0, "type": "range"},
                        {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                    ]},
                    {"id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                        {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                    ]},
                    {"id": "safetySystem", "name": "安全管理制度情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全管理规章制度健全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全管理规章制度较健全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全管理规章制度不健全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "workExperience", "name": "员工工作经验调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "员工工作经验整体较多", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "员工工作经验整体一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "员工工作经验整体较少", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                        {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                        {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                        {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                        {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                        {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "新保", "value": 1.0, "type": "fixed"},
                        {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                        {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                    ]},
                    {"id": "govInspection", "name": "政府安全检查情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "定期对企业进行安全生产检查", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对企业进行安全生产检查", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对企业进行安全生产检查", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "overtime", "name": "员工长时间加班情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "基本没有", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "偶尔有", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "经常有", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "operationCompliance", "name": "员工操作情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "严格按照安全生产制度操作", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "存在个别违反安全生产制度的情况", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在较多违反安全生产制度的情况", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "educationLevel", "name": "员工平均学历情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "automationLevel", "name": "机器设备自动化程度调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "机器设备自动化程度较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "机器设备自动化程度一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "机器设备自动化程度较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]}
                ]
            },
            "vD2026": {
                "label": "雇主责任险（D款2026版）费率",
                "premiumCap": 0.70,
                "baseRates": {
                    "fixed": {"class1": 0.00124, "class2": 0.00186, "class3": 0.00310},
                    "salary": {"class1": 0.00350, "class2": 0.00525, "class3": 0.00875}
                },
                "coefficients": [
                    {"id": "deathDisabilityMonths", "name": "死亡/伤残赔偿工资月数调整系数", "applicableTo": ["salary"], "rows": [
                        {"parameter": "36/48个月", "value": 1.0, "type": "fixed"},
                        {"parameter": "48/60个月", "value": 1.25, "type": "fixed"},
                        {"parameter": "60/72个月", "value": 1.4, "type": "fixed"},
                        {"parameter": "72/84个月", "value": 1.5, "type": "fixed"}
                    ]},
                    {"id": "disabilityDeathRatio", "name": "残疾/死亡赔偿限额比值调整系数", "applicableTo": ["fixed"], "note": "每人残疾赔偿限额÷每人死亡赔偿限额；未列明比例可按线性插值法计算", "rows": [
                        {"parameter": "≤0.8", "min": 0.85, "max": 0.9, "type": "range"},
                        {"parameter": "0.9", "value": 0.95, "type": "fixed"},
                        {"parameter": "1.0", "value": 1.0, "type": "fixed"},
                        {"parameter": "1.1", "value": 0.96, "type": "fixed"},
                        {"parameter": "≥1.2", "min": 0.85, "max": 0.92, "type": "range"}
                    ]},
                    {"id": "lostWorkDaysD", "name": "误工费赔偿天数调整系数", "applicableTo": ["fixed"], "note": "每人误工费用最高赔偿天数；未列明天数可按线性插值法计算", "rows": [
                        {"parameter": "≤90天", "min": 0.96, "max": 0.97, "type": "range"},
                        {"parameter": "180天", "value": 0.99, "type": "fixed"},
                        {"parameter": "365天", "value": 1.0, "type": "fixed"}
                    ]},
                    {"id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"], "note": "未列明人数可按线性插值法计算", "rows": [
                        {"parameter": "≤50人", "min": 1.05, "max": 1.1, "type": "range"},
                        {"parameter": "100人", "value": 1.0, "type": "fixed"},
                        {"parameter": "500人", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "medicalLimit", "name": "医疗费用赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "note": "每人医疗费用赔偿限额（万元）；未列明限额可按线性插值法计算", "rows": [
                        {"parameter": "≤1万元", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "2万元", "value": 1.0, "type": "fixed"},
                        {"parameter": "3万元", "value": 1.05, "type": "fixed"},
                        {"parameter": "4万元", "value": 1.1, "type": "fixed"},
                        {"parameter": "≥5万元", "min": 1.15, "max": 1.3, "type": "range"}
                    ]},
                    {"id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed", "salary"], "note": "累计赔偿限额÷每人每次事故赔偿限额；未列明比例可按线性插值法计算", "rows": [
                        {"parameter": "≤2倍", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "8倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥10倍", "min": 1.1, "max": 1.3, "type": "range"}
                    ]},
                    {"id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0", "value": 1.0, "type": "fixed"},
                        {"parameter": "10%", "value": 0.9, "type": "fixed"},
                        {"parameter": "20%", "value": 0.8, "type": "fixed"},
                        {"parameter": "30%", "value": 0.7, "type": "fixed"}
                    ]},
                    {"id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "note": "每次事故免赔额；未列明免赔额可按线性插值法计算", "rows": [
                        {"parameter": "0元", "value": 1.0, "type": "fixed"},
                        {"parameter": "2500元", "value": 0.9, "type": "fixed"},
                        {"parameter": "7500元", "value": 0.8, "type": "fixed"},
                        {"parameter": "≥10000元", "min": 0.7, "max": 0.75, "type": "range"}
                    ]},
                    {"id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "管理人员", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "生产人员", "min": 1.0, "max": 3.0, "type": "range"}
                    ]},
                    {"id": "historicalAccident", "name": "历史事故情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                        {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                        {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                        {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "managementLevel", "name": "管理水平调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "制度完善，无明显缺陷", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "较完善，存在个别缺陷", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "不完善或存在较多缺陷", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                        {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                        {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "＞95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "新保", "value": 1.0, "type": "fixed"},
                        {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                        {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                    ]}
                ]
            },
            "vF2026": {
                "label": "雇主责任险（F款2026版）费率",
                "premiumCap": 0.70,
                "fixedOnly": True,
                "baseRates": {
                    "fixed": {"class1": 0.0021, "class2": 0.0032, "class3": 0.0053}
                },
                "coefficients": [
                    {"id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"], "base": 500, "rows": [
                        {"parameter": "＜100人", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "[100, 500)人", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "[500, 1000)人", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                        {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "legalFeeLimit", "name": "法律费用限额调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "0元", "value": 0.95, "type": "fixed"},
                        {"parameter": "5000元", "value": 1.0, "type": "fixed"},
                        {"parameter": "10000元", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥15000元", "min": 1.10, "max": 1.20, "type": "range"}
                    ]},
                    {"id": "perAccidentRatio", "name": "每次事故赔偿限额调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "≤3倍", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "10倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥15倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "cumulativeRatio", "name": "累计赔偿限额调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "1倍", "value": 0.95, "type": "fixed"},
                        {"parameter": "2倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "3倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥4倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["fixed"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0", "value": 1.0, "type": "fixed"},
                        {"parameter": "10%", "value": 0.9, "type": "fixed"},
                        {"parameter": "20%", "value": 0.8, "type": "fixed"},
                        {"parameter": "30%", "value": 0.7, "type": "fixed"}
                    ]},
                    {"id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["fixed"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0元", "value": 1.0, "type": "fixed"},
                        {"parameter": "2000元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥4000元", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                        {"parameter": "后勤人员", "min": 0.8, "max": 1.0, "type": "range"},
                        {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                    ]},
                    {"id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                        {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                    ]},
                    {"id": "safetySystem", "name": "安全管理制度情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "安全管理规章制度健全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全管理规章制度较健全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全管理规章制度不健全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "workExperience", "name": "员工工作经验调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "员工工作经验整体较多", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "员工工作经验整体一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "员工工作经验整体较少", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                        {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                        {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "＞95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                        {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较多", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"},
                        {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "新保", "value": 1.0, "type": "fixed"},
                        {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                        {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                    ]},
                    {"id": "govInspection", "name": "政府安全检查情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "定期对企业进行安全生产检查", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对企业进行安全生产检查", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对企业进行安全生产检查", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "overtime", "name": "员工长时间加班情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "基本没有", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "偶尔有", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "经常有", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "operationCompliance", "name": "员工操作情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "严格按照安全生产制度操作", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "存在个别违反安全生产制度的情况", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在较多违反安全生产制度的情况", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "educationLevel", "name": "员工平均学历情况调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "automationLevel", "name": "机器设备自动化程度调整系数", "applicableTo": ["fixed"], "rows": [
                        {"parameter": "机器设备自动化程度较高", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "机器设备自动化程度一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "机器设备自动化程度较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]}
                ]
            },
            "vE2026": {
                "label": "雇主责任险（E款2026版）费率",
                "dualRate": True,
                "baseRates": {
                    "fixed": {
                        "death": {"class1": 0.00061, "class2": 0.00085, "class3": 0.00139},
                        "medical": {"class1": 0.00151, "class2": 0.00212, "class3": 0.00350}
                    },
                    "salary": {"class1": 0.00212, "class2": 0.00320, "class3": 0.00532}
                },
                "coefficients": [
                    {"id": "deathDisabilityLimit", "name": "每人死亡残疾责任限额调整系数", "applicableTo": ["fixed"], "note": "仅适用于死亡残疾责任；未列明限额可按线性插值法计算", "rows": [
                        {"parameter": "≤10万元", "min": 1.2, "max": 1.3, "type": "range"},
                        {"parameter": "30万元", "value": 1.1, "type": "fixed"},
                        {"parameter": "50万元", "value": 1.0, "type": "fixed"},
                        {"parameter": "80万元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥100万元", "min": 0.8, "max": 0.85, "type": "range"}
                    ]},
                    {"id": "medicalLimitE", "name": "每人医疗费用责任限额调整系数", "applicableTo": ["fixed"], "note": "仅适用于医疗费用责任；未列明限额可按线性插值法计算", "rows": [
                        {"parameter": "≤1万元", "min": 1.3, "max": 1.5, "type": "range"},
                        {"parameter": "2万元", "value": 1.0, "type": "fixed"},
                        {"parameter": "5万元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥8万元", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "employeeCount", "name": "承保人数调整系数", "applicableTo": ["fixed"], "note": "未列明人数可按线性插值法计算", "rows": [
                        {"parameter": "＜100人", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "[100, 500)人", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "[500, 1000)人", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                        {"parameter": "≥1000人", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "deathDisabilityMonths", "name": "每人人身伤亡责任限额调整系数（工资月数）", "applicableTo": ["salary"], "rows": [
                        {"parameter": "36/48个月", "value": 1.0, "type": "fixed"},
                        {"parameter": "48/60个月", "value": 1.25, "type": "fixed"},
                        {"parameter": "60/72个月", "value": 1.4, "type": "fixed"},
                        {"parameter": "72/84个月", "value": 1.5, "type": "fixed"}
                    ]},
                    {"id": "medicalLimitSalary", "name": "每人医疗费用责任限额调整系数", "applicableTo": ["salary"], "note": "每人医疗费用责任限额÷每人人身伤亡责任限额；未列明比例可按线性插值法计算", "rows": [
                        {"parameter": "≤5%", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "10%", "value": 1.0, "type": "fixed"},
                        {"parameter": "15%", "value": 1.05, "type": "fixed"},
                        {"parameter": "20%", "value": 1.1, "type": "fixed"},
                        {"parameter": "≥25%", "min": 1.15, "max": 1.3, "type": "range"}
                    ]},
                    {"id": "perAccidentRatio", "name": "每次事故责任限额调整系数", "applicableTo": ["fixed", "salary"], "note": "每次事故责任限额÷每人人身伤亡责任限额；未列明比例可按线性插值法计算", "rows": [
                        {"parameter": "≤3倍", "min": 0.9, "max": 0.95, "type": "range"},
                        {"parameter": "5倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "10倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥15倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "cumulativeRatio", "name": "累计责任限额调整系数", "applicableTo": ["fixed", "salary"], "note": "累计责任限额÷每次事故责任限额；未列明比例可按线性插值法计算", "rows": [
                        {"parameter": "1倍", "value": 0.95, "type": "fixed"},
                        {"parameter": "2倍", "value": 1.0, "type": "fixed"},
                        {"parameter": "3倍", "value": 1.05, "type": "fixed"},
                        {"parameter": "≥4倍", "min": 1.1, "max": 1.2, "type": "range"}
                    ]},
                    {"id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "rows": [
                        {"parameter": "0", "value": 1.0, "type": "fixed"},
                        {"parameter": "10%", "value": 0.9, "type": "fixed"},
                        {"parameter": "20%", "value": 0.8, "type": "fixed"},
                        {"parameter": "30%", "value": 0.7, "type": "fixed"}
                    ]},
                    {"id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["fixed", "salary"], "linkedGroup": "deductible", "note": "每次事故免赔额；未列明免赔额可按线性插值法计算", "rows": [
                        {"parameter": "0元", "value": 1.0, "type": "fixed"},
                        {"parameter": "2000元", "value": 0.9, "type": "fixed"},
                        {"parameter": "≥4000元", "min": 0.7, "max": 0.8, "type": "range"}
                    ]},
                    {"id": "employeeCategory", "name": "雇员类别调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "管理人员", "min": 0.7, "max": 0.8, "type": "range"},
                        {"parameter": "后勤人员", "min": 0.9, "max": 1.0, "type": "range"},
                        {"parameter": "一线操作人员", "min": 1.0, "max": 2.0, "type": "range"}
                    ]},
                    {"id": "managementLevel", "name": "管理水平调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "制度完善，无明显缺陷", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "较完善，存在个别缺陷", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "不完善或存在较多缺陷", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "[0, 20%]", "min": 0.5, "max": 0.6, "type": "range"},
                        {"parameter": "(20%, 45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                        {"parameter": "(45%, 70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "(70%, 95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "hazardInspection", "name": "企业隐患排查整改调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "无隐患", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "整改完成", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "存在重大隐患且未整改", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "historicalAccident", "name": "历史事故与损失情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                        {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                        {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                        {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                        {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyTraining", "name": "员工安全教育培训调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "每年定期对员工进行安全教育和培训", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "不定期对员工进行安全教育和培训", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "较少对员工进行安全教育和培训", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "safetyEquipment", "name": "安全设施和装备配置情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "安全设施和装备配置齐全", "min": 0.7, "max": 1.0, "type": "range"},
                        {"parameter": "安全设施和装备配置较齐全", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                        {"parameter": "安全设施和装备配置不齐全", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                    ]},
                    {"id": "renewal", "name": "续保调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "新保", "value": 1.0, "type": "fixed"},
                        {"parameter": "续保一年", "value": 0.95, "type": "fixed"},
                        {"parameter": "续保两年及以上", "min": 0.8, "max": 0.9, "type": "range"}
                    ]},
                    {"id": "workInjuryInsurance", "name": "工伤保险情况调整系数", "applicableTo": ["fixed", "salary"], "rows": [
                        {"parameter": "已投保工伤保险", "value": 1.0, "type": "fixed"},
                        {"parameter": "未投保工伤保险", "value": 1.2, "type": "fixed"}
                    ]}
                ]
            }
        }
    },
    "propertyAllRisk": {
        "productName": "财产一切险",
        "productType": "property",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": 0.70,
        "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算；短期承保保险费根据条款所附短期费率表计收",
        "versions": {
            "original": {
                "label": "财产一切险费率",
                "baseRates": { "default": 0.0020 },
                "coefficients": [
                    { "id": "industry", "name": "行业类别调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "农、林、牧、渔业", "min": 1.2, "max": 2.5, "type": "range" },
                        { "parameter": "采矿业", "min": 1.0, "max": 2.0, "type": "range" },
                        { "parameter": "制造业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "电力、热力、燃气及水生产和供应业，水利、环境和公共设施管理业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "建筑业", "min": 0.8, "max": 1.5, "type": "range" },
                        { "parameter": "批发和零售业，交通运输、仓储和邮政业", "min": 1.0, "max": 2.5, "type": "range" },
                        { "parameter": "住宿和餐饮业，文化、体育和娱乐业", "min": 0.6, "max": 1.5, "type": "range" },
                        { "parameter": "信息传输、软件和信息技术服务业，金融业，房地产业，租赁和商务服务业，科学研究和技术服务业，居民服务、修理和其他服务业，教育，卫生和社会工作，公共管理、社会保障和社会组织，国际组织", "min": 0.5, "max": 1.0, "type": "range" }
                    ]},
                    { "id": "insuredAmount", "name": "保险金额调整系数", "applicableTo": ["all"],
                      "note": "保险金额单位为亿元；上表范围内未列明的保险金额对应的调整系数可按线性插值法计算", "rows": [
                        { "parameter": "≤0.1亿元", "min": 1.2, "max": 1.3, "type": "range" },
                        { "parameter": "0.5亿元", "value": 1.1, "type": "fixed" },
                        { "parameter": "1亿元", "value": 1.0, "type": "fixed" },
                        { "parameter": "2亿元", "value": 0.95, "type": "fixed" },
                        { "parameter": "5亿元", "value": 0.9, "type": "fixed" },
                        { "parameter": "10亿元", "value": 0.8, "type": "fixed" },
                        { "parameter": "50亿元", "value": 0.7, "type": "fixed" },
                        { "parameter": "≥100亿元", "min": 0.5, "max": 0.6, "type": "range" }
                    ]},
                    { "id": "naturalDisaster", "name": "自然灾害风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.5, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "terrain", "name": "地势调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较高", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较低", "min": 1.2, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "fireExplosion", "name": "火灾爆炸隐患调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "firePrevention", "name": "防火措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "较完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "storageType", "name": "存储物品类型调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "易燃易爆物品", "value": 1.5, "type": "fixed" },
                        { "parameter": "可燃物", "value": 1.2, "type": "fixed" },
                        { "parameter": "难燃或不燃物", "value": 1.0, "type": "fixed" }
                    ]},
                    { "id": "explosiveStorage", "name": "易燃易爆物品存放位置调整系数", "applicableTo": ["all"],
                      "note": "若无易燃易爆物品则本调整系数取值为1.0", "rows": [
                        { "parameter": "是（存放在危险品仓库中）", "value": 1.0, "type": "fixed" },
                        { "parameter": "否（未存放在危险品仓库中）", "value": 1.2, "type": "fixed" }
                    ]},
                    { "id": "processRisk", "name": "生产工艺风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "safetySystem", "name": "安全生产制度与措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "安全生产制度与措施完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "安全生产制度与措施较完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "安全生产制度与措施不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "renewal", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保一年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保两年", "value": 0.85, "type": "fixed" },
                        { "parameter": "续保三年及以上", "value": 0.8, "type": "fixed" }
                    ]},
                    { "id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "每次事故免赔额（万元）；未列明的免赔额可按线性插值法计算；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "0.1万元", "value": 0.98, "type": "fixed" },
                        { "parameter": "0.2万元", "value": 0.97, "type": "fixed" },
                        { "parameter": "0.5万元", "value": 0.94, "type": "fixed" },
                        { "parameter": "1万元", "value": 0.9, "type": "fixed" },
                        { "parameter": "2万元", "value": 0.85, "type": "fixed" },
                        { "parameter": "5万元", "value": 0.78, "type": "fixed" },
                        { "parameter": "10万元", "value": 0.71, "type": "fixed" },
                        { "parameter": "20万元", "value": 0.64, "type": "fixed" },
                        { "parameter": "50万元", "value": 0.53, "type": "fixed" },
                        { "parameter": "≥100万元", "min": 0.4, "max": 0.46, "type": "range" }
                    ]},
                    { "id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "每次事故免赔率；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "10%", "value": 0.9, "type": "fixed" },
                        { "parameter": "20%", "value": 0.8, "type": "fixed" },
                        { "parameter": "30%", "value": 0.7, "type": "fixed" }
                    ]},
                    { "id": "lossHistory", "name": "历史事故与损失情况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极少", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "buildingStructure", "name": "建筑结构调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "钢筋混凝土结构", "min": 0.8, "max": 0.85, "type": "range" },
                        { "parameter": "钢结构", "min": 0.85, "max": 1.0, "type": "range" },
                        { "parameter": "砖木结构", "min": 1.05, "max": 1.1, "type": "range" },
                        { "parameter": "其他", "min": 1.1, "max": 1.2, "type": "range" }
                    ]},
                    { "id": "fireStation", "name": "公共消防队调整系数", "applicableTo": ["all"],
                      "note": "公共消防队到达标的所在地需要时间（分钟）", "rows": [
                        { "parameter": "≤15分钟", "min": 0.8, "max": 1.0, "type": "range" },
                        { "parameter": "（15，30]分钟", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞30分钟", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "propertyComprehensive": {
        "productName": "财产综合险",
        "productType": "property",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": 0.70,
        "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算；短期承保保险费根据条款所附短期费率表计收",
        "versions": {
            "original": {
                "label": "财产综合险费率",
                "baseRates": { "default": 0.0017 },
                "coefficients": [
                    { "id": "industry", "name": "行业类别调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "农、林、牧、渔业", "min": 1.2, "max": 2.5, "type": "range" },
                        { "parameter": "采矿业", "min": 1.0, "max": 2.0, "type": "range" },
                        { "parameter": "制造业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "电力、热力、燃气及水生产和供应业，水利、环境和公共设施管理业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "建筑业", "min": 0.8, "max": 1.5, "type": "range" },
                        { "parameter": "批发和零售业，交通运输、仓储和邮政业", "min": 1.0, "max": 2.5, "type": "range" },
                        { "parameter": "住宿和餐饮业，文化、体育和娱乐业", "min": 0.6, "max": 1.5, "type": "range" },
                        { "parameter": "信息传输、软件和信息技术服务业，金融业，房地产业，租赁和商务服务业，科学研究和技术服务业，居民服务、修理和其他服务业，教育，卫生和社会工作，公共管理、社会保障和社会组织，国际组织", "min": 0.5, "max": 1.0, "type": "range" }
                    ]},
                    { "id": "insuredAmount", "name": "保险金额调整系数", "applicableTo": ["all"],
                      "note": "保险金额单位为亿元；上表范围内未列明的保险金额对应的调整系数可按线性插值法计算", "rows": [
                        { "parameter": "≤0.1亿元", "min": 1.2, "max": 1.3, "type": "range" },
                        { "parameter": "0.5亿元", "value": 1.1, "type": "fixed" },
                        { "parameter": "1亿元", "value": 1.0, "type": "fixed" },
                        { "parameter": "2亿元", "value": 0.95, "type": "fixed" },
                        { "parameter": "5亿元", "value": 0.9, "type": "fixed" },
                        { "parameter": "10亿元", "value": 0.8, "type": "fixed" },
                        { "parameter": "50亿元", "value": 0.7, "type": "fixed" },
                        { "parameter": "≥100亿元", "min": 0.5, "max": 0.6, "type": "range" }
                    ]},
                    { "id": "naturalDisaster", "name": "自然灾害风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.5, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "terrain", "name": "地势调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较高", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较低", "min": 1.2, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "fireExplosion", "name": "火灾爆炸隐患调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "firePrevention", "name": "防火措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "较完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "storageType", "name": "存储物品类型调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "易燃易爆物品", "value": 1.5, "type": "fixed" },
                        { "parameter": "可燃物", "value": 1.2, "type": "fixed" },
                        { "parameter": "难燃或不燃物", "value": 1.0, "type": "fixed" }
                    ]},
                    { "id": "explosiveStorage", "name": "易燃易爆物品存放位置调整系数", "applicableTo": ["all"],
                      "note": "若无易燃易爆物品则本调整系数取值为1.0", "rows": [
                        { "parameter": "是（存放在危险品仓库中）", "value": 1.0, "type": "fixed" },
                        { "parameter": "否（未存放在危险品仓库中）", "value": 1.2, "type": "fixed" }
                    ]},
                    { "id": "processRisk", "name": "生产工艺风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "safetySystem", "name": "安全生产制度与措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "安全生产制度与措施完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "安全生产制度与措施较完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "安全生产制度与措施不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "renewal", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保一年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保两年", "value": 0.85, "type": "fixed" },
                        { "parameter": "续保三年及以上", "value": 0.8, "type": "fixed" }
                    ]},
                    { "id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "每次事故免赔额（万元）；未列明的免赔额可按线性插值法计算；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "0.1万元", "value": 0.98, "type": "fixed" },
                        { "parameter": "0.2万元", "value": 0.97, "type": "fixed" },
                        { "parameter": "0.5万元", "value": 0.94, "type": "fixed" },
                        { "parameter": "1万元", "value": 0.9, "type": "fixed" },
                        { "parameter": "2万元", "value": 0.85, "type": "fixed" },
                        { "parameter": "5万元", "value": 0.78, "type": "fixed" },
                        { "parameter": "10万元", "value": 0.71, "type": "fixed" },
                        { "parameter": "20万元", "value": 0.64, "type": "fixed" },
                        { "parameter": "50万元", "value": 0.53, "type": "fixed" },
                        { "parameter": "≥100万元", "min": 0.4, "max": 0.46, "type": "range" }
                    ]},
                    { "id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "每次事故免赔率；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "10%", "value": 0.9, "type": "fixed" },
                        { "parameter": "20%", "value": 0.8, "type": "fixed" },
                        { "parameter": "30%", "value": 0.7, "type": "fixed" }
                    ]},
                    { "id": "lossHistory", "name": "历史事故与损失情况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极少", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "buildingStructure", "name": "建筑结构调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "钢筋混凝土结构", "min": 0.8, "max": 0.85, "type": "range" },
                        { "parameter": "钢结构", "min": 0.85, "max": 1.0, "type": "range" },
                        { "parameter": "砖木结构", "min": 1.05, "max": 1.1, "type": "range" },
                        { "parameter": "其他", "min": 1.1, "max": 1.2, "type": "range" }
                    ]},
                    { "id": "fireStation", "name": "公共消防队调整系数", "applicableTo": ["all"],
                      "note": "公共消防队到达标的所在地需要时间（分钟）", "rows": [
                        { "parameter": "≤15分钟", "min": 0.8, "max": 1.0, "type": "range" },
                        { "parameter": "（15，30]分钟", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞30分钟", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "machineryBreakdown": {
        "productName": "机器损坏保险",
        "productType": "property",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": None,
        "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "短期承保保险费按条款所附短期费率表计收",
        "versions": {
            "original": {
                "label": "机器损坏保险费率",
                "baseRates": { "default": 0.0020 },
                "coefficients": [
                    { "id": "industry", "name": "行业调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "农、林、牧、渔业", "min": 1.2, "max": 2.5, "type": "range" },
                        { "parameter": "采矿业", "min": 1.0, "max": 2.0, "type": "range" },
                        { "parameter": "制造业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "电力、热力、燃气及水生产和供应业，水利、环境和公共设施管理业", "min": 0.5, "max": 1.5, "type": "range" },
                        { "parameter": "建筑业", "min": 0.8, "max": 1.5, "type": "range" },
                        { "parameter": "批发和零售业，交通运输、仓储和邮政业", "min": 1.0, "max": 2.5, "type": "range" },
                        { "parameter": "住宿和餐饮业，文化、体育和娱乐业", "min": 0.6, "max": 1.5, "type": "range" },
                        { "parameter": "信息传输、软件和信息技术服务业，金融业，房地产业，租赁和商务服务业，科学研究和技术服务业，居民服务、修理和其他服务业，教育，卫生和社会工作，公共管理、社会保障和社会组织，国际组织", "min": 0.5, "max": 1.0, "type": "range" }
                    ]},
                    { "id": "insuredAmount", "name": "保险金额调整系数", "applicableTo": ["all"],
                      "note": "保险金额单位为千万元", "rows": [
                        { "parameter": "＜1千万元", "value": 1.2, "type": "fixed" },
                        { "parameter": "[1，5）千万元", "value": 1.1, "type": "fixed" },
                        { "parameter": "[5，10）千万元", "value": 1.0, "type": "fixed" },
                        { "parameter": "[10，20）千万元", "value": 0.95, "type": "fixed" },
                        { "parameter": "[20，50）千万元", "value": 0.9, "type": "fixed" },
                        { "parameter": "[50，100）千万元", "value": 0.8, "type": "fixed" },
                        { "parameter": "[100，500）千万元", "value": 0.7, "type": "fixed" },
                        { "parameter": "≥500千万元", "value": 0.6, "type": "fixed" }
                    ]},
                    { "id": "lossRatio", "name": "赔付率调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "managementLevel", "name": "管理水平调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "管理制度和措施完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "管理制度和措施较为完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "管理制度和措施不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "operatorQuality", "name": "操作人员素质调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "操作人员素质高、操作经验丰富", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "操作人员素质较高、操作经验较为丰富", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "操作人员素质不高或者操作经验较少", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "deductibleAmount", "name": "免赔额调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "每次事故免赔额（万元）；未列明的免赔额可按线性插值法计算；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "0.1万元", "value": 0.98, "type": "fixed" },
                        { "parameter": "0.2万元", "value": 0.95, "type": "fixed" },
                        { "parameter": "0.5万元", "value": 0.9, "type": "fixed" },
                        { "parameter": "1万元", "value": 0.85, "type": "fixed" },
                        { "parameter": "2万元", "value": 0.8, "type": "fixed" },
                        { "parameter": "5万元", "value": 0.75, "type": "fixed" },
                        { "parameter": "10万元", "value": 0.7, "type": "fixed" }
                    ]},
                    { "id": "deductibleRate", "name": "免赔率调整系数", "applicableTo": ["all"], "linkedGroup": "deductible",
                      "note": "若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "5%", "value": 0.95, "type": "fixed" },
                        { "parameter": "10%", "value": 0.9, "type": "fixed" },
                        { "parameter": "15%", "value": 0.85, "type": "fixed" },
                        { "parameter": "20%", "value": 0.8, "type": "fixed" },
                        { "parameter": "25%", "value": 0.75, "type": "fixed" },
                        { "parameter": "30%", "value": 0.7, "type": "fixed" },
                        { "parameter": "35%", "value": 0.65, "type": "fixed" }
                    ]},
                    { "id": "renewal", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保一年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保两年及以上", "min": 0.8, "max": 0.85, "type": "range" }
                    ]},
                    { "id": "machineRisk", "name": "机器设备风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "损坏或故障的可能性较低", "min": 0.7, "max": 1.0, "maxExclusive": True, "type": "range" },
                        { "parameter": "损坏或故障的可能性中等", "min": 1.0, "max": 1.2, "maxExclusive": True, "type": "range" },
                        { "parameter": "损坏或故障的可能性较高", "min": 1.2, "max": 1.5, "type": "range" }
                    ]},
                    { "id": "maintenance", "name": "维修保养状况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "定期对设备进行维修保养", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "不定期对设备进行维修保养", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较少对设备进行维修保养", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "equipmentOrigin", "name": "设备产地调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "国产", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "进口", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "equipmentAge", "name": "设备使用时间调整系数", "applicableTo": ["all"],
                      "note": "设备已使用时间（年）", "rows": [
                        { "parameter": "≤5年", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "（5，10]年", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞10年", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "techMaturity", "name": "设备技术成熟度调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "设备所使用的技术非常成熟", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "设备所使用的技术较为成熟", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "设备所使用的技术不成熟", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "usageIntensity", "name": "设备使用强度调整系数", "applicableTo": ["all"],
                      "note": "设备平均每日使用时间占比", "rows": [
                        { "parameter": "≤20%", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "（20%，50%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞50%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "compositeHengli2025": {
        "productName": "突发和意外的直接物质损失一切险(恒力项目专用2025版)",
        "productType": "composite",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": None,
        "subRisks": ["materialDamage", "machineryBreakdown"],
        "subRiskLabels": { "materialDamage": "物质损失或损坏一切险", "machineryBreakdown": "机器损坏" },
        "formulaText": "保险费＝物质损失或损坏一切险保险费＋机器损坏保险费",
        "formulaNote": "物质损失或损坏一切险年保险费＝物质损失保险金额×物质损失基准费率×适用系数乘积；机器损坏年保险费＝机器损坏保险金额×机器损坏基准费率×适用系数乘积；短期承保保险费＝年保险费×保险期间天数÷365",
        "versions": {
            "hengli2025": {
                "label": "恒力2025版费率",
                "baseRates": { "materialDamage": 0.00089, "machineryBreakdown": 0.00362 },
                "coefficients": [
                    { "id": "lossRatioHL", "name": "赔付率调整系数", "applicableTo": ["materialDamage", "machineryBreakdown"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "riskMgmtHL", "name": "风险防范与管理水平调整系数", "applicableTo": ["materialDamage", "machineryBreakdown"], "rows": [
                        { "parameter": "较高", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "deductAmountHL", "name": "免赔额调整系数", "applicableTo": ["materialDamage", "machineryBreakdown"], "linkedGroup": "deductibleHL",
                      "note": "免赔额（元）；未列明的免赔额可按线性插值法计算；若同时设定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "1000元", "value": 0.98, "type": "fixed" },
                        { "parameter": "2000元", "value": 0.96, "type": "fixed" },
                        { "parameter": "5000元", "value": 0.93, "type": "fixed" },
                        { "parameter": "1万元", "value": 0.88, "type": "fixed" },
                        { "parameter": "2万元", "value": 0.83, "type": "fixed" },
                        { "parameter": "5万元", "value": 0.73, "type": "fixed" },
                        { "parameter": "10万元", "value": 0.63, "type": "fixed" }
                    ]},
                    { "id": "deductRateHL", "name": "免赔率调整系数", "applicableTo": ["materialDamage", "machineryBreakdown"], "linkedGroup": "deductibleHL",
                      "note": "若同时设定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "5%", "value": 0.95, "type": "fixed" },
                        { "parameter": "10%", "value": 0.9, "type": "fixed" },
                        { "parameter": "15%", "value": 0.85, "type": "fixed" },
                        { "parameter": "20%", "value": 0.8, "type": "fixed" },
                        { "parameter": "25%", "value": 0.75, "type": "fixed" },
                        { "parameter": "30%", "value": 0.7, "type": "fixed" },
                        { "parameter": "35%", "value": 0.65, "type": "fixed" }
                    ]},
                    { "id": "operatorHL", "name": "操作人员调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分", "rows": [
                        { "parameter": "较高", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "maintenanceHL", "name": "机器设备维修保养状况调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分；评估项：1.有无配备专业维修人员 2.是否进行日常维修保养 3.是否有检查维修保养计划 4.大修间隔是否低于半年 5.损坏/故障时是否有应急修理措施", "rows": [
                        { "parameter": "全部正向达标", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一项未达标", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "超过一项未达标", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "machineRiskHL", "name": "机器设备风险调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分", "rows": [
                        { "parameter": "损坏或故障的可能性较低", "min": 0.7, "max": 1.0, "maxExclusive": True, "type": "range" },
                        { "parameter": "损坏或故障的可能性中等", "min": 1.0, "max": 1.2, "maxExclusive": True, "type": "range" },
                        { "parameter": "损坏或故障的可能性较高", "min": 1.2, "max": 1.5, "type": "range" }
                    ]},
                    { "id": "equipOriginHL", "name": "设备产地调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分", "rows": [
                        { "parameter": "国产", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "进口", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "equipAgeHL", "name": "设备使用时间调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分；设备已使用时间（年）", "rows": [
                        { "parameter": "≤5年", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "（5，10]年", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞10年", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "techMaturityHL", "name": "设备技术成熟度调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分", "rows": [
                        { "parameter": "设备所使用的技术非常成熟", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "设备所使用的技术较为成熟", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "设备所使用的技术不成熟", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "usageIntensityHL", "name": "设备使用强度调整系数", "applicableTo": ["machineryBreakdown"],
                      "note": "仅适用于机器损坏部分；设备平均每日使用时间占比", "rows": [
                        { "parameter": "≤20%", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "（20%，50%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞50%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "disasterRiskHL", "name": "标的所在地灾害风险调整系数", "applicableTo": ["materialDamage"],
                      "note": "仅适用于物质损害或损坏一切险部分", "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "fireEquipHL", "name": "消防设施调整系数", "applicableTo": ["materialDamage"],
                      "note": "仅适用于物质损害或损坏一切险部分；评估项：1.防雷检测是否合格 2.防静电装置是否符合要求 3.是否安装自动检测和火灾报警系统 4.是否安装灭火喷淋系统 5.是否配置相应消防设备设施和灭火药剂", "rows": [
                        { "parameter": "全部正向达标", "min": 0.9, "max": 1.0, "type": "range" },
                        { "parameter": "一项未达标", "min": 1.0, "max": 1.1, "minExclusive": True, "type": "range" },
                        { "parameter": "超过一项未达标", "min": 1.1, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "fireProofHL", "name": "建筑防火设施规范调整系数", "applicableTo": ["materialDamage"],
                      "note": "仅适用于物质损害或损坏一切险部分；评估项：1.仓库防火等级是否达标 2.爆炸品等是否专库储存 3.通风和温度调节是否满足要求 4.垛距、通道、墙距是否达标", "rows": [
                        { "parameter": "全部正向达标", "min": 0.9, "max": 1.0, "type": "range" },
                        { "parameter": "一项未达标", "min": 1.0, "max": 1.1, "minExclusive": True, "type": "range" },
                        { "parameter": "超过一项未达标", "min": 1.1, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "electricHL", "name": "电气线路调整系数", "applicableTo": ["materialDamage"],
                      "note": "仅适用于物质损害或损坏一切险部分；评估项：1.有无电线乱接乱拉 2.电线电缆使用是否超过5年 3.线路布置是否合理 4.电气线路是否有老化现象 5.库房中是否有叉车充电现象", "rows": [
                        { "parameter": "全部正向达标", "min": 0.9, "max": 1.0, "type": "range" },
                        { "parameter": "一项未达标", "min": 1.0, "max": 1.1, "minExclusive": True, "type": "range" },
                        { "parameter": "超过一项未达标", "min": 1.1, "max": 1.3, "minExclusive": True, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "businessInterruption": {
        "productName": "营业中断保险",
        "productType": "interruption",
        "amountUnit": "元",
        "amountLabel": "毛利润损失保险金额",
        "premiumCap": None,
        "formulaText": "年保险费＝毛利润损失保险金额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "短期承保保险费按条款所附短期费率表计收",
        "versions": {
            "original": {
                "label": "营业中断保险费率",
                "baseRates": { "default": 0.0011 },
                "coefficients": [
                    { "id": "maxIndemnityPeriod", "name": "最大赔偿期调整系数", "applicableTo": ["all"],
                      "note": "上表范围内未列明的最大赔偿期对应的调整系数可按线性插值法计算", "rows": [
                        { "parameter": "≤1个月", "value": 0.5, "type": "fixed" },
                        { "parameter": "3个月", "value": 0.6, "type": "fixed" },
                        { "parameter": "6个月", "value": 0.8, "type": "fixed" },
                        { "parameter": "9个月", "value": 0.9, "type": "fixed" },
                        { "parameter": "12个月", "value": 1.0, "type": "fixed" },
                        { "parameter": "18个月", "value": 0.8, "type": "fixed" },
                        { "parameter": "24个月", "value": 0.7, "type": "fixed" },
                        { "parameter": "≥36个月", "value": 0.6, "type": "fixed" }
                    ]},
                    { "id": "materialContract", "name": "物质损失保险合同调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "财产基本险", "value": 0.7, "type": "fixed" },
                        { "parameter": "财产综合险", "value": 1.0, "type": "fixed" },
                        { "parameter": "财产一切险", "value": 1.1, "type": "fixed" },
                        { "parameter": "其他", "min": 0.7, "max": 1.3, "type": "range" }
                    ]},
                    { "id": "naturalDisasterBI", "name": "自然灾害风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极低", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很高", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "accidentRiskBI", "name": "意外事故风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极低", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很高", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "recoveryAbility", "name": "恢复生产能力调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较强", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较弱", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossRatioBI", "name": "赔付率调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "renewalBI", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保1年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保2年", "value": 0.8, "type": "fixed" },
                        { "parameter": "续保3年及以上", "value": 0.7, "type": "fixed" }
                    ]},
                    { "id": "deductAmountBI", "name": "免赔额调整系数", "applicableTo": ["all"], "linkedGroup": "deductibleBI",
                      "note": "免赔额以毛利润额天数计；未列明的可按线性插值法计算；若同时约定免赔额和免赔期，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.1, "type": "fixed" },
                        { "parameter": "4天的毛利润额", "value": 1.0, "type": "fixed" },
                        { "parameter": "8天的毛利润额", "value": 0.9, "type": "fixed" },
                        { "parameter": "12天及以上的毛利润额", "min": 0.7, "max": 0.8, "type": "range" }
                    ]},
                    { "id": "deductPeriodBI", "name": "免赔期调整系数", "applicableTo": ["all"], "linkedGroup": "deductibleBI",
                      "note": "免赔期（天）；若同时约定免赔额和免赔期，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0天", "value": 1.1, "type": "fixed" },
                        { "parameter": "4天", "value": 1.0, "type": "fixed" },
                        { "parameter": "8天", "value": 0.9, "type": "fixed" },
                        { "parameter": "≥12天", "min": 0.7, "max": 0.8, "type": "range" }
                    ]},
                    { "id": "riskPreventionBI", "name": "风险防范措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较好", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较差", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "auditExpenseBI", "name": "审计费用赔偿限额调整系数", "applicableTo": ["all"],
                      "note": "审计费用赔偿限额÷毛利润损失保险金额；未列明的可按线性插值法计算", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "5%", "value": 1.05, "type": "fixed" },
                        { "parameter": "≥10%", "min": 1.1, "max": 1.2, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "cashComprehensive": {
        "productName": "现金综合保险",
        "productType": "property",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": None,
        "extraFields": ["dailyCashTurnover"],
        "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "短期承保保险费＝年保险费×保险期间天数÷365",
        "versions": {
            "original": {
                "label": "现金综合保险费率",
                "baseRates": { "default": 0.00082 },
                "coefficients": [
                    { "id": "turnoverRatio", "name": "营业额调整系数", "applicableTo": ["all"],
                      "note": "保险金额÷平均每日现金营业额", "rows": [
                        { "parameter": "＜1", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range" },
                        { "parameter": "[1，2）", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "[2，3）", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range" },
                        { "parameter": "≥3", "min": 0.7, "max": 0.8, "type": "range" }
                    ]},
                    { "id": "subjectType", "name": "标的类型调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "否（不包含运输途中的现金）", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "是（包含运输途中的现金）", "min": 1.0, "max": 1.5, "type": "range" }
                    ]},
                    { "id": "transportRisk", "name": "运输风险调整系数", "applicableTo": ["all"],
                      "note": "若不承保运输途中的现金，则本调整系数取值为1.0", "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "safetyMgmt", "name": "安全管理制度与措施调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "安全管理制度与措施完善", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "安全管理制度与措施较完善", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "安全管理制度与措施不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossHistoryCash", "name": "历史事故与损失情况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极少", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "renewalCash", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保一年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保两年", "value": 0.85, "type": "fixed" },
                        { "parameter": "续保三年及以上", "value": 0.8, "type": "fixed" }
                    ]},
                    { "id": "deductAmountCash", "name": "免赔额调整系数", "applicableTo": ["all"], "linkedGroup": "deductibleCash",
                      "note": "每次事故免赔额（元）；未列明的可按线性插值法计算；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "2000元", "value": 0.9, "type": "fixed" },
                        { "parameter": "5000元", "value": 0.8, "type": "fixed" },
                        { "parameter": "≥10000元", "min": 0.6, "max": 0.7, "type": "range" }
                    ]},
                    { "id": "deductRateCash", "name": "免赔率调整系数", "applicableTo": ["all"], "linkedGroup": "deductibleCash",
                      "note": "每次事故免赔率；若同时约定免赔额和免赔率，以两者调整系数的低者取值", "rows": [
                        { "parameter": "0", "value": 1.0, "type": "fixed" },
                        { "parameter": "10%", "value": 0.9, "type": "fixed" },
                        { "parameter": "20%", "value": 0.8, "type": "fixed" },
                        { "parameter": "30%", "value": 0.7, "type": "fixed" }
                    ]}
                ]
            }
        }
    },
    "jewelryComprehensive": {
        "productName": "珠宝商综合保险",
        "productType": "jewelry",
        "amountUnit": "元",
        "amountLabel": "保险金额",
        "premiumCap": None,
        "merchantTypes": { "wholesale": "批发商/制造商", "retail": "零售商" },
        "coverageCategories": {
            "goodsStore": { "label": "货品损失-经营场所的货品及货品销售所得现金", "baseRates": { "wholesale": 0.00197, "retail": 0.00197 } },
            "goodsOffsite": { "label": "货品损失-场所外存储货品", "baseRates": { "wholesale": 0.00161, "retail": 0.00159 } },
            "goodsCustody": { "label": "货品损失-代保管货品", "baseRates": { "wholesale": 0.00151, "retail": 0.00151 } },
            "businessProperty": { "label": "营业财产损失保险-营业财产", "baseRates": { "wholesale": 0.00083, "retail": 0.00083 } }
        },
        "formulaText": "货品损失年保险费＝货品损失保险金额×货品损失基准费率×各项调整系数乘积",
        "formulaNote": "营业财产损失年保险费＝营业财产保险金额×营业财产基准费率×各项调整系数乘积；短期承保保险费＝年保险费×保险期间天数÷365",
        "versions": {
            "original": {
                "label": "珠宝商综合保险费率",
                "baseRates": {},
                "coefficients": [
                    { "id": "accidentRisk", "name": "意外事故风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "中等", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "naturalDisasterJewel", "name": "自然灾害风险调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较低", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "中等", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "deductibleJewel", "name": "免赔额调整系数", "applicableTo": ["all"],
                      "note": "每次事故免赔额（元）；未列明的可按线性插值法计算", "rows": [
                        { "parameter": "0", "value": 1.2, "type": "fixed" },
                        { "parameter": "1000元", "value": 1.0, "type": "fixed" },
                        { "parameter": "≥3000元", "min": 0.7, "max": 0.9, "type": "range" }
                    ]},
                    { "id": "managementJewel", "name": "管理水平调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "安全管理制度和措施完善，无明显缺陷", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "安全管理制度和措施较完善，但存在个别缺陷", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "安全管理制度和措施不完善或存在较多缺陷", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossRatioJewel", "name": "赔付率调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "[0，20%]", "min": 0.5, "max": 0.6, "type": "range" },
                        { "parameter": "（20%，45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range" },
                        { "parameter": "（45%，70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "（70%，95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "＞95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "renewalJewel", "name": "续保调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "新保", "value": 1.0, "type": "fixed" },
                        { "parameter": "续保一年", "value": 0.9, "type": "fixed" },
                        { "parameter": "续保两年", "value": 0.85, "type": "fixed" },
                        { "parameter": "续保三年及以上", "value": 0.8, "type": "fixed" }
                    ]},
                    { "id": "securityJewel", "name": "治安状况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "较好", "min": 0.7, "max": 1.0, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较差", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" }
                    ]},
                    { "id": "lossHistoryJewel", "name": "历史事故与损失情况调整系数", "applicableTo": ["all"], "rows": [
                        { "parameter": "极少", "min": 0.5, "max": 0.7, "type": "range" },
                        { "parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range" },
                        { "parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range" },
                        { "parameter": "较多", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range" },
                        { "parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range" }
                    ]}
                ]
            }
        }
    },
    "machineryEquipment": {
            "productName": "工程机械设备综合保险",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "短期承保保险费＝年保险费×保险期间天数÷365",
            "versions": {
                "original": {
                    "label": "工程机械设备综合保险费率",
                    "baseRates": {
                        "default": 0.0026
                    },
                    "coefficients": [
                        {
                            "id": "insuredAmountME",
                            "name": "保险金额调整系数",
                            "applicableTo": ["all"],
                            "note": "平均每台设备保险金额；未列明金额可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "≤10万元",
                                    "min": 1.3,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "20万元",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "50万元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "100万元",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥200万元",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "complexityME",
                            "name": "设备结构复杂程度调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "简单",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "复杂",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "特别复杂",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossHistoryME",
                            "name": "历史事故与损失情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "极少",
                                    "min": 0.5,
                                    "max": 0.7,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较少",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较多",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "很多",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "areaRiskME",
                            "name": "作业区域风险调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较低",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较高",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "originME",
                            "name": "设备产地调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "国内",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "国外",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementME",
                            "name": "管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "安全管理制度和措施完善，无明显缺陷",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "安全管理制度和措施较完善，但存在个别缺陷",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "安全管理制度和措施不完善或存在较多缺陷",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "renewalME",
                            "name": "续保调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "新保",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保一年",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保两年",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保三年及以上",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountME",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleME",
                            "note": "未列明免赔额可按线性插值法计算；若同时约定免赔额和免赔率，以低者作为免赔调整系数",
                            "rows": [
                                {
                                    "parameter": "0元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1000元",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2500元",
                                    "value": 0.91,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥10000元",
                                    "min": 0.7,
                                    "max": 0.77,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductRateME",
                            "name": "免赔率调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleME",
                            "note": "未列明免赔率可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "0",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "40%",
                                    "value": 0.6,
                                    "type": "fixed"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "hiTechProperty2025": {
            "productName": "高新技术企业财产保险（一切险2025版）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": 0.7,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算；短期承保保险费根据条款所附短期费率表计收",
            "versions": {
                "original": {
                    "label": "高新技术企业财产保险（一切险2025版）费率",
                    "baseRates": {
                        "default": 0.0038
                    },
                    "coefficients": [
                        {
                            "id": "industryHTP",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.0,
                                    "max": 1.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "软件(0201-0203)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 0.8,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.1,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionRiskHTP",
                            "name": "区域灾害风险调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "低",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "中",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "高",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "insuredAmountHTP",
                            "name": "保险金额调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜1000万",
                                    "min": 2.1,
                                    "max": 5.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000万，1亿）",
                                    "min": 0.8,
                                    "max": 2.1,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1亿，10亿）",
                                    "min": 0.5,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥10亿",
                                    "min": 0.2,
                                    "max": 0.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountHTP",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleHTP",
                            "note": "免赔额与保险金额交叉查表；未列明可按线性插值法计算；若同时约定免赔额和免赔率，以低者为准",
                            "rows": [
                                {
                                    "parameter": "1000元（保额≤1000万）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1000元（保额约1亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额≤1000万）",
                                    "value": 0.93,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额约1亿）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额约10亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额＞10亿）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额≤1000万）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额约1亿）",
                                    "value": 0.93,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额约10亿）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额＞10亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额≤1000万）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额约1亿）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额约10亿）",
                                    "value": 0.93,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额＞10亿）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额≤1000万）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额约1亿）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额约10亿）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额＞10亿）",
                                    "value": 0.93,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额≤1000万）",
                                    "value": 0.6,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额约1亿）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额约10亿）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额＞10亿）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额≤1000万）",
                                    "value": 0.52,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额约1亿）",
                                    "value": 0.6,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额约10亿）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额＞10亿）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额≤1000万）",
                                    "value": 0.45,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额约1亿）",
                                    "value": 0.52,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额约10亿）",
                                    "value": 0.6,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额＞10亿）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额约1亿）",
                                    "value": 0.45,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额约10亿）",
                                    "value": 0.52,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额＞10亿）",
                                    "value": 0.6,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductRateHTP",
                            "name": "免赔率调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleHTP",
                            "note": "未列明免赔率可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "0",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "buildingHTP",
                            "name": "建筑物结构调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "钢筋混凝土",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "钢结构",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "其他",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "riskMgmtHTP",
                            "name": "风险管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较高",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较低",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossRatioHTP",
                            "name": "赔付率调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "[0，20%]",
                                    "min": 0.5,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（20%，45%]",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（45%，70%]",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（70%，95%]",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞95%",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "hiTechRDEquip2025": {
            "productName": "高新技术企业关键研发设备保险（2025版）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "短期承保保险费按条款所附短期费率表计收",
            "versions": {
                "original": {
                    "label": "高新技术企业关键研发设备保险（2025版）费率",
                    "baseRates": {
                        "default": 0.0041
                    },
                    "coefficients": [
                        {
                            "id": "industryHTRD",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.2,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 1.5,
                                    "max": 5.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 1.1,
                                    "max": 2.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.0,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 0.8,
                                    "max": 4.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.2,
                                    "max": 3.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 3.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 0.8,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 0.8,
                                    "max": 2.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "insuredAmountHTRD",
                            "name": "保险金额调整系数",
                            "applicableTo": ["all"],
                            "note": "保险金额单位为万元",
                            "rows": [
                                {
                                    "parameter": "＜50万",
                                    "min": 2.1,
                                    "max": 5.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[50，100）万",
                                    "min": 1.2,
                                    "max": 2.1,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[100，500）万",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[500，1000）万",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥1000万",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountHTRD",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleHTRD",
                            "note": "免赔额与保险金额交叉查表；保险金额单位为万元",
                            "rows": [
                                {
                                    "parameter": "[1000,2000]元（保额＜50万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000,2000]元（保额[50,100)万）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额＜50万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额[50,100)万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额＜50万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额[50,100)万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额[100,500)万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额＜50万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额[50,100)万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额[100,500)万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额[500,1000)万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额＜50万）",
                                    "min": 0.84,
                                    "max": 0.9,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额[50,100)万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额[100,500)万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额[500,1000)万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额≥1000万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额＜50万）",
                                    "min": 0.67,
                                    "max": 0.84,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额[50,100)万）",
                                    "min": 0.84,
                                    "max": 0.9,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额[100,500)万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额[500,1000)万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额≥1000万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额[100,500)万）",
                                    "min": 0.84,
                                    "max": 0.9,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额[500,1000)万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额≥1000万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额[100,500)万）",
                                    "min": 0.67,
                                    "max": 0.84,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额[500,1000)万）",
                                    "min": 0.84,
                                    "max": 0.9,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额≥1000万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "riskMgmtHTRD",
                            "name": "风险管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较高",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较低",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossRatioHTRD",
                            "name": "赔付率调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "[0，20%]",
                                    "min": 0.5,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（20%，45%]",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（45%，70%]",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（70%，95%]",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞95%",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "hiTechBI2025": {
            "productName": "高新技术企业营业中断保险（2025版）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "短期承保保险费按条款所附短期费率表计收",
            "versions": {
                "original": {
                    "label": "高新技术企业营业中断保险（2025版）费率",
                    "baseRates": {
                        "default": 0.003
                    },
                    "coefficients": [
                        {
                            "id": "industryHTBI",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.0,
                                    "max": 1.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "软件(0201-0203)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 0.8,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.1,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionRiskHTBI",
                            "name": "区域灾害风险调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较低",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较高",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "rdTermHTBI",
                            "name": "研发期限调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜0.5年",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[0.5，1）年",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[1，2）年",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥2年",
                                    "min": 1.5,
                                    "max": 1.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountHTBI",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleHTBI",
                            "note": "免赔额与保险金额交叉查表；若同时约定免赔额和免赔率，以低者为准",
                            "rows": [
                                {
                                    "parameter": "[1000,2000]元（保额＜1000万）",
                                    "min": 0.93,
                                    "max": 0.96,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000,2000]元（保额[1000万,1亿)）",
                                    "min": 0.96,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额＜1000万）",
                                    "min": 0.87,
                                    "max": 0.93,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额[1000万,1亿)）",
                                    "min": 0.93,
                                    "max": 0.96,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额[1亿,10亿)）",
                                    "min": 0.96,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(2000,5000]元（保额≥10亿）",
                                    "min": 0.98,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额＜1000万）",
                                    "min": 0.8,
                                    "max": 0.87,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额[1000万,1亿)）",
                                    "min": 0.87,
                                    "max": 0.93,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额[1亿,10亿)）",
                                    "min": 0.93,
                                    "max": 0.96,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5000,1万]元（保额≥10亿）",
                                    "min": 0.96,
                                    "max": 0.98,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额＜1000万）",
                                    "min": 0.66,
                                    "max": 0.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额[1000万,1亿)）",
                                    "min": 0.8,
                                    "max": 0.87,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额[1亿,10亿)）",
                                    "min": 0.87,
                                    "max": 0.93,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(1万,3万]元（保额≥10亿）",
                                    "min": 0.93,
                                    "max": 0.96,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额＜1000万）",
                                    "min": 0.6,
                                    "max": 0.66,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额[1000万,1亿)）",
                                    "min": 0.66,
                                    "max": 0.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额[1亿,10亿)）",
                                    "min": 0.8,
                                    "max": 0.87,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(3万,5万]元（保额≥10亿）",
                                    "min": 0.87,
                                    "max": 0.93,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额＜1000万）",
                                    "min": 0.52,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额[1000万,1亿)）",
                                    "min": 0.6,
                                    "max": 0.66,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额[1亿,10亿)）",
                                    "min": 0.66,
                                    "max": 0.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(5万,10万]元（保额≥10亿）",
                                    "min": 0.8,
                                    "max": 0.87,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额＜1000万）",
                                    "min": 0.45,
                                    "max": 0.52,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额[1000万,1亿)）",
                                    "min": 0.52,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额[1亿,10亿)）",
                                    "min": 0.6,
                                    "max": 0.66,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(10万,20万]元（保额≥10亿）",
                                    "min": 0.66,
                                    "max": 0.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额[1000万,1亿)）",
                                    "min": 0.45,
                                    "max": 0.52,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额[1亿,10亿)）",
                                    "min": 0.52,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "(20万,30万]元（保额≥10亿）",
                                    "min": 0.6,
                                    "max": 0.66,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "riskMgmtHTBI",
                            "name": "风险管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较高",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较低",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossRatioHTBI",
                            "name": "赔付率调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "[0，20%]",
                                    "min": 0.5,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（20%，45%]",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（45%，70%]",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（70%，95%]",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞95%",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "techRDProtection2025": {
            "productName": "科技企业研发保障保险（2025版）",
            "productType": "composite",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "subRisks": ["rdExpense", "livingAllowance"],
            "subRiskLabels": {
                "rdExpense": "研发费用损失责任",
                "livingAllowance": "生活补贴责任"
            },
            "formulaText": "保险费＝∑（分项责任保险金额×分项责任基准费率×各项适用的费率调整系数的乘积）",
            "formulaNote": "",
            "versions": {
                "original": {
                    "label": "科技企业研发保障保险（2025版）费率",
                    "baseRates": {
                        "rdExpense": 0.0306,
                        "livingAllowance": 0.0332
                    },
                    "coefficients": [
                        {
                            "id": "projectRiskTRP",
                            "name": "项目风险调整系数",
                            "applicableTo": ["rdExpense", "livingAllowance"],
                            "rows": [
                                {
                                    "parameter": "较低",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较高",
                                    "min": 1.5,
                                    "max": 2.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "projectTermTRP",
                            "name": "项目期限调整系数",
                            "applicableTo": ["rdExpense", "livingAllowance"],
                            "rows": [
                                {
                                    "parameter": "≤1年",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（1，2]年",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（2，3]年",
                                    "min": 1.2,
                                    "max": 1.4,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞3年",
                                    "min": 1.4,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "rdExperienceTRP",
                            "name": "研发经验调整系数",
                            "applicableTo": ["rdExpense", "livingAllowance"],
                            "rows": [
                                {
                                    "parameter": "有丰富的项目研发经验",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "有一定的项目研发经验",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "无项目研发经验",
                                    "min": 1.3,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementTRP",
                            "name": "管理水平调整系数",
                            "applicableTo": ["rdExpense", "livingAllowance"],
                            "rows": [
                                {
                                    "parameter": "管理人员的管理水平较高",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "管理人员的管理水平一般",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "管理人员的管理水平较低",
                                    "min": 1.3,
                                    "max": 1.8,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductRateTRP",
                            "name": "免赔率调整系数",
                            "applicableTo": ["rdExpense", "livingAllowance"],
                            "linkedGroup": "deductibleTRP",
                            "note": "若同时约定免赔率与免赔额，以低者为准",
                            "rows": [
                                {
                                    "parameter": "0",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountRdTRP",
                            "name": "免赔额调整系数（研发费用损失）",
                            "applicableTo": ["rdExpense"],
                            "linkedGroup": "deductibleTRP",
                            "note": "未列明免赔额可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "0万元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元",
                                    "value": 0.99,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元",
                                    "value": 0.97,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥10万元",
                                    "min": 0.9,
                                    "max": 0.93,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountLivTRP",
                            "name": "免赔额调整系数（生活补贴）",
                            "applicableTo": ["livingAllowance"],
                            "linkedGroup": "deductibleTRP",
                            "note": "未列明免赔额可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "0万元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥2万元",
                                    "min": 0.85,
                                    "max": 0.9,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "energyStorage2025": {
            "productName": "储能设施综合运营保险（2025版）",
            "productType": "composite",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "subRisks": [
                "storageAllRisk",
                "equipDamage",
                "biAllRisk",
                "biEquip"
            ],
            "subRiskLabels": {
                "storageAllRisk": "运营期一切险",
                "equipDamage": "设备损失保险",
                "biAllRisk": "营业中断(一切险)",
                "biEquip": "营业中断(设备损失)"
            },
            "formulaText": "总保险费＝各项保险的年保险费之和",
            "formulaNote": "各项保险的年保险费＝各项保险的保险金额×对应的基准费率×各项适用的费率调整系数的乘积",
            "versions": {
                "original": {
                    "label": "储能设施综合运营保险（2025版）费率",
                    "baseRates": {
                        "storageAllRisk": 0.0021,
                        "equipDamage": 0.0023,
                        "biAllRisk": 0.0032,
                        "biEquip": 0.0058
                    },
                    "coefficients": [
                        {
                            "id": "insuredAmountES",
                            "name": "保险金额调整系数",
                            "applicableTo": [
                                "storageAllRisk",
                                "equipDamage",
                                "biAllRisk",
                                "biEquip"
                            ],
                            "rows": [
                                {
                                    "parameter": "＜1000万",
                                    "min": 1.1,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000万，5000万）",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000万，1亿）",
                                    "min": 0.95,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1亿，2亿）",
                                    "min": 0.9,
                                    "max": 0.95,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2亿，5亿）",
                                    "min": 0.85,
                                    "max": 0.9,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥5亿",
                                    "min": 0.8,
                                    "max": 0.85,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossRatioES",
                            "name": "赔付率调整系数",
                            "applicableTo": [
                                "storageAllRisk",
                                "equipDamage",
                                "biAllRisk",
                                "biEquip"
                            ],
                            "rows": [
                                {
                                    "parameter": "[0，20%]",
                                    "min": 0.5,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（20%，45%]",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（45%，70%]",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（70%，95%]",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞95%",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementES",
                            "name": "管理水平调整系数",
                            "applicableTo": [
                                "storageAllRisk",
                                "equipDamage",
                                "biAllRisk",
                                "biEquip"
                            ],
                            "rows": [
                                {
                                    "parameter": "管理制度和事故防范措施完善",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "管理制度和事故防范措施较为完善",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "管理制度和事故防范措施缺陷较多",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossHistoryES",
                            "name": "历史事故与损失情况调整系数",
                            "applicableTo": [
                                "storageAllRisk",
                                "equipDamage",
                                "biAllRisk",
                                "biEquip"
                            ],
                            "rows": [
                                {
                                    "parameter": "极少",
                                    "min": 0.5,
                                    "max": 0.7,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较少",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较多",
                                    "min": 1.3,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "很多",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "naturalDisasterES",
                            "name": "自然灾害风险调整系数",
                            "applicableTo": ["storageAllRisk", "biAllRisk"],
                            "rows": [
                                {
                                    "parameter": "低",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "中",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "高",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireExplosionES",
                            "name": "火灾爆炸风险调整系数",
                            "applicableTo": ["storageAllRisk", "biAllRisk"],
                            "rows": [
                                {
                                    "parameter": "较低",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较高",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireCapabilityES",
                            "name": "消防能力调整系数",
                            "applicableTo": ["storageAllRisk", "biAllRisk"],
                            "rows": [
                                {
                                    "parameter": "消防设施设备及人员完备，能快速处置",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "消防设施设备及人员较完备，能较快处置",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "消防设施设备及人员不足或无法及时处置",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountARES",
                            "name": "一切险免赔额调整系数",
                            "applicableTo": ["storageAllRisk"],
                            "linkedGroup": "deductibleAR_ES",
                            "note": "若同时约定免赔额和免赔率，以低者为准",
                            "rows": [
                                {
                                    "parameter": "0万元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "0.5万元",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1.0万元",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1.5万元",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥2.0万元",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductRateARES",
                            "name": "一切险免赔率调整系数",
                            "applicableTo": ["storageAllRisk"],
                            "linkedGroup": "deductibleAR_ES",
                            "rows": [
                                {
                                    "parameter": "0%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "equipAgeES",
                            "name": "主要设备已使用年数调整系数",
                            "applicableTo": ["equipDamage", "biEquip"],
                            "rows": [
                                {
                                    "parameter": "＜5年",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5，10）年",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10，15）年",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥15年",
                                    "min": 2.0,
                                    "max": 3.0,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "equipOriginES",
                            "name": "主要设备产地调整系数",
                            "applicableTo": ["equipDamage", "biEquip"],
                            "rows": [
                                {
                                    "parameter": "国产",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "进口",
                                    "min": 1.2,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountEDES",
                            "name": "设备损失免赔额调整系数",
                            "applicableTo": ["equipDamage"],
                            "linkedGroup": "deductibleED_ES",
                            "note": "若同时约定免赔额和免赔率，以低者为准",
                            "rows": [
                                {
                                    "parameter": "0万元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "0.5万元",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1.0万元",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1.5万元",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥2.0万元",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductRateEDES",
                            "name": "设备损失免赔率调整系数",
                            "applicableTo": ["equipDamage"],
                            "linkedGroup": "deductibleED_ES",
                            "rows": [
                                {
                                    "parameter": "0%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "compensationTermES",
                            "name": "赔偿期限调整系数",
                            "applicableTo": ["biAllRisk", "biEquip"],
                            "note": "未列明赔偿期限可按线性插值法计算",
                            "rows": [
                                {
                                    "parameter": "6个月",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "12个月",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "18个月",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "24个月",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountBIES",
                            "name": "营业中断免赔额调整系数",
                            "applicableTo": ["biAllRisk", "biEquip"],
                            "linkedGroup": "deductibleBI_ES",
                            "note": "免赔额÷营业中断保险金额；若同时约定免赔额、免赔率、免赔天数，以低者为准",
                            "rows": [
                                {
                                    "parameter": "0",
                                    "value": 1.3,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1%",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2%",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "6%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥9%",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductRateBIES",
                            "name": "营业中断免赔率调整系数",
                            "applicableTo": ["biAllRisk", "biEquip"],
                            "linkedGroup": "deductibleBI_ES",
                            "rows": [
                                {
                                    "parameter": "0%",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductDaysBIES",
                            "name": "营业中断免赔天数调整系数",
                            "applicableTo": ["biAllRisk", "biEquip"],
                            "linkedGroup": "deductibleBI_ES",
                            "rows": [
                                {
                                    "parameter": "0天",
                                    "value": 1.3,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5天",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10天",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "15天",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20天",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "25天",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥30天",
                                    "min": 0.7,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "techCorpProperty": {
            "productName": "科技型企业财产保险（一切险）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": 0.7,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算",
            "versions": {
                "original": {
                    "label": "科技型企业财产保险（一切险）费率",
                    "baseRates": {
                        "default": 0.0018
                    },
                    "coefficients": [
                        {
                            "id": "industryTCP",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.0,
                                    "max": 1.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "软件(0201-0203)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 0.8,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.1,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionStormTCP",
                            "name": "暴风雨区域调整系数(占比22%)",
                            "applicableTo": ["all"],
                            "note": "区域调整系数＝暴风雨22%×本值+台风7%×台风值+洪水4%×洪水值+其他67%×其他值",
                            "rows": [
                                {
                                    "parameter": "一类(京、陕、青、宁、津)",
                                    "min": 0.12,
                                    "max": 0.54,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(晋、蒙、辽、连、吉、黑、沪、浙、甬、闽、厦、鲁、青岛、豫、鄂、粤、深、琼、渝、黔、藏、甘、新)",
                                    "min": 0.52,
                                    "max": 1.34,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(冀、苏、皖、赣、湘、桂、滇、川)",
                                    "min": 1.27,
                                    "max": 1.74,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionTyphoonTCP",
                            "name": "台风区域调整系数(占比7%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、冀、晋、蒙、辽、连、吉、黑、皖、豫、鄂、川、渝、黔、滇、藏、陕、甘、青、宁、新)",
                                    "min": 0.01,
                                    "max": 0.05,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(鲁、青岛、赣、湘)",
                                    "min": 0.1,
                                    "max": 0.48,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(沪、苏、桂、琼)",
                                    "min": 0.6,
                                    "max": 2.44,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(闽、厦、浙、甬、粤、深)",
                                    "min": 2.61,
                                    "max": 5.13,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionFloodTCP",
                            "name": "洪水区域调整系数(占比4%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、沪、藏)",
                                    "min": 0.01,
                                    "max": 0.08,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(冀、晋、蒙、辽、连、吉、黑、苏、鲁、青岛、豫、粤、深、琼、陕、甘、青、宁、新、厦)",
                                    "min": 0.11,
                                    "max": 0.88,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(浙、甬、闽、川、渝、黔、滇)",
                                    "min": 1.04,
                                    "max": 1.92,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(皖、赣、湘、桂、鄂)",
                                    "min": 1.74,
                                    "max": 4.45,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionOtherTCP",
                            "name": "其他灾因区域调整系数(占比67%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、沪、陕、青、宁、吉、渝、川、闽、粤、琼、厦、深)",
                                    "min": 0.85,
                                    "max": 0.91,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(甘、新、晋、辽、黑、鄂、苏、皖、浙、黔、连、甬、藏)",
                                    "min": 0.95,
                                    "max": 1.05,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(赣、湘、桂、冀、鲁、豫、青岛)",
                                    "min": 1.08,
                                    "max": 1.14,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(蒙、滇)",
                                    "min": 1.17,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "insuredAmountTCP",
                            "name": "保险金额调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜1000万",
                                    "min": 2.1,
                                    "max": 5.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000万，1亿）",
                                    "min": 0.8,
                                    "max": 2.1,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1亿，10亿）",
                                    "min": 0.5,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥10亿",
                                    "min": 0.2,
                                    "max": 0.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountTCP",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleTCP",
                            "note": "免赔额与保险金额交叉查表；若同时约定免赔额和免赔率，以低者为准",
                            "rows": [
                                {
                                    "parameter": "1000元（保额≤1000万）",
                                    "value": 0.97,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1000元（保额约1亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额≤1000万）",
                                    "value": 0.93,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额约1亿）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额约10亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2000元（保额＞10亿）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额≤1000万）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额约1亿）",
                                    "value": 0.92,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额约10亿）",
                                    "value": 0.96,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5000元（保额＞10亿）",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额≤1000万）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额约1亿）",
                                    "value": 0.86,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额约10亿）",
                                    "value": 0.94,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1万元（保额＞10亿）",
                                    "value": 0.97,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额≤1000万）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额约1亿）",
                                    "value": 0.73,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额约10亿）",
                                    "value": 0.86,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3万元（保额＞10亿）",
                                    "value": 0.94,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额≤1000万）",
                                    "value": 0.6,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额约1亿）",
                                    "value": 0.66,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额约10亿）",
                                    "value": 0.82,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5万元（保额＞10亿）",
                                    "value": 0.91,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额≤1000万）",
                                    "value": 0.52,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额约1亿）",
                                    "value": 0.57,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额约10亿）",
                                    "value": 0.73,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10万元（保额＞10亿）",
                                    "value": 0.87,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额≤1000万）",
                                    "value": 0.45,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额约1亿）",
                                    "value": 0.49,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额约10亿）",
                                    "value": 0.64,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20万元（保额＞10亿）",
                                    "value": 0.81,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额约1亿）",
                                    "value": 0.45,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额约10亿）",
                                    "value": 0.59,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30万元（保额＞10亿）",
                                    "value": 0.77,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductRateTCP",
                            "name": "免赔率调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleTCP",
                            "rows": [
                                {
                                    "parameter": "0",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "20%",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30%",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "geoLocationTCP",
                            "name": "地理位置调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "位于地势低洼处、山坡/山脚下、或临近江河海湖水库",
                                    "min": 1.1,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "无上述情况",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "surroundingTCP",
                            "name": "周边环境调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "周围存在火灾爆炸隐患或毗邻高风险建筑",
                                    "min": 1.1,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "无上述情况",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "specialIndustryTCP",
                            "name": "特定行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "生产工艺对环境要求极严格（如无尘、无菌）",
                                    "min": 2.0,
                                    "max": 4.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "其他行业",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "buildingTCP",
                            "name": "建筑物结构调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "钢筋混凝土结构",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "钢结构",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "砖混结构",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "砖砌/石头",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "木质结构",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "其他",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "venueRiskTCP",
                            "name": "场所占用性质调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "风险密集型公共营业场所（小商品市场、家具灯具城等）",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "其他",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "lightningTCP",
                            "name": "防雷避雷设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "缺少必要的防雷避雷设施",
                                    "min": 1.05,
                                    "max": 1.3,
                                    "type": "range"
                                },
                                {
                                    "parameter": "无上述情况",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "fireFacilityTCP",
                            "name": "消防设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "消防设施齐备、有效",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "消防设施不太齐备或部分失效",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "消防设施不齐备或无效",
                                    "min": 1.3,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "publicFireTCP",
                            "name": "公共消防队调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "距离较近、级别较高、反应较快",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "有一定距离、级别中等、反应一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "距离较远、级别较低、反应较慢",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "floodFacilityTCP",
                            "name": "防洪设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "防洪设施齐备、有效",
                                    "min": 0.95,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "防洪设施不太齐备或部分失效",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "防洪设施不齐备或均无效",
                                    "min": 1.3,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementTCP",
                            "name": "管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较高",
                                    "min": 0.5,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较低",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossHistoryTCP",
                            "name": "历史事故与损失情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "极少",
                                    "min": 0.5,
                                    "max": 0.7,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较少",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较多",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "很多",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "renewalTCP",
                            "name": "续保调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "新保",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保一年",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保两年",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保三年及以上",
                                    "value": 0.92,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "riskDispersionTCP",
                            "name": "保险标的风险分散程度调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "分散程度较高",
                                    "min": 0.85,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "分散程度较低",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "waterResistTCP",
                            "name": "保险标的耐水性调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "较高",
                                    "min": 0.95,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较低",
                                    "min": 1.3,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "techCorpRDEquip": {
            "productName": "科技型企业关键研发设备保险（2025版）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": 0.7,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算",
            "versions": {
                "original": {
                    "label": "科技型企业关键研发设备保险（2025版）费率",
                    "baseRates": {
                        "default": 0.0023
                    },
                    "coefficients": [
                        {
                            "id": "industryTCRE",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.2,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 2.0,
                                    "max": 5.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 1.1,
                                    "max": 2.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.0,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 0.8,
                                    "max": 4.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.2,
                                    "max": 3.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 3.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 0.8,
                                    "max": 2.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 0.8,
                                    "max": 2.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "insuredAmountTCRE",
                            "name": "保险金额调整系数",
                            "applicableTo": ["all"],
                            "note": "保险金额单位为万元",
                            "rows": [
                                {
                                    "parameter": "≤50万",
                                    "min": 3.0,
                                    "max": 5.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（50，100]万",
                                    "min": 2.0,
                                    "max": 3.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（100，500]万",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（500，1000]万",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞1000万",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountTCRE",
                            "name": "免赔额调整系数",
                            "applicableTo": ["all"],
                            "linkedGroup": "deductibleTCRE",
                            "note": "免赔额与保险金额交叉查表；保险金额单位为万元",
                            "rows": [
                                {
                                    "parameter": "[1000,2000)元（保额≤50万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000,2000)元（保额(50,100]万）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额≤50万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额(50,100]万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额≤50万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额(50,100]万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额(100,500]万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额≤50万）",
                                    "min": 0.9,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额(50,100]万）",
                                    "min": 0.93,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额(100,500]万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额(500,1000]万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额≤50万）",
                                    "min": 0.84,
                                    "max": 0.9,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额(50,100]万）",
                                    "min": 0.89,
                                    "max": 0.93,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额(100,500]万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额(500,1000]万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额＞1000万）",
                                    "min": 0.99,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额≤50万）",
                                    "min": 0.67,
                                    "max": 0.84,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额(50,100]万）",
                                    "min": 0.78,
                                    "max": 0.89,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额(100,500]万）",
                                    "min": 0.95,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额(500,1000]万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额＞1000万）",
                                    "min": 0.98,
                                    "max": 0.99,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额(100,500]万）",
                                    "min": 0.89,
                                    "max": 0.95,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额(500,1000]万）",
                                    "min": 0.96,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额＞1000万）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额(100,500]万）",
                                    "min": 0.84,
                                    "max": 0.89,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额(500,1000]万）",
                                    "min": 0.93,
                                    "max": 0.96,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额＞1000万）",
                                    "min": 0.95,
                                    "max": 0.97,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "operatorTCRE",
                            "name": "操作人员培训情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "充分培训并持证上岗",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "上岗前必要培训但无持证要求",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "缺少必要培训",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "maintenanceTCRE",
                            "name": "维修保养情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "定期进行维修保养",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "不定期进行维修保养",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "缺少维修保养条件",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lightningTCRE",
                            "name": "防雷避雷设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "布设了必要的防雷避雷设施",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "布设了部分防雷避雷设施",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "未布设防雷避雷设施",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementTCRE",
                            "name": "管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "制定有相关管理规范并严格执行",
                                    "min": 0.5,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "制定有相关管理规范并基本得到执行",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "缺少相关管理规范",
                                    "min": 1.15,
                                    "max": 2.0,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossHistoryTCRE",
                            "name": "历史事故与损失情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "极少",
                                    "min": 0.5,
                                    "max": 0.7,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较少",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较多",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "很多",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "renewalTCRE",
                            "name": "续保调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "新保",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保1年",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保2年",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保3年以上",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "riskDispersionTCRE",
                            "name": "研发项目风险分散程度调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "高",
                                    "min": 0.8,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "中",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "低",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lossRatioTCRE",
                            "name": "赔付率调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "[0，20%]",
                                    "min": 0.5,
                                    "max": 0.6,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（20%，45%]",
                                    "min": 0.6,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（45%，70%]",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "（70%，95%]",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "＞95%",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "techCorpBI": {
            "productName": "科技型企业营业中断保险（2025版）",
            "productType": "property",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": 0.7,
            "formulaText": "年保险费＝保险金额×基准费率×各项费率调整系数的乘积",
            "formulaNote": "若基准费率与各项费率调整系数的乘积大于70%，则按70%参与保险费的计算",
            "versions": {
                "original": {
                    "label": "科技型企业营业中断保险（2025版）费率",
                    "baseRates": {
                        "default": 0.0029
                    },
                    "coefficients": [
                        {
                            "id": "industryTCBI",
                            "name": "行业调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "电子信息(0101-0106)",
                                    "min": 1.0,
                                    "max": 1.8,
                                    "type": "range"
                                },
                                {
                                    "parameter": "软件(0201-0203)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "航空航天(0301-0302)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "光机电一体化(0401-0402)",
                                    "min": 0.8,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "生物、医药和医疗器械(0501-0505)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新材料(0601-0604)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "新能源与高效节能(0701-0702)",
                                    "min": 1.2,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "环境保护(0801-0805)",
                                    "min": 1.0,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地球、空间与海洋(0901-0905)",
                                    "min": 1.1,
                                    "max": 2.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-优良动植物新品种(1101)",
                                    "min": 0.7,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-家畜良种胚胎(1102)",
                                    "min": 0.9,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-生物农药(1103)",
                                    "min": 0.9,
                                    "max": 1.2,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-诊断试剂与疫苗(1104)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-饲料及添加剂(1105)",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-新型肥料(1106)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农业工程设施(1107)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "农业-农副产品贮藏加工(1108)",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionStormTCBI",
                            "name": "暴风雨区域调整系数(占比31%)",
                            "applicableTo": ["all"],
                            "note": "区域调整系数＝暴风雨31%×本值+台风9%×台风值+洪水5%×洪水值+其他55%×其他值",
                            "rows": [
                                {
                                    "parameter": "一类(京、陕、青、宁、津)",
                                    "min": 0.12,
                                    "max": 0.54,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(晋、蒙、辽、连、吉、黑、沪、浙、甬、闽、厦、鲁、青岛、豫、鄂、粤、深、琼、渝、黔、藏、甘、新)",
                                    "min": 0.52,
                                    "max": 1.34,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(冀、苏、皖、赣、湘、桂、滇、川)",
                                    "min": 1.27,
                                    "max": 1.74,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionTyphoonTCBI",
                            "name": "台风区域调整系数(占比9%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、冀、晋、蒙、辽、连、吉、黑、皖、豫、鄂、川、渝、黔、滇、藏、陕、甘、青、宁、新)",
                                    "min": 0.01,
                                    "max": 0.05,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(鲁、青岛、赣、湘)",
                                    "min": 0.1,
                                    "max": 0.48,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(沪、苏、桂、琼)",
                                    "min": 0.6,
                                    "max": 2.44,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(闽、厦、浙、甬、粤、深)",
                                    "min": 2.61,
                                    "max": 5.13,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionFloodTCBI",
                            "name": "洪水区域调整系数(占比5%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、沪、藏)",
                                    "min": 0.01,
                                    "max": 0.08,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(冀、晋、蒙、辽、连、吉、黑、苏、鲁、青岛、豫、粤、深、琼、陕、甘、青、宁、新、厦)",
                                    "min": 0.11,
                                    "max": 0.88,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(浙、甬、闽、川、渝、黔、滇)",
                                    "min": 1.04,
                                    "max": 1.92,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(皖、赣、湘、桂、鄂)",
                                    "min": 1.74,
                                    "max": 4.45,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "regionOtherTCBI",
                            "name": "其他灾因区域调整系数(占比55%)",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "一类(京、津、沪、陕、青、宁、吉、渝、川、闽、粤、琼、厦、深)",
                                    "min": 0.85,
                                    "max": 0.91,
                                    "type": "range"
                                },
                                {
                                    "parameter": "二类(甘、新、晋、辽、黑、鄂、苏、皖、浙、黔、连、甬、藏)",
                                    "min": 0.95,
                                    "max": 1.05,
                                    "type": "range"
                                },
                                {
                                    "parameter": "三类(赣、湘、桂、冀、鲁、豫、青岛)",
                                    "min": 1.08,
                                    "max": 1.14,
                                    "type": "range"
                                },
                                {
                                    "parameter": "四类(蒙、滇)",
                                    "min": 1.17,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "rdTermTCBI",
                            "name": "研发期限调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "半年以下（含）",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "半年至1年（含）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "1-2年（含）",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "2-3年（含）",
                                    "value": 1.5,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "3年以上",
                                    "min": 1.6,
                                    "max": 5.0,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "deductAmountTCBI",
                            "name": "绝对免赔额调整系数",
                            "applicableTo": ["all"],
                            "note": "免赔额与保险金额交叉查表",
                            "rows": [
                                {
                                    "parameter": "[1000,2000)元（保额≤1000万）",
                                    "min": 0.93,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1000,2000)元（保额约1亿）",
                                    "min": 0.96,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额≤1000万）",
                                    "min": 0.87,
                                    "max": 0.93,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额约1亿）",
                                    "min": 0.92,
                                    "max": 0.96,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额约10亿）",
                                    "min": 0.96,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[2000,5000)元（保额＞10亿）",
                                    "min": 0.98,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额≤1000万）",
                                    "min": 0.8,
                                    "max": 0.87,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额约1亿）",
                                    "min": 0.86,
                                    "max": 0.92,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额约10亿）",
                                    "min": 0.94,
                                    "max": 0.96,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5000,1万)元（保额＞10亿）",
                                    "min": 0.97,
                                    "max": 0.98,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额≤1000万）",
                                    "min": 0.66,
                                    "max": 0.8,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额约1亿）",
                                    "min": 0.73,
                                    "max": 0.86,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额约10亿）",
                                    "min": 0.86,
                                    "max": 0.94,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[1万,3万)元（保额＞10亿）",
                                    "min": 0.94,
                                    "max": 0.97,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额≤1000万）",
                                    "min": 0.6,
                                    "max": 0.66,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额约1亿）",
                                    "min": 0.66,
                                    "max": 0.73,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额约10亿）",
                                    "min": 0.81,
                                    "max": 0.86,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[3万,5万)元（保额＞10亿）",
                                    "min": 0.91,
                                    "max": 0.94,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额≤1000万）",
                                    "min": 0.52,
                                    "max": 0.6,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额约1亿）",
                                    "min": 0.57,
                                    "max": 0.66,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额约10亿）",
                                    "min": 0.73,
                                    "max": 0.81,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[5万,10万)元（保额＞10亿）",
                                    "min": 0.87,
                                    "max": 0.91,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额≤1000万）",
                                    "min": 0.45,
                                    "max": 0.52,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额约1亿）",
                                    "min": 0.49,
                                    "max": 0.57,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额约10亿）",
                                    "min": 0.64,
                                    "max": 0.73,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[10万,20万)元（保额＞10亿）",
                                    "min": 0.81,
                                    "max": 0.87,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额约1亿）",
                                    "min": 0.45,
                                    "max": 0.49,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额约10亿）",
                                    "min": 0.59,
                                    "max": 0.64,
                                    "type": "range"
                                },
                                {
                                    "parameter": "[20万,30万]元（保额＞10亿）",
                                    "min": 0.77,
                                    "max": 0.81,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "geoLocationTCBI",
                            "name": "地理位置风险状况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "地势较高或远离江河海湖水库",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "地势低洼或临近江河海湖水库",
                                    "min": 1.0,
                                    "max": 1.3,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireRiskTCBI",
                            "name": "火灾爆炸风险调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "无隐患，防火灭火装置齐备有效",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "存在一定隐患，装置不太齐备或部分有效",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "隐患较多，装置不齐备且无效",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "buildingTCBI",
                            "name": "建筑物结构调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "钢筋混凝土结构",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "钢结构",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "砖混结构",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "砖砌/石头",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "木材",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "其他",
                                    "value": 1.2,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "venueRiskTCBI",
                            "name": "营业场所风险密集程度",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "低",
                                    "min": 0.8,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "中",
                                    "min": 1.0,
                                    "max": 1.5,
                                    "type": "range"
                                },
                                {
                                    "parameter": "高",
                                    "min": 1.5,
                                    "max": 3.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "lightningTCBI",
                            "name": "防雷避雷设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "布设了必要的防雷避雷设施",
                                    "min": 0.8,
                                    "max": 1.0,
                                    "type": "range"
                                },
                                {
                                    "parameter": "布设了部分防雷避雷设施",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "未布设防雷避雷设施",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "publicFireTCBI",
                            "name": "公共消防队调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "距离较近、级别较高、反应较快",
                                    "min": 0.8,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "有一定距离、级别中等、反应一般",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "距离较远、级别较低、反应较慢",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "floodFacilityTCBI",
                            "name": "防洪设施调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "防洪设施齐备、有效",
                                    "min": 0.8,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "防洪设施不太齐备，部分有效",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "防洪设施不齐备或均无效",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "managementTCBI",
                            "name": "企业风险管理水平调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "制定有相关管理规范并严格执行",
                                    "min": 0.5,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "制定有相关管理规范并基本得到执行",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "缺少相关管理规范",
                                    "min": 1.15,
                                    "max": 2.0,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "riskDispersionTCBI",
                            "name": "研发项目风险分散程度调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "高",
                                    "min": 0.8,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "中",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "低",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "waterRiskTCBI",
                            "name": "水损风险调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "不易受损",
                                    "min": 0.5,
                                    "max": 0.95,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.1,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较易受损",
                                    "min": 1.15,
                                    "max": 2.0,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "renewalTCBI",
                            "name": "续保调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "新保",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保1年",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保2年",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "续保3年以上",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "lossHistoryTCBI",
                            "name": "历史事故与损失情况调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "极少",
                                    "min": 0.5,
                                    "max": 0.7,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较少",
                                    "min": 0.7,
                                    "max": 1.0,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "一般",
                                    "min": 1.0,
                                    "max": 1.2,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "较多",
                                    "min": 1.2,
                                    "max": 1.5,
                                    "minExclusive": True,
                                    "type": "range"
                                },
                                {
                                    "parameter": "很多",
                                    "min": 1.5,
                                    "max": 2.0,
                                    "minExclusive": True,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "commercialBuilding": {
            "productName": "商业楼宇财产一切险",
            "productType": "multiRiskSum",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": 0.5,
            "formulaText": "年保费 = 保险金额 × Σ(风险基准费率 × 调整系数乘积)",
            "formulaNote": "各费率调整系数乘积不应低于0.50",
            "versions": {
                "original": {
                    "label": "商业楼宇财产一切险费率",
                    "risks": [
                        {
                            "id": "fireExplosion",
                            "name": "火灾、爆炸",
                            "baseRateTable": {
                                "type": "industryBuilding",
                                "buildingTypes": [
                                    {
                                        "id": "fireproof",
                                        "name": "防火建筑"
                                    },
                                    {
                                        "id": "nonflammable",
                                        "name": "不可燃建筑"
                                    },
                                    {
                                        "id": "combustible",
                                        "name": "可燃建筑"
                                    },
                                    {
                                        "id": "flammable",
                                        "name": "易燃建筑"
                                    }
                                ],
                                "labels": {
                                    "J6800": "J6800 银行业",
                                    "J700": "J700 保险业",
                                    "J6900": "J6900 证券业",
                                    "J7100": "J7100 其他金融活动",
                                    "L7300": "L7300 租赁业",
                                    "L7400": "L7400 商务服务",
                                    "L7410": "L7410 企业管理服务",
                                    "L7420": "L7420 法律服务",
                                    "L7430": "L7430 咨询与调查",
                                    "L7439": "L7439 其他专业咨询",
                                    "L7440": "L7440 广告业",
                                    "L7450": "L7450 知识产权服务",
                                    "L7460": "L7460 职业中介服务",
                                    "L7480": "L7480 旅行社",
                                    "L7491": "L7491 保安服务",
                                    "L7490": "L7490 其他服务",
                                    "I6600": "I6600 住宿业",
                                    "I6610": "I6610 旅游饭店",
                                    "I6620": "I6620 一般旅馆",
                                    "I6620-1": "I6620-1 星级旅馆",
                                    "I6690": "I6690 其他住宿服务",
                                    "I6700": "I6700 餐饮业",
                                    "I6710": "I6710 正餐服务",
                                    "I6720": "I6720 快餐服务",
                                    "I6730": "I6730 饮料及冷饮服务",
                                    "M7500": "M7500 研究与试验发展",
                                    "M7600": "M7600 专业技术服务",
                                    "M7700": "M7700 科技交流和推广服务",
                                    "N7900": "N7900 水利管理业",
                                    "N8000": "N8000 环境管理业",
                                    "N8010": "N8010 自然保护",
                                    "N8021": "N8021 城市环卫管理",
                                    "N8022": "N8022 水污染治理",
                                    "N8100": "N8100 公共设施管理",
                                    "O8220": "O8220 托儿所/养老院/保健服务",
                                    "O8240": "O8240 理发美容",
                                    "O8230": "O8230 洗染服务",
                                    "O8250": "O8250 洗浴服务",
                                    "O8270": "O8270 殡葬服务",
                                    "O8280": "O8280 摄影扩印服务",
                                    "O8290": "O8290 其他居民服务",
                                    "O8311": "O8311 汽车摩托车维护保养",
                                    "O8316": "O8316 电器修理",
                                    "O8317": "O8317 钟表珠宝维修",
                                    "O8318": "O8318 装潢家具维修",
                                    "O8319": "O8319 其他日用品维修",
                                    "O8320": "O8320 保洁服务",
                                    "P8410": "P8410 学前初等中等教育",
                                    "P8440": "P8440 高等教育",
                                    "P8490": "P8490 其他教育",
                                    "P8491": "P8491 图书馆和信息中心",
                                    "Q8510-1": "Q8510-1 医院卫生院",
                                    "Q8510-2": "Q8510-2 妇幼保健活动",
                                    "Q8570": "Q8570 疾病防控活动",
                                    "Q8511": "Q8511 疗养院",
                                    "Q8590": "Q8590 其他卫生活动",
                                    "Q8600": "Q8600 社会保障业",
                                    "Q8700": "Q8700 社会福利业",
                                    "R8810": "R8810 新闻出版业",
                                    "R8900": "R8900 广播电视影像制作",
                                    "R8931": "R8931 电影制作和发行",
                                    "R8933": "R8933 电影放映",
                                    "R9010": "R9010 文艺创作与表演",
                                    "R9020": "R9020 艺术表演场馆",
                                    "R9040": "R9040 文物与文化保护",
                                    "R9050": "R9050 博物馆纪念馆",
                                    "R9110": "R9110 体育组织",
                                    "R9120": "R9120 体育场馆",
                                    "R9210": "R9210 室内娱乐活动",
                                    "R9220": "R9220 游乐园",
                                    "R9230": "R9230 休闲健身活动",
                                    "R9200-1": "R9200-1 其他休闲娱乐",
                                    "F5521": "F5521 机场",
                                    "F5700": "F5700 运输服务业",
                                    "F5900": "F5900 邮政业",
                                    "F5990": "F5990 其他寄递服务",
                                    "G6010": "G6010 电信业",
                                    "G6020": "G6020 互联网服务",
                                    "G6040": "G6040 卫星传输服务",
                                    "G6030": "G6030 广播电视传输服务",
                                    "G6100": "G6100 计算机服务",
                                    "G6200": "G6200 软件业",
                                    "S9300": "S9300 中国共产党机关",
                                    "S9400": "S9400 国家及政府机关",
                                    "S9490": "S9490 公安局消防局",
                                    "S9500": "S9500 人民政协和民主党派",
                                    "S9600-1": "S9600-1 群众团体及社会团体",
                                    "S9600-2": "S9600-2 宗教组织",
                                    "S9600-3": "S9600-3 其他团体组织",
                                    "K": "K 房地产业",
                                    "H": "H 批发和零售业",
                                    "T": "T 国际组织"
                                },
                                "data": {
                                    "J6800": {
                                        "fireproof": 0.000127,
                                        "nonflammable": 0.000179,
                                        "combustible": 0.000242,
                                        "flammable": 0.000457
                                    },
                                    "J700": {
                                        "fireproof": 0.000127,
                                        "nonflammable": 0.000179,
                                        "combustible": 0.000242,
                                        "flammable": 0.000457
                                    },
                                    "J6900": {
                                        "fireproof": 0.000127,
                                        "nonflammable": 0.000179,
                                        "combustible": 0.000242,
                                        "flammable": 0.000457
                                    },
                                    "J7100": {
                                        "fireproof": 0.000139,
                                        "nonflammable": 0.000197,
                                        "combustible": 0.000265,
                                        "flammable": 0.000503
                                    },
                                    "L7300": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000338,
                                        "flammable": 0.000641
                                    },
                                    "L7400": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000338,
                                        "flammable": 0.000641
                                    },
                                    "L7410": {
                                        "fireproof": 0.000139,
                                        "nonflammable": 0.000197,
                                        "combustible": 0.000267,
                                        "flammable": 0.000504
                                    },
                                    "L7420": {
                                        "fireproof": 0.000127,
                                        "nonflammable": 0.000179,
                                        "combustible": 0.000242,
                                        "flammable": 0.000457
                                    },
                                    "L7430": {
                                        "fireproof": 0.000139,
                                        "nonflammable": 0.000197,
                                        "combustible": 0.000267,
                                        "flammable": 0.000504
                                    },
                                    "L7439": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000338,
                                        "flammable": 0.000641
                                    },
                                    "L7440": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000338,
                                        "flammable": 0.000641
                                    },
                                    "L7450": {
                                        "fireproof": 0.000165,
                                        "nonflammable": 0.000233,
                                        "combustible": 0.000313,
                                        "flammable": 0.000595
                                    },
                                    "L7460": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000339,
                                        "flammable": 0.000642
                                    },
                                    "L7480": {
                                        "fireproof": 0.000191,
                                        "nonflammable": 0.000268,
                                        "combustible": 0.000362,
                                        "flammable": 0.000687
                                    },
                                    "L7491": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000338,
                                        "flammable": 0.000641
                                    },
                                    "L7490": {
                                        "fireproof": 0.000191,
                                        "nonflammable": 0.000268,
                                        "combustible": 0.000362,
                                        "flammable": 0.000687
                                    },
                                    "I6600": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "I6610": {
                                        "fireproof": 0.000217,
                                        "nonflammable": 0.000304,
                                        "combustible": 0.00041,
                                        "flammable": 0.000777
                                    },
                                    "I6620": {
                                        "fireproof": 0.000217,
                                        "nonflammable": 0.000304,
                                        "combustible": 0.00041,
                                        "flammable": 0.000777
                                    },
                                    "I6620-1": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000339,
                                        "flammable": 0.000642
                                    },
                                    "I6690": {
                                        "fireproof": 0.000217,
                                        "nonflammable": 0.000304,
                                        "combustible": 0.00041,
                                        "flammable": 0.000777
                                    },
                                    "I6700": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "I6710": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "I6720": {
                                        "fireproof": 0.00028,
                                        "nonflammable": 0.000394,
                                        "combustible": 0.000532,
                                        "flammable": 0.001005
                                    },
                                    "I6730": {
                                        "fireproof": 0.000242,
                                        "nonflammable": 0.000341,
                                        "combustible": 0.00046,
                                        "flammable": 0.000871
                                    },
                                    "M7500": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "M7600": {
                                        "fireproof": 0.00023,
                                        "nonflammable": 0.000323,
                                        "combustible": 0.000436,
                                        "flammable": 0.000825
                                    },
                                    "M7700": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "N7900": {
                                        "fireproof": 0.000192,
                                        "nonflammable": 0.00027,
                                        "combustible": 0.000365,
                                        "flammable": 0.000691
                                    },
                                    "N8000": {
                                        "fireproof": 0.000176,
                                        "nonflammable": 0.000247,
                                        "combustible": 0.000332,
                                        "flammable": 0.00063
                                    },
                                    "N8010": {
                                        "fireproof": 0.000176,
                                        "nonflammable": 0.000247,
                                        "combustible": 0.000332,
                                        "flammable": 0.00063
                                    },
                                    "N8021": {
                                        "fireproof": 0.000189,
                                        "nonflammable": 0.000265,
                                        "combustible": 0.000359,
                                        "flammable": 0.000678
                                    },
                                    "N8022": {
                                        "fireproof": 0.000189,
                                        "nonflammable": 0.000265,
                                        "combustible": 0.000359,
                                        "flammable": 0.000678
                                    },
                                    "N8100": {
                                        "fireproof": 0.000192,
                                        "nonflammable": 0.00027,
                                        "combustible": 0.000365,
                                        "flammable": 0.000691
                                    },
                                    "O8220": {
                                        "fireproof": 0.000186,
                                        "nonflammable": 0.000262,
                                        "combustible": 0.000354,
                                        "flammable": 0.000671
                                    },
                                    "O8240": {
                                        "fireproof": 0.000248,
                                        "nonflammable": 0.00035,
                                        "combustible": 0.000472,
                                        "flammable": 0.000895
                                    },
                                    "O8230": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "O8250": {
                                        "fireproof": 0.000289,
                                        "nonflammable": 0.000406,
                                        "combustible": 0.000548,
                                        "flammable": 0.001039
                                    },
                                    "O8270": {
                                        "fireproof": 0.000313,
                                        "nonflammable": 0.000441,
                                        "combustible": 0.000595,
                                        "flammable": 0.001128
                                    },
                                    "O8280": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "O8290": {
                                        "fireproof": 0.000248,
                                        "nonflammable": 0.00035,
                                        "combustible": 0.000472,
                                        "flammable": 0.000895
                                    },
                                    "O8311": {
                                        "fireproof": 0.000289,
                                        "nonflammable": 0.000406,
                                        "combustible": 0.000548,
                                        "flammable": 0.001039
                                    },
                                    "O8316": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "O8317": {
                                        "fireproof": 0.00023,
                                        "nonflammable": 0.000323,
                                        "combustible": 0.000436,
                                        "flammable": 0.000825
                                    },
                                    "O8318": {
                                        "fireproof": 0.00023,
                                        "nonflammable": 0.000323,
                                        "combustible": 0.000436,
                                        "flammable": 0.000825
                                    },
                                    "O8319": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "O8320": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "P8410": {
                                        "fireproof": 0.00018,
                                        "nonflammable": 0.000254,
                                        "combustible": 0.000344,
                                        "flammable": 0.00065
                                    },
                                    "P8440": {
                                        "fireproof": 0.00018,
                                        "nonflammable": 0.000254,
                                        "combustible": 0.000344,
                                        "flammable": 0.00065
                                    },
                                    "P8490": {
                                        "fireproof": 0.000209,
                                        "nonflammable": 0.000294,
                                        "combustible": 0.000397,
                                        "flammable": 0.000753
                                    },
                                    "P8491": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000339,
                                        "flammable": 0.000642
                                    },
                                    "Q8510-1": {
                                        "fireproof": 0.000248,
                                        "nonflammable": 0.00035,
                                        "combustible": 0.000472,
                                        "flammable": 0.000895
                                    },
                                    "Q8510-2": {
                                        "fireproof": 0.000244,
                                        "nonflammable": 0.000342,
                                        "combustible": 0.000462,
                                        "flammable": 0.000875
                                    },
                                    "Q8570": {
                                        "fireproof": 0.000244,
                                        "nonflammable": 0.000342,
                                        "combustible": 0.000462,
                                        "flammable": 0.000875
                                    },
                                    "Q8511": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "Q8590": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "Q8600": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "Q8700": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "R8810": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "R8900": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "R8931": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "R8933": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "R9010": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "R9020": {
                                        "fireproof": 0.000238,
                                        "nonflammable": 0.000335,
                                        "combustible": 0.000451,
                                        "flammable": 0.000854
                                    },
                                    "R9040": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "R9050": {
                                        "fireproof": 0.000179,
                                        "nonflammable": 0.000251,
                                        "combustible": 0.000339,
                                        "flammable": 0.000642
                                    },
                                    "R9110": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "R9120": {
                                        "fireproof": 0.000268,
                                        "nonflammable": 0.000377,
                                        "combustible": 0.000509,
                                        "flammable": 0.000965
                                    },
                                    "R9210": {
                                        "fireproof": 0.000294,
                                        "nonflammable": 0.000412,
                                        "combustible": 0.000557,
                                        "flammable": 0.001055
                                    },
                                    "R9220": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "R9230": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "R9200-1": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "F5521": {
                                        "fireproof": 0.000192,
                                        "nonflammable": 0.00027,
                                        "combustible": 0.000365,
                                        "flammable": 0.000691
                                    },
                                    "F5700": {
                                        "fireproof": 0.00023,
                                        "nonflammable": 0.000324,
                                        "combustible": 0.000438,
                                        "flammable": 0.00083
                                    },
                                    "F5900": {
                                        "fireproof": 0.000254,
                                        "nonflammable": 0.000359,
                                        "combustible": 0.000483,
                                        "flammable": 0.000916
                                    },
                                    "F5990": {
                                        "fireproof": 0.000306,
                                        "nonflammable": 0.00043,
                                        "combustible": 0.00058,
                                        "flammable": 0.001099
                                    },
                                    "G6010": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000733
                                    },
                                    "G6020": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000731
                                    },
                                    "G6040": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000733
                                    },
                                    "G6030": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000733
                                    },
                                    "G6100": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000733
                                    },
                                    "G6200": {
                                        "fireproof": 0.000203,
                                        "nonflammable": 0.000286,
                                        "combustible": 0.000386,
                                        "flammable": 0.000733
                                    },
                                    "S9300": {
                                        "fireproof": 0.000176,
                                        "nonflammable": 0.000247,
                                        "combustible": 0.000332,
                                        "flammable": 0.00063
                                    },
                                    "S9400": {
                                        "fireproof": 0.000176,
                                        "nonflammable": 0.000247,
                                        "combustible": 0.000332,
                                        "flammable": 0.00063
                                    },
                                    "S9490": {
                                        "fireproof": 0.000157,
                                        "nonflammable": 0.000221,
                                        "combustible": 0.000298,
                                        "flammable": 0.000565
                                    },
                                    "S9500": {
                                        "fireproof": 0.000153,
                                        "nonflammable": 0.000215,
                                        "combustible": 0.000291,
                                        "flammable": 0.00055
                                    },
                                    "S9600-1": {
                                        "fireproof": 0.000159,
                                        "nonflammable": 0.000224,
                                        "combustible": 0.000303,
                                        "flammable": 0.000572
                                    },
                                    "S9600-2": {
                                        "fireproof": 0.000207,
                                        "nonflammable": 0.000291,
                                        "combustible": 0.000392,
                                        "flammable": 0.000744
                                    },
                                    "S9600-3": {
                                        "fireproof": 0.000176,
                                        "nonflammable": 0.000247,
                                        "combustible": 0.000332,
                                        "flammable": 0.00063
                                    },
                                    "K": {
                                        "fireproof": 0.000153,
                                        "nonflammable": 0.000215,
                                        "combustible": 0.000289,
                                        "flammable": 0.000548
                                    },
                                    "H": {
                                        "fireproof": 0.000282,
                                        "nonflammable": 0.000397,
                                        "combustible": 0.000535,
                                        "flammable": 0.001015
                                    },
                                    "T": {
                                        "fireproof": 0.000153,
                                        "nonflammable": 0.000215,
                                        "combustible": 0.000289,
                                        "flammable": 0.000548
                                    }
                                }
                            }
                        },
                        {
                            "id": "typhoonFloodRain",
                            "name": "台风、洪水、暴雨",
                            "baseRateTable": {
                                "type": "zoneStructure",
                                "zoneLabels": {
                                    "zone1": "一类区（北京、青海、宁夏、天津）",
                                    "zone2": "二类区（陕西、山西、内蒙古、辽宁、大连、吉林、黑龙江、上海、山东、青岛、河南、重庆、甘肃、新疆、西藏）",
                                    "zone3": "三类区（江苏、湖北、河北、贵州、云南、四川、广东、深圳、海南）",
                                    "zone4": "四类区（安徽、湖南、广西、江西、浙江、宁波、福建、厦门）"
                                },
                                "structures": [
                                    {
                                        "id": "steel",
                                        "name": "钢结构"
                                    },
                                    {
                                        "id": "rc",
                                        "name": "钢筋混凝土"
                                    },
                                    {
                                        "id": "mixed",
                                        "name": "混合结构"
                                    },
                                    {
                                        "id": "brickWood",
                                        "name": "砖木结构"
                                    },
                                    {
                                        "id": "other",
                                        "name": "其他结构"
                                    }
                                ],
                                "data": {
                                    "zone1": {
                                        "steel": 4.8e-05,
                                        "rc": 7.3e-05,
                                        "mixed": 0.000121,
                                        "brickWood": 0.000218,
                                        "other": 0.000363
                                    },
                                    "zone2": {
                                        "steel": 6.1e-05,
                                        "rc": 9.1e-05,
                                        "mixed": 0.000151,
                                        "brickWood": 0.000273,
                                        "other": 0.000454
                                    },
                                    "zone3": {
                                        "steel": 7.3e-05,
                                        "rc": 0.000109,
                                        "mixed": 0.000182,
                                        "brickWood": 0.000327,
                                        "other": 0.000545
                                    },
                                    "zone4": {
                                        "steel": 9.1e-05,
                                        "rc": 0.000136,
                                        "mixed": 0.000227,
                                        "brickWood": 0.000409,
                                        "other": 0.000681
                                    }
                                }
                            }
                        },
                        {
                            "id": "earthquakeTsunami",
                            "name": "地震、海啸",
                            "baseRateTable": {
                                "type": "zoneStructure",
                                "zoneLabels": {
                                    "zone1": "一类区（黑龙江、吉林、贵州、湖北、湖南、江西、浙江、宁波、广东、深圳、广西）",
                                    "zone2": "二类区（内蒙古、山东、青岛、河南、安徽、福建、厦门、江苏、上海、重庆）",
                                    "zone3": "三类区（山西、陕西、辽宁、大连、海南）",
                                    "zone4": "四类区（北京、天津、河北、宁夏、甘肃、青海、四川）",
                                    "zone5": "五类区（新疆、西藏、云南）"
                                },
                                "structures": [
                                    {
                                        "id": "steel",
                                        "name": "钢结构"
                                    },
                                    {
                                        "id": "rc",
                                        "name": "钢筋混凝土"
                                    },
                                    {
                                        "id": "mixed",
                                        "name": "混合结构"
                                    },
                                    {
                                        "id": "brickWood",
                                        "name": "砖木结构"
                                    },
                                    {
                                        "id": "other",
                                        "name": "其他结构"
                                    }
                                ],
                                "data": {
                                    "zone1": {
                                        "steel": 3e-05,
                                        "rc": 3.8e-05,
                                        "mixed": 4.2e-05,
                                        "brickWood": 8.5e-05,
                                        "other": 0.000129
                                    },
                                    "zone2": {
                                        "steel": 7.7e-05,
                                        "rc": 9.2e-05,
                                        "mixed": 0.000104,
                                        "brickWood": 0.000214,
                                        "other": 0.00032
                                    },
                                    "zone3": {
                                        "steel": 0.000117,
                                        "rc": 0.000136,
                                        "mixed": 0.000156,
                                        "brickWood": 0.00032,
                                        "other": 0.00048
                                    },
                                    "zone4": {
                                        "steel": 0.000139,
                                        "rc": 0.000165,
                                        "mixed": 0.000188,
                                        "brickWood": 0.000382,
                                        "other": 0.000575
                                    },
                                    "zone5": {
                                        "steel": 0.000233,
                                        "rc": 0.000274,
                                        "mixed": 0.000313,
                                        "brickWood": 0.000639,
                                        "other": 0.000959
                                    }
                                }
                            }
                        },
                        {
                            "id": "otherNatural",
                            "name": "其它自然灾害",
                            "baseRateTable": {
                                "type": "manual"
                            },
                            "defaultRate": 3.8e-05,
                            "rateRange": [4e-05, 0.00016]
                        },
                        {
                            "id": "otherAccident",
                            "name": "其它意外事故",
                            "baseRateTable": {
                                "type": "manual"
                            },
                            "defaultRate": 6.1e-05,
                            "rateRange": [5e-05, 0.00015]
                        }
                    ],
                    "coefficients": [
                        {
                            "id": "fireAmountAdj",
                            "name": "火灾爆炸-保额调整系数",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "＜0.5亿",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[0.5, 2)亿",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[2, 10)亿",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[10, 20)亿",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥20亿",
                                    "value": 0.85,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "typhoonAmountAdj",
                            "name": "台风洪水暴雨-保额调整系数",
                            "applicableTo": ["typhoonFloodRain"],
                            "rows": [
                                {
                                    "parameter": "＜0.5亿",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[0.5, 2)亿",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[2, 10)亿",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[10, 20)亿",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥20亿",
                                    "value": 0.85,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "quakeAmountAdj",
                            "name": "地震海啸-保额调整系数",
                            "applicableTo": ["earthquakeTsunami"],
                            "rows": [
                                {
                                    "parameter": "＜0.5亿",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[0.5, 2)亿",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[2, 10)亿",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[10, 20)亿",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥20亿",
                                    "value": 0.85,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "typhoonThreat",
                            "name": "台风侵袭可能性",
                            "applicableTo": ["typhoonFloodRain"],
                            "rows": [
                                {
                                    "parameter": "台风登陆区域",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "台风影响区域",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "台风无影响区域",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "floodThreat",
                            "name": "洪水侵袭可能性",
                            "applicableTo": ["typhoonFloodRain"],
                            "rows": [
                                {
                                    "parameter": "危险程度高地区",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "危险程度较高地区",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "危险程度较小地区",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "quakeThreat",
                            "name": "地震侵袭可能性",
                            "applicableTo": ["earthquakeTsunami"],
                            "rows": [
                                {
                                    "parameter": "高于当地抗震设计标准",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "符合当地抗震设计标准",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "低于当地抗震设计标准",
                                    "value": 1.05,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "otherNatThreat",
                            "name": "其他自然灾害威胁",
                            "applicableTo": ["otherNatural"],
                            "rows": [
                                {
                                    "parameter": "威胁高",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "一般威胁",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "威胁低",
                                    "value": 0.85,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "fireTeamTime",
                            "name": "消防队到达时间",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "5分钟以内",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "6-15分钟",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "16-30分钟",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "30分钟以上",
                                    "value": 1.1,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "lossRecord",
                            "name": "损失记录",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜40%（N=1年）",
                                    "value": 0.75,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "＜40%（N=2年）",
                                    "value": 0.7,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "＜40%（N≥3年）",
                                    "value": 0.65,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N=1年）",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N=2年）",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N≥3年）",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N=1年）",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N=2年）",
                                    "value": 1.02,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N≥3年）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N=1年）",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N=2年）",
                                    "value": 1.12,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N≥3年）",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[70%,90%)（任意年限）",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥90%（任意年限）",
                                    "min": 1.3,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "naturalDisasterMeasure",
                            "name": "自然灾害防范措施",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "有防范措施（效率高）",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "有防范措施（效率一般）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "无防范措施",
                                    "value": 1.15,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "safetyAwareness",
                            "name": "安全意识",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "好",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "一般",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "safetyFacility",
                            "name": "安全设施",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "有防范措施（效率高）",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "有防范措施（效率一般）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "无防范措施",
                                    "value": 1.15,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductibleAmount",
                            "name": "免赔额调整系数",
                            "linkedGroup": "deductible",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜1000元",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[1000, 2000)元",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[2000, 5000)元",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[5000, 8000)元",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[8000, 12000)元",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥12000元",
                                    "value": 0.8,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "deductibleRate",
                            "name": "免赔率调整系数",
                            "linkedGroup": "deductible",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜5%",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[5%, 10%)",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[10%, 15%)",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[15%, 20%)",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "≥20%",
                                    "value": 0.75,
                                    "type": "fixed"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    "petrochemical": {
            "productName": "石油化工企业财产一切险",
            "productType": "multiRiskSum",
            "amountUnit": "元",
            "amountLabel": "保险金额",
            "premiumCap": None,
            "formulaText": "年保费 = 保险金额 × Σ(灾因基准费率 × 灾因调整系数乘积 × 赔付率调整系数)",
            "formulaNote": "适用于石油化工企业整体投保；多产品企业按产量加权计算生产类型及规模系数",
            "versions": {
                "original": {
                    "label": "石油化工企业财产一切险费率",
                    "risks": [
                        {
                            "id": "fireExplosion",
                            "name": "火灾及爆炸",
                            "baseRateTable": {
                                "type": "plantType",
                                "data": {
                                    "炼油厂": 0.00061,
                                    "石油化工厂": 0.00045
                                }
                            },
                            "deductibleNote": "基准免赔额：50万元或损失的5%，两者以高者为准"
                        },
                        {
                            "id": "rainFlood",
                            "name": "暴雨及洪水",
                            "baseRateTable": {
                                "type": "plantType",
                                "data": {
                                    "炼油厂": 0.00027,
                                    "石油化工厂": 0.00023
                                }
                            },
                            "deductibleNote": "基准免赔额：50万或损失的5%，两者以高者为准"
                        },
                        {
                            "id": "windNatural",
                            "name": "风灾及其他自然灾害",
                            "baseRateTable": {
                                "type": "plantType",
                                "data": {
                                    "炼油厂": 0.00015,
                                    "石油化工厂": 0.00015
                                }
                            },
                            "deductibleNote": "基准免赔额：15万或损失的5%，两者以高者为准"
                        },
                        {
                            "id": "humanError",
                            "name": "人工疏忽及失误",
                            "baseRateTable": {
                                "type": "plantType",
                                "data": {
                                    "炼油厂": 0.0003,
                                    "石油化工厂": 0.00027
                                }
                            },
                            "deductibleNote": "基准免赔额：20万或损失的5%，两者以高者为准"
                        }
                    ],
                    "coefficients": [
                        {
                            "id": "productionTypeScale",
                            "name": "生产类型及规模",
                            "applicableTo": ["fireExplosion", "humanError"],
                            "rows": [
                                {
                                    "parameter": "炼油厂 年加工量＜300万吨",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "炼油厂 年加工量300-800万吨",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "炼油厂 年加工量≥800万吨",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合成树脂及塑料＜20万吨",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合成树脂及塑料≥20万吨",
                                    "value": 1.03,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合纤原料及聚合物/合成纤维＜20万吨",
                                    "value": 1.03,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合纤原料及聚合物/合成纤维≥20万吨",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合成橡胶＜20万吨",
                                    "value": 0.97,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 合成橡胶≥20万吨",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 基本有机原料＜30万吨",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 基本有机原料≥30万吨",
                                    "value": 1.03,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "石化厂 其它产品",
                                    "value": 1.05,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "ethyleneCracking",
                            "name": "乙烯裂解",
                            "applicableTo": ["fireExplosion", "humanError"],
                            "rows": [
                                {
                                    "parameter": "生产流程中包含乙烯裂解",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "生产流程中不包含乙烯裂解",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "equipAge",
                            "name": "设备新旧程度",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "投产后1年以下",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "投产后1-8年",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "投产后8年及以上",
                                    "value": 1.05,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "equipMaintenance",
                            "name": "设备大修情况",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "能按规定定期大修",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "不能按规定定期大修",
                                    "min": 1.1,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireAccident5yr",
                            "name": "5年内火灾事故",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "发生过两次及以上重特大火灾事故",
                                    "value": 1.3,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "发生过一次重特大火灾事故",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "未发生重特大火灾事故",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "未发生火灾事故",
                                    "value": 0.9,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "fireDeductAmount",
                            "name": "火灾爆炸-免赔额",
                            "linkedGroup": "fireDeductible",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "基准免赔额",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的2倍",
                                    "value": 0.925,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的5倍",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的10倍",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的30倍及以上",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "fireDeductRate",
                            "name": "火灾爆炸-免赔率",
                            "linkedGroup": "fireDeductible",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "5%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "enterpriseType",
                            "name": "企业类型",
                            "applicableTo": ["fireExplosion", "humanError"],
                            "rows": [
                                {
                                    "parameter": "国际跨国公司独资或合资大型石化企业",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "国有大型石油公司所属石化企业",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "其它国有石化企业",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "民营石化企业",
                                    "value": 1.05,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "lightningFacility",
                            "name": "避雷设施",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "符合安全生产相关规定",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "不符合安全生产相关规定",
                                    "min": 1.1,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireTeam",
                            "name": "专业消防队",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "企业内组建专业消防队并配备专业消防车",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "未组建专业消防队或未配备专业消防车",
                                    "min": 1.1,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "surroundingEnv",
                            "name": "周边环境",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "周围不存在火灾爆炸隐患",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "周围存在火灾爆炸隐患",
                                    "min": 1.1,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "fireSafety",
                            "name": "消防安全",
                            "applicableTo": ["fireExplosion"],
                            "rows": [
                                {
                                    "parameter": "完全满足安全生产相关规定",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "不能完全满足安全生产相关规定",
                                    "min": 1.1,
                                    "max": 1.3,
                                    "type": "range"
                                }
                            ]
                        },
                        {
                            "id": "rainfall20yr",
                            "name": "近20年最大日降水量",
                            "applicableTo": ["rainFlood"],
                            "rows": [
                                {
                                    "parameter": "500毫米以上",
                                    "value": 1.08,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "400-500毫米",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "200-400毫米",
                                    "value": 1.02,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "100-200毫米",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "50-100毫米",
                                    "value": 0.98,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "50毫米以下",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "waterProximity",
                            "name": "周边水域",
                            "applicableTo": ["rainFlood"],
                            "rows": [
                                {
                                    "parameter": "5公里范围内有水域",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "5公里范围内无水域",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "mudslide10yr",
                            "name": "近10年泥石流",
                            "applicableTo": ["rainFlood"],
                            "rows": [
                                {
                                    "parameter": "从未发生过",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "发生过一次",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "发生过两次及以上",
                                    "value": 1.2,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "rainDeductAmount",
                            "name": "暴雨洪水-免赔额",
                            "linkedGroup": "rainDeductible",
                            "applicableTo": ["rainFlood"],
                            "rows": [
                                {
                                    "parameter": "基准免赔额",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的2倍",
                                    "value": 0.925,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的5倍",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的10倍",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的30倍及以上",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "rainDeductRate",
                            "name": "暴雨洪水-免赔率",
                            "linkedGroup": "rainDeductible",
                            "applicableTo": ["rainFlood"],
                            "rows": [
                                {
                                    "parameter": "5%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "windProvince",
                            "name": "风灾-工厂所在省份",
                            "applicableTo": ["windNatural"],
                            "rows": [
                                {
                                    "parameter": "浙江、广东、福建、海南",
                                    "value": 1.5,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "广西、辽宁、江苏",
                                    "value": 1.3,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "湖南、吉林、黑龙江、山东",
                                    "value": 1.2,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "陕西、山西、四川、甘肃、云南、河南",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "其它省份",
                                    "value": 1.0,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "windDeductAmount",
                            "name": "风灾-免赔额",
                            "linkedGroup": "windDeductible",
                            "applicableTo": ["windNatural"],
                            "rows": [
                                {
                                    "parameter": "基准免赔额",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的2倍",
                                    "value": 0.925,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的5倍",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的10倍",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的30倍及以上",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "windDeductRate",
                            "name": "风灾-免赔率",
                            "linkedGroup": "windDeductible",
                            "applicableTo": ["windNatural"],
                            "rows": [
                                {
                                    "parameter": "5%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "humanDeductAmount",
                            "name": "人工疏忽-免赔额",
                            "linkedGroup": "humanDeductible",
                            "applicableTo": ["humanError"],
                            "rows": [
                                {
                                    "parameter": "基准免赔额",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的2倍",
                                    "value": 0.925,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的5倍",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的10倍",
                                    "value": 0.8,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "基准免赔额的30倍及以上",
                                    "value": 0.7,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "humanDeductRate",
                            "name": "人工疏忽-免赔率",
                            "linkedGroup": "humanDeductible",
                            "applicableTo": ["humanError"],
                            "rows": [
                                {
                                    "parameter": "5%",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "10%",
                                    "value": 0.95,
                                    "type": "fixed"
                                }
                            ]
                        },
                        {
                            "id": "lossRecord",
                            "name": "赔付率调整系数",
                            "applicableTo": ["all"],
                            "rows": [
                                {
                                    "parameter": "＜40%（N=1年）",
                                    "value": 0.75,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "＜40%（N=2年）",
                                    "value": 0.7,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "＜40%（N≥3年）",
                                    "value": 0.65,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N=1年）",
                                    "value": 0.95,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N=2年）",
                                    "value": 0.9,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[40%,50%)（N≥3年）",
                                    "value": 0.85,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N=1年）",
                                    "value": 1.05,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N=2年）",
                                    "value": 1.02,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[50%,60%)（N≥3年）",
                                    "value": 1.0,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N=1年）",
                                    "value": 1.15,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N=2年）",
                                    "value": 1.12,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[60%,70%)（N≥3年）",
                                    "value": 1.1,
                                    "type": "fixed"
                                },
                                {
                                    "parameter": "[70%,90%)（任意年限）",
                                    "min": 1.15,
                                    "max": 1.3,
                                    "type": "range"
                                },
                                {
                                    "parameter": "≥90%（任意年限）",
                                    "min": 1.3,
                                    "max": 1.5,
                                    "type": "range"
                                }
                            ]
                        }
                    ]
                }
            }
        },
    # =============================================
    # 网络安全保险系列（3个险种）
    # =============================================
    "cyberSecurityA2025": {
        "productName": "网络安全保险（2025版A款）",
        "productType": "multiRiskSum",
        "amountUnit": "元",
        "amountLabel": "各项责任赔偿限额",
        "formulaText": "各项责任年保险费＝赔偿限额×对应基准费率×适用系数乘积；总保险费为各项保险费之和",
        "formulaNote": "短期承保保险费按条款所附短期费率表计收；免赔期/免赔率同时约定时取低者作为免赔调整系数",
        "versions": {
            "original": {
                "label": "网络安全保险（2025版A款）费率",
                "risks": [
                    {"id": "emergencyDetection", "name": "应急响应-检测鉴定费用", "baseRate": 0.0098},
                    {"id": "emergencyPR", "name": "应急响应-名誉恢复公关费用", "baseRate": 0.0074},
                    {"id": "emergencyNotify", "name": "应急响应-通知费用", "baseRate": 0.0008},
                    {"id": "businessInterruption", "name": "营业中断损失", "baseRate": 0.0003},
                    {"id": "dataRecovery", "name": "数据恢复及硬件维修费用", "baseRate": 0.0009},
                    {"id": "cyberExtortion", "name": "网络勒索损失", "baseRate": 0.0002}
                ],
                "coefficients": [
                    {
                        "id": "csaSecurityLevel", "name": "信息安全管理水平调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "较高", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "一般", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "较低", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csaLossRatio", "name": "赔付率调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "[0,20%]", "min": 0.5, "max": 0.6, "type": "range"},
                            {"parameter": "(20%,45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                            {"parameter": "(45%,70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(70%,95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": ">95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csaBiDeductPeriod", "name": "免赔期调整系数（营业中断）",
                        "applicableTo": ["businessInterruption"],
                        "linkedGroup": "csaBiDeductible",
                        "note": "免赔期/免赔率同时约定时取低者；未列明可线性插值",
                        "rows": [
                            {"parameter": "0天", "value": 1.0, "type": "fixed"},
                            {"parameter": "5天", "value": 0.9, "type": "fixed"},
                            {"parameter": "10天", "value": 0.8, "type": "fixed"},
                            {"parameter": "≥15天", "min": 0.7, "max": 0.75, "type": "range"}
                        ]
                    },
                    {
                        "id": "csaBiDeductRate", "name": "免赔率调整系数（营业中断）",
                        "applicableTo": ["businessInterruption"],
                        "linkedGroup": "csaBiDeductible",
                        "note": "免赔期/免赔率同时约定时取低者",
                        "rows": [
                            {"parameter": "0", "value": 1.0, "type": "fixed"},
                            {"parameter": "10%", "value": 0.9, "type": "fixed"},
                            {"parameter": "20%", "value": 0.8, "type": "fixed"},
                            {"parameter": "30%", "value": 0.7, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "csaOtherDeductAmt", "name": "免赔额调整系数（其他责任）",
                        "applicableTo": ["emergencyDetection", "emergencyPR", "emergencyNotify", "dataRecovery", "cyberExtortion"],
                        "linkedGroup": "csaOtherDeductible",
                        "note": "免赔额/免赔率同时约定时取低者；未列明可线性插值",
                        "rows": [
                            {"parameter": "0万", "value": 1.0, "type": "fixed"},
                            {"parameter": "2万", "value": 0.9, "type": "fixed"},
                            {"parameter": "5万", "value": 0.8, "type": "fixed"},
                            {"parameter": "≥7万", "min": 0.7, "max": 0.75, "type": "range"}
                        ]
                    },
                    {
                        "id": "csaOtherDeductRate", "name": "免赔率调整系数（其他责任）",
                        "applicableTo": ["emergencyDetection", "emergencyPR", "emergencyNotify", "dataRecovery", "cyberExtortion"],
                        "linkedGroup": "csaOtherDeductible",
                        "note": "免赔额/免赔率同时约定时取低者",
                        "rows": [
                            {"parameter": "0", "value": 1.0, "type": "fixed"},
                            {"parameter": "10%", "value": 0.9, "type": "fixed"},
                            {"parameter": "20%", "value": 0.8, "type": "fixed"},
                            {"parameter": "30%", "value": 0.7, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "csaBiCompPeriod", "name": "赔偿期调整系数（营业中断）",
                        "applicableTo": ["businessInterruption"],
                        "note": "未列明赔偿期可按线性插值法计算",
                        "rows": [
                            {"parameter": "1个月", "value": 1.0, "type": "fixed"},
                            {"parameter": "3个月", "value": 1.2, "type": "fixed"},
                            {"parameter": "6个月", "value": 1.4, "type": "fixed"},
                            {"parameter": "9个月", "value": 1.6, "type": "fixed"},
                            {"parameter": "12个月", "value": 1.8, "type": "fixed"}
                        ]
                    }
                ]
            }
        }
    },
    "cyberSecurityB2025": {
        "productName": "企业网络安全保险（2025版B款）",
        "productType": "multiRiskSum",
        "amountUnit": "元",
        "amountLabel": "各项责任赔偿限额",
        "formulaText": "各项责任年保险费＝赔偿限额×基准费率×适用系数乘积；应对费用＝(IT责任费+网络中断费)×10%×减额赔付比例系数",
        "formulaNote": "短期承保保险费＝年保险费×保险期间天数÷365；总保险费为各项保险费之和",
        "versions": {
            "original": {
                "label": "企业网络安全保险（2025版B款）费率",
                "risks": [
                    {"id": "itBehavior", "name": "IT业务或IT用户行为责任", "baseRate": 0.000106},
                    {"id": "networkInterruption", "name": "网络中断责任", "baseRate": 0.000232},
                    {"id": "incidentResponse", "name": "网络安全事故应对费用责任",
                     "derivedFrom": ["itBehavior", "networkInterruption"], "derivedRate": 0.10},
                    {"id": "propertyLoss", "name": "被保险人管理下的财产损失责任", "baseRate": 0.000053}
                ],
                "coefficients": [
                    {
                        "id": "csbRetroperiod", "name": "追溯期调整系数",
                        "applicableTo": ["itBehavior"],
                        "note": "追溯期长度/保险期间长度；未列明可线性插值",
                        "rows": [
                            {"parameter": "0", "value": 1.0, "type": "fixed"},
                            {"parameter": "1", "value": 1.6, "type": "fixed"},
                            {"parameter": "2", "value": 2.1, "type": "fixed"},
                            {"parameter": "≥3", "min": 2.5, "max": 3.0, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbItDeductAmt", "name": "每次事故免赔额调整系数（IT责任）",
                        "applicableTo": ["itBehavior"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤1万", "min": 1.02, "max": 1.05, "type": "range"},
                            {"parameter": "2万", "value": 1.0, "type": "fixed"},
                            {"parameter": "5万", "value": 0.98, "type": "fixed"},
                            {"parameter": "≥10万", "min": 0.9, "max": 0.95, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbBusinessScope", "name": "经营范围调整系数",
                        "applicableTo": ["itBehavior"],
                        "rows": [
                            {"parameter": "无IT业务", "value": 1.0, "type": "fixed"},
                            {"parameter": "有IT业务", "min": 1.1, "max": 1.5, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbProfitCompPeriod", "name": "利润赔偿期间调整系数",
                        "applicableTo": ["networkInterruption"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤30天", "min": 0.9, "max": 1.0, "type": "range"},
                            {"parameter": "60天", "value": 1.15, "type": "fixed"},
                            {"parameter": "90天", "value": 1.3, "type": "fixed"},
                            {"parameter": "≥120天", "min": 1.4, "max": 1.7, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbContinuityPeriod", "name": "经营延续费用恢复期间调整系数",
                        "applicableTo": ["networkInterruption"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤15天", "min": 0.9, "max": 0.95, "type": "range"},
                            {"parameter": "30天", "value": 1.0, "type": "fixed"},
                            {"parameter": "45天", "value": 1.05, "type": "fixed"},
                            {"parameter": "≥60天", "min": 1.1, "max": 1.3, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbProfitDeductAmt", "name": "利润免赔额调整系数",
                        "applicableTo": ["networkInterruption"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤1万", "min": 1.02, "max": 1.05, "type": "range"},
                            {"parameter": "2万", "value": 1.0, "type": "fixed"},
                            {"parameter": "5万", "value": 0.98, "type": "fixed"},
                            {"parameter": "≥10万", "min": 0.9, "max": 0.95, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbContinuityDeductAmt", "name": "经营延续费用免赔额调整系数",
                        "applicableTo": ["networkInterruption"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤2000元", "min": 1.02, "max": 1.05, "type": "range"},
                            {"parameter": "5000元", "value": 1.0, "type": "fixed"},
                            {"parameter": "10000元", "value": 0.98, "type": "fixed"},
                            {"parameter": "≥15000元", "min": 0.9, "max": 0.95, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbWaitingPeriod", "name": "免赔期调整系数",
                        "applicableTo": ["networkInterruption"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "0天", "value": 1.5, "type": "fixed"},
                            {"parameter": "1天", "value": 1.4, "type": "fixed"},
                            {"parameter": "3天", "value": 1.2, "type": "fixed"},
                            {"parameter": "5天", "value": 1.0, "type": "fixed"},
                            {"parameter": "≥10天", "min": 0.7, "max": 0.8, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbInsuredRatio", "name": "投保比率调整系数",
                        "applicableTo": ["networkInterruption"],
                        "rows": [
                            {"parameter": "100%", "value": 1.0, "type": "fixed"},
                            {"parameter": "90%", "value": 0.9, "type": "fixed"},
                            {"parameter": "80%", "value": 0.8, "type": "fixed"},
                            {"parameter": "70%", "value": 0.7, "type": "fixed"},
                            {"parameter": "60%", "value": 0.6, "type": "fixed"},
                            {"parameter": "50%", "value": 0.5, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "csbReductionRatio", "name": "减额赔付比例调整系数",
                        "applicableTo": ["incidentResponse"],
                        "rows": [
                            {"parameter": "100%", "value": 1.0, "type": "fixed"},
                            {"parameter": "90%", "value": 0.9, "type": "fixed"},
                            {"parameter": "80%", "value": 0.8, "type": "fixed"},
                            {"parameter": "70%", "value": 0.7, "type": "fixed"},
                            {"parameter": "60%", "value": 0.6, "type": "fixed"},
                            {"parameter": "50%", "value": 0.5, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "csbPropertyDeductAmt", "name": "每次事故免赔额调整系数（财产损失）",
                        "applicableTo": ["propertyLoss"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤1万", "min": 1.05, "max": 1.1, "type": "range"},
                            {"parameter": "2万", "value": 1.0, "type": "fixed"},
                            {"parameter": "5万", "value": 0.95, "type": "fixed"},
                            {"parameter": "≥10万", "min": 0.8, "max": 0.9, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbSecurityLevel", "name": "信息安全管理水平调整系数",
                        "applicableTo": ["itBehavior", "networkInterruption", "propertyLoss"],
                        "rows": [
                            {"parameter": "制度和防护完善", "min": 0.7, "max": 0.9, "type": "range"},
                            {"parameter": "制度和防护较完善", "min": 0.9, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "制度和防护不完善", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbDataCategory", "name": "信息类别和数量调整系数",
                        "applicableTo": ["itBehavior", "networkInterruption", "propertyLoss"],
                        "rows": [
                            {"parameter": "数量较少/重要程度较低", "min": 0.7, "max": 1.0, "type": "range"},
                            {"parameter": "数量较大/重要程度较高", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": "数量大/重要程度高", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbHistoryLoss", "name": "历史事故与损失情况调整系数",
                        "applicableTo": ["itBehavior", "networkInterruption", "propertyLoss"],
                        "rows": [
                            {"parameter": "极少", "min": 0.5, "max": 0.7, "type": "range"},
                            {"parameter": "较少", "min": 0.7, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "一般", "min": 1.0, "max": 1.3, "minExclusive": True, "type": "range"},
                            {"parameter": "较多", "min": 1.3, "max": 1.5, "minExclusive": True, "type": "range"},
                            {"parameter": "很多", "min": 1.5, "max": 2.0, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbRevenue", "name": "业务收入调整系数",
                        "applicableTo": ["itBehavior", "networkInterruption", "propertyLoss"],
                        "rows": [
                            {"parameter": "低于同业平均", "min": 0.7, "max": 0.9, "type": "range"},
                            {"parameter": "接近同业平均", "min": 0.9, "max": 1.1, "minExclusive": True, "type": "range"},
                            {"parameter": "高于同业平均", "min": 1.1, "max": 1.3, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "csbLossRatio", "name": "赔付率调整系数",
                        "applicableTo": ["itBehavior", "networkInterruption", "propertyLoss"],
                        "rows": [
                            {"parameter": "[0,20%]", "min": 0.5, "max": 0.6, "type": "range"},
                            {"parameter": "(20%,45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                            {"parameter": "(45%,70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(70%,95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": ">95%", "min": 1.2, "max": 1.5, "minExclusive": True, "type": "range"}
                        ]
                    }
                ]
            }
        }
    },
    "cyberSecurityPH2025": {
        "productName": "网络安全保险（普惠2025版·上海）",
        "productType": "property",
        "amountUnit": "万元",
        "amountLabel": "累计赔偿限额",
        "formulaText": "年保险费＝累计赔偿限额×基准费率×各项费率调整系数的乘积",
        "formulaNote": "短期承保保险费按条款所附短期费率表计收；分期缴费每期保险费＝年保险费÷分期期数",
        "versions": {
            "original": {
                "label": "网络安全保险（普惠2025版·上海）费率",
                "baseRates": {
                    "default": 0.0104
                },
                "coefficients": [
                    {
                        "id": "cspDeductAmt", "name": "免赔额调整系数",
                        "applicableTo": ["all"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "0元", "value": 1.1, "type": "fixed"},
                            {"parameter": "1000元", "value": 1.0, "type": "fixed"},
                            {"parameter": "2000元", "value": 0.95, "type": "fixed"},
                            {"parameter": "5000元", "value": 0.9, "type": "fixed"},
                            {"parameter": "10000元", "value": 0.85, "type": "fixed"},
                            {"parameter": "15000元", "value": 0.8, "type": "fixed"},
                            {"parameter": "20000元", "value": 0.75, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cspPerAccidentLimit", "name": "每次事故赔偿限额调整系数",
                        "applicableTo": ["all"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤1000元", "value": 0.9, "type": "fixed"},
                            {"parameter": "2000元", "value": 1.0, "type": "fixed"},
                            {"parameter": "3000元", "value": 1.1, "type": "fixed"},
                            {"parameter": "5000元", "value": 1.2, "type": "fixed"},
                            {"parameter": "10000元", "value": 1.5, "type": "fixed"},
                            {"parameter": "≥100000元", "value": 2.0, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cspCumulativeLimit", "name": "累计赔偿限额调整系数",
                        "applicableTo": ["all"],
                        "note": "未列明可线性插值",
                        "rows": [
                            {"parameter": "≤1万", "value": 2.0, "type": "fixed"},
                            {"parameter": "2万", "value": 1.5, "type": "fixed"},
                            {"parameter": "5万", "value": 1.3, "type": "fixed"},
                            {"parameter": "10万", "value": 1.0, "type": "fixed"},
                            {"parameter": "50万", "value": 0.9, "type": "fixed"},
                            {"parameter": "100万", "value": 0.8, "type": "fixed"},
                            {"parameter": "≥500万", "value": 0.7, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cspLossRatio", "name": "赔付率调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "[0,20%]", "min": 0.5, "max": 0.6, "type": "range"},
                            {"parameter": "(20%,45%]", "min": 0.6, "max": 0.8, "minExclusive": True, "type": "range"},
                            {"parameter": "(45%,70%]", "min": 0.8, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(70%,95%]", "min": 1.0, "max": 1.2, "minExclusive": True, "type": "range"},
                            {"parameter": ">95%", "min": 1.2, "max": 2.0, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "cspRenewal", "name": "续保调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "新保", "value": 1.0, "type": "fixed"},
                            {"parameter": "续保一年", "value": 0.9, "type": "fixed"},
                            {"parameter": "续保两年或以上", "value": 0.8, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cspRiskMgmt", "name": "风险管理水平调整系数",
                        "applicableTo": ["all"],
                        "note": "根据风险管理规章制度健全程度、防灾防损设施完备程度、安全教育培训等综合判定",
                        "rows": [
                            {"parameter": "较好", "min": 0.7, "max": 0.9, "type": "range"},
                            {"parameter": "一般", "min": 0.9, "max": 1.1, "minExclusive": True, "type": "range"},
                            {"parameter": "较差", "min": 1.1, "max": 1.3, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "cspRevenue", "name": "营业收入调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "≤500万", "min": 0.7, "max": 0.8, "type": "range"},
                            {"parameter": "(500,1000]万", "min": 0.8, "max": 0.9, "minExclusive": True, "type": "range"},
                            {"parameter": "(1000,5000]万", "min": 0.9, "max": 1.0, "minExclusive": True, "type": "range"},
                            {"parameter": "(5000,10000]万", "min": 1.0, "max": 1.1, "minExclusive": True, "type": "range"},
                            {"parameter": ">10000万", "min": 1.1, "max": 1.2, "minExclusive": True, "type": "range"}
                        ]
                    },
                    {
                        "id": "cspIndustry", "name": "行业类型调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "电力、燃气及水的生产和供应业", "value": 1.15, "type": "fixed"},
                            {"parameter": "水利、环境和公共设施管理业", "value": 1.1, "type": "fixed"},
                            {"parameter": "信息传输、计算机服务和软件业", "value": 1.05, "type": "fixed"},
                            {"parameter": "住宿和餐饮业", "value": 1.0, "type": "fixed"},
                            {"parameter": "教育业", "value": 0.95, "type": "fixed"},
                            {"parameter": "金融业", "value": 0.9, "type": "fixed"},
                            {"parameter": "房地产业", "value": 0.85, "type": "fixed"},
                            {"parameter": "租赁和商务服务业", "value": 0.8, "type": "fixed"},
                            {"parameter": "卫生、社会保障和社会福利业", "value": 0.75, "type": "fixed"},
                            {"parameter": "批发和零售业", "value": 0.7, "type": "fixed"},
                            {"parameter": "制造业", "value": 0.65, "type": "fixed"},
                            {"parameter": "交通运输、仓储和邮政业", "value": 0.6, "type": "fixed"},
                            {"parameter": "居民服务和其他服务业", "value": 0.55, "type": "fixed"},
                            {"parameter": "其他", "value": 0.8, "type": "fixed"}
                        ]
                    },
                    {
                        "id": "cspInstallment", "name": "保费分期调整系数",
                        "applicableTo": ["all"],
                        "rows": [
                            {"parameter": "不分期", "value": 1.0, "type": "fixed"},
                            {"parameter": "分期", "value": 1.09, "type": "fixed"}
                        ]
                    }
                ]
            }
        }
    }
}

MC_DISABILITY_TABLES = {
    "table1": {
        "label": "附表1",
        "ratios": [
            {"level": "一级伤残", "pct": 100}, {"level": "二级伤残", "pct": 80},
            {"level": "三级伤残", "pct": 70}, {"level": "四级伤残", "pct": 60},
            {"level": "五级伤残", "pct": 50}, {"level": "六级伤残", "pct": 40},
            {"level": "七级伤残", "pct": 30}, {"level": "八级伤残", "pct": 20},
            {"level": "九级伤残", "pct": 10}, {"level": "十级伤残", "pct": 5}
        ]
    },
    "table2": {
        "label": "附表2",
        "ratios": [
            {"level": "一级伤残", "pct": 100}, {"level": "二级伤残", "pct": 80},
            {"level": "三级伤残", "pct": 65}, {"level": "四级伤残", "pct": 55},
            {"level": "五级伤残", "pct": 45}, {"level": "六级伤残", "pct": 25},
            {"level": "七级伤残", "pct": 15}, {"level": "八级伤残", "pct": 10},
            {"level": "九级伤残", "pct": 4}, {"level": "十级伤残", "pct": 1}
        ]
    },
    "table3": {
        "label": "附表3",
        "ratios": [
            {"level": "一级伤残", "pct": 100}, {"level": "二级伤残", "pct": 90},
            {"level": "三级伤残", "pct": 80}, {"level": "四级伤残", "pct": 70},
            {"level": "五级伤残", "pct": 60}, {"level": "六级伤残", "pct": 50},
            {"level": "七级伤残", "pct": 40}, {"level": "八级伤残", "pct": 30},
            {"level": "九级伤残", "pct": 20}, {"level": "十级伤残", "pct": 10}
        ]
    }
}

MC_DISABILITY_ADDON_OPTIONS = [
    {"group": "A", "label": "A组: 九级10% 十级5%", "p9": 10, "p10": 5, "coeff": {"table1": 1.000, "table2": 1.077, "table3": 0.924}},
    {"group": "A", "label": "A组: 九级8% 十级5%", "p9": 8, "p10": 5, "coeff": {"table1": 0.996, "table2": 1.073, "table3": 0.921}},
    {"group": "A", "label": "A组: 九级15% 十级5%", "p9": 15, "p10": 5, "coeff": {"table1": 1.009, "table2": 1.087, "table3": 0.933}},
    {"group": "A", "label": "A组: 九级10% 十级3%", "p9": 10, "p10": 3, "coeff": {"table1": 0.995, "table2": 1.072, "table3": 0.919}},
    {"group": "A", "label": "A组: 九级15% 十级3%", "p9": 15, "p10": 3, "coeff": {"table1": 1.004, "table2": 1.082, "table3": 0.928}},
    {"group": "B", "label": "B组: 九级4% 十级1%", "p9": 4, "p10": 1, "coeff": {"table1": 0.928, "table2": 1.000, "table3": 0.858}},
    {"group": "B", "label": "B组: 九级2% 十级1%", "p9": 2, "p10": 1, "coeff": {"table1": 0.924, "table2": 0.996, "table3": 0.854}},
    {"group": "B", "label": "B组: 九级3% 十级1%", "p9": 3, "p10": 1, "coeff": {"table1": 0.926, "table2": 0.998, "table3": 0.856}},
    {"group": "B", "label": "B组: 九级3% 十级2%", "p9": 3, "p10": 2, "coeff": {"table1": 0.929, "table2": 1.001, "table3": 0.858}},
    {"group": "B", "label": "B组: 九级4% 十级2%", "p9": 4, "p10": 2, "coeff": {"table1": 0.931, "table2": 1.003, "table3": 0.860}},
    {"group": "C", "label": "C组: 九级20% 十级10%", "p9": 20, "p10": 10, "coeff": {"table1": 1.082, "table2": 1.166, "table3": 1.000}},
    {"group": "C", "label": "C组: 九级15% 十级8%", "p9": 15, "p10": 8, "coeff": {"table1": 1.068, "table2": 1.151, "table3": 0.987}},
    {"group": "C", "label": "C组: 九级20% 十级8%", "p9": 20, "p10": 8, "coeff": {"table1": 1.077, "table2": 1.161, "table3": 0.995}},
    {"group": "C", "label": "C组: 九级15% 十级6%", "p9": 15, "p10": 6, "coeff": {"table1": 1.063, "table2": 1.145, "table3": 0.982}},
    {"group": "C", "label": "C组: 九级20% 十级6%", "p9": 20, "p10": 6, "coeff": {"table1": 1.072, "table2": 1.155, "table3": 0.991}}
]

MC_DISABILITY_GROUP_DESC = {
    "A": "二级80% 三级70% 四级60% 五级50% 六级40% 七级30% 八级20%",
    "B": "二级80% 三级65% 四级55% 五级45% 六级25% 七级15% 八级10%",
    "C": "二级90% 三级80% 四级70% 五级60% 六级50% 七级40% 八级30%"
}

MC_INDUSTRY_DATA = [
    {
        "className": "一类（低风险）", "classValue": "class1", "cssClass": "class1",
        "subs": [
            {"code": "1.1", "name": "金融保险业", "detail": "银行业、证券业、保险业、其他金融活动业"},
            {"code": "1.2", "name": "邮政电信业", "detail": "邮政业、电信和其他传输服务业"},
            {"code": "1.3", "name": "商务服务业", "detail": "租赁业、住宿业、餐饮业、批发业、零售业、仓储业"},
            {"code": "1.4", "name": "电子服务业", "detail": "计算机服务业、软件业"},
            {"code": "1.5", "name": "公共服务业", "detail": "城市公共交通业、社会保障业、社会福利业、教育、研究与试验发展、科技交流和推广服务业、专业技术业"},
            {"code": "1.6", "name": "影音服务业", "detail": "新闻出版业、广播、电视、电影和音像业、文化艺术业"},
            {"code": "1.7", "name": "其他服务业", "detail": "居民服务业、其他服务业"},
        ]
    },
    {
        "className": "二类（中风险）", "classValue": "class2", "cssClass": "class2",
        "subs": [
            {"code": "2.1", "name": "公共服务业", "detail": "房地产业、体育、娱乐业、水利管理业、环境管理业、公共设施管理业"},
            {"code": "2.2", "name": "食品加工业", "detail": "农副食品加工业、食品制造业、饮料制造业、烟草制品业"},
            {"code": "2.3", "name": "普通制造业", "detail": "纺织业、纺织服装/鞋/帽制造业、皮革/毛皮/羽绒及制品业、木材加工及木竹藤草制品业、家具制造业、造纸及纸制品业、印刷业、文教体育用品制造业、通信设备/计算机及其他电子设备制造业"},
            {"code": "2.4", "name": "农林牧渔业", "detail": "林业、农业、畜牧业、渔业、农林牧渔服务业"},
            {"code": "2.5", "name": "化工制造业", "detail": "化学纤维制造业、医药制造业、橡胶制品业、塑料制品业"},
            {"code": "2.6", "name": "机械制造业", "detail": "通用机械制造业、专用机械制造业、交通运输设备制造业、电气机械及器材制造业、仪器仪表及文化办公用机械制造业"},
            {"code": "2.7", "name": "金属及非金属制造业", "detail": "非金属矿物制品业、金属制品业"},
            {"code": "2.8", "name": "资源生产供应业", "detail": "废弃资源和废旧材料回收加工业、电力/热力的生产和供应业、燃气生产和供应业、水的生产和供应业"},
            {"code": "2.9", "name": "建筑安装及相关行业", "detail": "房屋和土木工程建筑业、建筑安装业、建筑装饰业、其他建筑业、地质勘查业"},
            {"code": "2.10", "name": "运输服务业", "detail": "铁路运输业、道路运输业、水上运输业、航空运输业、管道运输业、装卸搬运和其他运输服务业"},
        ]
    },
    {
        "className": "三类（高风险）", "classValue": "class3", "cssClass": "class3",
        "subs": [
            {"code": "3.1", "name": "石油化工加工业", "detail": "石油加工、炼焦及核燃料加工业、化学原料及化学制品制造业"},
            {"code": "3.2", "name": "金属冶炼及加工业", "detail": "黑色金属冶炼及压延加工业、有色金属冶炼及压延加工业"},
            {"code": "3.3", "name": "石油天然气及矿山开采业", "detail": "石油和天然气开采业、黑色金属矿采选业、有色金属矿采选业、非金属矿采选业、煤炭开采和洗选业、其他采矿业"},
        ]
    }
]


# =============================================
# 工具函数
# =============================================

def fmt_currency(num):
    """格式化货币"""
    if num is None or math.isnan(num):
        return "¥0.00"
    return f"¥{abs(num):,.2f}"


def fmt_num(num, digits=4):
    """格式化数字"""
    return f"{float(num):.{digits}f}"


def is_leap_year(year):
    """判断闰年"""
    return (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0)


# =============================================
# 通用样式
# =============================================

def get_common_styles():
    """返回通用控件样式"""
    return f"""
        QLabel {{
            color: {AnthropicColors.TEXT_PRIMARY};
        }}
        QComboBox {{
            background: {AnthropicColors.BG_PRIMARY};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px 12px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 13px;
            min-height: 20px;
            min-width: 100px;
        }}
        QComboBox:focus {{
            border-color: {AnthropicColors.ACCENT};
        }}
        QComboBox::drop-down {{
            border: none;
            width: 24px;
        }}
        QSpinBox, QDoubleSpinBox {{
            background: {AnthropicColors.BG_PRIMARY};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px 12px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 13px;
            min-height: 20px;
            min-width: 90px;
        }}
        QSpinBox:focus, QDoubleSpinBox:focus {{
            border-color: {AnthropicColors.ACCENT};
        }}
        QLineEdit {{
            background: {AnthropicColors.BG_PRIMARY};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px 12px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 13px;
        }}
        QLineEdit:focus {{
            border-color: {AnthropicColors.ACCENT};
        }}
        QPushButton {{
            background: {AnthropicColors.BG_CARD};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px 16px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-weight: 500;
            font-size: 13px;
        }}
        QPushButton:hover {{
            background: {AnthropicColors.BG_PRIMARY};
            border-color: {AnthropicColors.ACCENT};
        }}
        QTextEdit {{
            background: {AnthropicColors.BG_PRIMARY};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 12px;
        }}
        QListWidget {{
            color: {AnthropicColors.TEXT_PRIMARY};
        }}
        QListWidget::item {{
            color: {AnthropicColors.TEXT_PRIMARY};
        }}
        QScrollArea {{
            border: none;
            background: transparent;
        }}
        QDateEdit {{
            background: {AnthropicColors.BG_PRIMARY};
            border: 1px solid {AnthropicColors.BORDER};
            border-radius: 8px;
            padding: 8px 12px;
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 13px;
            min-height: 20px;
        }}
        QDateEdit:focus {{
            border-color: {AnthropicColors.ACCENT};
        }}
    """


def make_accent_button(text):
    """创建强调色按钮"""
    btn = QPushButton(text)
    btn.setCursor(Qt.PointingHandCursor)
    btn.setStyleSheet(f"""
        QPushButton {{
            background: {AnthropicColors.ACCENT};
            color: {AnthropicColors.TEXT_LIGHT};
            border: none;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 14px;
            font-weight: 600;
        }}
        QPushButton:hover {{
            background: {AnthropicColors.ACCENT_DARK};
        }}
    """)
    return btn


def make_success_button(text):
    """创建成功色按钮"""
    btn = QPushButton(text)
    btn.setCursor(Qt.PointingHandCursor)
    btn.setStyleSheet(f"""
        QPushButton {{
            background: #10b981;
            color: white;
            border: none;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 14px;
            font-weight: 600;
        }}
        QPushButton:hover {{
            background: #059669;
        }}
    """)
    return btn


# =============================================
# IndustryLookupDialog — 行业分类速查
# =============================================

class IndustryLookupDialog(QDialog):
    """行业分类速查对话框"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_class = None
        self.setWindowTitle("行业分类速查")
        self.setMinimumSize(520, 500)
        self.setStyleSheet(f"""
            QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; }}
        """)
        layout = QVBoxLayout(self)
        layout.setSpacing(10)

        # 搜索框
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("搜索行业名称或描述...")
        self.search_input.setStyleSheet(f"""
            QLineEdit {{ border: 1px solid {AnthropicColors.BORDER}; border-radius: 8px;
                padding: 8px 12px; font-size: 14px; background: white; }}
            QLineEdit:focus {{ border-color: {AnthropicColors.ACCENT}; }}
        """)
        self.search_input.textChanged.connect(self._filter)
        layout.addWidget(self.search_input)

        # 列表区
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        self.list_widget = QWidget()
        self.list_layout = QVBoxLayout(self.list_widget)
        self.list_layout.setContentsMargins(0, 0, 0, 0)
        self.list_layout.setSpacing(4)
        scroll.setWidget(self.list_widget)
        layout.addWidget(scroll, 1)

        self._render("")

    def _filter(self, text):
        self._render(text.strip().lower())

    def _render(self, keyword):
        while self.list_layout.count():
            item = self.list_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        risk_colors = {
            "class1": ("#fef9c3", "#854d0e"),
            "class2": ("#fed7aa", "#9a3412"),
            "class3": ("#fecaca", "#991b1b"),
        }

        for cls in MC_INDUSTRY_DATA:
            filtered = [s for s in cls["subs"]
                        if not keyword
                        or keyword in s["name"].lower()
                        or keyword in s["detail"].lower()
                        or keyword in s["code"]]
            if not filtered:
                continue

            bg_color, text_color = risk_colors.get(cls["cssClass"], ("#f0f0f0", "#333"))
            header = QLabel(f"{cls['className']}（共{len(cls['subs'])}个子类）")
            header.setStyleSheet(f"""
                QLabel {{ background: {bg_color}; color: {text_color}; font-weight: 600;
                    font-size: 13px; padding: 6px 12px; border-radius: 6px; }}
            """)
            self.list_layout.addWidget(header)

            for sub in filtered:
                btn = QPushButton(f"[{sub['code']}] {sub['name']}\n{sub['detail']}")
                btn.setCursor(Qt.PointingHandCursor)
                btn.setStyleSheet(f"""
                    QPushButton {{ text-align: left; padding: 8px 12px; border: 1px solid {AnthropicColors.BORDER};
                        border-radius: 6px; font-size: 12px; background: white; }}
                    QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; background: rgba(217, 119, 87, 0.05); }}
                """)
                btn.clicked.connect(lambda checked, cv=cls["classValue"]: self._select(cv))
                self.list_layout.addWidget(btn)

        self.list_layout.addStretch()

    def _select(self, class_value):
        self.selected_class = class_value
        self.accept()


# =============================================
# DisabilityAddonDialog — 伤残赔偿比例方案选择
# =============================================

class DisabilityAddonDialog(QDialog):
    """15列伤残赔偿比例方案选择对话框"""

    GROUP_DEFS = {
        "A": [100, 80, 70, 60, 50, 40, 30, 20],
        "B": [100, 80, 65, 55, 45, 25, 15, 10],
        "C": [100, 90, 80, 70, 60, 50, 40, 30],
    }
    GROUP_COLORS = {
        "A": ("#fef9c3", "#854d0e"),
        "B": ("#fed7aa", "#9a3412"),
        "C": ("#fecaca", "#991b1b"),
    }
    LEVELS = ["一级", "二级", "三级", "四级", "五级", "六级", "七级", "八级", "九级", "十级"]

    def __init__(self, table_key, current_option, parent=None):
        super().__init__(parent)
        self.table_key = table_key
        self.selected_option = current_option
        self.setWindowTitle("选择伤残赔偿比例方案")
        self.setMinimumSize(800, 480)
        self.setStyleSheet(f"QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}")
        layout = QVBoxLayout(self)

        hint = QLabel("点击任一列选择方案")
        hint.setAlignment(Qt.AlignCenter)
        hint.setStyleSheet(f"font-size: 13px; color: {AnthropicColors.TEXT_SECONDARY}; margin-bottom: 8px;")
        layout.addWidget(hint)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        table_widget = QWidget()
        table_layout = QGridLayout(table_widget)
        table_layout.setSpacing(1)
        scroll.setWidget(table_widget)
        layout.addWidget(scroll, 1)

        options = MC_DISABILITY_ADDON_OPTIONS
        col_count = len(options)

        # Header row
        header_label = QLabel("伤残等级")
        header_label.setStyleSheet("font-weight: 600; font-size: 11px; padding: 4px;")
        table_layout.addWidget(header_label, 0, 0)
        for ci, opt in enumerate(options):
            bg, fg = self.GROUP_COLORS[opt["group"]]
            is_sel = ci == current_option
            lbl = QLabel(f"{opt['group']}-{ci + 1}")
            lbl.setAlignment(Qt.AlignCenter)
            style = f"font-weight: 600; font-size: 11px; padding: 4px; background: {bg}; color: {fg};"
            if is_sel:
                style += f" border: 2px solid {AnthropicColors.ACCENT};"
            lbl.setStyleSheet(style)
            table_layout.addWidget(lbl, 0, ci + 1)

        # Data rows
        for li, level in enumerate(self.LEVELS):
            row_label = QLabel(f"{level}伤残")
            row_label.setStyleSheet("font-size: 11px; padding: 4px; font-weight: 500;")
            table_layout.addWidget(row_label, li + 1, 0)
            for ci, opt in enumerate(options):
                if li <= 7:
                    pct = self.GROUP_DEFS[opt["group"]][li]
                elif li == 8:
                    pct = opt["p9"]
                else:
                    pct = opt["p10"]
                bg, fg = self.GROUP_COLORS[opt["group"]]
                is_sel = ci == current_option
                btn = QPushButton(f"{pct}%")
                btn.setCursor(Qt.PointingHandCursor)
                style = f"""
                    QPushButton {{ font-size: 11px; padding: 3px; border: none; background: {bg}; color: {fg};
                        {'border: 2px solid ' + AnthropicColors.ACCENT + ';' if is_sel else ''} }}
                    QPushButton:hover {{ background: {AnthropicColors.ACCENT}; color: white; }}
                """
                btn.setStyleSheet(style)
                btn.clicked.connect(lambda checked, idx=ci: self._select_column(idx))
                table_layout.addWidget(btn, li + 1, ci + 1)

        # Coefficient rows
        coeff_keys = [("table1", "附表1系数"), ("table2", "附表2系数"), ("table3", "附表3系数")]
        for ki, (ck, ck_label) in enumerate(coeff_keys):
            row_label = QLabel(ck_label)
            row_label.setStyleSheet("font-size: 11px; padding: 4px; font-weight: 600;")
            table_layout.addWidget(row_label, len(self.LEVELS) + 1 + ki, 0)
            for ci, opt in enumerate(options):
                bg, fg = self.GROUP_COLORS[opt["group"]]
                is_sel = ci == current_option
                is_active_table = ck == table_key
                val = opt["coeff"][ck]
                btn = QPushButton(fmt_num(val, 3))
                btn.setCursor(Qt.PointingHandCursor)
                extra = ""
                if is_sel:
                    extra += f"border: 2px solid {AnthropicColors.ACCENT};"
                if is_active_table:
                    extra += "font-weight: 700;"
                btn.setStyleSheet(f"""
                    QPushButton {{ font-size: 11px; padding: 3px; border: none; background: {bg}; color: {fg}; {extra} }}
                    QPushButton:hover {{ background: {AnthropicColors.ACCENT}; color: white; }}
                """)
                btn.clicked.connect(lambda checked, idx=ci: self._select_column(idx))
                table_layout.addWidget(btn, len(self.LEVELS) + 1 + ki, ci + 1)

    def _select_column(self, col_idx):
        self.selected_option = col_idx
        self.accept()


# =============================================
# MainInsuranceTab — 主险计算器
# =============================================

class MainInsuranceTab(QWidget):
    """主险计算器 Tab"""
    premium_calculated = pyqtSignal(float, float)
    # 完整计算结果信号，传递给附加险使用
    full_result_calculated = pyqtSignal(object)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_product = "employerLiability"
        self.selected_version = "original"
        self.current_plan = MC_PRODUCTS["employerLiability"]["versions"]["original"]
        self.coeff_selections = {}
        self.selected_disability_table = "none"
        self.selected_disability_option = -1
        self.result = None
        self.mrs_widgets = {}
        self._setup_ui()

    def _setup_ui(self):
        self.setStyleSheet(get_common_styles())
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(12)
        main_layout.setContentsMargins(15, 10, 15, 10)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll_widget = QWidget()
        self.scroll_layout = QVBoxLayout(scroll_widget)
        self.scroll_layout.setSpacing(12)
        scroll.setWidget(scroll_widget)
        main_layout.addWidget(scroll, 1)

        self._build_control_bar()
        self._build_params_section()
        self._build_disability_section()
        self._build_coeff_section()
        self._build_action_buttons()
        self._build_result_section()
        self._build_log_section()
        self.scroll_layout.addStretch()

    def _build_control_bar(self):
        card = GlassCard()
        layout = QHBoxLayout(card)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.addWidget(QLabel("险种:"))
        self.product_combo = QComboBox()
        for pid, pdata in MC_PRODUCTS.items():
            self.product_combo.addItem(pdata["productName"], pid)
        self.product_combo.currentIndexChanged.connect(self._on_product_change)
        layout.addWidget(self.product_combo)
        layout.addWidget(QLabel("版本:"))
        self.version_combo = QComboBox()
        self._refresh_version_combo()
        self.version_combo.currentIndexChanged.connect(self._on_version_change)
        layout.addWidget(self.version_combo)
        layout.addStretch()
        import_btn = QPushButton("📂 导入费率方案")
        import_btn.setCursor(Qt.PointingHandCursor)
        import_btn.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        import_btn.setMinimumWidth(140)
        import_btn.clicked.connect(self._import_rate_plan)
        layout.addWidget(import_btn)
        self.scroll_layout.addWidget(card)

    def _refresh_version_combo(self, select_version=None):
        self.version_combo.blockSignals(True)
        self.version_combo.clear()
        product = MC_PRODUCTS.get(self.selected_product)
        if product:
            for vid, vdata in product["versions"].items():
                self.version_combo.addItem(vdata["label"], vid)
        if select_version:
            idx = self.version_combo.findData(select_version)
            if idx >= 0:
                self.version_combo.setCurrentIndex(idx)
        self.version_combo.blockSignals(False)
        self.selected_version = self.version_combo.currentData() or "original"
        self.current_plan = MC_PRODUCTS.get(self.selected_product, {}).get("versions", {}).get(self.selected_version, {})

    def _on_product_change(self):
        self.selected_product = self.product_combo.currentData()
        self._refresh_version_combo()
        self._update_params_visibility()
        self._on_version_change()
        self._log(f"切换险种: {MC_PRODUCTS[self.selected_product]['productName']}")

    def _on_version_change(self):
        self.selected_version = self.version_combo.currentData() or "original"
        product = MC_PRODUCTS.get(self.selected_product, {})
        self.current_plan = product.get("versions", {}).get(self.selected_version, {})
        self.coeff_selections = {}
        self.result = None
        self._clear_result()
        self._update_params_visibility()
        self._render_coefficients()
        if self.current_plan:
            self._log(f"切换版本: {self.current_plan.get('label', '')}")

    def _build_params_section(self):
        card = GlassCard()
        grid = QGridLayout(card)
        grid.setContentsMargins(16, 12, 16, 12)
        grid.setSpacing(10)

        self.industry_label_widget = QLabel("行业类别:")
        grid.addWidget(self.industry_label_widget, 0, 0)
        industry_row = QHBoxLayout()
        self.industry_combo = QComboBox()
        self.industry_combo.addItem("一类行业", "class1")
        self.industry_combo.addItem("二类行业", "class2")
        self.industry_combo.addItem("三类行业", "class3")
        self.industry_combo.currentIndexChanged.connect(self._on_industry_change)
        industry_row.addWidget(self.industry_combo, 1)
        self.industry_lookup_btn = QPushButton("速查")
        self.industry_lookup_btn.setCursor(Qt.PointingHandCursor)
        self.industry_lookup_btn.setFixedWidth(50)
        self.industry_lookup_btn.setStyleSheet(f"""
            QPushButton {{ background: {AnthropicColors.BG_CARD}; border: 1px solid {AnthropicColors.BORDER};
                border-radius: 6px; padding: 4px 8px; font-size: 12px; color: {AnthropicColors.ACCENT}; }}
            QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; }}
        """)
        self.industry_lookup_btn.clicked.connect(self._show_industry_lookup)
        industry_row.addWidget(self.industry_lookup_btn)
        grid.addLayout(industry_row, 0, 1)

        self.method_label_widget = QLabel("计费方式:")
        grid.addWidget(self.method_label_widget, 0, 2)
        self.method_combo = QComboBox()
        self.method_combo.addItem("固定限额", "fixed")
        self.method_combo.addItem("工资总额", "salary")
        self.method_combo.currentIndexChanged.connect(self._on_method_change)
        grid.addWidget(self.method_combo, 0, 3)

        self.limit_label = QLabel("每人限额(万元):")
        grid.addWidget(self.limit_label, 1, 0)
        self.limit_spin = QDoubleSpinBox()
        self.limit_spin.setRange(1, 10000)
        self.limit_spin.setValue(50)
        self.limit_spin.setDecimals(2)
        self.limit_spin.setSuffix(" 万元")
        grid.addWidget(self.limit_spin, 1, 1)

        self.death_limit_label = QLabel("死亡残疾限额(万元):")
        grid.addWidget(self.death_limit_label, 1, 2)
        self.death_limit_spin = QDoubleSpinBox()
        self.death_limit_spin.setRange(1, 10000)
        self.death_limit_spin.setValue(50)
        self.death_limit_spin.setDecimals(2)
        self.death_limit_spin.setSuffix(" 万元")
        grid.addWidget(self.death_limit_spin, 1, 3)
        self.death_limit_label.hide()
        self.death_limit_spin.hide()

        self.medical_limit_label = QLabel("医疗费用限额(万元):")
        grid.addWidget(self.medical_limit_label, 2, 0)
        self.medical_limit_spin = QDoubleSpinBox()
        self.medical_limit_spin.setRange(0.1, 10000)
        self.medical_limit_spin.setValue(5)
        self.medical_limit_spin.setDecimals(2)
        self.medical_limit_spin.setSuffix(" 万元")
        grid.addWidget(self.medical_limit_spin, 2, 1)
        self.medical_limit_label.hide()
        self.medical_limit_spin.hide()

        self.salary_label = QLabel("年度工资总额(元):")
        grid.addWidget(self.salary_label, 2, 2)
        self.salary_spin = QDoubleSpinBox()
        self.salary_spin.setRange(0, 999999999999)
        self.salary_spin.setValue(5000000)
        self.salary_spin.setDecimals(2)
        self.salary_spin.setSuffix(" 元")
        grid.addWidget(self.salary_spin, 2, 3)
        self.salary_label.hide()
        self.salary_spin.hide()

        self.count_label_widget = QLabel("承保人数:")
        grid.addWidget(self.count_label_widget, 3, 0)
        self.count_spin = QSpinBox()
        self.count_spin.setRange(1, 999999)
        self.count_spin.setValue(100)
        self.count_spin.setSuffix(" 人")
        grid.addWidget(self.count_spin, 3, 1)

        grid.addWidget(QLabel("保险期间:"), 3, 2)
        self.term_combo = QComboBox()
        self.term_combo.addItem("年度", "annual")
        self.term_combo.addItem("短期", "short")
        self.term_combo.currentIndexChanged.connect(self._on_term_change)
        grid.addWidget(self.term_combo, 3, 3)

        self.days_label = QLabel("保险天数:")
        grid.addWidget(self.days_label, 4, 0)
        self.days_spin = QSpinBox()
        self.days_spin.setRange(1, 365)
        self.days_spin.setValue(180)
        self.days_spin.setSuffix(" 天")
        grid.addWidget(self.days_spin, 4, 1)
        self.days_label.hide()
        self.days_spin.hide()

        # === 通用保险金额输入（property / interruption / composite / jewelry） ===
        self.amount_label = QLabel("保险金额(元):")
        grid.addWidget(self.amount_label, 5, 0)
        self.amount_spin = QDoubleSpinBox()
        self.amount_spin.setRange(0, 999999999999)
        self.amount_spin.setValue(10000000)
        self.amount_spin.setDecimals(2)
        self.amount_spin.setSuffix(" 元")
        grid.addWidget(self.amount_spin, 5, 1)
        self.amount_label.hide()
        self.amount_spin.hide()

        # === 恒力项目第二个保额（composite类型） ===
        self.sub_amount_label = QLabel("机器损坏保额:")
        grid.addWidget(self.sub_amount_label, 4, 2)
        self.sub_amount_spin = QDoubleSpinBox()
        self.sub_amount_spin.setRange(0, 999999999999)
        self.sub_amount_spin.setValue(5000000)
        self.sub_amount_spin.setDecimals(2)
        self.sub_amount_spin.setSuffix(" 元")
        grid.addWidget(self.sub_amount_spin, 4, 3)
        self.sub_amount_label.hide()
        self.sub_amount_spin.hide()

        # === 现金综合保险：平均每日现金营业额 ===
        self.daily_cash_label = QLabel("平均每日现金营业额:")
        grid.addWidget(self.daily_cash_label, 5, 0)
        self.daily_cash_spin = QDoubleSpinBox()
        self.daily_cash_spin.setRange(0, 999999999)
        self.daily_cash_spin.setValue(100000)
        self.daily_cash_spin.setDecimals(2)
        self.daily_cash_spin.setSuffix(" 元")
        grid.addWidget(self.daily_cash_spin, 5, 1)
        self.daily_cash_label.hide()
        self.daily_cash_spin.hide()

        # === 珠宝商：商户类型 + 保障类型 ===
        self.merchant_type_label = QLabel("商户类型:")
        grid.addWidget(self.merchant_type_label, 5, 2)
        self.merchant_type_combo = QComboBox()
        self.merchant_type_combo.addItem("批发商", "wholesale")
        self.merchant_type_combo.addItem("零售商", "retail")
        grid.addWidget(self.merchant_type_combo, 5, 3)
        self.merchant_type_label.hide()
        self.merchant_type_combo.hide()

        self.coverage_type_label = QLabel("保障类型:")
        grid.addWidget(self.coverage_type_label, 6, 0)
        self.coverage_type_combo = QComboBox()
        self.coverage_type_combo.addItem("货品损失-经营场所货品及现金", "goodsStore")
        self.coverage_type_combo.addItem("货品损失-场所外存储货品", "goodsOffsite")
        self.coverage_type_combo.addItem("货品损失-代保管货品", "goodsCustody")
        self.coverage_type_combo.addItem("营业财产损失保险", "businessProperty")
        grid.addWidget(self.coverage_type_combo, 6, 1)
        self.coverage_type_label.hide()
        self.coverage_type_combo.hide()

        # 基准费率和公式展示标签
        self.formula_label = QLabel("")
        self.formula_label.setWordWrap(True)
        self.formula_label.setStyleSheet("QLabel { background: rgba(59, 130, 246, 0.1); border: 1px solid rgba(59, 130, 246, 0.3); border-radius: 8px; padding: 10px 14px; font-size: 13px; color: #3b82f6; }")
        self.formula_label.setVisible(False)
        grid.addWidget(self.formula_label, 7, 0, 1, 4)

        # multiRiskSum 风险参数区
        self.mrs_container = QWidget()
        self.mrs_layout = QVBoxLayout(self.mrs_container)
        self.mrs_layout.setContentsMargins(0, 8, 0, 0)
        self.mrs_layout.setSpacing(6)
        self.mrs_container.setVisible(False)
        grid.addWidget(self.mrs_container, 8, 0, 1, 4)

        self.scroll_layout.addWidget(card)

    def _on_industry_change(self):
        self._update_formula_display()

    def _show_industry_lookup(self):
        dialog = IndustryLookupDialog(self)
        if dialog.exec_() == QDialog.Accepted and dialog.selected_class:
            class_map = {"class1": 0, "class2": 1, "class3": 2}
            idx = class_map.get(dialog.selected_class, 0)
            self.industry_combo.setCurrentIndex(idx)
            self._log(f"行业速查选择: {dialog.selected_class}")

    def _on_method_change(self):
        is_fixed = self.method_combo.currentData() == "fixed"
        self.limit_label.setVisible(is_fixed)
        self.limit_spin.setVisible(is_fixed)
        self.salary_label.setVisible(not is_fixed)
        self.salary_spin.setVisible(not is_fixed)
        self.coeff_selections = {}
        self._render_coefficients()
        self._update_formula_display()
        self._log(f"切换计费方式: {'固定限额' if is_fixed else '工资总额'}")

    def _on_term_change(self):
        is_short = self.term_combo.currentData() == "short"
        self.days_label.setVisible(is_short)
        self.days_spin.setVisible(is_short)

    def _get_product_type(self):
        """获取当前选中产品的 productType"""
        product = MC_PRODUCTS.get(self.selected_product, {})
        return product.get("productType", "liability")

    def _update_params_visibility(self):
        """根据产品类型动态显示/隐藏参数字段"""
        pt = self._get_product_type()
        product = MC_PRODUCTS.get(self.selected_product, {})
        is_liability = pt == "liability"
        is_property = pt == "property"
        is_composite = pt == "composite"
        is_interruption = pt == "interruption"
        is_jewelry = pt == "jewelry"
        is_multi_risk = pt == "multiRiskSum"

        # liability 专属字段
        is_fixed_only = is_liability and self.current_plan.get("fixedOnly", False)
        is_dual_rate = is_liability and self.current_plan.get("dualRate", False)
        is_fixed = self.method_combo.currentData() == "fixed"
        self.industry_label_widget.setVisible(is_liability)
        self.industry_combo.setVisible(is_liability)
        if hasattr(self, 'industry_lookup_btn'):
            self.industry_lookup_btn.setVisible(is_liability)
        # fixedOnly版本隐藏计费方式选择器
        self.method_label_widget.setVisible(is_liability and not is_fixed_only)
        self.method_combo.setVisible(is_liability and not is_fixed_only)
        if is_fixed_only:
            self.method_combo.setCurrentIndex(0)  # 强制fixed
        # 普通每人限额：非dualRate + fixed
        self.limit_label.setVisible(is_liability and is_fixed and not is_dual_rate)
        self.limit_spin.setVisible(is_liability and is_fixed and not is_dual_rate)
        # E款双限额
        self.death_limit_label.setVisible(is_liability and is_fixed and is_dual_rate)
        self.death_limit_spin.setVisible(is_liability and is_fixed and is_dual_rate)
        self.medical_limit_label.setVisible(is_liability and is_fixed and is_dual_rate)
        self.medical_limit_spin.setVisible(is_liability and is_fixed and is_dual_rate)
        self.salary_label.setVisible(is_liability and not is_fixed_only and self.method_combo.currentData() == "salary")
        self.salary_spin.setVisible(is_liability and not is_fixed_only and self.method_combo.currentData() == "salary")
        self.count_label_widget.setVisible(is_liability)
        self.count_spin.setVisible(is_liability)

        # 通用保险金额（元）
        show_amount = is_property or is_composite or is_interruption or is_jewelry or is_multi_risk
        self.amount_label.setVisible(show_amount)
        self.amount_spin.setVisible(show_amount)
        if show_amount:
            label_text = product.get("amountLabel", "保险金额")
            if is_composite:
                sub_risks = product.get("subRisks", ["materialDamage", "machineryBreakdown"])
                sub_labels = product.get("subRiskLabels", {})
                label_text = sub_labels.get(sub_risks[0], "物质损失或损坏一切险保额")
            self.amount_label.setText(f"{label_text}(元):")

        # composite 第二保额
        self.sub_amount_label.setVisible(is_composite)
        self.sub_amount_spin.setVisible(is_composite)

        # 现金综合保险额外字段
        has_daily_cash = "dailyCashTurnover" in product.get("extraFields", [])
        self.daily_cash_label.setVisible(has_daily_cash)
        self.daily_cash_spin.setVisible(has_daily_cash)

        # 珠宝商字段
        self.merchant_type_label.setVisible(is_jewelry)
        self.merchant_type_combo.setVisible(is_jewelry)
        self.coverage_type_label.setVisible(is_jewelry)
        self.coverage_type_combo.setVisible(is_jewelry)

        # multiRiskSum 风险参数区
        self._build_mrs_params(is_multi_risk)

        # 基准费率和公式展示
        self._update_formula_display()

        # 伤残赔偿比例区域：仅 liability 显示
        if hasattr(self, 'disability_card'):
            self.disability_card.setVisible(is_liability)

    def _update_formula_display(self):
        """更新基准费率和公式展示卡片"""
        if not hasattr(self, 'formula_label'):
            return
        pt = self._get_product_type()
        product = MC_PRODUCTS.get(self.selected_product, {})
        if not product.get("formulaText"):
            self.formula_label.setVisible(False)
            return
        ver = product.get("versions", {}).get(self.selected_version, {})
        base_rate_text = "基准费率："
        is_liability = pt == "liability"
        is_composite = pt == "composite"
        is_jewelry = pt == "jewelry"
        is_multi_risk = pt == "multiRiskSum"
        if is_liability and ver.get("baseRates"):
            method = self.method_combo.currentData()
            ind_class = self.industry_combo.currentData()
            method_label = "固定限额" if method == "fixed" else "工资总额"
            class_labels = {"class1": "一类", "class2": "二类", "class3": "三类"}
            rates = ver["baseRates"].get(method, {})
            if rates:
                current_rate = rates.get(ind_class, 0)
                base_rate_text += f"{method_label} · {class_labels.get(ind_class, ind_class)} {current_rate * 100:.2f}%"
                all_rates = [f"{class_labels.get(k, k)} {v * 100:.2f}%" for k, v in rates.items()]
                base_rate_text += f"（{method_label}全部：{' / '.join(all_rates)}）"
        elif is_composite and ver.get("baseRates") and product.get("subRisks"):
            sr_parts = []
            for rid in product["subRisks"]:
                lbl = product.get("subRiskLabels", {}).get(rid, rid)
                rate = ver["baseRates"].get(rid, 0)
                if rate:
                    sr_parts.append(f"{lbl} {rate * 100:.3f}%")
            base_rate_text += " | ".join(sr_parts)
        elif is_jewelry and product.get("coverageCategories"):
            parts = []
            for k, cat in product["coverageCategories"].items():
                parts.append(f"{cat['label'].split('-')[-1]} {cat['baseRates']['wholesale'] * 100:.3f}%")
            base_rate_text += " | ".join(parts)
        elif is_multi_risk and ver.get("risks"):
            r_parts = []
            for rk in ver["risks"]:
                if rk.get("baseRate") is not None:
                    r_parts.append(f"{rk['name']} {rk['baseRate'] * 100:.4f}%")
                else:
                    r_parts.append(f"{rk['name']} 查表")
            base_rate_text += " | ".join(r_parts)
        elif ver.get("baseRates", {}).get("default") is not None:
            base_rate_text += f"{ver['baseRates']['default'] * 100:.3f}%"
        formula_text = f"{base_rate_text}\n公式：{product['formulaText']}"
        if product.get("formulaNote"):
            formula_text += f"\n{product['formulaNote']}"
        self.formula_label.setText(formula_text)
        self.formula_label.setVisible(True)

    def _build_mrs_params(self, visible):
        """构建/更新 multiRiskSum 风险参数选择区"""
        if not hasattr(self, 'mrs_container'):
            return
        # 清除旧控件
        while self.mrs_layout.count():
            item = self.mrs_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.mrs_widgets = {}
        self.mrs_container.setVisible(visible)
        if not visible:
            return
        if not self.current_plan:
            return
        risks = self.current_plan.get("risks", [])
        if not risks:
            return
        title = QLabel("风险参数选择")
        title.setStyleSheet(f"font-weight: 600; font-size: 13px; color: {AnthropicColors.TEXT_PRIMARY};")
        self.mrs_layout.addWidget(title)
        industry_added = False
        for risk in risks:
            tbl = risk.get("baseRateTable")
            if not tbl:
                continue
            tbl_type = tbl.get("type", "")
            if tbl_type == "industryBuilding" and not industry_added:
                industry_added = True
                lbl = QLabel("行业代码:")
                self.mrs_layout.addWidget(lbl)
                combo = QComboBox()
                for code in tbl.get("data", {}).keys():
                    label = tbl.get("labels", {}).get(code, code)
                    combo.addItem(label, code)
                self.mrs_layout.addWidget(combo)
                self.mrs_widgets["industry"] = combo
                lbl2 = QLabel("建筑类型:")
                self.mrs_layout.addWidget(lbl2)
                combo2 = QComboBox()
                for bt in tbl.get("buildingTypes", []):
                    combo2.addItem(bt["name"], bt["id"])
                self.mrs_layout.addWidget(combo2)
                self.mrs_widgets["buildingType"] = combo2
            elif tbl_type == "zoneStructure":
                lbl = QLabel(f"{risk['name']} — 地区分区:")
                self.mrs_layout.addWidget(lbl)
                combo = QComboBox()
                for zone in tbl.get("data", {}).keys():
                    z_label = tbl.get("zoneLabels", {}).get(zone, zone)
                    combo.addItem(z_label, zone)
                self.mrs_layout.addWidget(combo)
                self.mrs_widgets[f"zone_{risk['id']}"] = combo
                lbl2 = QLabel(f"{risk['name']} — 建筑结构:")
                self.mrs_layout.addWidget(lbl2)
                combo2 = QComboBox()
                for s in tbl.get("structures", []):
                    combo2.addItem(s["name"], s["id"])
                self.mrs_layout.addWidget(combo2)
                self.mrs_widgets[f"struct_{risk['id']}"] = combo2
            elif tbl_type == "plantType":
                lbl = QLabel(f"{risk['name']} — 工厂类型:")
                self.mrs_layout.addWidget(lbl)
                combo = QComboBox()
                for k in tbl.get("data", {}).keys():
                    combo.addItem(k, k)
                self.mrs_layout.addWidget(combo)
                self.mrs_widgets["plantType"] = combo
            elif tbl_type == "manual":
                lbl = QLabel(f"{risk['name']} 基准费率（万分之）:")
                self.mrs_layout.addWidget(lbl)
                spin = QDoubleSpinBox()
                spin.setRange(0, 9999)
                spin.setDecimals(1)
                default_val = (risk.get("defaultRate", 0) * 10000)
                spin.setValue(default_val)
                spin.setSuffix(" ‱")
                self.mrs_layout.addWidget(spin)
                self.mrs_widgets[f"rate_{risk['id']}"] = spin
                if risk.get("rateRange"):
                    hint = QLabel(f"范围：{risk['rateRange'][0] * 10000:.1f}‱ ~ {risk['rateRange'][1] * 10000:.1f}‱")
                    hint.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_SECONDARY};")
                    self.mrs_layout.addWidget(hint)

    def _build_disability_section(self):
        self.disability_card = card = GlassCard()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(16, 12, 16, 12)

        # Tab 按钮行
        title_row = QHBoxLayout()
        title_row.addWidget(QLabel("伤残赔偿比例附表:"))
        title_row.addStretch()
        layout.addLayout(title_row)

        tab_row = QHBoxLayout()
        tab_row.setSpacing(4)
        self.disability_tab_btns = {}
        self.disability_combo = QComboBox()
        self.disability_combo.addItem("不使用", "none")
        for tid, tdata in MC_DISABILITY_TABLES.items():
            self.disability_combo.addItem(tdata["label"], tid)
        self.disability_combo.hide()

        for key, label in [("none", "不启用"), ("table1", "附表1"), ("table2", "附表2"), ("table3", "附表3")]:
            btn = QPushButton(label)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setFixedHeight(32)
            btn.clicked.connect(lambda checked, k=key: self._on_disability_tab_click(k))
            tab_row.addWidget(btn)
            self.disability_tab_btns[key] = btn
        tab_row.addStretch()
        layout.addLayout(tab_row)
        self._update_disability_tab_styles("none")

        # 伤残比例展示区
        self.disability_display = QLabel("")
        self.disability_display.setWordWrap(True)
        self.disability_display.setStyleSheet(f"""
            QLabel {{ background: rgba(59, 130, 246, 0.08); border: 1px solid rgba(59, 130, 246, 0.2);
                border-radius: 8px; padding: 10px 14px; font-size: 12px; color: #3b82f6; }}
        """)
        self.disability_display.hide()
        layout.addWidget(self.disability_display)

        # "是否附加" 提示区
        self.disability_prompt = QWidget()
        prompt_layout = QHBoxLayout(self.disability_prompt)
        prompt_layout.setContentsMargins(0, 8, 0, 0)
        prompt_layout.addWidget(QLabel("是否附加伤残赔偿金赔偿限额比例条款？"))
        yes_btn = QPushButton("是")
        yes_btn.setCursor(Qt.PointingHandCursor)
        yes_btn.setStyleSheet(f"""
            QPushButton {{ background: {AnthropicColors.ACCENT}; color: white; border: none;
                border-radius: 6px; padding: 6px 16px; font-size: 13px; }}
            QPushButton:hover {{ background: {AnthropicColors.ACCENT_DARK}; }}
        """)
        yes_btn.clicked.connect(self._show_addon_modal)
        prompt_layout.addWidget(yes_btn)
        no_btn = QPushButton("否")
        no_btn.setCursor(Qt.PointingHandCursor)
        no_btn.setStyleSheet(f"""
            QPushButton {{ background: {AnthropicColors.BG_CARD}; color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER}; border-radius: 6px; padding: 6px 16px; font-size: 13px; }}
            QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; }}
        """)
        no_btn.clicked.connect(self._on_addon_prompt_no)
        prompt_layout.addWidget(no_btn)
        prompt_layout.addStretch()
        self.disability_prompt.hide()
        layout.addWidget(self.disability_prompt)

        # 已选方案显示
        self.disability_selection_label = QLabel("")
        self.disability_selection_label.setWordWrap(True)
        self.disability_selection_label.setStyleSheet(f"""
            QLabel {{ background: rgba(16, 185, 129, 0.08); border: 1px solid rgba(16, 185, 129, 0.2);
                border-radius: 8px; padding: 8px 14px; font-size: 12px; color: #059669; }}
        """)
        self.disability_selection_label.hide()
        layout.addWidget(self.disability_selection_label)

        self.scroll_layout.addWidget(card)

    def _update_disability_tab_styles(self, active_key):
        for key, btn in self.disability_tab_btns.items():
            if key == active_key:
                btn.setStyleSheet(f"""
                    QPushButton {{ background: {AnthropicColors.ACCENT}; color: white; border: none;
                        border-radius: 6px; padding: 6px 16px; font-size: 13px; font-weight: 600; }}
                """)
            else:
                btn.setStyleSheet(f"""
                    QPushButton {{ background: {AnthropicColors.BG_PRIMARY}; color: {AnthropicColors.TEXT_PRIMARY};
                        border: 1px solid {AnthropicColors.BORDER}; border-radius: 6px; padding: 6px 16px; font-size: 13px; }}
                    QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; }}
                """)

    def _on_disability_tab_click(self, table_key):
        self.selected_disability_table = table_key
        self.selected_disability_option = -1
        self._update_disability_tab_styles(table_key)

        if table_key == "none":
            self.disability_display.hide()
            self.disability_prompt.hide()
            self.disability_selection_label.hide()
            self._render_coefficients()
            self._log("已关闭伤残赔偿比例附表")
            return

        tbl = MC_DISABILITY_TABLES[table_key]
        lines = [f"<b>{tbl['label']} · 伤残赔偿比例</b><br>"]
        for r in tbl["ratios"]:
            lines.append(f"{r['level']}: {r['pct']}%")
        self.disability_display.setText("<br>".join(lines))
        self.disability_display.show()
        self.disability_prompt.show()
        self.disability_selection_label.hide()
        self._render_coefficients()
        self._log(f"选择伤残赔偿比例: {tbl['label']}")

    def _on_disability_table_change(self):
        table_key = self.disability_combo.currentData()
        self._on_disability_tab_click(table_key)

    def _on_addon_prompt_no(self):
        self.selected_disability_option = -1
        self.disability_prompt.hide()
        self.disability_selection_label.hide()
        self._render_coefficients()
        self._log("未附加伤残赔偿比例条款")

    def _show_addon_modal(self):
        dialog = DisabilityAddonDialog(self.selected_disability_table, self.selected_disability_option, self)
        if dialog.exec_() == QDialog.Accepted and dialog.selected_option >= 0:
            self._select_disability_option(dialog.selected_option)

    def _select_disability_option(self, idx):
        self.selected_disability_option = idx
        self.disability_prompt.hide()
        opt = MC_DISABILITY_ADDON_OPTIONS[idx]
        coeff_val = opt["coeff"][self.selected_disability_table]
        tbl_label = MC_DISABILITY_TABLES[self.selected_disability_table]["label"]
        self.disability_selection_label.setText(
            f"已选方案: {opt['group']}组 · 九级{opt['p9']}% 十级{opt['p10']}% → 系数 {fmt_num(coeff_val, 3)} ({tbl_label})"
        )
        self.disability_selection_label.show()
        self._render_coefficients()
        self._log(f"选择伤残方案: {opt['label']} → 系数 {fmt_num(coeff_val, 3)}")

    def _build_coeff_section(self):
        self.coeff_container = QWidget()
        self.coeff_layout = QVBoxLayout(self.coeff_container)
        self.coeff_layout.setContentsMargins(0, 0, 0, 0)
        self.coeff_layout.setSpacing(8)
        self.scroll_layout.addWidget(self.coeff_container)
        self._render_coefficients()

    def _get_applicable_coefficients(self):
        """获取当前产品类型下适用的系数列表"""
        if not self.current_plan:
            return []
        pt = self._get_product_type()
        coefficients = self.current_plan.get("coefficients", [])
        if pt == "liability":
            method = self.method_combo.currentData()
            return [c for c in coefficients if method in c.get("applicableTo", [])]
        # 对于非 liability 类型，显示所有系数（它们的 applicableTo 可能包含 "all" 或子险种名称）
        return coefficients

    def _render_coefficients(self):
        while self.coeff_layout.count():
            item = self.coeff_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        if not self.current_plan:
            return
        applicable = self._get_applicable_coefficients()
        if not applicable and self.selected_disability_table == "none":
            lbl = QLabel("当前计费方式无可用系数表")
            lbl.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; padding: 16px;")
            self.coeff_layout.addWidget(lbl)
            return
        for coeff in applicable:
            card = self._create_coeff_card(coeff)
            self.coeff_layout.addWidget(card)

    def _create_coeff_card(self, coeff):
        card = GlassCard()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(6)
        sel = self.coeff_selections.get(coeff["id"])
        sel_value = sel["value"] if sel else None
        title_text = coeff["name"]
        title_text += f"  [{fmt_num(sel_value, 4)}]" if sel_value is not None else "  [未选]"
        title = QLabel(title_text)
        title.setStyleSheet(f"font-weight: 600; font-size: 13px; color: {AnthropicColors.TEXT_PRIMARY};")
        layout.addWidget(title)
        if coeff.get("note"):
            note = QLabel(f"注: {coeff['note']}")
            note.setWordWrap(True)
            note.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
            layout.addWidget(note)
        for ri, row in enumerate(coeff["rows"]):
            is_selected = sel and sel.get("rowIndex") == ri
            if row["type"] == "range":
                bl = "(" if row.get("minExclusive") else "["
                br = ")" if row.get("maxExclusive") else "]"
                val_text = f"{bl}{fmt_num(row['min'], 2)}, {fmt_num(row['max'], 2)}{br}"
            else:
                val_text = fmt_num(row["value"], 2)
            btn = QPushButton(f"{row['parameter']}    {val_text}")
            btn.setCursor(Qt.PointingHandCursor)
            bg = AnthropicColors.ACCENT if is_selected else AnthropicColors.BG_PRIMARY
            fg = AnthropicColors.TEXT_LIGHT if is_selected else AnthropicColors.TEXT_PRIMARY
            btn.setStyleSheet(f"""
                QPushButton {{ background: {bg}; color: {fg}; border: 1px solid {AnthropicColors.BORDER};
                    border-radius: 6px; padding: 6px 12px; font-size: 12px; text-align: left; }}
                QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; }}
            """)
            click_value = row["min"] if row["type"] == "range" else row["value"]
            btn.clicked.connect(lambda checked, cid=coeff["id"], ridx=ri, cv=click_value: self._select_coeff_row(cid, ridx, cv))
            layout.addWidget(btn)
        if sel and sel.get("rowIndex") is not None:
            row = coeff["rows"][sel["rowIndex"]]
            if row["type"] == "range":
                slider_layout = QHBoxLayout()
                slider_layout.addWidget(QLabel("精确系数: "))
                slider_label = QLabel(fmt_num(sel["value"], 4))
                slider_label.setStyleSheet(f"font-weight: 600; color: {AnthropicColors.ACCENT};")
                slider_layout.addWidget(slider_label)
                slider = QSlider(Qt.Horizontal)
                slider.setMinimum(int(row["min"] * 10000))
                slider.setMaximum(int(row["max"] * 10000))
                slider.setValue(int(sel["value"] * 10000))
                slider.setSingleStep(100)
                cid = coeff["id"]
                slider.valueChanged.connect(lambda v, c=cid, lr=slider_label: self._on_slider_change(c, v, lr))
                slider_layout.addWidget(slider, 1)
                range_info = QLabel(f"[{fmt_num(row['min'], 2)}, {fmt_num(row['max'], 2)}]")
                range_info.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
                slider_layout.addWidget(range_info)
                layout.addLayout(slider_layout)
        return card

    def _select_coeff_row(self, coeff_id, row_index, value):
        self.coeff_selections[coeff_id] = {"rowIndex": row_index, "value": value}
        self._render_coefficients()

    def _on_slider_change(self, coeff_id, int_value, label_widget):
        value = int_value / 10000.0
        if coeff_id in self.coeff_selections:
            self.coeff_selections[coeff_id]["value"] = value
        label_widget.setText(fmt_num(value, 4))

    def _build_action_buttons(self):
        row = QHBoxLayout()
        calc_btn = make_accent_button("🧮 计算主险保费")
        calc_btn.clicked.connect(self.calculate)
        row.addWidget(calc_btn)
        reset_btn = QPushButton("🔄 重置参数")
        reset_btn.setCursor(Qt.PointingHandCursor)
        reset_btn.clicked.connect(self._reset)
        row.addWidget(reset_btn)
        self.send_btn = make_success_button("📤 传入附加险计算")
        self.send_btn.clicked.connect(self._send_to_addon)
        self.send_btn.hide()
        row.addWidget(self.send_btn)
        row.addStretch()
        self.scroll_layout.addLayout(row)

    def _build_result_section(self):
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setMaximumHeight(300)
        self.result_display.hide()
        self.scroll_layout.addWidget(self.result_display)

    def _clear_result(self):
        self.result_display.clear()
        self.result_display.hide()
        self.send_btn.hide()

    def _build_log_section(self):
        self.log_display = QTextEdit()
        self.log_display.setReadOnly(True)
        self.log_display.setMaximumHeight(150)
        self.log_display.setStyleSheet(f"""
            QTextEdit {{ background: {AnthropicColors.BG_DARK}; color: {AnthropicColors.TEXT_LIGHT};
                border-radius: 8px; padding: 8px; font-size: 11px; font-family: monospace; }}
        """)
        self.scroll_layout.addWidget(self.log_display)

    def _log(self, msg, level="info"):
        from datetime import datetime
        time_str = datetime.now().strftime("%H:%M:%S")
        prefix = {"error": "❌", "warn": "⚠️", "success": "✅"}.get(level, "ℹ️")
        self.log_display.append(f"[{time_str}] {prefix} {msg}")

    def _calc_coeff_product(self, applicable):
        """计算系数乘积，返回 (coeff_product, coeff_details, unselected_count)
        支持 linkedGroup：同组系数取最低值而非相乘"""
        coeff_product = 1.0
        coeff_details = []
        unselected_count = 0
        linked_groups = {}
        standalone = []
        for coeff in applicable:
            if coeff.get("linkedGroup"):
                linked_groups.setdefault(coeff["linkedGroup"], []).append(coeff)
            else:
                standalone.append(coeff)
        # 处理独立系数
        for coeff in standalone:
            sel = self.coeff_selections.get(coeff["id"])
            if sel:
                coeff_product *= sel["value"]
                coeff_details.append({"name": coeff["name"], "value": sel["value"]})
                self._log(f"  系数 [{coeff['name']}] = {fmt_num(sel['value'], 4)}")
            else:
                coeff_details.append({"name": coeff["name"], "value": 1.0, "unselected": True})
                unselected_count += 1
        # 处理联动组：取组内最低值
        for group_key, group_coeffs in linked_groups.items():
            selected_vals = []
            for coeff in group_coeffs:
                sel = self.coeff_selections.get(coeff["id"])
                if sel:
                    selected_vals.append(sel["value"])
                    coeff_details.append({"name": coeff["name"], "value": sel["value"], "linkedGroup": group_key})
                    self._log(f"  系数 [{coeff['name']}] = {fmt_num(sel['value'], 4)} (联动组: {group_key})")
                else:
                    coeff_details.append({"name": coeff["name"], "value": 1.0, "unselected": True, "linkedGroup": group_key})
                    unselected_count += 1
            if selected_vals:
                min_val = min(selected_vals)
                coeff_product *= min_val
                self._log(f"  联动组 [{group_key}] 取最低值: {fmt_num(min_val, 4)}")
        if unselected_count > 0:
            self._log(f"  注意: {unselected_count} 个系数未选择，按基准 1.0 计算", "warn")
        self._log(f"系数乘积: {fmt_num(coeff_product, 6)}")
        return coeff_product, coeff_details, unselected_count

    def _apply_premium_cap(self, adjusted_rate, premium_cap):
        """应用费率封顶"""
        is_capped = False
        if premium_cap is not None and adjusted_rate > premium_cap:
            self._log(f"调整后费率 {adjusted_rate * 100:.4f}% 超过{premium_cap * 100:.0f}%封顶", "warn")
            adjusted_rate = premium_cap
            is_capped = True
        return adjusted_rate, is_capped

    def _calc_liability(self):
        """雇主责任险计算逻辑"""
        method = self.method_combo.currentData()
        industry_class = self.industry_combo.currentData()
        employee_count = self.count_spin.value()
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        if employee_count <= 0:
            self._log("计算失败: 承保人数无效", "error")
            return
        is_dual_rate = self.current_plan.get("dualRate") and method == "fixed"
        product_data = MC_PRODUCTS.get(self.selected_product, {})
        premium_cap = self.current_plan.get("premiumCap", product_data.get("premiumCap"))
        applicable = [c for c in self.current_plan.get("coefficients", []) if method in c.get("applicableTo", [])]
        coeff_product, coeff_details, _ = self._calc_coeff_product(applicable)
        per_person_premium = 0.0
        total_premium = 0.0
        formula = ""
        base_rate = 0.0
        adjusted_rate = 0.0
        is_capped = False
        if is_dual_rate:
            # E款双费率模式
            death_limit = getattr(self, 'death_limit_spin', None)
            medical_limit_val = getattr(self, 'medical_limit_spin', None)
            death_limit_wan = death_limit.value() if death_limit else 50
            medical_limit_wan = medical_limit_val.value() if medical_limit_val else 5
            death_limit_yuan = death_limit_wan * 10000
            medical_limit_yuan = medical_limit_wan * 10000
            fixed_rates = self.current_plan.get("baseRates", {}).get("fixed", {})
            death_rate = fixed_rates.get("death", {}).get(industry_class, 0)
            medical_rate = fixed_rates.get("medical", {}).get(industry_class, 0)
            self._log(f"版本: {self.current_plan.get('label', '')} | 计费: 固定限额（双费率） | 行业: {industry_class}")
            self._log(f"死亡残疾基准费率: {death_rate * 100:.4f}% | 医疗基准费率: {medical_rate * 100:.4f}%")
            base_premium = death_limit_yuan * death_rate + medical_limit_yuan * medical_rate
            per_person_premium = base_premium * coeff_product
            if term_type == "short":
                per_person_premium *= (days / 365)
            total_premium = per_person_premium * employee_count
            base_rate = base_premium / (death_limit_yuan + medical_limit_yuan) if (death_limit_yuan + medical_limit_yuan) > 0 else 0
            adjusted_rate = base_rate * coeff_product
            formula = f"每人基础保费 = {fmt_currency(death_limit_yuan)} × {death_rate * 100:.4f}% + {fmt_currency(medical_limit_yuan)} × {medical_rate * 100:.4f}% = {fmt_currency(base_premium)}"
            formula += f"\n每人保费 = {fmt_currency(base_premium)} × {fmt_num(coeff_product, 6)}"
            if term_type == "short":
                formula += f" × ({days}/365)"
            formula += f" = {fmt_currency(per_person_premium)}"
            formula += f"\n主险保费 = {fmt_currency(per_person_premium)} × {employee_count}人 = {fmt_currency(total_premium)}"
        else:
            # 标准模式
            base_rates = self.current_plan.get("baseRates", {}).get(method, {})
            base_rate = base_rates.get(industry_class)
            if not base_rate:
                self._log(f"计算失败: 基准费率不存在 method={method} class={industry_class}", "error")
                return
            self._log(f"版本: {self.current_plan.get('label', '')} | 计费: {'固定限额' if method == 'fixed' else '工资总额'} | 行业: {industry_class}")
            self._log(f"基准费率: {base_rate * 100:.4f}%")
            adjusted_rate = base_rate * coeff_product
            adjusted_rate, is_capped = self._apply_premium_cap(adjusted_rate, premium_cap)
            self._log(f"调整后费率: {adjusted_rate * 100:.4f}%{'（封顶）' if is_capped else ''}")
            if method == "fixed":
                limit_yuan = self.limit_spin.value() * 10000
                per_person_premium = limit_yuan * adjusted_rate
                if term_type == "short":
                    per_person_premium *= (days / 365)
                total_premium = per_person_premium * employee_count
                formula = f"每人保费 = {fmt_currency(limit_yuan)} × {adjusted_rate * 100:.4f}%"
                if term_type == "short":
                    formula += f" × ({days}/365)"
                formula += f" = {fmt_currency(per_person_premium)}"
                formula += f"\n主险保费 = {fmt_currency(per_person_premium)} × {employee_count}人 = {fmt_currency(total_premium)}"
            else:
                salary_yuan = self.salary_spin.value()
                total_premium = salary_yuan * adjusted_rate
                if term_type == "short":
                    total_premium *= (days / 365)
                per_person_premium = total_premium / employee_count if employee_count > 0 else 0
                formula = f"年保费 = {fmt_currency(salary_yuan)} × {adjusted_rate * 100:.4f}%"
                if term_type == "short":
                    formula += f" × ({days}/365)"
                formula += f" = {fmt_currency(total_premium)}"
                formula += f"\n每人均摊: {fmt_currency(total_premium)} / {employee_count}人 = {fmt_currency(per_person_premium)}"
        disability_coeff = 1.0
        disability_desc = ""
        if self.selected_disability_table != "none" and self.selected_disability_option >= 0:
            d_opt = MC_DISABILITY_ADDON_OPTIONS[self.selected_disability_option]
            disability_coeff = d_opt["coeff"][self.selected_disability_table]
            before_premium = total_premium
            total_premium *= disability_coeff
            per_person_premium *= disability_coeff
            tbl_label = MC_DISABILITY_TABLES[self.selected_disability_table]["label"]
            disability_desc = f"附加伤残赔偿比例({tbl_label} · {d_opt['label']})"
            formula += f"\n\n扩展伤残赔偿比例: {fmt_currency(before_premium)} × {fmt_num(disability_coeff, 3)} = {fmt_currency(total_premium)}"
            self._log(f"伤残赔偿比例调整: × {fmt_num(disability_coeff, 3)} ({disability_desc})")
        self.result = {
            "version": self.current_plan.get("label", ""), "method": method, "industryClass": industry_class,
            "baseRate": base_rate, "coeffProduct": coeff_product, "disabilityCoeff": disability_coeff,
            "disabilityDesc": disability_desc, "adjustedRate": adjusted_rate, "isCapped": is_capped,
            "perPersonPremium": per_person_premium, "totalPremium": total_premium,
            "employeeCount": employee_count, "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": coeff_details,
            "productType": "liability"
        }

    def _calc_property(self):
        """财产险类计算: 年保费 = 保险金额 × 基准费率 × 系数乘积"""
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        insured_amount = self.amount_spin.value()
        base_rate = self.current_plan.get("baseRates", {}).get("default")
        if not base_rate:
            self._log("计算失败: 基准费率不存在", "error")
            return
        self._log(f"版本: {self.current_plan.get('label', '')} | 保险金额: {fmt_currency(insured_amount)}")
        self._log(f"基准费率: {base_rate * 100:.4f}%")
        applicable = self._get_applicable_coefficients()
        coeff_product, coeff_details, _ = self._calc_coeff_product(applicable)
        adjusted_rate = base_rate * coeff_product
        product_data = MC_PRODUCTS.get(self.selected_product, {})
        premium_cap = product_data.get("premiumCap")
        adjusted_rate, is_capped = self._apply_premium_cap(adjusted_rate, premium_cap)
        self._log(f"调整后费率: {adjusted_rate * 100:.4f}%{'（封顶）' if is_capped else ''}")
        total_premium = insured_amount * adjusted_rate
        if term_type == "short":
            total_premium *= (days / 365)
        formula = f"年保费 = {fmt_currency(insured_amount)} × {adjusted_rate * 100:.4f}%"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(total_premium)}"
        self.result = {
            "version": self.current_plan.get("label", ""), "baseRate": base_rate,
            "coeffProduct": coeff_product, "adjustedRate": adjusted_rate, "isCapped": is_capped,
            "totalPremium": total_premium, "perPersonPremium": 0,
            "insuredAmount": insured_amount, "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": coeff_details,
            "productType": "property"
        }

    def _calc_composite(self):
        """恒力项目组合险计算: 分别计算物质损失和机器损坏保费后求和"""
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        material_amount = self.amount_spin.value()
        machinery_amount = self.sub_amount_spin.value()
        base_rates = self.current_plan.get("baseRates", {})
        material_rate = base_rates.get("materialDamage")
        machinery_rate = base_rates.get("machineryBreakdown")
        if not material_rate or not machinery_rate:
            self._log("计算失败: 基准费率不存在", "error")
            return
        self._log(f"版本: {self.current_plan.get('label', '')}")
        self._log(f"物质损失保额: {fmt_currency(material_amount)} | 基准费率: {material_rate * 100:.4f}%")
        self._log(f"机器损坏保额: {fmt_currency(machinery_amount)} | 基准费率: {machinery_rate * 100:.4f}%")
        coefficients = self.current_plan.get("coefficients", [])
        # 物质损失系数
        material_coeffs = [c for c in coefficients if "materialDamage" in c.get("applicableTo", []) or "all" in c.get("applicableTo", [])]
        machinery_coeffs = [c for c in coefficients if "machineryBreakdown" in c.get("applicableTo", []) or "all" in c.get("applicableTo", [])]
        self._log("--- 物质损失系数 ---")
        material_coeff_product, material_coeff_details, _ = self._calc_coeff_product(material_coeffs)
        self._log("--- 机器损坏系数 ---")
        machinery_coeff_product, machinery_coeff_details, _ = self._calc_coeff_product(machinery_coeffs)
        material_adj_rate = material_rate * material_coeff_product
        machinery_adj_rate = machinery_rate * machinery_coeff_product
        material_premium = material_amount * material_adj_rate
        machinery_premium = machinery_amount * machinery_adj_rate
        if term_type == "short":
            material_premium *= (days / 365)
            machinery_premium *= (days / 365)
        total_premium = material_premium + machinery_premium
        formula = f"物质损失保费 = {fmt_currency(material_amount)} × {material_adj_rate * 100:.4f}%"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(material_premium)}"
        formula += f"\n机器损坏保费 = {fmt_currency(machinery_amount)} × {machinery_adj_rate * 100:.4f}%"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(machinery_premium)}"
        formula += f"\n合计保费 = {fmt_currency(material_premium)} + {fmt_currency(machinery_premium)} = {fmt_currency(total_premium)}"
        all_coeff_details = material_coeff_details + machinery_coeff_details
        self.result = {
            "version": self.current_plan.get("label", ""),
            "baseRate": {"materialDamage": material_rate, "machineryBreakdown": machinery_rate},
            "coeffProduct": {"materialDamage": material_coeff_product, "machineryBreakdown": machinery_coeff_product},
            "adjustedRate": {"materialDamage": material_adj_rate, "machineryBreakdown": machinery_adj_rate},
            "isCapped": False, "totalPremium": total_premium, "perPersonPremium": 0,
            "materialAmount": material_amount, "machineryAmount": machinery_amount,
            "materialPremium": material_premium, "machineryPremium": machinery_premium,
            "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": all_coeff_details,
            "productType": "composite"
        }

    def _calc_interruption(self):
        """营业中断保险计算: 年保费 = 毛利润损失保额 × 基准费率 × 系数乘积"""
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        insured_amount = self.amount_spin.value()
        base_rate = self.current_plan.get("baseRates", {}).get("default")
        if not base_rate:
            self._log("计算失败: 基准费率不存在", "error")
            return
        self._log(f"版本: {self.current_plan.get('label', '')} | 毛利润损失保额: {fmt_currency(insured_amount)}")
        self._log(f"基准费率: {base_rate * 100:.4f}%")
        applicable = self._get_applicable_coefficients()
        coeff_product, coeff_details, _ = self._calc_coeff_product(applicable)
        adjusted_rate = base_rate * coeff_product
        self._log(f"调整后费率: {adjusted_rate * 100:.4f}%")
        total_premium = insured_amount * adjusted_rate
        if term_type == "short":
            total_premium *= (days / 365)
        formula = f"年保费 = {fmt_currency(insured_amount)} × {adjusted_rate * 100:.4f}%"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(total_premium)}"
        self.result = {
            "version": self.current_plan.get("label", ""), "baseRate": base_rate,
            "coeffProduct": coeff_product, "adjustedRate": adjusted_rate, "isCapped": False,
            "totalPremium": total_premium, "perPersonPremium": 0,
            "insuredAmount": insured_amount, "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": coeff_details,
            "productType": "interruption"
        }

    def _calc_jewelry(self):
        """珠宝商综合保险计算: 按选择的保障类型和商户类型计算"""
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        insured_amount = self.amount_spin.value()
        merchant_type = self.merchant_type_combo.currentData()
        coverage_type = self.coverage_type_combo.currentData()
        product_data = MC_PRODUCTS.get(self.selected_product, {})
        categories = product_data.get("coverageCategories", {})
        category = categories.get(coverage_type, {})
        base_rate = category.get("baseRates", {}).get(merchant_type)
        if not base_rate:
            self._log(f"计算失败: 基准费率不存在 merchant={merchant_type} coverage={coverage_type}", "error")
            return
        merchant_label = product_data.get("merchantTypes", {}).get(merchant_type, merchant_type)
        coverage_label = category.get("label", coverage_type)
        self._log(f"版本: {self.current_plan.get('label', '')} | {merchant_label} | {coverage_label}")
        self._log(f"保险金额: {fmt_currency(insured_amount)} | 基准费率: {base_rate * 100:.4f}%")
        applicable = self._get_applicable_coefficients()
        coeff_product, coeff_details, _ = self._calc_coeff_product(applicable)
        adjusted_rate = base_rate * coeff_product
        self._log(f"调整后费率: {adjusted_rate * 100:.4f}%")
        total_premium = insured_amount * adjusted_rate
        if term_type == "short":
            total_premium *= (days / 365)
        formula = f"保障类型: {coverage_label} ({merchant_label})"
        formula += f"\n年保费 = {fmt_currency(insured_amount)} × {adjusted_rate * 100:.4f}%"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(total_premium)}"
        self.result = {
            "version": self.current_plan.get("label", ""), "baseRate": base_rate,
            "coeffProduct": coeff_product, "adjustedRate": adjusted_rate, "isCapped": False,
            "totalPremium": total_premium, "perPersonPremium": 0,
            "insuredAmount": insured_amount, "merchantType": merchant_type, "coverageType": coverage_type,
            "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": coeff_details,
            "productType": "jewelry"
        }

    def _calc_multi_risk_sum(self):
        """多风险求和计算: 年保费 = 保险金额 × Σ(风险基准费率 × 调整系数乘积) + 派生风险保费"""
        term_type = self.term_combo.currentData()
        days = self.days_spin.value() if term_type == "short" else 365
        insured_amount = self.amount_spin.value()
        if insured_amount <= 0:
            self._log("计算失败: 保险金额必须大于0", "error")
            return
        risks = self.current_plan.get("risks", [])
        if not risks:
            self._log("计算失败: 未定义风险类别", "error")
            return
        self._log(f"版本: {self.current_plan.get('label', '')} | 保险金额: {fmt_currency(insured_amount)}")
        all_applicable = self._get_applicable_coefficients()
        total_rate = 0.0
        risk_details = []
        all_coeff_details = []
        risk_premiums = {}
        normal_risks = [r for r in risks if not r.get("derivedFrom")]
        derived_risks = [r for r in risks if r.get("derivedFrom")]
        for risk in normal_risks:
            base_rate = 0.0
            if risk.get("baseRate") is not None:
                base_rate = risk["baseRate"]
            elif risk.get("baseRateTable"):
                base_rate = self._lookup_multi_risk_base_rate(risk)
            if not base_rate:
                self._log(f"风险 {risk['name']}: 基准费率未选定或为0，跳过", "warn")
            risk_id = risk.get("id", "")
            risk_coeffs = [c for c in all_applicable
                           if not c.get("applicableTo")
                           or "all" in c.get("applicableTo", [])
                           or risk_id in c.get("applicableTo", [])]
            coeff_product, coeff_details, _ = self._calc_coeff_product(risk_coeffs)
            risk_rate = base_rate * coeff_product
            total_rate += risk_rate
            risk_premium = insured_amount * risk_rate
            risk_premiums[risk_id] = risk_premium
            risk_details.append({
                "id": risk_id, "name": risk["name"],
                "baseRate": base_rate, "coeffProduct": coeff_product,
                "riskRate": risk_rate, "coeffDetails": coeff_details
            })
            for d in coeff_details:
                if not any(e.get("name") == d.get("name") and e.get("value") == d.get("value") for e in all_coeff_details):
                    all_coeff_details.append(d)
            self._log(f"风险 {risk['name']}: 基准费率={base_rate * 100:.4f}% × 系数乘积={coeff_product:.4f} → {risk_rate * 100:.4f}%")
        derived_premium_total = 0.0
        for risk in derived_risks:
            risk_id = risk.get("id", "")
            source_ids = risk.get("derivedFrom", [])
            derived_rate = risk.get("derivedRate", 0.0)
            source_sum = sum(risk_premiums.get(sid, 0.0) for sid in source_ids)
            risk_coeffs = [c for c in all_applicable
                           if not c.get("applicableTo")
                           or "all" in c.get("applicableTo", [])
                           or risk_id in c.get("applicableTo", [])]
            coeff_product, coeff_details, _ = self._calc_coeff_product(risk_coeffs)
            derived_prem = source_sum * derived_rate * coeff_product
            derived_premium_total += derived_prem
            risk_premiums[risk_id] = derived_prem
            source_names = [r["name"] for r in normal_risks if r.get("id") in source_ids]
            risk_details.append({
                "id": risk_id, "name": risk["name"],
                "derivedFrom": source_ids, "derivedRate": derived_rate,
                "sourcePremium": source_sum, "coeffProduct": coeff_product,
                "derivedPremium": derived_prem, "coeffDetails": coeff_details
            })
            for d in coeff_details:
                if not any(e.get("name") == d.get("name") and e.get("value") == d.get("value") for e in all_coeff_details):
                    all_coeff_details.append(d)
            self._log(f"派生风险 {risk['name']}: ({'+'.join(source_names)}保费) × {derived_rate*100:.1f}% × 系数={coeff_product:.4f} → {fmt_currency(derived_prem)}")
        product_data = MC_PRODUCTS.get(self.selected_product, {})
        cap_note = ""
        if product_data.get("premiumCap"):
            for rd in risk_details:
                if rd.get("coeffProduct", 1) < product_data["premiumCap"]:
                    cap_note = f"注：各风险调整系数乘积不应低于 {product_data['premiumCap']}"
                    break
        total_premium = insured_amount * total_rate + derived_premium_total
        if term_type == "short":
            total_premium *= (days / 365)
        formula = "年保费 = 保险金额 × Σ(风险基准费率 × 调整系数乘积)"
        if derived_risks:
            formula += " + 派生风险保费"
        formula += "\n"
        for rd in risk_details:
            if rd.get("derivedFrom"):
                formula += f"{rd['name']}: ({'+'.join(rd['derivedFrom'])}保费 {fmt_currency(rd['sourcePremium'])}) × {rd['derivedRate']*100:.1f}% × {rd['coeffProduct']:.4f} = {fmt_currency(rd['derivedPremium'])}\n"
            else:
                formula += f"{rd['name']}: {rd['baseRate'] * 100:.4f}% × {rd['coeffProduct']:.4f} = {rd['riskRate'] * 100:.4f}%\n"
        if not derived_risks:
            formula += f"总费率: {total_rate * 100:.4f}%\n"
            formula += f"年保费 = {fmt_currency(insured_amount)} × {total_rate * 100:.4f}%"
        else:
            formula += f"基础风险总费率: {total_rate * 100:.4f}% → 基础保费 {fmt_currency(insured_amount * total_rate)}\n"
            formula += f"派生风险保费合计: {fmt_currency(derived_premium_total)}\n"
            formula += f"年保费 = {fmt_currency(insured_amount * total_rate)} + {fmt_currency(derived_premium_total)}"
        if term_type == "short":
            formula += f" × ({days}/365)"
        formula += f" = {fmt_currency(total_premium)}"
        if cap_note:
            formula += f"\n{cap_note}"
        self._log(f"总保费: {fmt_currency(total_premium)}", "success")
        self.result = {
            "version": self.current_plan.get("label", ""), "totalRate": total_rate,
            "totalPremium": total_premium, "insuredAmount": insured_amount,
            "riskDetails": risk_details, "termType": term_type, "days": days,
            "formulaBreakdown": formula, "coeffDetails": all_coeff_details,
            "productType": "multiRiskSum", "derivedPremium": derived_premium_total
        }

    def _lookup_multi_risk_base_rate(self, risk):
        """多风险基准费率查表"""
        table = risk.get("baseRateTable")
        if not table:
            return 0.0
        tbl_type = table.get("type", "")
        if tbl_type == "industryBuilding":
            industry_combo = self.mrs_widgets.get("industry")
            building_combo = self.mrs_widgets.get("buildingType")
            if not industry_combo or not building_combo:
                return 0.0
            code = industry_combo.currentData()
            b_type = building_combo.currentData()
            rates = table.get("data", {}).get(code, {})
            return rates.get(b_type, 0.0)
        if tbl_type == "zoneStructure":
            zone_combo = self.mrs_widgets.get(f"zone_{risk['id']}")
            struct_combo = self.mrs_widgets.get(f"struct_{risk['id']}")
            if not zone_combo or not struct_combo:
                return 0.0
            zone = zone_combo.currentData()
            struct = struct_combo.currentData()
            zone_data = table.get("data", {}).get(zone, {})
            return zone_data.get(struct, 0.0)
        if tbl_type == "plantType":
            pt_combo = self.mrs_widgets.get("plantType")
            if not pt_combo:
                return 0.0
            return table.get("data", {}).get(pt_combo.currentData(), 0.0)
        if tbl_type == "manual":
            rate_spin = self.mrs_widgets.get(f"rate_{risk['id']}")
            if not rate_spin:
                return 0.0
            return rate_spin.value() / 10000.0
        return 0.0

    def calculate(self):
        self._log("--- 开始计算 ---")
        pt = self._get_product_type()
        self.result = None
        if pt == "liability":
            self._calc_liability()
        elif pt == "property":
            self._calc_property()
        elif pt == "composite":
            self._calc_composite()
        elif pt == "interruption":
            self._calc_interruption()
        elif pt == "jewelry":
            self._calc_jewelry()
        elif pt == "multiRiskSum":
            self._calc_multi_risk_sum()
        else:
            self._log(f"未知产品类型: {pt}", "error")
            return
        if self.result:
            self._log(f"主险保费: {fmt_currency(self.result['totalPremium'])}", "success")
            self._log("--- 计算完成 ---", "success")
            self._render_result()
            self.send_btn.show()

    def _render_result(self):
        if not self.result:
            return
        r = self.result
        pt = r.get("productType", "liability")
        product_name = MC_PRODUCTS.get(self.selected_product, {}).get("productName", "")
        lines = [
            f"═══════════════ 📊 {product_name} 计算结果 ═══════════════", "",
            f"  主险保费:   {fmt_currency(r['totalPremium'])}",
        ]
        if pt == "liability":
            lines.append(f"  每人保费:   {fmt_currency(r['perPersonPremium'])}")
        lines.extend(["", "─────────── 公式分解 ───────────", r["formulaBreakdown"], ""])
        lines.append("─────────── 参数明细 ───────────")
        lines.append(f"  费率版本: {r.get('version', '')}")
        if pt == "liability":
            lines.append(f"  计费方式: {'固定限额' if r.get('method') == 'fixed' else '工资总额'}")
            lines.append(f"  行业类别: {r.get('industryClass', '')}")
            base_rate_val = r.get('baseRate', 0)
            lines.append(f"  基准费率: {base_rate_val * 100:.4f}%")
            lines.append(f"  系数乘积: {fmt_num(r.get('coeffProduct', 1), 6)}")
            lines.append(f"  调整后费率: {r.get('adjustedRate', 0) * 100:.4f}%{' (封顶)' if r.get('isCapped') else ''}")
            lines.append(f"  承保人数: {r.get('employeeCount', 0)}人")
        elif pt == "composite":
            br = r.get("baseRate", {})
            cp = r.get("coeffProduct", {})
            ar = r.get("adjustedRate", {})
            lines.append(f"  物质损失基准费率: {br.get('materialDamage', 0) * 100:.4f}%")
            lines.append(f"  机器损坏基准费率: {br.get('machineryBreakdown', 0) * 100:.4f}%")
            lines.append(f"  物质损失系数乘积: {fmt_num(cp.get('materialDamage', 1), 6)}")
            lines.append(f"  机器损坏系数乘积: {fmt_num(cp.get('machineryBreakdown', 1), 6)}")
            lines.append(f"  物质损失保额: {fmt_currency(r.get('materialAmount', 0))}")
            lines.append(f"  机器损坏保额: {fmt_currency(r.get('machineryAmount', 0))}")
        elif pt == "multiRiskSum":
            lines.append(f"  保险金额: {fmt_currency(r.get('insuredAmount', 0))}")
            lines.append(f"  基础风险总费率: {r.get('totalRate', 0) * 100:.4f}%")
            if r.get("derivedPremium"):
                lines.append(f"  派生风险保费合计: {fmt_currency(r['derivedPremium'])}")
            for rd in r.get("riskDetails", []):
                if rd.get("derivedFrom"):
                    lines.append(f"  ⤷ {rd['name']}（派生）:")
                    lines.append(f"    来源: {' + '.join(rd['derivedFrom'])}")
                    lines.append(f"    来源保费: {fmt_currency(rd.get('sourcePremium', 0))} × {rd.get('derivedRate', 0) * 100:.1f}% × 系数{rd.get('coeffProduct', 1):.4f}")
                    lines.append(f"    = {fmt_currency(rd.get('derivedPremium', 0))}")
                else:
                    lines.append(f"  {rd['name']}: 基准{rd.get('baseRate', 0) * 100:.4f}% × 系数{rd.get('coeffProduct', 1):.4f} = {rd.get('riskRate', 0) * 100:.4f}%")
        else:
            base_rate_val = r.get('baseRate', 0)
            lines.append(f"  基准费率: {base_rate_val * 100:.4f}%")
            lines.append(f"  系数乘积: {fmt_num(r.get('coeffProduct', 1), 6)}")
            adj_rate = r.get('adjustedRate', 0)
            lines.append(f"  调整后费率: {adj_rate * 100:.4f}%{' (封顶)' if r.get('isCapped') else ''}")
            if r.get("insuredAmount"):
                lines.append(f"  保险金额: {fmt_currency(r['insuredAmount'])}")
        lines.append(f"  保险期间: {'年度' if r.get('termType') == 'annual' else str(r.get('days', 365)) + '天'}")
        if pt == "liability" and r.get("disabilityCoeff", 1.0) != 1.0:
            lines.append(f"  伤残赔偿比例系数: {fmt_num(r['disabilityCoeff'], 3)}")
            lines.append(f"  {r.get('disabilityDesc', '')}")
        lines.extend(["", "─────────── 系数明细 ───────────"])
        for d in r.get("coeffDetails", []):
            suffix = "（未选，默认）" if d.get("unselected") else ""
            lines.append(f"  {d['name']}: {fmt_num(d['value'], 4)}{suffix}")
        self.result_display.setText("\n".join(lines))
        self.result_display.show()

    def _reset(self):
        self.coeff_selections = {}
        self.result = None
        self.industry_combo.setCurrentIndex(0)
        self.method_combo.setCurrentIndex(0)
        self.limit_spin.setValue(50)
        self.salary_spin.setValue(5000000)
        self.count_spin.setValue(100)
        self.term_combo.setCurrentIndex(0)
        self.days_spin.setValue(180)
        self.amount_spin.setValue(10000000)
        self.sub_amount_spin.setValue(5000000)
        self.daily_cash_spin.setValue(100000)
        self.merchant_type_combo.setCurrentIndex(0)
        self.coverage_type_combo.setCurrentIndex(0)
        self.selected_disability_table = "none"
        self.selected_disability_option = -1
        self._on_disability_tab_click("none")
        self._clear_result()
        self._update_params_visibility()
        self._render_coefficients()
        self._log("已重置参数和系数选择（险种/版本不变）")

    def _send_to_addon(self):
        if not self.result:
            return
        self.premium_calculated.emit(self.result["totalPremium"], self.result.get("perPersonPremium", 0))
        full_data = dict(self.result)
        full_data["selectedProduct"] = self.selected_product
        full_data["productType"] = self._get_product_type()
        if self._get_product_type() == "liability":
            full_data["perPersonLimit"] = self.limit_spin.value() * 10000 if self.method_combo.currentData() == "fixed" else 0
            full_data["annualSalary"] = self.salary_spin.value() if self.method_combo.currentData() == "salary" else 0
            full_data["disabilityTable"] = self.selected_disability_table
            full_data["disabilityOption"] = self.selected_disability_option
        self.full_result_calculated.emit(full_data)
        self._log(f"已将主险保费 {fmt_currency(self.result['totalPremium'])} 传入附加险计算", "success")

    def _import_rate_plan(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "导入费率方案", "", "费率方案文件 (*.json *.docx);;JSON (*.json);;Word (*.docx)")
        if not file_path:
            return
        if file_path.endswith(".json"):
            self._import_json(file_path)
        elif file_path.endswith(".docx"):
            self._import_docx(file_path)

    def _import_json(self, file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self._process_imported_data(data)
        except Exception as e:
            self._log(f"JSON 导入失败: {e}", "error")

    def _process_imported_data(self, data):
        if data.get("productName"):
            product_id = data.get("productId", f"custom_{id(data)}")
            product_name = data["productName"]
            versions = {}
            if isinstance(data.get("versions"), list):
                for idx, v in enumerate(data["versions"]):
                    vid = v.get("versionId", f"v{idx + 1}")
                    if not all(k in v for k in ("label", "baseRates", "coefficients")):
                        raise ValueError(f"版本 {vid} 缺少必要字段")
                    versions[vid] = {"label": v["label"], "baseRates": v["baseRates"], "coefficients": v["coefficients"]}
            else:
                raise ValueError("新格式 JSON 需包含 versions 数组")
        else:
            if not all(k in data for k in ("label", "baseRates", "coefficients")):
                raise ValueError("JSON 缺少必要字段: label, baseRates, coefficients")
            product_id = f"custom_{id(data)}"
            product_name = data["label"]
            versions = {"v1": {"label": data["label"], "baseRates": data["baseRates"], "coefficients": data["coefficients"]}}
        first_version = None
        if product_id in MC_PRODUCTS:
            existing = MC_PRODUCTS[product_id]
            for vid, vdata in versions.items():
                final_vid = vid if vid not in existing["versions"] else f"{vid}_{id(vdata)}"
                existing["versions"][final_vid] = vdata
                if not first_version:
                    first_version = final_vid
            self._log(f"向险种 [{existing['productName']}] 追加了 {len(versions)} 个新版本", "success")
        else:
            MC_PRODUCTS[product_id] = {"productName": product_name, "versions": versions}
            first_version = list(versions.keys())[0]
            self._log(f"导入新险种: {product_name}，包含 {len(versions)} 个版本", "success")
        self.product_combo.blockSignals(True)
        self.product_combo.clear()
        for pid, pdata in MC_PRODUCTS.items():
            self.product_combo.addItem(pdata["productName"], pid)
        idx = self.product_combo.findData(product_id)
        if idx >= 0:
            self.product_combo.setCurrentIndex(idx)
        self.product_combo.blockSignals(False)
        self.selected_product = product_id
        self._refresh_version_combo(first_version)
        self._on_version_change()

    def _import_docx(self, file_path):
        try:
            from docx import Document
        except ImportError:
            self._log("python-docx 未安装，请运行: pip install python-docx", "error")
            return
        try:
            doc = Document(file_path)
            text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            parsed = self._parse_rate_plan_text(text)
            reply = QMessageBox.question(
                self, "Docx 导入确认",
                f"险种名称: {parsed['productName']}\n基准费率: {len(parsed['baseRates'].get('fixed', {}))} 个固定 + {len(parsed['baseRates'].get('salary', {}))} 个工资\n系数表: {len(parsed['coefficients'])} 个\n\n确认导入?",
                QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                import_data = {"productName": parsed["productName"], "productId": f"docx_{id(parsed)}",
                               "versions": [{"versionId": "v1", "label": parsed["label"], "baseRates": parsed["baseRates"], "coefficients": parsed["coefficients"]}]}
                self._process_imported_data(import_data)
        except Exception as e:
            self._log(f"Docx 解析失败: {e}", "error")

    def _parse_rate_plan_text(self, text):
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        product_name = "未知险种"
        for line in lines[:5]:
            if "费率" in line or "保险" in line:
                product_name = re.sub(r"费率方案|费率表|附件[:：]?\s*", "", line).strip()[:20]
                break
        base_rates = {"fixed": {}, "salary": {}}
        class_map = {"一": "class1", "二": "class2", "三": "class3", "1": "class1", "2": "class2", "3": "class3"}
        full_text = "\n".join(lines)
        rate_pattern = re.compile(r"[第]?([一二三1-3])[类].*?(\d+\.?\d*)\s*[%‰％]")
        fixed_section = re.search(r"固定[赔偿]*限额[\s\S]*?(?=工资|$)", full_text, re.IGNORECASE)
        if fixed_section:
            for m in rate_pattern.finditer(fixed_section.group()):
                cls = class_map.get(m.group(1))
                if cls:
                    val = float(m.group(2))
                    base_rates["fixed"][cls] = val / 1000 if "‰" in m.group() else val / 100
        salary_section = re.search(r"工资[收入]*[\s\S]*?(?=费率调整|调整系数|$)", full_text, re.IGNORECASE)
        if salary_section:
            for m in rate_pattern.finditer(salary_section.group()):
                cls = class_map.get(m.group(1))
                if cls:
                    val = float(m.group(2))
                    base_rates["salary"][cls] = val / 1000 if "‰" in m.group() else val / 100
        if not base_rates["fixed"] and not base_rates["salary"]:
            raise ValueError("未能从文本中提取到基准费率数据")
        if not base_rates["fixed"]:
            base_rates["fixed"] = dict(base_rates["salary"])
        if not base_rates["salary"]:
            base_rates["salary"] = dict(base_rates["fixed"])
        return {"productName": product_name, "label": f"{product_name}费率", "baseRates": base_rates, "coefficients": []}


# =============================================
# 附加险分类常量
# =============================================

ADDON_TYPES = {
    "modifier_coeff": {"label": "主险系数调整", "color": "#dc2626"},
    "sudden_death": {"label": "突发疾病身故", "color": "#ef4444"},
    "per_person_rate": {"label": "每人费率", "color": "#f59e0b"},
    "per_person_base": {"label": "每人定额", "color": "#06b6d4"},
    "disability_adjust": {"label": "伤残调整", "color": "#10b981"},
    "property_loss": {"label": "财产损失", "color": "#0d9488"},
    "formula_sum": {"label": "求和公式", "color": "#3b82f6"},
    "deduction": {"label": "减收", "color": "#ec4899"},
    "no_calc": {"label": "无需计算", "color": "#6b7280"},
    "simple_percentage": {"label": "百分比", "color": "#d97706"},
    "table_coefficient": {"label": "系数表", "color": "#7c3aed"},
    "regulatory": {"label": "规范类", "color": "#9ca3af"},
    "included_in_main": {"label": "纳入主险", "color": "#64748b"},
    "daily_prorate": {"label": "按日计费", "color": "#ea580c"},
    "main_premium_modifier": {"label": "主险保费调整", "color": "#dc2626"},
    "base_rate_formula": {"label": "基准费率", "color": "#0891b2"},
    "formula_conditional": {"label": "条件公式", "color": "#6366f1"},
    "conditional_simple": {"label": "条件勾选", "color": "#8b5cf6"},
}

# 关键词映射：文件名关键词 → 附加险类型
ADDON_KEYWORD_MAP = [
    (["误工费"], "modifier_coeff"),
    (["突发疾病身故"], "sudden_death"),
    (["特定财产损失"], "property_loss"),
    (["工伤补充", "特定人员"], "per_person_rate"),
    (["药品服务", "药品费用"], "per_person_base"),
    (["劳务关系人员"], "disability_adjust"),
    (["雇主法律责任", "法律费用责任"], "formula_sum"),
    (["一次性伤残"], "formula_sum"),
    (["突发疾病除外", "猝死除外"], "deduction"),
    (["月申报", "员工自动承保", "每月申报"], "no_calc"),
    (["纳入主险保险金额"], "included_in_main"),
    (["按日比例计算", "按日比例收取"], "daily_prorate"),
    (["每次事故赔偿限额"], "main_premium_modifier"),
    (["基准费率"], "base_rate_formula"),
]

# 伤残调整系数 (劳务关系人员)
DISABILITY_ADJUST_COEFFS = {
    "table1": 0.995,
    "table2": 1.072,
    "table3": 0.919,
}


# =============================================
# AddonInsuranceTab — 附加险计算器 (重构版)
# =============================================

class AddonInsuranceTab(QWidget):
    """附加险计算器 Tab — 支持11种附加险分类和专属计算逻辑"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.rate_data = None
        self.filtered_entries = []
        self.selected_entry = None
        self.coeff_selections = {}
        self.custom_input_widgets = {}  # 自定义输入控件映射
        # 主险数据
        self.main_premium = 0.0
        self.main_sum_insured = 0.0
        self.per_person_premium = 0.0
        self.policy_days = 365  # 保单天数，默认365天
        self.full_main_data = None  # 完整主险计算结果
        # 保费汇总
        self.premium_items = []
        self._setup_ui()

    def _setup_ui(self):
        self.setStyleSheet(get_common_styles())
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(8)
        main_layout.setContentsMargins(10, 8, 10, 8)

        # 顶栏: 主险信息 + 导入按钮 (两行 GridLayout)
        top_bar = GlassCard()
        top_grid = QGridLayout(top_bar)
        top_grid.setContentsMargins(16, 10, 16, 10)
        top_grid.setHorizontalSpacing(8)
        top_grid.setVerticalSpacing(6)

        # Row 0: 参数输入
        top_grid.addWidget(QLabel("主险保费:"), 0, 0)
        self.main_premium_input = QDoubleSpinBox()
        self.main_premium_input.setRange(0, 999999999999)
        self.main_premium_input.setDecimals(2)
        self.main_premium_input.setSuffix(" 元")
        self.main_premium_input.setMinimumWidth(100)
        self.main_premium_input.valueChanged.connect(lambda v: setattr(self, 'main_premium', v))
        top_grid.addWidget(self.main_premium_input, 0, 1)

        top_grid.addWidget(QLabel("主险保额:"), 0, 2)
        self.main_sum_insured_input = QDoubleSpinBox()
        self.main_sum_insured_input.setRange(0, 999999999999)
        self.main_sum_insured_input.setDecimals(2)
        self.main_sum_insured_input.setSuffix(" 元")
        self.main_sum_insured_input.setMinimumWidth(100)
        self.main_sum_insured_input.valueChanged.connect(lambda v: setattr(self, 'main_sum_insured', v))
        top_grid.addWidget(self.main_sum_insured_input, 0, 3)

        top_grid.addWidget(QLabel("每人保费:"), 0, 4)
        self.per_person_input = QDoubleSpinBox()
        self.per_person_input.setRange(0, 999999999999)
        self.per_person_input.setDecimals(2)
        self.per_person_input.setSuffix(" 元")
        self.per_person_input.setMinimumWidth(100)
        self.per_person_input.valueChanged.connect(lambda v: setattr(self, 'per_person_premium', v))
        top_grid.addWidget(self.per_person_input, 0, 5)

        top_grid.addWidget(QLabel("保单天数:"), 0, 6)
        self.policy_days_input = QSpinBox()
        self.policy_days_input.setRange(1, 9999)
        self.policy_days_input.setValue(365)
        self.policy_days_input.setSuffix(" 天")
        self.policy_days_input.setMinimumWidth(100)
        self.policy_days_input.valueChanged.connect(lambda v: setattr(self, 'policy_days', v))
        top_grid.addWidget(self.policy_days_input, 0, 7)

        # Row 1: 保险类型 + 状态 + 按钮
        top_grid.addWidget(QLabel("保险类型:"), 1, 0)
        self.batch_insurance_combo = QComboBox()
        self.batch_insurance_combo.addItem("财产基本险")
        self.batch_insurance_combo.addItem("财产综合险")
        self.batch_insurance_combo.addItem("财产一切险")
        self.batch_insurance_combo.setCurrentIndex(2)  # Default to 财产一切险
        top_grid.addWidget(self.batch_insurance_combo, 1, 1)

        # 主险数据状态指示
        self.main_data_status = QLabel("⚪ 未接收主险数据")
        self.main_data_status.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
        top_grid.addWidget(self.main_data_status, 1, 2, 1, 2)

        folder_btn = QPushButton("📁 导入文件夹")
        folder_btn.setCursor(Qt.PointingHandCursor)
        folder_btn.clicked.connect(self._load_folder)
        top_grid.addWidget(folder_btn, 1, 5)

        json_btn = QPushButton("📂 导入JSON")
        json_btn.setCursor(Qt.PointingHandCursor)
        json_btn.clicked.connect(self._load_json)
        top_grid.addWidget(json_btn, 1, 6)

        inquiry_btn = QPushButton("📋 导入询价")
        inquiry_btn.setCursor(Qt.PointingHandCursor)
        inquiry_btn.clicked.connect(self._handle_inquiry_import)
        top_grid.addWidget(inquiry_btn, 1, 7)

        # 奇数列(输入框列)设置弹性拉伸
        for col in (1, 3, 5, 7):
            top_grid.setColumnStretch(col, 1)

        main_layout.addWidget(top_bar)

        # 三列布局 (QSplitter 支持拖拽调整)
        splitter = QSplitter(Qt.Horizontal)
        splitter.setHandleWidth(4)

        # 左列: 搜索 + 条款列表
        left_panel = QWidget()
        left_panel.setMinimumWidth(250)
        left_panel.setMaximumWidth(450)
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("🔍 搜索条款名称...")
        self.search_input.textChanged.connect(self._filter_entries)
        left_layout.addWidget(self.search_input)

        self.load_status = QLabel("未加载费率数据")
        self.load_status.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
        left_layout.addWidget(self.load_status)

        self.clause_list = QListWidget()
        self.clause_list.setStyleSheet(f"""
            QListWidget {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                font-size: 12px;
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item {{
                padding: 8px 12px;
                border-bottom: 1px solid {AnthropicColors.BORDER};
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item:selected {{
                background: {AnthropicColors.ACCENT};
                color: #ffffff;
            }}
            QListWidget::item:hover:!selected {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
        """)
        self.clause_list.currentRowChanged.connect(self._on_clause_selected)
        left_layout.addWidget(self.clause_list, 1)

        self.inquiry_status = QLabel("")
        self.inquiry_status.setWordWrap(True)
        self.inquiry_status.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
        left_layout.addWidget(self.inquiry_status)

        self.batch_calc_btn = make_accent_button("⚡ 一键批量计算")
        self.batch_calc_btn.clicked.connect(self._batch_calculate)
        self.batch_calc_btn.hide()
        left_layout.addWidget(self.batch_calc_btn)

        splitter.addWidget(left_panel)

        # 中列: 详情 + 计算
        mid_scroll = QScrollArea()
        mid_scroll.setWidgetResizable(True)
        mid_scroll.setFrameShape(QFrame.NoFrame)
        self.detail_widget = QWidget()
        self.detail_layout = QVBoxLayout(self.detail_widget)
        self.detail_layout.setContentsMargins(8, 0, 8, 0)
        self.detail_layout.setSpacing(8)
        mid_scroll.setWidget(self.detail_widget)

        self.detail_placeholder = QLabel("📊 请导入费率方案文件夹，然后从左侧选择条款")
        self.detail_placeholder.setAlignment(Qt.AlignCenter)
        self.detail_placeholder.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 16px; padding: 60px;")
        self.detail_layout.addWidget(self.detail_placeholder)
        self.detail_layout.addStretch()

        splitter.addWidget(mid_scroll)

        # 右列: 保费汇总
        right_panel = QWidget()
        right_panel.setMinimumWidth(220)
        right_panel.setMaximumWidth(400)
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(6)

        summary_title = QLabel("💰 保费汇总")
        summary_title.setStyleSheet(f"font-weight: 600; font-size: 14px; color: {AnthropicColors.TEXT_PRIMARY};")
        right_layout.addWidget(summary_title)

        self.premium_list_area = QScrollArea()
        self.premium_list_area.setWidgetResizable(True)
        self.premium_list_area.setFrameShape(QFrame.NoFrame)
        self.premium_list_widget = QWidget()
        self.premium_list_layout = QVBoxLayout(self.premium_list_widget)
        self.premium_list_layout.setContentsMargins(0, 0, 0, 0)
        self.premium_list_layout.setSpacing(4)
        self.premium_list_area.setWidget(self.premium_list_widget)
        right_layout.addWidget(self.premium_list_area, 1)

        self.premium_empty_label = QLabel("计算附加险保费后\n将自动添加到此列表")
        self.premium_empty_label.setAlignment(Qt.AlignCenter)
        self.premium_empty_label.setStyleSheet(f"color: {AnthropicColors.TEXT_TERTIARY}; font-size: 12px; padding: 20px;")
        self.premium_list_layout.addWidget(self.premium_empty_label)

        self.addon_total_label = QLabel("附加险合计: ¥0.00")
        self.addon_total_label.setStyleSheet(f"font-weight: 600; font-size: 13px; color: {AnthropicColors.ACCENT}; padding: 8px;")
        right_layout.addWidget(self.addon_total_label)

        self.annual_total_label = QLabel("保单预估年保费: ¥0.00")
        self.annual_total_label.setStyleSheet(f"font-weight: 700; font-size: 15px; color: #10b981; padding: 8px; background: #ecfdf5; border-radius: 8px;")
        right_layout.addWidget(self.annual_total_label)

        # 短期保费计算
        short_card = GlassCard()
        short_layout = QVBoxLayout(short_card)
        short_layout.setContentsMargins(12, 10, 12, 10)
        short_layout.setSpacing(6)
        short_layout.addWidget(QLabel("📅 短期保费计算"))

        date_row1 = QHBoxLayout()
        date_row1.addWidget(QLabel("起保日:"))
        self.start_date = QDateEdit()
        self.start_date.setCalendarPopup(True)
        self.start_date.setDate(QDate.currentDate())
        self.start_date.dateChanged.connect(self._calc_short_term)
        date_row1.addWidget(self.start_date)
        short_layout.addLayout(date_row1)

        date_row2 = QHBoxLayout()
        date_row2.addWidget(QLabel("终止日:"))
        self.end_date = QDateEdit()
        self.end_date.setCalendarPopup(True)
        self.end_date.setDate(QDate.currentDate().addYears(1))
        self.end_date.dateChanged.connect(self._calc_short_term)
        date_row2.addWidget(self.end_date)
        short_layout.addLayout(date_row2)

        self.short_term_result = QLabel("")
        self.short_term_result.setWordWrap(True)
        self.short_term_result.setStyleSheet(f"font-size: 12px; color: {AnthropicColors.TEXT_PRIMARY};")
        short_layout.addWidget(self.short_term_result)
        right_layout.addWidget(short_card)

        splitter.addWidget(right_panel)

        # Splitter 伸缩策略: 左右固定，中间弹性
        splitter.setStretchFactor(0, 0)  # 左列不自动伸展
        splitter.setStretchFactor(1, 1)  # 中间弹性
        splitter.setStretchFactor(2, 0)  # 右列不自动伸展
        splitter.setSizes([300, 500, 280])

        main_layout.addWidget(splitter, 1)

        # 底部日志
        self.log_display = QTextEdit()
        self.log_display.setReadOnly(True)
        self.log_display.setMaximumHeight(100)
        self.log_display.setStyleSheet(f"""
            QTextEdit {{ background: {AnthropicColors.BG_DARK}; color: {AnthropicColors.TEXT_LIGHT};
                border-radius: 8px; padding: 6px; font-size: 11px; font-family: monospace; }}
        """)
        main_layout.addWidget(self.log_display)

    # ---------- 信号接收 ----------
    def receive_main_premium(self, total, per_person):
        self.main_premium = total
        self.per_person_premium = per_person
        self.main_premium_input.setValue(total)
        self.per_person_input.setValue(per_person)
        self._log(f"收到主险保费: {fmt_currency(total)}，每人: {fmt_currency(per_person)}", "success")

    def receive_full_data(self, data):
        """接收主险完整计算数据"""
        self.full_main_data = data
        self.main_premium = data.get("totalPremium", 0)
        self.per_person_premium = data.get("perPersonPremium", 0)
        self.main_premium_input.setValue(self.main_premium)
        self.per_person_input.setValue(self.per_person_premium)
        # 设置总保额
        insured_amount = data.get("insuredAmount", 0)
        if insured_amount > 0:
            self.main_sum_insured = insured_amount
            self.main_sum_insured_input.setValue(insured_amount)
        method = data.get("method", "")
        base_rate = data.get("baseRate", 0)
        adjusted_rate = data.get("adjustedRate", 0)
        coeff_product = data.get("coeffProduct", 1)
        count = data.get("employeeCount", 0)
        industry = data.get("industryClass", "")
        limit_val = data.get("perPersonLimit", 0)
        salary_val = data.get("annualSalary", 0)
        dis_table = data.get("disabilityTable", "none")
        product_name = data.get("selectedProduct", "")
        # 格式化费率显示（可能是对象）
        def fmt_rate(r):
            if isinstance(r, dict):
                return " | ".join(f"{k}: {v*100:.4f}%" if isinstance(v, (int, float)) else f"{k}: {v}" for k, v in r.items())
            if isinstance(r, (int, float)):
                return f"{r*100:.4f}%"
            return str(r) if r else "—"
        status_parts = [f"🟢 {product_name or method}"]
        if industry:
            status_parts.append(f"{industry}类")
        status_parts.append(f"基准{fmt_rate(base_rate)}")
        if adjusted_rate:
            status_parts.append(f"调整后{fmt_rate(adjusted_rate)}")
        status_parts.append(f"系数积{coeff_product:.4f}")
        if insured_amount > 0:
            status_parts.append(f"保额{fmt_currency(insured_amount)}")
        if limit_val:
            status_parts.append(f"限额{fmt_currency(limit_val)}")
        elif salary_val:
            status_parts.append(f"工资{fmt_currency(salary_val)}")
        if count > 0:
            status_parts.append(f"{count}人")
        if dis_table != "none":
            status_parts.append(f"伤残{dis_table}")
        self.main_data_status.setText(" · ".join(status_parts))
        self.main_data_status.setStyleSheet(f"font-size: 11px; color: #10b981;")
        self._log(f"收到主险完整数据: 险种={product_name}, 基准费率={fmt_rate(base_rate)}, "
                  f"调整后费率={fmt_rate(adjusted_rate)}, 保额={fmt_currency(insured_amount)}", "success")

    # ---------- 日志 ----------
    def _log(self, msg, level="info"):
        from datetime import datetime
        time_str = datetime.now().strftime("%H:%M:%S")
        prefix = {"error": "❌", "warn": "⚠️", "success": "✅"}.get(level, "ℹ️")
        self.log_display.append(f"[{time_str}] {prefix} {msg}")

    # ---------- 数据加载 ----------
    def _load_json(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "加载费率数据", "", "JSON (*.json)")
        if not file_path:
            return
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if not isinstance(data.get("entries"), list):
                raise ValueError("JSON 格式无效: 缺少 entries 数组")
            self.rate_data = data
            self.filtered_entries = list(data["entries"])
            self.load_status.setText(f"已加载 {len(data['entries'])} 条 ({os.path.basename(file_path)})")
            self._render_clause_list()
            self._log(f"加载成功: {len(data['entries'])} 个费率方案", "success")
        except Exception as e:
            self._log(f"JSON 加载失败: {e}", "error")

    def _load_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择费率方案文件夹")
        if not folder:
            return
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self._log("python-docx 未安装，请运行: pip install python-docx", "error")
            return
        entries = []
        docx_files = [f for f in os.listdir(folder) if "费率方案" in f and f.endswith(".docx") and not f.startswith("~$")]
        if not docx_files:
            self._log("未找到费率方案 docx 文件", "warn")
            return
        self._log(f"发现 {len(docx_files)} 个费率方案文件，开始智能分类解析...")
        type_counts = {}
        for fname in sorted(docx_files):
            try:
                fpath = os.path.join(folder, fname)
                doc = DocxDocument(fpath)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                tables = []
                for tbl in doc.tables:
                    rows = []
                    for row in tbl.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    if rows:
                        tables.append(rows)
                entry = self._classify_entry(fname, paragraphs, tables)
                if entry:
                    entries.append(entry)
                    rt = entry.get("rateType", "unknown")
                    type_counts[rt] = type_counts.get(rt, 0) + 1
            except Exception as e:
                self._log(f"解析失败: {fname} - {e}", "warn")
        self.rate_data = {"entries": entries}
        self.filtered_entries = list(entries)
        self.load_status.setText(f"已加载 {len(entries)} 条 (文件夹)")
        self._render_clause_list()
        type_summary = ", ".join(f"{ADDON_TYPES.get(k, {}).get('label', k)}:{v}" for k, v in sorted(type_counts.items()))
        self._log(f"解析完成: {len(entries)} 个费率方案 [{type_summary}]", "success")

    def _classify_entry(self, filename, paragraphs, tables):
        """智能分类附加险类型 — 基于文件名关键词 + 文本内容分析"""
        name = filename.replace(".docx", "").replace("中国太平洋财产保险股份有限公司", "")
        clause_name = name.replace("费率方案", "").strip()
        m = re.match(r"附加(.+?)(?:条款|保险)?$", clause_name)
        if m:
            clause_name = "附加" + m.group(1)

        entry = {
            "clauseName": clause_name,
            "fullName": filename.replace(".docx", ""),
            "industry": "雇主责任保险",
            "sourceFile": filename,
        }
        full_text = " ".join(paragraphs)
        substantive = [p for p in paragraphs
                       if "中国太平洋财产保险股份有限公司" not in p
                       and not (p.endswith("费率方案") and len(p) < 100)]

        # ===== Step 1: 关键词匹配确定类型 =====
        detected_type = None
        for keywords, addon_type in ADDON_KEYWORD_MAP:
            if any(kw in clause_name or kw in full_text for kw in keywords):
                detected_type = addon_type
                break

        # ===== Step 2: 规范类检测 =====
        reg_keywords = ["不涉及保险费的调整", "属于规范类", "不涉及费率", "不另收保险费"]
        is_regulatory = (not tables and not detected_type and substantive and
                         all(any(kw in p for kw in reg_keywords) or "保单最终保险费" in p or "工资总额" in p
                             for p in substantive))
        if is_regulatory:
            return {**entry, "rateType": "regulatory", "description": substantive[0] if substantive else ""}

        # ===== Step 3: 按检测类型构建 entry =====

        # 解析所有表格（通用）
        coeff_tables = self._parse_tables(tables, full_text)

        # 提取百分比
        percentages = []
        for p in paragraphs:
            for pct_m in re.finditer(r"([\d.]+)\s*[%％]", p):
                percentages.append({"value": float(pct_m.group(1)), "context": p})

        # --- modifier_coeff: 误工费 ---
        if detected_type == "modifier_coeff":
            return {**entry, "rateType": "modifier_coeff",
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": "调整后主险保费 = 主险保费 × 免赔天数调整系数"}

        # --- sudden_death: 突发疾病身故 ---
        if detected_type == "sudden_death":
            base_pct = None
            for pi in percentages:
                if "基准保险费" in pi["context"] or "每人每次事故赔偿限额" in pi["context"]:
                    base_pct = pi["value"]
                    break
            if not base_pct:
                base_pct = 6.6  # 默认
            return {**entry, "rateType": "sudden_death",
                    "basePercent": base_pct,
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": f"基准保费 = 每人限额 × {base_pct}% × 人数，再乘以系数调整"}

        # --- per_person_rate: 工伤补充/特定人员 ---
        if detected_type == "per_person_rate":
            rate_info = {}
            for pi in percentages:
                ctx = pi["context"]
                if "已购买工伤保险" in ctx or "有工伤" in ctx:
                    rate_info["with_injury_insurance"] = pi["value"]
                elif "未购买工伤保险" in ctx or "无工伤" in ctx:
                    rate_info["without_injury_insurance"] = pi["value"]
                elif not rate_info:
                    rate_info["default"] = pi["value"]
            return {**entry, "rateType": "per_person_rate",
                    "rateInfo": rate_info,
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": "保费 = 每人保费 × 费率% × 人数 × 系数"}

        # --- property_loss: 员工个人特定财产损失 ---
        if detected_type == "property_loss":
            base_premium = 20  # 基本保险费=基准保险费=20元
            base_rate_pct = 1.5  # 基准费率1.5%
            for pi in percentages:
                base_rate_pct = pi["value"]
                break
            for p in paragraphs:
                amt_m = re.search(r"基[本准]保险费[=＝]?\s*(\d+)\s*元", p)
                if amt_m:
                    base_premium = int(amt_m.group(1))
                    break
            return {**entry, "rateType": "property_loss",
                    "basePremium": base_premium,
                    "baseRatePercent": base_rate_pct,
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": f"保费 = ({base_premium}元 + 每人赔偿限额 × {base_rate_pct}%) × 系数积 × 承保人数"}

        # --- per_person_base: 药品服务 ---
        if detected_type == "per_person_base":
            base_amount = 300  # 默认
            for p in paragraphs:
                amt_m = re.search(r"(\d+)\s*元[/／每]人", p)
                if amt_m:
                    base_amount = int(amt_m.group(1))
                    break
            return {**entry, "rateType": "per_person_base",
                    "baseAmount": base_amount,
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": f"保费 = {base_amount}元/人 × 系数 × 人数"}

        # --- disability_adjust: 劳务关系人员 ---
        if detected_type == "disability_adjust":
            return {**entry, "rateType": "disability_adjust",
                    "adjustCoeffs": dict(DISABILITY_ADJUST_COEFFS),
                    "description": substantive[0] if substantive else "",
                    "formula": "保费 = 每人保费 × 伤残调整系数 × 人数"}

        # --- formula_sum: 雇主法律责任/一次性伤残 ---
        if detected_type == "formula_sum":
            base_rate_factor = 1.0
            for pi in percentages:
                if "90" in str(pi["value"]):
                    base_rate_factor = 0.9
                elif "95" in str(pi["value"]):
                    base_rate_factor = 0.95
                elif "100" in str(pi["value"]):
                    base_rate_factor = 1.0
                elif "110" in str(pi["value"]):
                    base_rate_factor = 1.1
                elif "120" in str(pi["value"]):
                    base_rate_factor = 1.2
            if "90" in full_text and "一次性伤残" not in clause_name:
                base_rate_factor = 0.9
            elif "95" in full_text and "一次性伤残" not in clause_name:
                base_rate_factor = 0.95
            # 一次性伤残根据 ABCD 款
            if "一次性伤残" in clause_name:
                if "A款" in clause_name or "（A）" in clause_name:
                    base_rate_factor = 0.9
                elif "B款" in clause_name or "（B）" in clause_name:
                    base_rate_factor = 1.0
                elif "C款" in clause_name or "（C）" in clause_name:
                    base_rate_factor = 1.1
                elif "D款" in clause_name or "（D）" in clause_name:
                    base_rate_factor = 1.2
            return {**entry, "rateType": "formula_sum",
                    "baseRateFactor": base_rate_factor,
                    "coefficientTables": coeff_tables,
                    "description": substantive[0] if substantive else "",
                    "formula": f"保费 = Σ(每人限额 × 主险基准费率 × {base_rate_factor} × 人数 × 系数积)"}

        # --- deduction: 突发疾病除外 ---
        if detected_type == "deduction":
            deduct_pct = 5.0  # 默认5%
            for pi in percentages:
                deduct_pct = pi["value"]
                break
            return {**entry, "rateType": "deduction",
                    "deductPercent": deduct_pct,
                    "description": substantive[0] if substantive else "",
                    "formula": f"减收 = 主险保费 × {deduct_pct}%"}

        # --- no_calc: 月申报/员工自动承保 ---
        if detected_type == "no_calc":
            return {**entry, "rateType": "no_calc",
                    "description": substantive[0] if substantive else full_text[:300],
                    "formula": "本条款有计费说明但无需单独计算附加保费"}

        # ===== Step 4: 未匹配关键词，按内容分析 =====
        if coeff_tables:
            base_premium = {"description": "未找到基准保险费描述"}
            for p in paragraphs:
                pct_m = re.search(r"([\d.]+)\s*[%％]", p)
                if pct_m and ("基准保险费" in p or "主险保险费的" in p):
                    base_premium = {"description": p, "percentage": float(pct_m.group(1))}
                    break
                mult_m = re.search(r"主险保险费的\s*([\d.]+)\s*倍", p)
                if mult_m:
                    base_premium = {"description": p, "multiplier": float(mult_m.group(1))}
                    break
            formula = "保险费 = 基准保险费 × 各项费率调整系数的乘积"
            for p in paragraphs:
                if "保险费" in p and ("×" in p or "＝" in p or "乘积" in p):
                    formula = p
                    break
            return {**entry, "rateType": "table_coefficient", "basePremium": base_premium,
                    "coefficientTables": coeff_tables, "formula": formula,
                    "description": substantive[0] if substantive else ""}

        # 简单百分比
        for p in paragraphs:
            pct_m = re.search(r"([\d.]+)\s*[%％]", p)
            if pct_m:
                return {**entry, "rateType": "simple_percentage",
                        "percentage": float(pct_m.group(1)), "description": p}
            mult_m = re.search(r"主险保险费的\s*([\d.]+)\s*倍", p)
            if mult_m:
                mult = float(mult_m.group(1))
                return {**entry, "rateType": "simple_percentage",
                        "percentage": mult * 100, "multiplier": mult, "description": p}

        return {**entry, "rateType": "regulatory", "description": full_text[:200]}

    def _parse_tables(self, tables, full_text=""):
        """通用表格解析"""
        coeff_tables = []
        for raw_table in tables:
            if len(raw_table) < 2:
                continue
            header = raw_table[0]
            rows = []
            for i in range(1, len(raw_table)):
                if len(raw_table[i]) < 2:
                    continue
                param, coeff = raw_table[i][0], raw_table[i][1]
                if not param or not coeff:
                    continue
                rows.append({"parameter": param, "coefficient": coeff,
                             "parsedValue": self._parse_coeff_value(coeff)})
            if rows:
                coeff_tables.append({
                    "name": header[0] if header else "调整系数",
                    "headerRow": header,
                    "supportsInterpolation": "线性插值" in full_text or "插值" in full_text,
                    "rows": rows,
                })
        return coeff_tables

    def _parse_coeff_value(self, text):
        text = text.strip().replace("，", ",").replace("（", "(").replace("）", ")")
        range_m = re.match(r"[\[(]?\s*([\d.]+)\s*[,]\s*([\d.]+)\s*[\])]?", text)
        if range_m:
            return {"type": "range", "min": float(range_m.group(1)), "max": float(range_m.group(2)), "display": text}
        num_m = re.match(r"^([\d.]+)$", text)
        if num_m:
            return {"type": "fixed", "value": float(num_m.group(1)), "display": text}
        return {"type": "text", "display": text}

    # ---------- 搜索/筛选 ----------
    def _filter_entries(self, keyword=""):
        if not self.rate_data:
            return
        query = keyword.strip().lower()
        self.filtered_entries = [e for e in self.rate_data["entries"]
                                 if not query or query in e.get("clauseName", "").lower()
                                 or query in e.get("fullName", "").lower()]
        self._render_clause_list()

    def _render_clause_list(self):
        self.clause_list.clear()
        for entry in self.filtered_entries:
            rt = entry.get("rateType", "")
            type_info = ADDON_TYPES.get(rt, {"label": rt, "color": "#6b7280"})
            item = QListWidgetItem(f"{entry['clauseName']}  [{type_info['label']}]")
            item.setData(Qt.UserRole, entry)
            self.clause_list.addItem(item)

    def _on_clause_selected(self, row):
        if row < 0 or row >= len(self.filtered_entries):
            return
        self.selected_entry = self.filtered_entries[row]
        self.coeff_selections = {}
        self._render_detail()
        rt = self.selected_entry.get("rateType", "")
        type_label = ADDON_TYPES.get(rt, {}).get("label", rt)
        self._log(f"选中: {self.selected_entry['clauseName']} [{type_label}]")

    # ---------- 详情渲染 ----------
    def _render_detail(self):
        while self.detail_layout.count():
            item = self.detail_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        entry = self.selected_entry
        if not entry:
            placeholder = QLabel("📊 请导入费率方案文件夹，然后从左侧选择条款")
            placeholder.setAlignment(Qt.AlignCenter)
            placeholder.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 16px; padding: 60px;")
            self.detail_layout.addWidget(placeholder)
            self.detail_layout.addStretch()
            return

        rate_type = entry.get("rateType", "")
        type_info = ADDON_TYPES.get(rate_type, {"label": rate_type, "color": "#6b7280"})

        # 条款名称 + 类型标签
        header_row = QHBoxLayout()
        name_label = QLabel(entry["clauseName"])
        name_label.setStyleSheet(f"font-weight: 700; font-size: 16px; color: {AnthropicColors.TEXT_PRIMARY};")
        header_row.addWidget(name_label)
        type_badge = QLabel(type_info["label"])
        type_badge.setStyleSheet(f"background: {type_info['color']}; color: white; padding: 2px 10px; "
                                 f"border-radius: 10px; font-size: 11px; font-weight: 600;")
        type_badge.setFixedHeight(22)
        header_row.addWidget(type_badge)
        header_row.addStretch()
        header_widget = QWidget()
        header_widget.setLayout(header_row)
        self.detail_layout.addWidget(header_widget)

        # 描述
        if entry.get("description"):
            desc = QLabel(entry["description"])
            desc.setWordWrap(True)
            desc.setStyleSheet(f"font-size: 12px; color: {AnthropicColors.TEXT_SECONDARY};")
            self.detail_layout.addWidget(desc)

        # 公式说明
        if entry.get("formula"):
            formula_label = QLabel(f"📐 {entry['formula']}")
            formula_label.setWordWrap(True)
            formula_label.setStyleSheet(f"padding: 10px; background: #eff6ff; border-radius: 8px; font-size: 12px; color: #1e40af;")
            self.detail_layout.addWidget(formula_label)

        # ===== 按类型渲染专属输入区 =====

        if rate_type == "regulatory":
            reg_label = QLabel("📋 规范类条款\n本条款不涉及保险费的调整")
            reg_label.setAlignment(Qt.AlignCenter)
            reg_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 14px; padding: 30px;")
            self.detail_layout.addWidget(reg_label)
            self.detail_layout.addStretch()
            return

        if rate_type == "no_calc":
            no_label = QLabel("📋 本条款有计费规则说明，但无需单独计算附加保费\n其影响已包含在主险或其他条款中")
            no_label.setAlignment(Qt.AlignCenter)
            no_label.setWordWrap(True)
            no_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 13px; padding: 30px;")
            self.detail_layout.addWidget(no_label)
            self._add_manual_section()
            self.detail_layout.addStretch()
            return

        if rate_type == "included_in_main":
            inc_label = QLabel("📦 纳入主险保险金额\n本附加险承保的财产应纳入主险保险金额，按主险费率计收保险费，不另收附加保险费")
            inc_label.setAlignment(Qt.AlignCenter)
            inc_label.setWordWrap(True)
            inc_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 14px; padding: 30px;")
            self.detail_layout.addWidget(inc_label)
            self.detail_layout.addStretch()
            return

        if rate_type == "daily_prorate":
            dp_label = QLabel("📅 按日比例计算\n保费按日比例计算：保险金额 × 保单费率 × (天数 / 365)")
            dp_label.setAlignment(Qt.AlignCenter)
            dp_label.setWordWrap(True)
            dp_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 14px; padding: 30px;")
            self.detail_layout.addWidget(dp_label)
            self.detail_layout.addStretch()
            return

        if rate_type == "simple_percentage":
            # Check for insuranceTypeRates
            insurance_type_rates = entry.get("insuranceTypeRates")
            if insurance_type_rates:
                # Add insurance type combo
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("保险类型:"))
                self.insurance_type_combo = QComboBox()
                for ins_type in insurance_type_rates.keys():
                    self.insurance_type_combo.addItem(ins_type)
                type_row.addWidget(self.insurance_type_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)

                # Display rate table
                rate_card = GlassCard()
                rate_layout = QVBoxLayout(rate_card)
                rate_layout.setContentsMargins(12, 10, 12, 10)
                rate_card.setStyleSheet("background: #eff6ff; border: 1px solid #93c5fd; border-radius: 8px;")
                rate_title = QLabel("📊 保险类型费率表")
                rate_title.setStyleSheet("font-size: 13px; font-weight: 600; color: #1e40af;")
                rate_layout.addWidget(rate_title)
                for ins_type, rate in insurance_type_rates.items():
                    rate_label = QLabel(f"  {ins_type}: {rate}%")
                    rate_label.setStyleSheet("font-size: 12px; color: #1e3a8a;")
                    rate_layout.addWidget(rate_label)
                self.detail_layout.addWidget(rate_card)

                # ratioMultiplier support (customInputs)
                if not hasattr(self, 'custom_input_widgets'):
                    self.custom_input_widgets = {}
                custom_inputs = entry.get("customInputs", [])
                for ci_item in custom_inputs:
                    ci_key = ci_item.get("key", "")
                    ci_label = ci_item.get("label", "自定义输入")
                    ci_unit = ci_item.get("unit", "")
                    row = QHBoxLayout()
                    unit_text = f" ({ci_unit})" if ci_unit else ""
                    row.addWidget(QLabel(f"{ci_label}{unit_text}:"))
                    spin = QDoubleSpinBox()
                    spin.setRange(0, 999999999999)
                    spin.setDecimals(2)
                    if ci_unit == "元":
                        spin.setSuffix(" 元")
                    elif ci_unit:
                        spin.setSuffix(f" {ci_unit}")
                    row.addWidget(spin)
                    row.addStretch()
                    row_w = QWidget()
                    row_w.setLayout(row)
                    self.detail_layout.addWidget(row_w)
                    if ci_key:
                        self.custom_input_widgets[ci_key] = spin

                # ratioHint support
                if entry.get("ratioMultiplier") and entry["ratioMultiplier"].get("ratioHint"):
                    hint_label = QLabel(entry["ratioMultiplier"]["ratioHint"])
                    hint_label.setWordWrap(True)
                    hint_label.setStyleSheet("padding: 8px; background: #fef3c7; border: 1px solid #fbbf24; border-radius: 6px; font-size: 12px;")
                    self.detail_layout.addWidget(hint_label)
            else:
                # Original display
                pct = entry.get("percentage", 0)
                mult = entry.get("multiplier")
                if mult:
                    info = QLabel(f"费率: 主险保费 × {mult}")
                else:
                    info = QLabel(f"费率: 主险保费 × {pct}%")
                info.setStyleSheet(f"padding: 12px; background: #eff6ff; border-radius: 8px; font-size: 13px;")
                self.detail_layout.addWidget(info)

        elif rate_type == "modifier_coeff":
            warn_label = QLabel("⚠️ 此条款直接调整主险保费（优先计算）")
            warn_label.setStyleSheet("padding: 12px; background: #fef2f2; border: 2px solid #dc2626; "
                                     "border-radius: 8px; font-size: 14px; font-weight: 700; color: #dc2626;")
            self.detail_layout.addWidget(warn_label)
            hint = QLabel("选择免赔天数对应的调整系数后计算")
            hint.setStyleSheet("padding: 6px; font-size: 12px; color: #991b1b;")
            self.detail_layout.addWidget(hint)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "sudden_death":
            base_pct = entry.get("basePercent", 6.6)
            hint = QLabel(f"基准费率: 每人限额 × {base_pct}%\n"
                          "附加险限额 < 主险限额 → 减收\n"
                          "附加险限额 > 主险限额 → 加收")
            hint.setWordWrap(True)
            hint.setStyleSheet(f"padding: 10px; background: #fef2f2; border: 1px solid #fca5a5; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            # 限额输入
            grid = QGridLayout()
            grid.addWidget(QLabel("附加险每人限额(万元):"), 0, 0)
            self.addon_limit_input = QDoubleSpinBox()
            self.addon_limit_input.setRange(0, 9999)
            self.addon_limit_input.setDecimals(1)
            self.addon_limit_input.setSuffix(" 万元")
            grid.addWidget(self.addon_limit_input, 0, 1)
            grid.addWidget(QLabel("主险每人限额(万元):"), 1, 0)
            self.main_limit_display = QDoubleSpinBox()
            self.main_limit_display.setRange(0, 9999)
            self.main_limit_display.setDecimals(1)
            self.main_limit_display.setSuffix(" 万元")
            if self.full_main_data:
                limit_wan = self.full_main_data.get("perPersonLimit", 0) / 10000
                self.main_limit_display.setValue(limit_wan)
            grid.addWidget(self.main_limit_display, 1, 1)
            grid_w = QWidget()
            grid_w.setLayout(grid)
            self.detail_layout.addWidget(grid_w)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "per_person_rate":
            rate_info = entry.get("rateInfo", {})
            hint_text = "费率说明:\n"
            if rate_info.get("with_injury_insurance"):
                hint_text += f"  已购买工伤保险: {rate_info['with_injury_insurance']}%\n"
            if rate_info.get("without_injury_insurance"):
                hint_text += f"  未购买工伤保险: {rate_info['without_injury_insurance']}%\n"
            if rate_info.get("default"):
                hint_text += f"  默认费率: {rate_info['default']}%\n"
            hint = QLabel(hint_text.strip())
            hint.setWordWrap(True)
            hint.setStyleSheet(f"padding: 10px; background: #fffbeb; border: 1px solid #fbbf24; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            # 工伤保险状态选择
            self.injury_insurance_combo = QComboBox()
            self.injury_insurance_combo.addItem("已购买工伤保险", "with")
            self.injury_insurance_combo.addItem("未购买工伤保险", "without")
            self.detail_layout.addWidget(self.injury_insurance_combo)
            # 独立人数输入（不一定全员）
            default_count = self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1
            count_row = QHBoxLayout()
            count_row.addWidget(QLabel("适用人数:"))
            self.addon_count_input = QSpinBox()
            self.addon_count_input.setRange(1, 999999)
            self.addon_count_input.setValue(default_count)
            self.addon_count_input.setSuffix("")
            count_row.addWidget(self.addon_count_input)
            count_row.addStretch()
            count_w = QWidget()
            count_w.setLayout(count_row)
            self.detail_layout.addWidget(count_w)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "property_loss":
            bp = entry.get("basePremium", 20)
            br = entry.get("baseRatePercent", 1.5)
            hint = QLabel(f"公式: 保费 = (基本保险费 + 每次事故每人赔偿限额 × 基准费率) × 系数积 × 承保人数\n基本保险费: {bp}元，基准费率: {br}%")
            hint.setWordWrap(True)
            hint.setStyleSheet(f"padding: 10px; background: #f0fdfa; border: 1px solid #5eead4; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            grid = QGridLayout()
            grid.addWidget(QLabel("每次事故每人赔偿限额(万元):"), 0, 0)
            self.property_limit_input = QDoubleSpinBox()
            self.property_limit_input.setRange(0, 9999)
            self.property_limit_input.setDecimals(1)
            self.property_limit_input.setSuffix(" 万元")
            grid.addWidget(self.property_limit_input, 0, 1)
            default_count = self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1
            grid.addWidget(QLabel("适用人数:"), 1, 0)
            self.addon_count_input = QSpinBox()
            self.addon_count_input.setRange(1, 999999)
            self.addon_count_input.setValue(default_count)
            self.addon_count_input.setSuffix("")
            grid.addWidget(self.addon_count_input, 1, 1)
            grid_w = QWidget()
            grid_w.setLayout(grid)
            self.detail_layout.addWidget(grid_w)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "per_person_base":
            base_amt = entry.get("baseAmount", 300)
            hint = QLabel(f"基准: {base_amt}元/人，乘以调整系数后 × 人数")
            hint.setStyleSheet(f"padding: 10px; background: #ecfeff; border: 1px solid #67e8f9; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "disability_adjust":
            coeffs = entry.get("adjustCoeffs", DISABILITY_ADJUST_COEFFS)
            hint = QLabel(f"根据主险选择的伤残赔偿附表调整:\n"
                          f"  附表1: ×{coeffs.get('table1', 0.995)}\n"
                          f"  附表2: ×{coeffs.get('table2', 1.072)}\n"
                          f"  附表3: ×{coeffs.get('table3', 0.919)}")
            hint.setStyleSheet(f"padding: 10px; background: #ecfdf5; border: 1px solid #6ee7b7; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            if self.full_main_data:
                dis_table = self.full_main_data.get("disabilityTable", "none")
                if dis_table != "none":
                    auto_label = QLabel(f"🔗 已自动识别主险伤残附表: {dis_table}")
                    auto_label.setStyleSheet(f"color: #10b981; font-size: 12px; font-weight: 600;")
                    self.detail_layout.addWidget(auto_label)
            # 独立人数输入（不一定全员）
            default_count = self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1
            count_row = QHBoxLayout()
            count_row.addWidget(QLabel("适用人数:"))
            self.addon_count_input = QSpinBox()
            self.addon_count_input.setRange(1, 999999)
            self.addon_count_input.setValue(default_count)
            self.addon_count_input.setSuffix("")
            count_row.addWidget(self.addon_count_input)
            count_row.addStretch()
            count_w = QWidget()
            count_w.setLayout(count_row)
            self.detail_layout.addWidget(count_w)

        elif rate_type == "formula_sum":
            factor = entry.get("baseRateFactor", 1.0)
            hint = QLabel(f"求和公式: Σ(每人限额 × 主险基准费率 × {factor} × 人数 × 系数积)\n"
                          f"基准费率来源: 主险基准费率 × {factor}")
            hint.setWordWrap(True)
            hint.setStyleSheet(f"padding: 10px; background: #eff6ff; border: 1px solid #93c5fd; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "deduction":
            deduct_pct = entry.get("deductPercent", 5.0)
            hint = QLabel(f"减收: 主险保费 × {deduct_pct}%\n将从主险保费中扣减此金额")
            hint.setStyleSheet(f"padding: 10px; background: #fdf2f8; border: 1px solid #f9a8d4; border-radius: 8px; font-size: 12px;")
            self.detail_layout.addWidget(hint)

        elif rate_type == "main_premium_modifier":
            warn_label = QLabel("⚠️ 此条款直接调整主险保费（优先计算）")
            warn_label.setStyleSheet("padding: 12px; background: #fef2f2; border: 2px solid #dc2626; "
                                     "border-radius: 8px; font-size: 14px; font-weight: 700; color: #dc2626;")
            self.detail_layout.addWidget(warn_label)
            if entry.get("formula"):
                formula_mod = QLabel(entry["formula"])
                formula_mod.setStyleSheet("padding: 8px; font-size: 12px; color: #991b1b;")
                self.detail_layout.addWidget(formula_mod)

            # simple_deduction 处理
            if entry.get("modifierType") == "simple_deduction":
                deductions = entry.get("insuranceTypeDeductions", {})
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("主险类型:"))
                self.deduction_insurance_combo = QComboBox()
                for ins_type in deductions.keys():
                    self.deduction_insurance_combo.addItem(ins_type)
                type_row.addWidget(self.deduction_insurance_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)

                # 显示减免表
                deduct_label = QLabel("减免比例:")
                deduct_label.setStyleSheet("font-weight: 700; margin-top: 8px;")
                self.detail_layout.addWidget(deduct_label)
                deduct_text = "\n".join([f"  {k}: {v}%" for k, v in deductions.items()])
                deduct_info = QLabel(deduct_text)
                deduct_info.setStyleSheet("padding: 8px; background: #fef2f2; border-radius: 6px; font-size: 12px;")
                self.detail_layout.addWidget(deduct_info)

            elif entry.get("modifierType") == "water_level_deduction":
                wlt = entry.get("waterLevelTable", {})
                ins_types = list(wlt.get("rates", {}).keys())
                # 险种下拉
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("主险类型:"))
                self.insurance_type_combo = QComboBox()
                for ins_type in ins_types:
                    self.insurance_type_combo.addItem(ins_type)
                type_row.addWidget(self.insurance_type_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)
                # 水位输入
                wl_row = QHBoxLayout()
                unit = wlt.get("unit", "cm")
                wl_row.addWidget(QLabel(f"水位线高度 ({unit}):"))
                self.water_level_input = QDoubleSpinBox()
                self.water_level_input.setRange(0, 200)
                self.water_level_input.setDecimals(1)
                self.water_level_input.setSuffix(f" {unit}")
                wl_row.addWidget(self.water_level_input)
                wl_row.addStretch()
                wl_w = QWidget()
                wl_w.setLayout(wl_row)
                self.detail_layout.addWidget(wl_w)
                # 水位减费率表格
                heights = wlt.get("heights", [])
                table_text = f"水位({unit})" + "".join(f"  |  {t}" for t in ins_types) + "\n"
                for wi, h in enumerate(heights):
                    row_text = f"  {h}"
                    for t in ins_types:
                        r = wlt["rates"][t][wi]
                        row_text += f"  |  {'—' if r == 0 else f'{r}%'}"
                    table_text += row_text + "\n"
                table_label = QLabel(table_text.strip())
                table_label.setStyleSheet("padding: 8px; background: #dbeafe; border: 1px solid #93c5fd; "
                                          "border-radius: 6px; font-size: 11px; font-family: monospace;")
                self.detail_layout.addWidget(table_label)

            elif entry.get("modifierType") == "regional_deduction":
                regions = entry.get("regions", [])
                ins_types = ["财产基本险", "财产综合险", "财产一切险"]
                overrides = entry.get("insuranceTypeOverrides", {})
                # 险种下拉
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("主险类型:"))
                self.insurance_type_combo = QComboBox()
                for ins_type in ins_types:
                    suffix = " (不调整)" if overrides.get(ins_type) == 0 else ""
                    self.insurance_type_combo.addItem(f"{ins_type}{suffix}", ins_type)
                type_row.addWidget(self.insurance_type_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)
                # 地区单选
                from PyQt5.QtWidgets import QRadioButton, QButtonGroup
                region_label = QLabel("承保地点区域:")
                region_label.setStyleSheet("font-weight: 600; margin-top: 8px;")
                self.detail_layout.addWidget(region_label)
                self.region_combo = QComboBox()
                for reg in regions:
                    self.region_combo.addItem(f"{reg['label']} (减{reg['deductPct']}%)", reg["key"])
                self.detail_layout.addWidget(self.region_combo)
                # 减费表格
                table_text = "\n".join([f"  {r['label']}: 减{r['deductPct']}%" for r in regions])
                table_info = QLabel(table_text)
                table_info.setWordWrap(True)
                table_info.setStyleSheet("padding: 8px; background: #fef3c7; border: 1px solid #fbbf24; "
                                         "border-radius: 6px; font-size: 12px;")
                self.detail_layout.addWidget(table_info)

            else:
                # 原有的系数表逻辑
                table = (entry.get("coefficientTables") or [{}])[0]
                col_keys = table.get("columns", [])
                col_labels = table.get("columnLabels", [])
                if len(col_keys) > 1:
                    type_row = QHBoxLayout()
                    type_row.addWidget(QLabel("适用主险类型:"))
                    self.modifier_insurance_combo = QComboBox()
                    for ci in range(1, len(col_keys)):
                        self.modifier_insurance_combo.addItem(col_labels[ci] if ci < len(col_labels) else col_keys[ci], col_keys[ci])
                    type_row.addWidget(self.modifier_insurance_combo)
                    type_row.addStretch()
                    type_w = QWidget()
                    type_w.setLayout(type_row)
                    self.detail_layout.addWidget(type_w)
                ratio_row = QHBoxLayout()
                input_label_text = table.get("inputLabel", "限额÷保额比例")
                ratio_row.addWidget(QLabel(f"{input_label_text}:"))
                self.modifier_ratio_input = QDoubleSpinBox()
                self.modifier_ratio_input.setRange(0, 100)
                self.modifier_ratio_input.setDecimals(1)
                self.modifier_ratio_input.setSuffix(" %")
                ratio_row.addWidget(self.modifier_ratio_input)
                ratio_row.addStretch()
                ratio_w = QWidget()
                ratio_w.setLayout(ratio_row)
                self.detail_layout.addWidget(ratio_w)
                self._render_modifier_table(table)

        elif rate_type == "conditional_simple":
            # 勾选条件计算
            from PyQt5.QtWidgets import QCheckBox
            self.conditional_checkbox = QCheckBox(entry.get("checkboxLabel", "主险未包含本附加险责任"))
            self.conditional_checkbox.setChecked(entry.get("defaultChecked", False))
            self.conditional_checkbox.setStyleSheet("font-size: 13px; font-weight: 600; color: #5b21b6; padding: 8px; "
                                                     "background: #f5f3ff; border: 1px solid #a78bfa; border-radius: 8px;")
            self.detail_layout.addWidget(self.conditional_checkbox)
            # 险种选择（base_rate_division需要baseRates时）
            wc = entry.get("whenChecked", {})
            if wc.get("formulaType") == "base_rate_division" and wc.get("baseRates"):
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("主险类型:"))
                self.insurance_type_combo = QComboBox()
                for ins_type, br in wc["baseRates"].items():
                    self.insurance_type_combo.addItem(f"{ins_type} (基准费率 {br}%)", ins_type)
                type_row.addWidget(self.insurance_type_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)
            # 自定义输入框（customInputs）
            custom_inputs = entry.get("customInputs", [])
            self.conditional_custom_inputs = {}
            for ci in custom_inputs:
                ci_row = QHBoxLayout()
                ci_row.addWidget(QLabel(f"{ci['label']}:"))
                ci_spin = QDoubleSpinBox()
                ci_spin.setRange(0, 999999999999)
                ci_spin.setDecimals(2)
                if ci.get("unit"):
                    ci_spin.setSuffix(f" {ci['unit']}")
                ci_spin.setMinimumWidth(150)
                ci_row.addWidget(ci_spin)
                ci_row.addStretch()
                ci_w = QWidget()
                ci_w.setLayout(ci_row)
                self.detail_layout.addWidget(ci_w)
                self.conditional_custom_inputs[ci["key"]] = ci_spin
            # 公式说明
            if entry.get("formula"):
                formula_hint = QLabel(f"📐 {entry['formula']}")
                formula_hint.setWordWrap(True)
                formula_hint.setStyleSheet("padding: 10px; background: #eff6ff; border-radius: 8px; font-size: 12px; color: #1e40af;")
                self.detail_layout.addWidget(formula_hint)

        elif rate_type == "base_rate_formula":
            self._render_base_rate_formula(entry)

        elif rate_type == "formula_conditional":
            fc_desc = entry.get("description") or entry.get("formula") or ""
            if fc_desc:
                fc_label = QLabel(f"📐 条件公式\n{fc_desc}")
                fc_label.setWordWrap(True)
                fc_label.setStyleSheet("padding: 12px; background: #eef2ff; border: 1px solid #818cf8; "
                                       "border-radius: 8px; font-size: 12px; color: #3730a3;")
                self.detail_layout.addWidget(fc_label)
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                self._render_addon_coeff_table(table, ti)

        elif rate_type == "table_coefficient":
            # Check for insuranceTypeRates
            insurance_type_rates = entry.get("insuranceTypeRates")
            start_table_index = 0

            if insurance_type_rates:
                # Add insurance type combo
                type_row = QHBoxLayout()
                type_row.addWidget(QLabel("保险类型:"))
                self.insurance_type_combo = QComboBox()
                for ins_type in insurance_type_rates.keys():
                    self.insurance_type_combo.addItem(ins_type)
                type_row.addWidget(self.insurance_type_combo)
                type_row.addStretch()
                type_w = QWidget()
                type_w.setLayout(type_row)
                self.detail_layout.addWidget(type_w)

                # Display rate table
                rate_card = GlassCard()
                rate_layout = QVBoxLayout(rate_card)
                rate_layout.setContentsMargins(12, 10, 12, 10)
                rate_card.setStyleSheet("background: #eff6ff; border: 1px solid #93c5fd; border-radius: 8px;")
                rate_title = QLabel("📊 保险类型基准费率表")
                rate_title.setStyleSheet("font-size: 13px; font-weight: 600; color: #1e40af;")
                rate_layout.addWidget(rate_title)
                for ins_type, rate in insurance_type_rates.items():
                    rate_label = QLabel(f"  {ins_type}: {rate}%")
                    rate_label.setStyleSheet("font-size: 12px; color: #1e3a8a;")
                    rate_layout.addWidget(rate_label)
                self.detail_layout.addWidget(rate_card)

                # Skip first coefficient table
                start_table_index = 1
            else:
                # Original basePremium display
                if entry.get("basePremium"):
                    bp = entry["basePremium"]
                    bp_label = QLabel(f"基准保险费: {bp.get('description', '')}")
                    bp_label.setWordWrap(True)
                    bp_label.setStyleSheet(f"padding: 10px; background: #eff6ff; border-radius: 8px; font-size: 12px;")
                    self.detail_layout.addWidget(bp_label)

            # customInputs 支持
            if not hasattr(self, 'custom_input_widgets'):
                self.custom_input_widgets = {}
            custom_inputs = entry.get("customInputs", [])
            for ci_item in custom_inputs:
                ci_key = ci_item.get("key", "")
                ci_label = ci_item.get("label", "自定义输入")
                ci_unit = ci_item.get("unit", "")
                row = QHBoxLayout()
                unit_text = f" ({ci_unit})" if ci_unit else ""
                row.addWidget(QLabel(f"{ci_label}{unit_text}:"))
                spin = QDoubleSpinBox()
                spin.setRange(0, 999999999999)
                spin.setDecimals(2)
                if ci_unit == "元":
                    spin.setSuffix(" 元")
                elif ci_unit:
                    spin.setSuffix(f" {ci_unit}")
                row.addWidget(spin)
                row.addStretch()
                row_w = QWidget()
                row_w.setLayout(row)
                self.detail_layout.addWidget(row_w)
                if ci_key:
                    self.custom_input_widgets[ci_key] = spin

            # ratioHint 支持
            if entry.get("ratioMultiplier") and entry["ratioMultiplier"].get("ratioHint"):
                hint_label = QLabel(entry["ratioMultiplier"]["ratioHint"])
                hint_label.setWordWrap(True)
                hint_label.setStyleSheet("padding: 8px; background: #fef3c7; border: 1px solid #fbbf24; border-radius: 6px; font-size: 12px;")
                self.detail_layout.addWidget(hint_label)

            # Render coefficient tables (skip first one if insuranceTypeRates exists)
            coeff_tables = entry.get("coefficientTables", [])
            for ti in range(start_table_index, len(coeff_tables)):
                self._render_addon_coeff_table(coeff_tables[ti], ti)

        # 计算按钮（非展示类类型）
        if rate_type not in ("regulatory", "no_calc", "included_in_main", "daily_prorate", "formula_conditional"):
            calc_btn = make_accent_button("🧮 计算附加险保费")
            calc_btn.clicked.connect(self._calculate)
            self.detail_layout.addWidget(calc_btn)

        # 结果区
        self.addon_result_label = QLabel("")
        self.addon_result_label.setWordWrap(True)
        self.addon_result_label.setStyleSheet(f"font-size: 13px; padding: 10px;")
        self.addon_result_label.hide()
        self.detail_layout.addWidget(self.addon_result_label)

        # 核保经验计费
        self._add_manual_section()
        self.detail_layout.addStretch()

    def _add_manual_section(self):
        """添加核保经验计费区域"""
        manual_card = GlassCard()
        manual_layout = QVBoxLayout(manual_card)
        manual_layout.setContentsMargins(12, 10, 12, 10)
        manual_layout.addWidget(QLabel("✏️ 核保经验计费"))
        manual_hint = QLabel("手动输入附加险保费（覆盖公式计算结果）")
        manual_hint.setWordWrap(True)
        manual_hint.setStyleSheet(f"font-size: 11px; color: {AnthropicColors.TEXT_TERTIARY};")
        manual_layout.addWidget(manual_hint)
        manual_row = QHBoxLayout()
        self.manual_input = QDoubleSpinBox()
        self.manual_input.setRange(-999999999, 999999999)
        self.manual_input.setDecimals(2)
        self.manual_input.setSuffix(" 元")
        manual_row.addWidget(self.manual_input, 1)
        manual_btn = QPushButton("确认计入")
        manual_btn.setCursor(Qt.PointingHandCursor)
        manual_btn.setStyleSheet(f"QPushButton {{ background: #f59e0b; color: white; border: none; border-radius: 6px; padding: 8px 14px; }}")
        manual_btn.clicked.connect(self._add_manual_premium)
        manual_row.addWidget(manual_btn)
        manual_layout.addLayout(manual_row)
        self.detail_layout.addWidget(manual_card)

    def _render_addon_coeff_table(self, table, table_idx):
        card = GlassCard()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(4)
        title = QLabel(table.get("name", "调整系数"))
        title.setStyleSheet(f"font-weight: 600; font-size: 13px;")
        layout.addWidget(title)
        if table.get("supportsInterpolation"):
            interp = QLabel("支持线性插值")
            interp.setStyleSheet(f"font-size: 11px; color: #3b82f6;")
            layout.addWidget(interp)
        for ri, row in enumerate(table.get("rows", [])):
            sel = self.coeff_selections.get(table_idx)
            is_selected = sel and sel.get("rowIdx") == ri
            btn = QPushButton(f"{row['parameter']}    {row['coefficient']}")
            btn.setCursor(Qt.PointingHandCursor)
            bg = AnthropicColors.ACCENT if is_selected else AnthropicColors.BG_PRIMARY
            fg = AnthropicColors.TEXT_LIGHT if is_selected else AnthropicColors.TEXT_PRIMARY
            btn.setStyleSheet(f"""
                QPushButton {{ background: {bg}; color: {fg}; border: 1px solid {AnthropicColors.BORDER};
                    border-radius: 6px; padding: 5px 10px; font-size: 12px; text-align: left; }}
                QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; }}
            """)
            btn.clicked.connect(lambda checked, ti=table_idx, r=ri: self._select_addon_coeff_row(ti, r))
            layout.addWidget(btn)
        sel = self.coeff_selections.get(table_idx)
        if sel and sel.get("parsedValue", {}).get("type") == "range":
            pv = sel["parsedValue"]
            current_val = sel.get("value", pv["min"])
            slider_layout = QHBoxLayout()
            slider_label = QLabel(f"{current_val:.4f}")
            slider_label.setStyleSheet(f"font-weight: 600; color: {AnthropicColors.ACCENT};")
            slider_layout.addWidget(QLabel("精确系数:"))
            slider_layout.addWidget(slider_label)
            slider = QSlider(Qt.Horizontal)
            slider.setMinimum(int(pv["min"] * 1000))
            slider.setMaximum(int(pv["max"] * 1000))
            slider.setValue(int(current_val * 1000))
            ti_ref = table_idx
            slider.valueChanged.connect(lambda v, ti=ti_ref, lbl=slider_label: self._on_addon_slider_change(ti, v, lbl))
            slider_layout.addWidget(slider, 1)
            layout.addLayout(slider_layout)
        self.detail_layout.addWidget(card)

    def _select_addon_coeff_row(self, table_idx, row_idx):
        entry = self.selected_entry
        if not entry or not entry.get("coefficientTables"):
            return
        table = entry["coefficientTables"][table_idx]
        row = table["rows"][row_idx]
        pv = row["parsedValue"]
        value = pv.get("value", pv.get("min", 1.0)) if pv["type"] != "text" else 1.0
        self.coeff_selections[table_idx] = {
            "rowIdx": row_idx, "value": value, "parsedValue": pv,
            "parameter": row["parameter"], "coefficient": row["coefficient"],
        }
        self._render_detail()

    def _on_addon_slider_change(self, table_idx, int_value, label_widget):
        value = int_value / 1000.0
        if table_idx in self.coeff_selections:
            self.coeff_selections[table_idx]["value"] = value
        label_widget.setText(f"{value:.4f}")

    def _get_selected_insurance_type(self):
        """获取选择的保险类型（从多个可能的combo中）"""
        for attr in ['insurance_type_combo', 'deduction_insurance_combo', 'batch_insurance_combo']:
            combo = getattr(self, attr, None)
            if combo:
                return combo.currentText()
        return None

    def _get_coeff_product(self, entry, start_index=0):
        """计算所有已选系数表的系数乘积

        Args:
            entry: 条款数据
            start_index: 起始索引,当insuranceTypeRates存在时从1开始跳过第一个表
        """
        product = 1.0
        details = []
        coeff_tables = entry.get("coefficientTables", [])
        for ti in range(start_index, len(coeff_tables)):
            sel = self.coeff_selections.get(ti)
            if not sel:
                table_name = coeff_tables[ti]["name"]
                raise ValueError(f"请选择「{table_name}」的系数值")
            product *= sel["value"]
            details.append({"table": coeff_tables[ti]["name"],
                            "parameter": sel["parameter"], "value": sel["value"]})
        return product, details

    # ---------- 计算引擎 (重构版) ----------
    # ---------- 主险保费调整条款：多列系数表渲染 ----------
    def _render_base_rate_formula(self, entry):
        """渲染 base_rate_formula 类型的详情面板"""
        has_sub = bool(entry.get("subFormulas"))
        has_multi_instance = bool(entry.get("maxInstances") and entry.get("serviceTypes"))

        # 顶部信息框
        info_card = GlassCard()
        info_layout = QVBoxLayout(info_card)
        info_layout.setContentsMargins(12, 10, 12, 10)
        info_card.setStyleSheet("background: #ecfeff; border: 2px solid #0891b2; border-radius: 8px;")
        title = QLabel("📊 基准费率计算条款")
        title.setStyleSheet("font-size: 14px; font-weight: 700; color: #0891b2;")
        info_layout.addWidget(title)
        if has_multi_instance:
            max_inst = entry.get("maxInstances", 5)
            cap_text = f" | 费率上限: {entry['rateCap']}%" if entry.get("rateCap") else ""
            rate_label = QLabel(f"多实例服务 (最多 {max_inst} 个){cap_text}")
            rate_label.setStyleSheet("font-size: 12px; color: #155e75;")
            info_layout.addWidget(rate_label)
        elif not has_sub:
            rate_val = entry.get("baseRate")
            formula_type = entry.get("formulaType", "")
            if rate_val is not None:
                rate_text = f"{rate_val}%"
            elif entry.get("baseRateIsMainRate"):
                rate_text = "同主险基准费率"
            elif formula_type in ("policy_rate_prorate", "policy_rate_simple", "auto_appreciation_c"):
                rate_text = "保单费率(自动计算)"
            elif formula_type in ("auto_appreciation", "auto_appreciation_b"):
                rate_text = "主险保费×升值率×50%"
            else:
                rate_text = "见系数表"
            cap_text = f" | 费率上限: {entry['rateCap']}%" if entry.get("rateCap") else ""
            rate_label = QLabel(f"基准费率: {rate_text}{cap_text}")
            rate_label.setStyleSheet("font-size: 12px; color: #155e75;")
            info_layout.addWidget(rate_label)
        if entry.get("formula"):
            fl = QLabel(entry["formula"])
            fl.setWordWrap(True)
            fl.setStyleSheet("font-size: 12px; color: #0e7490; margin-top: 4px;")
            info_layout.addWidget(fl)
        self.detail_layout.addWidget(info_card)

        # 法律服务多实例模式
        if has_multi_instance:
            self._render_legal_service_multi(entry)
            # 系数表（跳过前几个服务项目表）
            for ti, table in enumerate(entry.get("coefficientTables", [])):
                if ti >= 4:
                    self._render_addon_coeff_table(table, ti)
            return

        # 多责任子公式选择
        self.brf_sub_inputs = []
        if has_sub:
            sub_row = QHBoxLayout()
            sub_row.addWidget(QLabel("选择责任类型:"))
            self.brf_sub_combo = QComboBox()
            for si, sf in enumerate(entry["subFormulas"]):
                label = sf["name"]
                if sf.get("baseRate") is not None:
                    label += f" (基准费率: {sf['baseRate']}%)"
                self.brf_sub_combo.addItem(label, si)
            self.brf_sub_combo.currentIndexChanged.connect(self._on_brf_sub_changed)
            sub_row.addWidget(self.brf_sub_combo)
            sub_row.addStretch()
            sub_w = QWidget()
            sub_w.setLayout(sub_row)
            self.detail_layout.addWidget(sub_w)

            # 子公式信息
            self.brf_sub_info_label = QLabel()
            self.brf_sub_info_label.setWordWrap(True)
            self.brf_sub_info_label.setStyleSheet("padding: 8px; background: #f0fdfa; border-radius: 6px; font-size: 12px; color: #0e7490;")
            self.detail_layout.addWidget(self.brf_sub_info_label)

            # 输入框容器
            self.brf_inputs_widget = QWidget()
            self.brf_inputs_layout = QVBoxLayout(self.brf_inputs_widget)
            self.brf_inputs_layout.setContentsMargins(0, 0, 0, 0)
            self.detail_layout.addWidget(self.brf_inputs_widget)
            self._on_brf_sub_changed(0)
        else:
            # 单一公式
            if entry.get("baseRateIsMainRate"):
                rate_row = QHBoxLayout()
                rate_row.addWidget(QLabel("主险基准费率 (%):"))
                self.brf_main_rate_input = QDoubleSpinBox()
                self.brf_main_rate_input.setRange(0, 100)
                self.brf_main_rate_input.setDecimals(3)
                self.brf_main_rate_input.setSuffix(" %")
                rate_row.addWidget(self.brf_main_rate_input)
                rate_row.addStretch()
                rate_w = QWidget()
                rate_w.setLayout(rate_row)
                self.detail_layout.addWidget(rate_w)
            else:
                self.brf_main_rate_input = None

            self.brf_inputs_widget = None
            custom_inputs = entry.get("customInputs", [])
            for ci, inp in enumerate(custom_inputs):
                row = QHBoxLayout()
                unit_text = f" ({inp['unit']})" if inp.get("unit") else ""
                row.addWidget(QLabel(f"{inp['label']}{unit_text}:"))
                spin = QDoubleSpinBox()
                spin.setRange(0, 999999999999)
                spin.setDecimals(2)
                if inp.get("unit") == "元":
                    spin.setSuffix(" 元")
                elif inp.get("unit"):
                    spin.setSuffix(f" {inp['unit']}")
                row.addWidget(spin)
                row.addStretch()
                row_w = QWidget()
                row_w.setLayout(row)
                self.detail_layout.addWidget(row_w)
                self.brf_sub_inputs.append(spin)

        # 系数表
        for ti, table in enumerate(entry.get("coefficientTables", [])):
            self._render_addon_coeff_table(table, ti)

    def _render_legal_service_multi(self, entry):
        """渲染法律服务多实例UI"""
        self.legal_service_instances = []
        self.legal_service_widgets = []
        self.legal_service_entry = entry

        # 实例容器
        self.legal_instances_container = QWidget()
        self.legal_instances_layout = QVBoxLayout(self.legal_instances_container)
        self.legal_instances_layout.setContentsMargins(0, 0, 0, 0)
        self.detail_layout.addWidget(self.legal_instances_container)

        # 添加按钮
        add_btn = QPushButton(f"+ 添加服务实例 (最多 {entry.get('maxInstances', 5)} 个)")
        add_btn.setStyleSheet(
            "QPushButton { background: #0891b2; color: white; border: none; "
            "border-radius: 6px; padding: 8px 16px; font-weight: 600; }"
            "QPushButton:hover { background: #0e7490; }"
        )
        add_btn.clicked.connect(self._add_legal_service_instance)
        self.detail_layout.addWidget(add_btn)

        # 默认添加一个实例
        self._add_legal_service_instance()

    def _add_legal_service_instance(self):
        """添加一个法律服务实例"""
        entry = self.legal_service_entry
        max_inst = entry.get("maxInstances", 5)
        if len(self.legal_service_instances) >= max_inst:
            self._log(f"已达最大实例数: {max_inst}", "warn")
            return

        idx = len(self.legal_service_instances)
        instance = {"serviceKey": entry["serviceTypes"][0]["key"], "liabilityLimit": 0, "serviceCount": 1}
        self.legal_service_instances.append(instance)

        # 实例UI卡片
        card = GlassCard()
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(10, 8, 10, 8)
        card.setStyleSheet("background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px;")

        # 标题行+删除按钮
        header = QHBoxLayout()
        header.addWidget(QLabel(f"<b>服务实例 #{idx + 1}</b>"))
        header.addStretch()
        del_btn = QPushButton("✕")
        del_btn.setFixedSize(24, 24)
        del_btn.setStyleSheet("QPushButton { background: none; border: none; color: #dc2626; font-size: 14px; }"
                              "QPushButton:hover { color: #991b1b; }")
        del_btn.clicked.connect(lambda checked, i=idx: self._remove_legal_service_instance(i))
        header.addWidget(del_btn)
        header_w = QWidget()
        header_w.setLayout(header)
        card_layout.addWidget(header_w)

        # 服务类型下拉
        svc_row = QHBoxLayout()
        svc_row.addWidget(QLabel("服务类型:"))
        svc_combo = QComboBox()
        for svc in entry["serviceTypes"]:
            svc_combo.addItem(f"{svc['label']} ({svc['baseRate']}%)", svc["key"])
        svc_combo.currentIndexChanged.connect(
            lambda ci, i=idx, cb=svc_combo: self._on_legal_svc_changed(i, cb))
        svc_row.addWidget(svc_combo)
        svc_w = QWidget()
        svc_w.setLayout(svc_row)
        card_layout.addWidget(svc_w)

        # 限额和次数
        inputs_row = QHBoxLayout()
        inputs_row.addWidget(QLabel("每次事故责任限额 (元):"))
        limit_spin = QDoubleSpinBox()
        limit_spin.setRange(0, 999999999999)
        limit_spin.setDecimals(2)
        limit_spin.setSuffix(" 元")
        limit_spin.valueChanged.connect(lambda v, i=idx: self._on_legal_limit_changed(i, v))
        inputs_row.addWidget(limit_spin)
        inputs_row.addWidget(QLabel("服务次数:"))
        count_spin = QSpinBox()
        count_spin.setRange(1, 9999)
        count_spin.setValue(1)
        count_spin.valueChanged.connect(lambda v, i=idx: self._on_legal_count_changed(i, v))
        inputs_row.addWidget(count_spin)
        inputs_w = QWidget()
        inputs_w.setLayout(inputs_row)
        card_layout.addWidget(inputs_w)

        self.legal_instances_layout.addWidget(card)
        self.legal_service_widgets.append({
            "card": card, "svc_combo": svc_combo,
            "limit_spin": limit_spin, "count_spin": count_spin
        })

    def _remove_legal_service_instance(self, idx):
        """删除法律服务实例"""
        if idx >= len(self.legal_service_instances):
            return
        self.legal_service_instances.pop(idx)
        w = self.legal_service_widgets.pop(idx)
        w["card"].setParent(None)
        w["card"].deleteLater()
        # 更新剩余实例的索引标题
        for i, wgt in enumerate(self.legal_service_widgets):
            for child in wgt["card"].findChildren(QLabel):
                if child.text().startswith("<b>服务实例"):
                    child.setText(f"<b>服务实例 #{i + 1}</b>")
                    break

    def _on_legal_svc_changed(self, idx, combo):
        if idx < len(self.legal_service_instances):
            self.legal_service_instances[idx]["serviceKey"] = combo.currentData()

    def _on_legal_limit_changed(self, idx, val):
        if idx < len(self.legal_service_instances):
            self.legal_service_instances[idx]["liabilityLimit"] = val

    def _on_legal_count_changed(self, idx, val):
        if idx < len(self.legal_service_instances):
            self.legal_service_instances[idx]["serviceCount"] = val

    def _calc_legal_service(self, entry):
        """计算法律服务多实例保费"""
        instances = getattr(self, "legal_service_instances", [])
        if not instances:
            raise ValueError("请至少添加一个法律服务实例")

        coeff_product = 1.0
        coeff_details = []
        tables = entry.get("coefficientTables", [])
        if tables and len(tables) > 5 and hasattr(self, "coeff_selections"):
            for ti, sel in self.coeff_selections.items():
                if ti >= 4 and sel is not None:
                    coeff_product *= float(sel)
                    coeff_details.append({"name": tables[ti].get("name", ""), "value": float(sel)})

        total_premium = 0.0
        formula_parts = []
        for i, inst in enumerate(instances):
            svc_type = None
            for s in entry["serviceTypes"]:
                if s["key"] == inst["serviceKey"]:
                    svc_type = s
                    break
            if not svc_type:
                raise ValueError(f"未找到服务类型: {inst['serviceKey']}")
            base_rate = svc_type["baseRate"] / 100
            limit_val = inst["liabilityLimit"]
            count_val = inst["serviceCount"]
            if not limit_val or limit_val <= 0:
                raise ValueError(f"服务实例 #{i + 1}: 请输入有效的责任限额")
            if not count_val or count_val <= 0:
                raise ValueError(f"服务实例 #{i + 1}: 请输入有效的服务次数")
            inst_premium = limit_val * count_val * base_rate * coeff_product
            total_premium += inst_premium
            formula_parts.append(
                f"#{i + 1} {svc_type['label']}: "
                f"{fmt_currency(limit_val)} × {count_val}次 × {svc_type['baseRate']}% = "
                f"{fmt_currency(limit_val * count_val * base_rate)}")

        if coeff_details:
            coeff_str = " × ".join(f"{d['value']:.4f}" for d in coeff_details)
            formula_parts.append(f"系数积: {coeff_str} = {coeff_product:.6f}")
        if entry.get("rateCap"):
            formula_parts.append(f"(费率上限: {entry['rateCap']}%)")
        formula_parts.append(f"合计保费 = {fmt_currency(total_premium)}")

        return {"type": "base_rate_formula", "premium": total_premium,
                "formulaDisplay": "\n".join(formula_parts)}

    def _on_brf_sub_changed(self, idx):
        """切换子公式时更新输入区"""
        entry = self.selected_entry
        if not entry or not entry.get("subFormulas"):
            return
        sf = entry["subFormulas"][idx]
        # 更新信息
        info_parts = []
        if sf.get("formula"):
            info_parts.append(sf["formula"])
        if sf.get("baseRate") is not None:
            info_parts.append(f"基准费率: {sf['baseRate']}%")
        cap = entry.get("rateCap")
        if cap:
            info_parts.append(f"费率上限: {cap}%")
        self.brf_sub_info_label.setText("\n".join(info_parts))

        # 清空旧输入
        while self.brf_inputs_layout.count():
            item = self.brf_inputs_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.brf_sub_inputs = []

        # 添加新输入
        for ci, inp in enumerate(sf.get("customInputs", [])):
            row = QHBoxLayout()
            unit_text = f" ({inp['unit']})" if inp.get("unit") else ""
            row.addWidget(QLabel(f"{inp['label']}{unit_text}:"))
            spin = QDoubleSpinBox()
            spin.setRange(0, 999999999999)
            spin.setDecimals(2)
            if inp.get("unit") == "元":
                spin.setSuffix(" 元")
            elif inp.get("unit"):
                spin.setSuffix(f" {inp['unit']}")
            row.addWidget(spin)
            row.addStretch()
            row_w = QWidget()
            row_w.setLayout(row)
            self.brf_inputs_layout.addWidget(row_w)
            self.brf_sub_inputs.append(spin)

    def _calc_base_rate_formula(self, entry):
        """计算 base_rate_formula 类型"""
        if entry.get("maxInstances") and entry.get("serviceTypes"):
            return self._calc_legal_service(entry)
        has_sub = bool(entry.get("subFormulas"))
        if has_sub:
            sub_idx = self.brf_sub_combo.currentData() if hasattr(self, "brf_sub_combo") else 0
            if sub_idx is None or sub_idx < 0 or sub_idx >= len(entry["subFormulas"]):
                sub_idx = 0
            sf = entry["subFormulas"][sub_idx]
            base_rate = sf.get("baseRate")
            custom_inputs = sf.get("customInputs", [])
        else:
            base_rate = entry.get("baseRate")
            custom_inputs = entry.get("customInputs", [])

        rate_cap = entry.get("rateCap")
        formula_type = entry.get("formulaType")

        # 如果基准费率来自主险
        if base_rate is None and entry.get("baseRateIsMainRate"):
            if hasattr(self, "brf_main_rate_input") and self.brf_main_rate_input:
                base_rate = self.brf_main_rate_input.value()
            if not base_rate or base_rate <= 0:
                raise ValueError("请输入有效的主险基准费率")

        # 获取自定义输入值
        input_values = []
        input_labels = []
        if not hasattr(self, "brf_sub_inputs"):
            self.brf_sub_inputs = []
        for ci, inp in enumerate(custom_inputs):
            if ci < len(self.brf_sub_inputs):
                val = self.brf_sub_inputs[ci].value()
            else:
                raise ValueError(f"缺少输入值: {inp['label']}")
            if val <= 0:
                raise ValueError(f"请输入有效的{inp['label']}")
            input_values.append(val)
            input_labels.append(inp["label"])

        # 获取系数乘积
        coeff_product = 1.0
        coeff_details = []
        tables = entry.get("coefficientTables", [])
        if tables and hasattr(self, "coeff_selections"):
            for ti, sel in self.coeff_selections.items():
                if sel is not None:
                    coeff_details.append({"name": tables[ti].get("name", ""), "value": float(sel)})
                    coeff_product *= float(sel)

        # 保单费率按天分摊
        if formula_type == "policy_rate_prorate":
            main_si = getattr(self, 'main_sum_insured', 0)
            if not main_si or main_si <= 0:
                raise ValueError("请输入主险保险金额")
            if self.main_premium <= 0:
                raise ValueError("请输入主险保险费")
            policy_rate = self.main_premium / main_si
            effective_rate = policy_rate * (entry.get("policyRateMultiplier") or 1.0)
            rate_was_capped = False
            rate_cap = entry.get("rateCap")
            if rate_cap and effective_rate > rate_cap / 100:
                effective_rate = rate_cap / 100
                rate_was_capped = True
            # Read custom inputs: first=amount, second=days
            if len(input_values) < 2:
                raise ValueError("需要金额和天数两个输入")
            amount = input_values[0]
            days = input_values[1]
            p_days = self.policy_days if entry.get("useGlobalPolicyDays") else 365
            if p_days <= 0:
                p_days = 365
            premium = amount * effective_rate * (days / p_days)
            parts = [f"保单费率 = {fmt_currency(self.main_premium)} ÷ {fmt_currency(main_si)} = {policy_rate*100:.6f}%"]
            mult = entry.get("policyRateMultiplier", 1.0)
            if mult != 1.0:
                parts.append(f"费率乘数: ×{mult}")
                parts.append(f"有效费率: {effective_rate*100:.6f}%")
            if rate_was_capped:
                parts.append(f"⚠️ 费率 {policy_rate*mult*100:.4f}% > 上限 {rate_cap}%，按 {rate_cap}% 计算")
            parts.append(f"{input_labels[0]}: {fmt_currency(amount)}")
            parts.append(f"{input_labels[1]}: {days}天")
            parts.append(f"保单天数: {p_days}天")
            parts.append(f"保费 = {fmt_currency(amount)} × {effective_rate*100:.6f}% × {days} ÷ {p_days} = {fmt_currency(premium)}")
            return {"type": "base_rate_formula", "premium": premium, "formulaDisplay": "\n".join(parts)}

        # policy_rate_simple: 金额 × 保单费率
        if formula_type == "policy_rate_simple":
            main_si = getattr(self, 'main_sum_insured', 0)
            if not main_si or main_si <= 0:
                raise ValueError("请输入主险保险金额")
            if self.main_premium <= 0:
                raise ValueError("请输入主险保险费")
            policy_rate = self.main_premium / main_si
            amount = input_values[0] if input_values else 0
            if amount <= 0:
                raise ValueError(f"请输入{input_labels[0] if input_labels else '金额'}")
            premium = amount * policy_rate
            lbl = input_labels[0] if input_labels else "金额"
            return {"type": "base_rate_formula", "premium": premium,
                    "formulaDisplay": f"保单费率 = {fmt_currency(self.main_premium)} ÷ {fmt_currency(main_si)} = {policy_rate*100:.6f}%\n"
                                     f"{lbl}: {fmt_currency(amount)}\n"
                                     f"保费 = {fmt_currency(amount)} × {policy_rate*100:.6f}% = {fmt_currency(premium)}"}

        # auto_appreciation: 主险保费 × 升值率 × 50%
        if formula_type == "auto_appreciation":
            rate = input_values[0] if input_values else 0
            if rate <= 0:
                raise ValueError("请输入升值率")
            premium = self.main_premium * (rate / 100) * 0.5
            return {"type": "base_rate_formula", "premium": premium,
                    "formulaDisplay": f"主险保费 × 升值率 × 50%\n"
                                     f"{fmt_currency(self.main_premium)} × {rate}% × 50% = {fmt_currency(premium)}"}

        # auto_appreciation_b: 主险保费 × 升值比例 × 50% × (总保额-存货)÷总保额
        if formula_type == "auto_appreciation_b":
            rate = input_values[0] if input_values else 0
            inventory = input_values[1] if len(input_values) > 1 else 0
            if rate <= 0:
                raise ValueError("请输入升值比例")
            main_si = getattr(self, 'main_sum_insured', 0)
            if not main_si or main_si <= 0:
                raise ValueError("请输入主险保险金额")
            ratio = max(0, (main_si - inventory) / main_si)
            premium = self.main_premium * (rate / 100) * 0.5 * ratio
            return {"type": "base_rate_formula", "premium": premium,
                    "formulaDisplay": f"主险保费 × 升值比例 × 50% × (总保额-存货)÷总保额\n"
                                     f"{fmt_currency(self.main_premium)} × {rate}% × 50% × "
                                     f"({fmt_currency(main_si)}-{fmt_currency(inventory)})÷{fmt_currency(main_si)}\n"
                                     f"= {fmt_currency(self.main_premium)} × {rate}% × 50% × {ratio*100:.2f}% = {fmt_currency(premium)}"}

        # auto_appreciation_c: 增值保额 × 保单费率 × 50%
        if formula_type == "auto_appreciation_c":
            main_si = getattr(self, 'main_sum_insured', 0)
            if not main_si or main_si <= 0:
                raise ValueError("请输入主险保险金额")
            if self.main_premium <= 0:
                raise ValueError("请输入主险保险费")
            amount = input_values[0] if input_values else 0
            if amount <= 0:
                raise ValueError("请输入增值部分的保险金额")
            policy_rate = self.main_premium / main_si
            premium = amount * policy_rate * 0.5
            return {"type": "base_rate_formula", "premium": premium,
                    "formulaDisplay": f"增值保额 × 保单费率 × 50%\n"
                                     f"保单费率 = {fmt_currency(self.main_premium)} ÷ {fmt_currency(main_si)} = {policy_rate*100:.6f}%\n"
                                     f"{fmt_currency(amount)} × {policy_rate*100:.6f}% × 50% = {fmt_currency(premium)}"}

        # 特殊公式：主险保费比例
        if formula_type == "main_premium_ratio":
            custom_amount = input_values[0] if input_values else 0
            main_si = self.main_sum_insured if hasattr(self, "main_sum_insured") else 0
            if main_si <= 0:
                raise ValueError("请输入主险保险金额")
            premium = self.main_premium * (base_rate / 100) * custom_amount / main_si
            formula_str = (f"{self._fmt_currency(self.main_premium)} × {base_rate}% × "
                          f"{self._fmt_currency(custom_amount)} ÷ {self._fmt_currency(main_si)} = {self._fmt_currency(premium)}")
            return {"type": "base_rate_formula", "premium": premium, "formulaDisplay": formula_str}

        # 标准公式
        if base_rate is None or base_rate <= 0:
            raise ValueError("基准费率无效，请检查数据或输入主险基准费率")
        custom_product = 1.0
        for v in input_values:
            custom_product *= v

        effective_rate = (base_rate / 100) * coeff_product
        if rate_cap and effective_rate > (rate_cap / 100):
            effective_rate = rate_cap / 100

        premium = custom_product * effective_rate

        parts = []
        for i, val in enumerate(input_values):
            parts.append(f"{input_labels[i]}: {self._fmt_currency(val)}")
        parts.append(f"基准费率: {base_rate}%")
        if coeff_details:
            coeff_str = " × ".join(f"{d['value']:.4f}" for d in coeff_details)
            parts.append(f"系数积: {coeff_str} = {coeff_product:.6f}")
        raw_rate = (base_rate / 100) * coeff_product
        if rate_cap and raw_rate > (rate_cap / 100):
            parts.append(f"费率 {raw_rate*100:.4f}% > 上限 {rate_cap}%，按 {rate_cap}% 计算")
        parts.append(f"保费 = {self._fmt_currency(premium)}")

        return {
            "type": "base_rate_formula",
            "premium": premium,
            "formulaDisplay": "\n".join(parts),
        }

    @staticmethod
    def _fmt_currency(val):
        """格式化货币"""
        if val >= 10000:
            return f"¥{val:,.2f}"
        return f"¥{val:.2f}"

    def _render_modifier_table(self, table):
        """渲染 main_premium_modifier 的多列系数参考表"""
        col_labels = table.get("columnLabels", [])
        col_keys = table.get("columns", [])
        rows = table.get("rows", [])
        if not rows:
            return
        card = GlassCard()
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 10, 12, 10)
        title = QLabel(table.get("name", "调整系数表"))
        title.setStyleSheet("font-weight: 600; font-size: 13px;")
        layout.addWidget(title)
        if table.get("supportsInterpolation"):
            interp = QLabel("支持线性插值（可输入表中未列明的比例值）")
            interp.setStyleSheet("font-size: 11px; color: #3b82f6;")
            layout.addWidget(interp)
        # 表头
        header_text = "  |  ".join(col_labels)
        header_label = QLabel(header_text)
        header_label.setStyleSheet("font-weight: 600; font-size: 11px; padding: 4px; "
                                   "background: #f3f4f6; border-radius: 4px;")
        layout.addWidget(header_label)
        # 数据行
        for row in rows:
            cells = [row.get("parameter", "")]
            for ci in range(1, len(col_keys)):
                val = row.get(col_keys[ci], "-")
                cells.append(str(val))
            row_label = QLabel("  |  ".join(cells))
            row_label.setStyleSheet("font-size: 11px; padding: 2px 4px; font-family: monospace;")
            layout.addWidget(row_label)
        self.detail_layout.addWidget(card)

    # ---------- 主险保费调整条款：插值查表 ----------
    @staticmethod
    def _interpolate_modifier_table(table, col_key, ratio_percent):
        """在多列系数表中按比例插值查表"""
        rows = table.get("rows", [])
        ratios = []
        for row in rows:
            pct_str = row.get("parameter", "").replace("%", "")
            try:
                pct = float(pct_str)
            except ValueError:
                continue
            val = row.get(col_key)
            if val is not None:
                ratios.append((pct, float(val)))
        ratios.sort(key=lambda x: x[0])
        if not ratios:
            raise ValueError(f"系数表中无有效数据列 '{col_key}'")
        if ratio_percent <= ratios[0][0]:
            return ratios[0][1]
        if ratio_percent >= ratios[-1][0]:
            return ratios[-1][1]
        for i in range(len(ratios) - 1):
            if ratios[i][0] <= ratio_percent <= ratios[i + 1][0]:
                t = (ratio_percent - ratios[i][0]) / (ratios[i + 1][0] - ratios[i][0])
                return ratios[i][1] + t * (ratios[i + 1][1] - ratios[i][1])
        return ratios[-1][1]

    # ---------- 主险保费调整条款：计算 ----------
    def _calc_main_premium_modifier(self, entry):
        """计算主险保费调整（冰雹/台风/暴雪/通用每次事故赔偿限额）"""
        # simple_deduction: 按险种直接减免
        if entry.get("modifierType") == "simple_deduction":
            selected_type = self._get_selected_insurance_type()
            if not selected_type:
                raise ValueError("请选择主险类型")
            deductions = entry.get("insuranceTypeDeductions", {})
            deduct_pct = deductions.get(selected_type)
            if deduct_pct is None:
                raise ValueError(f"未找到{selected_type}的减免比例")
            if deduct_pct == 0:
                return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                        "originalPremium": self.main_premium, "adjustedPremium": self.main_premium,
                        "formulaDisplay": f"【{selected_type}】不涉及保费调整"}
            adjusted = self.main_premium * (1 - deduct_pct / 100)
            return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                    "originalPremium": self.main_premium, "adjustedPremium": adjusted,
                    "formulaDisplay": f"【{selected_type}】减收 {deduct_pct}%\n"
                                     f"{fmt_currency(self.main_premium)} × (1 - {deduct_pct}%) = {fmt_currency(adjusted)}"}

        # water_level_deduction: 水位线减费
        if entry.get("modifierType") == "water_level_deduction":
            wlt = entry.get("waterLevelTable", {})
            ins_type = self._get_selected_insurance_type()
            if not ins_type:
                raise ValueError("请选择主险类型")
            water_level = getattr(self, 'water_level_input', None)
            water_val = water_level.value() if water_level else 0
            if water_val <= 0:
                raise ValueError("请输入有效的水位线高度")
            type_rates = wlt.get("rates", {}).get(ins_type, [])
            heights = wlt.get("heights", [])
            if not type_rates or all(r == 0 for r in type_rates):
                return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                        "originalPremium": self.main_premium, "adjustedPremium": self.main_premium,
                        "formulaDisplay": f"【{ins_type}】不涉及保费调整（减免比例为0）"}
            # 线性插值
            deduct_pct = 0.0
            if water_val <= heights[0]:
                deduct_pct = type_rates[0]
            elif water_val >= heights[-1]:
                deduct_pct = type_rates[-1]
            else:
                for wi in range(len(heights) - 1):
                    if heights[wi] <= water_val <= heights[wi + 1]:
                        t = (water_val - heights[wi]) / (heights[wi + 1] - heights[wi])
                        deduct_pct = type_rates[wi] + t * (type_rates[wi + 1] - type_rates[wi])
                        break
            adjusted = self.main_premium * (1 - deduct_pct / 100)
            unit = wlt.get("unit", "cm")
            return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                    "originalPremium": self.main_premium, "adjustedPremium": adjusted,
                    "formulaDisplay": f"【{ins_type}】水位线 {water_val}{unit} → 减免比例 {deduct_pct:.2f}%\n"
                                     f"{fmt_currency(self.main_premium)} × (1 - {deduct_pct:.2f}%) = {fmt_currency(adjusted)}"}

        # regional_deduction: 地区减费
        if entry.get("modifierType") == "regional_deduction":
            ins_type = self._get_selected_insurance_type()
            if not ins_type:
                raise ValueError("请选择主险类型")
            overrides = entry.get("insuranceTypeOverrides", {})
            if overrides.get(ins_type) == 0:
                return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                        "originalPremium": self.main_premium, "adjustedPremium": self.main_premium,
                        "formulaDisplay": f"【{ins_type}】不涉及保费调整"}
            region_combo = getattr(self, 'region_combo', None)
            if not region_combo:
                raise ValueError("请选择承保地点区域")
            region_key = region_combo.currentData()
            regions = entry.get("regions", [])
            region = next((r for r in regions if r["key"] == region_key), None)
            if not region:
                raise ValueError("未找到对应地区")
            deduct_pct = region["deductPct"]
            adjusted = self.main_premium * (1 - deduct_pct / 100)
            return {"type": "main_premium_modifier", "isMainModifier": True, "premium": 0,
                    "originalPremium": self.main_premium, "adjustedPremium": adjusted,
                    "formulaDisplay": f"【{ins_type}】{region['label']} → 减收 {deduct_pct}%\n"
                                     f"{fmt_currency(self.main_premium)} × (1 - {deduct_pct}%) = {fmt_currency(adjusted)}"}

        table = (entry.get("coefficientTables") or [{}])[0]
        if not table.get("rows"):
            raise ValueError("缺少系数表数据")
        combo = getattr(self, 'modifier_insurance_combo', None)
        col_key = combo.currentData() if combo else table.get("columns", ["", ""])[1]
        ratio_input = getattr(self, 'modifier_ratio_input', None)
        ratio_percent = ratio_input.value() if ratio_input else 0
        if ratio_percent <= 0 or ratio_percent > 100:
            raise ValueError("请输入有效的限额÷保额比例（0-100）")
        col_label = ""
        col_labels = table.get("columnLabels", [])
        col_keys = table.get("columns", [])
        for ci, ck in enumerate(col_keys):
            if ck == col_key:
                col_label = col_labels[ci] if ci < len(col_labels) else ck
                break
        lookup_value = self._interpolate_modifier_table(table, col_key, ratio_percent)
        modifier_type = entry.get("modifierType", "coefficient")
        if modifier_type == "deductionPercent":
            deduct_ratio = lookup_value / 100
            adjusted = self.main_premium * (1 - deduct_ratio)
            formula = (f"【{col_label}】限额÷保额={ratio_percent}% → 减少比例={lookup_value:.1f}%\n"
                       f"{fmt_currency(self.main_premium)} × (1 - {lookup_value:.1f}%) = {fmt_currency(adjusted)}")
        else:
            adjusted = self.main_premium * lookup_value
            formula = (f"【{col_label}】限额÷保额={ratio_percent}% → 调整系数={lookup_value:.5f}\n"
                       f"{fmt_currency(self.main_premium)} × {lookup_value:.5f} = {fmt_currency(adjusted)}")
        return {
            "type": "main_premium_modifier",
            "premium": 0,
            "adjustedPremium": adjusted,
            "originalPremium": self.main_premium,
            "coefficient": lookup_value,
            "isMainModifier": True,
            "formulaDisplay": formula,
        }

    def _calc_conditional_simple(self, entry):
        """计算 conditional_simple 类型（勾选条件计算）"""
        checkbox = getattr(self, 'conditional_checkbox', None)
        is_checked = checkbox.isChecked() if checkbox else False
        # 根据勾选状态选择对应的计算配置
        if is_checked:
            calc_config = entry.get("whenChecked", {})
        else:
            calc_config = entry.get("whenUnchecked", {})
        # 如果没有对应配置，则不涉及保费调整
        if not calc_config:
            return {"type": "conditional_simple", "premium": 0,
                    "formulaDisplay": "主险已包含本附加险责任，不涉及保险费调整"}
        return self._exec_conditional_formula(calc_config, entry)

    def _exec_conditional_formula(self, calc_config, entry):
        """执行 conditional_simple 的具体计算公式"""
        formula_type = calc_config.get("formulaType", "")
        if formula_type == "simple_pct":
            pct = calc_config["percentage"]
            premium = self.main_premium * (pct / 100)
            return {"type": "conditional_simple", "premium": premium,
                    "formulaDisplay": f"主险保费 × {pct}%\n"
                                     f"{fmt_currency(self.main_premium)} × {pct}% = {fmt_currency(premium)}"}
        if formula_type == "base_rate_division":
            # 支持 multiplier 模式（如第三人照管财产）
            multiplier = calc_config.get("multiplier")
            if multiplier is not None:
                return self._calc_conditional_multiplier(calc_config, entry, multiplier)
            # 原有 baseRates + numeratorPct 模式
            ins_type = self._get_selected_insurance_type()
            if not ins_type:
                raise ValueError("请选择主险类型")
            base_rates = calc_config.get("baseRates", {})
            base_rate = base_rates.get(ins_type)
            if not base_rate:
                raise ValueError(f"未找到{ins_type}的基准费率")
            base_rate_decimal = base_rate / 100
            numerator_pct = calc_config["numeratorPct"]
            premium = self.main_premium * (numerator_pct / 100) / base_rate_decimal
            return {"type": "conditional_simple", "premium": premium,
                    "formulaDisplay": f"【{ins_type}】主险保费 × {numerator_pct}% ÷ 基准费率({base_rate}%)\n"
                                     f"{fmt_currency(self.main_premium)} × {numerator_pct}% ÷ {base_rate}% = {fmt_currency(premium)}"}
        raise ValueError(f"未知的conditional_simple公式类型: {formula_type}")

    def _calc_conditional_multiplier(self, calc_config, entry, multiplier):
        """计算 conditional_simple multiplier 模式（金额×费率×系数）"""
        custom_inputs = getattr(self, 'conditional_custom_inputs', {})
        addon_amount_spin = custom_inputs.get("addonAmount")
        addon_amount = addon_amount_spin.value() if addon_amount_spin else 0
        if addon_amount <= 0:
            raise ValueError("请输入本附加险保险金额")
        main_si = getattr(self, 'main_sum_insured', 0)
        if main_si <= 0:
            raise ValueError("主险保额不能为0")
        policy_rate = self.main_premium / main_si
        premium = addon_amount * policy_rate * multiplier
        desc = calc_config.get("description", "")
        return {"type": "conditional_simple", "premium": premium,
                "formulaDisplay": f"{desc}\n"
                                  f"保险费 = {fmt_currency(addon_amount)} × "
                                  f"({fmt_currency(self.main_premium)} ÷ {fmt_currency(main_si)}) × {multiplier}\n"
                                  f"= {fmt_currency(premium)}"}

    def _calculate(self):
        entry = self.selected_entry
        if not entry:
            self._log("请先选择费率方案", "warn")
            return
        self.main_premium = self.main_premium_input.value()
        rate_type = entry.get("rateType", "")
        if rate_type in ("regulatory", "no_calc", "included_in_main", "daily_prorate"):
            return
        if self.main_premium <= 0 and rate_type not in ("per_person_base", "property_loss", "base_rate_formula", "conditional_simple"):
            self._log("请输入有效的主险保费", "warn")
            return
        try:
            calc_method = {
                "simple_percentage": self._calc_simple,
                "modifier_coeff": self._calc_modifier_coeff,
                "sudden_death": self._calc_sudden_death,
                "per_person_rate": self._calc_per_person_rate,
                "per_person_base": self._calc_per_person_base,
                "disability_adjust": self._calc_disability_adjust,
                "property_loss": self._calc_property_loss,
                "formula_sum": self._calc_formula_sum,
                "deduction": self._calc_deduction,
                "table_coefficient": self._calc_table,
                "main_premium_modifier": self._calc_main_premium_modifier,
                "base_rate_formula": self._calc_base_rate_formula,
                "conditional_simple": self._calc_conditional_simple,
            }.get(rate_type)
            if not calc_method:
                self._log(f"未知计算类型: {rate_type}", "error")
                return
            result = calc_method(entry)
        except Exception as e:
            self._log(f"计算错误: {e}", "error")
            return
        if result.get("isMainModifier"):
            adjusted = result["adjustedPremium"]
            original = result["originalPremium"]
            self.main_premium = adjusted
            self.main_premium_input.setValue(adjusted)
            self.addon_result_label.setText(
                f"✅ {result['formulaDisplay']}\n"
                f"原主险保费: {fmt_currency(original)} → 调整后: {fmt_currency(adjusted)}"
            )
            self.addon_result_label.setStyleSheet(
                "font-size: 13px; padding: 12px; background: #fef2f2; "
                "border-left: 4px solid #dc2626; border-radius: 8px; color: #991b1b;"
            )
            self.addon_result_label.show()
            self._add_premium_item(entry["clauseName"], 0, result["formulaDisplay"])
            self._log(f"主险保费已调整: {fmt_currency(original)} → {fmt_currency(adjusted)}", "success")
        else:
            premium_text = fmt_currency(result["premium"])
            if result["premium"] < 0:
                premium_text = f"-{fmt_currency(abs(result['premium']))}"
            self.addon_result_label.setText(f"✅ {result['formulaDisplay']}\n保费: {premium_text}")
            bg_color = "#fef2f2" if result["premium"] < 0 else "#ecfdf5"
            fg_color = "#991b1b" if result["premium"] < 0 else "#065f46"
            self.addon_result_label.setStyleSheet(f"font-size: 13px; padding: 12px; background: {bg_color}; border-radius: 8px; color: {fg_color};")
            self.addon_result_label.show()
            self._add_premium_item(entry["clauseName"], result["premium"], result["formulaDisplay"])
            self._log(f"计算完成 [{ADDON_TYPES.get(rate_type, {}).get('label', '')}]: {premium_text}", "success")

    def _calc_simple(self, entry):
        # Check for insuranceTypeRates first
        insurance_type_rates = entry.get("insuranceTypeRates")
        if insurance_type_rates:
            selected_type = self._get_selected_insurance_type()
            if not selected_type:
                raise ValueError("请选择保险类型")
            pct = insurance_type_rates.get(selected_type)
            if pct is None:
                raise ValueError(f"未找到{selected_type}的费率")
            if pct == 0:
                return {
                    "type": "simple_percentage",
                    "premium": 0,
                    "formulaDisplay": f"【{selected_type}】规范类，不涉及保险费的调整"
                }
            # Support ratioMultiplier if present
            ratio_mult = 1.0
            ratio_str = ""
            if entry.get("ratioMultiplier"):
                rm = entry["ratioMultiplier"]
                numerator_key = rm.get("numeratorKey", "")
                numerator = 0
                if hasattr(self, 'custom_input_widgets') and numerator_key in self.custom_input_widgets:
                    numerator = self.custom_input_widgets[numerator_key].value()
                if numerator <= 0:
                    raise ValueError(f"请输入{rm.get('label', '自定义金额')}")
                main_si = getattr(self, 'main_sum_insured', 0)
                if not main_si or main_si <= 0:
                    raise ValueError("请输入主险保险金额")
                ratio_mult = numerator / main_si
                ratio_str = f" × {fmt_currency(numerator)}÷{fmt_currency(main_si)}({ratio_mult*100:.2f}%)"
            rate = pct / 100
            premium = self.main_premium * rate * ratio_mult
            formula_str = f"【{selected_type}】{fmt_currency(self.main_premium)} × {pct}%{ratio_str} = {fmt_currency(premium)}"
            return {"type": "simple_percentage", "premium": premium, "formulaDisplay": formula_str}

        # Original logic
        pct = entry.get("percentage", 0)
        mult = entry.get("multiplier")
        if mult:
            rate = mult
            premium = self.main_premium * rate
            formula_str = f"{fmt_currency(self.main_premium)} × {mult} = {fmt_currency(premium)}"
        else:
            rate = pct / 100
            premium = self.main_premium * rate
            formula_str = f"{fmt_currency(self.main_premium)} × {pct}% = {fmt_currency(premium)}"
        return {"type": "simple_percentage", "premium": premium, "formulaDisplay": formula_str}

    def _calc_modifier_coeff(self, entry):
        """误工费: 调整主险保费系数（直接调整主险保费）"""
        product, details = self._get_coeff_product(entry)
        adjusted = self.main_premium * product
        coeff_str = " × ".join(f"{d['value']:.4f}" for d in details)
        formula_str = (f"调整系数: {coeff_str}\n"
                       f"{fmt_currency(self.main_premium)} × {coeff_str} = {fmt_currency(adjusted)}")
        return {
            "type": "modifier_coeff",
            "premium": 0,
            "adjustedPremium": adjusted,
            "originalPremium": self.main_premium,
            "coefficient": product,
            "isMainModifier": True,
            "formulaDisplay": formula_str,
        }

    def _calc_sudden_death(self, entry):
        """突发疾病身故: 6.6% × (限额差异)"""
        base_pct = entry.get("basePercent", 6.6) / 100
        addon_limit = getattr(self, 'addon_limit_input', None)
        main_limit_w = getattr(self, 'main_limit_display', None)
        addon_val = (addon_limit.value() * 10000) if addon_limit else 0
        main_val = (main_limit_w.value() * 10000) if main_limit_w else 0
        if not main_val:
            if self.full_main_data:
                main_val = self.full_main_data.get("perPersonLimit", 0)
            if not main_val:
                raise ValueError("请输入主险每人限额")
        if not addon_val:
            raise ValueError("请输入附加险每人限额")
        count = self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1
        # 系数
        product = 1.0
        coeff_str = ""
        if entry.get("coefficientTables"):
            product, details = self._get_coeff_product(entry)
            coeff_str = " × " + " × ".join(f"{d['value']:.4f}" for d in details)
        if addon_val < main_val:
            # 减收
            ratio = 1 - addon_val / main_val
            premium = -(main_val * base_pct * count * product * ratio)
            formula_str = (f"减收: {fmt_currency(main_val)} × {entry.get('basePercent', 6.6)}% × {count}人{coeff_str} × "
                           f"(1 - {addon_val}/{main_val}) = {fmt_currency(abs(premium))}")
        elif addon_val > main_val:
            ratio = addon_val / main_val - 1
            premium = main_val * base_pct * count * product * ratio
            formula_str = (f"加收: {fmt_currency(main_val)} × {entry.get('basePercent', 6.6)}% × {count}人{coeff_str} × "
                           f"({addon_val}/{main_val} - 1) = {fmt_currency(premium)}")
        else:
            premium = 0
            formula_str = "附加险限额 = 主险限额，不调整"
        return {"type": "sudden_death", "premium": premium, "formulaDisplay": formula_str}

    def _calc_per_person_rate(self, entry):
        """工伤补充/特定人员: 每人保费 × 费率% × 人数"""
        rate_info = entry.get("rateInfo", {})
        combo = getattr(self, 'injury_insurance_combo', None)
        if combo and combo.currentData() == "with":
            pct = rate_info.get("with_injury_insurance", rate_info.get("default", 0))
        else:
            pct = rate_info.get("without_injury_insurance", rate_info.get("default", 0))
        rate = pct / 100
        count_input = getattr(self, 'addon_count_input', None)
        count = count_input.value() if count_input else (self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1)
        product = 1.0
        coeff_str = ""
        if entry.get("coefficientTables"):
            product, details = self._get_coeff_product(entry)
            coeff_str = " × " + " × ".join(f"{d['value']:.4f}" for d in details)
        premium = self.per_person_premium * rate * count * product
        status = "已购买" if (combo and combo.currentData() == "with") else "未购买"
        formula_str = (f"{fmt_currency(self.per_person_premium)} × {pct}%({status}工伤) × {count}人{coeff_str} = "
                       f"{fmt_currency(premium)}")
        return {"type": "per_person_rate", "premium": premium, "formulaDisplay": formula_str}

    def _calc_per_person_base(self, entry):
        """药品服务: 定额/人 × 系数 × 人数"""
        base_amt = entry.get("baseAmount", 300)
        count = self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1
        product = 1.0
        coeff_str = ""
        if entry.get("coefficientTables"):
            product, details = self._get_coeff_product(entry)
            coeff_str = " × " + " × ".join(f"{d['value']:.4f}" for d in details)
        premium = base_amt * product * count
        formula_str = f"{base_amt}元/人 × {count}人{coeff_str} = {fmt_currency(premium)}"
        return {"type": "per_person_base", "premium": premium, "formulaDisplay": formula_str}

    def _calc_disability_adjust(self, entry):
        """劳务关系人员: 每人保费 × 伤残调整系数 × 人数"""
        coeffs = entry.get("adjustCoeffs", DISABILITY_ADJUST_COEFFS)
        dis_table = "none"
        if self.full_main_data:
            dis_table = self.full_main_data.get("disabilityTable", "none")
        if dis_table == "none" or dis_table not in coeffs:
            raise ValueError("需要主险选择伤残赔偿附表（请先在主险计算Tab选择伤残附表并传入）")
        coeff = coeffs[dis_table]
        count_input = getattr(self, 'addon_count_input', None)
        count = count_input.value() if count_input else (self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1)
        premium = self.per_person_premium * coeff * count
        formula_str = (f"{fmt_currency(self.per_person_premium)} × {coeff}({dis_table}) × {count}人 = "
                       f"{fmt_currency(premium)}")
        return {"type": "disability_adjust", "premium": premium, "formulaDisplay": formula_str}

    def _calc_property_loss(self, entry):
        """员工个人特定财产损失: (基本保险费 + 限额×费率%) × 系数积 × 承保人数"""
        base_premium = entry.get("basePremium", 20)
        base_rate_pct = entry.get("baseRatePercent", 1.5)
        limit_input = getattr(self, 'property_limit_input', None)
        limit_val = limit_input.value() if limit_input else 0
        count_input = getattr(self, 'addon_count_input', None)
        count = count_input.value() if count_input else (self.full_main_data.get("employeeCount", 1) if self.full_main_data else 1)
        if limit_val <= 0:
            raise ValueError("请输入每次事故每人赔偿限额")
        limit_yuan = limit_val * 10000
        rate_part = limit_yuan * base_rate_pct / 100
        per_person = base_premium + rate_part
        product = 1.0
        coeff_str = ""
        if entry.get("coefficientTables"):
            product, details = self._get_coeff_product(entry)
            coeff_str = " × " + " × ".join(f"{d['value']:.4f}" for d in details)
        premium = per_person * product * count
        formula_str = (f"({base_premium}元 + {limit_val}万×{base_rate_pct}%) × 系数{coeff_str.lstrip(' × ')} × {count}人 = "
                       f"({base_premium} + {fmt_currency(rate_part)}) × {product:.4f} × {count} = "
                       f"{fmt_currency(premium)}")
        return {"type": "property_loss", "premium": premium, "formulaDisplay": formula_str}

    def _calc_formula_sum(self, entry):
        """雇主法律责任/一次性伤残: Σ公式"""
        factor = entry.get("baseRateFactor", 1.0)
        if not self.full_main_data:
            raise ValueError("需要主险完整数据（请先在主险计算Tab计算并传入附加险）")
        base_rate = self.full_main_data.get("baseRate", 0)
        per_limit = self.full_main_data.get("perPersonLimit", 0)
        count = self.full_main_data.get("employeeCount", 1)
        main_coeff = self.full_main_data.get("coeffProduct", 1.0)
        if not base_rate:
            raise ValueError("主险基准费率为0，无法计算")
        if not per_limit:
            raise ValueError("主险每人限额为0（仅固定限额计费方式支持此类附加险）")
        # 附加险自身的系数
        addon_product = 1.0
        coeff_str = ""
        if entry.get("coefficientTables"):
            addon_product, details = self._get_coeff_product(entry)
            coeff_str = " × " + " × ".join(f"{d['value']:.4f}" for d in details)
        adjusted_rate = base_rate * factor
        premium = per_limit * adjusted_rate * count * main_coeff * addon_product
        formula_str = (f"{fmt_currency(per_limit)} × {base_rate:.6f} × {factor} × {count}人 × "
                       f"主险系数{main_coeff:.4f}{coeff_str} = {fmt_currency(premium)}")
        return {"type": "formula_sum", "premium": premium, "formulaDisplay": formula_str}

    def _calc_deduction(self, entry):
        """突发疾病除外: 减收主险保费的X%"""
        deduct_pct = entry.get("deductPercent", 5.0)
        premium = -(self.main_premium * deduct_pct / 100)
        formula_str = f"减收: {fmt_currency(self.main_premium)} × {deduct_pct}% = {fmt_currency(abs(premium))}"
        return {"type": "deduction", "premium": premium, "formulaDisplay": formula_str}

    def _calc_table(self, entry):
        """通用系数表计算"""
        # Check for insuranceTypeRates first
        insurance_type_rates = entry.get("insuranceTypeRates")
        start_index = 0
        base_str_prefix = ""

        if insurance_type_rates:
            selected_type = self._get_selected_insurance_type()
            if not selected_type:
                raise ValueError("请选择保险类型")
            pct = insurance_type_rates.get(selected_type)
            if pct is None:
                raise ValueError(f"未找到{selected_type}的费率")
            if pct == 0:
                return {
                    "type": "table_coefficient",
                    "premium": 0,
                    "formulaDisplay": f"【{selected_type}】规范类，不涉及保险费的调整"
                }
            base_premium = self.main_premium * (pct / 100)
            base_str_prefix = f"【{selected_type}】"
            base_str = f"{fmt_currency(self.main_premium)} × {pct}%"
            start_index = 1  # Skip first coefficient table
        else:
            # Original basePremium logic
            base_premium = self.main_premium
            bp = entry.get("basePremium", {})
            if bp.get("multiplier"):
                base_premium = self.main_premium * bp["multiplier"]
                base_str = f"{fmt_currency(self.main_premium)} × {bp['multiplier']}"
            elif bp.get("percentage"):
                base_premium = self.main_premium * (bp["percentage"] / 100)
                base_str = f"{fmt_currency(self.main_premium)} × {bp['percentage']}%"
            else:
                base_str = fmt_currency(self.main_premium)

        # ratioMultiplier 支持
        ratio_mult = 1.0
        ratio_str = ""
        if entry.get("ratioMultiplier"):
            rm = entry["ratioMultiplier"]
            numerator_key = rm.get("numeratorKey", "")
            # Try to find the custom input widget
            numerator = 0
            if hasattr(self, 'custom_input_widgets') and numerator_key in self.custom_input_widgets:
                numerator = self.custom_input_widgets[numerator_key].value()
            if numerator <= 0:
                raise ValueError(f"请输入{rm.get('label', '自定义金额')}")
            main_si = getattr(self, 'main_sum_insured', 0)
            if not main_si or main_si <= 0:
                raise ValueError("请输入主险保险金额")
            ratio_mult = numerator / main_si
            ratio_str = f" × {fmt_currency(numerator)}÷{fmt_currency(main_si)}({ratio_mult*100:.2f}%)"

        product, details = self._get_coeff_product(entry, start_index)
        premium = base_premium * ratio_mult * product
        coeff_str = " × ".join(f"{c['value']:.4f}" for c in details) if details else "1.0000"
        return {"type": "table_coefficient", "premium": premium,
                "formulaDisplay": f"{base_str_prefix}基准 {base_str} = {fmt_currency(base_premium)}{ratio_str} × 系数 ({coeff_str}) = {fmt_currency(premium)}"}

    # ---------- 保费汇总管理 ----------
    def _add_premium_item(self, clause_name, premium, formula):
        existing_idx = next((i for i, item in enumerate(self.premium_items) if item["clauseName"] == clause_name), -1)
        new_item = {"id": id(formula), "clauseName": clause_name, "premium": premium, "formula": formula}
        if existing_idx >= 0:
            self.premium_items[existing_idx] = new_item
            self._log(f"已更新: {clause_name} 的保费")
        else:
            self.premium_items.append(new_item)
        self._render_premium_summary()

    def _remove_premium_item(self, item_id):
        self.premium_items = [item for item in self.premium_items if item["id"] != item_id]
        self._render_premium_summary()
        self._log("已移除一项附加险保费")

    def _render_premium_summary(self):
        while self.premium_list_layout.count():
            item = self.premium_list_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not self.premium_items:
            empty = QLabel("计算附加险保费后\n将自动添加到此列表")
            empty.setAlignment(Qt.AlignCenter)
            empty.setStyleSheet(f"color: {AnthropicColors.TEXT_TERTIARY}; font-size: 12px; padding: 20px;")
            self.premium_list_layout.addWidget(empty)
            self.addon_total_label.setText("附加险合计: ¥0.00")
            self.annual_total_label.setText("保单预估年保费: ¥0.00")
            return

        addon_total = 0.0
        for item in self.premium_items:
            addon_total += item["premium"]
            row_widget = QWidget()
            row_layout = QHBoxLayout(row_widget)
            row_layout.setContentsMargins(8, 4, 8, 4)
            info_layout = QVBoxLayout()
            name_lbl = QLabel(item["clauseName"])
            name_lbl.setStyleSheet(f"font-size: 12px; font-weight: 500;")
            info_layout.addWidget(name_lbl)
            amount_color = "#ef4444" if item["premium"] < 0 else AnthropicColors.ACCENT
            amount_text = f"-{fmt_currency(abs(item['premium']))}" if item["premium"] < 0 else fmt_currency(item["premium"])
            amount_lbl = QLabel(amount_text)
            amount_lbl.setStyleSheet(f"font-size: 11px; color: {amount_color};")
            info_layout.addWidget(amount_lbl)
            row_layout.addLayout(info_layout, 1)
            del_btn = QPushButton("×")
            del_btn.setFixedSize(24, 24)
            del_btn.setCursor(Qt.PointingHandCursor)
            del_btn.setStyleSheet(f"QPushButton {{ background: transparent; color: {AnthropicColors.TEXT_TERTIARY}; border: none; font-size: 16px; }} QPushButton:hover {{ color: #ef4444; }}")
            item_id = item["id"]
            del_btn.clicked.connect(lambda checked, iid=item_id: self._remove_premium_item(iid))
            row_layout.addWidget(del_btn)
            row_widget.setStyleSheet(f"background: {AnthropicColors.BG_CARD}; border-radius: 6px;")
            self.premium_list_layout.addWidget(row_widget)

        addon_total_text = f"-{fmt_currency(abs(addon_total))}" if addon_total < 0 else fmt_currency(addon_total)
        self.addon_total_label.setText(f"附加险合计: {addon_total_text}")
        main_val = self.main_premium_input.value()
        annual_total = main_val + addon_total
        annual_text = f"-{fmt_currency(abs(annual_total))}" if annual_total < 0 else fmt_currency(annual_total)
        self.annual_total_label.setText(f"保单预估年保费: {annual_text}")
        self._calc_short_term()

    # ---------- 短期保费计算 ----------
    def _calc_short_term(self):
        start = self.start_date.date()
        end = self.end_date.date()
        if end <= start:
            self.short_term_result.setText("终止日须晚于起保日")
            self.short_term_result.setStyleSheet(f"color: #ef4444; font-size: 11px;")
            return
        insurance_days = start.daysTo(end)
        start_year = start.year()
        year_days = 366 if is_leap_year(start_year) else 365
        main_val = self.main_premium_input.value()
        addon_total = sum(item["premium"] for item in self.premium_items)
        annual_total = main_val + addon_total
        short_premium = annual_total / year_days * insurance_days
        leap_text = f"（闰年 {year_days}天）" if year_days == 366 else f"（平年 {year_days}天）"
        short_text = f"-{fmt_currency(abs(short_premium))}" if short_premium < 0 else fmt_currency(short_premium)
        self.short_term_result.setText(
            f"保险天数: {insurance_days} 天 · {start_year}年{leap_text}\n"
            f"{fmt_currency(abs(annual_total))} ÷ {year_days} × {insurance_days}\n"
            f"短期保费: {short_text}")
        self.short_term_result.setStyleSheet(f"font-size: 12px; color: {AnthropicColors.ACCENT}; font-weight: 600;")

    # ---------- 核保经验计费 ----------
    def _add_manual_premium(self):
        if not self.selected_entry:
            self._log("请先选择条款", "warn")
            return
        manual_val = self.manual_input.value()
        self._add_premium_item(self.selected_entry["clauseName"], manual_val, f"核保经验计费: {fmt_currency(manual_val)}")
        self._log(f"核保经验计费: {self.selected_entry['clauseName']} → {fmt_currency(manual_val)}", "success")

    # ---------- 询价导入 ----------
    def _handle_inquiry_import(self):
        if not self.rate_data or not self.rate_data.get("entries"):
            self._log("请先加载费率方案数据", "warn")
            return
        file_path, _ = QFileDialog.getOpenFileName(self, "导入询价文件", "", "询价文件 (*.xlsx *.docx);;Excel (*.xlsx);;Word (*.docx)")
        if not file_path:
            return
        if file_path.endswith(".xlsx"):
            self._parse_inquiry_excel(file_path)
        elif file_path.endswith(".docx"):
            self._parse_inquiry_docx(file_path)

    def _parse_inquiry_excel(self, file_path):
        try:
            from openpyxl import load_workbook
        except ImportError:
            self._log("openpyxl 未安装，请运行: pip install openpyxl", "error")
            return
        try:
            wb = load_workbook(file_path, read_only=True)
            ws = wb.active
            clause_names = []
            for row in ws.iter_rows():
                if len(row) > 5 and row[5].value:
                    val = str(row[5].value).strip()
                    if len(val) > 2 and val not in ("附加条款", "条款名称") and "附加" in val:
                        clause_names.append(val)
            if not clause_names:
                for row in ws.iter_rows():
                    for cell in row:
                        val = str(cell.value or "").strip()
                        if len(val) > 4 and "附加" in val and val not in clause_names:
                            clause_names.append(val)
            wb.close()
            self._match_inquiry_clauses(clause_names, os.path.basename(file_path))
        except Exception as e:
            self._log(f"Excel 解析失败: {e}", "error")

    def _parse_inquiry_docx(self, file_path):
        try:
            from docx import Document
        except ImportError:
            self._log("python-docx 未安装", "error")
            return
        try:
            doc = Document(file_path)
            clause_names = []
            for p in doc.paragraphs:
                for run in p.runs:
                    if run.font.color and run.font.color.rgb:
                        rgb = str(run.font.color.rgb)
                        if rgb.startswith("0000") or rgb.lower() in ("0000ff", "0000cd", "0000ee"):
                            text = run.text.strip()
                            if len(text) > 2:
                                clause_names.append(text)
            if not clause_names:
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if "附加" in text and 4 < len(text) < 60:
                        clause_names.append(text)
            self._match_inquiry_clauses(clause_names, os.path.basename(file_path))
        except Exception as e:
            self._log(f"Docx 解析失败: {e}", "error")

    def _match_clause_name(self, imported_name, entries):
        normalized = imported_name.replace(" ", "").replace("（", "(").replace("）", ")")
        for e in entries:
            entry_norm = e["clauseName"].replace(" ", "").replace("（", "(").replace("）", ")")
            if entry_norm == normalized:
                return e
        for e in entries:
            entry_norm = e["clauseName"].replace(" ", "").replace("（", "(").replace("）", ")")
            if normalized in entry_norm or entry_norm in normalized:
                return e
        core = re.sub(r"^附加", "", normalized)
        core = re.sub(r"条款$|扩展$", "", core)
        if len(core) < 3:
            return None
        for e in entries:
            entry_core = re.sub(r"^附加", "", e["clauseName"].replace(" ", ""))
            entry_core = re.sub(r"条款$|扩展$", "", entry_core)
            if core in entry_core or entry_core in core:
                return e
        return None

    def _match_inquiry_clauses(self, clause_names, file_name):
        if not clause_names:
            self.inquiry_status.setText("未识别到条款名称")
            return
        entries = self.rate_data["entries"]
        matched = []
        unmatched = []
        seen = set()
        for name in clause_names:
            entry = self._match_clause_name(name, entries)
            if entry and entry["sourceFile"] not in seen:
                matched.append({"importedName": name, "entry": entry})
                seen.add(entry["sourceFile"])
            elif not entry:
                unmatched.append(name)
        self.inquiry_status.setText(f"{file_name} → 识别 {len(clause_names)} 条，匹配 {len(matched)} 条")
        if unmatched:
            self._log(f"未匹配: {', '.join(unmatched[:5])}", "warn")
        if matched:
            self.batch_calc_btn.show()
            self.batch_calc_btn.setText(f"⚡ 一键计算全部（{len(matched)} 条）")
            self._batch_matched = matched
        else:
            self.batch_calc_btn.hide()
        self._log(f"导入 {file_name}：识别 {len(clause_names)} 条，匹配 {len(matched)} 条")

    def _batch_calculate(self):
        matched = getattr(self, '_batch_matched', [])
        if not matched:
            self._log("无匹配条款可计算", "warn")
            return
        original_premium = self.main_premium_input.value()
        if original_premium <= 0:
            self._log("请先输入主险保费", "warn")
            return
        self.main_premium = original_premium
        calc_count = 0
        skip_count = 0
        modifier_count = 0

        # 第一轮：优先计算主险保费调整类型（main_premium_modifier + modifier_coeff）
        for item in matched:
            entry = item["entry"]
            rt = entry.get("rateType", "")
            if rt == "modifier_coeff":
                # modifier_coeff 需要系数表选择
                if not self.coeff_selections:
                    self._add_premium_item(entry["clauseName"], 0,
                                           "⚠️ 需手动计算 [主险系数调整] — 请单选此条款后选择系数")
                    skip_count += 1
                    continue
                try:
                    result = self._calc_modifier_coeff(entry)
                    self.main_premium = result["adjustedPremium"]
                    self._add_premium_item(entry["clauseName"], 0, result["formulaDisplay"])
                    modifier_count += 1
                    calc_count += 1
                except Exception as e:
                    self._add_premium_item(entry["clauseName"], 0, f"⚠️ 计算失败: {e}")
                    skip_count += 1
                continue
            if rt != "main_premium_modifier":
                continue
            # water_level_deduction / regional_deduction: 需手动参数
            if entry.get("modifierType") == "water_level_deduction":
                self._add_premium_item(entry["clauseName"], 0,
                                       "⚠️ 需手动计算 [水位线减费] — 请单选此条款后输入水位高度")
                skip_count += 1
                continue
            if entry.get("modifierType") == "regional_deduction":
                self._add_premium_item(entry["clauseName"], 0,
                                       "⚠️ 需手动计算 [地区减费] — 请单选此条款后选择地区")
                skip_count += 1
                continue
            combo = getattr(self, 'modifier_insurance_combo', None)
            ratio_input = getattr(self, 'modifier_ratio_input', None)
            col_key = combo.currentData() if combo else None
            ratio_val = ratio_input.value() if ratio_input else 0
            if not col_key or ratio_val <= 0:
                self._add_premium_item(entry["clauseName"], 0,
                                       "⚠️ 需手动计算 [主险保费调整] — 请单选此条款后设置参数")
                skip_count += 1
                continue
            try:
                result = self._calc_main_premium_modifier(entry)
                self.main_premium = result["adjustedPremium"]
                self._add_premium_item(entry["clauseName"], 0, result["formulaDisplay"])
                modifier_count += 1
                calc_count += 1
            except Exception as e:
                self._add_premium_item(entry["clauseName"], 0, f"⚠️ 计算失败: {e}")
                skip_count += 1

        if modifier_count > 0:
            self.main_premium_input.setValue(self.main_premium)
            self._log(f"主险保费已调整: {fmt_currency(original_premium)} → "
                      f"{fmt_currency(self.main_premium)}（{modifier_count} 条调整条款）", "info")

        # 第二轮：用调整后的主险保费计算其余附加险
        for item in matched:
            entry = item["entry"]
            rt = entry.get("rateType", "")
            if rt in ("main_premium_modifier", "modifier_coeff"):
                continue
            if rt in ("regulatory", "no_calc", "included_in_main", "daily_prorate", "formula_conditional"):
                skip_count += 1
                continue
            if rt == "conditional_simple":
                self._add_premium_item(entry["clauseName"], 0,
                                       "条件勾选类 — 默认主险已包含，不加收（需手动勾选后单独计算）")
                skip_count += 1
                continue
            if rt == "base_rate_formula" and entry.get("maxInstances"):
                self._add_premium_item(entry["clauseName"], 0,
                                       "⚠️ 需手动计算 [多服务实例] — 请单选此条款后配置各服务实例")
                skip_count += 1
                continue
            batch_calc_fn = {
                "simple_percentage": self._calc_simple,
                "deduction": self._calc_deduction,
                "table_coefficient": self._calc_table,
                "base_rate_formula": self._calc_base_rate_formula,
            }.get(rt)
            if batch_calc_fn:
                try:
                    result = batch_calc_fn(entry)
                    self._add_premium_item(entry["clauseName"], result["premium"], result["formulaDisplay"])
                    calc_count += 1
                except Exception as e:
                    label = ADDON_TYPES.get(rt, {}).get("label", rt)
                    self._add_premium_item(entry["clauseName"], 0, f"⚠️ 需手动计算 [{label}]: {e}")
                    skip_count += 1
            else:
                self._add_premium_item(entry["clauseName"], 0,
                                       f"需手动计算 [{ADDON_TYPES.get(rt, {}).get('label', rt)}]")
                skip_count += 1
        msg = f"批量计算完成: {calc_count} 条已计算, {skip_count} 条需手动处理"
        if modifier_count > 0:
            msg += f" (含 {modifier_count} 条主险调整)"
        self._log(msg, "success")
