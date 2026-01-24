# -*- coding: utf-8 -*-
"""
货运保险工具 ProMax - Windows 版本
Author: Dachi_Yijin
"""
import sys
import os
import platform
import subprocess
import re
import copy
import math
import ast
import operator
import logging
from datetime import datetime, timedelta

# Windows 控制台编码修复（仅当有控制台时）
if platform.system() == 'Windows':
    import io
    # GUI 应用（--windowed）没有控制台，stdout/stderr 可能是 None
    if sys.stdout is not None and hasattr(sys.stdout, 'buffer'):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    if sys.stderr is not None and hasattr(sys.stderr, 'buffer'):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PyPDF2 import PdfMerger  # 确保已安装: pip install PyPDF2
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QProgressBar, QCheckBox,
    QTextEdit, QMessageBox, QGroupBox, QFrame, QTabWidget,
    QListWidget, QListWidgetItem, QAbstractItemView, QRadioButton, QButtonGroup,
    QComboBox, QDateEdit, QScrollArea, QSizePolicy, QGraphicsDropShadowEffect,
    QDialog, QMenu
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QDate, QTimer
from PyQt6.QtGui import QFont, QPainter, QPen, QColor

# 导入配置管理器
from customer_config import get_config_manager

# 抑制 Qt 相关的系统日志消息
os.environ['QT_LOGGING_RULES'] = '*.debug=false;qt.qpa.*=false'

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ==========================================
# 安全表达式计算器（替代 eval）
# ==========================================
_SAFE_OPERATORS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.USub: operator.neg,
    ast.UAdd: operator.pos,
}


def _safe_eval_node(node):
    """递归计算 AST 节点"""
    if isinstance(node, ast.Constant):
        if isinstance(node.value, (int, float)):
            return node.value
        raise ValueError(f"不支持的常量类型: {type(node.value)}")
    elif isinstance(node, ast.BinOp):
        left = _safe_eval_node(node.left)
        right = _safe_eval_node(node.right)
        op_func = _SAFE_OPERATORS.get(type(node.op))
        if op_func is None:
            raise ValueError(f"不支持的运算符: {type(node.op)}")
        return op_func(left, right)
    elif isinstance(node, ast.UnaryOp):
        operand = _safe_eval_node(node.operand)
        op_func = _SAFE_OPERATORS.get(type(node.op))
        if op_func is None:
            raise ValueError(f"不支持的一元运算符: {type(node.op)}")
        return op_func(operand)
    elif isinstance(node, ast.Num):  # Python 3.7 兼容
        return node.n
    else:
        raise ValueError(f"不支持的节点类型: {type(node)}")


def safe_eval_expr(expr: str):
    """
    安全地计算数学表达式（替代 eval）
    只支持: 数字、+、-、*、/、括号
    """
    try:
        tree = ast.parse(expr, mode='eval')
        return _safe_eval_node(tree.body)
    except (ValueError, SyntaxError, TypeError, ZeroDivisionError):
        return None


APP_NAME = '货运保险工具 ProMax'
APP_AUTHOR = 'Dachi_Yijin'
APP_VERSION = 'ProMax'

# ASCII Art Logo
APP_LOGO = """
╔═══════════════════════════════════════════════════════════════════╗
║          ██████╗  ██████╗██╗   ██╗     ██╗██╗███╗   ██╗           ║
║          ██╔══██╗██╔════╝╚██╗ ██╔╝     ██║██║████╗  ██║           ║
║          ██║  ██║██║      ╚████╔╝      ██║██║██╔██╗ ██║           ║
║          ██║  ██║██║       ╚██╔╝  ██   ██║██║██║╚██╗██║           ║
║          ██████╔╝╚██████╗   ██║   ╚█████╔╝██║██║ ╚████║           ║
║          ╚═════╝  ╚═════╝   ╚═╝    ╚════╝ ╚═╝╚═╝  ╚═══╝           ║
║                 🚀 货运保险工具 ProMax 🚀                         ║
║                     Author: Dachi_Yijin                           ║
╚═══════════════════════════════════════════════════════════════════╝
"""
# 打印Logo
print(APP_LOGO)

# ==========================================
# 协议编号映射（恒力能源销售）
# ==========================================
AGREEMENT_CODES = {
    '恒力能源（苏州）有限公司': 'CSHHHYX2025Q000337',
    '苏州恒力精细化工销售有限公司': 'CSHHHYX2025Q000360',
    '恒力石化销售有限公司': 'CSHHHYX2025Q000356',
    '恒力油品销售（苏州）有限公司': 'CSHHHYX2025Q000361',
    '恒力华南石化销售有限公司': 'CSHHHYX2025Q000358',
}

# PDF导出分组规则（恒力能源销售）
PDF_EXPORT_GROUPS = {
    '能源苏州': lambda name: '能源苏州' in name,
    '华南石化': lambda name: '华南石化' in name,
    '其他业务': lambda name: any(k in name for k in ['精细化工', '恒力石化', '油品销售']),
}

# 营业执照映射（投保人 -> 营业执照号）
LICENSE_MAP = {
    '恒力能源（苏州）有限公司': '91320594MA20E5BE5B',
    '苏州恒力精细化工销售有限公司': '91320509MAC49YP38G',
    '恒力石化销售有限公司': '91310120MA1HQKU28C',
    '恒力油品销售（苏州）有限公司': '91320594MA20TM6J2A',
    '恒力华南石化销售有限公司': '91440300MA5FL0K791',
    '恒力石化（大连）有限公司': '91210244550622058M',
    '恒力石化（惠州）有限公司': '91441300MA556RMB75',
    '浙江卓航多式联运科技有限公司': '91330109MAC4A84UXN',
    '康辉大连新材料科技有限公司': '91210244MA10YYBP1T',
    '康辉新材料科技有限公司': '91210800580717031A',
    '康辉国际贸易（江苏）有限公司': '91320509061869594L',
}


# ==========================================
# 工具函数
# ==========================================

def cn_currency(value):
    """人民币数字转大写汉字"""
    if isinstance(value, str):
        try:
            value = float(value.replace(',', ''))
        except (ValueError, TypeError):
            return value
            
    n = round(value, 2)
    integer_part = int(n)
    fraction_part = int(round((n - integer_part) * 100))
    
    canvas = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    unit = ['元', '拾', '佰', '仟', '万', '拾', '佰', '仟', '亿', '拾', '佰', '仟']
    
    s_int = str(integer_part)
    result = ""
    if integer_part == 0:
        result = "零"
    else:
        s_int = s_int[::-1]
        for i, digit in enumerate(s_int):
            num = int(digit)
            result = canvas[num] + unit[i] + result
            
    result = re.sub(r'零[拾佰仟]', '零', result)
    result = re.sub(r'零万', '万', result)
    result = re.sub(r'零亿', '亿', result)
    result = re.sub(r'亿万', '亿', result)
    result = re.sub(r'零+', '零', result)
    result = re.sub(r'零元', '元', result)
    if result.endswith('零') and len(result) > 1:
        result = result[:-1]
    
    if result == '元': 
        result = '零元'
        
    if integer_part == 0 and fraction_part == 0:
        return '零元整'
        
    if fraction_part == 0:
        result += "整"
    else:
        jiao = fraction_part // 10
        fen = fraction_part % 10
        if jiao > 0:
            result += canvas[jiao] + "角"
        elif integer_part > 0 and fen > 0:
            result += "零"
            
        if fen > 0:
            result += canvas[fen] + "分"
            
    return result

def calc_text_width(text):
    if not text:
        return 0
    return sum(1.8 if '\u4e00' <= c <= '\u9fff' else 1 for c in str(text))

def auto_fit_column_width(ws, col_idx, min_width=6, max_width=50, sample_rows=100):
    col_letter = get_column_letter(col_idx)
    if col_idx == 1:
        return 8
    
    header = None
    for r in [1, 2]:
        val = ws.cell(r, col_idx).value
        if val and any(kw in str(val) for kw in ['序', '货', '船', '日期', '保', '费', '金额', '吨', '备注', '申报']):
            header = str(val).replace('\n', '')
            break
    
    if header:
        if '序号' in header or '序列' in header: return 4
        if '共同被保险人' in header or '货主' in header or '申报公司' in header: return 29.03
        if '货值' in header: return 7.5
        if any(kw in header for kw in ['日期', '船期', '起运', '报险', '报预', '报正式']): return 12
        if '货种' in header or '保险货物' in header: return 15  # 默认宽度，实际宽度由format_sheet根据sheet位置设置
        if '船名' in header: return 14
        if '流向' in header: return 18
        if '金额' in header: return 14
        if '吨位' in header or '数量' in header: return 12
        if '费率' in header: return 10
        if '保费' in header: return 12
        if '包袋' in header or '件' in header: return 10
        if '备注' in header: return 20
    
    max_len = 0
    for i, cell in enumerate(ws[col_letter]):
        if i >= sample_rows: break
        if cell.value is None: continue
        lines = str(cell.value).split('\n')
        for line in lines:
            width = calc_text_width(line)
            max_len = max(max_len, width)
    return min(max(max_len * 1.1 + 2, min_width), max_width)

def auto_fit_row_height(ws, row_idx, base_height=15, font_size=10, header_row=None):
    max_lines = 1
    for cell in ws[row_idx]:
        if cell.value:
            val = str(cell.value)
            # 检查显式换行符
            if '\n' in val:
                max_lines = max(max_lines, val.count('\n') + 1)
            else:
                # 对于长文本，根据文本长度和列宽估算行数
                text_width = calc_text_width(val)
                col_letter = get_column_letter(cell.column)
                # 获取列宽：优先使用已设置的列宽，否则使用预期列宽
                col_width = ws.column_dimensions[col_letter].width
                if not col_width or col_width < 5:
                    # 如果列宽未设置，使用预期列宽
                    col_width = auto_fit_column_width(ws, cell.column, min_width=8, max_width=40)
                # 只有当文本宽度超过列宽时才估算换行
                if text_width > col_width * 1.2:
                    chars_per_line = max(col_width * 1.2, 8)  # 估算每行字符数
                    estimated_lines = max(1, int(text_width / chars_per_line) + 1)
                    max_lines = max(max_lines, estimated_lines)
    return max(base_height, max_lines * (font_size + 4))

def show_completion_dialog(parent, title, message, output_dir):
    msg_box = QMessageBox(parent)
    msg_box.setWindowTitle(title)
    msg_box.setText(message)
    msg_box.setIcon(QMessageBox.Icon.Information)
    btn_open = msg_box.addButton("📂 打开文件夹", QMessageBox.ButtonRole.ActionRole)
    btn_close = msg_box.addButton("关闭", QMessageBox.ButtonRole.RejectRole)
    msg_box.exec()
    if msg_box.clickedButton() == btn_open:
        # 跨平台打开文件夹
        if platform.system() == 'Windows':
            os.startfile(output_dir)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.run(['open', output_dir])
        else:  # Linux
            subprocess.run(['xdg-open', output_dir])

def find_header_row(ws, max_search=10, header_keywords=None):
    """查找表头行（包含指定关键词的行）

    Args:
        ws: 工作表对象
        max_search: 最大搜索行数
        header_keywords: 表头关键词列表，如 ["序号", "编号"]，默认为 ["序号", "序列"]
    """
    if header_keywords is None:
        header_keywords = ["序号", "序列"]

    for r in range(1, max_search + 1):
        for c in range(1, min(20, ws.max_column + 1)):
            val = ws.cell(r, c).value
            if val:
                val_str = str(val)
                for kw in header_keywords:
                    if kw in val_str:
                        return r
    return 2

def find_total_row(ws, header_row, total_keywords=None):
    """查找合计行

    Args:
        ws: 工作表对象
        header_row: 表头行号
        total_keywords: 合计行关键词列表，如 ["合计", "总计"]，默认为 ["合计"]
    """
    if total_keywords is None:
        total_keywords = ["合计"]

    for r in range(header_row + 1, ws.max_row + 1):
        val = ws.cell(r, 1).value
        if val:
            val_str = str(val)
            for kw in total_keywords:
                if kw in val_str:
                    return r
    return None

def find_column_indices(ws, header_row):
    columns = {'cargo_type': None, 'tonnage': None, 'insurance_amount': None, 'premium': None, 'rate': None}
    keywords = {
        'cargo_type': ['货种', '保险货物'],
        'tonnage': ['实载吨位', '实载', '吨位', '数量（吨）', '数量'],
        'insurance_amount': ['保险金额', '货物金额'],
        'premium': ['保费'],
        'rate': ['费率'],
    }
    for c in range(1, ws.max_column + 1):
        val = ws.cell(header_row, c).value
        if not val: continue
        val_clean = str(val).replace('\n', '').strip()
        for key, kw_list in keywords.items():
            if columns[key] is None:
                for kw in kw_list:
                    if kw in val_clean:
                        columns[key] = c
                        break
    return columns


def find_column_by_keywords(col_map, *keywords):
    """
    在列映射字典中根据关键字查找列号

    Args:
        col_map: 列名到列号的字典映射 {header_name: column_index}
        *keywords: 要查找的关键字（按优先级排序）

    Returns:
        匹配的列号，如果没找到返回 None
    """
    for kw in keywords:
        for h, c in col_map.items():
            if kw in h:
                return c
    return None


def build_column_map(ws, header_row):
    """
    构建工作表的列名到列号映射

    Args:
        ws: 工作表对象
        header_row: 表头行号

    Returns:
        字典 {列名: 列号}
    """
    col_map = {}
    for c in range(1, ws.max_column + 1):
        h_val = str(ws.cell(header_row, c).value or '').replace('\n', '').strip()
        if h_val:
            col_map[h_val] = c
    return col_map


# ==========================================
# 恒力能源销售专用工具函数
# ==========================================

def clean_material_name(name):
    """清理物料名称，去除后缀"""
    if not name:
        return ""
    name_str = str(name).strip()
    if "共聚甲醛树脂" in name_str:
        return "共聚甲醛树脂"
    if name_str == "工业用双酚A":
        return name_str
    matches = list(re.finditer(r'[\u4e00-\u9fa5]', name_str))
    if matches:
        return name_str[:matches[-1].end()]
    return name_str


def _parse_date_value(value):
    """
    内部函数：解析日期值为datetime对象
    支持格式：datetime对象、YYYY/MM/DD、YYYY-MM-DD、YYYYMMDD、YYYY年MM月DD日
    返回：(datetime对象或None, 标准化的日期字符串或原始字符串)
    """
    if not value:
        return None, ""
    if isinstance(value, datetime):
        return value, value.strftime("%Y/%m/%d")
    try:
        s_val = str(value).strip()
        if " " in s_val:
            s_val = s_val.split(" ")[0]
        s_val = s_val.replace("-", "/").replace("年", "/").replace("月", "/").replace("日", "")
        for fmt in ["%Y/%m/%d", "%Y-%m-%d", "%Y%m%d"]:
            try:
                dt = datetime.strptime(s_val, fmt)
                return dt, dt.strftime("%Y/%m/%d")
            except ValueError:
                continue
        return None, s_val
    except Exception:
        return None, str(value) if value else ""


def format_date_slashes(value):
    """格式化日期为 YYYY/MM/DD 格式"""
    _, formatted = _parse_date_value(value)
    return formatted


def parse_date_for_compare(date_val):
    """将日期值转换为可比较的datetime对象"""
    dt, _ = _parse_date_value(date_val)
    return dt


def safe_float(val):
    """安全转换为浮点数"""
    try:
        return float(val)
    except (ValueError, TypeError):
        return 0.0


def find_header_row_energy(ws, max_search=6):
    """查找恒力能源销售的表头行（包含"申报公司名称"或"车船号"）"""
    for r in range(1, max_search):
        vals = [str(c.value).strip() if c.value else "" for c in ws[r]]
        if "申报公司名称" in vals or "车船号" in vals:
            col_map = {}
            for cell in ws[r]:
                if cell.value:
                    key = str(cell.value).strip().replace("（", "(").replace("）", ")")
                    col_map[key] = cell.col_idx
            return r, col_map
    return None, {}


def find_total_row_energy(ws, header_row, max_col=10):
    """查找恒力能源销售的合计行"""
    for r in range(header_row + 1, ws.max_row + 2):
        for c in range(1, max_col):
            val = ws.cell(row=r, column=c).value
            if val and "合计" in str(val):
                return r, c
    return None, 1


def get_column_value(ws, row, col_map, keys):
    """从列映射中获取指定行的值"""
    for key in keys:
        if key in col_map:
            return safe_float(ws.cell(row=row, column=col_map[key]).value)
    return 0.0


def extract_hengli_energy_data(ws, sheet_name):
    """提取恒力能源销售表格数据"""
    header_row, col_map = find_header_row_energy(ws)
    if not header_row:
        return None

    total_row, _ = find_total_row_energy(ws, header_row)

    # 遍历所有数据行，找到最早和最晚发货日期，以及业务笔数
    qty_keys = ["开单量", "数量"]
    earliest_date = None
    earliest_row = None
    latest_date = None
    business_count = 0

    end_loop = total_row if total_row else ws.max_row + 1
    for r in range(header_row + 1, end_loop):
        qty = get_column_value(ws, r, col_map, qty_keys)
        if qty > 0:
            # 获取该行发货日期
            date_val = ws.cell(r, col_map.get("发货日期", 1)).value if "发货日期" in col_map else None
            parsed_date = parse_date_for_compare(date_val)

            if parsed_date:
                # 找最早日期
                if earliest_date is None or parsed_date < earliest_date:
                    earliest_date = parsed_date
                    earliest_row = r
                # 找最晚日期
                if latest_date is None or parsed_date > latest_date:
                    latest_date = parsed_date

    # 业务笔数 = 合计行上一行A列的序号数字
    if total_row and total_row > header_row + 1:
        seq_val = ws.cell(row=total_row - 1, column=1).value
        try:
            business_count = int(seq_val) if seq_val else 0
        except (ValueError, TypeError):
            business_count = 0

    # 从最早日期行获取信息
    info = {"comp": "", "no": "", "date": "", "mat": ""}
    if earliest_row:
        info = {
            "comp": ws.cell(earliest_row, col_map.get("申报公司名称", 1)).value if "申报公司名称" in col_map else "",
            "no": ws.cell(earliest_row, col_map.get("车船号", 1)).value if "车船号" in col_map else "",
            "date": ws.cell(earliest_row, col_map.get("发货日期", 1)).value if "发货日期" in col_map else "",
            "mat": ws.cell(earliest_row, col_map.get("物料名称", 1)).value if "物料名称" in col_map else ""
        }
    else:
        # 回退：如果没有找到有效日期行，使用第一条有效数据行
        for r in range(header_row + 1, end_loop):
            qty = get_column_value(ws, r, col_map, qty_keys)
            if qty > 0:
                info = {
                    "comp": ws.cell(r, col_map.get("申报公司名称", 1)).value if "申报公司名称" in col_map else "",
                    "no": ws.cell(r, col_map.get("车船号", 1)).value if "车船号" in col_map else "",
                    "date": ws.cell(r, col_map.get("发货日期", 1)).value if "发货日期" in col_map else "",
                    "mat": ws.cell(r, col_map.get("物料名称", 1)).value if "物料名称" in col_map else ""
                }
                break

    amt_keys = ["金额(元)", "金额（元）", "金额"]
    prem_keys = ["保费(元)", "保费（元）", "保费"]

    if total_row:
        final_qty = get_column_value(ws, total_row, col_map, qty_keys)
        final_amt = get_column_value(ws, total_row, col_map, amt_keys)
        final_prem = get_column_value(ws, total_row, col_map, prem_keys)
    else:
        final_qty = final_amt = final_prem = 0

    # 如果合计行没有数据，遍历累加
    if final_amt == 0 and final_prem == 0:
        for r in range(header_row + 1, end_loop):
            final_qty += get_column_value(ws, r, col_map, qty_keys)
            final_amt += get_column_value(ws, r, col_map, amt_keys)
            final_prem += get_column_value(ws, r, col_map, prem_keys)

    return {
        'sheet_name': sheet_name,
        'comp': info["comp"],
        'no': info["no"],
        'date': info["date"],
        'mat': clean_material_name(info["mat"]),
        'amt': final_qty,
        'money': final_amt,
        'prem': final_prem,
        'header_row': header_row,
        'total_row': total_row,
        'col_map': col_map,
        'business_count': business_count,
        'latest_date': format_date_slashes(latest_date) if latest_date else "",
    }


def process_hengli_energy_sheet(ws, data, policy_label="投保单号"):
    """处理恒力能源销售表格：格式化、添加投保单号行"""
    header_row = data['header_row']
    total_row = data['total_row']
    col_map = data.get('col_map', {})

    # 获取表头行的填充色
    header_fills = {
        cell.column: copy.copy(cell.fill)
        for cell in ws[header_row]
        if cell.fill and cell.fill.fill_type
    }

    # 找到保费列的索引（用于去除粉色填充）
    prem_col = None
    for k, v in col_map.items():
        if "保费" in k:
            prem_col = v
            break

    # 找到物料名称列的索引
    mat_col = col_map.get("物料名称")

    # 格式化数据行（第3行及以下所有行高设为18）
    end_row = (total_row + 1) if total_row else ws.max_row
    for r in range(header_row + 1, end_row + 2):
        ws.row_dimensions[r].height = 18
        for cell in ws[r]:
            if r == total_row:
                fill = header_fills.get(cell.column, PatternFill(fill_type=None))
            else:
                fill = PatternFill(fill_type=None)

            # 应用样式
            font_size = 12
            # 检查物料名称是否含'丙烯腈-丁二烯-苯乙烯(ABS)树脂'
            if mat_col and cell.column == mat_col:
                mat_val = str(cell.value) if cell.value else ""
                if "丙烯腈-丁二烯-苯乙烯(ABS)树脂" in mat_val or "丙烯腈-丁二烯-苯乙烯（ABS）树脂" in mat_val:
                    font_size = 10
            cell.font = Font(name='Times New Roman', size=font_size)
            cell.alignment = Alignment(
                horizontal=cell.alignment.horizontal,
                vertical='center',
                wrap_text=False
            )
            if fill.fill_type:
                cell.fill = fill

    # 去除保费列中粉色填充的单元格（通常在合计行往上一行）
    if prem_col and total_row:
        for r in range(header_row + 1, total_row):
            cell = ws.cell(row=r, column=prem_col)
            if cell.fill and cell.fill.fill_type:
                # 检查是否为粉色填充（fgColor 为粉色系）
                fg = cell.fill.fgColor
                if fg and fg.rgb:
                    rgb_str = str(fg.rgb).upper()
                    # 粉色系通常RGB值：红色分量高，蓝色和绿色分量相对较低
                    # 常见粉色：FFFF00FF, FFFF99CC, FFFF69B4 等
                    # 简单判断：如果是有填充色，就清除
                    cell.fill = PatternFill(fill_type=None)

    # 格式化表头行
    for cell in ws[header_row]:
        if cell.value:
            val = str(cell.value)
            if '单价' in val and ('元/吨' in val or '元／吨' in val):
                new_val = val.replace('(', '\n(').replace('（', '\n（')
                cell.value = new_val
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[header_row].height = 32

    # 第2行行高改为35，字号改为12
    ws.row_dimensions[2].height = 35
    for cell in ws[2]:
        if cell.value:
            cell.font = Font(name='Times New Roman', size=12)
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # 添加投保单号行
    if total_row:
        total_col = 1
        for c in range(1, 10):
            if ws.cell(row=total_row, column=c).value and "合计" in str(ws.cell(row=total_row, column=c).value):
                total_col = c
                break

        target_cell = ws.cell(row=total_row + 1, column=total_col)
        target_cell.value = policy_label
        target_cell.font = Font(name='Songti SC', size=12)
        target_cell.alignment = Alignment(horizontal='right', vertical='center')

        # 投保单号右边3个单元格合并后左对齐
        policy_row = total_row + 1
        merge_start_col = total_col + 1
        merge_end_col = total_col + 3
        from openpyxl.utils import get_column_letter
        merge_start_letter = get_column_letter(merge_start_col)
        merge_end_letter = get_column_letter(merge_end_col)
        try:
            ws.unmerge_cells(f'{merge_start_letter}{policy_row}:{merge_end_letter}{policy_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        ws.merge_cells(f'{merge_start_letter}{policy_row}:{merge_end_letter}{policy_row}')
        ws.cell(policy_row, merge_start_col).value = ''
        ws.cell(policy_row, merge_start_col).alignment = Alignment(horizontal='left', vertical='center')

    # 设置列宽 - 根据用户要求更新
    FIXED_WIDTHS = {
        "企业交货单": 21, "申报公司名称": 42, "发货日期": 18,
        "序号": 6, "车船号": 14, "航次号": 10, "开单量": 13,
        "数量": 13, "单价": 12, "物料名称": 38,
        # 用户指定的列宽
        "到站(港)": 15.6, "到站（港）": 15.6,
        "金额(元)": 22, "金额（元）": 22, "金额": 22,
        "费率": 10,
        "保费(元)": 15.25, "保费（元）": 15.25, "保费": 15.25,
        "箱号": 17.25,
    }

    for col_idx in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_idx)
        col_name = ""
        for k, v in col_map.items():
            if v == col_idx:
                col_name = k
                break

        fixed_width = None
        # 精确匹配优先
        if col_name in FIXED_WIDTHS:
            fixed_width = FIXED_WIDTHS[col_name]
        else:
            # 模糊匹配
            for key, width in FIXED_WIDTHS.items():
                if key in col_name:
                    fixed_width = width
                    break

        if fixed_width:
            ws.column_dimensions[col_letter].width = fixed_width
        else:
            ws.column_dimensions[col_letter].width = auto_fit_column_width(ws, col_idx, min_width=12)

    # 设置打印设置
    ws.page_setup.orientation = 'landscape'
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToHeight = False
    ws.page_setup.fitToWidth = 1
    ws.print_area = None
    # 设置页边距：上和右1.2cm（约0.47英寸），左和下1.5cm（约0.59英寸）
    ws.page_margins.left = 0.7
    ws.page_margins.right = 0.47
    ws.page_margins.top = 0.47
    ws.page_margins.bottom = 0.7
    # 移除打印页码
    ws.oddFooter.center.text = ""
    ws.oddFooter.left.text = ""
    ws.oddFooter.right.text = ""

def extract_note(ws, total_row):
    note = None
    for r in range(total_row + 1, min(total_row + 5, ws.max_row + 1)):
        for c in range(1, 5):
            val = ws.cell(r, c).value
            if val and '备注' in str(val):
                full_note = str(val)
                if '特约：' in full_note:
                    idx = full_note.find('特约：')
                    note = full_note[idx + 3:]
                elif '特约:' in full_note:
                    idx = full_note.find('特约:')
                    note = full_note[idx + 3:]
                elif '特约' in full_note:
                    idx = full_note.find('特约')
                    note = full_note[idx + 2:]
                else:
                    note = full_note
                return note.strip() if note else note
    return note

def set_run_font_standard(run, size_pt=14):
    """
    设置run的字体为标准格式：中文Songti SC，英文Times New Roman
    确保LibreOffice转PDF时字体正确显示
    """
    run.font.size = Pt(size_pt)
    run.font.name = 'Times New Roman'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Songti SC')
    rFonts.set(qn('w:cs'), 'Times New Roman')


def apply_mixed_font(run):
    """应用混合字体格式：12pt，中文Songti SC，英文Times New Roman"""
    set_run_font_standard(run, 12)


def apply_payment_notice_font(run):
    """应用付款通知书的字体格式：四号(14pt)，中文Songti SC，英文Times New Roman"""
    set_run_font_standard(run, 14)

def find_sheet_match_in_paragraph(paragraph, sheet_names):
    """
    在段落中查找与sheet名匹配的文本
    返回: (matched_sheet_name, match_position) 或 (None, -1)
    """
    text = paragraph.text
    for sheet_name in sheet_names:
        if sheet_name in text:
            return (sheet_name, text.find(sheet_name))
    return (None, -1)

def replace_text_preserve_format(paragraph, pattern, replacement):
    """
    在段落中替换文本，同时保留原有格式（包括下划线、字体）
    用于付款通知书中的金额填充
    """
    if pattern not in paragraph.text:
        return False

    # 保存段落的对齐方式和格式
    original_alignment = paragraph.alignment

    # 保存第一个run的格式信息（如果存在）
    original_font_size = None
    original_font_name = None
    original_underline = None
    original_bold = None

    if paragraph.runs:
        first_run = paragraph.runs[0]
        original_font_size = first_run.font.size
        original_font_name = first_run.font.name
        original_underline = first_run.font.underline
        original_bold = first_run.font.bold

    # 获取原始文本并执行替换
    full_text = paragraph.text
    new_text = full_text.replace(pattern, replacement)

    # 清空段落内容
    paragraph.clear()

    # 添加新文本并应用格式
    run = paragraph.add_run(new_text)

    # 应用付款通知书的基本字体格式（四号，中文宋体，英文Times New Roman）- 完整设置确保PDF正确
    set_run_font_standard(run, 14)

    # 如果原来有下划线，保留下划线
    if original_underline:
        run.font.underline = original_underline

    # 如果原来是粗体，保留粗体
    if original_bold:
        run.font.bold = original_bold

    # 恢复段落对齐
    if original_alignment:
        paragraph.alignment = original_alignment

    return True

def replace_text_in_paragraph(paragraph, context, preserve_font_size=False):
    if "{" not in paragraph.text:
        return
    text = paragraph.text
    original_text = text

    # 保存原始字体大小（默认12pt）
    original_font_size = Pt(12)
    if preserve_font_size and paragraph.runs:
        for run in paragraph.runs:
            if run.font.size:
                original_font_size = run.font.size
                break

    for key, value in context.items():
        if key in text:
            text = text.replace(key, str(value))
    if text != original_text:
        paragraph.clear()
        run = paragraph.add_run(text)
        # 统一字体：Songti SC（中文）+ Times New Roman（英文）
        run.font.size = original_font_size
        run.font.name = 'Times New Roman'
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Songti SC')
        rFonts.set(qn('w:cs'), 'Times New Roman')

def replace_regex_in_paragraph(paragraph, regex_map):
    text = paragraph.text
    changed = False
    for pattern, replacement in regex_map.items():
        if re.search(pattern, text):
            try:
                text = re.sub(pattern, replacement, text)
                changed = True
            except re.error:
                pass  # 正则表达式错误
    if changed:
        paragraph.clear()
        run = paragraph.add_run(text)
        apply_mixed_font(run)

def remove_images_from_doc(doc):
    """移除文档中的所有图片（用于生成无签名Word版）"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            drawings = run._element.findall('.//' + qn('w:drawing'))
            for drawing in drawings:
                drawing.getparent().remove(drawing)
            picts = run._element.findall('.//' + qn('w:pict'))
            for pict in picts:
                pict.getparent().remove(pict)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        drawings = run._element.findall('.//' + qn('w:drawing'))
                        for drawing in drawings:
                            drawing.getparent().remove(drawing)
                        picts = run._element.findall('.//' + qn('w:pict'))
                        for pict in picts:
                            pict.getparent().remove(pict)

def convert_to_pdf(docx_path, pdf_path):
    try:
        import platform
        if platform.system() == 'Darwin':
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                '/usr/local/bin/soffice',
                'soffice'
            ]
            output_dir = os.path.dirname(pdf_path)
            for lo_path in libreoffice_paths:
                try:
                    subprocess.run([
                        lo_path, '--headless', '--convert-to', 'pdf',
                        '--outdir', output_dir, docx_path
                    ], capture_output=True, timeout=60)
                    
                    generated_pdf = os.path.join(
                        output_dir, 
                        os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
                    )
                    if os.path.exists(generated_pdf):
                        if generated_pdf != pdf_path:
                            if os.path.exists(pdf_path): os.remove(pdf_path)
                            os.rename(generated_pdf, pdf_path)
                        return True
                except (subprocess.SubprocessError, FileNotFoundError, OSError, subprocess.TimeoutExpired):
                    continue
        elif platform.system() == 'Windows':
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                return True
            except (ImportError, OSError):
                pass
        return False
    except Exception as e:
        print(f"PDF 转换失败: {e}")
        return False

def convert_excel_to_pdf(excel_path, pdf_path):
    try:
        import platform
        output_dir = os.path.dirname(pdf_path)
        if platform.system() == 'Darwin':
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                '/usr/local/bin/soffice',
                'soffice'
            ]
            for lo_path in libreoffice_paths:
                try:
                    subprocess.run([
                        lo_path, '--headless', '--convert-to', 'pdf',
                        '--outdir', output_dir, excel_path
                    ], capture_output=True, timeout=180)
                    generated_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(excel_path))[0] + '.pdf')
                    if os.path.exists(generated_pdf):
                        if generated_pdf != pdf_path:
                            if os.path.exists(pdf_path): os.remove(pdf_path)
                            os.rename(generated_pdf, pdf_path)
                        return True
                except (subprocess.SubprocessError, FileNotFoundError, OSError, subprocess.TimeoutExpired):
                    continue
        elif platform.system() == 'Windows':
            try:
                import comtypes.client
                excel = comtypes.client.CreateObject('Excel.Application')
                excel.Visible = False
                wb = excel.Workbooks.Open(os.path.abspath(excel_path))
                wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
                wb.Close(False)
                excel.Quit()
                return True
            except (ImportError, OSError, AttributeError):
                pass
        return False
    except Exception as e:
        print(f"Excel 转 PDF 失败: {e}")
        return False


# ==========================================
# Excel 格式处理
# ==========================================

class ExcelFormatter:
    FONT_CN = Font(name='Songti SC', size=10)
    FONT_EN = Font(name='Times New Roman', size=10)
    FONT_HEADER = Font(name='Songti SC', size=10, bold=True)
    FONT_TITLE = Font(name='Songti SC', size=18, bold=True)
    ALIGN_CENTER = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ALIGN_LEFT = Alignment(horizontal='left', vertical='center', wrap_text=True)
    ALIGN_RIGHT = Alignment(horizontal='right', vertical='center', wrap_text=True)

    @staticmethod
    def get_visual_settings(customer_type=None):
        """获取可视化设置"""
        config_manager = get_config_manager()
        return config_manager.get_visual_settings(customer_type)

    @staticmethod
    def create_font_from_settings(vs, font_type='normal'):
        """根据设置创建字体对象"""
        font_settings = vs.get('font', {})
        font_name = font_settings.get('name', 'Songti SC')

        if font_type == 'title':
            size = font_settings.get('title_size', 18)
            return Font(name=font_name, size=size, bold=True)
        elif font_type == 'header':
            size = font_settings.get('header_size', 10)
            bold = vs.get('header_style', {}).get('bold', True)
            return Font(name=font_name, size=size, bold=bold)
        else:
            size = font_settings.get('size', 10)
            return Font(name=font_name, size=size)

    @staticmethod
    def format_sheet(ws, header_row, total_row, sheet_name="", is_hengli=False, page_orientation="landscape", customer_type=None, is_first_sheet=False):
        # 获取可视化设置
        vs = ExcelFormatter.get_visual_settings(customer_type)
        font_settings = vs.get('font', {})
        row_height_settings = vs.get('row_height', {})
        header_style = vs.get('header_style', {})

        # 获取行高配置（新配置优先，否则使用旧配置）
        from customer_config import get_config_manager
        config_mgr = get_config_manager()
        customer_config = config_mgr.get_customer_config(customer_type) if customer_type else None
        row_heights_config = customer_config.row_heights if customer_config and customer_config.row_heights else None
        special_columns = customer_config.special_columns if customer_config else []

        # 从配置创建字体
        font_name = font_settings.get('name', 'Songti SC')
        font_size = font_settings.get('size', 10)
        title_size = font_settings.get('title_size', 18)
        header_size = font_settings.get('header_size', 10)
        header_bold = header_style.get('bold', True)

        # 创建字体对象
        font_normal = Font(name=font_name, size=font_size)
        font_title = Font(name=font_name, size=title_size, bold=True)
        font_header = Font(name=font_name, size=header_size, bold=header_bold)

        # 表头背景色
        header_bg_color = header_style.get('background_color')
        header_fill = PatternFill(start_color=header_bg_color, end_color=header_bg_color, fill_type='solid') if header_bg_color else None

        # 查找车牌号列和货种列（如果有）
        plate_col = None
        cargo_type_col = None
        max_col = min(ws.max_column, 20)  # 限制最大列数，避免处理空列
        for c in range(1, max_col + 1):
            hdr = str(ws.cell(header_row, c).value or '').replace('\n', '')
            if '车牌' in hdr:
                plate_col = c
            if '货种' in hdr or '保险货物' in hdr:
                cargo_type_col = c

        # 优化：只处理到 total_row + 5 行，避免处理大量空行
        max_row = min(total_row + 5, ws.max_row)

        for r in range(1, max_row + 1):
            for c in range(1, max_col + 1):
                cell = ws.cell(r, c)
                if r == 1:
                    # 惠州PTA：第一行使用等线18号，不修改（已在process_hengli_sheet中设置）
                    if is_hengli and customer_type == "惠州PTA":
                        pass  # 保持process_hengli_sheet中设置的等线18号
                    # 恒力PTA：标题12号，编号9号左对齐
                    elif is_hengli:
                        if cell.value and '编号' in str(cell.value):
                            cell.font = Font(name=font_name, size=9)
                            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                        elif cell.value:
                            cell.font = Font(name=font_name, size=12, bold=True)
                            cell.alignment = ExcelFormatter.ALIGN_CENTER
                    else:
                        cell.font = font_title
                        cell.alignment = ExcelFormatter.ALIGN_CENTER
                elif r == header_row:
                    cell.font = font_header
                    cell.alignment = ExcelFormatter.ALIGN_CENTER
                    if header_fill:
                        cell.fill = header_fill
                elif r > total_row:
                    val = str(cell.value) if cell.value else ""
                    has_chinese = any('\u4e00' <= ch <= '\u9fff' for ch in val)
                    cell.font = ExcelFormatter.FONT_CN if has_chinese else ExcelFormatter.FONT_EN
                    # === 核心修改开始：专门检测"投保单号"或"保单号"并强制右对齐 ===
                    if '投保单号' in val or '保单号' in val:
                        cell.alignment = ExcelFormatter.ALIGN_RIGHT
                    else:
                        cell.alignment = ExcelFormatter.ALIGN_LEFT
                    # === 核心修改结束 ===
                else:
                    val = str(cell.value) if cell.value else ""
                    has_chinese = any('\u4e00' <= ch <= '\u9fff' for ch in val)
                    cell.font = ExcelFormatter.FONT_CN if has_chinese else ExcelFormatter.FONT_EN
                    # 货种列特殊处理：启用wrap_text以支持长文本显示
                    if cargo_type_col and c == cargo_type_col and len(val) > 10:
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    # 恒力PTA车牌号列：启用wrap_text自动换行，左对齐
                    elif is_hengli and customer_type == "恒力PTA" and plate_col and c == plate_col and r >= 3 and r < total_row:
                        cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                    elif isinstance(cell.value, (int, float)):
                        cell.alignment = ExcelFormatter.ALIGN_RIGHT
                    else:
                        cell.alignment = ExcelFormatter.ALIGN_CENTER

            # 行高设置（优先使用新配置，否则使用旧配置）
            if row_heights_config:
                title_height = row_heights_config.title_row
                second_height = row_heights_config.second_row
                header_height = row_heights_config.header_row
                data_height = row_heights_config.data_row
                total_height = row_heights_config.total_row
                footer_height = row_heights_config.footer_row
                auto_fit = row_heights_config.auto_fit
            else:
                title_height = row_height_settings.get('title', 39)
                second_height = row_height_settings.get('second', 33)
                header_height = row_height_settings.get('header', 20)
                data_height = row_height_settings.get('data', 15)
                total_height = row_height_settings.get('total', 18)
                footer_height = row_height_settings.get('footer', 18)
                auto_fit = True

            # 惠州PTA：前3行保持原格式不变，直接跳过
            if customer_type == "惠州PTA" and r <= 3:
                pass  # 不修改行高
            elif r == 1:
                ws.row_dimensions[r].height = title_height
            elif r == 2 and not is_hengli:
                # 多式联运：第2行行高使用配置值
                ws.row_dimensions[r].height = second_height
            elif r == header_row:
                # 恒力PTA/惠州PTA表头行高设为32，其他类型使用配置值
                if is_hengli and customer_type in ("恒力PTA", "惠州PTA"):
                    ws.row_dimensions[r].height = 32
                else:
                    ws.row_dimensions[r].height = header_height
            elif customer_type == "恒力能源销售" and r >= 3:
                # 恒力能源销售：第3行及以下所有行高设为18
                ws.row_dimensions[r].height = 18
            elif is_hengli and customer_type == "恒力PTA" and r >= 3 and r < total_row:
                # 恒力PTA：根据车牌号内容计算行高（每行17）
                if plate_col:
                    plate_val = str(ws.cell(r, plate_col).value or '')
                    col_width = ws.column_dimensions[get_column_letter(plate_col)].width or 40
                    # 计算显示宽度：中文字符占2个单位，ASCII字符占1个单位
                    display_width = sum(2 if '\u4e00' <= ch <= '\u9fff' else 1 for ch in plate_val)
                    # 计算显示行数
                    lines = max(1, math.ceil(display_width / col_width)) if plate_val else 1
                    ws.row_dimensions[r].height = lines * 17
                else:
                    ws.row_dimensions[r].height = 17
            elif is_hengli and customer_type == "惠州PTA" and r > header_row and r <= total_row:
                # 惠州PTA：第5行及以下（header_row+1起）行高设为18，前3行保持原格式不变
                ws.row_dimensions[r].height = 18
            elif is_hengli and r >= 3 and r <= total_row:
                # 其他恒力类型：使用默认数据行高
                ws.row_dimensions[r].height = data_height
            elif r == total_row:
                # 多式联运第一个sheet合计行行高设为32
                if not is_hengli and is_first_sheet:
                    ws.row_dimensions[r].height = 32
                else:
                    ws.row_dimensions[r].height = total_height
            elif r > total_row:
                # 检查是否为备注行，备注行需要更大的行高
                row_text = str(ws.cell(r, 1).value or '')
                if '备注' in row_text:
                    # 计算备注行需要的行高（根据换行数）
                    newlines = row_text.count('\n')
                    note_height = max(footer_height, (newlines + 1) * 15)  # 每行约15
                    ws.row_dimensions[r].height = note_height
                else:
                    ws.row_dimensions[r].height = footer_height  # 使用页脚行高配置
            else:
                if auto_fit:
                    ws.row_dimensions[r].height = auto_fit_row_height(ws, r)
                else:
                    ws.row_dimensions[r].height = data_height
        
        # 列宽设置（使用max_col限制范围）
        if is_hengli:
            # 惠州PTA使用专门的列宽配置（在process_hengli_sheet中已设置），这里不覆盖
            if customer_type == "惠州PTA":
                # 惠州PTA：只设置A列宽度，其他列宽已在process_hengli_sheet中设置
                ws.column_dimensions['A'].width = 8
            else:
                # 恒力PTA专用列宽
                for c in range(1, max_col + 1):
                    hdr = str(ws.cell(header_row, c).value or '').replace('\n', '')
                    if '费率' in hdr:
                        ws.column_dimensions[get_column_letter(c)].width = 11
                    elif '保费' in hdr:
                        ws.column_dimensions[get_column_letter(c)].width = 13.4
                    elif hdr == '航次':
                        ws.column_dimensions[get_column_letter(c)].width = 7
                    elif '离港' in hdr or '运输日期' in hdr:
                        ws.column_dimensions[get_column_letter(c)].width = 12.5
                    elif '保险金额' in hdr:
                        ws.column_dimensions[get_column_letter(c)].width = 16.6
                    else:
                        ws.column_dimensions[get_column_letter(c)].width = auto_fit_column_width(ws, c, min_width=4, max_width=40)
                ws.column_dimensions['A'].width = 8
        else:
            # 多式联运模式 - 使用特殊列配置
            special_col_map = {}
            for sc in special_columns:
                special_col_map[sc.column_name] = sc

            for c in range(1, max_col + 1):
                hdr = str(ws.cell(header_row, c).value or '').replace('\n', '')
                col_letter = get_column_letter(c)

                # 检查是否有特殊列配置
                if hdr in special_col_map:
                    sc = special_col_map[hdr]
                    if is_first_sheet and sc.first_sheet_width:
                        ws.column_dimensions[col_letter].width = sc.first_sheet_width
                    elif not is_first_sheet and sc.other_sheet_width:
                        ws.column_dimensions[col_letter].width = sc.other_sheet_width
                    else:
                        ws.column_dimensions[col_letter].width = sc.width
                # 保持兼容：货种/保险货物列
                elif '货种' in hdr or '保险货物' in hdr:
                    ws.column_dimensions[col_letter].width = 25 if is_first_sheet else 13.5
                else:
                    ws.column_dimensions[col_letter].width = auto_fit_column_width(ws, c, min_width=4, max_width=40)
            ws.column_dimensions['A'].width = 8
        
        # 根据配置设置打印方向
        ws.page_setup.orientation = page_orientation
        
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        # 设置页边距：上和右1.2cm（约0.47英寸），左和下1.5cm（约0.59英寸）
        ws.page_margins.left = 0.7
        ws.page_margins.right = 0.47
        ws.page_margins.top = 0.47
        ws.page_margins.bottom = 0.7
        # 移除打印页码
        ws.oddFooter.center.text = ""
        ws.oddFooter.left.text = ""
        ws.oddFooter.right.text = ""

        # 多式联运第一个sheet：合计行的货种单元格（C:E）设置为左对齐并自动换行
        if not is_hengli and is_first_sheet:
            # 检查合计行是否有C:E合并单元格
            for merged_range in ws.merged_cells.ranges:
                if merged_range.min_row == total_row and merged_range.min_col == 3:
                    # 找到合计行的C列合并单元格，设置为左对齐、自动换行
                    ws.cell(total_row, 3).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                    # 根据内容长度自动调整行高
                    cell_value = str(ws.cell(total_row, 3).value or '')
                    if len(cell_value) > 30:
                        # 内容较长时增加行高
                        ws.row_dimensions[total_row].height = max(32, len(cell_value) // 15 * 15)
                    break

        # =======================================================
        # === 核心修改：仅针对'PTA船运'，精准定位打印截止行 ===
        # =======================================================
        if 'PTA船运' in ws.title:
            # 1. 锁定列范围：维持原表格的列数（不乱动宽度）
            max_col_letter = get_column_letter(ws.max_column-1)
            
            # 2. 寻找截止行：从合计行往下扫，找到包含目标文字的那一行
            target_row = ws.max_row  # 默认保底
            
            # 设定扫描范围：从合计行往下找 20 行足够了
            start_scan = total_row
            end_scan = min(ws.max_row + 1, total_row + 20)
            
            for r in range(start_scan, end_scan):
                # 拼接该行前 15 列的内容进行检查
                row_text = ""
                for c in range(1, 16):
                    row_text += str(ws.cell(r, c).value or "")
                
                # 只要这行字里包含“以实际开票金额为准”，它就是最后一行
                if "以实际开票金额为准" in row_text:
                    target_row = r
                    break
            
            # 3. 设置打印区域：A1 到 (最大列, 目标行)
            ws.print_area = f"A1:{max_col_letter}{target_row}"
        
# ==========================================
# 数据提取和处理
# ==========================================

def extract_sheet_data(ws, sheet_name):
    header_row = find_header_row(ws)
    total_row = find_total_row(ws, header_row)
    if not total_row: return None
    col_indices = find_column_indices(ws, header_row)
    
    tonnage = 0
    insurance_amount = 0
    premium = 0
    if col_indices['tonnage']: tonnage = ws.cell(total_row, col_indices['tonnage']).value
    if col_indices['insurance_amount']: insurance_amount = ws.cell(total_row, col_indices['insurance_amount']).value
    if col_indices['premium']: premium = ws.cell(total_row, col_indices['premium']).value
    
    cargo_type = None
    for c in [3, 4, 5]:
        val = ws.cell(total_row, c).value
        if val and str(val).strip() and '合计' not in str(val):
            cargo_type = str(val)
            break
    if not cargo_type and col_indices['cargo_type']:
        for r in range(header_row + 1, total_row):
            val = ws.cell(r, col_indices['cargo_type']).value
            if val:
                cargo_type = str(val)
                break
    
    ship_voyage = None
    departure_date = None
    data_row = header_row + 1
    for c in range(1, ws.max_column + 1):
        header_val = ws.cell(header_row, c).value
        if header_val and '船名' in str(header_val):
            ship_voyage = ws.cell(data_row, c).value
            break
    
    for c in range(1, ws.max_column + 1):
        header_val = ws.cell(header_row, c).value
        if header_val and '起运' in str(header_val) and '日期' in str(header_val):
            date_val = ws.cell(data_row, c).value
            if date_val:
                if isinstance(date_val, (int, float)):
                    try:
                        base_date = datetime(1899, 12, 30)
                        actual_date = base_date + timedelta(days=int(date_val))
                        departure_date = actual_date.strftime('%Y/%m/%d')
                    except (ValueError, OverflowError):
                        departure_date = str(date_val)
                elif isinstance(date_val, datetime):
                    departure_date = date_val.strftime('%Y/%m/%d')
                else:
                    departure_date = str(date_val)
            break
    
    special_terms = extract_note(ws, total_row)

    # 业务笔数：从合计行往上找最后一个有效序号，或者计算有效数据行数
    business_count = 0
    # 方法1：尝试从合计行上一行获取序号
    for r in range(total_row - 1, header_row, -1):
        seq_val = ws.cell(r, 1).value
        if seq_val is not None and seq_val != '':
            try:
                # 尝试转换为整数（支持浮点数格式如 1.0, 2.0）
                business_count = int(float(str(seq_val).strip()))
                break
            except (ValueError, TypeError):
                continue

    # 方法2：如果方法1失败，计算有效数据行数（非空行数）
    if business_count == 0:
        for r in range(header_row + 1, total_row):
            # 检查该行是否有有效数据（检查前几列）
            has_data = False
            for c in range(1, min(6, ws.max_column + 1)):
                val = ws.cell(r, c).value
                if val is not None and str(val).strip() and '合计' not in str(val):
                    has_data = True
                    break
            if has_data:
                business_count += 1
    
    try:
        tonnage = float(tonnage) if tonnage else 0
        insurance_amount = float(insurance_amount) if insurance_amount else 0
        premium = float(premium) if premium else 0
    except (ValueError, TypeError):
        tonnage = insurance_amount = premium = 0
    
    if insurance_amount > 0 and premium > 0:
        rate = premium / insurance_amount
        rate_rounded = round(rate, 8)
        new_premium = round(rate_rounded * insurance_amount, 2)
    else:
        rate_rounded = 0
        new_premium = 0
    
    return {
        'sheet_name': sheet_name,
        'cargo_type': cargo_type,
        'tonnage': tonnage,
        'insurance_amount': insurance_amount,
        'original_premium': premium,
        'rate': rate_rounded,
        'new_premium': new_premium,
        'header_row': header_row,
        'total_row': total_row,
        'col_indices': col_indices,
        'special_terms': special_terms,
        'ship_voyage': ship_voyage,
        'departure_date': departure_date,
        'business_count': business_count
    }

def update_sheet_with_rate(ws, data):
    total_row = data['total_row']
    col_indices = data.get('col_indices', {})
    premium_col = col_indices.get('premium')
    if premium_col:
        premium_cell = ws.cell(total_row, premium_col)
        premium_cell.value = data['new_premium']
        premium_cell.number_format = '#,##0.00'


def _extend_print_area_to_row(ws, target_row):
    """
    扩展打印区域以包含指定行。
    如果原打印区域的最大行小于 target_row，则扩展到 target_row。
    """
    if not ws.print_area:
        return  # 没有设置打印区域，不需要扩展

    try:
        # 解析打印区域，格式可能是 "A1:K20" 或 "$A$1:$K$20" 或 "'Sheet1'!$A$1:$K$20"
        area = ws.print_area
        if '!' in area:
            area = area.split('!')[1]
        area = area.replace('$', '')

        match = re.match(r'([A-Z]+)(\d+):([A-Z]+)(\d+)', area)
        if match:
            start_col, start_row, end_col, end_row = match.groups()
            start_row = int(start_row)
            end_row = int(end_row)

            # 如果目标行超出当前打印区域，扩展它
            if target_row > end_row:
                ws.print_area = f"{start_col}{start_row}:{end_col}{target_row}"
    except Exception:
        pass  # 解析失败，保持原打印区域不变


def set_safe_value(ws, row, col, value, alignment=None):
    cell = ws.cell(row, col)
    from openpyxl.cell.cell import MergedCell
    if isinstance(cell, MergedCell):
        for rng in ws.merged_cells.ranges:
            if row >= rng.min_row and row <= rng.max_row and col >= rng.min_col and col <= rng.max_col:
                cell = ws.cell(rng.min_row, rng.min_col)
                break
    cell.value = value
    if alignment: cell.alignment = alignment

def process_multimodal_sheet(ws, data):
    update_sheet_with_rate(ws, data)
    total_row = data['total_row']
    found_note = False
    note_row = -1
    for r in range(total_row + 1, total_row + 6):
        for c in range(1, 5):
            val = str(ws.cell(r, c).value or '')
            if '备注' in val:
                note_row = r
                found_note = True
                break
        if found_note: break
    target_row = (note_row + 1) if note_row > 0 else (total_row + 2)
    # 合并C:D列用于显示"保单号："，确保右对齐
    try:
        ws.unmerge_cells(f'C{target_row}:D{target_row}')
    except (KeyError, ValueError):
        pass  # 单元格可能未合并
    ws.merge_cells(f'C{target_row}:D{target_row}')
    ws.cell(target_row, 3).value = '保单号：'
    ws.cell(target_row, 3).alignment = Alignment(horizontal='right', vertical='center')
    ws.row_dimensions[target_row].height = 18  # 行高设为18

    # 保单号右边3个单元格合并后左对齐（E:F:G）
    try:
        ws.unmerge_cells(f'E{target_row}:G{target_row}')
    except (KeyError, ValueError):
        pass  # 单元格可能未合并
    ws.merge_cells(f'E{target_row}:G{target_row}')
    ws.cell(target_row, 5).value = ''
    ws.cell(target_row, 5).alignment = Alignment(horizontal='left', vertical='center')

def extract_hengli_data(ws, sheet_name):
    """提取恒力PTA表格数据，支持多种列名格式"""
    header_row = find_header_row(ws)
    if not header_row:
        return None

    # 从第一行识别康辉公司（去除空格后匹配）
    comp = None
    license_no = None
    row1_val = ''
    for c in range(1, min(ws.max_column + 1, 10)):
        cell_val = ws.cell(1, c).value
        if cell_val:
            row1_val += str(cell_val)
    # 去除所有空格进行匹配
    row1_normalized = row1_val.replace(' ', '').replace('\u3000', '')
    if '康辉大连新材料科技有限公司' in row1_normalized:
        comp = '康辉大连新材料科技有限公司'
        license_no = '91210244MA10YYBP1T'
    elif '康辉新材料科技有限公司' in row1_normalized:
        comp = '康辉新材料科技有限公司'
        license_no = '91210800580717031A'
    elif '康辉国际贸易（江苏）有限公司' in row1_normalized or '康辉国际贸易(江苏)有限公司' in row1_normalized:
        comp = '康辉国际贸易（江苏）有限公司'
        license_no = '91320509061869594L'

    # 查找"合计"行或"含税金额"行（去除空格后匹配"合计"）
    total_row = None
    for r in range(header_row + 1, ws.max_row + 1):
        for c in range(1, min(ws.max_column + 1, 10)):
            val = str(ws.cell(r, c).value or '').strip()
            # 去除空格后匹配"合计"
            val_normalized = val.replace(' ', '').replace('\u3000', '')
            if '合计' in val_normalized or '含税金额' in val_normalized:
                total_row = r
                break
        if total_row:
            break

    if not total_row:
        return None

    # 建立列名到列号的映射
    col_map = build_column_map(ws, header_row)

    # 创建局部查找函数（使用通用辅助函数）
    def find_col(*keywords):
        return find_column_by_keywords(col_map, *keywords)

    # 获取第一行数据行
    data_row = header_row + 1

    # A. 起运日期：离港时间 或 运输日期
    departure_date = None
    date_col = find_col('离港', '运输日期')
    if date_col:
        val = ws.cell(data_row, date_col).value
        if val:
            if isinstance(val, (int, float)):
                try:
                    base_date = datetime(1899, 12, 30)
                    dt = base_date + timedelta(days=int(val))
                    departure_date = dt.strftime('%Y/%m/%d')
                except (ValueError, OverflowError):
                    departure_date = str(val)
            elif isinstance(val, datetime):
                departure_date = val.strftime('%Y/%m/%d')
            else:
                departure_date = str(val)
    
    # B. 运输工具：船名 或 车牌号（车牌号只取第一个）
    transport_tool = None
    transport_col = find_col('船名', '车牌')
    if transport_col:
        val = ws.cell(data_row, transport_col).value
        if val:
            transport_str = str(val)
            # 如果是车牌号（包含/分隔），只取第一个
            if '/' in transport_str:
                transport_tool = transport_str.split('/')[0].strip()
            else:
                transport_tool = transport_str
    
    # C. 装货数量列（支持"装货重量（吨）"等多种表头格式）
    col_tonnage = find_col('装货数量', '装货重量', '数量（吨）', '实载')
    tonnage = 0
    if col_tonnage:
        tonnage_val = ws.cell(total_row, col_tonnage).value
        if tonnage_val and tonnage_val != '' and not isinstance(tonnage_val, str):
            tonnage = tonnage_val
        else:
            # 如果合计行没有数据，需要计算求和
            total_sum = 0
            for r in range(header_row + 1, total_row):
                cell_val = ws.cell(r, col_tonnage).value
                if cell_val and isinstance(cell_val, (int, float)):
                    total_sum += float(cell_val)
            tonnage = total_sum
    
    # D. 保险金额列
    col_amount = find_col('保险金额')
    insurance_amount = 0
    if col_amount:
        insurance_amount = ws.cell(total_row, col_amount).value or 0
    
    # E. 保费列
    col_premium = find_col('保费')
    original_premium = 0
    if col_premium:
        original_premium = ws.cell(total_row, col_premium).value or 0
    
    # 业务笔数：合计行上一行的序号（需要遍历查找最后一个有效序号）
    business_count = 0
    for r in range(total_row - 1, header_row, -1):
        seq_val = ws.cell(r, 1).value
        if seq_val is not None and seq_val != '':
            try:
                business_count = int(str(seq_val).strip())
                break
            except (ValueError, TypeError):
                pass
    
    # 转换为数值
    try:
        tonnage = float(tonnage) if tonnage else 0
        insurance_amount = float(insurance_amount) if insurance_amount else 0
        original_premium = float(original_premium) if original_premium else 0
    except (ValueError, TypeError):
        tonnage = insurance_amount = original_premium = 0
    
    # 使用 0.0083% 费率计算新保费
    target_rate = 0.000083  # 0.0083%
    new_premium = round(insurance_amount * target_rate, 2)
    
    # 查找不含税金额行和税额行位置
    untax_row = None
    tax_row = None
    for r in range(total_row + 1, min(total_row + 10, ws.max_row + 1)):
        val_a = str(ws.cell(r, 1).value or '').strip()
        if '不含税金额' in val_a:
            untax_row = r
        if val_a == '税额' or (val_a.endswith('税额') and len(val_a) < 10):
            tax_row = r

    # 提取来源字段（用于汇总表统计）
    # 优先查找"收货单位"、"收货方"、"客户名称"、"发货方"等列
    source = ''
    source_keywords = ['收货单位', '收货方', '客户名称', '发货方', '货主', '客户']
    source_col = None
    for kw in source_keywords:
        source_col = find_col(kw)
        if source_col:
            break
    if source_col:
        source_val = ws.cell(data_row, source_col).value
        if source_val:
            source = str(source_val).strip()

    return {
        'sheet_name': sheet_name,
        'departure_date': departure_date,
        'transport_tool': transport_tool,
        'tonnage': tonnage,
        'insurance_amount': insurance_amount,
        'original_premium': original_premium,
        'new_premium': new_premium,
        'rate': target_rate,
        'business_count': business_count,
        'header_row': header_row,
        'total_row': total_row,
        'col_premium': col_premium,
        'col_amount': col_amount,
        'col_tonnage': col_tonnage,
        'untax_row': untax_row,
        'tax_row': tax_row,
        'source': source,
        'comp': comp,
        'license_no': license_no
    }

def process_hengli_sheet(ws, data, is_huizhou=False, column_widths=None, policy_label="投保单号"):
    """处理恒力PTA/惠州PTA表格：格式化、调整保费差额、更新税额相关行

    Args:
        ws: 工作表对象
        data: 数据字典
        is_huizhou: 是否为惠州PTA
        column_widths: 列宽配置列表，每项为 {"column": "A", "width": 10.0} 格式
        policy_label: 单号标签，"投保单号" 或 "保单号"
    """
    from openpyxl.utils import get_column_letter
    
    total_row = data['total_row']
    col_premium = data.get('col_premium')
    header_row = data.get('header_row', 2)
    untax_row = data.get('untax_row')
    tax_row = data.get('tax_row')

    # ========== 惠州PTA特殊格式化 ==========
    if is_huizhou:
        # 1. 第一行：字体等线18号，并替换年度（2024-2025 -> 2025-2026）
        for c in range(1, ws.max_column + 1):
            cell = ws.cell(1, c)
            if cell.value:
                # 替换年度
                val = str(cell.value)
                if '2024-2025年度PTA' in val:
                    cell.value = val.replace('2024-2025年度PTA', '2025-2026年度PTA')
                # 设置字体为等线18号
                cell.font = Font(name='DengXian', size=18, bold=cell.font.bold if cell.font else False)

        # 2. 表头行（第4行/header_row）修复'单价            （元/吨）'为'单价（元/吨）'，并设置行高为32
        ws.row_dimensions[header_row].height = 32  # 表头行行高设置为32
        for c in range(1, ws.max_column + 1):
            cell = ws.cell(header_row, c)
            if cell.value:
                val = str(cell.value)
                if '单价' in val and '元/吨' in val:
                    # 移除多余空格
                    new_val = re.sub(r'\s+', '', val)
                    cell.value = new_val

        # 3. 找到最后一个有数据的列，用于设置打印区域和隐藏空白列
        last_data_col = 1
        for c in range(1, ws.max_column + 1):
            cell_val = ws.cell(header_row, c).value
            if cell_val and str(cell_val).strip():
                last_data_col = c
        last_col_letter = get_column_letter(last_data_col)

        # 4. 设置列宽（仅应用配置中指定的列宽，未配置的列保持自动宽度）
        if column_widths:
            for cw in column_widths:
                col = cw.get('column', '') if isinstance(cw, dict) else cw.column
                width = cw.get('width', 10.0) if isinstance(cw, dict) else cw.width
                if col:
                    ws.column_dimensions[col].width = width

        # 5. 隐藏空白列（L,M,N,O,P等，即last_data_col之后的列）
        for c in range(last_data_col + 1, ws.max_column + 1):
            col_letter = get_column_letter(c)
            ws.column_dimensions[col_letter].hidden = True

        # 6. 在合计行下一行的E列添加单号标签，右对齐
        policy_row = total_row + 1
        ws.cell(policy_row, 5).value = f'{policy_label}:'
        ws.cell(policy_row, 5).alignment = Alignment(horizontal='right', vertical='center')
        ws.cell(policy_row, 5).font = Font(name='Songti SC', size=10)

        # 投保单号右边3个单元格合并后左对齐（F:G:H）
        try:
            ws.unmerge_cells(f'F{policy_row}:H{policy_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        ws.merge_cells(f'F{policy_row}:H{policy_row}')
        ws.cell(policy_row, 6).value = ''
        ws.cell(policy_row, 6).alignment = Alignment(horizontal='left', vertical='center')

        # 7. 设置打印区域：仅包含有数据的列
        print_end_row = total_row + 1 if total_row else ws.max_row
        ws.print_area = f"A1:{last_col_letter}{print_end_row}"

    # ========== 恒力PTA格式化（仅恒力PTA执行）==========
    if not is_huizhou:
        # 1. 格式化第一行编号单元格（K或I列）：字体9号，自动换行，左对齐
        for c in range(9, ws.max_column + 1):  # 从I列开始查找
            cell = ws.cell(1, c)
            if cell.value and '编号' in str(cell.value):
                # 在编号和年月之间添加换行（如果还没有）
                val = str(cell.value)
                if '\n' not in val and '20' in val:
                    # 在年份前添加换行
                    val = re.sub(r'\s+(20\d{2}年)', r'\n\1', val)
                    cell.value = val
                cell.font = Font(name='Songti SC', size=9)
                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                break

    # 保费调整逻辑（两种类型都执行）
    if col_premium:
        original_premium = data.get('original_premium', 0)
        new_premium = data.get('new_premium', 0)
        diff = new_premium - original_premium
        
        # E/F. 如有差额，在合计行上一行的保费列调整
        if abs(diff) > 0.005:
            # 查找合计行上一行有保费数据的行
            last_data_row = total_row - 1
            for r in range(total_row - 1, header_row, -1):
                val = ws.cell(r, col_premium).value
                if val is not None and val != '':
                    last_data_row = r
                    break
            
            old_val = ws.cell(last_data_row, col_premium).value
            try:
                old_val = float(old_val) if old_val else 0
                set_safe_value(ws, last_data_row, col_premium, round(old_val + diff, 2))
            except (ValueError, TypeError):
                pass
            set_safe_value(ws, total_row, col_premium, new_premium)
            
            # 调整不含税金额和税额（按比例）
            if original_premium > 0 and untax_row and tax_row:
                ratio = new_premium / original_premium
                # 找到保费值所在列，不含税和税额的值应在同一列
                for r in [untax_row, tax_row]:
                    val_cell = ws.cell(r, col_premium)
                    if val_cell.value and isinstance(val_cell.value, (int, float)):
                        new_val = round(float(val_cell.value) * ratio, 2)
                        set_safe_value(ws, r, col_premium, new_val)
    
    # 2-5. 格式化不含税金额、税额行和说明文字行（仅恒力PTA）
    if not is_huizhou and untax_row and tax_row:
        # 2. 确保不含税金额和税额文字在A:B合并单元格
        # 先取消可能存在的合并
        try:
            ws.unmerge_cells(f'A{untax_row}:B{untax_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        try:
            ws.unmerge_cells(f'A{tax_row}:B{tax_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并

        # 重新合并并设置内容
        ws.merge_cells(f'A{untax_row}:B{untax_row}')
        ws.merge_cells(f'A{tax_row}:B{tax_row}')
        ws.cell(untax_row, 1).value = '不含税金额'
        ws.cell(untax_row, 1).alignment = Alignment(horizontal='left', vertical='center')
        ws.cell(tax_row, 1).value = '税额'
        ws.cell(tax_row, 1).alignment = Alignment(horizontal='left', vertical='center')

        # 3. 行高固定15
        ws.row_dimensions[untax_row].height = 15
        ws.row_dimensions[tax_row].height = 15

        # 找到或创建说明文字行（税额下方一行）
        note_row = tax_row + 1

        # 说明文字放在A:B:C:D合并单元格，左对齐
        try:
            ws.unmerge_cells(f'A{note_row}:D{note_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        ws.merge_cells(f'A{note_row}:D{note_row}')
        ws.cell(note_row, 1).value = '*不含税金额和税额以实际开票金额为准'
        ws.cell(note_row, 1).alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[note_row].height = 15

        # 4. F:G合并填入"投保单号："右对齐，H:I合并左对齐留空
        try:
            ws.unmerge_cells(f'F{note_row}:G{note_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        try:
            ws.unmerge_cells(f'H{note_row}:I{note_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并

        ws.merge_cells(f'F{note_row}:G{note_row}')
        ws.cell(note_row, 6).value = f'{policy_label}:'
        ws.cell(note_row, 6).alignment = Alignment(horizontal='right', vertical='center')

        # 单号右边3个单元格合并后左对齐（H:I:J）
        try:
            ws.unmerge_cells(f'H{note_row}:J{note_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        ws.merge_cells(f'H{note_row}:J{note_row}')
        ws.cell(note_row, 8).value = ''
        ws.cell(note_row, 8).alignment = Alignment(horizontal='left', vertical='center')

        # 更新打印区域以包含新添加的说明文字行
        _extend_print_area_to_row(ws, note_row)

    # 恒力PTA但没有不含税金额/税额行时（如康辉），直接在合计行下一行添加单号标签
    elif not is_huizhou and not untax_row and not tax_row:
        policy_row = total_row + 1
        # E列添加单号标签，右对齐
        ws.cell(policy_row, 5).value = f'{policy_label}:'
        ws.cell(policy_row, 5).alignment = Alignment(horizontal='right', vertical='center')
        ws.cell(policy_row, 5).font = Font(name='Songti SC', size=10)
        ws.row_dimensions[policy_row].height = 18

        # 投保单号右边3个单元格合并后左对齐（F:G:H）
        try:
            ws.unmerge_cells(f'F{policy_row}:H{policy_row}')
        except (KeyError, ValueError):
            pass  # 单元格可能未合并
        ws.merge_cells(f'F{policy_row}:H{policy_row}')
        ws.cell(policy_row, 6).value = ''
        ws.cell(policy_row, 6).alignment = Alignment(horizontal='left', vertical='center')

        # 更新打印区域以包含新添加的单号标签行
        _extend_print_area_to_row(ws, policy_row)

# ==========================================
# 工作线程
# ==========================================


class CellWithFallback:
    """单元格包装器，支持 data_only 值回退到公式值，并能计算简单公式"""

    def __init__(self, data_cell, formula_cell, ws_wrapper=None):
        self._data_cell = data_cell
        self._formula_cell = formula_cell
        self._ws_wrapper = ws_wrapper  # 用于计算公式时访问其他单元格

    @property
    def value(self):
        """获取单元格值：优先使用 data_only 的计算值，如果为 None 则尝试计算公式"""
        val = self._data_cell.value
        if val is None and self._formula_cell.value is not None:
            formula = self._formula_cell.value
            # 如果是公式，尝试计算
            if isinstance(formula, str) and formula.startswith('='):
                val = self._evaluate_formula(formula)
            else:
                val = formula
        return val

    def _evaluate_formula(self, formula):
        """尝试计算简单的 Excel 公式"""
        formula = formula.strip()
        if not formula.startswith('='):
            return formula

        formula_body = formula[1:].strip()

        # 处理 SUM 公式: =SUM(A1:A10)
        sum_match = re.match(r'^SUM\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)$', formula_body, re.IGNORECASE)
        if sum_match and self._ws_wrapper:
            col_start, row_start, col_end, row_end = sum_match.groups()
            row_start, row_end = int(row_start), int(row_end)
            col_start_idx = self._col_letter_to_idx(col_start)
            col_end_idx = self._col_letter_to_idx(col_end)

            total = 0.0
            for r in range(row_start, row_end + 1):
                for c in range(col_start_idx, col_end_idx + 1):
                    cell_val = self._ws_wrapper.cell(row=r, column=c).value
                    if isinstance(cell_val, (int, float)):
                        total += cell_val
            return total

        # 处理 ROUND 公式: =ROUND(expression, digits)
        round_match = re.match(r'^ROUND\((.+),\s*(\d+)\)$', formula_body, re.IGNORECASE)
        if round_match:
            expr, digits = round_match.groups()
            digits = int(digits)
            result = self._evaluate_expression(expr)
            if isinstance(result, (int, float)):
                return round(result, digits)

        # 处理 ROW()-N 公式: =ROW()-4
        row_match = re.match(r'^ROW\(\)\s*-\s*(\d+)$', formula_body, re.IGNORECASE)
        if row_match:
            offset = int(row_match.group(1))
            return self._data_cell.row - offset

        # 处理简单的算术表达式: =A1*B1, =A1+B1 等
        simple_expr = self._evaluate_expression(formula_body)
        if simple_expr is not None:
            return simple_expr

        # 无法计算，返回原始公式
        return formula

    def _evaluate_expression(self, expr):
        """计算简单的算术表达式"""
        if not self._ws_wrapper:
            return None

        # 替换单元格引用为实际值
        def replace_cell_ref(match):
            col_letter = match.group(1)
            row_num = int(match.group(2))
            col_idx = self._col_letter_to_idx(col_letter)
            cell_val = self._ws_wrapper.cell(row=row_num, column=col_idx).value
            if isinstance(cell_val, (int, float)):
                return str(cell_val)
            return '0'

        # 替换单元格引用
        expr_with_values = re.sub(r'([A-Z]+)(\d+)', replace_cell_ref, expr, flags=re.IGNORECASE)

        # 使用安全的 AST 计算器（替代 eval）
        result = safe_eval_expr(expr_with_values)
        if result is not None:
            return result

        return None

    def _col_letter_to_idx(self, col_letter):
        """将列字母转换为列索引（1-based）"""
        result = 0
        for char in col_letter.upper():
            result = result * 26 + (ord(char) - ord('A') + 1)
        return result

    @property
    def row(self):
        return self._data_cell.row

    @property
    def column(self):
        return self._data_cell.column

    @property
    def col_idx(self):
        return self._data_cell.column

    def __getattr__(self, name):
        # 其他属性从 data_cell 获取
        return getattr(self._data_cell, name)


class WorksheetWithFallback:
    """工作表包装器，支持 data_only 值自动回退到公式值，并能计算简单公式

    用于处理 openpyxl 生成但未被 Excel 打开保存过的文件，
    这些文件的公式计算值不会被缓存，data_only=True 会返回 None。
    支持的公式：SUM, ROUND, ROW()-N, 简单算术运算 (+, -, *, /)
    """

    def __init__(self, ws_data, ws_formula):
        """
        Args:
            ws_data: data_only=True 加载的工作表（用于获取计算值）
            ws_formula: data_only=False 加载的工作表（用于获取公式作为回退）
        """
        self._ws_data = ws_data
        self._ws_formula = ws_formula

    def cell(self, row=None, column=None):
        """获取单元格，自动处理 data_only 值回退和公式计算"""
        data_cell = self._ws_data.cell(row=row, column=column)
        formula_cell = self._ws_formula.cell(row=row, column=column)
        return CellWithFallback(data_cell, formula_cell, ws_wrapper=self)

    def __getitem__(self, key):
        """支持 ws[row] 访问方式"""
        if isinstance(key, int):
            # 返回行的单元格列表
            data_row = self._ws_data[key]
            formula_row = self._ws_formula[key]
            return [CellWithFallback(d, f, ws_wrapper=self) for d, f in zip(data_row, formula_row)]
        return self._ws_data[key]

    def iter_rows(self, min_row=None, max_row=None, min_col=None, max_col=None):
        """迭代行，自动处理回退"""
        data_rows = self._ws_data.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col)
        formula_rows = self._ws_formula.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col)
        for data_row, formula_row in zip(data_rows, formula_rows):
            yield [CellWithFallback(d, f, ws_wrapper=self) for d, f in zip(data_row, formula_row)]

    @property
    def max_row(self):
        return self._ws_data.max_row

    @property
    def max_column(self):
        return self._ws_data.max_column

    @property
    def sheetnames(self):
        return self._ws_data.parent.sheetnames

    def __getattr__(self, name):
        # 其他属性从 ws_data 获取
        return getattr(self._ws_data, name)


class ProcessWorker(QThread):
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(bool, str, list)

    def __init__(self, excel_files, output_dir, customer_type="多式联运", policy_label="投保单号"):
        super().__init__()
        self.excel_files = excel_files
        self.output_dir = output_dir
        self.customer_type = customer_type
        self.policy_label = policy_label  # "投保单号" 或 "保单号"
        self._is_stopped = False

    def stop(self):
        self._is_stopped = True

    def run(self):
        try:
            all_data = []
            total_files = len(self.excel_files)
            for idx, excel_path in enumerate(self.excel_files):
                if self._is_stopped:
                    self.log.emit("⏹️ 已停止处理")
                    self.finished.emit(False, "已停止", all_data)
                    return
                file_name = os.path.basename(excel_path)
                self.log.emit(f"📊 处理: {file_name}")

                # 同时加载两个工作簿：data_only 获取计算值，普通模式获取公式（作为回退）
                wb_data = openpyxl.load_workbook(excel_path, data_only=True)
                wb_formula = openpyxl.load_workbook(excel_path, data_only=False)

                sheet_data_list = []
                for sheet_name in wb_data.sheetnames:
                    if self._is_stopped:
                        wb_data.close()
                        wb_formula.close()
                        self.log.emit("⏹️ 已停止处理")
                        self.finished.emit(False, "已停止", all_data)
                        return

                    ws_data = wb_data[sheet_name]
                    ws_formula = wb_formula[sheet_name]

                    # 创建带回退功能的工作表包装器
                    ws = WorksheetWithFallback(ws_data, ws_formula)

                    # 根据客户类型选择提取函数
                    if self.customer_type in ("恒力PTA", "惠州PTA"):
                        data = extract_hengli_data(ws, sheet_name)
                    elif self.customer_type == "恒力能源销售":
                        data = extract_hengli_energy_data(ws, sheet_name)
                    else:
                        data = extract_sheet_data(ws, sheet_name)

                    if data:
                        data['file_name'] = file_name
                        sheet_data_list.append(data)
                        all_data.append(data)

                        # 恒力能源销售使用不同的日志格式
                        if self.customer_type == "恒力能源销售":
                            self.log.emit(f"  ✅ {sheet_name}: 金额={data.get('money', 0):,.2f}, 保费={data.get('prem', 0):,.2f}")
                        else:
                            rate_disp = data['rate'] * 1000 if self.customer_type in ("恒力PTA", "惠州PTA") else data['rate'] * 100
                            unit = "‰" if self.customer_type in ("恒力PTA", "惠州PTA") else "%"
                            self.log.emit(f"  ✅ {sheet_name}: 费率={rate_disp:.3f}{unit}, 保费={data['new_premium']:,.2f}")
                    else:
                        self.log.emit(f"  ⚠️ 跳过 {sheet_name}: 未找到有效数据结构（需要表头行和合计行）")

                wb_data.close()
                wb_formula.close()
                
                if self._is_stopped: return
                self.log.emit(f"  📝 正在格式化 {len(sheet_data_list)} 个工作表...")
                wb = openpyxl.load_workbook(excel_path)
                for i, data in enumerate(sheet_data_list):
                    ws = wb[data['sheet_name']]
                    # 根据客户类型选择处理函数
                    if self.customer_type in ("恒力PTA", "惠州PTA"):
                        is_huizhou = (self.customer_type == "惠州PTA")
                        # 获取配置中的列宽
                        config_manager = get_config_manager()
                        customer_config = config_manager.get_customer_config(self.customer_type)
                        col_widths = None
                        if customer_config and customer_config.processed_column_widths:
                            col_widths = [{"column": cw.column, "width": cw.width} for cw in customer_config.processed_column_widths]
                        process_hengli_sheet(ws, data, is_huizhou=is_huizhou, column_widths=col_widths, policy_label=self.policy_label)
                        is_hengli = True
                    elif self.customer_type == "恒力能源销售":
                        process_hengli_energy_sheet(ws, data, policy_label=self.policy_label)
                        is_hengli = False  # 恒力能源销售使用不同的格式化逻辑
                    else:
                        process_multimodal_sheet(ws, data)
                        is_hengli = False

                    # 恒力能源销售不使用 ExcelFormatter.format_sheet，因为它有自己的格式化
                    if self.customer_type != "恒力能源销售":
                        # 从配置中获取打印方向，默认横向
                        config_manager = get_config_manager()
                        customer_config = config_manager.get_customer_config(self.customer_type)
                        page_orientation = customer_config.page_orientation if customer_config else "landscape"
                        ExcelFormatter.format_sheet(ws, data['header_row'], data['total_row'], data['sheet_name'], is_hengli, page_orientation, self.customer_type, is_first_sheet=(i == 0))
                    # 更新进度
                    sub_progress = int((idx + (i + 1) / len(sheet_data_list)) / total_files * 70)
                    self.progress.emit(sub_progress)

                output_name = os.path.splitext(file_name)[0] + "_processed.xlsx"
                output_path = os.path.join(self.output_dir, output_name)
                self.log.emit(f"  💾 保存文件...")
                wb.save(output_path)
                wb.close()
                self.log.emit(f"  ✅ 完成: {output_name}")
            
            if all_data:
                self.log.emit("📋 生成汇总表...")
                self._create_summary(all_data)
            self.progress.emit(100)
            self.finished.emit(True, f"处理完成！共 {len(all_data)} 个 sheet", all_data)
        except Exception as e:
            import traceback
            self.finished.emit(False, f"处理出错:\n{traceback.format_exc()}", [])
    
    def _create_summary(self, all_data):
        from collections import Counter
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "汇总"

        # 辅助函数：从日期字符串中提取年月
        def extract_year_month(date_str):
            if not date_str:
                return None
            date_str = str(date_str)
            # 支持格式：2025/09/01, 2025-09-01, 2025.09.01
            match = re.match(r'(\d{4})[/\-.](\d{1,2})', date_str)
            if match:
                return (int(match.group(1)), int(match.group(2)))
            return None

        # 辅助函数：计算众数年月并格式化为"申报周期：xxxx年xx月"
        def get_majority_period_mark(all_data, date_key):
            year_months = []
            for d in all_data:
                date_val = d.get(date_key, '') or d.get('date', '')
                ym = extract_year_month(date_val)
                if ym:
                    year_months.append(ym)
            if not year_months:
                return "申报周期："
            # 找众数
            counter = Counter(year_months)
            most_common_ym = counter.most_common(1)[0][0]
            return f"申报周期：{most_common_ym[0]}年{most_common_ym[1]:02d}月"

        # 辅助函数：根据单个日期生成申报周期标记
        def get_period_mark_from_date(date_val):
            """从起运日期提取年月，生成 '申报周期：xxxx年xx月' 格式"""
            ym = extract_year_month(date_val)
            if ym:
                return f"申报周期：{ym[0]}年{ym[1]:02d}月"
            return "申报周期："

        # 1. 设置表头（根据客户类型）- 营业执照列在投保人和标记之间
        if self.customer_type == "惠州PTA":
            # Sheet名(1), 序号(2), 文件名(3), 投保人(4), 营业执照(5), 标记(6), 运输工具(7), 业务笔数(8), 起运日期(9), 装货数量(10), 保险金额(11), 保费(12), 人民币保费(13), 投保单号(14), 保险单号(15)
            headers = ["Sheet名", "序号", "文件名", "投保人", "营业执照", "标记", "运输工具", "业务笔数", "起运日期", "装货数量（吨）", "保险金额", "保费", "人民币保费", "投保单号", "保险单号"]
            period_mark = None  # 每行根据起运日期单独计算
        elif self.customer_type == "恒力PTA":
            # Sheet名(1), 序号(2), 文件名(3), 投保人(4), 营业执照(5), 标记(6), 运输工具(7), 业务笔数(8), 起运日期(9), 装货数量(10), 保险金额(11), 保费(12), 人民币保费(13), 投保单号(14), 保险单号(15)
            headers = ["Sheet名", "序号", "文件名", "投保人", "营业执照", "标记", "运输工具", "业务笔数", "起运日期", "装货数量（吨）", "保险金额", "保费", "人民币保费", "投保单号", "保险单号"]
            period_mark = None  # 每行根据起运日期单独计算
        elif self.customer_type == "恒力能源销售":
            # Sheet名(1), 序号(2), 物料名称(3), 投保人(4), 营业执照(5), 标记(6), 车船号(7), 业务笔数(8), 发货日期(9), 开单量(10), 金额(11), 保费(12), 申报止期(13), 投保单号(14), 保险单号(15)
            headers = ["Sheet名", "序号", "物料名称", "投保人", "营业执照", "标记", "车船号", "业务笔数", "发货日期", "开单量", "金额（元）", "保费", "申报止期", "投保单号", "保险单号"]
            period_mark = None  # 恒力能源销售的标记是每行不同
        else:
            # 多式联运：Sheet名(1), 货种(2), 非标准化特约(3), 投保人(4), 营业执照(5), 标记(6), 船名/航次(7), 业务笔数(8), 起运日期(9), 实载吨位(10), 保险金额(11), 保费(12), 千分费率(13), 投保单号(14), 保险单号(15)
            headers = ["Sheet名", "货种", "非标准化特约", "投保人", "营业执照", "标记", "船名/航次", "业务笔数", "起运日期", "实载吨位", "保险金额", "保费", "千分费率", "投保单号", "保险单号"]
            period_mark = get_majority_period_mark(all_data, 'departure_date')

        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(name='Songti SC', size=10, bold=True)
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 2. 填充数据并计算总和
        total_premium_val = 0.0
        row_num = 1  # 序号从1开始

        for d in all_data:
            if self.customer_type == "恒力能源销售":
                premium = d.get('prem', 0)
                total_premium_val += premium
                # 标记格式："申报周期：发货日期-申报止期"
                ship_date = format_date_slashes(d.get('date', ''))
                end_date = d.get('latest_date', '')
                row_mark = f"申报周期：{ship_date}-{end_date}"
                comp = d.get('comp', '')
                license_no = LICENSE_MAP.get(comp, '')
                # Sheet名(1), 序号(2), 物料名称(3), 投保人(4), 营业执照(5), 标记(6), 车船号(7), 业务笔数(8), 发货日期(9), 开单量(10), 金额(11), 保费(12), 申报止期(13), 投保单号(14), 保险单号(15)
                ws.append([
                    d.get('sheet_name', ''), row_num, d.get('mat', ''), comp, license_no,
                    row_mark, d.get('no', ''), d.get('business_count', 0), ship_date,
                    d.get('amt', 0), d.get('money', 0), premium, end_date, '', ''
                ])
            else:
                premium = d.get('new_premium', 0)
                total_premium_val += premium
                rmb_premium = cn_currency(premium)

                if self.customer_type == "惠州PTA":
                    comp = '恒力石化（惠州）有限公司'
                    license_no = LICENSE_MAP.get(comp, '')
                    # 根据该行的起运日期生成申报周期标记
                    row_period_mark = get_period_mark_from_date(d.get('departure_date', ''))
                    # Sheet名(1), 序号(2), 文件名(3), 投保人(4), 营业执照(5), 标记(6), 运输工具(7), 业务笔数(8), 起运日期(9), 装货数量(10), 保险金额(11), 保费(12), 人民币保费(13), 投保单号(14), 保险单号(15)
                    ws.append([
                        d['sheet_name'], row_num, d['file_name'], comp, license_no,
                        row_period_mark, d.get('transport_tool', ''), d.get('business_count', 0), d.get('departure_date', ''),
                        d.get('tonnage', 0), d.get('insurance_amount', 0), premium, rmb_premium, '', ''
                    ])
                elif self.customer_type == "恒力PTA":
                    # 优先使用从Excel第一行识别的康辉公司信息，否则使用默认值
                    comp = d.get('comp') or '恒力石化（大连）有限公司'
                    license_no = d.get('license_no') or LICENSE_MAP.get(comp, '')
                    # 根据该行的起运日期生成申报周期标记
                    row_period_mark = get_period_mark_from_date(d.get('departure_date', ''))
                    # Sheet名(1), 序号(2), 文件名(3), 投保人(4), 营业执照(5), 标记(6), 运输工具(7), 业务笔数(8), 起运日期(9), 装货数量(10), 保险金额(11), 保费(12), 人民币保费(13), 投保单号(14), 保险单号(15)
                    ws.append([
                        d['sheet_name'], row_num, d['file_name'], comp, license_no,
                        row_period_mark, d.get('transport_tool', ''), d.get('business_count', 0), d.get('departure_date', ''),
                        d.get('tonnage', 0), d.get('insurance_amount', 0), premium, rmb_premium, '', ''
                    ])
                else:
                    # 多式联运：Sheet名(1), 货种(2), 非标准化特约(3), 投保人(4), 营业执照(5), 标记(6), 船名/航次(7), 业务笔数(8), 起运日期(9), 实载吨位(10), 保险金额(11), 保费(12), 千分费率(13), 投保单号(14), 保险单号(15)
                    comp = '浙江卓航多式联运科技有限公司'
                    license_no = LICENSE_MAP.get(comp, '')
                    rate_val = d['rate']
                    rate_permille_str = f"{rate_val * 1000:.6f}".rstrip('0').rstrip('.')
                    ws.append([
                        d['sheet_name'], d['cargo_type'], d.get('special_terms', ''), comp, license_no,
                        period_mark, d.get('ship_voyage', ''), d.get('business_count', 0), d.get('departure_date', ''),
                        d['tonnage'], d['insurance_amount'], d['new_premium'], rate_permille_str, '', ''
                    ])
            row_num += 1

        # 3. 设置合计行
        row_count = len(all_data)
        sum_row = row_count + 2
        ws.cell(sum_row, 1).value = "合计"
        ws.cell(sum_row, 1).font = Font(name='Songti SC', size=10, bold=True)

        # 计算总保费的大写
        total_rmb_str = cn_currency(total_premium_val)

        # 配置汇总表格式：{数量列格式, 金额列格式, 保费列格式, 是否显示总大写保费, 列宽}
        SUMMARY_FORMAT_CONFIG = {
            "惠州PTA": {
                'qty_fmt': '#,##0.00', 'amt_fmt': '#,##0.00', 'prem_fmt': '#,##0.00',
                'show_rmb_total': True, 'rmb_col': 13,
                'col_widths': [20, 6, 25, 28, 22, 22, 10, 10, 12, 15, 18, 15, 22, 20, 20]
            },
            "恒力PTA": {
                'qty_fmt': '#,##0.00', 'amt_fmt': '#,##0.00', 'prem_fmt': '#,##0.00',
                'show_rmb_total': True, 'rmb_col': 13,
                'col_widths': [20, 6, 25, 28, 22, 22, 10, 10, 12, 15, 18, 15, 22, 20, 20]
            },
            "恒力能源销售": {
                'qty_fmt': '0.000', 'amt_fmt': '#,##0.00', 'prem_fmt': '#,##0.00',
                'show_rmb_total': False,
                'col_widths': [20, 6, 20, 38, 22, 30, 14, 10, 14, 14, 18, 16, 14, 20, 20]
            },
            "多式联运": {
                'qty_fmt': '#,##0.000', 'amt_fmt': '#,##0.00', 'prem_fmt': '#,##0.00',
                'show_rmb_total': False, 'special_col3_wrap': True,
                'col_widths': [30, 50, 60, 28, 22, 22, 15, 10, 12, 15, 18, 15, 10, 20, 20]
            },
        }

        config = SUMMARY_FORMAT_CONFIG.get(self.customer_type, SUMMARY_FORMAT_CONFIG["多式联运"])

        # 应用数据行格式
        for r in range(2, row_count + 2):
            ws.cell(r, 10).number_format = config['qty_fmt']
            ws.cell(r, 11).number_format = config['amt_fmt']
            ws.cell(r, 12).number_format = config['prem_fmt']
            if config.get('special_col3_wrap'):
                ws.cell(r, 3).alignment = Alignment(wrap_text=True, vertical='top')

        # 设置合计行公式
        ws.cell(sum_row, 8).value = f"=SUM(H2:H{sum_row-1})"   # 业务笔数
        ws.cell(sum_row, 10).value = f"=SUM(J2:J{sum_row-1})"  # 数量/吨位
        ws.cell(sum_row, 11).value = f"=SUM(K2:K{sum_row-1})"  # 金额
        ws.cell(sum_row, 12).value = f"=SUM(L2:L{sum_row-1})"  # 保费

        # 合计行数字格式
        ws.cell(sum_row, 10).number_format = config['qty_fmt']
        ws.cell(sum_row, 11).number_format = config['amt_fmt']
        ws.cell(sum_row, 12).number_format = config['prem_fmt']

        # 显示人民币大写总额（仅部分客户类型需要）
        if config.get('show_rmb_total'):
            ws.cell(sum_row, config['rmb_col']).value = total_rmb_str

        # 应用列宽
        for i, w in enumerate(config['col_widths'], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # 每个客户类型生成各自的汇总表
        summary_name = f"汇总表_{self.customer_type}.xlsx"
        wb.save(os.path.join(self.output_dir, summary_name))


# ==========================================
# Excel Sheet 复制辅助函数
# ==========================================

def parse_print_area(ws):
    """
    解析工作表的打印区域，返回 (min_row, max_row, min_col, max_col)
    如果没有设置打印区域，自动计算合理的打印范围：
    - 行：从第1行到包含"投保单号"的行（或"*不含税金额和税额以实际开票金额为准"行）
    - 列：排除隐藏列，最多到L列（第12列）
    """
    from openpyxl.utils import range_boundaries
    min_row, min_col = 1, 1
    max_row, max_col = ws.max_row, ws.max_column

    if ws.print_area:
        try:
            # print_area 可能是 "A1:K20" 或 "$A$1:$K$20" 或带Sheet名 "'Sheet1'!$A$1:$K$20"
            area = ws.print_area
            # 移除Sheet名前缀
            if '!' in area:
                area = area.split('!')[-1]
            area = area.replace('$', '')
            min_col, min_row, max_col, max_row = range_boundaries(area)
        except Exception:
            pass  # 解析失败则使用自动计算
    else:
        # 没有设置打印区域，自动计算
        # 1. 查找包含"投保单号"或"*不含税金额和税额以实际开票金额为准"的行作为最后一行
        target_row = None
        for r in range(1, min(ws.max_row + 1, 100)):
            for c in range(1, min(ws.max_column + 1, 15)):
                val = str(ws.cell(r, c).value or '')
                if '投保单号' in val or '以实际开票金额为准' in val:
                    target_row = r
                    break
            if target_row:
                break

        if target_row:
            max_row = target_row
        else:
            # 没找到目标行，查找合计行后的3行
            for r in range(1, ws.max_row + 1):
                val = str(ws.cell(r, 1).value or '')
                if '合计' in val or '含税金额' in val:
                    max_row = min(r + 3, ws.max_row)
                    break

        # 2. 列范围：最多到L列（第12列），排除隐藏列
        max_col = min(12, ws.max_column)  # L列 = 第12列
        # 检查是否有更早的隐藏列需要作为边界
        for c in range(1, max_col + 1):
            col_letter = get_column_letter(c)
            if col_letter in ws.column_dimensions and ws.column_dimensions[col_letter].hidden:
                # 如果某列是隐藏的，不包含它及之后的列
                max_col = c - 1
                break

    return min_row, max_row, min_col, max_col


def copy_sheet_cells(ws_src, ws_dst, unify_font=None, print_area_only=False, ws_data=None):
    """
    复制工作表单元格数据和样式

    Args:
        ws_src: 源工作表（用于获取样式）
        ws_dst: 目标工作表
        unify_font: 如果指定，统一所有单元格的字体名称（如 'Songti SC'）
        print_area_only: 如果为True，仅复制打印区域内的单元格
        ws_data: 数据源工作表（用于获取计算后的值，如果为None则使用ws_src）
    """
    from datetime import datetime, timedelta

    # Excel 内置中文日期格式 ID 映射（openpyxl 无法正确识别这些格式）
    # 参考：https://docs.microsoft.com/en-us/dotnet/api/documentformat.openxml.spreadsheet.numberingformat
    BUILTIN_DATE_FORMATS = {
        27: 'yyyy"年"m"月"',           # 中文年月
        28: 'm"月"d"日"',              # 中文月日
        29: 'm"月"d"日"',              # 中文月日
        30: 'm/d/yy',                  # 短日期
        31: 'yyyy"年"m"月"d"日"',      # 中文完整日期
        32: 'h"时"mm"分"',             # 中文时间
        33: 'h"时"mm"分"ss"秒"',       # 中文完整时间
        34: 'yyyy/m/d',                # 日期
        35: 'yyyy/m/d',                # 日期
        36: 'yyyy/m/d',                # 日期
        50: 'yyyy"年"m"月"',           # 中文年月
        51: 'm"月"d"日"',              # 中文月日
        52: 'yyyy"年"m"月"',           # 中文年月
        53: 'm"月"d"日"',              # 中文月日
        54: 'm"月"d"日"',              # 中文月日
        55: 'yyyy"年"m"月"',           # 中文年月
        56: 'm"月"d"日"',              # 中文月日
        57: 'yyyy"年"m"月"',           # 中文年月 (常见)
        58: 'm"月"d"日"',              # 中文月日
    }

    def get_real_number_format(cell):
        """获取单元格的真实数字格式，处理 openpyxl 无法识别的内置格式"""
        number_format = cell.number_format

        # 如果 openpyxl 返回 'General'，检查是否有内置格式 ID
        if number_format == 'General':
            try:
                # 尝试获取单元格的 numFmtId
                num_fmt_id = cell._style.numFmtId
                if num_fmt_id in BUILTIN_DATE_FORMATS:
                    return BUILTIN_DATE_FORMATS[num_fmt_id]
            except (AttributeError, TypeError):
                pass

        return number_format

    # 如果没有提供数据源，使用样式源
    if ws_data is None:
        ws_data = ws_src

    # 获取复制范围
    if print_area_only:
        min_row, max_row, min_col, max_col = parse_print_area(ws_src)
    else:
        min_row, max_row, min_col, max_col = 1, ws_src.max_row, 1, ws_src.max_column

    for row in ws_src.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
        for cell in row:
            # 计算目标单元格位置（如果是打印区域，调整为从1,1开始）
            dst_row = cell.row - min_row + 1 if print_area_only else cell.row
            dst_col = cell.column - min_col + 1 if print_area_only else cell.column

            # 从数据源获取值（计算后的值），从样式源获取样式
            data_cell = ws_data.cell(row=cell.row, column=cell.column)
            cell_value = data_cell.value

            # 如果 data_only 模式返回 None（公式未缓存），回退使用原始值
            # 这种情况发生在文件由 openpyxl 生成但未被 Excel 打开保存过
            if cell_value is None and cell.value is not None:
                cell_value = cell.value

            # 获取真实的 number_format（处理内置中文日期格式）
            number_format = get_real_number_format(cell)

            # 处理日期：
            # 1. 如果值是 datetime 对象，保持不变
            # 2. 如果值是日期序列号（40000-60000范围），转换为 datetime
            # 3. 保留原始 number_format（如 'yyyy"年"m"月"'、'yyyy/m/d' 等）
            if isinstance(cell_value, (int, float)) and not isinstance(cell_value, bool):
                # 检查是否可能是日期序列号
                if 40000 <= cell_value <= 60000:
                    # 检查 number_format 是否包含日期/时间指示符
                    has_date_format = any(x in str(number_format) for x in ['y', 'm', 'd', 'h', 's', '年', '月', '日'])

                    if has_date_format or number_format == 'General':
                        # 转换为 datetime 对象
                        try:
                            excel_epoch = datetime(1899, 12, 30)
                            cell_value = excel_epoch + timedelta(days=cell_value)

                            # 如果原格式是 General，设置一个合理的默认日期格式
                            if number_format == 'General':
                                number_format = 'yyyy/m/d'
                            # 否则保留原始格式（如 'yyyy"年"m"月"'）
                        except (ValueError, OverflowError):
                            pass  # 转换失败，保持原值

            new_cell = ws_dst.cell(row=dst_row, column=dst_col, value=cell_value)
            if cell.has_style:
                if unify_font:
                    src_font = cell.font
                    new_cell.font = Font(
                        name=unify_font,
                        size=src_font.size,
                        bold=src_font.bold,
                        italic=src_font.italic,
                        underline=src_font.underline,
                        color=src_font.color
                    )
                else:
                    new_cell.font = copy.copy(cell.font)
                new_cell.fill = copy.copy(cell.fill)
                new_cell.alignment = copy.copy(cell.alignment)
                new_cell.border = copy.copy(cell.border)
                new_cell.number_format = number_format


def copy_sheet_dimensions(ws_src, ws_dst, print_area_only=False):
    """复制工作表的列宽和行高"""
    if print_area_only:
        min_row, max_row, min_col, max_col = parse_print_area(ws_src)
        # 复制打印区域内的列宽
        for c in range(min_col, max_col + 1):
            src_letter = get_column_letter(c)
            dst_letter = get_column_letter(c - min_col + 1)
            if src_letter in ws_src.column_dimensions:
                ws_dst.column_dimensions[dst_letter].width = ws_src.column_dimensions[src_letter].width
        # 复制打印区域内的行高
        for r in range(min_row, max_row + 1):
            dst_row = r - min_row + 1
            if r in ws_src.row_dimensions:
                ws_dst.row_dimensions[dst_row].height = ws_src.row_dimensions[r].height
    else:
        for col_letter, dim in ws_src.column_dimensions.items():
            ws_dst.column_dimensions[col_letter].width = dim.width
        for row_num, dim in ws_src.row_dimensions.items():
            ws_dst.row_dimensions[row_num].height = dim.height


def copy_page_margins(ws_src, ws_dst):
    """复制工作表的页边距设置"""
    ws_dst.page_margins.left = ws_src.page_margins.left
    ws_dst.page_margins.right = ws_src.page_margins.right
    ws_dst.page_margins.top = ws_src.page_margins.top
    ws_dst.page_margins.bottom = ws_src.page_margins.bottom
    ws_dst.page_margins.header = ws_src.page_margins.header
    ws_dst.page_margins.footer = ws_src.page_margins.footer


def copy_merged_cells(ws_src, ws_dst, print_area_only=False):
    """复制工作表的合并单元格设置"""
    if print_area_only:
        min_row, max_row, min_col, max_col = parse_print_area(ws_src)
        for merged_range in ws_src.merged_cells.ranges:
            # 检查合并区域是否在打印区域内
            if (merged_range.min_row >= min_row and merged_range.max_row <= max_row and
                merged_range.min_col >= min_col and merged_range.max_col <= max_col):
                # 调整合并区域坐标
                new_min_row = merged_range.min_row - min_row + 1
                new_max_row = merged_range.max_row - min_row + 1
                new_min_col = merged_range.min_col - min_col + 1
                new_max_col = merged_range.max_col - min_col + 1
                new_range = f"{get_column_letter(new_min_col)}{new_min_row}:{get_column_letter(new_max_col)}{new_max_row}"
                ws_dst.merge_cells(new_range)
    else:
        for merged_range in ws_src.merged_cells.ranges:
            ws_dst.merge_cells(str(merged_range))


class PdfExportWorker(QThread):
    """
    PDF 导出工作线程：将所有 Sheet 的打印区域合并导出为单个 PDF 文件
    保持原 Excel 文件的字体、字号、页边距等格式
    """
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(bool, str, str)

    def __init__(self, excel_path, output_dir, orientation='landscape'):
        super().__init__()
        self.excel_path = excel_path
        self.output_dir = output_dir
        self.orientation = orientation
        self._is_stopped = False

    def stop(self): self._is_stopped = True

    def run(self):
        try:
            file_name = os.path.basename(self.excel_path)
            self.log.emit(f"📁 来源: {file_name}")

            # 加载两个工作簿：一个获取样式，一个获取计算后的值
            wb_style = openpyxl.load_workbook(self.excel_path)  # 样式源
            wb_data = openpyxl.load_workbook(self.excel_path, data_only=True)  # 数据源（计算值）
            total_sheets = len(wb_style.sheetnames)

            # 创建一个临时工作簿，将所有 Sheet 的打印区域复制进去
            wb_temp = openpyxl.Workbook()
            wb_temp.remove(wb_temp.active)

            sheet_count = 0
            for idx, sheet_name in enumerate(wb_style.sheetnames):
                if self._is_stopped:
                    wb_temp.close()
                    wb_style.close()
                    wb_data.close()
                    self.finished.emit(False, "已停止", self.output_dir)
                    return

                ws_style = wb_style[sheet_name]  # 样式源
                ws_data = wb_data[sheet_name]    # 数据源
                self.log.emit(f"  📄 处理: {sheet_name}")

                # 检查是否有打印区域
                print_area = ws_style.print_area
                if print_area:
                    self.log.emit(f"      打印区域: {print_area}")

                ws_dst = wb_temp.create_sheet(sheet_name)

                # 复制打印区域内的数据（从ws_data获取计算值，从ws_style获取样式）
                copy_sheet_cells(ws_style, ws_dst, print_area_only=True, ws_data=ws_data)
                copy_merged_cells(ws_style, ws_dst, print_area_only=True)
                copy_sheet_dimensions(ws_style, ws_dst, print_area_only=True)

                # 复制完整的页面设置（保留原始缩放比例，确保PDF格式一致）
                ws_dst.page_setup.orientation = ws_style.page_setup.orientation or self.orientation
                ws_dst.page_setup.paperSize = ws_style.page_setup.paperSize
                ws_dst.page_setup.scale = ws_style.page_setup.scale

                # 复制 fitToPage 设置
                if ws_style.sheet_properties.pageSetUpPr and ws_style.sheet_properties.pageSetUpPr.fitToPage:
                    ws_dst.sheet_properties.pageSetUpPr.fitToPage = True
                    ws_dst.page_setup.fitToWidth = ws_style.page_setup.fitToWidth
                    ws_dst.page_setup.fitToHeight = ws_style.page_setup.fitToHeight

                # 复制页边距
                copy_page_margins(ws_style, ws_dst)

                sheet_count += 1
                self.progress.emit(int((idx + 1) / total_sheets * 80))

            wb_style.close()
            wb_data.close()

            if sheet_count == 0:
                wb_temp.close()
                self.finished.emit(False, "没有可导出的工作表", self.output_dir)
                return

            # 保存临时 Excel 文件
            base_name = os.path.splitext(file_name)[0]
            temp_xlsx = os.path.join(self.output_dir, f'_temp_{base_name}.xlsx')
            wb_temp.save(temp_xlsx)
            wb_temp.close()

            # 转换为单个 PDF 文件
            self.log.emit(f"📑 正在转换为 PDF（共 {sheet_count} 个工作表）...")
            pdf_name = f"{base_name}.pdf"
            pdf_path = os.path.join(self.output_dir, pdf_name)

            if convert_excel_to_pdf(temp_xlsx, pdf_path):
                self.log.emit(f"✅ 已生成: {pdf_name}")
                self.progress.emit(100)
                # 删除临时文件
                if os.path.exists(temp_xlsx):
                    os.remove(temp_xlsx)
                self.finished.emit(True, f"完成！已将 {sheet_count} 个工作表导出为单个 PDF", self.output_dir)
            else:
                # 转换失败，保留临时文件供调试
                self.log.emit(f"❌ PDF 转换失败")
                self.finished.emit(False, "PDF 转换失败，请检查是否已安装 LibreOffice", self.output_dir)

        except Exception as e:
            import traceback
            self.finished.emit(False, f"出错:\n{traceback.format_exc()}", self.output_dir)


class PdfGroupedExportWorker(QThread):
    """按公司分组导出 PDF 的工作线程（用于恒力能源销售）"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(bool, str, str)

    def __init__(self, excel_path, output_dir, orientation='landscape'):
        super().__init__()
        self.excel_path = excel_path
        self.output_dir = output_dir
        self.orientation = orientation
        self._is_stopped = False

    def stop(self):
        self._is_stopped = True

    def run(self):
        try:
            self.log.emit(f"📁 来源: {os.path.basename(self.excel_path)}")
            self.progress.emit(10)

            wb = openpyxl.load_workbook(self.excel_path)
            results = {}

            # 遍历每个分组
            for group_name, match_func in PDF_EXPORT_GROUPS.items():
                if self._is_stopped:
                    break

                # 查找匹配的 sheet
                matched_sheets = [s for s in wb.sheetnames if match_func(s)]
                if not matched_sheets:
                    self.log.emit(f"  ⚠️ {group_name}: 无匹配的 Sheet")
                    continue

                self.log.emit(f"  📋 {group_name}: 找到 {len(matched_sheets)} 个 Sheet")

                # 创建临时工作簿
                wb_temp = openpyxl.Workbook()
                wb_temp.remove(wb_temp.active)

                for sheet_name in matched_sheets:
                    ws_src = wb[sheet_name]
                    ws_dst = wb_temp.create_sheet(sheet_name)

                    # 使用辅助函数复制数据（统一字体为 Songti SC）
                    copy_merged_cells(ws_src, ws_dst)
                    copy_sheet_cells(ws_src, ws_dst, unify_font='Songti SC')
                    copy_sheet_dimensions(ws_src, ws_dst)

                    # 设置页面布局
                    ws_dst.page_setup.orientation = self.orientation
                    ws_dst.sheet_properties.pageSetUpPr.fitToPage = True
                    ws_dst.page_setup.fitToWidth = 1
                    ws_dst.page_setup.fitToHeight = 0

                    copy_page_margins(ws_src, ws_dst)

                # 保存临时文件并转换为 PDF
                temp_xlsx = os.path.join(self.output_dir, f'_temp_{group_name}.xlsx')
                wb_temp.save(temp_xlsx)
                wb_temp.close()

                pdf_path = os.path.join(self.output_dir, f'{group_name}.pdf')
                if convert_excel_to_pdf(temp_xlsx, pdf_path):
                    results[group_name] = pdf_path
                    self.log.emit(f"  ✅ 已生成: {group_name}.pdf")
                else:
                    self.log.emit(f"  ❌ 转换失败: {group_name}")

                # 删除临时文件
                if os.path.exists(temp_xlsx):
                    os.remove(temp_xlsx)

            wb.close()
            self.progress.emit(100)

            if results:
                msg = f"完成！生成 {len(results)} 个 PDF:\n" + "\n".join(f"• {k}.pdf" for k in results.keys())
            else:
                msg = "⚠️ 无匹配的 Sheet"

            self.finished.emit(True, msg, self.output_dir)
        except Exception as e:
            import traceback
            self.finished.emit(False, f"出错:\n{traceback.format_exc()}", self.output_dir)


class PolicyBackfillWorker(QThread):
    """从汇总表回填投保单号和保单号到processed文件的工作线程"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(bool, str)

    def __init__(self, summary_path, target_files, customer_type):
        super().__init__()
        self.summary_path = summary_path
        self.target_files = target_files  # 可以是多个文件
        self.customer_type = customer_type
        self._is_stopped = False

    def stop(self):
        self._is_stopped = True

    def run(self):
        try:
            self.log.emit(f"📊 加载汇总表: {os.path.basename(self.summary_path)}")

            # 1. 解析汇总表，建立 (文件名, Sheet名) -> (投保单号, 保单号) 的映射
            wb_summary = openpyxl.load_workbook(self.summary_path, data_only=True)
            ws_summary = wb_summary.active

            # 获取表头列索引
            headers = {}
            for c in range(1, ws_summary.max_column + 1):
                h = ws_summary.cell(1, c).value
                if h:
                    headers[h] = c

            # 检查必要的列是否存在
            file_col = headers.get('文件名')
            sheet_col = headers.get('Sheet名')
            policy_col = headers.get('投保单号')
            # 支持"保单号"和"保险单号"两种列名
            cert_col = headers.get('保单号') or headers.get('保险单号')

            # Sheet名是必须的，文件名可选（恒力能源销售等单文件模式没有文件名列）
            if not sheet_col:
                self.finished.emit(False, "汇总表缺少'Sheet名'列")
                return

            has_file_col = file_col is not None
            if not has_file_col:
                self.log.emit("  ℹ️ 汇总表无'文件名'列，将仅通过Sheet名匹配")

            # 建立映射
            # 如果有文件名列: {(文件名, Sheet名): {'投保单号': xxx, '保单号': xxx}}
            # 如果无文件名列: {Sheet名: {'投保单号': xxx, '保单号': xxx}}
            policy_map = {}
            for r in range(2, ws_summary.max_row + 1):
                sheet_name = ws_summary.cell(r, sheet_col).value
                if not sheet_name:
                    continue

                if has_file_col:
                    file_name = ws_summary.cell(r, file_col).value
                    if not file_name:
                        continue
                    map_key = (file_name, sheet_name)
                else:
                    map_key = sheet_name

                policy_no = ws_summary.cell(r, policy_col).value if policy_col else None
                cert_no = ws_summary.cell(r, cert_col).value if cert_col else None

                if policy_no or cert_no:
                    policy_map[map_key] = {
                        '投保单号': policy_no or '',
                        '保单号': cert_no or ''
                    }

            wb_summary.close()
            self.log.emit(f"  ✅ 解析到 {len(policy_map)} 条单号记录")

            if not policy_map:
                self.finished.emit(False, "汇总表中没有找到有效的单号数据")
                return

            # 2. 遍历目标文件，回填单号
            total_files = len(self.target_files)
            success_count = 0
            sheet_count = 0

            for idx, target_path in enumerate(self.target_files):
                if self._is_stopped:
                    self.finished.emit(False, "已停止")
                    return

                target_name = os.path.basename(target_path)
                # 从processed文件名推断原始文件名
                original_name = target_name.replace('_processed', '')

                self.log.emit(f"📁 处理: {target_name}")

                wb = openpyxl.load_workbook(target_path)
                file_modified = False

                for sheet_name in wb.sheetnames:
                    if self._is_stopped:
                        break

                    # 尝试匹配
                    if has_file_col:
                        # 有文件名列：先尝试 (文件名, Sheet名) 匹配
                        key = (original_name, sheet_name)
                        if key not in policy_map:
                            # 尝试其他可能的文件名格式
                            key = None
                            for k in policy_map.keys():
                                if isinstance(k, tuple) and k[1] == sheet_name:
                                    key = k
                                    break
                        if not key or key not in policy_map:
                            continue
                    else:
                        # 无文件名列：直接用Sheet名匹配
                        if sheet_name not in policy_map:
                            continue
                        key = sheet_name

                    ws = wb[sheet_name]
                    numbers = policy_map[key]

                    # 查找并回填投保单号和保单号
                    filled = self._fill_policy_numbers(ws, numbers)
                    if filled:
                        file_modified = True
                        sheet_count += 1
                        self.log.emit(f"  ✅ {sheet_name}: 已回填")

                if file_modified:
                    wb.save(target_path)
                    success_count += 1

                wb.close()
                self.progress.emit(int((idx + 1) / total_files * 100))

            self.progress.emit(100)

            # 3. 针对恒力能源销售，更新汇总表添加开票备注列
            if self.customer_type == "恒力能源销售":
                self._add_invoice_remark_column()

            self.finished.emit(True, f"完成！处理 {success_count} 个文件，回填 {sheet_count} 个工作表")

        except Exception as e:
            import traceback
            self.finished.emit(False, f"出错:\n{traceback.format_exc()}")

    def _add_invoice_remark_column(self):
        """为恒力能源销售的汇总表添加开票备注列"""
        try:
            wb_summary = openpyxl.load_workbook(self.summary_path)
            ws_summary = wb_summary.active

            # 获取表头列索引
            headers = {}
            for c in range(1, ws_summary.max_column + 1):
                h = ws_summary.cell(1, c).value
                if h:
                    headers[h] = c

            # 找到投保人列和投保单号列
            comp_col = headers.get('投保人')
            policy_col = headers.get('投保单号')

            if not comp_col or not policy_col:
                self.log.emit("  ⚠️ 汇总表缺少投保人或投保单号列，跳过开票备注")
                wb_summary.close()
                return

            # 检查是否已有开票备注列
            remark_col = headers.get('开票备注')
            if not remark_col:
                # 在最后一列之后添加开票备注表头（动态获取列号）
                remark_col = ws_summary.max_column + 1
                ws_summary.cell(1, remark_col).value = "开票备注"
                ws_summary.cell(1, remark_col).font = Font(name='Songti SC', size=10, bold=True)
                ws_summary.cell(1, remark_col).alignment = Alignment(horizontal='center', vertical='center')
                ws_summary.column_dimensions[get_column_letter(remark_col)].width = 40
                self.log.emit("  📝 已添加'开票备注'列")

            # 遍历数据行，检查条件并填入备注
            remark_count = 0
            for r in range(2, ws_summary.max_row + 1):
                comp = ws_summary.cell(r, comp_col).value
                policy_no = ws_summary.cell(r, policy_col).value

                if not comp or not policy_no:
                    continue

                # 检查条件：投保人是"恒力华南石化销售有限公司"且投保单号第8-10位是"041"
                policy_str = str(policy_no)
                if comp == '恒力华南石化销售有限公司' and len(policy_str) >= 10:
                    # 投保单号第8-10位（1-based），即索引7-9（0-based）
                    if policy_str[7:10] == '041':
                        ws_summary.cell(r, remark_col).value = "备注请在完整保单号后加4个字  车船联运"
                        remark_count += 1

            if remark_count > 0:
                self.log.emit(f"  ✅ 已为 {remark_count} 行添加开票备注")

            wb_summary.save(self.summary_path)
            wb_summary.close()

        except Exception as e:
            self.log.emit(f"  ⚠️ 添加开票备注失败: {e}")

    def _fill_policy_numbers(self, ws, numbers):
        """在工作表中查找并回填投保单号和保单号"""
        filled = False
        policy_no = numbers.get('投保单号', '')
        cert_no = numbers.get('保单号', '')

        # 遍历查找投保单号和保单号的位置
        for r in range(1, ws.max_row + 1):
            for c in range(1, min(15, ws.max_column + 1)):
                cell_val = str(ws.cell(r, c).value or '')

                # 查找投保单号标签
                if '投保单号' in cell_val and policy_no:
                    # 找到投保单号右边的合并单元格并填入值
                    fill_col = self._find_fill_column(ws, r, c)
                    if fill_col:
                        ws.cell(r, fill_col).value = policy_no
                        filled = True

                # 查找保单号标签（多式联运使用）
                elif '保单号' in cell_val and cert_no:
                    fill_col = self._find_fill_column(ws, r, c)
                    if fill_col:
                        ws.cell(r, fill_col).value = cert_no
                        filled = True

        return filled

    def _find_fill_column(self, ws, row, label_col):
        """找到标签右边适合填入投保单号的合并单元格的起始列"""
        # 查找标签所在的合并单元格范围，确定标签的实际结束列
        label_end_col = label_col
        for merged_range in ws.merged_cells.ranges:
            if merged_range.min_row == row and merged_range.min_col <= label_col <= merged_range.max_col:
                label_end_col = merged_range.max_col
                break

        # 查找该行中所有在标签右边的合并单元格，按位置排序
        candidate_ranges = []
        for merged_range in ws.merged_cells.ranges:
            if merged_range.min_row == row and merged_range.min_col > label_end_col:
                candidate_ranges.append(merged_range)

        if candidate_ranges:
            # 按 min_col 排序，从最近的开始检查
            candidate_ranges.sort(key=lambda x: x.min_col)

            for merged_range in candidate_ranges:
                cell_val = ws.cell(merged_range.min_row, merged_range.min_col).value
                # 跳过包含"不含税金额"或"税额"等标签的单元格
                if cell_val and isinstance(cell_val, str):
                    if '不含税' in cell_val or '税额' in cell_val:
                        continue
                # 返回第一个可用的合并单元格（空的、或已有投保单号值的）
                return merged_range.min_col

        # 如果没有合适的合并单元格，返回标签结束列的下一列
        return label_end_col + 1


class WordGenWorker(QThread):
    """生成投保单的工作线程（支持 Word 无签和 PDF 有签合并，支持恒力混合生成）"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(bool, str, str)

    def __init__(self, template_path, data_list, output_dir, extra_args=None, output_format='word'):
        super().__init__()
        self.template_path = template_path # 可能是空字符串
        self.data_list = data_list
        self.output_dir = output_dir
        self.extra_args = extra_args or {}
        self.output_format = output_format  # 'word' or 'pdf'
        self._is_stopped = False

    def stop(self):
        self._is_stopped = True

    def run(self):
        try:
            total_files = 0
            msg_list = []
            
            # 获取申报周期
            period_str = self.extra_args.get('period', '')

            # === 任务 1：恒力/惠州付款通知书 (如果选择了模板且是恒力/惠州模式) ===
            if self.extra_args.get('customer_type') in ("恒力PTA", "惠州PTA"):
                notice_template = self.extra_args.get('notice_template')
                if notice_template:
                    self.log.emit("🚀 开始生成恒力付款通知书...")
                    count = self._generate_hengli_docs(self.data_list, notice_template, self.output_dir, period_str)
                    total_files += count
                    msg_list.append(f"付款通知书及对账单: {count} 个")

            # === 任务 2：通用投保单生成 (只要有 Policy 模板就生成) ===
            # 恒力模式下，如果选了投保单模板，也会进入这里，按 Excel 内容逐行填充
            if self.template_path and os.path.exists(self.template_path):
                self.log.emit("🚀 开始生成投保单...")
                count = self._generate_policy_docs(period_str)
                total_files += count
                msg_list.append(f"投保单: {count} 个")
            
            final_msg = "处理完成！\n" + "\n".join(msg_list)
            self.progress.emit(100)
            self.finished.emit(True, final_msg, self.output_dir)
            
        except Exception as e:
            import traceback
            self.finished.emit(False, f"出错:\n{traceback.format_exc()}", self.output_dir)

    def _generate_policy_docs(self, period_str):
        """标准投保单生成逻辑"""
        import tempfile
        import shutil
        import calendar

        total = len(self.data_list)
        success_count = 0
        pdf_files = []

        # PDF模式：使用临时目录，最后只保留合并文件
        # Word模式：保存到Word投保单子文件夹
        if self.output_format == 'pdf':
            target_dir = tempfile.mkdtemp(prefix="policy_pdf_")
        else:
            target_dir = os.path.join(self.output_dir, "Word投保单")
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)

        # 计算申报周期的起始和截止日期
        period_start = ""
        period_end = ""
        if period_str:
            try:
                # 从 "2025年01月" 格式解析
                match = re.match(r'(\d{4})年(\d{1,2})月', period_str)
                if match:
                    year = int(match.group(1))
                    month = int(match.group(2))
                    period_start = f"{year}/{month}/1"
                    last_day = calendar.monthrange(year, month)[1]
                    period_end = f"{year}/{month}/{last_day}"
            except (ValueError, AttributeError):
                pass

        for idx, data in enumerate(self.data_list):
            if self._is_stopped: break
            sheet_name = data.get('sheet_name', f'Unknown_{idx}')

            try:
                doc = Document(self.template_path)
                rate_val = data.get('rate', 0)
                rate_permille_str = f"{rate_val * 1000:.6f}".rstrip('0').rstrip('.')

                # 兼容恒力数据里的字段名
                ship_name = str(data.get('ship_voyage', '') or '')
                if not ship_name and 'transport_tool' in data:
                    ship_name = str(data.get('transport_tool', '') or '')
                if not ship_name and 'no' in data:  # 恒力能源销售的车船号
                    ship_name = str(data.get('no', '') or '')

                # 查找协议编号（恒力能源销售）
                comp_name = str(data.get('comp', '') or '')
                agreement_code = ""
                for key, code in AGREEMENT_CODES.items():
                    if key in comp_name or comp_name in key:
                        agreement_code = code
                        break

                # 恒力能源销售的保费和金额使用不同字段
                premium_val = data.get('new_premium', data.get('prem', 0))
                insurance_amount = data.get('insurance_amount', data.get('money', 0))

                # 计算标记值（从实际数据的日期中提取年月）
                customer_type = self.extra_args.get('customer_type', '')
                if customer_type == "恒力能源销售":
                    # 恒力能源销售：申报周期：发货日期-申报止期
                    ship_date_mark = format_date_slashes(data.get('date', ''))
                    end_date_mark = data.get('latest_date', '')
                    mark_value = f"申报周期：{ship_date_mark}-{end_date_mark}"
                else:
                    # 惠州PTA/恒力PTA/多式联运：从起运日期提取年月
                    date_str = str(data.get('departure_date', '') or data.get('date', '') or '')
                    import re as re_mark
                    match_mark = re_mark.match(r'(\d{4})[/\-.](\d{1,2})', date_str)
                    if match_mark:
                        mark_value = f"申报周期：{match_mark.group(1)}年{int(match_mark.group(2)):02d}月"
                    elif period_str:
                        mark_value = f"申报周期：{period_str}"
                    else:
                        mark_value = "申报周期："

                # 计算投保人值（优先使用数据中的comp字段，支持康辉等公司）
                if customer_type == "恒力能源销售":
                    policyholder = comp_name  # 申报公司名称
                elif customer_type == "惠州PTA":
                    policyholder = data.get('comp') or "恒力石化（惠州）有限公司"
                elif customer_type == "恒力PTA":
                    # 优先使用从Excel识别的康辉公司名称，否则使用默认值
                    policyholder = data.get('comp') or "恒力石化（大连）有限公司"
                else:
                    policyholder = data.get('comp') or "浙江卓航多式联运科技有限公司"

                context = {
                    "{Sheet名}": sheet_name,
                    "{船名/航次}": ship_name,
                    "{车船号}": str(data.get('no', '') or ship_name),
                    "{业务笔数}": str(data.get('business_count', 0)),
                    "{起运日期}": str(data.get('departure_date', '') or format_date_slashes(data.get('date', ''))),
                    "{发货日期}": format_date_slashes(data.get('date', '')),
                    "{货种}": str(data.get('cargo_type', '') or data.get('mat', '') or ''),
                    "{物料名称}": str(data.get('mat', '') or data.get('cargo_type', '') or ''),
                    "{保险货物描述}": str(data.get('cargo_type', '') or data.get('mat', '') or ''),
                    "{实载吨位}": f"{data.get('tonnage', data.get('amt', 0)):,.3f}",
                    "{装货数量（吨）}": f"{data.get('tonnage', data.get('amt', 0)):,.2f}",
                    "{装货数量}": f"{data.get('tonnage', data.get('amt', 0)):,.2f}",
                    "{开单量}": f"{data.get('amt', 0):,.3f}",
                    "{保险金额}": f"{insurance_amount:,.2f}",
                    "{金额}": f"{insurance_amount:,.2f}",
                    "{金额（元）}": f"{insurance_amount:,.2f}",  # 全角括号
                    "{金额(元)}": f"{insurance_amount:,.2f}",   # 半角括号
                    "{总保额}": f"{insurance_amount:,.2f}",
                    "{综合费率}": f"{rate_val*100:.6f}%",
                    "{千分费率}": rate_permille_str,
                    "{综合费率（‰）}": rate_permille_str,
                    "{保费}": f"{premium_val:,.2f}",
                    "{保费（元）}": f"{premium_val:,.2f}",  # 全角括号
                    "{保费(元)}": f"{premium_val:,.2f}",   # 半角括号
                    "{总保费}": f"{premium_val:,.2f}",
                    "{保费大写}": cn_currency(premium_val),
                    "{非标准化特约}": str(data.get('special_terms', '') or ''),
                    "{特约}": str(data.get('special_terms', '') or ''),
                    "{申报周期}": period_str,
                    "{起始日期}": period_start,
                    "{截止日期}": period_end,
                    "{申报止期}": format_date_slashes(data.get('latest_date', '')),
                    "{申报公司名称}": comp_name,
                    "{申报公司}": comp_name,
                    "{协议编号}": agreement_code,
                    "{标记}": mark_value,
                    "{投保人}": policyholder,
                }
                
                # 多式联运需要保留原始字体大小
                preserve_font = (customer_type == "多式联运")
                for p in doc.paragraphs: replace_text_in_paragraph(p, context, preserve_font_size=preserve_font)
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: replace_text_in_paragraph(p, context, preserve_font_size=preserve_font)
                
                # 文件名格式: 投保人_Sheet名 (避免重名)
                safe_sheet = sheet_name.replace('/', '_').replace('\\', '_')
                safe_sheet = re.sub(r'[<>:"/\\|?*]', '_', safe_sheet)
                safe_policyholder = policyholder.replace('/', '_').replace('\\', '_')
                safe_policyholder = re.sub(r'[<>:"/\\|?*]', '_', safe_policyholder)
                safe_name = f"{safe_policyholder}_{safe_sheet}"

                if self.output_format == 'word':
                    remove_images_from_doc(doc)
                    out_path = os.path.join(target_dir, f"{safe_name}.docx")
                    doc.save(out_path)
                    success_count += 1
                    self.log.emit(f"  ✅ 生成Word: {safe_name}.docx")
                elif self.output_format == 'pdf':
                    temp_docx = os.path.join(target_dir, f"_temp_{safe_name}.docx")
                    doc.save(temp_docx)
                    self.log.emit(f"  📄 正在转换PDF: {safe_name}...")
                    pdf_path = os.path.join(target_dir, f"{safe_name}.pdf")
                    if convert_to_pdf(temp_docx, pdf_path):
                        pdf_files.append(pdf_path)
                        success_count += 1
                        self.log.emit(f"  ✅ PDF转换成功: {safe_name}.pdf")
                    else:
                        self.log.emit(f"  ❌ PDF转换失败: {safe_name} (请确保已安装LibreOffice)")
                    if os.path.exists(temp_docx): os.remove(temp_docx)

            except Exception as e:
                self.log.emit(f"    ❌ 失败 {sheet_name}: {e}")
            
            self.progress.emit(int((idx + 1) / total * 90))
        
        # PDF 合并
        if self.output_format == 'pdf':
            if pdf_files:
                self.log.emit(f"📑 正在合并 {len(pdf_files)} 个投保单 PDF...")
                try:
                    merger = PdfMerger()
                    for p in pdf_files:
                        if os.path.exists(p):
                            merger.append(p)
                    timestamp = QDate.currentDate().toString("yyyyMMdd")
                    merged_name = f"投保单合并_{timestamp}.pdf"
                    merged_path = os.path.join(self.output_dir, merged_name)
                    merger.write(merged_path)
                    merger.close()
                    self.log.emit(f"✅ 已合并保存: {merged_name}")
                    self.log.emit(f"📁 保存位置: {self.output_dir}")

                    # 删除临时目录和所有临时文件
                    try:
                        shutil.rmtree(target_dir)
                    except (OSError, PermissionError):
                        pass  # 文件可能被占用

                    # 返回1表示只生成了1个合并文件
                    return 1
                except Exception as e:
                    self.log.emit(f"❌ 合并失败: {e}")
                    # 清理临时目录
                    try:
                        shutil.rmtree(target_dir)
                    except (OSError, PermissionError):
                        pass  # 文件可能被占用
            else:
                self.log.emit("⚠️ 没有成功生成任何投保单PDF，请检查LibreOffice是否已正确安装")
                self.log.emit("💡 提示：可以先选择Word格式生成，确认内容正确后再手动转换为PDF")

        return success_count

    def _generate_hengli_docs(self, data_list, template_path, output_dir, period_str):
        """生成恒力付款通知书 - 增强版：支持Sheet名匹配填充保费"""
        # 从配置中获取对账单分组规则
        customer_type = self.extra_args.get('customer_type', '')
        config_manager = get_config_manager()
        customer_config = config_manager.get_customer_config(customer_type)

        # 构建分组字典
        groups = {}
        statement_groups_config = {}
        if customer_config and customer_config.statement_groups:
            statement_groups_config = customer_config.statement_groups
            for group_name in statement_groups_config.keys():
                groups[group_name] = {'data': [], 'sum': 0}
        else:
            # 默认分组（兼容旧配置）
            groups = {'PTA船运': {'data': [], 'sum': 0}, 'BA': {'data': [], 'sum': 0}, 'PTA车运': {'data': [], 'sum': 0}}
            statement_groups_config = {
                'PTA船运': {'match_type': 'contains', 'patterns': ['PTA船运']},
                'BA': {'match_type': 'contains', 'patterns': ['BA']},
                'PTA车运': {'match_type': 'contains', 'patterns': ['PTA车运']}
            }

        grand_total = 0
        month_str = ""

        # 尝试从申报周期提取月份（如 "2025年10月" -> "10月"）
        if period_str and "年" in period_str and "月" in period_str:
             month_str = period_str.split('年')[1]

        # 匹配函数
        def match_group(sheet_name, rule):
            match_type = rule.get('match_type', 'contains')
            patterns = rule.get('patterns', [])
            if not patterns or (len(patterns) == 1 and patterns[0] == ''):
                return True  # 空模式匹配所有
            if match_type == 'contains':
                return any(p in sheet_name for p in patterns if p)
            elif match_type == 'contains_any':
                return any(p in sheet_name for p in patterns if p)
            elif match_type == 'startswith':
                return any(sheet_name.startswith(p) for p in patterns if p)
            elif match_type == 'endswith':
                return any(sheet_name.endswith(p) for p in patterns if p)
            return False

        for d in data_list:
            sheet = d['sheet_name']
            premium = d.get('new_premium', 0)
            grand_total += premium
            # 如果没设申报周期，尝试从数据提取
            if not month_str and d.get('departure_date'):
                try: month_str = str(d['departure_date']).split('/')[1] + "月"
                except (IndexError, AttributeError): pass

            # 使用配置的分组规则进行匹配
            matched = False
            for key, rule in statement_groups_config.items():
                if match_group(sheet, rule):
                    groups[key]['data'].append(d)
                    groups[key]['sum'] += premium
                    matched = True
                    break

            # 如果没有匹配到任何分组，放入第一个分组
            if not matched and groups:
                first_key = list(groups.keys())[0]
                groups[first_key]['data'].append(d)
                groups[first_key]['sum'] += premium

        start_date = self.extra_args.get('deadline_date', '')
        end_date = self.extra_args.get('issue_date', '')

        # 日期格式化：2025-12-21 -> "2025年12月21日"
        def fmt_date_chinese(ymd):
            try:
                y, m, d = ymd.split('-')
                return f"{y}年{int(m)}月{int(d)}日"
            except (ValueError, AttributeError):
                return ymd

        # 日期格式化（旧格式，用于兼容）：2025-12-21 -> "2025    12    21"
        def fmt_date_spaced(ymd):
            try:
                y, m, d = ymd.split('-')
                return f"{y}    {int(m)}    {int(d)}"
            except (ValueError, AttributeError):
                return ymd

        d_deadline = fmt_date_chinese(start_date)  # 往来单位对账单用中文格式
        d_issue = fmt_date_chinese(end_date)       # 往来单位对账单用中文格式
        d_deadline_spaced = fmt_date_spaced(start_date)  # 付款通知书用空格格式（兼容）
        d_issue_spaced = fmt_date_spaced(end_date)       # 付款通知书用空格格式（兼容）

        # 准备各分类通知书文档（在内存中处理，不保存单独文件）
        notice_docs = []  # 保存每个分类通知书的Document对象

        for key, info in groups.items():
            if info['data']:
                doc = Document(template_path)
                total_premium = info['sum']
                total_premium_cn = cn_currency(total_premium)

                # === 在付款通知书模版中查找Sheet名并填入保费 ===
                self._fill_payment_notice_by_sheet_match(doc, key, total_premium, total_premium_cn)

                # 原有的占位符替换逻辑保持不变
                context = {
                    "{月}": month_str,
                    "{总保费}": f"{total_premium:,.2f}",
                    "{大写总保费}": total_premium_cn,
                    "{截止日期}": d_deadline,
                    "{落款日期}": d_issue,
                    "{运输工具}": f"{key.replace('运','')}/{key.replace('运','')}",
                    "{申报周期}": period_str,
                }

                for p in doc.paragraphs:
                    replace_text_in_paragraph(p, context)
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_text_in_paragraph(p, context)

                # 应用标准字体（中文宋体，英文Times New Roman）到标题以下的所有内容
                self._apply_font_to_content(doc)

                # 只在内存中保存Document对象，不保存文件
                notice_docs.append((key, doc))

        # 生成合并对账单 - 直接使用内存中的Document对象
        merged_doc = self._create_merged_statement(
            notice_docs, grand_total,
            grand_total_upper := cn_currency(grand_total),
            d_deadline, d_issue, period_str
        )

        # 保存合并对账单
        word_name = "合并对账单.docx"
        word_path = os.path.join(output_dir, word_name)
        merged_doc.save(word_path)

        # 如果是PDF模式，将合并对账单转换为PDF
        if self.output_format == 'pdf':
            self.log.emit("📑 正在将合并对账单转换为 PDF...")
            pdf_name = "合并对账单.pdf"
            pdf_path = os.path.join(output_dir, pdf_name)
            if convert_to_pdf(word_path, pdf_path):
                self.log.emit(f"✅ 已生成: {pdf_name}")
                # 删除Word版本，只保留PDF
                try:
                    os.remove(word_path)
                except (OSError, PermissionError):
                    pass  # 文件可能被占用
            else:
                self.log.emit(f"❌ PDF转换失败，保留Word版本")

        return 1  # 只生成1个文件

    def _create_merged_statement(self, notice_docs, grand_total,
                                   grand_total_upper, d_deadline, d_issue, period_str):
        """
        创建合并对账单文档：
        - 前3页：根据分组名称从模板中找到对应页面
        - 最后一页：往来单位对账单
        """
        # 创建新文档
        merged_doc = Document()

        # 1. 添加各分类通知书的对应页面内容
        for idx, (key, notice_doc) in enumerate(notice_docs):
            # 直接使用内存中的Document对象
            # 关键修改：找到模板中包含该分组名称的"收款事由"所在的页面

            from docx.oxml.ns import qn

            # 将文档按分页符分割成多个页面
            pages = []  # 每个元素是一个页面的elements列表
            current_page = []

            for element in notice_doc.element.body:
                is_page_break = False
                if element.tag.endswith('p'):
                    # 检查段落中是否包含分页符
                    for child in element:
                        if child.tag.endswith('pPr'):
                            for prop in child:
                                if prop.tag.endswith('pageBreakBefore'):
                                    is_page_break = True
                                    break
                        if child.tag.endswith('r'):  # run
                            for run_child in child:
                                if run_child.tag.endswith('br'):
                                    br_type = run_child.get(qn('w:type'))
                                    if br_type == 'page':
                                        is_page_break = True
                                        break

                    if is_page_break:
                        # 保存当前页面，开始新页面
                        if current_page:
                            pages.append(current_page)
                        current_page = []
                        continue

                current_page.append(element)

            # 保存最后一个页面
            if current_page:
                pages.append(current_page)

            # 找到包含该分组名称（如'PTA船运'、'BA'、'PTA车运'）的页面
            target_page_idx = 0  # 默认第一页
            for page_idx, page_elements in enumerate(pages):
                for element in page_elements:
                    if element.tag.endswith('p'):
                        # 获取段落文本
                        text = ''.join(node.text or '' for node in element.iter() if node.text)
                        # 检查是否包含收款事由和分组名称
                        if '收款事由' in text and key in text:
                            target_page_idx = page_idx
                            break
                if target_page_idx == page_idx and page_idx > 0:
                    break

            # 复制目标页面的内容
            if target_page_idx < len(pages):
                for element in pages[target_page_idx]:
                    if element.tag.endswith('p'):
                        new_para = merged_doc.add_paragraph()
                        new_para._element.getparent().replace(new_para._element, copy.deepcopy(element))
                    elif element.tag.endswith('tbl'):
                        merged_doc.element.body.append(copy.deepcopy(element))

            # 在每个分类通知书页面后添加分页符
            if idx < len(notice_docs) - 1:
                merged_doc.add_page_break()

        # 2. 添加最后一页的往来单位对账单
        merged_doc.add_page_break()

        # 从第一个notice_doc中提取模板（因为都是基于同一个模板生成的）
        if notice_docs:
            # 使用第一个文档作为模板来获取对账单页面
            first_notice_doc = notice_docs[0][1]
            self._add_statement_page_to_merged(merged_doc, first_notice_doc, grand_total,
                                               grand_total_upper, d_deadline, d_issue)

        # 3. 设置合并对账单的标题为二号字体（22pt）
        self._set_merged_doc_title_font(merged_doc)

        return merged_doc

    def _set_merged_doc_title_font(self, doc):
        """
        将合并对账单的标题设置为二号字体（22pt）
        只设置真正的标题：'付款通知书&收据' 和 '往来单位对账单'
        不影响其他包含这些关键词的普通文本
        """
        for para in doc.paragraphs:
            text = para.text.strip()
            # 只有当段落几乎完全是标题时才设置二号字体
            # 避免把"实际收到保费后本付款通知书可视同收据生效"这样的句子也改了
            if text == '付款通知书&收据' or text == '往来单位对账单' or \
               (len(text) < 20 and ('付款通知书' in text and '&' in text and '收据' in text)):
                for run in para.runs:
                    run.font.size = Pt(22)  # 二号 = 22pt
                    run.font.bold = True

    def _apply_font_to_content(self, doc):
        """
        将文档中所有内容（包括标题）应用标准字体格式（中文宋体，英文Times New Roman）
        """
        for para in doc.paragraphs:
            for run in para.runs:
                # 保留原有的字号和粗体设置
                original_size = run.font.size
                original_bold = run.font.bold

                # 应用标准字体
                set_run_font_standard(run, 14)

                # 恢复原有的字号（如果有）
                if original_size:
                    run.font.size = original_size
                # 恢复粗体设置
                if original_bold:
                    run.font.bold = original_bold

        # 处理表格中的内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            original_size = run.font.size
                            original_bold = run.font.bold
                            set_run_font_standard(run, 14)
                            if original_size:
                                run.font.size = original_size
                            if original_bold:
                                run.font.bold = original_bold

    def _add_statement_page_to_merged(self, merged_doc, template_doc, grand_total,
                                       grand_total_upper, d_deadline, d_issue):
        """
        从模板中提取往来单位对账单页面，并添加到合并文档中
        """
        # 查找模板中包含"往来单位对账单"的页面
        found_statement_section = False

        for element in template_doc.element.body:
            # 检查段落中是否包含"往来单位对账单"
            if element.tag.endswith('p'):
                # 正确获取段落文本
                para_text = ''.join(node.text for node in element.iter() if node.text)
                if '往来单位对账单' in para_text:
                    found_statement_section = True

                if found_statement_section:
                    # 复制段落到合并文档 - 修复方法
                    new_para = merged_doc.add_paragraph()
                    new_para._element.getparent().replace(new_para._element, copy.deepcopy(element))

            # 复制表格
            elif element.tag.endswith('tbl') and found_statement_section:
                new_table_element = copy.deepcopy(element)
                merged_doc.element.body.append(new_table_element)

        # 填充对账单内容
        self._fill_statement_page(merged_doc, grand_total, grand_total_upper, d_deadline, d_issue)

        # 应用字体格式
        for para in merged_doc.paragraphs:
            if '往来单位对账单' not in para.text:  # 标题除外
                for run in para.runs:
                    set_run_font_standard(run, 14)

    def _fill_payment_notice_by_sheet_match(self, doc, sheet_key, premium_amount, premium_cn):
        """
        在付款通知书模版中查找匹配的Sheet名（如'PTA船运'），
        并在对应分类的'人民币（大写）：￥ '后填入金额和大写
        注意：模板中"人民币（大写）"字段在"收款事由"之前，所以需要向前搜索
        """
        # 先找到包含sheet_key的"收款事由"段落的索引
        found_section_idx = -1
        paragraphs = list(doc.paragraphs)

        for i, para in enumerate(paragraphs):
            text = para.text
            if '收款事由' in text and sheet_key in text:
                found_section_idx = i
                break

        # 如果找到了对应分类，从该段落往前找"人民币（大写）：￥"（因为金额字段在收款事由之前）
        if found_section_idx >= 0:
            for i in range(found_section_idx, -1, -1):  # 向前搜索
                para = paragraphs[i]
                text = para.text

                if '人民币（大写）：￥' in text or '人民币(大写)：￥' in text:
                    # 构建替换文本：金额数字 + 空格 + 大写金额
                    amount_text = f"{premium_amount:,.2f} {premium_cn}"

                    # 使用保留格式的替换
                    if '人民币（大写）：￥ ' in text:
                        replace_text_preserve_format(para, '人民币（大写）：￥ ', f'人民币（大写）：￥ {amount_text}')
                    elif '人民币(大写)：￥ ' in text:
                        replace_text_preserve_format(para, '人民币(大写)：￥ ', f'人民币(大写)：￥ {amount_text}')

                    return  # 找到并处理完成后退出

        # 如果段落中没找到，检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text

                        if '人民币（大写）：￥' in text or '人民币(大写)：￥' in text:
                            amount_text = f"{premium_amount:,.2f} {premium_cn}"

                            if '人民币（大写）：￥ ' in text:
                                replace_text_preserve_format(para, '人民币（大写）：￥ ', f'人民币（大写）：￥ {amount_text}')
                            elif '人民币(大写)：￥ ' in text:
                                replace_text_preserve_format(para, '人民币(大写)：￥ ', f'人民币(大写)：￥ {amount_text}')

                            return

    def _fill_statement_page(self, doc, total_amount, total_amount_cn, deadline_date, issue_date):
        """
        在付款通知书最后一页的'往来单位对账单'中填入日期和金额
        """
        from docx.oxml.shared import OxmlElement

        # 遍历文档查找"往来单位对账单"
        found_statement = False

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text

            # 找到"往来单位对账单"标题
            if '往来单位对账单' in text:
                found_statement = True
                continue

            # 在找到对账单后，查找需要填充的字段
            if found_statement:
                # 填充截止日期：查找包含"截止"的行
                if '截止' in text and '年' in text and '月' in text and '日' in text:
                    # 尝试多种日期占位符格式
                    new_text = text
                    # 格式1: 截止____年____月____日
                    new_text = re.sub(r'截止____年____月____日', f'截止{deadline_date}', new_text)
                    # 格式2: 截止    年    月    日
                    new_text = re.sub(r'截止\s+年\s+月\s+日', f'截止{deadline_date}', new_text)
                    # 格式3: 截止 年 月 日
                    new_text = re.sub(r'截止\s*年\s*月\s*日', f'截止{deadline_date}', new_text)
                    # 格式4: 已有日期的情况
                    new_text = re.sub(r'截止\d+年\d+月\d+日', f'截止{deadline_date}', new_text)

                    if new_text != text:
                        replace_text_preserve_format(para, text, new_text)

                # 填充落款日期：查找以"日"结尾的行（但不包含"截止"、"金额"、"余额"、"本对账单"）
                elif (text.strip().endswith('日') and '年' in text and '月' in text and
                      '截止' not in text and '金额' not in text and '余额' not in text and
                      '本对账单' not in text):
                    # 尝试多种日期占位符格式
                    new_text = text
                    # 格式1: ____年____月____日
                    new_text = re.sub(r'____年____月____日', issue_date, new_text)
                    # 格式2:     年    月    日
                    new_text = re.sub(r'\s+年\s+月\s+日', issue_date, new_text)
                    # 格式3: 年 月 日
                    new_text = re.sub(r'\s*年\s*月\s*日', issue_date, new_text)
                    # 格式4: 已有日期的情况
                    new_text = re.sub(r'\d+年\d+月\d+日', issue_date, new_text)

                    if new_text != text:
                        replace_text_preserve_format(para, text, new_text)

                # 填充金额：查找"金额"或"余额"字样
                if ('金额' in text or '余额' in text) and '￥' in text:
                    # 找到￥的位置，分离前后部分
                    parts = text.split('￥', 1)
                    if len(parts) == 2:
                        before_yuan = parts[0] + '￥'

                        # 构建金额文本（只用两个空格）
                        amount_text = f"{total_amount:,.2f}  {total_amount_cn}"

                        # 清空段落并重新构建，只对金额部分添加下划线
                        para.clear()

                        # 添加￥之前的部分（不带下划线）
                        run1 = para.add_run(before_yuan)
                        set_run_font_standard(run1, 14)

                        # 添加金额部分（带下划线）
                        run2 = para.add_run(amount_text)
                        set_run_font_standard(run2, 14)
                        run2.font.underline = True  # 只给金额添加下划线

        # 同样处理表格中的对账单
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text

                    # 检查是否包含"往来单位对账单"
                    if '往来单位对账单' in cell_text:
                        # 在这个表格中查找并填充
                        for para in cell.paragraphs:
                            text = para.text

                            # 填充截止日期
                            if '截止' in text and '年' in text and '月' in text and '日' in text:
                                new_text = text
                                new_text = re.sub(r'截止____年____月____日', f'截止{deadline_date}', new_text)
                                new_text = re.sub(r'截止\s+年\s+月\s+日', f'截止{deadline_date}', new_text)
                                new_text = re.sub(r'截止\s*年\s*月\s*日', f'截止{deadline_date}', new_text)
                                new_text = re.sub(r'截止\d+年\d+月\d+日', f'截止{deadline_date}', new_text)

                                if new_text != text:
                                    replace_text_preserve_format(para, text, new_text)

                            # 填充落款日期
                            elif (text.strip().endswith('日') and '年' in text and '月' in text and
                                  '截止' not in text and '金额' not in text and '余额' not in text and
                                  '本对账单' not in text):
                                new_text = text
                                new_text = re.sub(r'____年____月____日', issue_date, new_text)
                                new_text = re.sub(r'\s+年\s+月\s+日', issue_date, new_text)
                                new_text = re.sub(r'\s*年\s*月\s*日', issue_date, new_text)
                                new_text = re.sub(r'\d+年\d+月\d+日', issue_date, new_text)

                                if new_text != text:
                                    replace_text_preserve_format(para, text, new_text)

                            # 填充金额
                            if ('金额' in text or '余额' in text) and '￥' in text:
                                parts = text.split('￥', 1)
                                if len(parts) == 2:
                                    before_yuan = parts[0] + '￥'
                                    amount_text = f"{total_amount:,.2f}  {total_amount_cn}"

                                    para.clear()

                                    # ￥之前的部分（不带下划线）
                                    run1 = para.add_run(before_yuan)
                                    set_run_font_standard(run1, 14)

                                    # 金额部分（带下划线）
                                    run2 = para.add_run(amount_text)
                                    set_run_font_standard(run2, 14)
                                    run2.font.underline = True

# ==========================================
# UI 样式 - Anthropic 官方设计系统
# ==========================================

class AnthropicColors:
    """Anthropic 官方色彩系统"""
    # 背景色
    BG_PRIMARY = "#faf9f5"      # 主背景/奶油白
    BG_CARD = "#f0eee6"         # 卡片背景/浅米色
    BG_MINT = "#bcd1ca"         # 特殊卡片/薄荷绿
    BG_LAVENDER = "#cbcadb"     # 特殊卡片/淡紫色
    BG_DARK = "#141413"         # 深色区域

    # 强调色
    ACCENT = "#d97757"          # 主强调色/陶土色
    ACCENT_DARK = "#c6613f"     # 次强调色/深赭红
    ACCENT_HOVER = "#e8956f"    # 悬停色

    # 文字色
    TEXT_PRIMARY = "#141413"    # 主要文字
    TEXT_SECONDARY = "#b0aea5"  # 次要文字（仅用于装饰性文字）
    TEXT_MUTED = "#6b6a65"      # 中等对比度文字（用于按钮/标签）
    TEXT_LIGHT = "#faf9f5"      # 深色背景上的文字

    # 状态色
    SUCCESS = "#5a9a7a"         # 成功/绿色
    WARNING = "#d9a557"         # 警告/金色
    ERROR = "#c75050"           # 错误/红色
    INFO = "#5a7a9a"            # 信息/蓝灰

    # 边框色
    BORDER = "#e5e3db"          # 浅边框
    BORDER_DARK = "#d0cec6"     # 深边框


class AnthropicFonts:
    """Anthropic 字体配置"""
    # 标题字体
    TITLE_LARGE = ("Söhne", 28)
    TITLE = ("Söhne", 22)
    TITLE_SMALL = ("Söhne", 16)

    # 正文字体
    BODY = ("Söhne", 14)
    BODY_SMALL = ("Söhne", 12)

    # UI 元素
    BUTTON = ("Söhne", 14)
    LABEL = ("Söhne", 13)

    # 代码字体
    CODE = ("JetBrains Mono", 12)

    # 中文回退
    CN_FALLBACK = "PingFang SC"


STYLE_SHEET = f"""
/* ==========================================
   货运保险工具 ProMax - Anthropic 官方设计系统
   主色：{AnthropicColors.ACCENT} (陶土色)
   背景：{AnthropicColors.BG_PRIMARY} (奶油白)
   卡片：{AnthropicColors.BG_CARD} (浅米色)
   文字：{AnthropicColors.TEXT_PRIMARY} / {AnthropicColors.TEXT_MUTED}
   边框：{AnthropicColors.BORDER}
   ========================================== */

/* ========== 主窗口背景 - Anthropic奶油白 ========== */
QMainWindow {{
    background: {AnthropicColors.BG_PRIMARY};
}}
QWidget#centralWidget {{
    background: transparent;
}}

/* ========== 选项卡 - Anthropic风格 ========== */
QTabWidget::pane {{
    border: none;
    background: {AnthropicColors.BG_PRIMARY};
}}
QTabBar::tab {{
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.TEXT_SECONDARY};
    border: none;
    padding: 14px 40px;
    margin-right: 8px;
    border-radius: 8px 8px 0 0;
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
    font-weight: 600;
    min-width: 100px;
}}
QTabBar::tab:selected {{
    background: {AnthropicColors.BG_DARK};
    color: {AnthropicColors.TEXT_LIGHT};
}}
QTabBar::tab:hover:!selected {{
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.TEXT_PRIMARY};
}}

/* ========== 分组框/卡片 - Anthropic设计 ========== */
QGroupBox {{
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-weight: 600;
    font-size: 14px;
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 12px;
    margin-top: 24px;
    padding: 28px 24px 24px 24px;
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.TEXT_PRIMARY};
}}
QGroupBox::title {{
    subcontrol-origin: margin;
    left: 20px;
    padding: 6px 16px;
    color: {AnthropicColors.ACCENT};
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    font-weight: 600;
    font-size: 13px;
}}

/* ========== 按钮通用 - Anthropic浅色按钮 ========== */
QPushButton {{
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    padding: 12px 20px;
    font-size: 14px;
    font-weight: 500;
    background: {AnthropicColors.BG_PRIMARY};
    color: {AnthropicColors.TEXT_PRIMARY};
}}
QPushButton:hover {{
    background: {AnthropicColors.BG_CARD};
    border-color: {AnthropicColors.ACCENT};
    color: {AnthropicColors.ACCENT};
}}
QPushButton:pressed {{
    background: {AnthropicColors.BG_MINT};
}}
QPushButton:disabled {{
    background: {AnthropicColors.BG_CARD};
    border-color: {AnthropicColors.BORDER};
    color: {AnthropicColors.TEXT_SECONDARY};
}}

/* ========== 文件选择按钮 - Anthropic虚线设计 ========== */
QPushButton#fileBtn {{
    text-align: left;
    padding: 16px 20px;
    border: 2px dashed {AnthropicColors.BORDER};
    background: {AnthropicColors.BG_PRIMARY};
    border-radius: 12px;
    color: {AnthropicColors.TEXT_MUTED};
    font-weight: 500;
}}
QPushButton#fileBtn:hover {{
    border: 2px dashed {AnthropicColors.ACCENT};
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.ACCENT};
}}
QPushButton#fileBtn[selected="true"] {{
    background: {AnthropicColors.BG_MINT};
    color: {AnthropicColors.TEXT_PRIMARY};
    border: 2px solid {AnthropicColors.SUCCESS};
}}

/* ========== 主运行按钮 - Anthropic深色主按钮 ========== */
QPushButton#runBtn {{
    background: {AnthropicColors.BG_DARK};
    color: {AnthropicColors.TEXT_LIGHT};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-weight: 600;
    padding: 14px 32px;
    font-size: 14px;
    border: none;
    border-radius: 8px;
}}
QPushButton#runBtn:hover {{
    background: {AnthropicColors.ACCENT};
}}
QPushButton#runBtn:pressed {{
    background: {AnthropicColors.ACCENT_DARK};
}}
QPushButton#runBtn:disabled {{
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.TEXT_SECONDARY};
}}

/* ========== 强调按钮 - Anthropic陶土色 ========== */
QPushButton#accentBtn {{
    background: {AnthropicColors.ACCENT};
    color: {AnthropicColors.TEXT_LIGHT};
    font-weight: 600;
    border: none;
    border-radius: 8px;
    padding: 12px 24px;
}}
QPushButton#accentBtn:hover {{
    background: {AnthropicColors.ACCENT_DARK};
}}

/* ========== 进度条 - Anthropic简洁风格 ========== */
QProgressBar {{
    border: none;
    border-radius: 4px;
    background: {AnthropicColors.BG_CARD};
    height: 8px;
    color: transparent;
    text-align: center;
}}
QProgressBar::chunk {{
    background: {AnthropicColors.ACCENT};
    border-radius: 4px;
}}

/* ========== 文本编辑框(日志) - Anthropic风格 ========== */
QTextEdit {{
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 12px;
    padding: 16px;
    background: {AnthropicColors.BG_PRIMARY};
    font-family: 'JetBrains Mono', 'SF Mono', 'Menlo', 'PingFang SC', monospace;
    font-size: 12px;
    line-height: 1.6;
    color: {AnthropicColors.TEXT_PRIMARY};
    selection-background-color: rgba(217, 119, 87, 0.25);
}}

/* ========== 列表控件 - Anthropic列表设计 ========== */
QListWidget {{
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 12px;
    padding: 8px;
    background: {AnthropicColors.BG_PRIMARY};
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne Mono', 'SF Mono', 'Menlo', 'PingFang SC', monospace;
    font-size: 13px;
    outline: none;
}}
QListWidget::item {{
    padding: 10px 14px;
    border-radius: 6px;
    margin: 2px 0;
}}
QListWidget::item:hover {{
    background: {AnthropicColors.BG_CARD};
}}
QListWidget::item:selected {{
    background: {AnthropicColors.BG_MINT};
    color: {AnthropicColors.TEXT_PRIMARY};
}}

/* ========== 标签 - Anthropic文字设计 ========== */
QLabel {{
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}

/* ========== 单选按钮 - Anthropic风格 ========== */
QRadioButton {{
    color: {AnthropicColors.TEXT_PRIMARY};
    spacing: 10px;
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QRadioButton::indicator {{
    width: 18px;
    height: 18px;
    border-radius: 9px;
    border: 2px solid {AnthropicColors.BORDER};
    background: {AnthropicColors.BG_PRIMARY};
}}
QRadioButton::indicator:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QRadioButton::indicator:checked {{
    background: {AnthropicColors.ACCENT};
    border-color: {AnthropicColors.ACCENT};
}}

/* ========== 复选框 - Anthropic风格 ========== */
QCheckBox {{
    color: {AnthropicColors.TEXT_PRIMARY};
    spacing: 10px;
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QCheckBox::indicator {{
    width: 18px;
    height: 18px;
    border-radius: 4px;
    border: 2px solid {AnthropicColors.BORDER};
    background: {AnthropicColors.BG_PRIMARY};
}}
QCheckBox::indicator:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QCheckBox::indicator:checked {{
    background: {AnthropicColors.ACCENT};
    border-color: {AnthropicColors.ACCENT};
}}

/* ========== 下拉框 - Anthropic设计 ========== */
QComboBox {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    padding: 12px 16px;
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
    min-height: 20px;
}}
QComboBox:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QComboBox:focus {{
    border-color: {AnthropicColors.ACCENT};
}}
QComboBox::drop-down {{
    border: none;
    width: 28px;
}}
QComboBox::down-arrow {{
    image: none;
    border-left: 5px solid transparent;
    border-right: 5px solid transparent;
    border-top: 6px solid {AnthropicColors.TEXT_MUTED};
}}
QComboBox QAbstractItemView {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    selection-background-color: {AnthropicColors.BG_MINT};
    color: {AnthropicColors.TEXT_PRIMARY};
    padding: 6px;
    outline: none;
}}
QComboBox QAbstractItemView::item {{
    padding: 10px 14px;
    border-radius: 4px;
}}
QComboBox QAbstractItemView::item:hover {{
    background: {AnthropicColors.BG_CARD};
}}

/* ========== 日期选择框 - Anthropic设计 ========== */
QDateEdit {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    padding: 12px 16px;
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QDateEdit:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QDateEdit:focus {{
    border-color: {AnthropicColors.ACCENT};
}}
QDateEdit::drop-down {{
    border: none;
    width: 28px;
}}

/* ========== 滚动区域 ========== */
QScrollArea {{
    border: none;
    background: transparent;
}}
QScrollArea > QWidget > QWidget {{
    background: transparent;
}}

/* ========== 滚动条 - Anthropic细窄设计 ========== */
QScrollBar:vertical {{
    background: {AnthropicColors.BG_CARD};
    width: 8px;
    border-radius: 4px;
    margin: 2px;
}}
QScrollBar::handle:vertical {{
    background: {AnthropicColors.BORDER_DARK};
    border-radius: 4px;
    min-height: 40px;
}}
QScrollBar::handle:vertical:hover {{
    background: {AnthropicColors.ACCENT};
}}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
    height: 0;
}}
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
    background: none;
}}

QScrollBar:horizontal {{
    background: {AnthropicColors.BG_CARD};
    height: 8px;
    border-radius: 4px;
    margin: 2px;
}}
QScrollBar::handle:horizontal {{
    background: {AnthropicColors.BORDER_DARK};
    border-radius: 4px;
    min-width: 40px;
}}
QScrollBar::handle:horizontal:hover {{
    background: {AnthropicColors.ACCENT};
}}
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
    width: 0;
}}

/* ========== 工具提示 - Anthropic设计 ========== */
QToolTip {{
    background: {AnthropicColors.BG_DARK};
    color: {AnthropicColors.TEXT_LIGHT};
    border: none;
    border-radius: 6px;
    padding: 8px 12px;
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 12px;
}}

/* ========== 消息框 - Anthropic设计 ========== */
QMessageBox {{
    background: {AnthropicColors.BG_PRIMARY};
}}
QMessageBox QLabel {{
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QMessageBox QPushButton {{
    min-width: 90px;
    padding: 10px 20px;
}}

/* ========== 输入框 - Anthropic设计 ========== */
QLineEdit {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    padding: 12px 16px;
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QLineEdit:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QLineEdit:focus {{
    border-color: {AnthropicColors.ACCENT};
}}

/* ========== SpinBox - Anthropic设计 ========== */
QSpinBox, QDoubleSpinBox {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 8px;
    padding: 10px 14px;
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
    font-size: 14px;
}}
QSpinBox:hover, QDoubleSpinBox:hover {{
    border-color: {AnthropicColors.ACCENT};
}}
QSpinBox:focus, QDoubleSpinBox:focus {{
    border-color: {AnthropicColors.ACCENT};
}}

/* ========== 表格控件 - Anthropic设计 ========== */
QTableWidget {{
    background: {AnthropicColors.BG_PRIMARY};
    border: 1px solid {AnthropicColors.BORDER};
    border-radius: 12px;
    gridline-color: {AnthropicColors.BORDER};
    color: {AnthropicColors.TEXT_PRIMARY};
    font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
}}
QTableWidget::item {{
    padding: 10px;
}}
QTableWidget::item:selected {{
    background: {AnthropicColors.BG_MINT};
    color: {AnthropicColors.TEXT_PRIMARY};
}}
QHeaderView::section {{
    background: {AnthropicColors.BG_CARD};
    color: {AnthropicColors.TEXT_PRIMARY};
    padding: 12px;
    border: none;
    border-bottom: 1px solid {AnthropicColors.BORDER};
    font-weight: 600;
}}
"""

# ==========================================
# UI 组件 - Anthropic 风格
# ==========================================

class GlassCard(QFrame):
    """Anthropic 风格的卡片组件 - 带柔和阴影"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setStyleSheet(f"""
            GlassCard {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 12px;
            }}
        """)
        # 添加柔和阴影
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(20)
        shadow.setColor(QColor(0, 0, 0, 25))
        shadow.setOffset(0, 4)
        self.setGraphicsEffect(shadow)


class AnimatedBorderFrame(QFrame):
    """Anthropic 风格的边框容器 - 陶土色渐变动画"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._glow_intensity = 0.0
        self._animation_direction = 1
        self._gradient_offset = 0.0

        # 设置定时器驱动动画
        self._timer = QTimer(self)
        self._timer.timeout.connect(self._update_glow)
        self._timer.start(50)  # 20fps

        self.setStyleSheet("background: transparent;")

    def _update_glow(self):
        # 更新发光强度 (脉冲效果)
        self._glow_intensity += 0.02 * self._animation_direction
        if self._glow_intensity >= 1.0:
            self._animation_direction = -1
        elif self._glow_intensity <= 0.3:
            self._animation_direction = 1

        # 渐变偏移动画
        self._gradient_offset += 0.006
        if self._gradient_offset >= 1.0:
            self._gradient_offset = 0.0

        self.update()

    def paintEvent(self, event):
        super().paintEvent(event)
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)

        rect = self.rect().adjusted(4, 4, -4, -4)

        # Anthropic 陶土色渐变 #d97757 → #e8956f → #c6613f
        base_alpha = int(50 + 60 * self._glow_intensity)  # 50-110 (更柔和)

        # 主边框颜色 - Anthropic陶土色渐变
        if self._gradient_offset < 0.5:
            t = self._gradient_offset * 2
            r = int(217 + (232 - 217) * t)   # 217→232
            g = int(119 + (149 - 119) * t)   # 119→149
            b = int(87 + (111 - 87) * t)     # 87→111
        else:
            t = (self._gradient_offset - 0.5) * 2
            r = int(232 + (198 - 232) * t)   # 232→198
            g = int(149 + (97 - 149) * t)    # 149→97
            b = int(111 + (63 - 111) * t)    # 111→63

        border_color = QColor(r, g, b, base_alpha)
        pen = QPen(border_color)
        pen.setWidth(2)
        painter.setPen(pen)
        painter.drawRoundedRect(rect, 12, 12)

        # 外发光效果
        glow_alpha = int(20 * self._glow_intensity)
        glow_color = QColor(217, 119, 87, glow_alpha)  # Anthropic ACCENT
        glow_pen = QPen(glow_color)
        glow_pen.setWidth(3)
        painter.setPen(glow_pen)
        painter.drawRoundedRect(rect.adjusted(-2, -2, 2, 2), 14, 14)


class AnimatedButton(QPushButton):
    """Anthropic 风格的普通按钮"""
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setStyleSheet(f"""
            AnimatedButton {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 12px 20px;
                font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
                font-size: 14px;
                font-weight: 500;
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            AnimatedButton:hover {{
                background: {AnthropicColors.BG_CARD};
                border-color: {AnthropicColors.ACCENT};
                color: {AnthropicColors.ACCENT};
            }}
            AnimatedButton:pressed {{
                background: {AnthropicColors.BG_MINT};
            }}
            AnimatedButton:disabled {{
                background: {AnthropicColors.BG_CARD};
                border-color: {AnthropicColors.BORDER};
                color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)


class AnimatedRunButton(QPushButton):
    """Anthropic 风格的主操作按钮 - 深色背景"""
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setStyleSheet(f"""
            AnimatedRunButton {{
                background: {AnthropicColors.BG_DARK};
                border: none;
                border-radius: 8px;
                padding: 14px 32px;
                font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
                font-size: 14px;
                font-weight: 600;
                color: {AnthropicColors.TEXT_LIGHT};
                min-width: 120px;
            }}
            AnimatedRunButton:hover {{
                background: {AnthropicColors.ACCENT};
            }}
            AnimatedRunButton:pressed {{
                background: {AnthropicColors.ACCENT_DARK};
            }}
            AnimatedRunButton:disabled {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)

    def stop_animation(self):
        """保留接口兼容性"""
        pass


class AccentButton(QPushButton):
    """Anthropic 风格的强调按钮 - 陶土色背景"""
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setStyleSheet(f"""
            AccentButton {{
                background: {AnthropicColors.ACCENT};
                border: none;
                border-radius: 8px;
                padding: 12px 24px;
                font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
                font-size: 14px;
                font-weight: 600;
                color: {AnthropicColors.TEXT_LIGHT};
            }}
            AccentButton:hover {{
                background: {AnthropicColors.ACCENT_DARK};
            }}
            AccentButton:pressed {{
                background: #b85535;
            }}
            AccentButton:disabled {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)
        # 添加柔和阴影
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(12)
        shadow.setColor(QColor(217, 119, 87, 60))
        shadow.setOffset(0, 2)
        self.setGraphicsEffect(shadow)

# ==========================================
# 页面类
# ==========================================

class MainPage(QWidget):
    def __init__(self):
        super().__init__()
        self.excel_files = []
        self.worker = None
        self.extracted_data = []
        self._init_ui()

    def _init_ui(self):
        # 主布局
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)

        # 滚动区域内容容器
        scroll_content = QWidget()
        layout = QVBoxLayout(scroll_content)
        layout.setSpacing(10)
        layout.setContentsMargins(14, 14, 14, 14)

        # === 文件选择区 ===
        g1 = QGroupBox("📁 选择 Excel 文件")
        v1 = QVBoxLayout()
        v1.setSpacing(8)
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)
        self.btn_add = AnimatedButton("➕ 添加文件")
        self.btn_add.clicked.connect(self._add_files)
        self.btn_clear = AnimatedButton("🗑️ 清空")
        self.btn_clear.clicked.connect(self._clear_files)
        btn_layout.addWidget(self.btn_add)
        btn_layout.addWidget(self.btn_clear)
        btn_layout.addStretch()
        v1.addLayout(btn_layout)
        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(80)
        self.file_list.setMaximumHeight(100)
        v1.addWidget(self.file_list)
        g1.setLayout(v1)
        layout.addWidget(g1)

        # === 客户类型选择 ===
        g_cust = QGroupBox("👤 客户类型")
        v_cust = QVBoxLayout()
        v_cust.setSpacing(8)
        self.combo_customer = QComboBox()
        self._load_customer_types()
        v_cust.addWidget(self.combo_customer)
        # 单号标签选择（放在客户类型下方）
        h_policy = QHBoxLayout()
        h_policy.setSpacing(12)
        h_policy.setContentsMargins(0, 4, 0, 0)
        lbl_policy = QLabel("单号标签:")
        lbl_policy.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 12px;")
        self.policy_label_group = QButtonGroup(self)
        self.radio_policy_no = QRadioButton("投保单号")
        self.radio_cert_no = QRadioButton("保单号")
        self.radio_policy_no.setChecked(True)
        self.policy_label_group.addButton(self.radio_policy_no)
        self.policy_label_group.addButton(self.radio_cert_no)
        h_policy.addWidget(lbl_policy)
        h_policy.addWidget(self.radio_policy_no)
        h_policy.addWidget(self.radio_cert_no)
        h_policy.addStretch()
        v_cust.addLayout(h_policy)
        g_cust.setLayout(v_cust)
        layout.addWidget(g_cust)

        # === 操作按钮区 ===
        g_action = QGroupBox("⚡ 操作")
        v_action = QVBoxLayout()
        v_action.setSpacing(10)
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        self.btn_run = AnimatedRunButton("🚀 开始处理")
        self.btn_run.setObjectName("runBtn")
        self.btn_run.clicked.connect(self._run)
        self.btn_stop = AnimatedButton("⏹️ 停止")
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self._stop)
        self.btn_backfill = AnimatedButton("📝 回填单号")
        self.btn_backfill.clicked.connect(self._backfill_policy_numbers)
        btn_row.addWidget(self.btn_run)
        btn_row.addWidget(self.btn_stop)
        btn_row.addWidget(self.btn_backfill)
        v_action.addLayout(btn_row)
        g_action.setLayout(v_action)
        layout.addWidget(g_action)

        # === 进度区 ===
        g_progress = QGroupBox("📊 处理进度")
        v_progress = QVBoxLayout()
        v_progress.setSpacing(8)
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        v_progress.addWidget(self.progress_bar)
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMinimumHeight(100)
        self.log_area.setMaximumHeight(150)
        self.log_area.setPlaceholderText("处理日志将显示在这里...")
        v_progress.addWidget(self.log_area)
        g_progress.setLayout(v_progress)
        layout.addWidget(g_progress)

        layout.addStretch()

        # 设置滚动区域
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)

    def _add_files(self):
        # 确保事件循环被处理，解决macOS上文件对话框有时无法点击的问题
        QApplication.processEvents()
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择 Excel",
            "",
            "Excel (*.xlsx *.xls)"
        )
        if files:
            for f in files:
                if f not in self.excel_files:
                    self.excel_files.append(f)
                    self.file_list.addItem(os.path.basename(f))

    def _clear_files(self):
        self.excel_files.clear()
        self.file_list.clear()

    def _run(self):
        if not self.excel_files: return
        output_dir = os.path.dirname(self.excel_files[0])
        self.btn_run.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.progress_bar.setValue(0)
        self.log_area.clear()
        policy_label = "投保单号" if self.radio_policy_no.isChecked() else "保单号"
        self.worker = ProcessWorker(self.excel_files, output_dir, self.combo_customer.currentText(), policy_label)
        self.worker.log.connect(self.log_area.append)
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _stop(self):
        if self.worker: self.worker.stop()

    def _on_finished(self, success, msg, data):
        self.btn_run.setEnabled(True)
        self.btn_run.stop_animation()
        self.btn_stop.setEnabled(False)
        self.extracted_data = data
        if success: QMessageBox.information(self, "完成", msg)
        elif msg != "已停止": QMessageBox.critical(self, "错误", msg)

    def _load_customer_types(self):
        """从配置加载客户类型到下拉框"""
        config_manager = get_config_manager()
        current = self.combo_customer.currentText()
        self.combo_customer.clear()
        customer_types = config_manager.get_customer_types()
        if customer_types:
            self.combo_customer.addItems(customer_types)
        else:
            # 后备默认值
            self.combo_customer.addItems(["多式联运", "恒力PTA", "恒力能源销售"])
        # 尝试恢复之前的选择
        if current:
            idx = self.combo_customer.findText(current)
            if idx >= 0:
                self.combo_customer.setCurrentIndex(idx)

    def refresh_customer_types(self):
        """刷新客户类型（供设置页面调用）"""
        self._load_customer_types()

    def _backfill_policy_numbers(self):
        """回填单号功能"""
        QApplication.processEvents()

        # 1. 选择汇总表
        summary_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择汇总表 Excel",
            "",
            "Excel (*.xlsx)"
        )
        if not summary_path:
            return

        # 2. 选择要回填的processed文件（可多选）
        target_files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择要回填的 Processed Excel 文件（可多选）",
            os.path.dirname(summary_path),
            "Excel (*.xlsx)"
        )
        if not target_files:
            return

        # 3. 执行回填
        self.btn_backfill.setEnabled(False)
        self.btn_run.setEnabled(False)
        self.progress_bar.setValue(0)
        self.log_area.clear()
        self.log_area.append("📝 开始回填单号...")

        self.backfill_worker = PolicyBackfillWorker(
            summary_path,
            target_files,
            self.combo_customer.currentText()
        )
        self.backfill_worker.log.connect(self.log_area.append)
        self.backfill_worker.progress.connect(self.progress_bar.setValue)
        self.backfill_worker.finished.connect(self._on_backfill_finished)
        self.backfill_worker.start()

    def _on_backfill_finished(self, success, msg):
        """回填完成回调"""
        self.btn_backfill.setEnabled(True)
        self.btn_run.setEnabled(True)
        if success:
            QMessageBox.information(self, "完成", msg)
        else:
            QMessageBox.critical(self, "错误", msg)


class PdfPage(QWidget):
    def __init__(self):
        super().__init__()
        self.excel_path = None
        self.worker = None
        self._init_ui()

    def _init_ui(self):
        # 主布局
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)

        # 滚动区域内容容器
        scroll_content = QWidget()
        layout = QVBoxLayout(scroll_content)
        layout.setSpacing(10)
        layout.setContentsMargins(14, 14, 14, 14)

        # === 文件选择 ===
        g1 = QGroupBox("📁 选择 Excel 文件")
        v1 = QVBoxLayout()
        v1.setSpacing(6)
        self.btn_excel = QPushButton("📊 点击选择 Excel 文件")
        self.btn_excel.setObjectName("fileBtn")
        self.btn_excel.clicked.connect(self._select_excel)
        v1.addWidget(self.btn_excel)
        g1.setLayout(v1)
        layout.addWidget(g1)

        # === 导出模式 ===
        g_mode = QGroupBox("📤 导出模式")
        v_mode = QVBoxLayout()
        v_mode.setSpacing(8)
        self.mode_group = QButtonGroup(self)
        self.radio_per_sheet = QRadioButton("📄 按 Sheet 逐个导出")
        self.radio_grouped = QRadioButton("📁 按公司分组导出")
        self.radio_per_sheet.setChecked(True)
        self.mode_group.addButton(self.radio_per_sheet)
        self.mode_group.addButton(self.radio_grouped)
        v_mode.addWidget(self.radio_per_sheet)
        v_mode.addWidget(self.radio_grouped)
        g_mode.setLayout(v_mode)
        layout.addWidget(g_mode)

        # === 打印方向 ===
        g2 = QGroupBox("🖨️ 打印方向")
        v2 = QVBoxLayout()
        v2.setSpacing(8)
        self.btn_group = QButtonGroup(self)
        self.radio_landscape = QRadioButton("↔️ 横向打印")
        self.radio_portrait = QRadioButton("↕️ 纵向打印")
        self.radio_landscape.setChecked(True)
        self.btn_group.addButton(self.radio_landscape)
        self.btn_group.addButton(self.radio_portrait)
        v2.addWidget(self.radio_landscape)
        v2.addWidget(self.radio_portrait)
        g2.setLayout(v2)
        layout.addWidget(g2)

        # === 操作按钮 ===
        g_action = QGroupBox("⚡ 操作")
        v_action = QVBoxLayout()
        v_action.setSpacing(10)
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        self.btn_run = AnimatedRunButton("📥 导出 PDF")
        self.btn_run.setObjectName("runBtn")
        self.btn_run.clicked.connect(self._run)
        self.btn_stop = AnimatedButton("⏹️ 停止")
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self._stop)
        btn_row.addWidget(self.btn_run)
        btn_row.addWidget(self.btn_stop)
        v_action.addLayout(btn_row)
        g_action.setLayout(v_action)
        layout.addWidget(g_action)

        # === 进度区 ===
        g_progress = QGroupBox("📊 导出进度")
        v_progress = QVBoxLayout()
        v_progress.setSpacing(8)
        self.progress_bar = QProgressBar()
        v_progress.addWidget(self.progress_bar)
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMinimumHeight(80)
        self.log_area.setMaximumHeight(120)
        self.log_area.setPlaceholderText("导出日志将显示在这里...")
        v_progress.addWidget(self.log_area)
        g_progress.setLayout(v_progress)
        layout.addWidget(g_progress)

        layout.addStretch()

        # 设置滚动区域
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)

    def _select_excel(self):
        QApplication.processEvents()
        path, _ = QFileDialog.getOpenFileName(self, "Excel", "", "Excel (*.xlsx *.xls)")
        if path:
            self.excel_path = path
            self.btn_excel.setText(os.path.basename(path))
            self.btn_excel.setProperty("selected", "true")
            self.btn_excel.style().unpolish(self.btn_excel)
            self.btn_excel.style().polish(self.btn_excel)

    def _run(self):
        if not self.excel_path: return
        self.btn_run.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.progress_bar.setValue(0)
        self.log_area.clear()
        output_dir = os.path.dirname(self.excel_path)
        orientation = 'landscape' if self.radio_landscape.isChecked() else 'portrait'

        # 根据导出模式选择不同的 Worker
        if self.radio_grouped.isChecked():
            self.worker = PdfGroupedExportWorker(self.excel_path, output_dir, orientation)
        else:
            self.worker = PdfExportWorker(self.excel_path, output_dir, orientation)

        self.worker.log.connect(self.log_area.append)
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _stop(self):
        if self.worker: self.worker.stop()

    def _on_finished(self, success, msg, output_dir):
        self.btn_run.setEnabled(True)
        self.btn_run.stop_animation()
        self.btn_stop.setEnabled(False)
        if success: show_completion_dialog(self, "完成", msg, output_dir)
        elif msg != "已停止": QMessageBox.critical(self, "错误", msg)

class WordPage(QWidget):
    def __init__(self, main_page):
        super().__init__()
        self.main_page = main_page
        self.template_path = None
        self.notice_template_path = None
        self.excel_path = None
        self.extracted_data = []
        self.worker = None
        self._init_ui()

    def _init_ui(self):
        # 主布局
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)

        # 滚动区域内容容器
        scroll_content = QWidget()
        layout = QVBoxLayout(scroll_content)
        layout.setSpacing(10)
        layout.setContentsMargins(14, 14, 14, 14)

        # === 模板选择 ===
        g1 = QGroupBox("📄 选择模板（右键取消选择）")
        v1 = QVBoxLayout()
        v1.setSpacing(8)
        self.btn_template = QPushButton("📋 投保单模板")
        self.btn_template.setObjectName("fileBtn")
        self.btn_template.clicked.connect(self._select_template)
        self.btn_template.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.btn_template.customContextMenuRequested.connect(lambda pos: self._show_template_context_menu(pos, 'template'))
        v1.addWidget(self.btn_template)
        self.btn_notice_template = QPushButton("💰 付款通知书模板")
        self.btn_notice_template.setObjectName("fileBtn")
        self.btn_notice_template.clicked.connect(self._select_notice_template)
        self.btn_notice_template.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.btn_notice_template.customContextMenuRequested.connect(lambda pos: self._show_template_context_menu(pos, 'notice'))
        v1.addWidget(self.btn_notice_template)
        g1.setLayout(v1)
        layout.addWidget(g1)

        # === 数据来源 ===
        g_data = QGroupBox("📊 数据来源")
        v_data = QVBoxLayout()
        v_data.setSpacing(8)
        self.radio_from_main = QRadioButton("🔗 使用对账处理页面数据")
        self.radio_from_excel = QRadioButton("📁 选择 Excel 文件")
        self.radio_from_main.setChecked(True)
        self.radio_from_excel.toggled.connect(self._on_radio_changed)
        v_data.addWidget(self.radio_from_main)
        v_data.addWidget(self.radio_from_excel)
        self.btn_excel = QPushButton("📊 选择 Excel 文件")
        self.btn_excel.setObjectName("fileBtn")
        self.btn_excel.clicked.connect(self._select_excel)
        self.btn_excel.setEnabled(False)
        v_data.addWidget(self.btn_excel)
        g_data.setLayout(v_data)
        layout.addWidget(g_data)

        # === 输出格式 ===
        g_fmt = QGroupBox("📤 输出格式")
        h_fmt = QHBoxLayout()
        h_fmt.setSpacing(12)
        self.group_fmt = QButtonGroup(self)
        self.radio_word = QRadioButton("📝 Word (无签名)")
        self.radio_pdf = QRadioButton("📕 PDF (含签名)")
        self.radio_word.setChecked(True)
        self.group_fmt.addButton(self.radio_word)
        self.group_fmt.addButton(self.radio_pdf)
        h_fmt.addWidget(self.radio_word)
        h_fmt.addWidget(self.radio_pdf)
        h_fmt.addStretch()
        g_fmt.setLayout(h_fmt)
        layout.addWidget(g_fmt)

        # === 日期设置 ===
        g_date = QGroupBox("📅 付款通知书日期")
        v_date = QVBoxLayout()
        v_date.setSpacing(10)

        # 日期行1
        date_row1 = QHBoxLayout()
        date_row1.setSpacing(8)
        date_row1.addWidget(QLabel("周期:"))
        self.date_period = QDateEdit()
        self.date_period.setCalendarPopup(True)
        self.date_period.setDate(QDate.currentDate())
        self.date_period.setDisplayFormat("yyyy年MM月")
        date_row1.addWidget(self.date_period)
        date_row1.addStretch()
        v_date.addLayout(date_row1)

        # 日期行2
        date_row2 = QHBoxLayout()
        date_row2.setSpacing(8)
        date_row2.addWidget(QLabel("截止:"))
        self.date_deadline = QDateEdit()
        self.date_deadline.setCalendarPopup(True)
        self.date_deadline.setDate(QDate.currentDate())
        self.date_deadline.setDisplayFormat("yyyy-MM-dd")
        date_row2.addWidget(self.date_deadline)
        date_row2.addWidget(QLabel("落款:"))
        self.date_issue = QDateEdit()
        self.date_issue.setCalendarPopup(True)
        self.date_issue.setDate(QDate.currentDate())
        self.date_issue.setDisplayFormat("yyyy-MM-dd")
        date_row2.addWidget(self.date_issue)
        date_row2.addStretch()
        v_date.addLayout(date_row2)

        g_date.setLayout(v_date)
        layout.addWidget(g_date)

        # === 操作按钮 ===
        g_action = QGroupBox("⚡ 操作")
        v_action = QVBoxLayout()
        v_action.setSpacing(10)
        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)
        self.btn_run = AnimatedRunButton("📝 生成投保单")
        self.btn_run.setObjectName("runBtn")
        self.btn_run.clicked.connect(self._run)
        self.btn_stop = AnimatedButton("⏹️ 停止")
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self._stop)
        btn_row.addWidget(self.btn_run)
        btn_row.addWidget(self.btn_stop)
        v_action.addLayout(btn_row)
        g_action.setLayout(v_action)
        layout.addWidget(g_action)

        # === 进度区 ===
        g_progress = QGroupBox("📊 生成进度")
        v_progress = QVBoxLayout()
        v_progress.setSpacing(8)
        self.progress_bar = QProgressBar()
        v_progress.addWidget(self.progress_bar)
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMinimumHeight(70)
        self.log_area.setMaximumHeight(100)
        self.log_area.setPlaceholderText("生成日志将显示在这里...")
        v_progress.addWidget(self.log_area)
        g_progress.setLayout(v_progress)
        layout.addWidget(g_progress)

        layout.addStretch()

        # 设置滚动区域
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area)

    def _on_radio_changed(self, checked):
        self.btn_excel.setEnabled(checked)

    def _select_template(self):
        QApplication.processEvents()
        path, _ = QFileDialog.getOpenFileName(self, "Word", "", "Word (*.docx)")
        if path:
            self.template_path = path
            self.btn_template.setText(os.path.basename(path))
            self.btn_template.setProperty("selected", "true")
            self.btn_template.style().unpolish(self.btn_template)
            self.btn_template.style().polish(self.btn_template)

    def _select_notice_template(self):
        QApplication.processEvents()
        path, _ = QFileDialog.getOpenFileName(self, "Word", "", "Word (*.docx)")
        if path:
            self.notice_template_path = path
            self.btn_notice_template.setText(os.path.basename(path))
            self.btn_notice_template.setProperty("selected", "true")
            self.btn_notice_template.style().unpolish(self.btn_notice_template)
            self.btn_notice_template.style().polish(self.btn_notice_template)

    def _show_template_context_menu(self, pos, template_type):
        """显示模板选择的右键菜单"""
        menu = QMenu(self)
        menu.setStyleSheet(f"""
            QMenu {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 4px;
            }}
            QMenu::item {{
                padding: 8px 20px;
                border-radius: 4px;
            }}
            QMenu::item:selected {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.ACCENT};
            }}
        """)

        if template_type == 'template':
            if self.template_path:
                clear_action = menu.addAction("❌ 取消选择投保单模板")
                clear_action.triggered.connect(self._clear_template)
            else:
                menu.addAction("ℹ️ 尚未选择模板").setEnabled(False)
        elif template_type == 'notice':
            if self.notice_template_path:
                clear_action = menu.addAction("❌ 取消选择付款通知书模板")
                clear_action.triggered.connect(self._clear_notice_template)
            else:
                menu.addAction("ℹ️ 尚未选择模板").setEnabled(False)

        # 获取发送信号的按钮
        sender_btn = self.btn_template if template_type == 'template' else self.btn_notice_template
        menu.exec(sender_btn.mapToGlobal(pos))

    def _clear_template(self):
        """取消选择投保单模板"""
        self.template_path = None
        self.btn_template.setText("📋 投保单模板")
        self.btn_template.setProperty("selected", "false")
        self.btn_template.style().unpolish(self.btn_template)
        self.btn_template.style().polish(self.btn_template)
        self.log_area.append("ℹ️ 已取消选择投保单模板")

    def _clear_notice_template(self):
        """取消选择付款通知书模板"""
        self.notice_template_path = None
        self.btn_notice_template.setText("💰 付款通知书模板")
        self.btn_notice_template.setProperty("selected", "false")
        self.btn_notice_template.style().unpolish(self.btn_notice_template)
        self.btn_notice_template.style().polish(self.btn_notice_template)
        self.log_area.append("ℹ️ 已取消选择付款通知书模板")

    def _select_excel(self):
        QApplication.processEvents()  # 解决macOS上文件对话框有时无法点击的问题
        path, _ = QFileDialog.getOpenFileName(self, "Excel", "", "Excel (*.xlsx *.xls)")
        if path:
            self.excel_path = path
            self.btn_excel.setText(f"📊 {os.path.basename(path)}")
            self.log_area.append(f"已选择: {os.path.basename(path)}")
            self._load_excel(path)

    def _load_excel(self, excel_path):
        """加载汇总表格式的 Excel 文件，将每一行转换为投保单数据"""
        try:
            self.log_area.append("📊 正在加载汇总表...")
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            self.extracted_data = []
            
            # 尝试查找"汇总"sheet，否则使用第一个 sheet
            ws = None
            for sn in wb.sheetnames:
                if '汇总' in sn:
                    ws = wb[sn]
                    self.log_area.append(f"  📄 使用工作表: {sn}")
                    break
            if ws is None:
                ws = wb.active
                self.log_area.append(f"  📄 使用工作表: {ws.title}")
            
            # 读取表头（第一行）
            headers = []
            for col in range(1, ws.max_column + 1):
                val = ws.cell(1, col).value
                headers.append(str(val).strip() if val else "")
            
            # 建立列名到索引的映射
            col_map = {}
            for idx, h in enumerate(headers):
                col_map[h] = idx + 1  # 1-indexed
            
            # 从第 2 行开始读取数据，直到遇到"合计"或空行
            for row in range(2, ws.max_row + 1):
                first_cell = ws.cell(row, 1).value
                if first_cell is None or str(first_cell).strip() == "":
                    continue
                if "合计" in str(first_cell):
                    break
                
                # 提取各字段（支持多式联运和恒力PTA两种格式）
                def get_val(possible_names):
                    for name in possible_names:
                        if name in col_map:
                            return ws.cell(row, col_map[name]).value
                    return None
                
                sheet_name = get_val(["Sheet名", "sheet名", "Sheet"]) or f"数据_{row-1}"
                
                # 解析费率
                rate_val = 0
                rate_raw = get_val(["综合费率", "费率"])
                rate_permille = get_val(["千分费率"])
                if rate_raw:
                    try:
                        rate_val = float(rate_raw)
                    except (ValueError, TypeError):
                        pass
                elif rate_permille:
                    try:
                        rate_val = float(str(rate_permille).replace('‰', '')) / 1000
                    except (ValueError, TypeError):
                        pass

                # 解析数值字段
                def parse_num(val):
                    if val is None:
                        return 0
                    if isinstance(val, (int, float)):
                        return float(val)
                    try:
                        return float(str(val).replace(',', '').replace(' ', ''))
                    except (ValueError, TypeError):
                        return 0
                
                # 获取运输工具字段（用于恒力PTA和多式联运）
                transport_tool_val = get_val(["运输工具", "船名/航次", "船名航次"]) or ''

                # 恒力能源销售专用字段（汇总表列名为"投保人"）
                comp_val = get_val(["投保人", "申报公司名称", "申报公司"]) or ''
                no_val = get_val(["车船号"]) or ''
                date_val = get_val(["发货日期"]) or ''
                latest_date_val = get_val(["申报止期"]) or ''
                mat_val = get_val(["物料名称"]) or ''
                amt_val = parse_num(get_val(["开单量"]))
                money_val = parse_num(get_val(["金额（元）", "金额(元)", "金额"]))
                prem_val = parse_num(get_val(["保费（元）", "保费(元)", "保费", "总保费", "人民币保费"]))

                data = {
                    'sheet_name': str(sheet_name) if sheet_name else f"数据_{row-1}",
                    'file_name': os.path.basename(excel_path),
                    'ship_voyage': get_val(["船名/航次", "船名航次"]) or transport_tool_val or no_val,  # 多式联运用
                    'transport_tool': transport_tool_val,  # 恒力PTA用
                    'business_count': int(parse_num(get_val(["业务笔数"]))),
                    'departure_date': str(get_val(["起运日期"]) or date_val or ''),
                    'cargo_type': get_val(["货种", "保险货物描述"]) or mat_val or '',
                    'tonnage': parse_num(get_val(["实载吨位", "装货数量（吨）", "装货数量"])) or amt_val,
                    'insurance_amount': parse_num(get_val(["保险金额"])) or money_val,
                    'rate': rate_val,
                    'new_premium': prem_val,
                    'special_terms': get_val(["非标准化特约", "特约"]) or '',
                    # 恒力能源销售专用字段
                    'comp': comp_val,
                    'no': no_val,
                    'date': date_val,
                    'latest_date': latest_date_val,
                    'mat': mat_val,
                    'amt': amt_val,
                    'money': money_val,
                    'prem': prem_val,
                }
                
                self.extracted_data.append(data)
                self.log_area.append(f"  ✅ 行 {row}: {data['sheet_name']}")
            
            wb.close()
            self.log_area.append(f"✅ 共加载 {len(self.extracted_data)} 条数据")
        except Exception as e:
            import traceback
            self.log_area.append(f"❌ 加载失败: {e}\n{traceback.format_exc()}")

    def _run(self):
        # 1. 获取数据
        if self.radio_from_main.isChecked():
            data_list = self.main_page.extracted_data
            if not data_list:
                QMessageBox.warning(self, "提示", "请先在「对账处理」页面处理 Excel 文件！")
                return
            # 使用对账处理页面的客户类型
            customer_type = self.main_page.combo_customer.currentText()
        else:
            data_list = self.extracted_data
            if not data_list:
                QMessageBox.warning(self, "提示", "请先选择 Excel 文件！")
                return
            # 从Excel导入数据时，根据数据特征判断客户类型
            # 如果数据中有'运输工具'字段且包含'PTA'或'BA'，则为恒力PTA
            customer_type = self._detect_customer_type(data_list)

        # 2. 检查模板 (逻辑修改：允许任选其一)
        has_template = bool(self.template_path)
        has_notice = bool(self.notice_template_path)

        if not has_template and not has_notice:
            QMessageBox.warning(self, "提示", "请至少选择一个模板文件（投保单模板 或 付款通知书模板）！")
            return

        # 3. 选择输出目录
        QApplication.processEvents()  # 解决macOS上文件对话框有时无法点击的问题
        output_dir = QFileDialog.getExistingDirectory(self, "选择保存目录")
        if not output_dir:
            return

        self.btn_run.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.log_area.clear()

        # 4. 收集额外参数
        extra_args = {
            'customer_type': customer_type,  # 使用检测到的客户类型
            'notice_template': self.notice_template_path,
            'deadline_date': self.date_deadline.date().toString("yyyy-MM-dd"),
            'issue_date': self.date_issue.date().toString("yyyy-MM-dd"),
            'period': self.date_period.text()
        }

        fmt = 'pdf' if self.radio_pdf.isChecked() else 'word'

        # 如果没选 template_path，传空字符串
        self.worker = WordGenWorker(self.template_path if self.template_path else "", data_list, output_dir, extra_args, output_format=fmt)
        self.worker.log.connect(self.log_area.append)
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _detect_customer_type(self, data_list):
        """
        根据数据特征自动检测客户类型
        """
        if not data_list:
            return "多式联运"

        # 检查前几条数据的特征
        for data in data_list[:5]:
            sheet_name = data.get('sheet_name', '')
            transport_tool = data.get('transport_tool', '')
            comp_name = data.get('comp', '')

            # 恒力能源销售：包含特定公司名称
            if any(key in comp_name for key in AGREEMENT_CODES.keys()):
                return "恒力能源销售"
            if any(keyword in sheet_name for keyword in ['能源苏州', '华南石化', '精细化工', '油品销售', '恒力石化']):
                return "恒力能源销售"

            # 恒力PTA：包含PTA、BA等关键词
            if any(keyword in sheet_name for keyword in ['PTA船运', 'PTA车运', 'BA']):
                return "恒力PTA"
            if any(keyword in transport_tool for keyword in ['PTA', 'BA']):
                return "恒力PTA"

        # 默认为多式联运
        return "多式联运"

    def _stop(self):
        if self.worker: self.worker.stop()

    def _on_finished(self, success, msg, output_dir):
        self.btn_run.setEnabled(True)
        self.btn_run.stop_animation()
        self.btn_stop.setEnabled(False)
        if success: show_completion_dialog(self, "完成", msg, output_dir)
        elif msg != "已停止": QMessageBox.critical(self, "错误", msg)

class DonateDialog(QDialog):
    """捐赠对话框 - Anthropic 风格"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("支持作者")
        self.setFixedSize(420, 520)
        self.setStyleSheet(f"""
            QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; }}
            QPushButton {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT}; border: none; border-radius: 8px;
                padding: 12px 24px; font-weight: 600; font-size: 14px;
            }}
            QPushButton:hover {{ background: {AnthropicColors.ACCENT}; }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(30, 30, 30, 30)

        # 标题
        title = QLabel("💝 支持作者")
        title.setStyleSheet(f'''
            color: {AnthropicColors.ACCENT};
            font-size: 22px; font-weight: bold;
            font-family: 'Söhne', 'SF Pro Display', -apple-system, 'PingFang SC', sans-serif;
        ''')
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        # 描述
        desc = QLabel("如果这个工具对您有帮助，欢迎请作者喝杯咖啡 ☕")
        desc.setStyleSheet(f'color: {AnthropicColors.TEXT_MUTED}; font-size: 13px;')
        desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # 二维码区域
        qr_layout = QHBoxLayout()
        qr_layout.setSpacing(30)

        # 微信支付
        wechat_box = QVBoxLayout()
        wechat_label = QLabel("微信支付")
        wechat_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #07C160;')
        wechat_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        wechat_box.addWidget(wechat_label)

        wechat_qr = QLabel("[ 微信二维码 ]")
        wechat_qr.setFixedSize(140, 140)
        wechat_qr.setAlignment(Qt.AlignmentFlag.AlignCenter)
        wechat_qr.setStyleSheet(f'''
            font-size: 14px; background-color: {AnthropicColors.BG_CARD};
            border-radius: 12px; border: 1px solid {AnthropicColors.BORDER};
        ''')
        wechat_box.addWidget(wechat_qr)
        qr_layout.addLayout(wechat_box)

        # 支付宝
        alipay_box = QVBoxLayout()
        alipay_label = QLabel("支付宝")
        alipay_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #1677FF;')
        alipay_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        alipay_box.addWidget(alipay_label)

        alipay_qr = QLabel("[ 支付宝二维码 ]")
        alipay_qr.setFixedSize(140, 140)
        alipay_qr.setAlignment(Qt.AlignmentFlag.AlignCenter)
        alipay_qr.setStyleSheet(f'''
            font-size: 14px; background-color: {AnthropicColors.BG_CARD};
            border-radius: 12px; border: 1px solid {AnthropicColors.BORDER};
        ''')
        alipay_box.addWidget(alipay_qr)
        qr_layout.addLayout(alipay_box)

        layout.addLayout(qr_layout)

        # 感谢语
        thanks = QLabel("感谢您的支持！🙏")
        thanks.setStyleSheet(f'''
            color: {AnthropicColors.TEXT_PRIMARY}; font-size: 16px;
            font-weight: 600; padding: 15px 0 5px 0;
        ''')
        thanks.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(thanks)

        # 作者信息
        author = QLabel("Made with ❤️ by Dachi Yijin")
        author.setStyleSheet(f'color: {AnthropicColors.TEXT_SECONDARY}; font-size: 11px;')
        author.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(author)

        layout.addStretch()

        # 关闭按钮
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)


class AppWindow(QMainWindow):
    """主窗口 - Anthropic 官方风格"""
    def __init__(self):
        super().__init__()
        self.setWindowTitle("货运保险工具 ProMax")
        self.resize(600, 850)
        self.setMinimumSize(500, 650)
        self.setStyleSheet(STYLE_SHEET)

        self._setup_ui()

    def _setup_ui(self):
        # 主容器
        central = QWidget()
        central.setObjectName("centralWidget")
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(12)
        layout.setContentsMargins(30, 20, 30, 20)

        # ==========================================
        # 标题栏 - Anthropic 风格
        # ==========================================
        header_layout = QHBoxLayout()

        title = QLabel("🚢 货运保险工具")
        title.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 26px; font-weight: bold;")
        header_layout.addWidget(title)

        header_layout.addStretch()

        # 版本信息
        subtitle = QLabel("ProMax · 对账处理 · PDF导出 · 投保单")
        subtitle.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 12px;")
        header_layout.addWidget(subtitle)

        # 支持作者按钮 - Anthropic 强调色风格
        self.donate_btn = QPushButton("💝 支持作者")
        self.donate_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.donate_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.ACCENT};
                color: {AnthropicColors.TEXT_LIGHT};
                border: none;
                border-radius: 15px;
                padding: 6px 16px;
                font-size: 12px;
                font-weight: 500;
                margin-left: 15px;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.ACCENT_DARK};
            }}
        """)
        self.donate_btn.clicked.connect(self._show_donate_dialog)

        # 添加柔和阴影
        donate_shadow = QGraphicsDropShadowEffect()
        donate_shadow.setBlurRadius(12)
        donate_shadow.setColor(QColor(217, 119, 87, 80))
        donate_shadow.setOffset(0, 2)
        self.donate_btn.setGraphicsEffect(donate_shadow)

        header_layout.addWidget(self.donate_btn)
        layout.addLayout(header_layout)

        # ==========================================
        # 主Tab区域 - Anthropic风格
        # ==========================================
        self.main_page = MainPage()
        self.tabs = QTabWidget()
        self.tabs.addTab(self.main_page, "📋 对账处理")
        self.tabs.addTab(PdfPage(), "📄 PDF导出")
        self.tabs.addTab(WordPage(self.main_page), "📝 投保单")
        layout.addWidget(self.tabs, 1)

        # ==========================================
        # 底部版本信息
        # ==========================================
        version = QLabel("ProMax Edition · Made with ❤️ by Dachi Yijin")
        version.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 11px;")
        layout.addWidget(version)

    def _show_donate_dialog(self):
        """显示捐赠对话框"""
        dialog = DonateDialog(self)
        dialog.exec()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    font = app.font()
    # Anthropic 风格字体
    font.setFamily("PingFang SC")
    font.setPointSize(13)
    app.setFont(font)
    window = AppWindow()
    window.show()
    sys.exit(app.exec())