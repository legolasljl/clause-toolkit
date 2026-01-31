#!/usr/bin/env python3
"""
费率方案提取脚本 v2.1
从附加险费率方案 docx/doc 文件中提取结构化数据，输出 JSON。
支持雇主责任保险、企业财产保险等多行业附加条款。

用法:
    python3 extract_rates.py "/path/to/附加险及费率" -o rate_data.json

费率方案分类（与 HTML 端一致）:
  - simple_percentage: 主险保费的固定加收百分比
  - deduction: 减收主险保费的百分比
  - modifier_coeff: 调整系数（如误工费）
  - sudden_death: 突发疾病身故类（基准费率×限额×人数）
  - per_person_rate: 每人费率（区分工伤保险状态）
  - per_person_base: 每人定额（如药品服务）
  - disability_adjust: 伤残赔偿比例调整
  - property_loss: 财产损失（基本保险费+限额×费率）
  - formula_sum: 求和公式（雇主法律责任/一次性伤残）
  - table_coefficient: 多系数表相乘
  - no_calc: 有计费说明但无需单独计算
  - regulatory: 不涉及费用调整
  - included_in_main: 纳入主险保险金额计收保险费
  - daily_prorate: 按日比例计算保费
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx: pip3 install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# 噪音过滤
# ---------------------------------------------------------------------------
NOISE_PATTERNS = [
    re.compile(r'PAGE\s+', re.IGNORECASE),
    re.compile(r'NUMPAGES', re.IGNORECASE),
    re.compile(r'MERGEFORMAT', re.IGNORECASE),
    re.compile(r'第\s*.*页\s*共\s*.*页'),
    re.compile(r'^-\s*\d+\s*-$'),
    re.compile(r'^\d+$'),
]

COMPANY_PREFIX = "中国太平洋财产保险股份有限公司"


def is_noise(text):
    """判断段落是否为噪音行。"""
    stripped = text.strip()
    if not stripped:
        return True
    for pat in NOISE_PATTERNS:
        if pat.search(stripped):
            return True
    return False


def clean_paragraphs(doc):
    """提取并清理文档段落文本。"""
    result = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and not is_noise(text):
            result.append(text)
    return result


# ---------------------------------------------------------------------------
# 表格解析
# ---------------------------------------------------------------------------
def parse_tables(doc):
    """提取文档中的所有表格为 list[list[list[str]]]。"""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)
    return tables


def parse_coefficient_value(text):
    """解析系数值文本，返回结构化数据。"""
    text = text.strip()
    # 替换全角字符
    text = text.replace('，', ',').replace('（', '(').replace('）', ')')

    # 范围格式: [0.7, 0.8] 或 (0.7, 0.8] 等
    range_match = re.match(
        r'[\[(\（\[]?\s*([\d.]+)\s*[,，]\s*([\d.]+)\s*[\])\）\]]?', text
    )
    if range_match:
        lo = float(range_match.group(1))
        hi = float(range_match.group(2))
        return {"type": "range", "min": lo, "max": hi, "display": text}

    # 固定数值
    num_match = re.match(r'^([\d.]+)$', text)
    if num_match:
        return {"type": "fixed", "value": float(num_match.group(1)), "display": text}

    # 其他文本
    return {"type": "text", "display": text}


def build_coefficient_table(raw_table, paragraphs):
    """从原始表格数据构建系数表结构。"""
    if len(raw_table) < 2:
        return None

    header = raw_table[0]
    table_name = header[0] if header[0] else "调整系数"

    # 检测是否支持线性插值
    supports_interpolation = any(
        '线性插值' in p or '插值' in p for p in paragraphs
    )

    rows = []
    for row_data in raw_table[1:]:
        if len(row_data) < 2:
            continue
        param = row_data[0].strip()
        coeff_text = row_data[1].strip()
        if not param or not coeff_text:
            continue
        parsed = parse_coefficient_value(coeff_text)
        rows.append({
            "parameter": param,
            "coefficient": coeff_text,
            "parsedValue": parsed
        })

    if not rows:
        return None

    return {
        "name": table_name,
        "headerRow": header,
        "supportsInterpolation": supports_interpolation,
        "rows": rows
    }


# ---------------------------------------------------------------------------
# 费率类型分类
# ---------------------------------------------------------------------------
def extract_percentage(text):
    """从文本中提取百分比数值。"""
    # 匹配 "5%" 或 "5％" 或 "百分之五"
    m = re.search(r'([\d.]+)\s*[%％]', text)
    if m:
        return float(m.group(1))
    return None


def extract_multiplier(text):
    """从文本中提取倍数。"""
    m = re.search(r'主险保险费的\s*([\d.]+)\s*倍', text)
    if m:
        return float(m.group(1))
    return None


def classify_and_extract(paragraphs, tables, filename):
    """根据段落和表格内容分类费率类型并提取数据（v2.0，与 HTML 端一致）。"""
    full_text = ' '.join(paragraphs)

    # 过滤掉公司名和标题行，获取实质段落
    substantive = [
        p for p in paragraphs
        if COMPANY_PREFIX not in p
        and not (p.endswith('费率方案') and len(p) < 100)
    ]

    # 从文件名中提取条款简称
    clause_name = filename.replace('.docx', '').replace('.doc', '')
    clause_name = clause_name.replace(COMPANY_PREFIX, '')

    # ---- regulatory ----
    regulatory_keywords = [
        '不涉及保险费的调整',
        '属于规范类',
        '不涉及费率',
        '不另收保险费',
        '不增减主险的保险费',
        '不增减主险保险费',
    ]
    is_purely_regulatory = (
        not tables
        and substantive
        and all(
            any(kw in p for kw in regulatory_keywords)
            or '保单最终保险费' in p
            or '工资总额' in p
            for p in substantive
        )
    )
    if is_purely_regulatory:
        return {
            "rateType": "regulatory",
            "description": substantive[0][:300] if substantive else ""
        }

    # ---- 构建系数表 ----
    coeff_tables = []
    if tables:
        for raw_table in tables:
            ct = build_coefficient_table(raw_table, paragraphs)
            if ct:
                coeff_tables.append(ct)

    # ---- 提取百分比 ----
    percentages = []
    for p in paragraphs:
        for m in re.finditer(r'([\d.]+)\s*[%％]', p):
            percentages.append({"value": float(m.group(1)), "context": p})

    # ---- 关键词检测费率类型（与 HTML 端 RC_ADDON_KEYWORD_MAP 一致）----
    keyword_map = [
        (['误工费'], 'modifier_coeff'),
        (['突发疾病身故'], 'sudden_death'),
        (['工伤补充', '特定人员'], 'per_person_rate'),
        (['药品服务', '药品费用'], 'per_person_base'),
        (['劳务关系人员'], 'disability_adjust'),
        (['特定财产损失'], 'property_loss'),
        (['雇主法律责任', '法律费用责任'], 'formula_sum'),
        (['一次性伤残'], 'formula_sum'),
        (['突发疾病除外', '猝死除外'], 'deduction'),
        (['月申报', '员工自动承保', '每月申报'], 'no_calc'),
    ]

    detected_type = None
    for keywords, rate_type in keyword_map:
        if any(kw in clause_name for kw in keywords):
            detected_type = rate_type
            break

    # ---- 按检测类型构建 entry ----

    if detected_type == 'modifier_coeff':
        return {
            "rateType": "modifier_coeff",
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": "调整后主险保费 = 主险保费 × 免赔天数调整系数"
        }

    if detected_type == 'sudden_death':
        base_pct = 6.6
        for pi in percentages:
            if '基准保险费' in pi["context"] or '每人每次事故赔偿限额' in pi["context"]:
                base_pct = pi["value"]
                break
        return {
            "rateType": "sudden_death",
            "basePercent": base_pct,
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": f"基准保费 = 每人限额 × {base_pct}% × 人数，再乘以系数调整"
        }

    if detected_type == 'per_person_rate':
        rate_info = {}
        for pi in percentages:
            ctx = pi["context"]
            if '已购买工伤保险' in ctx or '有工伤' in ctx:
                rate_info["with_injury_insurance"] = pi["value"]
            elif '未购买工伤保险' in ctx or '无工伤' in ctx:
                rate_info["without_injury_insurance"] = pi["value"]
            elif "default" not in rate_info:
                rate_info["default"] = pi["value"]
        return {
            "rateType": "per_person_rate",
            "rateInfo": rate_info,
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": "保费 = 每人保费 × 费率% × 人数 × 系数"
        }

    if detected_type == 'per_person_base':
        base_amount = 300
        for p in paragraphs:
            amt_m = re.search(r'(\d+)\s*元[/／每]人', p)
            if amt_m:
                base_amount = int(amt_m.group(1))
                break
        return {
            "rateType": "per_person_base",
            "baseAmount": base_amount,
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": f"保费 = {base_amount}元/人 × 系数 × 人数"
        }

    if detected_type == 'disability_adjust':
        return {
            "rateType": "disability_adjust",
            "adjustCoeffs": {"table1": 0.995, "table2": 1.072, "table3": 0.919},
            "description": substantive[0] if substantive else "",
            "formula": "保费 = 每人保费 × 伤残调整系数 × 人数"
        }

    if detected_type == 'property_loss':
        base_premium = 20
        base_rate_pct = 1.5
        for pi in percentages:
            base_rate_pct = pi["value"]
            break
        for p in paragraphs:
            amt_m = re.search(r'基[本准]保险费[=＝]?\s*(\d+)\s*元', p)
            if amt_m:
                base_premium = int(amt_m.group(1))
                break
        return {
            "rateType": "property_loss",
            "basePremium": base_premium,
            "baseRatePercent": base_rate_pct,
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": f"保费 = ({base_premium}元 + 每人赔偿限额 × {base_rate_pct}%) × 系数积 × 承保人数"
        }

    if detected_type == 'formula_sum':
        base_rate_factor = 1.0
        if '一次性伤残' in clause_name:
            if 'A款' in clause_name or '（A）' in clause_name:
                base_rate_factor = 0.9
            elif 'B款' in clause_name or '（B）' in clause_name:
                base_rate_factor = 1.0
            elif 'C款' in clause_name or '（C）' in clause_name:
                base_rate_factor = 1.1
            elif 'D款' in clause_name or '（D）' in clause_name:
                base_rate_factor = 1.2
        else:
            if '90' in full_text:
                base_rate_factor = 0.9
            elif '95' in full_text:
                base_rate_factor = 0.95
        return {
            "rateType": "formula_sum",
            "baseRateFactor": base_rate_factor,
            "coefficientTables": coeff_tables,
            "description": substantive[0] if substantive else "",
            "formula": f"保费 = Σ(每人限额 × 主险基准费率 × {base_rate_factor} × 人数 × 系数积)"
        }

    if detected_type == 'deduction':
        deduct_pct = 5.0
        if percentages:
            deduct_pct = percentages[0]["value"]
        return {
            "rateType": "deduction",
            "deductPercent": deduct_pct,
            "description": substantive[0] if substantive else "",
            "formula": f"减收 = 主险保费 × {deduct_pct}%"
        }

    if detected_type == 'no_calc':
        return {
            "rateType": "no_calc",
            "description": substantive[0] if substantive else full_text[:300],
            "formula": "本条款有计费说明但无需单独计算附加保费"
        }

    # ---- 新类型: included_in_main（纳入主险保险金额）----
    included_keywords = ['纳入主险保险金额计收保险费', '纳入主险保险金额', '应纳入主险保险金额']
    if any(kw in full_text for kw in included_keywords) and not tables:
        return {
            "rateType": "included_in_main",
            "description": next(
                (p for p in substantive if any(kw in p for kw in included_keywords)),
                ""
            ),
            "formula": "纳入主险保险金额，按主险费率计收保险费，不另收附加保险费"
        }

    # ---- 新类型: daily_prorate（按日比例计算）----
    daily_keywords = ['按日比例计算', '按日比例收取', '从发生损失之日起至保险止期按日比例']
    if any(kw in full_text for kw in daily_keywords):
        formula_text = next(
            (p for p in paragraphs if any(kw in p for kw in daily_keywords)),
            ""
        )
        return {
            "rateType": "daily_prorate",
            "description": formula_text[:300],
            "formula": "按日比例计算保费 = 保险金额 × 保单费率 × (天数 / 365)"
        }

    # ---- 未匹配关键词，走通用检测 ----

    # table_coefficient
    if coeff_tables:
        base_premium = _extract_base_premium(paragraphs)
        formula = _extract_formula(paragraphs)
        result = {
            "rateType": "table_coefficient",
            "basePremium": base_premium,
            "coefficientTables": coeff_tables,
            "formula": formula,
            "description": _find_description(paragraphs)
        }
        extra_rules = _extract_extra_rules(paragraphs)
        if extra_rules:
            result["extraRules"] = extra_rules
        return result

    # 检测减收/减少类（deduction）
    for p in substantive:
        if ('减少' in p or '减收' in p) and extract_percentage(p) is not None:
            pct = extract_percentage(p)
            return {
                "rateType": "deduction",
                "deductPercent": pct,
                "description": p,
                "formula": f"减收 = 主险保费 × {pct}%"
            }

    # formula_conditional
    has_formula_sign = any('＝' in p or '×' in p for p in substantive)
    has_condition = any(
        any(kw in p for kw in ['若', '如果'])
        for p in substantive
    )

    if has_condition or (has_formula_sign and not tables):
        conditions = _extract_conditions(paragraphs)
        base_rate = None
        for p in paragraphs:
            pct = extract_percentage(p)
            if pct:
                base_rate = pct
                break
        return {
            "rateType": "formula_conditional",
            "baseRatePercent": base_rate,
            "conditions": conditions,
            "description": _find_description(paragraphs)
        }

    # simple_percentage
    for p in paragraphs:
        pct = extract_percentage(p)
        if pct is not None:
            return {
                "rateType": "simple_percentage",
                "percentage": pct,
                "description": p
            }

    # 倍数形式
    for p in paragraphs:
        mult = extract_multiplier(p)
        if mult is not None:
            return {
                "rateType": "simple_percentage",
                "percentage": mult * 100,
                "multiplier": mult,
                "description": p
            }

    # regulatory fallback
    adjustment_keywords = ['保单最终保险费', '工资总额进行调整', '保费调整', '按保险费率补缴保险费']
    for kw in adjustment_keywords:
        if kw in full_text:
            return {
                "rateType": "regulatory",
                "description": next(
                    (p for p in paragraphs if kw in p or '调整' in p),
                    full_text[:200]
                )
            }

    # fallback
    return {
        "rateType": "unknown",
        "description": full_text[:300] if full_text else "无法识别费率类型"
    }


def _extract_base_premium(paragraphs):
    """提取基准保险费描述。"""
    for p in paragraphs:
        if '基准保险费' in p or '基准保费' in p:
            pct = extract_percentage(p)
            mult = extract_multiplier(p)
            # 优先匹配 "主险保险费的X%" 格式（企业财产险常见）
            m = re.search(r'主险保险费的\s*([\d.]+)\s*[%％]', p)
            if m:
                return {
                    "description": p,
                    "percentage": float(m.group(1))
                }
            result = {"description": p}
            if pct:
                result["percentage"] = pct
            if mult:
                result["multiplier"] = mult
            return result
    # 回退: 查找 "主险保险费的X%"
    for p in paragraphs:
        m = re.search(r'主险保险费的\s*([\d.]+)\s*[%％]', p)
        if m:
            return {
                "description": p,
                "percentage": float(m.group(1))
            }
        m2 = re.search(r'主险保险费的\s*([\d.]+)\s*倍', p)
        if m2:
            return {
                "description": p,
                "multiplier": float(m2.group(1))
            }
    return {"description": "未找到基准保险费描述"}


def _extract_formula(paragraphs):
    """提取计算公式。"""
    for p in paragraphs:
        if '保险费' in p and ('×' in p or '＝' in p or '乘积' in p):
            return p
    return "保险费 = 基准保险费 × 各项费率调整系数的乘积"


def _extract_conditions(paragraphs):
    """提取条件公式列表。"""
    conditions = []
    for p in paragraphs:
        if any(kw in p for kw in ['若', '如果']) and (
            '＝' in p or '×' in p or '%' in p or '％' in p or
            '减收' in p or '加收' in p or '减少' in p or
            '不调整' in p or '不涉及' in p
        ):
            # 提取条件部分和公式部分
            condition_part = p
            formula_part = p
            # 尝试按逗号分割
            for sep in ['，', ',', '：', ':']:
                if sep in p:
                    parts = p.split(sep, 1)
                    if len(parts) == 2:
                        condition_part = parts[0].strip()
                        formula_part = parts[1].strip()
                    break
            conditions.append({
                "condition": condition_part,
                "formula": formula_part,
                "fullText": p
            })
        elif '不涉及保险费的调整' in p or '则不涉及' in p or '则不调整保险费' in p:
            conditions.append({
                "condition": p,
                "formula": "不调整",
                "fullText": p
            })
        # 企业财产险: "增加的保险费＝Σ（保险金额×费率×天数/365）" 模式
        elif ('Σ' in p or '∑' in p or '求和' in p) and ('×' in p or '＝' in p):
            conditions.append({
                "condition": "求和公式",
                "formula": p,
                "fullText": p
            })
    return conditions


def _extract_extra_rules(paragraphs):
    """提取附加规则（如未购买工伤保险等条件）。"""
    rules = []
    for p in paragraphs:
        if '未购买' in p or '未参加' in p or '未投保' in p:
            rules.append(p)
    return rules if rules else None


def _find_description(paragraphs):
    """获取费率方案的简要描述。"""
    # 跳过公司名和标题,取第一个实质段落
    for p in paragraphs:
        if COMPANY_PREFIX in p:
            continue
        if '费率方案' in p and len(p) < 80:
            continue
        if p.strip():
            return p[:300]
    return ""


# ---------------------------------------------------------------------------
# 文件名解析
# ---------------------------------------------------------------------------
def parse_filename(filename):
    """从文件名中提取条款名称和行业信息。"""
    name = filename.replace('.docx', '').replace('.doc', '')

    # 去掉公司前缀
    name = name.replace(COMPANY_PREFIX, '')

    # 提取行业信息
    industry = "雇主责任保险"
    industry_patterns = [
        (r'企业财产保险', '企业财产保险'),
        (r'四川省建筑施工企业', '四川省建筑施工企业雇主责任保险'),
        (r'网约配送员', '网约配送员雇主责任保险'),
        (r'石油石化行业[A-Z]?款?', '石油石化行业雇主责任保险'),
        (r'集成电路行业', '集成电路行业雇主责任保险'),
        (r'特种设备安全责任', '特种设备安全责任保险'),
        (r'护理服务责任', '护理服务责任保险'),
        (r'民宿经营责任', '民宿经营责任保险'),
        (r'燃气气瓶安全责任', '燃气气瓶安全责任保险'),
        (r'安全生产责任', '安全生产责任保险'),
        (r'食品安全责任', '食品安全责任保险'),
    ]
    for pat, ind in industry_patterns:
        if re.search(pat, name):
            industry = ind
            break

    # 提取条款简称
    clause_name = name.replace('费率方案', '').strip()
    # 去掉行业前缀，提取附加险名称
    m = re.search(r'附加(.+?)(?:条款|保险)?$', clause_name)
    if m:
        clause_name = '附加' + m.group(1)
        if not clause_name.endswith('条款') and not clause_name.endswith('保险'):
            clause_name = clause_name.rstrip('（）()')

    return {
        "clauseName": clause_name.strip(),
        "fullName": filename.replace('.docx', '').replace('.doc', ''),
        "industry": industry
    }


# ---------------------------------------------------------------------------
# 主处理
# ---------------------------------------------------------------------------
def _read_doc_file(filepath):
    """尝试读取 .doc 文件，返回 (paragraphs, tables) 或 None。"""
    import subprocess
    import shutil

    # 尝试 antiword → catdoc 的顺序
    for cmd in ['antiword', 'catdoc']:
        if not shutil.which(cmd):
            continue
        try:
            result = subprocess.run(
                [cmd, filepath],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                lines = result.stdout.strip().split('\n')
                paragraphs = [
                    line.strip() for line in lines
                    if line.strip() and not is_noise(line.strip())
                ]
                return paragraphs, []
        except Exception:
            continue

    # 尝试 textract (需 pip install textract)
    try:
        import textract
        raw = textract.process(filepath).decode('utf-8', errors='ignore')
        lines = raw.strip().split('\n')
        paragraphs = [
            line.strip() for line in lines
            if line.strip() and not is_noise(line.strip())
        ]
        return paragraphs, []
    except Exception:
        pass

    return None


def process_file(filepath):
    """处理单个 docx/doc 文件，返回费率条目或 None。"""
    filename = os.path.basename(filepath)

    if not filename.endswith('.docx') and not filename.endswith('.doc'):
        return None

    # 处理 .doc 文件（非 .docx）
    if filename.endswith('.doc') and not filename.endswith('.docx'):
        result = _read_doc_file(filepath)
        if result is None:
            print(f"  ⚠️  .doc 文件需要 antiword/catdoc/textract: {filename}", file=sys.stderr)
            return None
        paragraphs, tables = result
        file_info = parse_filename(filename)
        rate_info = classify_and_extract(paragraphs, tables, filename)
        return {
            **file_info,
            **rate_info,
            "sourceFile": filename
        }

    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ⚠️  无法打开: {filename} - {e}", file=sys.stderr)
        return None

    paragraphs = clean_paragraphs(doc)
    tables = parse_tables(doc)

    # 文件名解析
    file_info = parse_filename(filename)

    # 分类并提取
    rate_info = classify_and_extract(paragraphs, tables, filename)

    return {
        **file_info,
        **rate_info,
        "sourceFile": filename
    }


def process_directory(dir_path, verbose=False):
    """处理目录中所有费率方案 docx 文件。"""
    if not os.path.isdir(dir_path):
        print(f"错误: 目录不存在 - {dir_path}", file=sys.stderr)
        sys.exit(1)

    # 处理费率方案文件（.docx 和 .doc）
    files = [
        f for f in sorted(os.listdir(dir_path))
        if '费率方案' in f and (f.endswith('.docx') or f.endswith('.doc'))
    ]

    print(f"📂 扫描目录: {dir_path}")
    print(f"📄 找到费率方案文件: {len(files)} 个")

    entries = []
    stats = {}
    errors = 0

    for i, fname in enumerate(files, 1):
        fpath = os.path.join(dir_path, fname)
        if verbose:
            print(f"  [{i}/{len(files)}] {fname}")

        entry = process_file(fpath)
        if entry:
            rate_type = entry.get("rateType", "unknown")
            stats[rate_type] = stats.get(rate_type, 0) + 1
            entries.append(entry)
            if verbose:
                print(f"    → {rate_type}")
        else:
            errors += 1

    print(f"\n✅ 处理完成:")
    for rt in sorted(stats.keys()):
        print(f"   {rt:25s} {stats[rt]}")
    if errors:
        print(f"   {'errors':25s} {errors}")
    print(f"   {'总计':24s} {len(entries)}")

    return {
        "version": "2.1",
        "generatedAt": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sourceDirectory": dir_path,
        "totalEntries": len(entries),
        "stats": stats,
        "entries": entries
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="从费率方案 docx 文件提取结构化 JSON 数据"
    )
    parser.add_argument(
        "directory",
        help="包含费率方案 docx 文件的目录路径"
    )
    parser.add_argument(
        "-o", "--output",
        default="rate_data.json",
        help="输出 JSON 文件路径 (默认: rate_data.json)"
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="显示详细处理信息"
    )

    args = parser.parse_args()

    result = process_directory(args.directory, verbose=args.verbose)

    output_path = args.output
    if not os.path.isabs(output_path):
        output_path = os.path.join(os.getcwd(), output_path)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n📁 输出文件: {output_path}")


if __name__ == "__main__":
    main()
