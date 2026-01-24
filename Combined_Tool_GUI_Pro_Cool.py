"""
Excel/Word 自动化多功能工具 - Pro 优化版 V5.0
新增功能：
1. [设置] 数值格式化开关 - 可选千分位/原始格式
2. [设置] 递归扫描子目录
3. [设置] 单元格范围支持 (A1:A10)
4. [设置] 深色/浅色主题切换
5. [设置] 最近路径记忆
6. [功能] 拖拽文件/文件夹支持
7. [功能] 数据预览
8. [功能] 导出日志
9. [功能] 二维码生成支持 (需要 qrcode 库)
10. [功能] 打赏二维码显示真实图片

"""

import sys
import os
import platform
import subprocess
import re
import json
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime
from io import BytesIO
import base64

# macOS Qt插件路径修复 - 更健壮的版本
if sys.platform == 'darwin':
    # 尝试多个可能的路径
    possible_paths = [
        Path(__file__).resolve().parent / '.venv' / 'lib' / 'python3.11' / 'site-packages' / 'PyQt5' / 'Qt5' / 'plugins',
        Path(__file__).resolve().parent / '.venv' / 'lib' / 'python3.10' / 'site-packages' / 'PyQt5' / 'Qt5' / 'plugins',
        Path(__file__).resolve().parent / '.venv' / 'lib' / 'python3.12' / 'site-packages' / 'PyQt5' / 'Qt5' / 'plugins',
        Path(sys.prefix) / 'lib' / 'python3.11' / 'site-packages' / 'PyQt5' / 'Qt5' / 'plugins',
    ]
    for _pyqt_path in possible_paths:
        if _pyqt_path.exists():
            os.environ['QT_PLUGIN_PATH'] = str(_pyqt_path)
            break
    else:
        # 如果都找不到，尝试动态查找
        try:
            import PyQt5
            pyqt_dir = Path(PyQt5.__file__).parent
            plugin_path = pyqt_dir / 'Qt5' / 'plugins'
            if plugin_path.exists():
                os.environ['QT_PLUGIN_PATH'] = str(plugin_path)
            else:
                plugin_path = pyqt_dir / 'Qt' / 'plugins'
                if plugin_path.exists():
                    os.environ['QT_PLUGIN_PATH'] = str(plugin_path)
        except:
            pass

import pandas as pd
import openpyxl
from openpyxl.utils import column_index_from_string
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QFileDialog, QGroupBox,
    QComboBox, QMessageBox, QProgressBar, QCheckBox, QInputDialog,
    QTabWidget, QGridLayout, QFrame, QDialog, QSpinBox, QTableWidget,
    QTableWidgetItem, QHeaderView, QSplitter
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor, QPalette, QPixmap, QDragEnterEvent, QDropEvent

# 可选：二维码支持
try:
    import qrcode
    HAS_QRCODE = True
except ImportError:
    HAS_QRCODE = False


# ==================== 常量定义 ====================
APP_NAME = 'Excel/Word 自动化工具 Pro V5.0'
APP_AUTHOR = 'Dachi_Yijin'
APP_VERSION = 'V5.0'
DEFAULT_OUTPUT_DIR = 'Word_Output'
DEFAULT_EXTRACTION_FILE = 'Extraction_Result.xlsx'
CONFIG_FILE = 'tool_config_v5.json'

# ASCII Art Logo
APP_LOGO = """
╔══════════════════════════════════════════════════════════════════════════════╗
║   █████╗ ██╗   ██╗████████╗ ██████╗ ██╗    ██╗ ██████╗ ██████╗ ██████╗       ║
║  ██╔══██╗██║   ██║╚══██╔══╝██╔═══██╗██║    ██║██╔═══██╗██╔══██╗██╔══██╗      ║
║  ███████║██║   ██║   ██║   ██║   ██║██║ █╗ ██║██║   ██║██████╔╝██║  ██║      ║
║  ██╔══██║██║   ██║   ██║   ██║   ██║██║███╗██║██║   ██║██╔══██╗██║  ██║      ║
║  ██║  ██║╚██████╔╝   ██║   ╚██████╔╝╚███╔███╔╝╚██████╔╝██║  ██║██████╔╝      ║
║  ╚═╝  ╚═╝ ╚═════╝    ╚═╝    ╚═════╝  ╚══╝╚══╝  ╚═════╝ ╚═╝  ╚═╝╚═════╝       ║
║                    🚀 自动化办公神器 Pro V5.0 🚀                             ║
║                        Author: Dachi_Yijin                                   ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""


# ==================== 颜色主题 ====================
class Colors:
    """浅色主题"""
    PRIMARY = '#2563eb'
    PRIMARY_DARK = '#1d4ed8'
    PRIMARY_LIGHT = '#3b82f6'
    SUCCESS = '#10b981'
    WARNING = '#f59e0b'
    ERROR = '#ef4444'
    BG_LIGHT = '#f8fafc'
    BG_CARD = '#ffffff'
    TEXT_PRIMARY = '#1e293b'
    TEXT_SECONDARY = '#64748b'
    BORDER = '#e2e8f0'


class DarkColors:
    """深色主题"""
    PRIMARY = '#3b82f6'
    PRIMARY_DARK = '#2563eb'
    PRIMARY_LIGHT = '#60a5fa'
    SUCCESS = '#34d399'
    WARNING = '#fbbf24'
    ERROR = '#f87171'
    BG_LIGHT = '#1e293b'
    BG_CARD = '#334155'
    TEXT_PRIMARY = '#f1f5f9'
    TEXT_SECONDARY = '#94a3b8'
    BORDER = '#475569'


# ==================== 工具函数 ====================
def open_folder(path: str) -> bool:
    """跨平台打开文件夹"""
    if not os.path.exists(path):
        return False
    folder = os.path.dirname(path) if os.path.isfile(path) else path
    try:
        system = platform.system()
        if system == 'Darwin':
            subprocess.run(['open', folder])
        elif system == 'Windows':
            os.startfile(folder)
        else:
            subprocess.run(['xdg-open', folder])
        return True
    except Exception as e:
        print(f'打开目录失败: {e}')
        return False


def format_number(value, use_formatting: bool = True) -> str:
    """数值格式化：千分位+两位小数（可选）"""
    if value is None or (isinstance(value, str) and not value.strip()):
        return ''
    if pd.isna(value):
        return ''
    try:
        if use_formatting:
            return f'{float(value):,.2f}'
        else:
            return str(value)
    except (ValueError, TypeError):
        return str(value)


def is_chinese_char(char: str) -> bool:
    """判断字符是否为中文"""
    if len(char) != 1:
        return False
    code = ord(char)
    return (0x4E00 <= code <= 0x9FFF or
            0x3400 <= code <= 0x4DBF or
            0x20000 <= code <= 0x2A6DF or
            0xF900 <= code <= 0xFAFF or
            0x2F00 <= code <= 0x2FDF or
            0x3000 <= code <= 0x303F or
            0xFF00 <= code <= 0xFFEF)


def split_text_by_language(text: str) -> List[Tuple[str, bool]]:
    """将文本按中英文分割"""
    if not text:
        return []
    
    segments = []
    current_segment = ''
    current_is_chinese = None
    
    for char in text:
        char_is_chinese = is_chinese_char(char)
        
        if current_is_chinese is None:
            current_is_chinese = char_is_chinese
            current_segment = char
        elif char_is_chinese == current_is_chinese:
            current_segment += char
        else:
            if current_segment:
                segments.append((current_segment, current_is_chinese))
            current_segment = char
            current_is_chinese = char_is_chinese
    
    if current_segment:
        segments.append((current_segment, current_is_chinese))
    
    return segments


def generate_qrcode_image(data: str, box_size: int = 10, border: int = 2) -> BytesIO:
    """生成二维码图片"""
    if not HAS_QRCODE:
        raise ImportError("需要安装 qrcode 库: pip install qrcode[pil]")
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=box_size,
        border=border,
    )
    qr.add_data(data)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    return img_bytes


def generate_qrcode_base64(data: str, box_size: int = 6, border: int = 1) -> str:
    """生成二维码的 base64 字符串（用于 UI 显示）"""
    if not HAS_QRCODE:
        return ""
    try:
        img_bytes = generate_qrcode_image(data, box_size, border)
        return base64.b64encode(img_bytes.getvalue()).decode()
    except:
        return ""


# ==================== Word合并功能 ====================
def merge_word_documents(word_files: List[str], output_path: str, log_callback=None) -> bool:
    """
    合并多个Word文档为一个文件，每个文档之间使用分页符隔开
    使用XML级别的元素复制，完整保留原文档的格式和结构

    Args:
        word_files: Word文件路径列表
        output_path: 输出文件路径
        log_callback: 日志回调函数

    Returns:
        是否成功
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    if not word_files:
        return False

    try:
        # 使用第一个文档作为基础
        merged_doc = Document(word_files[0])

        if log_callback:
            log_callback(f'📄 开始合并 {len(word_files)} 个文档...')
            log_callback(f'✅ [1/{len(word_files)}] {os.path.basename(word_files[0])}')

        # 获取合并文档的body元素
        merged_body = merged_doc.element.body

        # 找到并保存sectPr（section属性），需要保持在文档最后
        sectPr = merged_body.find(qn('w:sectPr'))
        if sectPr is not None:
            final_sectPr = deepcopy(sectPr)
            merged_body.remove(sectPr)
        else:
            final_sectPr = None

        # 从第二个文档开始逐个合并
        for i, word_file in enumerate(word_files[1:], start=2):
            try:
                # 添加分页符段落
                page_break_para = OxmlElement('w:p')
                page_break_run = OxmlElement('w:r')
                page_break_br = OxmlElement('w:br')
                page_break_br.set(qn('w:type'), 'page')
                page_break_run.append(page_break_br)
                page_break_para.append(page_break_run)
                merged_body.append(page_break_para)

                # 读取源文档
                source_doc = Document(word_file)
                source_body = source_doc.element.body

                # 复制源文档的所有body子元素（除了sectPr）
                for child in source_body:
                    # 跳过sectPr元素（section属性）
                    if child.tag == qn('w:sectPr'):
                        continue
                    # 深拷贝元素并添加到目标body
                    new_child = deepcopy(child)
                    merged_body.append(new_child)

                if log_callback:
                    log_callback(f'✅ [{i}/{len(word_files)}] {os.path.basename(word_file)}')

            except Exception as e:
                if log_callback:
                    log_callback(f'⚠️ [{i}/{len(word_files)}] 合并失败: {os.path.basename(word_file)} - {e}')

        # 将sectPr放回文档最后（Word要求sectPr必须在body的最后）
        if final_sectPr is not None:
            merged_body.append(final_sectPr)

        # 保存合并后的文档
        merged_doc.save(output_path)

        if log_callback:
            log_callback(f'🎉 合并完成: {os.path.basename(output_path)}')

        return True

    except Exception as e:
        if log_callback:
            log_callback(f'❌ 合并失败: {e}')
        return False


# ==================== 样式管理 ====================
def get_stylesheet(colors) -> str:
    """生成样式表"""
    return f"""
/* 全局样式 */
QMainWindow, QWidget {{
    background-color: {colors.BG_LIGHT};
    color: {colors.TEXT_PRIMARY};
    font-family: "Microsoft YaHei", "PingFang SC", "Helvetica Neue", Arial, sans-serif;
}}

/* 标签页样式 */
QTabWidget::pane {{
    border: 1px solid {colors.BORDER};
    border-radius: 8px;
    background-color: {colors.BG_CARD};
    padding: 15px;
}}

QTabBar::tab {{
    background-color: {colors.BG_LIGHT};
    color: {colors.TEXT_SECONDARY};
    padding: 12px 24px;
    margin-right: 4px;
    border-top-left-radius: 8px;
    border-top-right-radius: 8px;
    font-size: 13px;
    font-weight: 500;
}}

QTabBar::tab:selected {{
    background-color: {colors.BG_CARD};
    color: {colors.PRIMARY};
    border-bottom: 2px solid {colors.PRIMARY};
}}

/* 分组框样式 */
QGroupBox {{
    font-size: 14px;
    font-weight: 600;
    color: {colors.TEXT_PRIMARY};
    border: 1px solid {colors.BORDER};
    border-radius: 10px;
    margin-top: 12px;
    padding-top: 16px;
    background-color: {colors.BG_CARD};
}}

QGroupBox::title {{
    subcontrol-origin: margin;
    subcontrol-position: top left;
    left: 16px;
    padding: 0 8px;
    background-color: {colors.BG_CARD};
    color: {colors.PRIMARY};
}}

/* 按钮样式 */
QPushButton {{
    background-color: {colors.PRIMARY};
    color: white;
    border: none;
    border-radius: 6px;
    padding: 8px 16px;
    font-size: 13px;
    font-weight: 500;
    min-height: 20px;
}}

QPushButton:hover {{
    background-color: {colors.PRIMARY_DARK};
}}

QPushButton:disabled {{
    background-color: #94a3b8;
}}

/* 主操作按钮 */
QPushButton#primaryButton {{
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 {colors.PRIMARY}, stop:1 {colors.PRIMARY_LIGHT});
    font-size: 14px;
    font-weight: 600;
    padding: 12px 24px;
}}

/* 次要按钮 */
QPushButton#secondaryButton {{
    background-color: {colors.BG_LIGHT};
    color: {colors.TEXT_PRIMARY};
    border: 1px solid {colors.BORDER};
}}

QPushButton#secondaryButton:hover {{
    background-color: {colors.BORDER};
    border-color: {colors.PRIMARY};
}}

/* 输入框样式 */
QLineEdit {{
    border: 1px solid {colors.BORDER};
    border-radius: 6px;
    padding: 8px 12px;
    background-color: {colors.BG_CARD};
    font-size: 13px;
    color: {colors.TEXT_PRIMARY};
}}

QLineEdit:focus {{
    border-color: {colors.PRIMARY};
}}

/* 下拉框样式 */
QComboBox {{
    border: 1px solid {colors.BORDER};
    border-radius: 6px;
    padding: 8px 12px;
    background-color: {colors.BG_CARD};
    font-size: 13px;
    color: {colors.TEXT_PRIMARY};
}}

QComboBox:focus {{
    border-color: {colors.PRIMARY};
}}

QComboBox QAbstractItemView {{
    background-color: {colors.BG_CARD};
    color: {colors.TEXT_PRIMARY};
    selection-background-color: {colors.PRIMARY_LIGHT};
}}

/* 复选框样式 */
QCheckBox {{
    font-size: 13px;
    spacing: 8px;
    color: {colors.TEXT_PRIMARY};
}}

QCheckBox::indicator {{
    width: 18px;
    height: 18px;
    border: 2px solid {colors.BORDER};
    border-radius: 4px;
    background-color: {colors.BG_CARD};
}}

QCheckBox::indicator:checked {{
    background-color: {colors.PRIMARY};
    border-color: {colors.PRIMARY};
}}

/* 进度条样式 */
QProgressBar {{
    border: none;
    border-radius: 6px;
    background-color: {colors.BORDER};
    height: 8px;
}}

QProgressBar::chunk {{
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
        stop:0 {colors.PRIMARY}, stop:1 {colors.SUCCESS});
    border-radius: 6px;
}}

/* 文本编辑框样式 */
QTextEdit {{
    border: 1px solid {colors.BORDER};
    border-radius: 8px;
    padding: 12px;
    background-color: {colors.BG_CARD};
    font-family: "SF Mono", "Menlo", "Consolas", monospace;
    font-size: 12px;
    color: {colors.TEXT_PRIMARY};
}}

/* 标签样式 */
QLabel {{
    font-size: 13px;
    color: {colors.TEXT_PRIMARY};
}}

QLabel#titleLabel {{
    font-size: 16px;
    font-weight: 600;
    color: {colors.PRIMARY};
}}

/* 滚动条样式 */
QScrollBar:vertical {{
    border: none;
    background-color: {colors.BG_LIGHT};
    width: 10px;
    border-radius: 5px;
}}

QScrollBar::handle:vertical {{
    background-color: {colors.BORDER};
    border-radius: 5px;
    min-height: 30px;
}}

/* SpinBox 样式 */
QSpinBox {{
    border: 1px solid {colors.BORDER};
    border-radius: 6px;
    padding: 6px 10px;
    background-color: {colors.BG_CARD};
    color: {colors.TEXT_PRIMARY};
}}

/* 表格样式 */
QTableWidget {{
    border: 1px solid {colors.BORDER};
    border-radius: 6px;
    background-color: {colors.BG_CARD};
    color: {colors.TEXT_PRIMARY};
    gridline-color: {colors.BORDER};
}}

QTableWidget::item {{
    padding: 6px;
}}

QHeaderView::section {{
    background-color: {colors.BG_LIGHT};
    color: {colors.TEXT_PRIMARY};
    padding: 8px;
    border: none;
    border-bottom: 1px solid {colors.BORDER};
    font-weight: 600;
}}
"""


def apply_theme(app: QApplication, theme: str = 'light'):
    """应用主题"""
    app.setStyle('Fusion')
    colors = Colors if theme == 'light' else DarkColors
    app.setStyleSheet(get_stylesheet(colors))
    
    font = QFont()
    font.setFamily("Microsoft YaHei" if platform.system() == 'Windows' else "PingFang SC")
    font.setPointSize(10)
    app.setFont(font)


# ==================== 配置管理器 ====================
class ConfigManager:
    """配置管理器"""
    
    DEFAULT_CONFIG = {
        'theme': 'light',
        'use_number_formatting': True,
        'recursive_scan': False,
        'last_input_dir': '',
        'last_output_dir': '',
        'last_template_dir': '',
        'qr_size_cm': 3.0,
        'extraction_configs': {}
    }
    
    def __init__(self, config_file: str = CONFIG_FILE):
        self.config_file = config_file
        self.config = self._load()
    
    def _load(self) -> Dict:
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    loaded = json.load(f)
                    return {**self.DEFAULT_CONFIG, **loaded}
        except Exception:
            pass
        return self.DEFAULT_CONFIG.copy()
    
    def save(self):
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置失败: {e}")
    
    def get(self, key: str, default=None):
        return self.config.get(key, default)
    
    def set(self, key: str, value):
        self.config[key] = value
        self.save()
    
    # 提取配置专用
    def get_extraction_config(self, name: str) -> Dict:
        return self.config.get('extraction_configs', {}).get(name, {})
    
    def set_extraction_config(self, name: str, data: Dict):
        if 'extraction_configs' not in self.config:
            self.config['extraction_configs'] = {}
        self.config['extraction_configs'][name] = data
        self.save()
    
    def delete_extraction_config(self, name: str):
        if name in self.config.get('extraction_configs', {}):
            del self.config['extraction_configs'][name]
            self.save()
    
    @property
    def extraction_config_names(self) -> List[str]:
        return list(self.config.get('extraction_configs', {}).keys())


# ==================== 拖拽输入框 ====================
class DragDropLineEdit(QLineEdit):
    """支持拖拽的输入框"""
    
    path_dropped = pyqtSignal(str)
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setAcceptDrops(True)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    
    def dropEvent(self, event: QDropEvent):
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            self.setText(path)
            self.path_dropped.emit(path)


# ==================== 单元格解析器 ====================
class CellRangeParser:
    """单元格范围解析器 - 支持 A1, A1:A10, B4-8 等格式"""
    
    @staticmethod
    def parse(cell_text: str) -> List[Tuple[int, int]]:
        """解析单元格表达式，返回 (row, col) 列表"""
        parsed_cells = []
        parts = re.split(r'[,;，；]', cell_text)
        
        for part in parts:
            part = part.strip().upper()
            if not part:
                continue
            
            # 格式1: B4-8 (同列多行)
            range_match = re.match(r'^([A-Z]+)(\d+)-(\d+)$', part)
            if range_match:
                col_str, start, end = range_match.groups()
                try:
                    col_idx = column_index_from_string(col_str)
                    for r in range(int(start), int(end) + 1):
                        if (r, col_idx) not in parsed_cells:
                            parsed_cells.append((r, col_idx))
                except Exception:
                    pass
                continue
            
            # 格式2: A1:C3 (矩形范围)
            rect_match = re.match(r'^([A-Z]+)(\d+):([A-Z]+)(\d+)$', part)
            if rect_match:
                col1_str, row1_str, col2_str, row2_str = rect_match.groups()
                try:
                    col1 = column_index_from_string(col1_str)
                    col2 = column_index_from_string(col2_str)
                    row1, row2 = int(row1_str), int(row2_str)
                    for r in range(min(row1, row2), max(row1, row2) + 1):
                        for c in range(min(col1, col2), max(col1, col2) + 1):
                            if (r, c) not in parsed_cells:
                                parsed_cells.append((r, c))
                except Exception:
                    pass
                continue
            
            # 格式3: 单个单元格 B4
            single_match = re.match(r'^([A-Z]+)(\d+)$', part)
            if single_match:
                col_str, row_str = single_match.groups()
                try:
                    col_idx = column_index_from_string(col_str)
                    if (int(row_str), col_idx) not in parsed_cells:
                        parsed_cells.append((int(row_str), col_idx))
                except Exception:
                    pass
        
        return parsed_cells


# ==================== 预览对话框 ====================
class PreviewDialog(QDialog):
    """数据预览对话框"""
    
    def __init__(self, data: List[List], headers: List[str], parent=None):
        super().__init__(parent)
        self.setWindowTitle('数据预览')
        self.setMinimumSize(700, 450)
        
        layout = QVBoxLayout(self)
        
        table = QTableWidget()
        if data:
            table.setColumnCount(len(headers))
            table.setHorizontalHeaderLabels(headers)
            table.setRowCount(min(10, len(data)))
            
            for row_idx, row_data in enumerate(data[:10]):
                for col_idx, val in enumerate(row_data):
                    item = QTableWidgetItem(str(val) if val else '')
                    table.setItem(row_idx, col_idx, item)
            
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        
        layout.addWidget(table)
        layout.addWidget(QLabel(f'共 {len(data)} 行数据，预览前 10 行'))
        
        btn_close = QPushButton('关闭')
        btn_close.clicked.connect(self.close)
        layout.addWidget(btn_close)


# ==================== 工作线程 ====================
class ExtractionThread(QThread):
    """Excel数据提取线程"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished_signal = pyqtSignal(str, bool)
    
    def __init__(self, input_paths: List[str], output_path: str, 
                 sheet_name: str, cells_text: str, headers_text: str, 
                 skip_header: bool, use_formatting: bool = True):
        super().__init__()
        self.input_paths = input_paths
        self.output_path = output_path
        self.sheet_name = sheet_name
        self.cells_text = cells_text
        self.headers_text = headers_text
        self.skip_header = skip_header
        self.use_formatting = use_formatting
        self.is_running = True
    
    def run(self):
        self.log.emit('🚀 开始批量提取...')
        
        cells = CellRangeParser.parse(self.cells_text)
        if not cells:
            self.log.emit('❌ 单元格位置列表为空或解析失败')
            self.finished_signal.emit('', False)
            return
        
        self.log.emit(f'📍 解析到 {len(cells)} 个单元格位置')
        
        # 准备表头
        headers = [h.strip() for h in re.split(r'[,;，；]', self.headers_text) if h.strip()]
        while len(headers) < len(cells):
            headers.append(f'Column_{len(headers) + 1}')
        
        all_data = []
        if not self.skip_header:
            all_data.append(['源文件路径'] + headers[:len(cells)])
        
        # 处理每个文件
        total = len(self.input_paths)
        for i, file_path in enumerate(self.input_paths):
            if not self.is_running:
                self.log.emit('⏸️ 任务已中断')
                self.finished_signal.emit('', False)
                return
            
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                ws = wb[self.sheet_name] if self.sheet_name and self.sheet_name in wb.sheetnames else wb.active
                
                row_data = [file_path]
                for row, col in cells:
                    value = ws.cell(row=row, column=col).value
                    formatted = format_number(value, self.use_formatting) if isinstance(value, (int, float)) else ('' if value is None else str(value))
                    row_data.append(formatted)
                
                all_data.append(row_data)
                self.log.emit(f'✅ {os.path.basename(file_path)}')
                wb.close()
            except Exception as e:
                self.log.emit(f'❌ {os.path.basename(file_path)}: {e}')
            
            self.progress.emit(int((i + 1) / total * 100))
        
        # 保存结果
        try:
            self.log.emit('💾 保存汇总文件...')
            os.makedirs(os.path.dirname(os.path.abspath(self.output_path)), exist_ok=True)
            pd.DataFrame(all_data).to_excel(self.output_path, index=False, header=False)
            self.log.emit('🎉 提取完成！')
            self.finished_signal.emit(self.output_path, True)
        except Exception as e:
            self.log.emit(f'❌ 保存失败: {e}')
            self.finished_signal.emit('', False)


class WordGenerationThread(QThread):
    """Word文档生成线程"""
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    finished = pyqtSignal(int, int)
    
    FONT_SIZE_WUHAO = Pt(10.5)
    
    def __init__(self, excel_path: str, template_path: str, output_dir: str, 
                 placeholder_fmt: str, filename_template: str,
                 use_formatting: bool = True, qr_size_cm: float = 3.0):
        super().__init__()
        self.excel_path = excel_path
        self.template_path = template_path
        self.output_dir = output_dir
        self.placeholder_fmt = placeholder_fmt
        self.filename_template = filename_template
        self.use_formatting = use_formatting
        self.qr_size_cm = qr_size_cm
        self.is_running = True
    
    def apply_mixed_font_to_run(self, run, is_chinese: bool, font_size=None, 
                                  preserve_bold=False, preserve_italic=False, preserve_underline=False):
        """
        给单个run应用混合字体
        - 保留：加粗、斜体、下划线
        - 清除：字体颜色（改为黑色）、底纹颜色
        """
        # 设置字号
        if font_size is not None:
            run.font.size = font_size
        
        # 设置字体
        if is_chinese:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # 保留格式
        if preserve_bold:
            run.font.bold = True
        if preserve_italic:
            run.font.italic = True
        if preserve_underline:
            run.font.underline = True
        
        # 清除字体颜色（设为黑色）
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 清除底纹/高亮
        run.font.highlight_color = None
        # 清除shading（底纹）
        try:
            rPr = run._element.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)
        except:
            pass
    
    def replace_in_paragraph_preserve_format(self, paragraph, row_data: Dict, doc, apply_font_size=None):
        """
        替换段落中的占位符，精确保留原有格式（加粗、斜体、下划线）
        清除颜色和底纹
        """
        # 检查段落是否包含分页符，需要保留
        has_page_break = False
        page_break_elements = []
        for run in paragraph.runs:
            for child in run._element:
                if child.tag == qn('w:br'):
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        has_page_break = True
        
        original_text = paragraph.text
        if not original_text:
            # 即使没有文字，也要保留分页符
            return
        
        # 检查是否有占位符需要替换
        new_text = original_text
        has_replacement = False
        
        # 处理二维码占位符 {QR:字段名}
        qr_data_list = []
        qr_matches = re.findall(r'\{QR:([^}]+)\}', new_text)
        for qr_field in qr_matches:
            if qr_field in row_data:
                qr_data = str(row_data[qr_field]) if pd.notna(row_data[qr_field]) else ''
            else:
                qr_data = qr_field
            
            placeholder = f'{{QR:{qr_field}}}'
            new_text = new_text.replace(placeholder, '')
            has_replacement = True
            if qr_data:
                qr_data_list.append(qr_data)
        
        # 处理普通占位符，记录替换信息
        replacements = []
        for col, val in row_data.items():
            placeholder = self.placeholder_fmt.replace('ColumnName', str(col))
            if placeholder in new_text:
                if isinstance(val, (int, float)) and not pd.isna(val):
                    val_str = format_number(val, self.use_formatting)
                else:
                    val_str = '' if pd.isna(val) else str(val)
                replacements.append((placeholder, val_str))
                has_replacement = True
        
        # 如果没有任何替换，只需要处理格式（清除颜色、应用字号字体）
        if not has_replacement:
            for run in paragraph.runs:
                # 保留原有的加粗、斜体、下划线
                is_bold = run.font.bold
                is_italic = run.font.italic
                is_underline = run.font.underline
                
                # 应用字体
                text = run.text
                if text:
                    for char in text:
                        is_chinese = is_chinese_char(char)
                        break
                    else:
                        is_chinese = False
                    
                    if is_chinese:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # 应用字号
                if apply_font_size is not None:
                    run.font.size = apply_font_size
                
                # 保留格式
                if is_bold:
                    run.font.bold = True
                if is_italic:
                    run.font.italic = True
                if is_underline:
                    run.font.underline = True
                
                # 清除颜色和底纹
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.highlight_color = None
                try:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is not None:
                        rPr.remove(shd)
                except:
                    pass
            return
        
        # 收集原有runs的详细格式信息（按字符位置）
        char_formats = []  # 每个字符的格式
        char_pos = 0
        for run in paragraph.runs:
            run_text = run.text
            is_bold = run.font.bold if run.font.bold else False
            is_italic = run.font.italic if run.font.italic else False
            is_underline = run.font.underline if run.font.underline else False
            
            for i, char in enumerate(run_text):
                char_formats.append({
                    'char': char,
                    'pos': char_pos + i,
                    'bold': is_bold,
                    'italic': is_italic,
                    'underline': is_underline,
                })
            char_pos += len(run_text)
        
        # 执行替换，并建立新旧字符位置的映射
        # 为了精确映射，我们需要逐个处理替换
        result_chars = []  # [(char, bold, italic, underline), ...]
        
        temp_text = original_text
        offset = 0
        
        for placeholder, val_str in replacements:
            idx = temp_text.find(placeholder)
            while idx != -1:
                # 添加占位符之前的字符（保持原格式）
                for i in range(offset, offset + idx):
                    if i < len(char_formats):
                        fmt = char_formats[i]
                        result_chars.append((fmt['char'], fmt['bold'], fmt['italic'], fmt['underline']))
                
                # 添加替换值（继承占位符第一个字符的格式）
                placeholder_start = offset + idx
                if placeholder_start < len(char_formats):
                    inherit_fmt = char_formats[placeholder_start]
                else:
                    inherit_fmt = {'bold': False, 'italic': False, 'underline': False}
                
                for char in val_str:
                    result_chars.append((char, inherit_fmt['bold'], inherit_fmt['italic'], inherit_fmt['underline']))
                
                # 更新offset和temp_text
                offset = offset + idx + len(placeholder)
                temp_text = temp_text[idx + len(placeholder):]
                idx = temp_text.find(placeholder)
        
        # 添加剩余字符
        for i in range(offset, len(char_formats)):
            fmt = char_formats[i]
            result_chars.append((fmt['char'], fmt['bold'], fmt['italic'], fmt['underline']))
        
        # 处理QR占位符（从result_chars中移除）
        final_text = ''.join([c[0] for c in result_chars])
        for qr_field in qr_matches:
            placeholder = f'{{QR:{qr_field}}}'
            final_text = final_text.replace(placeholder, '')
        
        # 重建result_chars
        if qr_matches:
            new_result_chars = []
            temp = ''.join([c[0] for c in result_chars])
            i = 0
            j = 0
            while i < len(temp) and j < len(final_text):
                if temp[i] == final_text[j]:
                    new_result_chars.append(result_chars[i])
                    i += 1
                    j += 1
                else:
                    i += 1
            result_chars = new_result_chars
        
        # 清除段落内容（保留段落元素）
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        
        # 按格式分组重建runs
        if result_chars:
            current_fmt = (result_chars[0][1], result_chars[0][2], result_chars[0][3])
            current_text = result_chars[0][0]
            
            for char, bold, italic, underline in result_chars[1:]:
                fmt = (bold, italic, underline)
                if fmt == current_fmt:
                    current_text += char
                else:
                    # 创建新run
                    self._add_formatted_run(paragraph, current_text, current_fmt[0], current_fmt[1], current_fmt[2], apply_font_size)
                    current_fmt = fmt
                    current_text = char
            
            # 添加最后一个run
            if current_text:
                self._add_formatted_run(paragraph, current_text, current_fmt[0], current_fmt[1], current_fmt[2], apply_font_size)
        
        # 添加分页符（如果原来有）
        if has_page_break:
            run = paragraph.add_run()
            run._element.append(OxmlElement('w:br'))
            run._element[-1].set(qn('w:type'), 'page')
        
        # 添加二维码
        for qr_data in qr_data_list:
            if HAS_QRCODE:
                try:
                    qr_img = generate_qrcode_image(qr_data)
                    run = paragraph.add_run()
                    run.add_picture(qr_img, width=Cm(self.qr_size_cm))
                except Exception as e:
                    self.log.emit(f'⚠️ 二维码生成失败: {e}')
    
    def _add_formatted_run(self, paragraph, text: str, bold: bool, italic: bool, underline: bool, font_size=None):
        """添加带格式的run，按中英文分段"""
        segments = split_text_by_language(text)
        for segment_text, is_chinese in segments:
            run = paragraph.add_run(segment_text)
            self.apply_mixed_font_to_run(run, is_chinese, font_size, bold, italic, underline)
    
    def apply_mixed_font(self, paragraph, text: str, font_size=None):
        """应用混合字体（简单版本，用于新建内容）"""
        paragraph.clear()
        
        segments = split_text_by_language(text)
        for segment_text, is_chinese in segments:
            run = paragraph.add_run(segment_text)
            self.apply_mixed_font_to_run(run, is_chinese, font_size)
    
    def replace_in_paragraph(self, paragraph, row_data: Dict, doc, apply_font_size=None):
        """替换段落中的占位符（调用保留格式版本）"""
        self.replace_in_paragraph_preserve_format(paragraph, row_data, doc, apply_font_size)
    
    def find_first_page_break_index(self, doc) -> int:
        """查找第一个分页符的位置"""
        for i, para in enumerate(doc.paragraphs):
            for run in para.runs:
                if run._element.xml.find('w:br') != -1 and 'w:type="page"' in run._element.xml:
                    return i + 1
            if para._element.xml.find('w:pageBreakBefore') != -1:
                return i
        return -1
    
    def run(self):
        try:
            self.log.emit('📄 读取Excel数据源...')
            df = pd.read_excel(self.excel_path)
            df.columns = [str(c).strip() for c in df.columns]
            
            total = len(df)
            success_count = 0
            
            os.makedirs(self.output_dir, exist_ok=True)
            
            for idx, row in df.iterrows():
                if not self.is_running:
                    break
                
                try:
                    row_data = row.to_dict()
                    doc = Document(self.template_path)
                    
                    page_break_index = self.find_first_page_break_index(doc)
                    
                    for i, para in enumerate(doc.paragraphs):
                        is_cover_page = (page_break_index == -1) or (i < page_break_index)
                        
                        if is_cover_page:
                            self.replace_in_paragraph(para, row_data, doc, apply_font_size=None)
                        else:
                            self.replace_in_paragraph(para, row_data, doc, apply_font_size=self.FONT_SIZE_WUHAO)
                    
                    for table in doc.tables:
                        for row_cells in table.rows:
                            for cell in row_cells.cells:
                                for para in cell.paragraphs:
                                    self.replace_in_paragraph(para, row_data, doc, apply_font_size=self.FONT_SIZE_WUHAO)
                    
                    # 生成文件名
                    filename = self.filename_template
                    for col, val in row_data.items():
                        val_str = str(val) if pd.notna(val) else ''
                        filename = filename.replace(f'{{{col}}}', val_str)
                    filename = re.sub(r'[\\/:*?"<>|]', '_', filename)
                    if not filename.endswith('.docx'):
                        filename += '.docx'
                    
                    output_path = os.path.join(self.output_dir, filename)
                    doc.save(output_path)
                    self.log.emit(f'✅ {filename}')
                    success_count += 1
                    
                except Exception as e:
                    self.log.emit(f'❌ 第{idx + 1}行失败: {e}')
                
                self.progress.emit(int((idx + 1) / total * 100))
            
            self.finished.emit(success_count, total - success_count)
            
        except Exception as e:
            self.log.emit(f'❌ 严重错误: {e}')
            self.finished.emit(0, 0)


# ==================== 打赏对话框 ====================
class DonateDialog(QDialog):
    """打赏对话框 - 微信和支付宝双二维码"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('💝 支持作者')
        self.setFixedSize(520, 500)
        self._setup_ui()
    
    def _get_qr_image_path(self, name: str) -> str:
        """获取二维码图片路径，支持多种可能的位置"""
        # 可能的路径列表
        possible_paths = [
            # 与脚本同目录
            os.path.join(os.path.dirname(os.path.abspath(__file__)), name),
            # 当前工作目录
            os.path.join(os.getcwd(), name),
            # Resources 子目录
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Resources', name),
            os.path.join(os.getcwd(), 'Resources', name),
            # assets 子目录
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', name),
            os.path.join(os.getcwd(), 'assets', name),
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        return ""
    
    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(30, 25, 30, 25)
        
        # 标题
        title = QLabel('感谢您的支持！')
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet(f'''
            font-size: 20px;
            font-weight: bold;
            color: {Colors.PRIMARY};
            padding: 5px;
        ''')
        layout.addWidget(title)
        
        desc = QLabel('如果这个工具对您有帮助，欢迎请作者喝杯咖啡 ☕')
        desc.setAlignment(Qt.AlignCenter)
        desc.setStyleSheet(f'color: {Colors.TEXT_SECONDARY}; font-size: 13px;')
        layout.addWidget(desc)
        
        # 打赏区域
        donate_container = QHBoxLayout()
        donate_container.setSpacing(25)
        
        # 微信支付
        wechat_widget = QWidget()
        wechat_layout = QVBoxLayout(wechat_widget)
        wechat_layout.setAlignment(Qt.AlignCenter)
        wechat_layout.setSpacing(8)
        
        wechat_label = QLabel('微信支付')
        wechat_label.setAlignment(Qt.AlignCenter)
        wechat_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #07C160;')
        wechat_layout.addWidget(wechat_label)
        
        # 微信二维码
        wechat_qr_label = QLabel()
        wechat_qr_label.setFixedSize(160, 160)
        wechat_qr_label.setAlignment(Qt.AlignCenter)
        wechat_qr_label.setStyleSheet('''
            background-color: white;
            border-radius: 10px;
            border: 3px solid #07C160;
        ''')
        
        # 尝试加载微信二维码图片
        wx_path = self._get_qr_image_path('wx.jpg')
        if wx_path:
            pixmap = QPixmap(wx_path)
            if not pixmap.isNull():
                wechat_qr_label.setPixmap(pixmap.scaled(154, 154, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            else:
                wechat_qr_label.setText('💚\n微信扫码')
                wechat_qr_label.setStyleSheet('''
                    font-size: 20px; background-color: white; border-radius: 10px;
                    border: 3px solid #07C160; color: #07C160;
                ''')
        else:
            wechat_qr_label.setText('💚\n微信扫码')
            wechat_qr_label.setStyleSheet('''
                font-size: 20px; background-color: white; border-radius: 10px;
                border: 3px solid #07C160; color: #07C160;
            ''')
        
        wechat_layout.addWidget(wechat_qr_label, alignment=Qt.AlignCenter)
        
        wechat_hint = QLabel('微信扫一扫')
        wechat_hint.setAlignment(Qt.AlignCenter)
        wechat_hint.setStyleSheet('font-size: 12px; color: #07C160;')
        wechat_layout.addWidget(wechat_hint)
        
        donate_container.addWidget(wechat_widget)
        
        # 分隔线
        separator = QFrame()
        separator.setFrameShape(QFrame.VLine)
        separator.setStyleSheet(f'background-color: {Colors.BORDER};')
        donate_container.addWidget(separator)
        
        # 支付宝
        alipay_widget = QWidget()
        alipay_layout = QVBoxLayout(alipay_widget)
        alipay_layout.setAlignment(Qt.AlignCenter)
        alipay_layout.setSpacing(8)
        
        alipay_label = QLabel('支付宝')
        alipay_label.setAlignment(Qt.AlignCenter)
        alipay_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #1677FF;')
        alipay_layout.addWidget(alipay_label)
        
        # 支付宝二维码
        alipay_qr_label = QLabel()
        alipay_qr_label.setFixedSize(160, 160)
        alipay_qr_label.setAlignment(Qt.AlignCenter)
        alipay_qr_label.setStyleSheet('''
            background-color: white;
            border-radius: 10px;
            border: 3px solid #1677FF;
        ''')
        
        # 尝试加载支付宝二维码图片
        zfb_path = self._get_qr_image_path('zfb.jpg')
        if zfb_path:
            pixmap = QPixmap(zfb_path)
            if not pixmap.isNull():
                alipay_qr_label.setPixmap(pixmap.scaled(154, 154, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            else:
                alipay_qr_label.setText('💙\n支付宝扫码')
                alipay_qr_label.setStyleSheet('''
                    font-size: 20px; background-color: white; border-radius: 10px;
                    border: 3px solid #1677FF; color: #1677FF;
                ''')
        else:
            alipay_qr_label.setText('💙\n支付宝扫码')
            alipay_qr_label.setStyleSheet('''
                font-size: 20px; background-color: white; border-radius: 10px;
                border: 3px solid #1677FF; color: #1677FF;
            ''')
        
        alipay_layout.addWidget(alipay_qr_label, alignment=Qt.AlignCenter)
        
        alipay_hint = QLabel('支付宝扫一扫')
        alipay_hint.setAlignment(Qt.AlignCenter)
        alipay_hint.setStyleSheet('font-size: 12px; color: #1677FF;')
        alipay_layout.addWidget(alipay_hint)
        
        donate_container.addWidget(alipay_widget)
        layout.addLayout(donate_container)
        
        # 感谢语
        thanks_label = QLabel('感谢您对大鑽戒基金會的支持')
        thanks_label.setAlignment(Qt.AlignCenter)
        thanks_label.setStyleSheet(f'''
            font-size: 15px; font-weight: 500; color: {Colors.PRIMARY};
            padding: 15px 0 5px 0;
        ''')
        layout.addWidget(thanks_label)
        
        # 作者信息
        author_info = QLabel(f'Author: {APP_AUTHOR}  |  Version: {APP_VERSION}')
        author_info.setAlignment(Qt.AlignCenter)
        author_info.setStyleSheet(f'color: {Colors.TEXT_SECONDARY}; font-size: 11px;')
        layout.addWidget(author_info)
        
        # 关闭按钮
        close_btn = QPushButton('关闭')
        close_btn.setFixedWidth(120)
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn, alignment=Qt.AlignCenter)


# ==================== Excel 提取页面 ====================
class ExtractionWidget(QWidget):
    """Excel提取整合页面"""
    
    def __init__(self, config_manager: ConfigManager, parent=None):
        super().__init__(parent)
        self.config_manager = config_manager
        self.thread = None
        self._setup_ui()
        self._update_config_combo()
    
    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 文件路径区域
        path_group = QGroupBox('📁 文件路径')
        path_layout = QVBoxLayout()
        path_layout.setSpacing(12)
        path_layout.setContentsMargins(15, 20, 15, 15)
        
        # Excel文件夹
        row1 = QHBoxLayout()
        label1 = QLabel('Excel文件夹:')
        label1.setFixedWidth(85)
        row1.addWidget(label1)
        self.input_dir_edit = DragDropLineEdit()
        self.input_dir_edit.setPlaceholderText('选择目录或拖拽到此处...')
        last_dir = self.config_manager.get('last_input_dir', '')
        if last_dir and os.path.isdir(last_dir):
            self.input_dir_edit.setText(last_dir)
        row1.addWidget(self.input_dir_edit, 1)
        btn_input = QPushButton('浏览')
        btn_input.setObjectName('secondaryButton')
        btn_input.setFixedWidth(70)
        btn_input.clicked.connect(self._select_input_dir)
        row1.addWidget(btn_input)
        path_layout.addLayout(row1)
        
        # 输出文件
        row2 = QHBoxLayout()
        label2 = QLabel('输出文件:')
        label2.setFixedWidth(85)
        row2.addWidget(label2)
        self.output_file_edit = DragDropLineEdit()
        self.output_file_edit.setText(os.path.join(os.getcwd(), DEFAULT_EXTRACTION_FILE))
        row2.addWidget(self.output_file_edit, 1)
        btn_output = QPushButton('浏览')
        btn_output.setObjectName('secondaryButton')
        btn_output.setFixedWidth(70)
        btn_output.clicked.connect(self._select_output_file)
        row2.addWidget(btn_output)
        path_layout.addLayout(row2)
        
        path_group.setLayout(path_layout)
        layout.addWidget(path_group)
        
        # 添加间距
        layout.addSpacing(10)
        
        # 设置区域
        config_group = QGroupBox('  ⚙️ 设置')  # 前面加空格让标题显示完整
        config_layout = QVBoxLayout()
        config_layout.setSpacing(20)
        config_layout.setContentsMargins(15, 30, 15, 20)
        
        # 第一行：已保存配置 + 保存/删除按钮
        config_row1 = QHBoxLayout()
        label_cfg = QLabel('已保存配置:')
        label_cfg.setFixedWidth(85)
        config_row1.addWidget(label_cfg)
        self.config_combo = QComboBox()
        self.config_combo.setMinimumWidth(150)
        self.config_combo.setFixedHeight(36)
        self.config_combo.currentTextChanged.connect(self._load_config)
        config_row1.addWidget(self.config_combo, 1)
        config_row1.addSpacing(15)
        btn_save_config = QPushButton('💾 保存')
        btn_save_config.setObjectName('secondaryButton')
        btn_save_config.setFixedWidth(85)
        btn_save_config.setFixedHeight(36)
        btn_save_config.clicked.connect(self._save_config)
        config_row1.addWidget(btn_save_config)
        btn_del_config = QPushButton('🗑️ 删除')
        btn_del_config.setObjectName('secondaryButton')
        btn_del_config.setFixedWidth(85)
        btn_del_config.setFixedHeight(36)
        btn_del_config.clicked.connect(self._delete_config)
        config_row1.addWidget(btn_del_config)
        config_layout.addLayout(config_row1)
        
        # 第二行：工作表名称 + 单元格位置（两列）
        config_row2 = QHBoxLayout()
        config_row2.setSpacing(30)
        
        # 左侧：工作表名称
        left_col = QHBoxLayout()
        label_sheet = QLabel('工作表名称:')
        label_sheet.setFixedWidth(85)
        left_col.addWidget(label_sheet)
        self.sheet_edit = QLineEdit()
        self.sheet_edit.setPlaceholderText('留空则读取第一个Sheet')
        self.sheet_edit.setFixedHeight(36)
        left_col.addWidget(self.sheet_edit, 1)
        config_row2.addLayout(left_col, 1)
        
        # 右侧：单元格位置
        right_col = QHBoxLayout()
        label_cells = QLabel('单元格位置:')
        label_cells.setFixedWidth(85)
        right_col.addWidget(label_cells)
        self.cells_edit = QLineEdit()
        self.cells_edit.setPlaceholderText('例如: B4, C14-18, D5, A1:A10')
        self.cells_edit.setFixedHeight(36)
        right_col.addWidget(self.cells_edit, 1)
        config_row2.addLayout(right_col, 1)
        
        config_layout.addLayout(config_row2)
        
        # 第三行：列标题
        config_row3 = QHBoxLayout()
        label_headers = QLabel('列标题:')
        label_headers.setFixedWidth(85)
        config_row3.addWidget(label_headers)
        self.headers_edit = QLineEdit()
        self.headers_edit.setPlaceholderText('用逗号分隔，例如: 姓名, 金额, 日期')
        self.headers_edit.setFixedHeight(36)
        config_row3.addWidget(self.headers_edit, 1)
        config_layout.addLayout(config_row3)
        
        config_group.setLayout(config_layout)
        layout.addWidget(config_group, 2)  # 给设置区域更大的伸展空间
        
        # 添加间距
        layout.addSpacing(10)
        
        # 执行区域
        exec_layout = QHBoxLayout()
        exec_layout.setSpacing(12)
        
        btn_preview = QPushButton('👁 预览数据')
        btn_preview.setObjectName('secondaryButton')
        btn_preview.setFixedHeight(38)
        btn_preview.clicked.connect(self._preview_data)
        exec_layout.addWidget(btn_preview)
        
        self.btn_run = QPushButton('🚀 开始提取整合')
        self.btn_run.setObjectName('primaryButton')
        self.btn_run.setFixedHeight(38)
        self.btn_run.clicked.connect(self._run_extraction)
        exec_layout.addWidget(self.btn_run, 2)
        
        self.btn_open_folder = QPushButton('📂 打开输出目录')
        self.btn_open_folder.setObjectName('secondaryButton')
        self.btn_open_folder.setFixedHeight(38)
        self.btn_open_folder.clicked.connect(self._open_output_folder)
        exec_layout.addWidget(self.btn_open_folder)
        
        layout.addLayout(exec_layout)
        
        # 进度条
        self.progress = QProgressBar()
        self.progress.setTextVisible(False)
        self.progress.setFixedHeight(6)
        layout.addWidget(self.progress)
        
        # 日志区域
        log_header = QHBoxLayout()
        log_label = QLabel('📋 运行日志')
        log_label.setObjectName('titleLabel')
        log_header.addWidget(log_label)
        log_header.addStretch()
        btn_export_log = QPushButton('📝 导出日志')
        btn_export_log.setObjectName('secondaryButton')
        btn_export_log.clicked.connect(self._export_log)
        log_header.addWidget(btn_export_log)
        layout.addLayout(log_header)
        
        self.log_view = QTextEdit()
        self.log_view.setReadOnly(True)
        self.log_view.setFixedHeight(80)
        layout.addWidget(self.log_view)
    
    def _get_excel_files(self, input_dir: str) -> List[str]:
        """获取 Excel 文件列表"""
        excel_files = [
            os.path.join(input_dir, f) for f in os.listdir(input_dir)
            if f.endswith(('.xlsx', '.xls')) and not f.startswith('~')
        ]
        return sorted(excel_files)
    
    def _select_input_dir(self):
        start_dir = self.config_manager.get('last_input_dir', '')
        path = QFileDialog.getExistingDirectory(self, '选择目录', start_dir)
        if path:
            self.input_dir_edit.setText(path)
            self.config_manager.set('last_input_dir', path)
    
    def _select_output_file(self):
        path, _ = QFileDialog.getSaveFileName(
            self, '保存文件', DEFAULT_EXTRACTION_FILE, 'Excel Files (*.xlsx)')
        if path:
            self.output_file_edit.setText(path)
    
    def _open_output_folder(self):
        if not open_folder(self.output_file_edit.text()):
            QMessageBox.warning(self, '提示', '目录不存在，请先执行提取操作')
    
    def _update_config_combo(self):
        self.config_combo.blockSignals(True)
        self.config_combo.clear()
        self.config_combo.addItems(self.config_manager.extraction_config_names)
        self.config_combo.blockSignals(False)
    
    def _save_config(self):
        name, ok = QInputDialog.getText(self, '保存配置', '请输入配置名称:')
        if ok and name:
            self.config_manager.set_extraction_config(name, {
                'sheet': self.sheet_edit.text(),
                'cells': self.cells_edit.text(),
                'headers': self.headers_edit.text()
            })
            self._update_config_combo()
            self.config_combo.setCurrentText(name)
            self._log('💾 配置已保存')
    
    def _load_config(self, name: str):
        if not name:
            return
        data = self.config_manager.get_extraction_config(name)
        self.sheet_edit.setText(data.get('sheet', ''))
        self.cells_edit.setText(data.get('cells', ''))
        self.headers_edit.setText(data.get('headers', ''))
    
    def _delete_config(self):
        name = self.config_combo.currentText()
        if name:
            self.config_manager.delete_extraction_config(name)
            self._update_config_combo()
            self.sheet_edit.clear()
            self.cells_edit.clear()
            self.headers_edit.clear()
            self._log('🗑️ 配置已删除')
    
    def _preview_data(self):
        """预览数据"""
        input_dir = self.input_dir_edit.text()
        if not os.path.isdir(input_dir):
            QMessageBox.warning(self, '错误', '请选择有效的输入目录')
            return
        
        files = self._get_excel_files(input_dir)[:5]
        if not files:
            QMessageBox.warning(self, '错误', '目录下没有Excel文件')
            return
        
        cells = CellRangeParser.parse(self.cells_edit.text())
        if not cells:
            QMessageBox.warning(self, '错误', '请先配置单元格位置')
            return
        
        headers = [h.strip() for h in re.split(r'[,;，；]', self.headers_edit.text()) if h.strip()]
        while len(headers) < len(cells):
            headers.append(f'Column_{len(headers) + 1}')
        
        preview_data = []
        for file_path in files:
            try:
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheet_name = self.sheet_edit.text()
                ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
                
                row_data = [os.path.basename(file_path)]
                for row, col in cells:
                    value = ws.cell(row=row, column=col).value
                    row_data.append('' if value is None else str(value))
                
                preview_data.append(row_data)
                wb.close()
            except Exception as e:
                preview_data.append([os.path.basename(file_path), f'错误: {e}'])
        
        dialog = PreviewDialog(preview_data, ['源文件'] + headers[:len(cells)], self)
        dialog.exec_()
    
    def _export_log(self):
        """导出日志"""
        log_content = self.log_view.toPlainText()
        if not log_content:
            QMessageBox.information(self, '提示', '日志为空')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, '导出日志', f'extraction_log_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt',
            'Text Files (*.txt)'
        )
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(log_content)
            QMessageBox.information(self, '成功', f'日志已导出')
    
    def _log(self, msg: str):
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.log_view.append(f'[{timestamp}] {msg}')
    
    def _run_extraction(self):
        input_dir = self.input_dir_edit.text()
        if not os.path.isdir(input_dir):
            QMessageBox.warning(self, '错误', '请选择有效的输入目录')
            return
        
        files = self._get_excel_files(input_dir)
        if not files:
            QMessageBox.warning(self, '错误', '目录下没有Excel文件')
            return
        
        self.btn_run.setEnabled(False)
        self.log_view.clear()
        self.progress.setValue(0)
        
        self.thread = ExtractionThread(
            files,
            self.output_file_edit.text(),
            self.sheet_edit.text(),
            self.cells_edit.text(),
            self.headers_edit.text(),
            False,  # skip_header 默认 False
            True    # use_formatting 默认 True
        )
        self.thread.log.connect(self._log)
        self.thread.progress.connect(self.progress.setValue)
        self.thread.finished_signal.connect(self._on_finished)
        self.thread.start()
    
    def _on_finished(self, output_path: str, success: bool):
        self.btn_run.setEnabled(True)
        if success:
            main_win = self.window()
            if isinstance(main_win, MainWindow):
                main_win.on_extraction_complete(output_path)


# ==================== Word 生成页面 ====================
class WordGenWidget(QWidget):
    """Word生成页面"""
    
    def __init__(self, config_manager: ConfigManager, parent=None):
        super().__init__(parent)
        self.config_manager = config_manager
        self.thread = None
        self._setup_ui()
    
    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 文件设置区域
        file_group = QGroupBox('📁 文件设置')
        file_layout = QVBoxLayout()
        file_layout.setSpacing(12)
        file_layout.setContentsMargins(15, 20, 15, 15)
        
        # Excel数据源
        row1 = QHBoxLayout()
        label1 = QLabel('Excel数据源:')
        label1.setFixedWidth(85)
        row1.addWidget(label1)
        self.excel_edit = DragDropLineEdit()
        self.excel_edit.setPlaceholderText('选择Excel数据文件或拖拽到此处...')
        row1.addWidget(self.excel_edit, 1)
        btn_excel = QPushButton('浏览')
        btn_excel.setObjectName('secondaryButton')
        btn_excel.setFixedWidth(70)
        btn_excel.clicked.connect(lambda: self._select_file(self.excel_edit, 'Excel (*.xlsx)'))
        row1.addWidget(btn_excel)
        file_layout.addLayout(row1)
        
        # Word模板
        row2 = QHBoxLayout()
        label2 = QLabel('Word模板:')
        label2.setFixedWidth(85)
        row2.addWidget(label2)
        self.template_edit = DragDropLineEdit()
        self.template_edit.setPlaceholderText('选择Word模板文件或拖拽到此处...')
        row2.addWidget(self.template_edit, 1)
        btn_template = QPushButton('浏览')
        btn_template.setObjectName('secondaryButton')
        btn_template.setFixedWidth(70)
        btn_template.clicked.connect(lambda: self._select_file(self.template_edit, 'Word (*.docx)'))
        row2.addWidget(btn_template)
        file_layout.addLayout(row2)
        
        # 输出目录
        row3 = QHBoxLayout()
        label3 = QLabel('输出目录:')
        label3.setFixedWidth(85)
        row3.addWidget(label3)
        self.output_dir_edit = DragDropLineEdit()
        self.output_dir_edit.setText(os.path.join(os.getcwd(), DEFAULT_OUTPUT_DIR))
        row3.addWidget(self.output_dir_edit, 1)
        btn_output = QPushButton('浏览')
        btn_output.setObjectName('secondaryButton')
        btn_output.setFixedWidth(70)
        btn_output.clicked.connect(self._select_output_dir)
        row3.addWidget(btn_output)
        file_layout.addLayout(row3)
        
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        # 添加间距
        layout.addSpacing(10)
        
        # 设置区域
        rule_group = QGroupBox('  ⚙️ 设置')  # 前面加空格让标题显示完整
        rule_layout = QVBoxLayout()
        rule_layout.setSpacing(20)
        rule_layout.setContentsMargins(15, 30, 15, 20)
        
        # 第一行：占位符格式 + 文件名模板（两列）
        rule_row1 = QHBoxLayout()
        rule_row1.setSpacing(30)
        
        # 左侧：占位符格式
        left_col = QHBoxLayout()
        label_fmt = QLabel('占位符格式:')
        label_fmt.setFixedWidth(85)
        left_col.addWidget(label_fmt)
        self.placeholder_combo = QComboBox()
        self.placeholder_combo.addItems([
            '{ColumnName}', '[[ColumnName]]', '##ColumnName##', '$ColumnName$'
        ])
        self.placeholder_combo.setFixedHeight(36)
        left_col.addWidget(self.placeholder_combo, 1)
        rule_row1.addLayout(left_col, 1)
        
        # 右侧：文件名模板
        right_col = QHBoxLayout()
        label_fname = QLabel('文件名模板:')
        label_fname.setFixedWidth(85)
        right_col.addWidget(label_fname)
        self.filename_edit = QLineEdit()
        self.filename_edit.setText('{名称}_合同')
        self.filename_edit.setPlaceholderText('使用 {列名} 作为变量')
        self.filename_edit.setFixedHeight(36)
        right_col.addWidget(self.filename_edit, 1)
        rule_row1.addLayout(right_col, 1)
        
        rule_layout.addLayout(rule_row1)

        # 第二行：合并选项
        rule_row2 = QHBoxLayout()
        rule_row2.setSpacing(30)

        # 左侧：合并文档选项
        merge_col = QHBoxLayout()
        self.merge_checkbox = QCheckBox('生成完成后合并所有Word文档为一个文件')
        self.merge_checkbox.setChecked(False)
        merge_col.addWidget(self.merge_checkbox)
        merge_col.addStretch()
        rule_row2.addLayout(merge_col, 1)

        rule_layout.addLayout(rule_row2)

        rule_group.setLayout(rule_layout)
        layout.addWidget(rule_group, 2)  # 给设置区域更大的伸展空间
        
        # 添加间距
        layout.addSpacing(10)
        
        # 执行区域
        exec_layout = QHBoxLayout()
        exec_layout.setSpacing(12)
        
        self.btn_run = QPushButton('🚀 开始生成Word')
        self.btn_run.setObjectName('primaryButton')
        self.btn_run.setFixedHeight(38)
        self.btn_run.clicked.connect(self._start_generation)
        
        self.btn_open_folder = QPushButton('📂 打开输出目录')
        self.btn_open_folder.setObjectName('secondaryButton')
        self.btn_open_folder.setFixedHeight(38)
        self.btn_open_folder.clicked.connect(self._open_output_folder)
        
        exec_layout.addWidget(self.btn_run, 2)
        exec_layout.addWidget(self.btn_open_folder, 1)
        layout.addLayout(exec_layout)
        
        # 进度条
        self.progress = QProgressBar()
        self.progress.setTextVisible(False)
        self.progress.setFixedHeight(6)
        layout.addWidget(self.progress)
        
        # 日志区域
        log_header = QHBoxLayout()
        log_label = QLabel('📋 运行日志')
        log_label.setObjectName('titleLabel')
        log_header.addWidget(log_label)
        log_header.addStretch()
        btn_export_log = QPushButton('📝 导出日志')
        btn_export_log.setObjectName('secondaryButton')
        btn_export_log.clicked.connect(self._export_log)
        log_header.addWidget(btn_export_log)
        layout.addLayout(log_header)
        
        self.log_view = QTextEdit()
        self.log_view.setReadOnly(True)
        self.log_view.setFixedHeight(80)
        layout.addWidget(self.log_view)
    
    def _select_file(self, edit: QLineEdit, file_filter: str):
        path, _ = QFileDialog.getOpenFileName(self, '选择文件', '', file_filter)
        if path:
            edit.setText(path)
    
    def _select_output_dir(self):
        path = QFileDialog.getExistingDirectory(self, '选择目录')
        if path:
            self.output_dir_edit.setText(path)
    
    def _open_output_folder(self):
        if not open_folder(self.output_dir_edit.text()):
            QMessageBox.warning(self, '提示', '目录不存在，请先执行生成操作')
    
    def set_source(self, path: str):
        """设置数据源（从提取页面调用）"""
        self.excel_edit.setText(path)
        self.log_view.append(f'✅ 已自动加载整合文件: {path}')
    
    def _export_log(self):
        log_content = self.log_view.toPlainText()
        if not log_content:
            QMessageBox.information(self, '提示', '日志为空')
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, '导出日志', f'word_gen_log_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt',
            'Text Files (*.txt)'
        )
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(log_content)
            QMessageBox.information(self, '成功', f'日志已导出')
    
    def _log(self, msg: str):
        timestamp = datetime.now().strftime('%H:%M:%S')
        self.log_view.append(f'[{timestamp}] {msg}')
    
    def _start_generation(self):
        if not all([self.excel_edit.text(), self.template_edit.text(), self.output_dir_edit.text()]):
            QMessageBox.warning(self, '错误', '请完善所有文件路径')
            return
        
        self.btn_run.setEnabled(False)
        self.log_view.clear()
        self.progress.setValue(0)
        
        self.thread = WordGenerationThread(
            self.excel_edit.text(),
            self.template_edit.text(),
            self.output_dir_edit.text(),
            self.placeholder_combo.currentText(),
            self.filename_edit.text(),
            True,  # use_formatting 默认 True
            3      # qr_size 默认 3cm
        )
        self.thread.log.connect(self._log)
        self.thread.progress.connect(self.progress.setValue)
        self.thread.finished.connect(self._on_finished)
        self.thread.start()
    
    def _on_finished(self, success: int, failed: int):
        self.btn_run.setEnabled(True)

        # 如果勾选了合并选项且有成功生成的文件，则执行合并
        merged_file = None
        if self.merge_checkbox.isChecked() and success > 0:
            self._log('📦 正在合并所有Word文档...')
            merged_file = self._merge_generated_files()

        # 显示完成消息
        if merged_file:
            QMessageBox.information(
                self, '完成',
                f'生成完成！\n✅ 成功: {success} 个\n❌ 失败: {failed} 个\n\n📦 已合并为: {os.path.basename(merged_file)}'
            )
        else:
            QMessageBox.information(
                self, '完成',
                f'生成完成！\n✅ 成功: {success} 个\n❌ 失败: {failed} 个'
            )

    def _merge_generated_files(self) -> Optional[str]:
        """合并生成的Word文件"""
        output_dir = self.output_dir_edit.text()
        if not os.path.exists(output_dir):
            self._log('❌ 输出目录不存在，无法合并')
            return None

        # 获取所有生成的docx文件
        word_files = []
        for f in sorted(os.listdir(output_dir)):
            if f.endswith('.docx') and not f.startswith('~$') and not f.startswith('合并_'):
                word_files.append(os.path.join(output_dir, f))

        if len(word_files) < 2:
            self._log('⚠️ 文件数量不足，无需合并')
            return None

        # 生成合并文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        merged_filename = f'合并_{timestamp}.docx'
        merged_path = os.path.join(output_dir, merged_filename)

        # 执行合并
        success = merge_word_documents(word_files, merged_path, self._log)

        if success:
            return merged_path
        return None


# ==================== 设置页面 ====================
class SettingsWidget(QWidget):
    """设置页面"""
    theme_changed = pyqtSignal(str)
    
    def __init__(self, config_manager: ConfigManager, parent=None):
        super().__init__(parent)
        self.config_manager = config_manager
        self._setup_ui()
    
    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # 外观设置
        appearance_group = QGroupBox('🎨 外观设置')
        appearance_layout = QGridLayout()
        
        appearance_layout.addWidget(QLabel('主题:'), 0, 0)
        self.theme_combo = QComboBox()
        self.theme_combo.addItems(['浅色', '深色'])
        current_theme = self.config_manager.get('theme', 'light')
        self.theme_combo.setCurrentIndex(0 if current_theme == 'light' else 1)
        self.theme_combo.currentIndexChanged.connect(self._on_theme_changed)
        appearance_layout.addWidget(self.theme_combo, 0, 1)
        
        appearance_group.setLayout(appearance_layout)
        layout.addWidget(appearance_group)
        
        # 默认值设置
        defaults_group = QGroupBox('⚙️ 默认设置')
        defaults_layout = QGridLayout()
        
        defaults_layout.addWidget(QLabel('数值格式化:'), 0, 0)
        self.default_format_cb = QCheckBox('启用千分位格式')
        self.default_format_cb.setChecked(self.config_manager.get('use_number_formatting', True))
        defaults_layout.addWidget(self.default_format_cb, 0, 1)
        
        defaults_layout.addWidget(QLabel('递归扫描:'), 1, 0)
        self.default_recursive_cb = QCheckBox('默认启用')
        self.default_recursive_cb.setChecked(self.config_manager.get('recursive_scan', False))
        defaults_layout.addWidget(self.default_recursive_cb, 1, 1)
        
        if HAS_QRCODE:
            defaults_layout.addWidget(QLabel('二维码尺寸(cm):'), 2, 0)
            self.default_qr_spin = QSpinBox()
            self.default_qr_spin.setRange(1, 10)
            self.default_qr_spin.setValue(int(self.config_manager.get('qr_size_cm', 3)))
            defaults_layout.addWidget(self.default_qr_spin, 2, 1)
        
        defaults_group.setLayout(defaults_layout)
        layout.addWidget(defaults_group)
        
        # 保存按钮
        btn_save = QPushButton('💾 保存设置')
        btn_save.clicked.connect(self._save_settings)
        layout.addWidget(btn_save)
        
        # 关于
        about_group = QGroupBox('ℹ️ 关于')
        about_layout = QVBoxLayout()
        about_text = QLabel(
            f'{APP_NAME}\n\n'
            '功能：\n'
            '• Excel 批量数据提取 (支持范围、递归)\n'
            '• Word 批量文档生成 (支持二维码)\n'
            '• 自定义文件名模板\n'
            '• 深色/浅色主题\n'
            '• 拖拽支持、数据预览、日志导出\n\n'
            f'二维码支持: {"✅ 已安装" if HAS_QRCODE else "❌ 未安装 (pip install qrcode[pil])"}\n\n'
            f'作者: {APP_AUTHOR}'
        )
        about_layout.addWidget(about_text)
        about_group.setLayout(about_layout)
        layout.addWidget(about_group)
        
        layout.addStretch()
    
    def _on_theme_changed(self, index: int):
        theme = 'light' if index == 0 else 'dark'
        self.config_manager.set('theme', theme)
        self.theme_changed.emit(theme)
    
    def _save_settings(self):
        self.config_manager.set('use_number_formatting', self.default_format_cb.isChecked())
        self.config_manager.set('recursive_scan', self.default_recursive_cb.isChecked())
        if HAS_QRCODE:
            self.config_manager.set('qr_size_cm', self.default_qr_spin.value())
        QMessageBox.information(self, '成功', '设置已保存')


# ==================== 主窗口 ====================
class MainWindow(QMainWindow):
    """主窗口"""
    extraction_finished = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        self.config_manager = ConfigManager()
        self.setWindowTitle(APP_NAME)
        self.setMinimumSize(900, 750)
        self.resize(1000, 850)
        
        # 主容器
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        # 顶部标题栏
        header = self._create_header()
        main_layout.addWidget(header)
        
        # 创建标签页
        self.tabs = QTabWidget()
        
        self.extraction_tab = ExtractionWidget(self.config_manager, self)
        self.word_tab = WordGenWidget(self.config_manager, self)
        self.settings_tab = SettingsWidget(self.config_manager, self)
        
        self.tabs.addTab(self.extraction_tab, '📊 Excel提取整合')
        self.tabs.addTab(self.word_tab, '📝 Word批量生成')
        self.tabs.addTab(self.settings_tab, '⚙️ 设置')
        
        main_layout.addWidget(self.tabs)
        
        # 底部版权信息
        footer = QLabel('大鑽戒基金會版權所有 © 2025')
        footer.setAlignment(Qt.AlignCenter)
        footer.setStyleSheet(f'''
            color: {Colors.TEXT_SECONDARY}; 
            font-size: 11px; 
            padding: 8px 0;
            background-color: {Colors.BG_LIGHT};
        ''')
        main_layout.addWidget(footer)
        
        # 连接信号
        self.extraction_finished.connect(self.word_tab.set_source)
        self.settings_tab.theme_changed.connect(self._apply_theme)
        
        # 打印Logo
        print(APP_LOGO)
    
    def _create_header(self) -> QWidget:
        """创建顶部标题栏"""
        header = QWidget()
        header.setFixedHeight(80)
        header.setStyleSheet(f'''
            QWidget {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 {Colors.PRIMARY_DARK},
                    stop:0.5 {Colors.PRIMARY},
                    stop:1 {Colors.PRIMARY_LIGHT});
            }}
        ''')
        
        layout = QHBoxLayout(header)
        layout.setContentsMargins(25, 10, 25, 10)
        
        # 左侧Logo和标题
        left_layout = QHBoxLayout()
        left_layout.setSpacing(15)
        
        logo_label = QLabel('📊')
        logo_label.setStyleSheet('font-size: 36px; background: transparent;')
        left_layout.addWidget(logo_label)
        
        title_layout = QVBoxLayout()
        title_layout.setSpacing(2)
        
        title = QLabel('Excel/Word 自动化工具')
        title.setStyleSheet('font-size: 20px; font-weight: bold; color: white; background: transparent;')
        title_layout.addWidget(title)
        
        subtitle = QLabel(f'Pro {APP_VERSION}  |  by {APP_AUTHOR}')
        subtitle.setStyleSheet('font-size: 12px; color: rgba(255, 255, 255, 0.85); background: transparent;')
        title_layout.addWidget(subtitle)
        
        left_layout.addLayout(title_layout)
        layout.addLayout(left_layout)
        
        layout.addStretch()
        
        # 右侧功能区
        right_layout = QHBoxLayout()
        right_layout.setSpacing(12)
        
        features = QLabel('🚀 批量处理  |  📄 模板替换  |  🎨 智能字体  |  📱 二维码')
        features.setStyleSheet('''
            font-size: 11px; color: rgba(255, 255, 255, 0.9); background: transparent;
            padding: 5px 10px; border: 1px solid rgba(255, 255, 255, 0.3); border-radius: 15px;
        ''')
        right_layout.addWidget(features)
        
        donate_btn = QPushButton('💝 支持作者')
        donate_btn.setStyleSheet('''
            QPushButton {
                background-color: rgba(255, 255, 255, 0.2); color: white;
                border: 1px solid rgba(255, 255, 255, 0.4); border-radius: 18px;
                padding: 8px 18px; font-size: 12px; font-weight: 500;
            }
            QPushButton:hover {
                background-color: rgba(255, 255, 255, 0.35); border-color: white;
            }
        ''')
        donate_btn.setCursor(Qt.PointingHandCursor)
        donate_btn.clicked.connect(self._show_donate_dialog)
        right_layout.addWidget(donate_btn)
        
        layout.addLayout(right_layout)
        
        return header
    
    def _show_donate_dialog(self):
        dialog = DonateDialog(self)
        dialog.exec_()
    
    def _apply_theme(self, theme: str):
        apply_theme(QApplication.instance(), theme)
    
    def on_extraction_complete(self, output_path: str):
        """提取完成后的处理"""
        reply = QMessageBox.question(
            self, '流程衔接',
            'Excel整合完成！\n\n是否立即使用该文件生成Word文档？',
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply == QMessageBox.Yes:
            self.tabs.setCurrentIndex(1)
            self.extraction_finished.emit(output_path)


# ==================== 主程序入口 ====================
def main():
    app = QApplication(sys.argv)
    
    # 加载配置并应用主题
    config = ConfigManager()
    theme = config.get('theme', 'light')
    apply_theme(app, theme)
    
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()