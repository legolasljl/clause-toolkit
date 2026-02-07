# -*- coding: utf-8 -*-
"""
Clause Comparison Assistant V18.9 Windows Edition
智能条款工具箱
- [性能] 预处理索引加速匹配 5-10x
- [算法] 编辑距离容错 + 混合相似度
- [重构] 多级匹配策略拆分
- [功能] 批量处理多文件
- [健壮] 完善异常处理和日志
- [配置] 外部化JSON配置
- [新增] 用户自定义映射管理（单条/批量）
- [新增] 导出时使用库内条款名
- [v17.0] 中文分词支持 (jieba)
- [v17.0] 扩展英中映射表 (200+条款)
- [v17.0] 中英混合条款智能分离匹配
- [v17.0] TF-IDF向量快速候选筛选
- [v17.0] 动态权重调整
- [v17.0] 扩展语义别名和关键词库
- [v17.1] 多结果匹配（每条客户条款返回最多3条匹配供选择）
- [v17.1] 除外条款智能过滤（除非客户明确包含"除外"）
- [v17.1] 条款查询功能（仅查询标题，支持模糊搜索）
- [v17.1] 用户映射优先（有映射时只返回映射的那一条）
- [V18.0] Tab页面布局（条款提取/条款比对/条款输出）
- [V18.0] 条款提取功能（支持docx/pdf，文件夹智能分类）
- [V18.0] 条款输出功能（Excel比对报告转Word文档）
- [V18.0] 文件夹分类ZIP打包导出
- [V18.0] UI优化：紧凑型统计面板
- [V18.0] Tab标签显示优化（加宽+字体调整）
- [V18.0] 分类预览框样式修复（字体渲染）
- [V18.0] .doc自动转换为.docx功能（macOS textutil/LibreOffice）
- [V18.0] 统计栏水平对齐优化（分隔符布局）
- [V18.0] 文件列表字体颜色修复（高对比度）
- [V18.0] Excel导出Anthropic风格美化
- [V18.0] 条款输出Tab完整功能实现
- [V18.0] 支持从条款提取结果或Excel文件加载数据
- [V18.0] 三种输出模式：按条款逐个/按分类合并/全部合并
- [V18.0] Word样式自定义：标题字号/正文字号/行距/注册号显示
- [V18.0] 条款预览列表支持多选/全选
- [V18.0] 智能Excel列识别（自动匹配条款名称/注册号/内容列）
- [V18.0] Word文档Anthropic配色方案
- [V18.1] 特殊规则匹配：支持自定义条款匹配规则和提示信息
- [V18.5] 从报告导入映射显示详细统计（新增/更新/相同/跳过）
- [V18.5] 已映射条款名称优先识别为标题
- [V18.5] 修复排除列表对 Heading 样式的优先级问题
- [V18.5] 代码质量优化（预编译正则、常量定义、类型注解、辅助方法）
- [V18.8] 精准识别模式：勾选后仅提取蓝色字体的条款，适用于干扰项多的文档
- [V18.9] 加粗格式保留：条款库中的加粗文本在比对报告、Word输出、录单版全流程保留

Author: Dachi Yijin
Date: 2025-12-23
Updated: 2026-01-27 (V18.9 Bold Format Preservation)
"""

import sys
import os

# Windows UTF-8 编码适配（解决 cp936 无法编码 emoji/Unicode 的问题）
if sys.platform == 'win32':
    os.environ.setdefault('PYTHONIOENCODING', 'utf-8')
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import re
import difflib
import traceback
import logging
import subprocess
from typing import List, Dict, Tuple, Optional, Set, Any
from dataclasses import dataclass, field
from enum import Enum
from collections import defaultdict
from functools import lru_cache
from pathlib import Path
from datetime import datetime
import json
import platform
import pandas as pd
from docx import Document

# ASCII Art Logo
APP_LOGO = """
╔═══════════════════════════════════════════════════════════════════╗
║          ██████╗  ██████╗██╗   ██╗     ██╗██╗███╗   ██╗           ║
║          ██╔══██╗██╔════╝╚██╗ ██╔╝     ██║██║████╗  ██║           ║
║          ██║  ██║██║      ╚████╔╝      ██║██║██╔██╗ ██║           ║
║          ██║  ██║██║       ╚██╔╝  ██   ██║██║██║╚██╗██║           ║
║          ██████╔╝╚██████╗   ██║   ╚█████╔╝██║██║ ╚████║           ║
║          ╚═════╝  ╚═════╝   ╚═╝    ╚════╝ ╚═╝╚═╝  ╚═══╝           ║
║                    🚀 智能条款比对工具 🚀                         ║
║                     Author: Dachi_Yijin                           ║
╚═══════════════════════════════════════════════════════════════════╝
"""
# 打印Logo
print(APP_LOGO)

# ==========================================
# 中文分词支持
# ==========================================
try:
    import jieba
    jieba.setLogLevel(logging.WARNING)  # 减少jieba日志输出
    HAS_JIEBA = True
except ImportError:
    HAS_JIEBA = False

# ==========================================
# TF-IDF向量匹配支持
# ==========================================
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

# ==========================================
# PDF解析支持
# ==========================================
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    try:
        import PyPDF2
        HAS_PYPDF2 = True
    except ImportError:
        HAS_PYPDF2 = False

# ==========================================
# ZIP打包支持
# ==========================================
import zipfile
import shutil

# ==========================================
# 日志配置
# ==========================================
# PyInstaller 打包时日志放在 exe 同级目录，而非临时解压目录
if getattr(sys, 'frozen', False):
    LOG_DIR = Path(sys.executable).parent / "logs"
else:
    LOG_DIR = Path(__file__).parent / "logs"
LOG_DIR.mkdir(exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler(LOG_DIR / f"Clause_Comparison_Assistant_{datetime.now():%Y%m%d}.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ==========================================
# 导入配置管理器
# ==========================================
try:
    from clause_config_manager import get_config, ClauseConfigManager
    HAS_CONFIG_MANAGER = True
except ImportError:
    HAS_CONFIG_MANAGER = False
    logger.warning("未找到 clause_config_manager，使用内置配置")

# 导入映射管理器
try:
    from clause_mapping_manager import ClauseMappingManager, get_mapping_manager
    from clause_mapping_dialog import ClauseMappingDialog
    HAS_MAPPING_MANAGER = True
except ImportError:
    HAS_MAPPING_MANAGER = False
    logger.warning("未找到 clause_mapping_manager，映射管理功能不可用")

# ==========================================
# Windows PyQt5 Plugin Fix & Console Encoding
# ==========================================
if platform.system() == 'Windows':
    try:
        import ctypes
        ctypes.windll.kernel32.SetConsoleOutputCP(65001)
    except Exception:
        pass
try:
    import PyQt5
    plugin_path = os.path.join(os.path.dirname(PyQt5.__file__), 'Qt5', 'plugins')
    if not os.path.exists(plugin_path):
        plugin_path = os.path.join(os.path.dirname(PyQt5.__file__), 'Qt', 'plugins')
    os.environ['QT_QPA_PLATFORM_PLUGIN_PATH'] = plugin_path
except ImportError:
    pass

HAS_TRANSLATOR = False

try:
    from insurance_calculator import MainInsuranceTab, AddonInsuranceTab
    HAS_INSURANCE_CALC = True
except ImportError:
    HAS_INSURANCE_CALC = False

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QProgressBar, QTextEdit,
    QFileDialog, QMessageBox, QFrame, QGraphicsDropShadowEffect,
    QDialog, QFormLayout, QListWidget, QListWidgetItem, QCheckBox,
    QTabWidget, QSpinBox, QDoubleSpinBox, QGroupBox, QComboBox,
    QScrollArea
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QUrl, QTimer, QPropertyAnimation, QEasingCurve
from PyQt5.QtGui import QFont, QColor, QDesktopServices, QTextCursor

# ==========================================
# Windows 打包防闪退 (PyInstaller --noconsole)
# ==========================================
class NullWriter:
    def write(self, text): pass
    def flush(self): pass

if getattr(sys, 'frozen', False):
    if sys.stdout is None:
        sys.stdout = NullWriter()
    if sys.stderr is None:
        sys.stderr = NullWriter()

def global_exception_handler(exctype, value, tb):
    error_msg = "".join(traceback.format_exception(exctype, value, tb))
    logger.error(f"未捕获异常: {error_msg}")
    try:
        msg_box = QMessageBox()
        msg_box.setIcon(QMessageBox.Critical)
        msg_box.setText("程序发生意外错误")
        msg_box.setInformativeText(str(value))
        msg_box.setDetailedText(error_msg)
        msg_box.exec_()
    except Exception as e:
        logger.error(f"无法显示错误对话框: {e}")

sys.excepthook = global_exception_handler


# ==========================================
# 安全工具函数
# ==========================================
def validate_file_path(file_path: str, allowed_extensions: list = None) -> bool:
    """
    验证文件路径安全性，防止路径遍历攻击

    Args:
        file_path: 要验证的文件路径
        allowed_extensions: 允许的文件扩展名列表 (如 ['.docx', '.xlsx'])

    Returns:
        True 如果路径安全，False 否则
    """
    if not file_path:
        return False

    # 转换为绝对路径并规范化
    try:
        abs_path = os.path.abspath(os.path.normpath(file_path))
    except (TypeError, ValueError):
        return False

    # 检查路径遍历攻击 (.. 序列)
    if '..' in file_path:
        logger.warning(f"检测到路径遍历尝试: {file_path}")
        return False

    # 检查是否访问敏感系统目录
    sensitive_dirs = ['/etc', '/usr', '/bin', '/sbin', '/var', '/root', '/System', '/Library']
    for sensitive in sensitive_dirs:
        if abs_path.startswith(sensitive):
            logger.warning(f"检测到敏感目录访问: {file_path}")
            return False

    # 检查文件扩展名
    if allowed_extensions:
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in [e.lower() for e in allowed_extensions]:
            return False

    return True


def sanitize_error_message(error: Exception) -> str:
    """
    清理错误信息，移除敏感路径和系统信息

    Args:
        error: 异常对象

    Returns:
        清理后的错误信息
    """
    error_str = str(error)

    # 移除完整文件路径，只保留文件名
    import re
    # 匹配类Unix路径
    error_str = re.sub(r'/(?:Users|home)/[^/]+/[^\s\'"]+', '<路径已隐藏>', error_str)
    # 匹配Windows路径
    error_str = re.sub(r'[A-Z]:\\[^\s\'"]+', '<路径已隐藏>', error_str)

    return error_str


# ==========================================
# Anthropic UI 设计系统
# ==========================================
class AnthropicColors:
    """Anthropic 官方色彩系统"""
    # 背景色
    BG_PRIMARY = "#faf9f5"      # 主背景/奶油白
    BG_CARD = "#f0eee6"         # 卡片背景/浅米色
    BG_MINT = "#bcd1ca"         # 特殊卡片/薄荷绿
    BG_LAVENDER = "#cbcadb"     # 特殊卡片/淡紫色
    BG_DARK = "#141413"         # 深色区域

    # 强调色
    ACCENT = "#d97757"          # 主强调色/陶土色
    ACCENT_DARK = "#c6613f"     # 次强调色/深赭红
    ACCENT_HOVER = "#e8956f"    # 悬停色

    # 文字色
    TEXT_PRIMARY = "#141413"    # 主要文字
    TEXT_SECONDARY = "#6b6960"  # 次要文字
    TEXT_MUTED = "#57554e"      # 中等对比度文字（用于按钮/标签）
    TEXT_TERTIARY = "#57554e"   # 第三级文字（保险计算器模块使用）
    TEXT_LIGHT = "#ffffff"      # 深色背景上的文字

    # 状态色
    SUCCESS = "#5a9a7a"         # 成功/绿色
    WARNING = "#d9a557"         # 警告/金色
    ERROR = "#c75050"           # 错误/红色
    INFO = "#5a7a9a"            # 信息/蓝灰

    # 边框色
    BORDER = "#e5e3db"          # 浅边框
    BORDER_DARK = "#d0cec6"     # 深边框


def get_anthropic_scrollbar_style():
    """Anthropic 风格细滚动条样式"""
    return f"""
        QScrollArea {{
            border: none;
            background: transparent;
        }}
        QScrollBar:vertical {{
            background: transparent;
            width: 8px;
            margin: 0;
        }}
        QScrollBar::handle:vertical {{
            background: {AnthropicColors.BORDER_DARK};
            border-radius: 4px;
            min-height: 30px;
        }}
        QScrollBar::handle:vertical:hover {{
            background: {AnthropicColors.TEXT_SECONDARY};
        }}
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
            height: 0;
        }}
        QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
            background: transparent;
        }}
        QScrollBar:horizontal {{
            background: transparent;
            height: 8px;
            margin: 0;
        }}
        QScrollBar::handle:horizontal {{
            background: {AnthropicColors.BORDER_DARK};
            border-radius: 4px;
            min-width: 30px;
        }}
        QScrollBar::handle:horizontal:hover {{
            background: {AnthropicColors.TEXT_SECONDARY};
        }}
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
            width: 0;
        }}
        QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {{
            background: transparent;
        }}
    """


class AnthropicFonts:
    """Anthropic 字体配置"""
    # 标题字体
    TITLE_LARGE = ("Anthropic Sans", 28)
    TITLE = ("Anthropic Sans", 22)
    TITLE_SMALL = ("Anthropic Sans", 16)

    # 正文字体
    BODY = ("Anthropic Serif", 14)
    BODY_SMALL = ("Anthropic Serif", 12)

    # UI 元素
    BUTTON = ("Anthropic Sans", 14)
    LABEL = ("Anthropic Sans", 13)

    # 代码字体
    CODE = ("JetBrains Mono", 12)

    # 中文回退（Anthropic 字体不含中文）
    CN_FALLBACK = "Microsoft YaHei"


# ==========================================
# 常量定义
# ==========================================
class ExcelColumns:
    """Excel列名常量 - v17.1支持多结果匹配"""
    SEQ = '序号'
    CLIENT_ORIG = '客户条款(原)'
    CLIENT_TRANS = '客户条款(译)'
    CLIENT_CONTENT = '客户原始内容'
    LIMIT_INFO = '约定的限额'  # v18.15: 提取的限额/约定信息

    # 多结果匹配列 (v17.1)
    # 匹配1
    MATCH1_NAME = '匹配1_条款名称'
    MATCH1_REG = '匹配1_注册号'
    MATCH1_CONTENT = '匹配1_条款内容'
    MATCH1_SCORE = '匹配1_匹配度'
    MATCH1_LEVEL = '匹配1_匹配级别'
    # 匹配2
    MATCH2_NAME = '匹配2_条款名称'
    MATCH2_REG = '匹配2_注册号'
    MATCH2_CONTENT = '匹配2_条款内容'
    MATCH2_SCORE = '匹配2_匹配度'
    MATCH2_LEVEL = '匹配2_匹配级别'
    # 匹配3
    MATCH3_NAME = '匹配3_条款名称'
    MATCH3_REG = '匹配3_注册号'
    MATCH3_CONTENT = '匹配3_条款内容'
    MATCH3_SCORE = '匹配3_匹配度'
    MATCH3_LEVEL = '匹配3_匹配级别'

    # 保留旧列名以兼容（主匹配结果）
    MATCHED_NAME = '匹配条款库名称'
    REG_NO = '产品注册号'
    MATCHED_CONTENT = '匹配条款库内容'
    SCORE = '综合匹配度'
    MATCH_LEVEL = '匹配级别'
    DIFF_ANALYSIS = '保障差异提示'
    TITLE_SCORE = '标题相似度'
    CONTENT_SCORE = '内容相似度'

    # 列索引（1-based, 需根据新格式调整）
    SCORE_COL_IDX = 8  # 匹配1_匹配度
    LEVEL_COL_IDX = 9  # 匹配1_匹配级别


# ==========================================
# 数据结构
# ==========================================
class MatchLevel(Enum):
    """匹配级别"""
    EXACT = "精确匹配"
    SEMANTIC = "语义匹配"
    KEYWORD = "关键词匹配"
    FUZZY = "模糊匹配"
    NONE = "无匹配"

@dataclass
class MatchThresholds:
    """匹配阈值"""
    exact_min: float = 0.98
    semantic_min: float = 0.85
    keyword_min: float = 0.60
    fuzzy_min: float = 0.40
    accept_min: float = 0.15

@dataclass
class ClauseItem:
    """条款项"""
    title: str
    content: str
    original_title: str = ""


@dataclass
class MatchResult:
    """匹配结果"""
    matched_name: str = ""
    matched_content: str = ""
    matched_reg: str = ""
    score: float = 0.0
    title_score: float = 0.0
    content_score: float = 0.0
    match_level: MatchLevel = MatchLevel.NONE
    diff_analysis: str = ""

@dataclass
class LibraryIndex:
    """条款库索引结构"""
    by_name_norm: Dict[str, int] = field(default_factory=dict)
    by_keyword: Dict[str, List[int]] = field(default_factory=lambda: defaultdict(list))
    cleaned_cache: Dict[int, Dict[str, str]] = field(default_factory=dict)
    data: List[Dict] = field(default_factory=list)


# ==========================================
# 内置默认配置（当配置管理器不可用时）
# ==========================================
class DefaultConfig:
    """默认配置 - v17.0 扩展版"""

    # ========================================
    # 英中条款映射表 (基于TOC.docx扩展，200+条目)
    # ========================================
    CLIENT_EN_CN_MAP = {
        # === 通用条款 ===
        "30 days notice of cancellation clause": "30天注销保单通知条款",
        "30 days notice of cancellation": "30天注销保单通知条款",
        "60 days non-renewal notice clause": "60天不续保通知条款",
        "60 days non-renewal notice": "60天不续保通知条款",
        "72 hours clause": "72小时条款",
        "72 hours": "72小时条款",
        "time adjustment": "72小时条款",
        "50/50 clause": "50/50条款",
        "85% co-insurance": "85％扩展条款",

        # === 索赔与控制 ===
        "claims control clause": "理赔控制条款",
        "claims control": "理赔控制条款",
        "joint-insured clause": "共同被保险人条款",
        "joint insured clause": "共同被保险人条款",
        "joint insured": "共同被保险人条款",
        "jurisdiction clause": "司法管辖权条款",
        "jurisdiction": "司法管辖权条款",
        "loss adjusters clause": "指定公估人条款",
        "loss adjusters": "指定公估人条款",
        "nomination of loss adjusters clause": "指定公估人条款",
        "loss adjuster clause": "指定公估人条款",
        "loss notification clause": "损失通知条款",
        "loss notification": "损失通知条款",

        # === 控制与取消 ===
        "no control clause": "不受控制条款",
        "no control": "不受控制条款",
        "non-cancellation clause": "不可注销保单条款",
        "non cancellation clause": "不可注销保单条款",
        "non-invalidation clause": "不使失效条款",
        "non invalidation clause": "不使失效条款",
        "non invalidation": "不使失效条款",

        # === 付款与费用 ===
        "payment on account clause": "预付赔款条款",
        "payment on account": "预付赔款条款",
        "premium installment clause": "分期付费条款",
        "premium instalment clause": "分期付费条款",
        "premium adjustment clause": "保费调整条款",
        "premium adjustment": "保费调整条款",
        "professional fees clause": "专业费用扩展条款",
        "professional fees": "专业费用扩展条款",
        "professional fee extension clause": "专业费用扩展条款",

        # === 代位与利益 ===
        "severability of interest clause": "利益可分性条款",
        "severability of interest": "利益可分性条款",
        "waiver of subrogation clause": "放弃代位追偿扩展条款",
        "waiver of subrogation extension clause": "放弃代位追偿扩展条款",
        "waiver of subrogation": "放弃代位追偿扩展条款",

        # === 自然灾害 ===
        "earthquake extension clause": "地震扩展条款",
        "earthquake extension": "地震扩展条款",
        "earthquake and tsunami": "地震扩展条款",
        "earthquake": "地震扩展条款",
        "flood extension clause": "洪水扩展条款",
        "flood extension": "洪水扩展条款",
        "flood and inundation exclusion clause": "洪水除外条款",
        "flood exclusion": "洪水除外条款",
        "flood prevention warranty clause": "防洪保证条款",
        "typhoon and hurricane extension clause": "台风、飓风扩展条款",
        "typhoon and hurricane": "台风、飓风扩展条款",
        "typhoon extension": "台风、飓风扩展条款",
        "hurricane extension": "台风、飓风扩展条款",
        "tornado extension clause": "龙卷风扩展条款",
        "tornado exclusion clause": "龙卷风除外条款",
        "hailstone extension clause": "冰雹扩展条款",
        "hailstone extension": "冰雹扩展条款",
        "snowstorm and icicle extension clause": "暴雪、冰凌扩展条款",
        "snowstorm extension": "暴雪、冰凌扩展条款",
        "storm and tempest extension clause": "暴风雨扩展条款",
        "storm and tempest exclusion clause": "暴风雨除外条款",
        "sandstorm extension clause": "沙尘暴扩展条款",
        "sandstorm exclusion clause": "沙尘暴除外条款",
        "lightning extension clause": "雷电扩展条款",
        "lightning extension": "雷电扩展条款",

        # === 地质灾害 ===
        "accidental subsidence of ground extension clause": "地面突然下陷下沉扩展条款",
        "subsidence extension": "地面突然下陷下沉扩展条款",
        "landslip & subsidence clause": "地崩及地陷条款",
        "landslip and subsidence": "地崩及地陷条款",
        "mud-rock flow, avalanche and sudden landslip extension clause": "泥石流、崩塌、突发性滑坡扩展条款",
        "mudslide extension": "泥石流、崩塌、突发性滑坡扩展条款",

        # === 盗窃与安全 ===
        "theft, burglary and robbery extension clause": "盗窃、抢劫扩展条款",
        "theft and robbery": "盗窃、抢劫扩展条款",
        "theft extension": "盗窃、抢劫扩展条款",
        "burglary insurance clause": "盗窃险条款",
        "burglary extension": "盗窃险条款",
        "malicious damage extension clause": "恶意破坏扩展条款",
        "malicious damage": "恶意破坏扩展条款",

        # === 罢工与暴乱 ===
        "strike riot and civil commotion extension clause": "罢工、暴乱及民众骚乱扩展条款",
        "strike, riot and civil commotion": "罢工、暴乱及民众骚乱扩展条款",
        "strike riot civil commotion": "罢工、暴乱及民众骚乱扩展条款",
        "srcc": "罢工、暴乱及民众骚乱扩展条款",
        "terrorism extension clause": "恐怖活动扩展条款",
        "terrorism extension": "恐怖活动扩展条款",
        "act of terrorism extension clause": "恐怖活动扩展条款",

        # === 价值与金额 ===
        "reinstatement value clause": "重置价值条款",
        "reinstatement value": "重置价值条款",
        "agreed value insurance clause": "定值保险条款",
        "agreed value": "定值保险条款",
        "automatic reinstatement of sum insured clause": "自动恢复保险金额条款",
        "automatic reinstatement of sum insured": "自动恢复保险金额条款",
        "automatic reinstatement": "自动恢复保险金额条款",
        "escalation extension clause": "自动升值扩展条款",
        "escalation extension": "自动升值扩展条款",

        # === 费用扩展 ===
        "removal of debris clause": "清理残骸费用扩展条款",
        "removal of debris": "清理残骸费用扩展条款",
        "debris removal expenses extension clause": "清理残骸费用扩展条款",
        "debris removal": "清理残骸费用扩展条款",
        "fire fighting cost extension clause": "灭火费用扩展条款",
        "fire fighting cost": "灭火费用扩展条款",
        "fire brigade charges extension clause": "消防队灭火费用扩展条款",
        "air freight fee extension clause": "空运费扩展条款",
        "air freight extension": "空运费扩展条款",
        "airfreight clause": "空运费扩展条款",
        "extra charges extension clause": "特别费用扩展条款",
        "extra charges clause": "特别费用扩展条款",
        "extra charges": "特别费用扩展条款",

        # === 财产与建筑 ===
        "all other contents extension clause": "其他物品扩展条款",
        "alteration of building clause": "建筑物变动扩展条款",
        "building alterations clause": "建筑物改变条款",
        "capital additions extension clause": "增加资产扩展条款",
        "capital additions": "增加资产扩展条款",
        "contract price extension clause": "合同价格扩展条款",
        "designation of property clause": "财物种别条款",
        "foundation exclusion clause": "地基除外条款",
        "simple building exclusion": "简易建筑除外条款",
        "property in the open or simple building extension clause": "露天存放及简易建筑内财产扩展条款",
        "property in the open": "露天存放及简易建筑内财产扩展条款",
        "off premises property clause": "场所外财产条款",
        "outside ancillary devices of building extension clause": "建筑物外部附属设施扩展条款",

        # === 设备与机械 ===
        "boilers and pressure vessels extension clause": "锅炉、压力容器扩展条款",
        "boiler explosion": "锅炉爆炸责任条款",
        "breakage of glass extension clause": "玻璃破碎扩展条款",
        "breakage of glass clause": "玻璃破碎扩展条款",
        "glass breakage": "玻璃破碎扩展条款",
        "bursting of water tank or water pipe extension clause": "水箱、水管爆裂扩展条款",
        "water damage": "水箱、水管爆裂扩展条款",
        "hoisting and transport machinery extension clause": "起重、运输机械扩展条款",
        "locomotive extension clause": "铁路机车车辆扩展条款",
        "refrigerating plants extension clause": "冷库扩展条款",
        "sprinkler leakage damage extension clause": "自动喷淋系统水损扩展条款",
        "sprinkler leakage": "自动喷淋系统水损扩展条款",
        "portable devices on premises extension clause": "便携式设备扩展条款",

        # === 运输与移动 ===
        "inland transit extension clause": "内陆运输扩展条款",
        "inland transit clause": "内陆运输扩展条款",
        "inland transit": "内陆运输扩展条款",
        "transit clause": "运输条款",
        "temporary removal extension clause": "临时移动扩展条款",
        "temporary removal": "临时移动扩展条款",
        "temporary removal between factories extension clause": "厂区间临时移动扩展条款",
        "loaded property extension clause": "车辆装载物扩展条款",

        # === 责任与赔偿 ===
        "public authority extension clause": "公共当局扩展条款",
        "public authority": "公共当局扩展条款",
        "civil authorities clause": "公共当局扩展条款",
        "civil authorities": "公共当局扩展条款",
        "error and omissions clause": "错误和遗漏条款",
        "errors and omissions clause": "错误和遗漏条款",
        "errors and omissions": "错误和遗漏条款",
        "breach of conditions clause": "违反条件条款",
        "breach of conditions": "违反条件条款",
        "cross liability clause": "交叉责任条款",
        "cross liability": "交叉责任条款",
        "contractual liability clause": "契约责任扩展条款",
        "contractual liability": "契约责任扩展条款",

        # === 其他扩展 ===
        "automatic cover clause": "自动承保条款",
        "automatic cover": "自动承保条款",
        "average relief clause": "分摊豁免条款",
        "brand & trademark clause": "商标条款",
        "brand and trademark": "商标条款",
        "cost of duplication extension clause": "复制费用扩展条款",
        "documents clause": "索赔单据条款",
        "emergency rescue clause": "紧急抢险条款",
        "falling of flying objects extension clause": "飞行物体及其他空中运行物体坠落扩展条款",
        "fire prevention facilities warranty clause": "消防保证条款",
        "impact damage extension clause": "碰撞扩展条款",
        "impact damage exclusion clause": "碰撞除外条款",
        "inhibition clause": "阻止条款",
        "it clarification clause": "数据损失澄清条款",
        "legal requirements warranty": "遵守法律规定保证条款",
        "loss payee clause": "赔款接受人条款",
        "mortgage clause": "抵押权条款",
        "mortgagee clause": "抵押条款",
        "non occupying landlord clause": "非占用者业主条款",
        "oil or gas pipeline damage extension clause": "油气管道损坏扩展条款",
        "out-sourcing processing extension clause": "委托加工扩展条款",
        "pair & set clause": "成对或成套设备条款",
        "personal effects of employees extension clause": "雇员个人物品扩展条款",
        "smoke damage extension clause": "烟熏扩展条款",
        "spontaneous combustion extension clause": "自燃扩展条款",
        "spontaneous combustion exclusion clause": "自燃除外条款",
        "stock declaration and adjustment clause": "仓储财产申报条款",
        "storage warranty": "存放保证条款",
        "supply failure extension clause": "供应中断扩展条款",
        "supply suspension extension clause": "供应中断扩展条款",
        "temporary protection extension clause": "临时保护措施扩展条款",
        "undamaged building extra charges extension clause": "建筑物未受损部分额外费用扩展条款",
        "workmen clause": "装修工人条款",
        "assignment clause": "权益转让条款",

        # === 工程险条款 ===
        "camps and stores clause": "工棚、库房特别条款",
        "cement storage warranty": "水泥存储保证条款",
        "construction material clause": "建筑材料特别条款",
        "construction plant, equipment and machinery clause": "施工用机具特别条款",
        "construction machinery clause": "建筑、安装施工机具、设备扩展条款",
        "erection machinery clause": "建筑、安装施工机具、设备扩展条款",
        "contract works taken over or put into service clause": "工程完工部分扩展条款",
        "cost for decontamination clause": "清除污染费用扩展条款",
        "customs duties clause": "海关关税条款",
        "dams and water reservoirs clause": "大坝、水库工程除外特别条款",
        "defective design, materials and workmanship": "设计错误、原材料缺陷及工艺不善条款",
        "defects liability period clause": "扩展责任保证期扩展条款",
        "extended maintenance clause": "扩展责任保证期扩展条款",
        "designer's risk clause": "设计师风险扩展条款",
        "drilling work for water wells clause": "钻井工程特别条款",
        "employer's property extension": "雇主财产财产扩展条款",
        "escalation clause": "10％增值条款",
        "existing structures and surrounding property clause": "原有建筑物及周围财产扩展条款",
        "extinguishing expenses clause": "灭火费用条款",
        "fire-fighting facilities clause": "防火设施特别条款",
        "free issue materials clause": "免费提供物料扩展条款",
        "guarantee period clause": "保证期特别扩展条款",
        "hoisting, cranes and unregistered vehicles liability clause": "起重机、未登记车辆责任扩展条款",
        "hydrocarbon processing industries clause": "碳氢化合物生产业特别条款",
        "indemnity to principals clause": "业主保障条款",
        "laying pipelines, ducts and cables clause": "铺设管道、电缆特别条款",
        "laying water supply and sewer pipes clause": "铺设供水、污水管特别条款",
        "leak search costs when laying pipelines clause": "埋管查漏费用特别条款",
        "maintenance & inspection clause": "检查维护条款",
        "maintenance visits clause": "有限责任保证期扩展条款",
        "marine cargo insurance clause": "运输险、工程险责任分摊条款",
        "marine work special condition": "海工特别条款",
        "non-negligent indemnity": "非疏忽过错赔偿条款",
        "nuclear fuel elements clause": "核燃料组件条款",
        "obstruction & nuisance clause": "阻碍或妨害条款",
        "offsite storage clause": "工地外储存物特别条款",
        "plans and documents clause": "工程图纸、文件特别条款",
        "principal's property clause": "业主财产扩展条款",
        "quarterly declaration clause": "季度申报条款",
        "reactor pressure vessel with internals clause": "压力反应堆特别扩展条款",
        "removal of debris from landslides clause": "清除滑坡土石方特别除外条款",
        "run off clause": "保单延续条款",
        "safety precautions clause": "安全防范条款",
        "structures in earthquake zones clause": "地震地区建筑物特别条款",
        "testing & commissioning clause": "试车条款",
        "time adjustment clause": "时间调整特别条款",
        "time schedule clause": "建筑、安装工程时间进度特别条款",
        "tunnels and galleries clause": "隧道工程特别除外条款",
        "underground cables, pipes and other facilities clause": "地下电缆、管道及设施特别条款",
        "underground service clause": "地下服务设施条款",
        "underground works clause": "地下工程条款",
        "unexploded bombs clause": "地下炸弹特别条款",
        "used machinery clause": "旧设备除外条款",
        "vibration, removal or weakening of support clause": "震动、移动或减弱支撑扩展条款",
        "sue & labor clause": "诉讼及劳务费用特别条款",
        "manufacturer's risks clause": "制造商风险扩展条款",
        "piling, foundation and retaining wall construction work": "打桩及挡土墙除外条款",
        "burning & welding clause": "烧焊条款",

        # === 责任险条款 ===
        "accidental pollution clause": "意外污染条款",
        "advertising signs and decorations liability clause": "广告及装饰装置责任条款",
        "car park liability clause": "停车场责任条款",
        "car park service clause": "泊车服务条款",
        "contractors contingent liability clause": "承包人意外责任条款",
        "defective sanitary installation clause": "有缺陷的卫生装置条款",
        "delivery goods extension": "运输货物扩展条款",
        "elevator and escalator clause": "电梯责任条款",
        "lifts, elevators and escalators liability clause": "电梯、升降机责任扩展条款",
        "employees temporary working overseas": "海外公干条款",
        "employees temporarily working oversea clause": "员工公（劳）务出国条款",
        "exhibition and sales demonstration": "展览和销售演示条款",
        "fire & explosion extension clause": "火灾和爆炸责任条款",
        "fire brigade and water damage clause": "灭火及水损责任条款",
        "first aid liability clause": "急救责任条款",
        "first aid treatment clause": "急救费用条款",
        "food and drink clause": "食品、饮料责任条款",
        "goods and services clause": "提供物品及服务条款",
        "guest's property clause": "客人财产责任条款",
        "hire and non-owned automobiles liability clause": "租用及非拥有机动车辆责任条款",
        "hoists, cranes and unregistered vehicles liability clause": "起重机及起重设备责任条款",
        "indemnity to landlord clause": "房东保障条款",
        "independent contractors liability clause": "独立承建商责任条款",
        "laundry liability clause": "洗衣房责任条款",
        "loading and unloading of vehicles clause": "车辆装卸责任条款",
        "maintenance, repair and decoration of the premises clause": "修改、修理及保养责任条款",
        "motor contingent liability clause": "租用汽车责任条款",
        "personal injury liability clause": "人身侵害责任条款",
        "social and welfare club clause": "联谊及康乐活动责任附加条款",
        "swimming pool liability clause": "游泳池责任条款",
        "tenant's liability clause": "出租人责任条款",
        "third party liability of directors and executives clause": "董事及高级管理人员个人第三者责任条款",
        "catering facilities clause": "膳食条款",
        "extraordinary weather condition clause": "反常天气条款",
        "social activities clause": "社会活动条款",

        # === 产品责任险条款 ===
        "absolute asbestos exclusion": "石棉除外条款",
        "absolute pollution exclusion": "污染除外条款",
        "allergy exclusion clause": "过敏除外条款",
        "batch clause": "同一批次产品条款",
        "circuit board & battery exclusion clause": "电路板、电池除外条款",
        "claim made basis clause": "以索赔提出为基础条款",
        "defense cost within the limit of indemnity": "抗辩费用条款",
        "designated vendor liability": "指定经销商责任条款",
        "efficacy exclusion clause": "功效除外条款",
        "electromagnetic radiation exclusion": "电磁辐射、无线电波除外条款",
        "employees bodily injury exclusion": "雇员人身伤害除外条款",
        "genetically modified organisms exclusion": "转基因体除外条款",
        "gmo exclusion": "转基因体除外条款",
        "lead exclusion": "铅物质除外条款",
        "nuclear energy liability exclusion": "核能责任除外条款",
        "occurrence basis clause": "以发生为基础条款",
        "product-completed operation": "完工操作风险条款",
        "punitive damage exclusion": "惩罚性赔偿除外条款",
        "exemplary damage exclusion": "惩罚性赔偿除外条款",
        "silica exclusion": "硅除外条款",
        "us canada domiciled operations exclusion clause": "美加地区操作除外条款",
        "vendor broad form liability": "列明经销商扩展条款",
        "war and terrorism exclusion": "战争及恐怖主义除外条款",

        # === 营业中断险条款 ===
        "accumulated stocks clause": "货物累积条款",
        "bomb scare extension": "炸弹恐吓条款",
        "denial of access": "通道堵塞条款",
        "departmental clause": "部门条款",
        "inclusion of all turnover": "包括全部营业额条款",
        "infectious disease murder and closure clause": "谋杀等条款",
        "public utilities extension": "公共事业设备扩展条款",
        "reinstatement of sum insured clause": "恢复保险金额条款",
        "uninsured standing charges clause": "未保险的维持费用条款",
        "waiver of excess clause": "免赔额豁免条款",
        "loss of book debts clause": "遗失欠款帐册条款",

        # === 机损险条款 ===
        "overhaul of electric motors": "电动马达检修条款",
        "overhaul of steam, water and gas turbines": "蒸气、水、气体涡轮机及涡轮发电机条款",

        # === 通用简写 ===
        "interpretation & headings": "通译和标题条款",
        "year 2000 problem exclusion clause": "财产险2000年问题除外责任条款",

        # === 基本风险术语 (来自史带财产保险条款) ===
        "fire": "火灾",
        "lightning": "闪电",
        "explosion": "爆炸",
        "aircraft": "飞机",
        "aircraft damage": "飞机损坏",

        # === 特别列明风险 (SPECIAL PERILS) ===
        "volcanic eruption": "火山爆发",
        "subterranean fire": "地下火",
        "cyclone": "飓风",
        "windstorm": "风暴",
        "bursting or overflowing": "爆裂或溢出",
        "leakage from fire extinguishing appliance": "灭火设备渗漏",
        "impact by vehicle": "车辆撞击",
        "impact by animal": "动物撞击",
        "smoke": "烟熏",
        "frost": "霜冻",
        "weight of snow or ice": "积雪积冰",
        "avalanche": "雪崩",
        "falling trees": "树木倒塌",
        "falling trees or parts thereof": "树木或其部分倒塌",
        "riot and strike": "暴动及罢工",

        # === 理赔基础术语 (BASIS OF INDEMNIFICATION) ===
        "reinstatement": "重置",
        "actual cash value": "实际现金价值",
        "feedstock": "原料",
        "catalysts": "催化剂",
        "sum insured": "保险金额",
        "indemnified": "赔偿",
        "indemnification": "赔偿",
        "basis of indemnification": "理赔基础",

        # === 除外财产术语 (PROPERTY EXCLUDED) ===
        "cash": "现金",
        "bank notes": "钞票",
        "securities": "有价证券",
        "deeds": "契约",
        "bonds": "债券",
        "bills of exchange": "汇票",
        "promissory notes": "本票",
        "cheques": "支票",
        "jewellery": "珠宝",
        "precious stones": "宝石",
        "curiosities": "古玩",
        "rare books": "珍本书籍",
        "works of art": "艺术品",
        "livestock": "牲畜",
        "growing crops": "生长中的农作物",
        "standing timber": "活立木",
        "motor vehicles": "机动车辆",
        "watercraft": "船舶",
        "railway locomotives": "铁路机车",
        "rolling stock": "机车车辆",
        "mining property": "采矿财产",
        "dams": "堤坝",
        "dikes": "河堤",
        "reservoirs": "蓄水池",
        "property excluded": "除外财产",

        # === 通用除外条款 (GENERAL EXCEPTED CAUSES) ===
        "wilful act": "故意行为",
        "gross negligence": "重大过失",
        "wear and tear": "自然磨损",
        "gradual deterioration": "逐渐老化",
        "rust": "锈蚀",
        "oxidisation": "氧化",
        "oxidization": "氧化",
        "mould": "霉变",
        "mold": "霉变",
        "contamination": "污染",
        "inherent vice": "内在缺陷",
        "latent defect": "潜在缺陷",
        "mechanical breakdown": "机械故障",
        "electrical breakdown": "电气故障",
        "war": "战争",
        "invasion": "入侵",
        "civil war": "内战",
        "rebellion": "暴动",
        "revolution": "革命",
        "insurrection": "起义",
        "military power": "军事力量",
        "usurped power": "篡权",
        "confiscation": "没收",
        "nationalisation": "国有化",
        "nationalization": "国有化",
        "requisition": "征用",
        "destruction by order of government": "政府命令销毁",
        "nuclear reaction": "核反应",
        "nuclear radiation": "核辐射",
        "radioactive contamination": "放射性污染",
        "faulty design": "设计缺陷",
        "faulty workmanship": "工艺缺陷",
        "shortage of inventory": "盘点短缺",
        "unexplained disappearance": "不明原因失踪",
        "cessation of work": "停工",
        "general excepted causes": "通用除外条款",
        "excluded causes": "除外责任",

        # === 通用条款术语 (GENERAL CONDITIONS) ===
        "misrepresentation": "错误陈述",
        "non-disclosure": "隐瞒",
        "fraud": "欺诈",
        "breach of warranty": "保证条款违反",
        "notification of loss": "损失通知",
        "claim": "索赔",
        "particulars": "详情",
        "evidence": "证据",
        "abandonment": "委付",
        "salvage": "救助",
        "contribution": "分摊",
        "arbitration": "仲裁",
        "observance of terms": "遵守条款",
        "rights of insurer": "保险人权利",
        "other insurances": "其他保险",
        "reasonable precautions": "合理预防措施",
        "cancellation": "取消",
        "subrogation": "代位追偿",
        "premium": "保险费",
        "general conditions": "通用条款",

        # === 特别规定术语 (SPECIAL PROVISIONS) ===
        "immediate repairs": "即时修复",
        "interests of other parties": "其他方保险利益",
        "interests of mortgagees": "抵押权人保险利益",
        "premium payment warranty": "保费支付保证",
        "premium payment": "保费支付",
        "policy": "保单",
        "insured": "被保险人",
        "insurer": "保险人",
        "schedule": "明细表",
        "endorsement": "批单",
        "special provisions": "特别规定",

        # === 保险条款结构术语 ===
        "insuring clause": "保险条款",
        "preamble": "序言",
        "definitions": "定义",
        "coverage": "承保范围",
        "perils insured": "承保风险",
        "exclusions": "除外责任",
        "conditions": "条款条件",
        "memorandum": "备忘录",
        "warranties": "保证条款",
        "limit of liability": "责任限额",
        "deductible": "免赔额",
        "excess": "免赔额",
        "period of insurance": "保险期间",
        "territorial limits": "地域范围",
    }

    # ========================================
    # 语义别名映射表 (扩展版)
    # ========================================
    SEMANTIC_ALIAS_MAP = {
        # === 污染相关 ===
        "污染保险": "意外污染责任",
        "污染责任": "意外污染责任",
        "污染条款": "意外污染条款",
        "环境污染": "意外污染责任",

        # === 财产存放 ===
        "露天财产": "露天存放及简易建筑内财产",
        "露天物品": "露天存放及简易建筑内财产",
        "简易建筑": "露天存放及简易建筑内财产",
        "临时建筑": "露天存放及简易建筑内财产",

        # === 施救费用 ===
        "损害防止": "阻止损失",
        "施救费用": "阻止损失",
        "救援费用": "阻止损失",
        "抢险费用": "紧急抢险",

        # === 地质灾害 ===
        "崩塌沉降": "地面突然下陷下沉",
        "地面下陷": "地面突然下陷下沉",
        "地陷": "地面突然下陷下沉",
        "地面沉降": "地面突然下陷下沉",
        "山体滑坡": "泥石流、崩塌、突发性滑坡",
        "滑坡": "泥石流、崩塌、突发性滑坡",
        "泥石流": "泥石流、崩塌、突发性滑坡",

        # === 盗窃相关 ===
        "盗窃险": "盗窃、抢劫扩展",
        "盗抢险": "盗窃、抢劫扩展",
        "抢劫险": "盗窃、抢劫扩展",
        "入室盗窃": "盗窃、抢劫扩展",

        # === 自然灾害 ===
        "地震海啸": "地震扩展",
        "震动": "地震扩展",
        "台风": "台风、飓风扩展",
        "飓风": "台风、飓风扩展",
        "暴风": "暴风雨扩展",
        "暴雨": "暴风雨扩展",
        "水灾": "洪水扩展",
        "水淹": "洪水扩展",
        "内涝": "洪水扩展",
        "雷击": "雷电扩展",
        "雷电": "雷电扩展",
        "冰雹": "冰雹扩展",
        "雪灾": "暴雪、冰凌扩展",
        "冰凌": "暴雪、冰凌扩展",

        # === 机械设备 ===
        "锅炉爆炸": "锅炉、压力容器扩展",
        "压力容器": "锅炉、压力容器扩展",
        "玻璃破损": "玻璃破碎扩展",
        "玻璃险": "玻璃破碎扩展",
        "水管爆裂": "水箱、水管爆裂扩展",
        "水管破裂": "水箱、水管爆裂扩展",
        "喷淋系统": "自动喷淋系统水损扩展",
        "消防喷淋": "自动喷淋系统水损扩展",

        # === 责任相关 ===
        "公共责任": "公众责任",
        "第三者责任": "公众责任",
        "雇主责任": "雇员责任",
        "工伤责任": "雇员责任",
        "产品责任": "产品责任",
        "职业责任": "专业责任",

        # === 费用相关 ===
        "残骸清理": "清理残骸费用",
        "清除残骸": "清理残骸费用",
        "灭火费用": "灭火费用扩展",
        "消防费用": "灭火费用扩展",
        "空运费": "空运费扩展",
        "加急运费": "空运费扩展",
        "专业费": "专业费用扩展",
        "公估费": "专业费用扩展",

        # === 价值相关 ===
        "重置价值": "重置价值条款",
        "重建价值": "重置价值条款",
        "恢复保额": "自动恢复保险金额",
        "自动恢复": "自动恢复保险金额",

        # === 罢工暴乱 ===
        "罢工": "罢工、暴乱及民众骚乱",
        "暴乱": "罢工、暴乱及民众骚乱",
        "民众骚乱": "罢工、暴乱及民众骚乱",
        "骚乱": "罢工、暴乱及民众骚乱",
        "恐怖活动": "恐怖活动扩展",
        "恐怖袭击": "恐怖活动扩展",

        # === 运输相关 ===
        "内陆运输": "内陆运输扩展",
        "陆上运输": "内陆运输扩展",
        "临时移动": "临时移动扩展",
        "厂区运输": "厂区间临时移动扩展",

        # === v19.0: 雇主责任险相关 ===
        "自动保障新增雇员": "员工自动承保",
        "新增雇员自动保障": "员工自动承保",
        "保费调整": "保费调整条款",
        "就餐时间": "员工食堂",
        "上下班途中": "通勤",
    }

    # ========================================
    # 关键词映射表 (扩展版)
    # ========================================
    KEYWORD_MAP = {
        # === 自然灾害 ===
        "地震": ["地震", "震动", "earthquake", "seismic"],
        "海啸": ["海啸", "tsunami"],
        "洪水": ["洪水", "水灾", "水淹", "内涝", "flood", "inundation"],
        "台风": ["台风", "飓风", "typhoon", "hurricane", "cyclone"],
        "龙卷风": ["龙卷风", "tornado", "twister"],
        "暴风雨": ["暴风", "暴雨", "storm", "tempest"],
        "雷电": ["雷电", "雷击", "闪电", "lightning"],
        "冰雹": ["冰雹", "hail", "hailstone"],
        "暴雪": ["暴雪", "雪灾", "冰凌", "snowstorm", "icicle"],
        "沙尘暴": ["沙尘暴", "sandstorm", "dust storm"],

        # === 地质灾害 ===
        "滑坡": ["滑坡", "崩塌", "泥石流", "landslip", "landslide", "mudslide", "avalanche"],
        "地陷": ["地陷", "下陷", "沉降", "subsidence", "sinkhole"],

        # === 盗窃相关 ===
        "盗窃": ["盗窃", "盗抢", "抢劫", "入室", "burglary", "theft", "robbery"],
        "恶意破坏": ["恶意", "蓄意", "malicious", "vandalism"],

        # === 罢工暴乱 ===
        "罢工": ["罢工", "strike"],
        "暴乱": ["暴乱", "暴动", "riot"],
        "骚乱": ["骚乱", "民众骚乱", "civil commotion"],
        "恐怖": ["恐怖", "terrorism", "terrorist"],

        # === 污染相关 ===
        "污染": ["污染", "意外污染", "环境污染", "pollution", "contamination"],

        # === 设备相关 ===
        "锅炉": ["锅炉", "boiler"],
        "压力容器": ["压力容器", "pressure vessel"],
        "玻璃": ["玻璃", "glass"],
        "水管": ["水管", "水箱", "水损", "water pipe", "water tank"],
        "喷淋": ["喷淋", "消防喷淋", "sprinkler"],
        "电梯": ["电梯", "升降机", "扶梯", "elevator", "escalator", "lift"],
        "起重机": ["起重机", "起重", "吊车", "crane", "hoist"],

        # === 火灾相关 ===
        "火灾": ["火灾", "火险", "fire"],
        "自燃": ["自燃", "spontaneous combustion"],
        "爆炸": ["爆炸", "explosion"],
        "烟熏": ["烟熏", "smoke"],

        # === 价值相关 ===
        "重置": ["重置", "重建", "reinstatement", "replacement"],
        "定值": ["定值", "约定价值", "agreed value"],
        "恢复保额": ["恢复保额", "恢复保险金额", "reinstatement of sum"],
        "升值": ["升值", "增值", "escalation"],

        # === 费用相关 ===
        "残骸": ["残骸", "清理残骸", "debris", "removal of debris"],
        "灭火": ["灭火", "消防", "fire fighting", "fire brigade"],
        "空运费": ["空运费", "空运", "air freight", "airfreight"],
        "专业费用": ["专业费用", "公估", "professional fee"],
        "施救": ["施救", "救援", "抢险", "sue and labor", "sue & labor"],

        # === 责任相关 ===
        "公众责任": ["公众责任", "第三者", "public liability", "third party"],
        "产品责任": ["产品责任", "product liability"],
        "雇主责任": ["雇主责任", "雇员责任", "employer", "employee liability"],
        "交叉责任": ["交叉责任", "cross liability"],
        "契约责任": ["契约责任", "合同责任", "contractual liability"],

        # === 运输相关 ===
        "运输": ["运输", "transit", "transport"],
        "内陆运输": ["内陆运输", "inland transit"],
        "临时移动": ["临时移动", "temporary removal"],

        # === 工程相关 ===
        "工程": ["工程", "construction", "erection"],
        "试车": ["试车", "testing", "commissioning"],
        "保证期": ["保证期", "维护期", "maintenance", "defects liability"],
        "隧道": ["隧道", "tunnel"],
        "打桩": ["打桩", "桩基", "piling"],

        # === 其他 ===
        "72小时": ["72小时", "时间调整", "72 hours", "time adjustment"],
        "代位追偿": ["代位追偿", "代位", "subrogation"],
        "共同被保险人": ["共同被保险人", "joint insured"],
        "免赔额": ["免赔额", "免赔", "deductible", "excess"],
    }

    PENALTY_KEYWORDS = ["打孔盗气"]

    NOISE_WORDS = [
        "中国太平洋财产保险股份有限公司",
        "企业财产保险", "附加", "条款",
        "（A款）", "（B款）", "(A款)", "(B款)",
        "2025版", "2024版", "2023版", "2022版", "2026版",
        "clause", "extension", "cover",
    ]

    # v19.0: 险种类别前缀（用于上下文感知匹配）
    CATEGORY_PREFIXES = {
        "property": ["企业财产保险", "财产一切险", "财产综合险", "财产基本险",
                      "机器损坏保险", "锅炉及压力容器", "营业中断保险", "珠宝商综合保险"],
        "liability": ["雇主责任保险", "公众责任保险", "职业责任保险"],
    }

    # ========================================
    # 语义聚类（用于更智能的匹配）
    # ========================================
    SEMANTIC_CLUSTERS = {
        "地震类": ["地震", "震动", "地震海啸", "地震扩展", "earthquake"],
        "水灾类": ["洪水", "水灾", "暴雨", "水淹", "内涝", "flood", "inundation"],
        "盗窃类": ["盗窃", "盗抢", "抢劫", "入室盗窃", "burglary", "theft", "robbery"],
        "施救类": ["施救费用", "损害防止", "阻止损失", "救援费用", "sue and labor"],
        "台风类": ["台风", "飓风", "热带风暴", "typhoon", "hurricane"],
        "火灾类": ["火灾", "火险", "燃烧", "fire"],
        "罢工类": ["罢工", "暴乱", "骚乱", "民众骚乱", "strike", "riot", "civil commotion"],
        "责任类": ["责任", "赔偿", "liability", "indemnity"],
    }

    # ========================================
    # 特殊规则（v18.1）
    # 当客户条款名称匹配特定模式时，返回预定义的提示信息
    # 格式: {
    #   "patterns": [匹配模式列表],
    #   "matched_name": "显示的匹配名称",
    #   "message": "提示信息",
    #   "match_level": "匹配级别"
    # }
    # ========================================
    SPECIAL_RULES = [
        {
            # 制造商/供应商担保条款 - 考虑各种变体
            "patterns": [
                "制造商/供应商担保条款",
                "制造商／供应商担保条款",  # 全角斜杠
                "制造商 / 供应商担保条款",  # 带空格
                "制造商/ 供应商担保条款",
                "制造商 /供应商担保条款",
                "制造商供应商担保条款",  # 无分隔符
                "manufacturer/supplier warranty",
                "manufacturer / supplier warranty",
                "manufacturer's warranty",
                "supplier's warranty",
            ],
            "matched_name": "主条款相关约定",
            "message": "主条款已有相关约定：被保险人已经从有关责任方取得赔偿的，保险人赔偿保险金时，可以相应扣减被保险人已从有关责任方取得的赔偿金额。",
            "match_level": "精确匹配",
        },
        {
            # 合同争议解决
            "patterns": [
                "合同争议解决",
                "争议解决",
                "合同争议",
            ],
            "matched_name": "主条款已有相关约定",
            "message": "主条款已有相关约定：因履行本合同发生的争议，由当事人协商解决，协商不成的，依法向保险标的所在地法院起诉。",
            "match_level": "精确匹配",
        },
        {
            # 责任免除第七条修改 - 除外责任明晰条款
            "patterns": [
                "责任免除第七条",
                "责任免除第七条（七）修改",
                "责任免除第七条(七)修改",
                "兹经双方同意，责任免除第七条",
                "但因此造成其他财产的损失不在此限",
                "造成其他财产的损失不在此限",
            ],
            "matched_name": "企业财产保险附加除外责任明晰条款",
            "message": "匹配条款：企业财产保险附加除外责任明晰条款。该条款对责任免除第七条（七）进行了修改，明确\"但因此造成其他财产的损失不在此限\"。",
            "match_level": "精确匹配",
        },
        {
            # "三停"损失保险 - 供应水电气中断
            "patterns": [
                "由于供应水、电、气",
                "供应水、电、气及其他能源",
                "供应发生故障或中断",
                "三停",
                "公共设施当局",
            ],
            "matched_name": "企业财产保险附加'三停'损失保险",
            "message": "匹配条款：企业财产保险附加'三停'损失保险。该条款承保因供应水、电、气等能源中断造成的损失。",
            "match_level": "精确匹配",
        },
    ]


# ==========================================
# 编辑距离算法
# ==========================================
@lru_cache(maxsize=10000)
def levenshtein_distance(s1: str, s2: str) -> int:
    """计算编辑距离（带缓存）"""
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)

    if len(s2) == 0:
        return len(s1)

    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row

    return previous_row[-1]


def levenshtein_ratio(s1: str, s2: str) -> float:
    """计算编辑距离相似度"""
    if not s1 or not s2:
        return 0.0

    # 长度差异过大直接返回低分
    len_diff = abs(len(s1) - len(s2))
    max_len = max(len(s1), len(s2))
    if len_diff > max_len * 0.6:
        return 0.0

    distance = levenshtein_distance(s1, s2)
    return 1 - (distance / max_len)


# ==========================================
# 核心匹配逻辑（重构版）
# ==========================================
class ClauseMatcherLogic:
    """条款匹配核心逻辑 - 优化版"""

    # ===== v18.5: 常量定义 =====
    # 条款标题最大长度
    MAX_TITLE_LENGTH_DEFAULT = 150   # 中文条款标题通常较短
    MAX_TITLE_LENGTH_ENGLISH = 250   # 英文条款标题可能较长，包含完整描述

    # ===== v18.5: 预编译正则表达式（性能优化）=====
    _RE_CLAUSE_KEYWORDS = re.compile(r'\b(Clause|Extension|Coverage|Endorsement)\b', re.IGNORECASE)
    _RE_MONEY_PATTERN = re.compile(
        r'(RMB|CNY|人民币|美元|USD|EUR|HKD|港币)?\s*\d+[\d,\.]*\s*(万元|元|万|亿|千元)',
        re.IGNORECASE
    )
    _RE_SUB_NUMBER = re.compile(r'^\d+\.[A-Z]')  # 子编号格式: 1.REINSTATEMENT
    _RE_LEADING_NUMBER = re.compile(r'^\d+[\.\s、]+')  # 开头编号
    _RE_PARENTHESIS_NUMBER = re.compile(r'^[\(（]\s*\d+\s*[\)）]')  # (1), （2）
    _RE_PARENTHESIS_LETTER = re.compile(r'^[\(（]\s*[a-zA-Z]\s*[\)）]')  # (a), (b)
    _RE_LETTER_PAREN = re.compile(r'^[a-z]\)')  # a), b)
    _RE_ROMAN_NUMBER = re.compile(r'^[ivxIVX]+[\.\)]')  # i., ii.
    _RE_CONTENT_STARTER = re.compile(
        r'^\d+[\.\s]+\s*(The|It|In|Any|This|Where|If|When|Unless|Subject)\s',
        re.IGNORECASE
    )

    # v18.4: 排除词汇缓存（完全匹配时排除，忽略编号和大小写）
    _excluded_titles: Optional[set] = None

    @classmethod
    def _load_excluded_titles(cls) -> set:
        """加载排除词汇列表"""
        if cls._excluded_titles is not None:
            return cls._excluded_titles

        cls._excluded_titles = set()
        config_path = Path(__file__).parent / "excluded_titles.json"

        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    titles = data.get('titles', [])
                    # 转换为大写存储，便于比较
                    cls._excluded_titles = {t.upper().strip() for t in titles if t}
                    logger.info(f"加载排除词汇 {len(cls._excluded_titles)} 条")
            except Exception as e:
                logger.error(f"加载排除词汇失败: {e}")

        return cls._excluded_titles

    @staticmethod
    def _is_blue_text(run) -> bool:
        """
        v18.9: 检测文字是否为蓝色（用于精准识别模式）
        支持多种蓝色格式：纯蓝色、深蓝、浅蓝等
        """
        try:
            from docx.shared import RGBColor

            # 检查 run 的字体颜色
            if run.font and run.font.color:
                color = run.font.color

                # 方式1: RGB 颜色
                if color.rgb:
                    r, g, b = color.rgb[0], color.rgb[1], color.rgb[2]
                    # 判断是否为蓝色系：B分量明显大于R和G
                    # 宽松匹配：B > 100 且 B > R 且 B > G
                    if b > 100 and b > r and b > g:
                        return True
                    # 也支持深蓝色 (0, 0, 128) 等
                    if b >= 128 and r < 100 and g < 100:
                        return True

                # 方式2: 主题颜色（如 MSO_THEME_COLOR.ACCENT_1 等）
                # Word 中的蓝色主题通常是 ACCENT_1 或 ACCENT_5
                if color.theme_color is not None:
                    # 主题颜色索引：1, 5 通常是蓝色系
                    # 但这取决于文档主题，保守处理
                    pass

            return False
        except Exception as e:
            logger.debug(f"检测蓝色文字失败: {e}")
            return False

    @staticmethod
    def _extract_blue_text_from_paragraph(para) -> str:
        """
        v18.9: 从段落中提取蓝色文字
        返回所有蓝色run的文本拼接
        """
        blue_texts = []
        for run in para.runs:
            if ClauseMatcherLogic._is_blue_text(run):
                text = run.text.strip()
                if text:
                    blue_texts.append(text)
        return ''.join(blue_texts).strip()

    @staticmethod
    def _extract_blue_text_from_cell(cell) -> str:
        """
        v18.9: 从表格单元格中提取蓝色文字
        """
        blue_texts = []
        for para in cell.paragraphs:
            text = ClauseMatcherLogic._extract_blue_text_from_paragraph(para)
            if text:
                blue_texts.append(text)
        return '\n'.join(blue_texts).strip()

    @staticmethod
    def _remove_leading_number(text: str) -> str:
        """去除开头的编号，如 '1.', '（一）', '(1)' 等"""
        text = text.strip()
        # 去除各种编号格式
        patterns = [
            r'^[\(（]\s*[一二三四五六七八九十\d]+\s*[\)）]\s*',  # (一)、（1）
            r'^[一二三四五六七八九十]+[、\.．]\s*',  # 一、二、
            r'^\d+[、\.．\s]\s*',  # 1、2.
            r'^[A-Za-z]\)\s*',  # a)、A)
        ]
        for pattern in patterns:
            text = re.sub(pattern, '', text)
        return text.strip()

    @staticmethod
    def _is_valid_clause_line(text: str) -> bool:
        """
        v18.6: 宽松的条款行验证（用于从表格条款区域提取）
        已经确定在"附加条款/Extension"区域，只需排除明显不是条款的内容
        """
        if not text or len(text) < 3:
            return False

        # 排除太长的行（可能是正文内容）
        # v18.6: 但如果以条款关键词开头，放宽到300字符
        max_len = 200
        if re.search(r'\b(CLAUSE|EXTENSION|COVER|INSURANCE|条款)\b', text, re.IGNORECASE):
            max_len = 300  # 条款标题可能包含 Limit 说明，放宽限制
        if len(text) > max_len:
            return False

        # 排除以句号结尾的长句（正文内容）
        if text.endswith(('。', '；')) and len(text) > 50:
            return False

        # 排除以小写字母开头的英文句子（正文内容）
        if text and text[0].islower() and len(text) > 30:
            return False

        # 排除明显的正文开头
        content_starts = (
            '本条款', '本保险', '本附加', '保险人', '被保险人', '投保人',
            '如果', '若', '当', '在', '对于', '经双方', '兹经', '因履行',
            '但', '无论', '特别条件', '重置价值是指', '交付日期', '每次事故免赔额',
            '被保险财产若', '中华人民共和国法律',
            'The insurer', 'The insured', 'If ', 'When ', 'Where ',
            'Subject to', 'Provided that', 'It is agreed', 'It is further',
            'It is hereby', 'It is understood', 'The limit', 'The deductible',
            'The amount', 'All the terms', 'Any breach', 'Any disputes',
            'Limit of indemnity', 'Headings have', 'Sedgwick', 'McLarens', 'Charles Taylor',
        )
        if text.startswith(content_starts):
            return False

        # v18.7.2: 排除以中文分号结尾的内容（正文句子）
        if text.endswith('；'):
            return False

        # v18.7: 排除以英文句号结尾的长句（条款内容）
        if text.endswith('.') and len(text) > 80:
            return False

        # v18.7.2: 排除以冒号结尾的定义行
        if text.endswith(('：', ':')):
            return False

        # 排除纯数字或金额
        if re.match(r'^[\d,\.\s]+$', text):
            return False
        if re.match(r'^(RMB|CNY|USD|EUR)\s*[\d,\.]+', text, re.IGNORECASE):
            return False

        # v18.7.3: 排除中文编号开头的子项（但保留包含"条款"关键词的）
        # "（1）.", "(一）", "①", "1、"
        has_clause_keyword = '条款' in text or 'clause' in text.lower() or 'extension' in text.lower()
        if re.match(r'^[\(（]\s*[\d一二三四五六七八九十]+\s*[\)）][\.\s、]?', text) and not has_clause_keyword:
            return False
        if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text) and not has_clause_keyword:
            return False
        if re.match(r'^\d+[、]', text) and not has_clause_keyword:  # "1、保单文本..." - 移除\s*
            return False

        # 排除编号开头的子项（如 "1. xxx", "(a) xxx", "1)xxx", "1.1 xxx"）
        if re.match(r'^\d+\.\s+[a-z]', text):  # "1. the liability..."
            return False
        if re.match(r'^[\(（]\s*[a-zA-Z\d]+\s*[\)）]\s+[a-z]', text):  # "(a) the..."
            return False
        if re.match(r'^\d+\)\s*[a-z]', text):  # "1)theft..."
            return False
        if re.match(r'^\d+\.\d+\s', text):  # "1.1 Damage..."
            return False

        # v18.7: 排除括号编号后跟"The said"等内容
        if re.match(r'^[\(（]\s*[a-zA-Z]\s*[\)）]\s+(The said|In the event)', text):
            return False

        # v18.7.3: 排除公司名（含 Ltd/Co./有限公司）
        # 只过滤主要是公司名的行，保留包含条款关键词的行
        if ('Ltd' in text or 'Co.' in text or '有限公司' in text) and '条款' not in text and 'clause' not in text.lower():
            return False

        return True

    # 条款库中的常见样板内容（这些内容不影响匹配度计算）
    BOILERPLATE_PHRASES = [
        "本条款所述费用在本保险单明细表所列保险金额之外另行赔付。",
        "本附加险根据保险单的约定收取保险费。",
        "本保险单所载其他条件不变。",
        "本附加条款与主条款内容相悖之处，以本附加条款为准；未尽之处，以主条款为准。",
        "限额由保险双方约定并在保险单中载明。",
        "本条款未尽事宜，以主保险合同的条款为准。",
        "本附加险条款与主险条款相抵触之处，以本附加险条款为准。",
        "本保险合同所载其他条款、条件和除外责任不变。",
        "本附加险保费按主险保费的一定比例收取。",
        "本条款中任何未定义的词语或术语具有主保险合同中规定的含义。",
    ]

    def __init__(self):
        """初始化匹配器"""
        # 加载配置
        if HAS_CONFIG_MANAGER:
            self.config = get_config()
            self._use_external_config = True
        else:
            self.config = None
            self._use_external_config = False

        self.thresholds = MatchThresholds()
        self._index: Optional[LibraryIndex] = None

        # v17.0: TF-IDF向量索引
        self._tfidf_vectorizer = None
        self._tfidf_vectors = None
        self._tfidf_names = []

        # v19.0: 险种上下文（由 MatchWorker 设置）
        self._current_category: str = ""

        logger.info(f"匹配器初始化完成，外部配置: {self._use_external_config}")
        logger.info(f"jieba分词: {HAS_JIEBA}, sklearn(TF-IDF): {HAS_SKLEARN}")

    @staticmethod
    def detect_category_from_sheet(sheet_name: str) -> str:
        """v19.0: 从 sheet 名称检测险种类别"""
        if not sheet_name:
            return ""
        sheet_lower = sheet_name.lower()
        if any(k in sheet_lower for k in ["liability", "责任", "雇主"]):
            return "liability"
        if any(k in sheet_lower for k in ["property", "财产", "企业"]):
            return "property"
        return ""

    def _detect_lib_category(self, lib_name: str) -> str:
        """v19.0: 从条款库名称检测险种类别"""
        if not lib_name:
            return ""
        for category, prefixes in DefaultConfig.CATEGORY_PREFIXES.items():
            for prefix in prefixes:
                if prefix in lib_name:
                    return category
        return ""

    @classmethod
    def remove_boilerplate(cls, content: str) -> str:
        """
        从内容中移除样板文字，用于更准确的相似度计算
        """
        if not content:
            return ""
        result = content
        for phrase in cls.BOILERPLATE_PHRASES:
            result = result.replace(phrase, "")
        # 移除多余的空白和换行
        result = re.sub(r'\s+', ' ', result).strip()
        return result

    @staticmethod
    def _normalize_for_special_rules(text: str) -> str:
        """
        标准化文本用于特殊规则匹配
        - 全角转半角
        - 移除空格
        - 转小写
        """
        if not text:
            return ""

        result = []
        for char in text:
            code = ord(char)
            # 全角空格
            if code == 0x3000:
                continue  # 移除空格
            # 全角字符范围 (！到～)
            elif 0xFF01 <= code <= 0xFF5E:
                result.append(chr(code - 0xFEE0))
            # 普通空格
            elif char == ' ':
                continue  # 移除空格
            else:
                result.append(char)

        return ''.join(result).lower()

    def check_special_rules(self, clause_title: str) -> Optional[MatchResult]:
        """
        检查条款是否匹配特殊规则
        返回 MatchResult 如果匹配，否则返回 None
        """
        if not clause_title:
            return None

        normalized_title = self._normalize_for_special_rules(clause_title)

        for rule in DefaultConfig.SPECIAL_RULES:
            patterns = rule.get("patterns", [])

            for pattern in patterns:
                normalized_pattern = self._normalize_for_special_rules(pattern)

                # 包含匹配（任一方向）
                if normalized_pattern in normalized_title or normalized_title in normalized_pattern:
                    # 匹配成功，返回特殊结果
                    match_level_str = rule.get("match_level", "精确匹配")
                    match_level = MatchLevel.EXACT
                    if "语义" in match_level_str:
                        match_level = MatchLevel.SEMANTIC
                    elif "关键词" in match_level_str:
                        match_level = MatchLevel.KEYWORD

                    logger.info(f"特殊规则匹配: '{clause_title}' -> '{rule.get('matched_name')}'")

                    return MatchResult(
                        matched_name=rule.get("matched_name", "特殊规则匹配"),
                        matched_content=rule.get("message", ""),
                        matched_reg="",
                        score=1.0,
                        title_score=1.0,
                        content_score=0.0,
                        match_level=match_level,
                        diff_analysis=rule.get("message", ""),
                    )

        return None

    # ========================================
    # 配置访问方法
    # ========================================

    def _get_client_mapping(self, term: str) -> Optional[str]:
        """获取英中映射"""
        if self._use_external_config:
            return self.config.get_client_mapping(term)
        return DefaultConfig.CLIENT_EN_CN_MAP.get(term.lower())

    def _get_semantic_alias(self, text: str) -> Optional[str]:
        """获取语义别名"""
        alias_map = (self.config.semantic_alias_map if self._use_external_config
                     else DefaultConfig.SEMANTIC_ALIAS_MAP)
        for alias, target in alias_map.items():
            if alias in text:
                return target
        return None

    def _get_keywords(self, text: str) -> Set[str]:
        """提取关键词"""
        keywords = set()
        text_lower = text.lower()
        keyword_map = (self.config.keyword_extract_map if self._use_external_config
                       else DefaultConfig.KEYWORD_MAP)
        for core, variants in keyword_map.items():
            for v in variants:
                if v.lower() in text_lower:
                    keywords.add(core)
                    break
        return keywords

    def _is_penalty_keyword(self, text: str) -> bool:
        """检查惩罚关键词"""
        penalty_list = (self.config.penalty_keywords if self._use_external_config
                        else DefaultConfig.PENALTY_KEYWORDS)
        return any(kw in text for kw in penalty_list)

    def _get_noise_words(self) -> List[str]:
        """获取噪音词列表"""
        return (self.config.noise_words if self._use_external_config
                else DefaultConfig.NOISE_WORDS)

    # ========================================
    # 文本处理方法
    # ========================================

    @staticmethod
    def normalize_text(text: str) -> str:
        """标准化文本"""
        if not isinstance(text, str):
            return ""
        text = text.lower().strip()
        text = re.sub(r"['\"\'\'\"\"\(\)（）\[\]【】]", '', text)
        text = re.sub(r'\s+', ' ', text)
        return text

    def clean_title(self, text: str) -> str:
        """清理标题

        v18.9: 保留版本标识符（如A款、B款等），避免索引冲突
        例如: "企业财产保险附加公共当局扩展条款（A款）" -> "企业财产保险附加公共当局扩展条款a款"
        """
        if not isinstance(text, str):
            return ""
        # 提取版本标识符（A款、B款、C款等）
        version_match = re.search(r'[（(]([A-Za-z]款)[）)]', text)
        version_suffix = version_match.group(1).lower() if version_match else ""

        # 移除所有括号内容
        text = re.sub(r'[\(（].*?[\)）]', '', text)
        for w in self._get_noise_words():
            text = text.replace(w, "").replace(w.lower(), "")
        text = re.sub(r'[0-9\s]+', '', text)

        # 重新添加版本标识
        result = text.strip()
        if version_suffix:
            result = result + version_suffix
        return result

    @classmethod
    def clean_content(cls, text: str) -> str:
        """
        清理内容，用于相似度计算
        1. 移除样板文字
        2. 移除括号内容
        3. 移除空白和数字
        """
        if not isinstance(text, str):
            return ""
        # 先移除样板文字
        text = cls.remove_boilerplate(text)
        # 移除括号及其内容
        text = re.sub(r'[\(（].*?[\)）]', '', text)
        # 移除空白
        text = re.sub(r'\s+', '', text)
        # 移除数字
        text = re.sub(r'[0-9]+', '', text)
        return text

    # ========================================
    # 除外条款检测 (v17.1)
    # ========================================

    @staticmethod
    def is_exclusion_clause(lib_name: str) -> bool:
        """
        v17.1: 判断条款库中的条款是否为除外条款
        例如: "企业财产保险附加洪水除外条款（A款）"
        """
        if not lib_name:
            return False
        return '除外' in lib_name

    @staticmethod
    def client_wants_exclusion(client_title: str) -> bool:
        """
        v17.1: 判断客户条款名称是否明确包含"除外"
        只有当客户条款明确包含"除外"时，才应匹配除外类条款
        """
        if not client_title:
            return False
        return '除外' in client_title

    @staticmethod
    def extract_extra_info(text: str) -> str:
        """提取括号内额外信息"""
        if not isinstance(text, str):
            return ""
        matches = re.findall(r'([\(（].*?[\)）])', text)
        return " ".join(matches) if matches else ""

    @staticmethod
    def is_english(text: str) -> bool:
        """判断是否为英文"""
        if not isinstance(text, str) or len(text) <= 3:
            return False
        zh_count = len(re.findall(r'[\u4e00-\u9fa5]', text))
        return zh_count < len(text) * 0.15

    @staticmethod
    def is_bilingual(text: str) -> bool:
        """
        v17.0: 判断是否为中英混合文本
        如: "Earthquake Extension Clause 地震扩展条款"
        """
        if not isinstance(text, str) or len(text) < 5:
            return False
        zh_count = len(re.findall(r'[\u4e00-\u9fa5]', text))
        en_count = len(re.findall(r'[a-zA-Z]', text))
        total = len(text)
        # 中英文各占15%以上视为双语
        return (zh_count / total >= 0.15) and (en_count / total >= 0.15)

    @staticmethod
    def split_bilingual(text: str) -> Tuple[str, str]:
        """
        v17.0: 分离中英文部分
        返回: (中文部分, 英文部分)
        """
        if not isinstance(text, str):
            return "", ""
        # 提取中文（包括中文标点）
        cn_chars = re.findall(r'[\u4e00-\u9fa5\u3000-\u303f\uff00-\uffef]+', text)
        cn_part = ''.join(cn_chars)
        # 提取英文单词
        en_words = re.findall(r'[a-zA-Z]+(?:\s+[a-zA-Z]+)*', text)
        en_part = ' '.join(en_words)
        return cn_part.strip(), en_part.strip().lower()

    @staticmethod
    def tokenize_chinese(text: str) -> List[str]:
        """
        v17.0: 中文分词
        使用jieba分词，如果不可用则使用字符级别分割
        """
        if not text:
            return []
        if HAS_JIEBA:
            # 使用jieba分词，过滤单字符和标点
            words = list(jieba.cut(text))
            return [w for w in words if len(w) > 1 and re.search(r'[\u4e00-\u9fa5a-zA-Z]', w)]
        else:
            # 降级：使用2-gram字符分割
            chars = re.findall(r'[\u4e00-\u9fa5]', text)
            if len(chars) < 2:
                return chars
            return [chars[i] + chars[i+1] for i in range(len(chars) - 1)]

    # ========================================
    # 相似度计算（v17.0 增强版）
    # ========================================

    @staticmethod
    def calculate_similarity(text1: str, text2: str) -> float:
        """
        混合相似度计算：
        - SequenceMatcher（序列匹配）
        - Levenshtein（编辑距离）
        取较高值
        """
        if not text1 or not text2:
            return 0.0

        # 序列匹配
        seq_ratio = difflib.SequenceMatcher(None, text1, text2).ratio()

        # 编辑距离（仅对较短文本使用，避免性能问题）
        if len(text1) <= 100 and len(text2) <= 100:
            lev_ratio = levenshtein_ratio(text1, text2)
            return max(seq_ratio, lev_ratio)

        return seq_ratio

    @classmethod
    def calculate_similarity_chinese(cls, text1: str, text2: str) -> float:
        """
        v17.0: 中文增强相似度计算
        结合词级别Jaccard相似度和字符级别相似度
        """
        if not text1 or not text2:
            return 0.0

        # 1. 字符级别相似度（基础）
        char_sim = cls.calculate_similarity(text1, text2)

        # 2. 词级别Jaccard相似度（如果jieba可用）
        if HAS_JIEBA:
            words1 = set(cls.tokenize_chinese(text1))
            words2 = set(cls.tokenize_chinese(text2))

            if words1 and words2:
                intersection = words1 & words2
                union = words1 | words2
                jaccard_sim = len(intersection) / len(union) if union else 0

                # 加权组合：词级别权重更高
                combined_sim = 0.6 * jaccard_sim + 0.4 * char_sim
                return max(combined_sim, char_sim)  # 取较高值

        return char_sim

    def calculate_bilingual_similarity(self, text1: str, text2: str) -> float:
        """
        v17.0: 双语条款相似度计算
        分别计算中英文部分的相似度，然后加权组合
        """
        if not text1 or not text2:
            return 0.0

        # 检查是否为双语文本
        is_bi1 = self.is_bilingual(text1)
        is_bi2 = self.is_bilingual(text2)

        if not is_bi1 and not is_bi2:
            # 都不是双语，使用标准相似度
            return self.calculate_similarity_chinese(text1, text2)

        # 分离中英文
        cn1, en1 = self.split_bilingual(text1)
        cn2, en2 = self.split_bilingual(text2)

        scores = []

        # 计算中文部分相似度
        if cn1 and cn2:
            cn_sim = self.calculate_similarity_chinese(cn1, cn2)
            scores.append(('cn', cn_sim, len(cn1) + len(cn2)))

        # 计算英文部分相似度
        if en1 and en2:
            en_sim = self.calculate_similarity(en1, en2)
            scores.append(('en', en_sim, len(en1) + len(en2)))

        # 尝试英文映射匹配
        if en1:
            mapped = self._get_client_mapping(en1)
            if mapped and cn2:
                map_sim = self.calculate_similarity_chinese(mapped, cn2)
                if map_sim > 0.7:
                    scores.append(('map', map_sim, 100))  # 高权重

        if not scores:
            return self.calculate_similarity(text1, text2)

        # 按长度加权平均
        total_weight = sum(s[2] for s in scores)
        weighted_sim = sum(s[1] * s[2] for s in scores) / total_weight if total_weight > 0 else 0

        # 取加权结果和最高单项的较大值
        max_sim = max(s[1] for s in scores)
        return max(weighted_sim, max_sim * 0.9)

    # ========================================
    # TF-IDF 向量匹配 (v17.0)
    # ========================================

    def build_tfidf_index(self, lib_data: List[Dict]) -> None:
        """
        v17.0: 构建TF-IDF向量索引，用于快速候选筛选
        """
        if not HAS_SKLEARN:
            logger.warning("sklearn不可用，跳过TF-IDF索引构建")
            return

        names = []
        for lib in lib_data:
            name = str(lib.get('条款名称', ''))
            if name.strip():
                # 对中文进行分词处理
                if HAS_JIEBA:
                    tokens = ' '.join(self.tokenize_chinese(name))
                    names.append(tokens if tokens else name)
                else:
                    names.append(name)

        if not names:
            return

        try:
            # 使用字符n-gram，适合中文
            self._tfidf_vectorizer = TfidfVectorizer(
                analyzer='char',
                ngram_range=(2, 4),
                max_features=5000,
                min_df=1
            )
            self._tfidf_vectors = self._tfidf_vectorizer.fit_transform(names)
            self._tfidf_names = names
            logger.info(f"TF-IDF索引构建完成，向量维度: {self._tfidf_vectors.shape}")
        except Exception as e:
            logger.warning(f"TF-IDF索引构建失败: {e}")
            self._tfidf_vectorizer = None
            self._tfidf_vectors = None

    def find_tfidf_candidates(self, query: str, top_k: int = 10) -> List[Tuple[int, float]]:
        """
        v17.0: 使用TF-IDF快速找到候选条款
        返回: [(索引, 相似度分数), ...]
        """
        if not HAS_SKLEARN or self._tfidf_vectorizer is None or self._tfidf_vectors is None:
            return []

        try:
            # 对查询进行同样的预处理
            if HAS_JIEBA:
                query_tokens = ' '.join(self.tokenize_chinese(query))
                query_text = query_tokens if query_tokens else query
            else:
                query_text = query

            query_vec = self._tfidf_vectorizer.transform([query_text])
            similarities = cosine_similarity(query_vec, self._tfidf_vectors).flatten()

            # 获取top_k个最相似的索引
            top_indices = np.argsort(similarities)[-top_k:][::-1]
            results = [(int(idx), float(similarities[idx])) for idx in top_indices if similarities[idx] > 0.1]

            return results
        except Exception as e:
            logger.debug(f"TF-IDF候选查找失败: {e}")
            return []

    # ========================================
    # 动态权重计算 (v17.0)
    # ========================================

    def calculate_dynamic_weight(self, title: str, content: str) -> Tuple[float, float]:
        """
        v17.0: 根据条款特征动态调整标题/内容权重
        返回: (标题权重, 内容权重)
        """
        # 默认权重
        title_weight = 0.7
        content_weight = 0.3

        title_len = len(title) if title else 0
        content_len = len(content) if content else 0

        # 情况1: 标题很短且内容丰富 -> 增加内容权重
        if title_len < 10 and content_len > 100:
            title_weight = 0.4
            content_weight = 0.6

        # 情况2: 标题很长（可能包含详细描述）-> 增加标题权重
        elif title_len > 30:
            title_weight = 0.8
            content_weight = 0.2

        # 情况3: 无内容 -> 全部使用标题
        elif content_len < 10:
            title_weight = 1.0
            content_weight = 0.0

        # 情况4: 标题包含特定关键词（表示具体条款类型）-> 增加标题权重
        specific_keywords = ['扩展条款', '除外条款', '特别条款', '附加险', 'extension', 'exclusion']
        if any(kw in title.lower() for kw in specific_keywords):
            title_weight = min(title_weight + 0.1, 0.9)
            content_weight = 1.0 - title_weight

        return title_weight, content_weight

    # ========================================
    # 索引构建（性能优化核心）
    # ========================================

    def build_index(self, lib_data: List[Dict]) -> LibraryIndex:
        """
        预构建条款库索引，加速匹配
        时间复杂度从 O(n*m) 降至 O(n + m)
        """
        logger.info(f"开始构建索引，条款数: {len(lib_data)}")

        index = LibraryIndex(data=lib_data)

        for i, lib in enumerate(lib_data):
            name = str(lib.get('条款名称', ''))
            if not name.strip():
                continue

            # 预计算清理结果（避免重复计算）
            name_norm = self.normalize_text(name)
            name_clean = self.clean_title(name)

            index.cleaned_cache[i] = {
                'norm': name_norm,
                'clean': name_clean,
                'original': name,
            }

            # 名称索引（精确匹配用）
            index.by_name_norm[name_norm] = i
            index.by_name_norm[name_clean] = i

            # 关键词倒排索引
            keywords = self._get_keywords(name)
            for kw in keywords:
                index.by_keyword[kw].append(i)

        logger.info(f"索引构建完成: {len(index.by_name_norm)} 名称, {len(index.by_keyword)} 关键词")
        self._index = index

        # v17.0: 构建TF-IDF索引
        self.build_tfidf_index(lib_data)

        return index

    @staticmethod
    def _fullwidth_to_halfwidth(text: str) -> str:
        """全角字符转半角"""
        result = []
        for char in text:
            code = ord(char)
            if code == 0x3000:  # 全角空格
                result.append(' ')
            elif 0xFF01 <= code <= 0xFF5E:  # 全角字符范围
                result.append(chr(code - 0xFEE0))
            else:
                result.append(char)
        return ''.join(result)

    def find_library_entry_by_name(self, target_name: str, index: LibraryIndex) -> Optional[Dict]:
        """
        根据名称在条款库中查找条目
        支持全角/半角字符模糊匹配
        """
        if not target_name:
            return None

        # 标准化目标名称
        target_norm = self._fullwidth_to_halfwidth(target_name.lower().strip())
        target_clean = re.sub(r'[^\u4e00-\u9fa5a-z0-9%]', '', target_norm)

        best_match_idx = -1
        best_score = 0.0

        for i, cached in index.cleaned_cache.items():
            lib_name = cached['original']
            lib_norm = self._fullwidth_to_halfwidth(lib_name.lower().strip())
            lib_clean = re.sub(r'[^\u4e00-\u9fa5a-z0-9%]', '', lib_norm)

            # 精确匹配（标准化后）
            if target_clean == lib_clean:
                return index.data[i]

            # 包含匹配
            if target_clean in lib_clean or lib_clean in target_clean:
                score = len(target_clean) / max(len(lib_clean), 1)
                if score > best_score:
                    best_score = score
                    best_match_idx = i

            # 相似度匹配
            sim = self.calculate_similarity(target_clean, lib_clean)
            if sim > best_score and sim > 0.8:
                best_score = sim
                best_match_idx = i

        if best_match_idx >= 0:
            return index.data[best_match_idx]

        return None

    @staticmethod
    def clean_reg_number(reg: str) -> str:
        """清理注册号，移除前缀"""
        if not reg:
            return ""
        # 移除 "产品注册号：" 等前缀
        reg = re.sub(r'^(产品)?注册号[：:]\s*', '', str(reg).strip())
        return reg

    # ========================================
    # 多级匹配策略（拆分重构）
    # ========================================

    def _try_exact_match(self, title_norm: str, title_clean: str,
                         index: LibraryIndex, original_title: str = "") -> Optional[Tuple[int, float]]:
        """
        级别1: 精确匹配 (v17.0 增强)
        - 标准化名称匹配
        - 英中映射表匹配
        - 双语条款分离匹配
        """
        # 标准化名称精确匹配
        if title_norm in index.by_name_norm:
            return index.by_name_norm[title_norm], 1.0

        # 清理后名称精确匹配
        if title_clean in index.by_name_norm:
            return index.by_name_norm[title_clean], self.thresholds.exact_min

        # v17.0: 英中映射表匹配
        if original_title:
            # 提取英文部分尝试映射
            _, en_part = self.split_bilingual(original_title)
            if en_part:
                mapped_cn = self._get_client_mapping(en_part)
                if mapped_cn:
                    mapped_norm = self.normalize_text(mapped_cn)
                    mapped_clean = self.clean_title(mapped_cn)
                    if mapped_norm in index.by_name_norm:
                        return index.by_name_norm[mapped_norm], 0.95
                    if mapped_clean in index.by_name_norm:
                        return index.by_name_norm[mapped_clean], 0.93
                    # 部分匹配
                    for i, cached in index.cleaned_cache.items():
                        if mapped_cn in cached['original'] or cached['original'] in mapped_cn:
                            return i, 0.90

        return None

    def _try_semantic_match(self, title: str, index: LibraryIndex) -> Optional[Tuple[int, float]]:
        """级别2: 语义别名匹配"""
        semantic_target = self._get_semantic_alias(title)
        if not semantic_target:
            return None

        # 在索引中查找目标
        for i, cached in index.cleaned_cache.items():
            if semantic_target in cached['original']:
                return i, self.thresholds.semantic_min

        return None

    def _try_keyword_match(self, title: str, index: LibraryIndex) -> Optional[Tuple[int, float]]:
        """级别3: 关键词匹配"""
        c_keywords = self._get_keywords(title)
        if not c_keywords:
            return None

        # 统计候选项得分
        candidate_scores: Dict[int, float] = defaultdict(float)

        for kw in c_keywords:
            if kw in index.by_keyword:
                for idx in index.by_keyword[kw]:
                    candidate_scores[idx] += 1

        if not candidate_scores:
            return None

        # 找最高分候选
        best_idx = max(candidate_scores, key=candidate_scores.get)
        best_count = candidate_scores[best_idx]

        # 计算关键词匹配度
        l_keywords = self._get_keywords(index.cleaned_cache[best_idx]['original'])
        if l_keywords:
            keyword_ratio = best_count / max(len(c_keywords), len(l_keywords))
            if keyword_ratio >= 0.5:
                score = self.thresholds.keyword_min + keyword_ratio * 0.2
                return best_idx, score

        return None

    def _try_fuzzy_match(self, title_clean: str, content: str,
                         index: LibraryIndex, is_title_only: bool,
                         original_title: str = "", max_results: int = 1) -> Any:
        """
        级别4: 模糊匹配 (v17.1 增强版)
        - 使用TF-IDF快速候选筛选
        - 使用中文分词增强相似度
        - 支持双语匹配
        - 动态权重调整
        - v17.1: 支持返回多条匹配结果
        - v17.1: 除外条款过滤

        Args:
            max_results: 返回结果数量，1为单个结果(兼容旧接口)，>1为多个结果列表

        Returns:
            当max_results=1时: Tuple[int, float, float, float] - (idx, score, title_sim, content_sim)
            当max_results>1时: List[Tuple[int, float, float, float]] - 候选列表
        """
        # v17.1: 检查客户是否需要除外类条款
        wants_exclusion = self.client_wants_exclusion(original_title)

        candidates = []

        # v17.0: 计算动态权重
        title_weight, content_weight = self.calculate_dynamic_weight(title_clean, content)

        # v17.0: 使用TF-IDF快速筛选候选（如果可用）
        candidate_indices = set()
        tfidf_candidates = self.find_tfidf_candidates(original_title or title_clean, top_k=20)
        if tfidf_candidates:
            candidate_indices = {idx for idx, _ in tfidf_candidates}
            # 同时也检查所有条款（以防TF-IDF遗漏）
            # 但优先处理TF-IDF候选
        else:
            candidate_indices = set(index.cleaned_cache.keys())

        # 如果TF-IDF候选较少，添加所有条款确保覆盖
        if len(candidate_indices) < 10:
            candidate_indices = set(index.cleaned_cache.keys())

        for i in candidate_indices:
            if i not in index.cleaned_cache:
                continue
            cached = index.cleaned_cache[i]
            l_name_clean = cached['clean']
            l_name_original = cached['original']

            # v17.1: 除外条款过滤 - 除非客户明确需要除外条款，否则跳过库内的除外条款
            if not wants_exclusion and self.is_exclusion_clause(l_name_original):
                continue

            # v17.0: 使用增强相似度计算
            # 先检查是否为双语匹配
            if original_title and (self.is_bilingual(original_title) or self.is_bilingual(l_name_original)):
                title_sim = self.calculate_bilingual_similarity(original_title, l_name_original)
            else:
                # 使用中文增强相似度
                title_sim = self.calculate_similarity_chinese(title_clean, l_name_clean)

            # 内容相似度
            content_sim = 0.0
            if not is_title_only and content.strip():
                c_content_clean = self.clean_content(content)
                l_content = str(index.data[i].get('条款内容', ''))
                l_content_clean = self.clean_content(l_content)
                if c_content_clean and l_content_clean:
                    # v17.0: 对内容也使用中文增强相似度
                    content_sim = self.calculate_similarity_chinese(c_content_clean, l_content_clean)

            # v17.0: 使用动态权重加权得分
            if is_title_only or not content.strip():
                score = title_sim
            else:
                score = title_weight * title_sim + content_weight * content_sim

            # 惩罚项
            if self._is_penalty_keyword(cached['original']) and not self._is_penalty_keyword(title_clean):
                score -= 0.5

            # v19.0: 险种上下文感知 - 同险种加分，跨险种减分
            if self._current_category:
                lib_category = self._detect_lib_category(l_name_original)
                if lib_category == self._current_category:
                    score += 0.15
                elif lib_category and lib_category != self._current_category:
                    score -= 0.25

            if score > self.thresholds.accept_min:
                candidates.append((i, score, title_sim, content_sim))

        # 按分数降序排序
        candidates.sort(key=lambda x: x[1], reverse=True)

        # v17.1: 根据max_results返回不同格式
        if max_results == 1:
            # 兼容旧接口
            if candidates:
                return candidates[0]
            return (-1, 0.0, 0.0, 0.0)
        else:
            # 返回多条结果
            return candidates[:max_results]

    def extract_limit_info(self, clause_name: str) -> tuple:
        """v18.15: 提取条款名称末尾的限额/约定信息

        返回: (去除限额后的名称, 限额信息)

        保留的括号格式（不作为限额提取）:
        - (A款), (B款), (2025版), (简易版), (通用), (甲类), (乙类)

        提取的限额格式:
        - (Limit: xxx), (Contract Limit: xxx)
        - (72Hours), (48小时)
        - (80%), (15% of xxx)
        - (RMB xxx), (USD xxx)
        - (per occurrence), (per accident)
        - (on stock), (World-wide)
        - (limit of xxx)
        """
        if not clause_name:
            return ('', '')

        # 保留模式 - 这些括号内容不作为限额提取
        preserve_patterns = [
            r'[（\(]\s*[A-Z甲乙丙丁]\s*款\s*[）\)]',  # (A款), (甲款)
            r'[（\(]\s*\d{4}\s*版?\s*[）\)]',  # (2025版), (2025)
            r'[（\(]\s*简易版?\s*[）\)]',  # (简易版)
            r'[（\(]\s*通用\s*[）\)]',  # (通用)
            r'[（\(]\s*[甲乙丙丁]类\s*[）\)]',  # (甲类)
            r'[（\(]\s*标准版?\s*[）\)]',  # (标准版)
        ]

        # 检查末尾是否是需要保留的括号
        for pattern in preserve_patterns:
            if re.search(pattern + r'\s*$', clause_name, re.IGNORECASE):
                return (clause_name, '')

        # 限额提取模式
        limit_patterns = [
            # Limit 相关
            r'\s*[（\(]\s*(?:Contract\s+)?Limit[：:\s]*[^）\)]+[）\)]\s*$',
            r'\s*[（\(]\s*limit\s+of[^）\)]+[）\)]\s*$',
            # 时间限制
            r'\s*[（\(]\s*\d+\s*[Hh]ours?\s*[）\)]\s*$',
            r'\s*[（\(]\s*\d+\s*小时\s*[）\)]\s*$',
            r'\s*[（\(]\s*\d+\s*[Dd]ays?\s*[）\)]\s*$',
            r'\s*[（\(]\s*\d+\s*天\s*[）\)]\s*$',
            # 百分比
            r'\s*[（\(]\s*\d+\.?\d*\s*%(?:\s*of[^）\)]*)?[）\)]\s*$',
            # 金额
            r'\s*[（\(]\s*(?:RMB|CNY|USD|EUR|HKD|人民币|美元)[\s\d,\.万亿元]+[^）\)]*[）\)]\s*$',
            # per occurrence/accident
            r'\s*[（\(]\s*per\s+(?:occurrence|accident|event|loss|claim)[^）\)]*[）\)]\s*$',
            # on stock, World-wide 等
            r'\s*[（\(]\s*on\s+stock[^）\)]*[）\)]\s*$',
            r'\s*[（\(]\s*[Ww]orld-?\s*wide[^）\)]*[）\)]\s*$',
            # 每次/每年
            r'\s*[（\(]\s*每[次年月][^）\)]*[）\)]\s*$',
            # 最高/最低
            r'\s*[（\(]\s*最[高低][^）\)]*[）\)]\s*$',
            # 免赔
            r'\s*[（\(]\s*免赔[^）\)]*[）\)]\s*$',
            r'\s*[（\(]\s*[Dd]eductible[^）\)]*[）\)]\s*$',
        ]

        # 尝试匹配并提取限额
        for pattern in limit_patterns:
            match = re.search(pattern, clause_name, re.IGNORECASE)
            if match:
                limit_info = match.group(0).strip()
                # 清理限额信息中的括号
                limit_info = re.sub(r'^[（\(]\s*', '', limit_info)
                limit_info = re.sub(r'\s*[）\)]\s*$', '', limit_info)
                base_name = clause_name[:match.start()].strip()
                return (base_name, limit_info)

        return (clause_name, '')

    def match_clause(self, clause: ClauseItem, index: LibraryIndex,
                     is_title_only: bool) -> MatchResult:
        """
        主匹配入口 - 多级策略
        优先级: 精确 > 语义 > 关键词 > 模糊
        """
        result = MatchResult()
        title = clause.title
        content = clause.content

        title_clean = self.clean_title(title)
        title_norm = self.normalize_text(title)

        matched_idx = -1
        match_level = MatchLevel.NONE
        score = 0.0
        title_score = 0.0
        content_score = 0.0

        # === 级别1: 精确匹配 (v17.0: 传递原始标题用于英中映射) ===
        exact_result = self._try_exact_match(title_norm, title_clean, index, original_title=title)
        if exact_result:
            matched_idx, score = exact_result
            match_level = MatchLevel.EXACT
            title_score = score

        # === 级别2: 语义匹配 ===
        if matched_idx < 0:
            semantic_result = self._try_semantic_match(title, index)
            if semantic_result:
                matched_idx, score = semantic_result
                match_level = MatchLevel.SEMANTIC
                title_score = score

        # === 级别3: 关键词匹配 ===
        if matched_idx < 0:
            keyword_result = self._try_keyword_match(title, index)
            if keyword_result:
                matched_idx, score = keyword_result
                match_level = MatchLevel.KEYWORD
                title_score = score

        # === 级别4: 模糊匹配 (v17.0 增强) ===
        if matched_idx < 0:
            fuzzy_idx, fuzzy_score, t_sim, c_sim = self._try_fuzzy_match(
                title_clean, content, index, is_title_only,
                original_title=title  # v17.0: 传递原始标题用于双语匹配
            )
            if fuzzy_score > self.thresholds.accept_min:
                matched_idx = fuzzy_idx
                score = fuzzy_score
                match_level = MatchLevel.FUZZY
                title_score = t_sim
                content_score = c_sim

        # 构建结果
        if matched_idx >= 0 and score > self.thresholds.accept_min:
            lib = index.data[matched_idx]
            base_name = lib.get('条款名称', '')
            extra_params = self.extract_extra_info(clause.original_title or clause.title)

            result.matched_name = f"{base_name} {extra_params}".strip() if extra_params else base_name
            result.matched_content = lib.get('条款内容', '')
            result.matched_reg = lib.get('产品注册号', lib.get('注册号', ''))
            result.score = max(0, score)
            result.title_score = title_score
            result.content_score = content_score
            result.match_level = match_level

            # 差异分析（低分时）
            if score < 0.6:
                result.diff_analysis = self.analyze_difference(content, result.matched_content)

        return result

    def match_clause_multiple(self, clause: ClauseItem, index: LibraryIndex,
                               is_title_only: bool, max_results: int = 3) -> List[MatchResult]:
        """
        v17.1: 多结果匹配入口
        返回最多max_results条匹配结果供用户选择

        Args:
            clause: 待匹配的条款
            index: 条款库索引
            is_title_only: 是否仅匹配标题
            max_results: 最多返回结果数，默认3条

        Returns:
            List[MatchResult]: 匹配结果列表，按分数降序排列
        """
        title = clause.title
        content = clause.content
        original_title = clause.original_title or title

        # v18.1: 首先检查特殊规则
        special_result = self.check_special_rules(original_title)
        if special_result is None and title != original_title:
            # 如果原标题没匹配，也检查翻译后的标题
            special_result = self.check_special_rules(title)

        if special_result:
            return [special_result]

        title_clean = self.clean_title(title)

        results = []
        seen_names = set()

        # 获取多条模糊匹配候选
        fuzzy_candidates = self._try_fuzzy_match(
            title_clean, content, index, is_title_only,
            original_title=original_title,
            max_results=max_results + 5  # 多获取一些以便去重
        )

        # fuzzy_candidates是列表: [(idx, score, title_sim, content_sim), ...]
        if isinstance(fuzzy_candidates, tuple):
            # 单结果模式返回的tuple
            if fuzzy_candidates[0] >= 0:
                fuzzy_candidates = [fuzzy_candidates]
            else:
                fuzzy_candidates = []

        for idx, score, title_sim, content_sim in fuzzy_candidates:
            if len(results) >= max_results:
                break

            if score <= self.thresholds.accept_min:
                continue

            lib = index.data[idx]
            base_name = lib.get('条款名称', '')

            # 去重
            if base_name in seen_names:
                continue
            seen_names.add(base_name)

            extra_params = self.extract_extra_info(original_title)

            result = MatchResult(
                matched_name=f"{base_name} {extra_params}".strip() if extra_params else base_name,
                matched_content=lib.get('条款内容', ''),
                matched_reg=lib.get('产品注册号', lib.get('注册号', '')),
                score=max(0, score),
                title_score=title_sim,
                content_score=content_sim,
                match_level=MatchLevel.FUZZY,
                diff_analysis=""
            )

            # 差异分析（低分时）
            if score < 0.6 and content:
                result.diff_analysis = self.analyze_difference(content, result.matched_content)

            results.append(result)

        # 如果没有任何匹配，返回空匹配结果
        if not results:
            results.append(MatchResult())

        return results

    def search_library_titles(self, query: str, index: LibraryIndex,
                               max_results: int = 5) -> List[Dict]:
        """
        v17.1: 条款查询功能 - 仅查询条款标题
        用于快速模糊查询条款库中的条款

        Args:
            query: 查询字符串（条款名称或关键词）
            index: 条款库索引
            max_results: 最多返回结果数，默认5条

        Returns:
            List[Dict]: 查询结果列表，每项包含:
                - name: 条款名称
                - content: 条款内容
                - reg: 注册号
                - score: 匹配分数
                - matchType: 匹配类型
        """
        if not query or not index.data:
            return []

        # 检查查询是否需要除外条款
        wants_exclusion = self.client_wants_exclusion(query)

        query_lower = query.lower()
        query_clean = self.clean_title(query)
        query_norm = self.normalize_text(query)

        candidates = []

        for i, cached in index.cleaned_cache.items():
            lib_name = cached['original']
            lib_name_clean = cached['clean']
            lib_name_norm = cached['norm']
            lib_name_lower = lib_name.lower()

            # 除外条款过滤
            if not wants_exclusion and self.is_exclusion_clause(lib_name):
                continue

            match_type = ""
            score = 0.0

            # 精确匹配
            if query_norm == lib_name_norm or query_clean == lib_name_clean:
                match_type = "exact"
                score = 1.0
            # 包含匹配
            elif query_lower in lib_name_lower or lib_name_lower in query_lower:
                match_type = "contain"
                score = 0.9
            elif query_clean in lib_name_clean or lib_name_clean in query_clean:
                match_type = "contain"
                score = 0.85
            else:
                # 模糊相似度匹配
                title_sim = self.calculate_similarity_chinese(query_clean, lib_name_clean)
                if title_sim > 0.3:  # 较低的阈值以便显示更多可能的匹配
                    match_type = "fuzzy"
                    score = title_sim

            if score > 0:
                lib = index.data[i]
                candidates.append({
                    'name': lib_name,
                    'content': lib.get('条款内容', ''),
                    'reg': lib.get('产品注册号', lib.get('注册号', '')),
                    'score': score,
                    'matchType': match_type
                })

        # 按分数降序排序
        candidates.sort(key=lambda x: x['score'], reverse=True)

        # 去重并限制结果数量
        results = []
        seen_names = set()
        for c in candidates:
            if len(results) >= max_results:
                break
            if c['name'] not in seen_names:
                seen_names.add(c['name'])
                results.append(c)

        return results

    def create_user_mapping_result(self, lib_entry: Dict, user_library_name: str) -> MatchResult:
        """
        v18.5: 根据用户映射创建匹配结果（提取重复代码）

        Args:
            lib_entry: 条款库条目
            user_library_name: 用户映射的条款库名称

        Returns:
            MatchResult: 匹配结果
        """
        if lib_entry:
            return MatchResult(
                matched_name=lib_entry.get('条款名称', user_library_name),
                matched_reg=self.clean_reg_number(lib_entry.get('产品注册号', lib_entry.get('注册号', ''))),
                matched_content=lib_entry.get('条款内容', ''),
                score=1.0,
                match_level=MatchLevel.EXACT,
                diff_analysis="用户自定义映射",
                title_score=1.0,
                content_score=0.0,
            )
        else:
            # 映射的条款在库中不存在
            return MatchResult(
                matched_name=user_library_name,
                matched_reg="",
                matched_content="",
                score=1.0,
                match_level=MatchLevel.EXACT,
                diff_analysis="用户映射（条款库中未找到）",
                title_score=1.0,
                content_score=0.0,
            )

    # ========================================
    # 翻译和差异分析
    # ========================================

    def translate_title(self, title: str) -> Tuple[str, bool]:
        """翻译英文标题"""
        if not self.is_english(title):
            return title, False

        title_norm = self.normalize_text(title)

        # 1. 查询映射
        mapped = self._get_client_mapping(title_norm)
        if mapped:
            return mapped, True

        # 2. 部分匹配
        client_map = (self.config.client_en_cn_map if self._use_external_config
                      else DefaultConfig.CLIENT_EN_CN_MAP)
        for eng, chn in client_map.items():
            if eng in title_norm or title_norm in eng:
                return chn, True

        return title, False

    @staticmethod
    def analyze_difference(c_content: str, l_content: str) -> str:
        """分析保障差异"""
        c_text, l_text = str(c_content), str(l_content)
        if not c_text.strip():
            return ""

        analysis = []
        keywords = {
            "限额": ["Limit", "限额", "最高", "limit"],
            "免赔": ["Deductible", "Excess", "免赔", "deductible"],
            "除外": ["Exclusion", "除外", "不负责", "exclusion"],
            "观察期": ["Waiting Period", "观察期", "等待期"],
            "赔偿期": ["Indemnity Period", "赔偿期间"],
        }

        for key, words in keywords.items():
            c_has = any(w.lower() in c_text.lower() for w in words)
            l_has = any(w.lower() in l_text.lower() for w in words)
            if c_has and not l_has:
                analysis.append(f"⚠️ 客户提及[{key}]但库内未提及")
            elif not c_has and l_has:
                analysis.append(f"ℹ️ 库内包含[{key}]但客户未提及")

        return " | ".join(analysis)

    # ========================================
    # 文档解析
    # ========================================

    @classmethod
    def is_likely_title(cls, text: str) -> bool:
        """
        判断是否像标题（严格模式）
        只有明确符合标题特征的才返回True
        v17.1: 增强过滤规则
        """
        if not text or len(text) < 3:
            return False

        # ===== v18.4: 排除词汇检查（最高优先级）=====
        # 去除编号后完全匹配（忽略大小写）则排除
        excluded_titles = ClauseMatcherLogic._load_excluded_titles()
        if excluded_titles:
            cleaned_text = ClauseMatcherLogic._remove_leading_number(text)
            if cleaned_text.upper() in excluded_titles:
                return False

        # ===== v18.2: 特殊长条款识别（在长度检查之前）=====
        # 这些是特殊的长文本条款，需要被识别为条款标题
        special_long_clause_patterns = [
            '兹经双方同意，责任免除第七条',  # 除外责任明晰条款
            '责任免除第七条（七）修改',
            '责任免除第七条(七)修改',
            '由于供应水、电、气',  # "三停"损失保险
            '供应水、电、气及其他能源',
        ]
        for pattern in special_long_clause_patterns:
            if pattern in text:
                return True

        # 太长的不是标题
        # v18.5: 使用类常量替代硬编码值
        max_length = cls.MAX_TITLE_LENGTH_DEFAULT
        if cls._RE_CLAUSE_KEYWORDS.search(text):
            max_length = cls.MAX_TITLE_LENGTH_ENGLISH
        if len(text) > max_length:
            return False

        # 以句号等结尾的通常是内容（但排除 ":" 和 "）"，这些在条款标题中常见）
        if text.endswith(('。', '；', '.', ';', '，', ',')):
            # 但如果包含条款关键词，可能是标题带了额外说明
            if not cls._RE_CLAUSE_KEYWORDS.search(text):
                return False

        # ===== v18.2: 特殊标题关键词（优先检查）=====
        # 这些短标题虽然不含"条款"但确实是条款名称
        special_title_keywords = [
            '合同争议解决', '争议解决', '合同争议',
            '自动恢复保险金额', '恢复保险金额',
            '通译和标题', '错误和遗漏', '错误与遗漏',
            '权益保障', '损失通知', '不受控制',
            '品牌和商标', '合同价格',
        ]
        for kw in special_title_keywords:
            if kw in text:
                return True

        # ===== v18.4: 英文特殊条款关键词（无Clause/Extension但确实是条款）=====
        english_special_keywords = [
            'Burglary', 'Theft', 'Robbery',  # 盗窃抢劫
            'Strike', 'Riot', 'Civil Commotion',  # 罢工暴动
            'Works of Arts', 'Work of Art',  # 艺术品
            'Cancellation by Insurer', 'Cancellation by Insured',  # 注销条款
            'Notice of Cancellation',  # 注销通知
            'Property in the Open',  # 露天财产
            'Unnamed location', 'Unnamed Location',  # 未指定地点
            'Miscellaneous',  # 杂项（但不在excluded中时）
        ]
        # 检查是否包含英文特殊关键词（需要至少匹配一个）
        for kw in english_special_keywords:
            if kw.lower() in text.lower():
                # 额外检查：排除明显是正文的情况
                content_starts = ('the ', 'this ', 'if ', 'when ', 'where ', 'by ', 'and ', 'or ',
                                  'provided ', 'subject ', 'in ', 'for ', 'any ', 'all ', 'such ')
                if text.lower().startswith(content_starts):
                    continue  # 跳过这个关键词，继续检查其他
                # 额外检查：以小写字母开头的通常是正文
                if text and text[0].islower():
                    continue
                return True

        # ===== v17.1: 优先检查是否为标题（"条款"关键词最优先）=====

        # 包含"条款"关键词，但排除以"本条款"、"本扩展条款"、"本附加条款"开头的内容句
        # 这个检查必须在 descriptive_keywords 之前！否则"恢复保险金额条款"会被错误排除
        if '条款' in text:
            if text.startswith(('本条款', '本扩展条款', '本附加条款')):
                return False
            return True

        # v18.2: 包含"附加"和"保险"的也可能是条款标题（如"企业财产保险附加自动恢复保险金额保险"）
        if '附加' in text and '保险' in text:
            # 排除以"本附加"开头的内容句
            if not text.startswith(('本附加', '在附加')):
                return True

        # ===== 排除明确不是标题的内容 =====

        # 1. 排除包含金额的内容（如 "RMB50万元"、"CNY5000万元"、"人民币100万"）
        money_pattern = r'(RMB|CNY|人民币|美元|USD|EUR|HKD|港币)?\s*\d+[\d,\.]*\s*(万元|元|万|亿|千元)'
        if re.search(money_pattern, text, re.IGNORECASE):
            return False

        # 2. 排除包含"赔偿限额"、"保险金额"等描述性文字的内容
        # 注意：如果包含"条款"关键词，上面已经返回True，不会到达这里
        descriptive_keywords = ['赔偿限额', '保险金额', '责任限额', '每次事故', '累计赔偿',
                                '免赔额', '自负额', '保险费', '费率', '保险期间']
        if any(kw in text for kw in descriptive_keywords):
            return False

        # 3. 排除特定的内容句（完整匹配或开头匹配）
        excluded_exact = [
            '本扩展条款受下列条件限制',
            '特约扩展责任',
        ]
        if text in excluded_exact:
            return False

        # ===== 其他标题模式检查 =====

        # 带数字编号的条款标题（如 "35、码头吊机、铁路车辆第三者责任险"）
        # 支持格式：1、xxx, 1.xxx, 1）xxx, (1) xxx, 一、xxx 等
        numbered_title_pattern = r'^(\d+|[一二三四五六七八九十]+)[、\.．）\)]'
        if re.match(numbered_title_pattern, text):
            # 但如果后面是描述性内容则排除
            title_part = re.sub(numbered_title_pattern, '', text).strip()
            if title_part and len(title_part) > 3 and not title_part.endswith(('。', '；', '，')):
                # 检查是否包含"险"、"条款"等标志性词汇
                if any(kw in title_part for kw in ['险', '条款', '责任', '扩展', '附加']):
                    return True

        # 附加保险条款，以 "(XXXX版)" 结尾（无"条款"字样）
        # 如：平安产险企业财产保险附加提前60天通知解除保单保险（2025版）
        if '附加' in text and '保险' in text and re.search(r'[（(]\d{4}版?[）)]$', text):
            return True

        # ===== v18.3: 英文条款关键词优先检查（在排除检查之前）=====
        # 包含 Clause/Extension/Coverage/Cover/Insurance 的英文文本通常是条款标题
        # 注意：Clauses 是复数形式，Cover 是 Coverage 的简写
        if re.search(r'\b(Clauses?|Extensions?|Coverage|Cover|Endorsement|Insurance)\b', text, re.IGNORECASE):
            # v18.4 修复1: 排除保险公司名称（包含 "Insurance Company" 或 "Insurance Co."）
            if re.search(r'Insurance\s+(Company|Co\.?)\b', text, re.IGNORECASE):
                return False

            # v18.4 修复2: 排除 "this/the + 关键词" 形式（条款正文内容）
            # 如 "this Clause", "the Policy", "this extension", "this Endorsement"
            if re.search(r'\b(this|the|such|that)\s+(Clause|Extension|Policy|Insurance|Cover|Endorsement)\b', text, re.IGNORECASE):
                return False

            # v18.4 修复3: 排除编号开头的内容（条款正文的子项）
            # v18.5: 使用预编译正则提升性能
            if cls._RE_PARENTHESIS_NUMBER.match(text):  # (1), (2), （1）
                return False
            if cls._RE_PARENTHESIS_LETTER.match(text):  # (a), (b), (c), (A), (B)
                return False
            if cls._RE_LETTER_PAREN.match(text):  # a), b), c)
                return False
            if cls._RE_ROMAN_NUMBER.match(text):  # i., ii., iii.
                return False

            # v18.5 修复8: 排除"数字+点+紧跟大写字母（无空格）"的子编号内容
            # 如 "1.REINSTATEMENT VALUE CLAUSE" - 这是条款正文的子项，不是独立条款
            if cls._RE_SUB_NUMBER.match(text):
                return False

            # v18.4 修复5: 排除"数字+点+The/It/In/Any..."开头的子项内容
            # 如 "1. The liability of...", "2. It is agreed that..."
            if cls._RE_CONTENT_STARTER.match(text):
                return False

            # v18.4 修复6: 排除以正文开头词开始的内容
            # 如 "Provided that...", "If the sum...", "by fire caused..."
            content_starters = (
                'Provided ', 'If ', 'Where ', 'When ', 'Unless ', 'Subject to ',
                'In the event ', 'In respect ', 'For the purpose ', 'Notwithstanding ',
                'by ', 'and ', 'or ', 'but ', 'that ', 'which ', 'who ', 'whose ',
            )
            if text.startswith(content_starters):
                return False

            # v18.4 修复7: 以小写字母开头的通常是正文内容
            if text and text[0].islower():
                return False

            # v18.4 修复4: 排除以冒号结尾的全大写文本（如 WARRANTED:）
            if text.isupper() and text.rstrip().endswith(':'):
                return False

            # 排除其他明显不是标题的情况
            if not text.startswith(('All the terms',)):
                return True

        # ===== 明确是内容的模式（不是标题）=====
        content_start_patterns = [
            # 条款内容常见开头
            r'^经双方同意',
            r'^兹经双方同意',
            r'^兹经保险',
            r'^兹经合同',
            r'^发生.*损失',
            r'^如果.*保险',
            r'^本保单',
            r'^本保险',
            r'^本条款',
            r'^本款项',
            r'^本公司',
            r'^本扩展条款',  # v17.1
            r'^本附加条款',  # v17.1
            r'^保险人',
            r'^被保险人',
            r'^投保人',
            r'^对于',
            r'^若',
            r'^但',
            r'^在保',
            r'^上述',
            r'^该',
            r'^其中',  # v17.1
            r'^此',
            r'^当',
            r'^财产险',
            r'^除',
            r'^凡',
            r'^任何',
            r'^无论',
            r'^特别条件',
            r'^重置价值是指',
            # 金额和免赔额描述（不是条款标题）
            r'^每次事故免赔额',
            r'^每次事故赔偿限额',
            r'^每次及累计',
            r'^累计赔偿限额',
            r'^RMB\s*[\d,]+',
            r'^\d+[\.,]\d+',  # 纯数字开头
            # 公司名称（不是条款标题）- v18.3: 只排除明确的公司名，不要太宽泛
            r'^Charles\s+Taylor',
            r'^McLarens',
            r'^Sedgwick',
            r'^Crawford',
            # 交付日期等说明
            r'^交付日期',
            r'^分期数',
            # 列表项（子条目，不是新条款）
            r'^[\(（]\s*[一二三四五六七八九十]+\s*[\)）]',  # (一)、（二）
            r'^[一二三四五六七八九十]+[、\.．]',  # 一、二、
            r'^\d+[、\.．\s](?![\.．\s]*[^\d].*条款)',  # 1、2、但不匹配 "1. xxx条款"
            r'^[\(（]\s*\d+\s*[\)）]',  # (1)、（2）
            r'^①|^②|^③|^④|^⑤',  # 圈数字
        ]

        for pattern in content_start_patterns:
            if re.match(pattern, text):
                return False

        # ===== 其他标题模式（已通过内容排除检查）=====
        # 全大写英文（可能是英文条款名）
        if text.isupper() and len(text) > 5 and re.search(r'[A-Z]{3,}', text):
            # v18.4: 排除以冒号结尾的（如 WARRANTED:）
            if text.rstrip().endswith(':'):
                return False
            return True

        # 默认不是标题（保守策略）
        return False

    def parse_docx(self, doc_path: str, precise_mode: bool = False) -> Tuple[List[ClauseItem], bool]:
        """
        解析Word文档 - 智能识别表格中的条款列表

        Args:
            doc_path: Word文档路径
            precise_mode: v18.9 精准识别模式 - 只提取蓝色字体的条款

        Returns:
            (条款列表, 是否为纯标题模式)
        """
        logger.info(f"解析文档: {doc_path}, 精准模式: {precise_mode}")

        try:
            doc = Document(doc_path)
        except Exception as e:
            logger.error(f"文档打开失败: {e}")
            raise ValueError(f"无法打开文档: {e}")

        # ===== v18.9: 精准识别模式 - 只提取蓝色文字 =====
        if precise_mode:
            return self._parse_docx_precise_mode(doc)

        # 1. 读取普通段落，同时记录样式信息
        # v18.4: 使用 Heading 样式作为条款标题的强识别信号
        all_lines = []
        heading_lines = set()  # 记录哪些行是 Heading 样式

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            all_lines.append(text)

            # 检查是否是 Heading 样式（条款标题通常使用 Heading 样式）
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name or 'title' in style_name:
                    if text:  # 只记录非空的 Heading
                        heading_lines.add(i)

        # 2. 智能读取表格内容 - 特别处理"附加条款"列
        table_clauses = []  # 从"附加条款"单元格提取的条款
        table_lines = []    # 其他表格内容

        # v18.8: 预加载已映射的客户条款名称（用于表格条款提取时的优先识别）
        mapped_client_names_for_table = set()
        try:
            if HAS_MAPPING_MANAGER:
                mapping_mgr = get_mapping_manager()
                if mapping_mgr:
                    for mapping in mapping_mgr.get_all_mappings():
                        if mapping.client_name:
                            mapped_client_names_for_table.add(mapping.client_name.strip())
                            # 也添加去除编号后的名称
                            cleaned = re.sub(r'^[\d\(\)（）]+[\.\s、]*', '', mapping.client_name).strip()
                            if cleaned:
                                mapped_client_names_for_table.add(cleaned)
                    if mapped_client_names_for_table:
                        logger.info(f"表格提取: 已加载 {len(mapped_client_names_for_table)} 个已映射条款名称")
        except Exception as e:
            logger.warning(f"加载映射条款名称失败: {e}")

        def is_mapped_clause(line: str) -> bool:
            """检查是否是已映射的条款（优先识别）"""
            if not mapped_client_names_for_table:
                return False
            # 精确匹配
            if line in mapped_client_names_for_table:
                return True
            # 去除编号后匹配
            cleaned = re.sub(r'^[\d\(\)（）]+[\.\s、]*', '', line).strip()
            if cleaned and cleaned in mapped_client_names_for_table:
                return True
            return False

        # 定义条款列的关键词（中英文，不区分大小写）
        # v18.6: 扩展英文关键词支持
        clause_row_keywords_cn = ['附加条款', '除外条款', '特别条款', '扩展条款', '承保条款', '特别约定']
        # v18.7: 移除 coverage/coverages，太宽泛会误匹配"保障范围"
        clause_row_keywords_en = ['extension', 'extensions', 'exclusion', 'exclusions',
                                   'special provisions', 'special provision',
                                   'conditions', 'condition']

        # v18.9: 检测"纯条款列表表格"（每行都是条款标题，没有区域标记）
        def is_clause_list_table(table) -> bool:
            """检测表格是否是纯条款列表（每行一个条款）"""
            if len(table.rows) < 5:  # 至少5行才考虑
                return False
            if len(table.columns) > 3:  # 最多3列（英文|中文|备注）
                return False

            clause_suffix_count = 0
            total_rows = 0

            for row in table.rows:
                # 获取主要内容（优先取中文列）
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells:
                    continue

                # 选择最后一个有内容的单元格（通常是中文）
                text = cells[-1] if len(cells) > 1 else cells[0]
                if not text or len(text) < 3:
                    continue

                total_rows += 1

                # 检查是否以条款相关关键词结尾
                clause_suffixes = ('条款', '扩展', '除外', '责任', '保障', '保险', '险',
                                   'clause', 'extension', 'exclusion', 'coverage')
                if any(text.lower().endswith(suffix) for suffix in clause_suffixes):
                    clause_suffix_count += 1
                # 或者包含条款关键词
                elif '条款' in text or 'clause' in text.lower():
                    clause_suffix_count += 1

            # 如果超过60%的行以条款关键词结尾，认为是纯条款列表
            if total_rows > 0 and clause_suffix_count / total_rows > 0.6:
                return True
            return False

        def extract_from_clause_list_table(table) -> List[str]:
            """从纯条款列表表格中提取所有条款"""
            clauses = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells:
                    continue

                # 如果是双列表格（英文|中文），取中文列
                if len(cells) >= 2:
                    # 检查是否是英中对照
                    if re.search(r'[a-zA-Z]', cells[0]) and re.search(r'[\u4e00-\u9fff]', cells[-1]):
                        text = cells[-1]  # 取中文
                    else:
                        text = cells[0]
                else:
                    text = cells[0]

                if text and len(text) >= 3 and len(text) <= 200:
                    # 基本过滤
                    if not text.startswith(('备注', 'Note', 'Remark', '说明')):
                        clauses.append(text)
            return clauses

        # v18.9: 先检测是否有纯条款列表表格
        for table in doc.tables:
            if is_clause_list_table(table):
                extracted = extract_from_clause_list_table(table)
                if extracted:
                    logger.info(f"检测到纯条款列表表格，提取到 {len(extracted)} 个条款")
                    table_clauses.extend(extracted)

        # 如果从纯条款列表表格中提取到了条款，直接返回
        if table_clauses:
            clauses = [ClauseItem(title=t, content="", original_title=t) for t in table_clauses]
            return clauses, True

        for table in doc.tables:
            in_clause_region = False  # v18.6: 标记是否在条款区域内
            clause_content_col = -1   # v18.7: 记录条款内容所在的列索引

            for row in table.rows:
                first_cell_text = row.cells[0].text.strip()

                # 检查是否是条款列表行（中英文关键词匹配）
                first_cell_lower = first_cell_text.lower()
                is_clause_row = (
                    any(kw in first_cell_text for kw in clause_row_keywords_cn) or
                    any(kw in first_cell_lower for kw in clause_row_keywords_en)
                )

                if is_clause_row:
                    in_clause_region = True  # 进入条款区域
                    # v18.7: 查找包含条款的列（从后往前找第一个有内容且不是标签的列）
                    for col_idx in range(len(row.cells) - 1, -1, -1):
                        cell_text = row.cells[col_idx].text.strip()
                        # 跳过标签单元格和分隔符
                        if cell_text and cell_text != first_cell_text and cell_text not in ['：', ':', '']:
                            clause_content_col = col_idx  # 记录条款内容列
                            # 按换行分割，提取所有条款
                            lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
                            for line in lines:
                                # v18.8: 已映射的条款优先识别，跳过常规验证
                                if is_mapped_clause(line) or self._is_valid_clause_line(line):
                                    table_clauses.append(line)
                            break

                elif in_clause_region:
                    # v18.7: 检查是否遇到新的区域标记（退出条款区域）
                    exit_keywords = ['备注', 'remark', 'note', '免赔', 'deductible', 'excess',
                                     '费率', 'rate', '保费', 'premium']
                    if any(kw in first_cell_lower for kw in exit_keywords):
                        in_clause_region = False
                        clause_content_col = -1
                        continue

                    # v18.7: 在条款区域内，处理"每行一个条款"的结构
                    content_found = False

                    # 方案1: 使用已知的条款内容列
                    if clause_content_col > 0 and clause_content_col < len(row.cells):
                        cell_text = row.cells[clause_content_col].text.strip()
                        if cell_text:
                            lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
                            for line in lines:
                                # v18.8: 已映射的条款优先识别，跳过常规验证
                                if is_mapped_clause(line) or self._is_valid_clause_line(line):
                                    table_clauses.append(line)
                                    content_found = True

                    # 方案2: 如果已知列没有内容，从后往前找有内容的列
                    if not content_found:
                        for col_idx in range(len(row.cells) - 1, 0, -1):
                            cell_text = row.cells[col_idx].text.strip()
                            if cell_text and len(cell_text) > 3:
                                lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
                                for line in lines:
                                    # v18.8: 已映射的条款优先识别，跳过常规验证
                                    if is_mapped_clause(line) or self._is_valid_clause_line(line):
                                        table_clauses.append(line)
                                        content_found = True
                                if content_found:
                                    break

                else:
                    # 其他行正常处理
                    row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_lines.append(row_text)

        # 如果从表格中提取到条款，优先使用这些条款
        if table_clauses:
            logger.info(f"从表格条款列提取到 {len(table_clauses)} 个条款")
            clauses = [ClauseItem(title=t, content="", original_title=t) for t in table_clauses]
            return clauses, True  # 纯标题模式

        # 如果没有提取到条款，使用原来的逻辑
        # 构建带格式信息的行列表: [(text, is_heading), ...]
        non_empty_lines_with_info = []

        # 先添加段落内容（保留Heading信息）
        non_empty_paragraphs = [(line, i in heading_lines) for i, line in enumerate(all_lines) if line]

        # 如果表格有内容且段落基本为空，优先使用表格内容
        if table_lines and len(non_empty_paragraphs) < len(table_lines):
            logger.info(f"检测到表格内容: {len(table_lines)} 行，优先使用表格")
            non_empty_lines_with_info = [(line, False) for line in table_lines if line]
        elif table_lines:
            logger.info(f"合并段落({len(non_empty_paragraphs)})和表格({len(table_lines)})内容")
            non_empty_lines_with_info = non_empty_paragraphs + [(line, False) for line in table_lines if line]
        else:
            non_empty_lines_with_info = non_empty_paragraphs

        heading_count = sum(1 for _, is_h in non_empty_lines_with_info if is_h)
        logger.info(f"非空行数: {len(non_empty_lines_with_info)}, Heading行数: {heading_count}")

        # 3. 基于标题识别进行分割（不再依赖空行）
        # v18.4: 使用 Heading 样式作为条款标题的强识别信号
        # v18.5: 已映射的条款名称优先识别为标题

        # v18.5: 获取已映射的客户条款名称（用于优先识别）
        mapped_client_names = set()
        try:
            if HAS_MAPPING_MANAGER:
                mapping_mgr = get_mapping_manager()
                if mapping_mgr:
                    for mapping in mapping_mgr.get_all_mappings():
                        if mapping.client_name:
                            mapped_client_names.add(mapping.client_name.strip())
                            # 也添加去除编号后的名称
                            cleaned = re.sub(r'^\d+[\.\s、]+', '', mapping.client_name).strip()
                            if cleaned:
                                mapped_client_names.add(cleaned)
            if mapped_client_names:
                logger.info(f"已加载 {len(mapped_client_names)} 个已映射条款名称用于优先识别")
        except Exception as e:
            logger.warning(f"获取映射条款名称失败: {e}")

        clauses = []
        current_title = None
        current_content = []

        for line, is_heading in non_empty_lines_with_info:
            # 判断是否是条款标题：
            # 1. is_likely_title 返回 True，或者
            # 2. 是 Heading 样式且不是明显的子编号内容，或者
            # 3. v18.5: 匹配已映射的条款名称
            is_title = self.is_likely_title(line)

            # v18.5: 检查是否在排除列表中（用于后续的 Heading/映射识别）
            is_excluded = False
            if HAS_MAPPING_MANAGER:
                excluded_titles = ClauseMatcherLogic._load_excluded_titles()
                if excluded_titles:
                    line_cleaned_for_exclude = ClauseMatcherLogic._remove_leading_number(line)
                    if line_cleaned_for_exclude.upper() in excluded_titles:
                        is_excluded = True
                        logger.debug(f"排除列表跳过: {line[:50]}")

            # v18.5: 已映射的条款名称优先识别为标题（但排除列表优先）
            if not is_title and not is_excluded and mapped_client_names:
                # 精确匹配
                if line in mapped_client_names:
                    is_title = True
                    logger.debug(f"已映射条款识别为标题: {line[:50]}")
                else:
                    # 去除编号后匹配
                    line_cleaned = re.sub(r'^\d+[\.\s、]+', '', line).strip()
                    if line_cleaned and line_cleaned in mapped_client_names:
                        is_title = True
                        logger.debug(f"已映射条款识别为标题(去编号): {line[:50]}")

            # v18.4: Heading 样式的段落优先识别为标题（但排除列表优先）
            if is_heading and not is_title and not is_excluded:
                # Heading 样式，但 is_likely_title 返回 False
                # 检查是否是子编号内容（如 "1.REINSTATEMENT VALUE CLAUSE"）
                # 子编号格式通常以 "数字.大写" 紧密连接，没有空格
                if not re.match(r'^\d+\.[A-Z]', line):
                    is_title = True
                    logger.debug(f"Heading样式识别为标题: {line[:50]}")

            if is_title:
                # 保存前一个条款
                if current_title is not None:
                    clauses.append(ClauseItem(
                        title=current_title,
                        content="\n".join(current_content),
                        original_title=current_title
                    ))
                # 开始新条款
                current_title = line
                current_content = []
            else:
                # 内容行
                if current_title is not None:
                    current_content.append(line)
                # v18.4修复: 在第一个标题之前的内容直接跳过，不再作为独立条款
                # 这避免了excluded_titles排除标题后，前置内容变成大量"条款"的问题

        # 保存最后一个条款
        if current_title is not None:
            clauses.append(ClauseItem(
                title=current_title,
                content="\n".join(current_content),
                original_title=current_title
            ))

        is_title_only = all(not c.content for c in clauses)
        logger.info(f"解析完成: {len(clauses)} 条款, 纯标题模式: {is_title_only}")

        return clauses, is_title_only

    def _parse_docx_precise_mode(self, doc) -> Tuple[List[ClauseItem], bool]:
        """
        v18.9: 精准识别模式 - 只提取蓝色字体的文字作为条款

        用户将条款内容标记为蓝色，以便在干扰项较多的文档中精准提取。
        每个蓝色文本块被视为一个独立的条款标题。
        """
        logger.info("使用精准识别模式（仅蓝色文字）")

        blue_clauses = []

        # 1. 从段落中提取蓝色文字
        for para in doc.paragraphs:
            blue_text = self._extract_blue_text_from_paragraph(para)
            if blue_text:
                # 按换行分割，每行可能是一个条款
                lines = [line.strip() for line in blue_text.split('\n') if line.strip()]
                for line in lines:
                    # 基本验证：长度合理
                    if 3 <= len(line) <= 300:
                        blue_clauses.append(line)

        # 2. 从表格中提取蓝色文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    blue_text = self._extract_blue_text_from_cell(cell)
                    if blue_text:
                        lines = [line.strip() for line in blue_text.split('\n') if line.strip()]
                        for line in lines:
                            if 3 <= len(line) <= 300:
                                # 避免重复（表格单元格可能合并导致重复读取）
                                if line not in blue_clauses:
                                    blue_clauses.append(line)

        # 去重并保持顺序
        seen = set()
        unique_clauses = []
        for clause in blue_clauses:
            if clause not in seen:
                seen.add(clause)
                unique_clauses.append(clause)

        logger.info(f"精准模式提取到 {len(unique_clauses)} 条蓝色文字条款")

        if not unique_clauses:
            logger.warning("未找到蓝色文字，请确认文档中已将条款标记为蓝色")

        # 转换为ClauseItem列表
        clauses = [
            ClauseItem(title=title, content="", original_title=title)
            for title in unique_clauses
        ]

        # 精准模式下始终是纯标题模式
        return clauses, True


# ==========================================
# 条款库加载器
# ==========================================
class LibraryLoader:
    """条款库加载器 - 支持自动列名识别和多Sheet选择"""

    @staticmethod
    def get_sheet_names(excel_path: str) -> List[str]:
        """
        获取Excel文件中所有Sheet名称
        """
        try:
            xl = pd.ExcelFile(excel_path)
            return xl.sheet_names
        except Exception as e:
            logger.warning(f"读取Sheet列表失败: {e}")
            return []

    @staticmethod
    def _extract_rich_text(cell) -> str:
        """v18.15: 从Excel单元格提取富文本，保留加粗格式

        使用 <b>...</b> 标记加粗文本
        """
        if cell.value is None:
            return ''

        # 检查是否有富文本
        try:
            from openpyxl.cell.rich_text import CellRichText
            if isinstance(cell.value, CellRichText):
                result = []
                for block in cell.value:
                    if hasattr(block, 'font') and block.font and block.font.b:
                        # 加粗文本
                        result.append(f'<b>{block.text}</b>')
                    elif hasattr(block, 'text'):
                        result.append(block.text)
                    else:
                        result.append(str(block))
                return ''.join(result)
        except (ImportError, AttributeError):
            pass

        return str(cell.value) if cell.value else ''

    @staticmethod
    def load_excel(excel_path: str, header_row: int = None, sheet_name: str = None) -> List[Dict]:
        """
        加载Excel条款库 - v18.15: 保留加粗格式
        自动识别列名和表头行

        Args:
            excel_path: Excel文件路径
            header_row: 表头行索引（自动检测时为None）
            sheet_name: Sheet名称（None时使用第一个Sheet）
        """
        logger.info(f"加载条款库: {excel_path}, Sheet: {sheet_name or '默认'}")

        try:
            # 使用 openpyxl 直接读取以保留富文本格式
            wb = openpyxl.load_workbook(excel_path, rich_text=True)
            ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active

            # 获取所有行数据
            rows = list(ws.iter_rows())
            if not rows:
                wb.close()
                return []

            # 自动检测表头行
            if header_row is None:
                header_row = 0
                for i in range(min(3, len(rows))):
                    row_values = [str(cell.value).lower() if cell.value else '' for cell in rows[i]]
                    if any('条款' in v or 'name' in v or '名称' in v for v in row_values):
                        header_row = i
                        break
                logger.info(f"自动检测表头行: {header_row}")

            # 获取表头
            header_cells = rows[header_row]
            columns = [str(cell.value).strip() if cell.value else f'Col{i}' for i, cell in enumerate(header_cells)]

        except FileNotFoundError:
            raise ValueError(f"文件不存在: {excel_path}")
        except Exception as e:
            raise ValueError(f"Excel读取失败: {e}")

        # 自动识别列名
        name_col_idx = None
        content_col_idx = None
        reg_col_idx = None

        for i, col in enumerate(columns):
            col_lower = col.lower()
            if name_col_idx is None and ('条款名称' in col or '名称' in col or 'name' in col_lower):
                name_col_idx = i
            elif content_col_idx is None and ('条款内容' in col or '内容' in col or 'content' in col_lower):
                content_col_idx = i
            elif reg_col_idx is None and ('注册号' in col or '产品' in col or 'reg' in col_lower):
                reg_col_idx = i

        # 回退到位置
        if name_col_idx is None and len(columns) > 0:
            name_col_idx = 0
        if content_col_idx is None and len(columns) > 2:
            content_col_idx = 2
        if reg_col_idx is None and len(columns) > 1:
            reg_col_idx = 1

        logger.info(f"列索引识别: 名称={name_col_idx}, 内容={content_col_idx}, 注册号={reg_col_idx}")

        # 构建数据（从表头下一行开始）
        lib_data = []
        for row in rows[header_row + 1:]:
            if name_col_idx is not None and name_col_idx < len(row):
                name = LibraryLoader._extract_rich_text(row[name_col_idx])
            else:
                name = ''

            if not name.strip():
                continue

            content = ''
            if content_col_idx is not None and content_col_idx < len(row):
                content = LibraryLoader._extract_rich_text(row[content_col_idx])

            reg_no = ''
            if reg_col_idx is not None and reg_col_idx < len(row):
                reg_no = LibraryLoader._extract_rich_text(row[reg_col_idx])

            lib_data.append({
                '条款名称': name,
                '条款内容': content,
                '产品注册号': reg_no,
            })

        wb.close()
        logger.info(f"加载完成: {len(lib_data)} 条有效记录")
        return lib_data


# ==========================================
# Excel样式器
# ==========================================
class ExcelStyler:
    """Excel样式应用器"""

    FILLS = {
        'green': PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
        'yellow': PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        'red': PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
        'blue': PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid"),
        'header': PatternFill(start_color="667eea", end_color="667eea", fill_type="solid"),
    }

    BORDER = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC')
    )

    # v18.15: 新格式列宽（支持3组匹配结果 + 约定的限额）
    # A=序号, B=客户条款(原), C=客户条款(译), D=客户原始内容, E=约定的限额
    # F-J=匹配1, K-O=匹配2, P-T=匹配3
    WIDTHS = {
        'A': 6, 'B': 35, 'C': 30, 'D': 45, 'E': 35,  # E=约定的限额
        # 匹配1
        'F': 40, 'G': 25, 'H': 50, 'I': 10, 'J': 12,
        # 匹配2
        'K': 40, 'L': 25, 'M': 50, 'N': 10, 'O': 12,
        # 匹配3
        'P': 40, 'Q': 25, 'R': 50, 'S': 10, 'T': 12,
    }

    # v18.15: 内容列索引（需要处理富文本的列）
    # D=4=客户原始内容, H=8=匹配1内容, M=13=匹配2内容, R=18=匹配3内容
    CONTENT_COLS = {4, 8, 13, 18}

    @staticmethod
    def _convert_to_rich_text(text: str):
        """v18.9: 将含<b>标记的文本转为CellRichText

        优化：将换行符合并到相邻文本块中，避免Excel不渲染单独的换行块
        """
        if not text or '<b>' not in str(text):
            return text

        try:
            from openpyxl.cell.rich_text import CellRichText, TextBlock
            from openpyxl.cell.text import InlineFont

            text = str(text)
            rich_text = CellRichText()
            pattern = re.compile(r'<b>(.*?)</b>', re.DOTALL)
            last_end = 0
            pending_whitespace = ''  # 待处理的空白/换行

            for match in pattern.finditer(text):
                # 处理当前匹配之前的非加粗部分
                if match.start() > last_end:
                    normal_text = text[last_end:match.start()]
                    if normal_text:
                        # 如果只是空白/换行，先保存待后续处理
                        if normal_text.strip() == '':
                            pending_whitespace = normal_text
                        else:
                            # 有实际内容的非加粗文本
                            rich_text.append(pending_whitespace + normal_text)
                            pending_whitespace = ''

                # 添加加粗部分（包含前置的换行）
                bold_text = match.group(1)
                if bold_text:
                    full_bold = pending_whitespace + bold_text
                    rich_text.append(TextBlock(InlineFont(b=True), full_bold))
                    pending_whitespace = ''

                last_end = match.end()

            # 添加最后的非加粗部分
            if last_end < len(text):
                remaining = text[last_end:]
                if remaining:
                    rich_text.append(pending_whitespace + remaining)
                    pending_whitespace = ''

            # 处理末尾剩余的空白
            if pending_whitespace:
                rich_text.append(pending_whitespace)

            return rich_text if rich_text else text

        except ImportError:
            # 不支持富文本时，返回去除标记的纯文本
            return re.sub(r'</?b>', '', text)

    @classmethod
    def apply_styles(cls, output_path: str):
        """应用Excel样式"""
        wb = openpyxl.load_workbook(output_path)
        ws = wb.active

        # 表头
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF", size=11)
            cell.fill = cls.FILLS['header']
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = cls.BORDER

        # 列宽
        for col, width in cls.WIDTHS.items():
            ws.column_dimensions[col].width = width

        # 数据行
        # v18.15: 新格式匹配度和匹配级别列索引（增加了E列约定的限额）
        # 匹配1: I(9)=匹配度, J(10)=级别
        # 匹配2: N(14)=匹配度, O(15)=级别
        # 匹配3: S(19)=匹配度, T(20)=级别
        score_cols = {9, 14, 19}  # 匹配度列索引
        level_cols = {10, 15, 20}  # 匹配级别列索引

        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical='top')
                cell.border = cls.BORDER

                # v17.1: 匹配度着色（支持3组）
                if cell.col_idx in score_cols:
                    try:
                        val = float(cell.value) if cell.value else 0
                        if val >= 0.8:
                            cell.fill = cls.FILLS['green']
                        elif val >= 0.5:
                            cell.fill = cls.FILLS['yellow']
                        elif val > 0:
                            cell.fill = cls.FILLS['red']
                    except (ValueError, TypeError):
                        pass

                # v17.1: 匹配级别着色（支持3组）
                if cell.col_idx in level_cols:
                    val = str(cell.value) if cell.value else ""
                    if "精确" in val:
                        cell.fill = cls.FILLS['green']
                    elif "语义" in val:
                        cell.fill = cls.FILLS['blue']
                    elif "关键词" in val:
                        cell.fill = cls.FILLS['yellow']

                # v18.15: 内容列转换为富文本（保留加粗格式）
                if cell.col_idx in cls.CONTENT_COLS and cell.value:
                    rich_value = cls._convert_to_rich_text(cell.value)
                    if rich_value != cell.value:
                        cell.value = rich_value

        # 冻结首行
        ws.freeze_panes = 'A2'

        wb.save(output_path)
        logger.info(f"Excel样式已应用: {output_path}")


# ==========================================
# 工作线程
# ==========================================
class MatchWorker(QThread):
    """单文件匹配工作线程"""
    log_signal = pyqtSignal(str, str)
    progress_signal = pyqtSignal(int, int)
    finished_signal = pyqtSignal(bool, str)

    def __init__(self, doc_path: str, excel_path: str, output_path: str, sheet_name: str = None,
                 match_mode: str = "auto", precise_mode: bool = False):
        super().__init__()
        self.doc_path = doc_path
        self.excel_path = excel_path
        self.output_path = output_path
        self.sheet_name = sheet_name  # 指定的Sheet名称
        self.match_mode = match_mode  # v18.3: 匹配模式 (auto/title/content)
        self.precise_mode = precise_mode  # v18.9: 精准识别模式（仅蓝色文字）
        self._cancelled = False  # v18.4: 取消标志

    def cancel(self):
        """v18.4: 取消比对"""
        self._cancelled = True

    def is_cancelled(self) -> bool:
        """v18.4: 检查是否已取消"""
        return self._cancelled

    def run(self):
        try:
            logic = ClauseMatcherLogic()

            # 状态信息
            self.log_signal.emit(f"📊 配置: 外部={logic._use_external_config}", "info")

            # v18.9: 精准识别模式提示
            if self.precise_mode:
                self.log_signal.emit("🎯 精准识别模式: 仅提取蓝色文字", "info")

            # 解析文档
            self.log_signal.emit("⏳ 正在解析文档...", "info")
            clauses, auto_detected_mode = logic.parse_docx(self.doc_path, precise_mode=self.precise_mode)

            # v18.3: 根据用户选择的模式决定 is_title_only
            if self.match_mode == "auto":
                is_title_only = auto_detected_mode
                mode_str = "自动检测→纯标题模式" if is_title_only else "自动检测→完整内容模式"
            elif self.match_mode == "title":
                is_title_only = True
                mode_str = "手动指定→纯标题模式"
            else:  # content
                is_title_only = False
                mode_str = "手动指定→完整内容模式"

            self.log_signal.emit(f"📖 [{mode_str}] 提取到 {len(clauses)} 条", "success")

            # 加载条款库
            sheet_info = f" [{self.sheet_name}]" if self.sheet_name else ""
            self.log_signal.emit(f"📚 加载条款库{sheet_info}...", "info")
            lib_data = LibraryLoader.load_excel(self.excel_path, sheet_name=self.sheet_name)
            self.log_signal.emit(f"✓ 条款库 {len(lib_data)} 条", "success")

            # v19.0: 设置险种上下文
            logic._current_category = logic.detect_category_from_sheet(self.sheet_name)
            if logic._current_category:
                self.log_signal.emit(f"🏷️ 检测到险种类别: {logic._current_category}", "info")

            # 构建索引
            self.log_signal.emit("🔧 构建索引...", "info")
            index = logic.build_index(lib_data)
            self.log_signal.emit(f"✓ 索引完成", "success")

            # 开始匹配 (v17.1 多结果匹配)
            self.log_signal.emit("🧠 开始智能匹配（v18.8 多结果模式）...", "info")
            results = []
            stats = {'exact': 0, 'semantic': 0, 'keyword': 0, 'fuzzy': 0, 'none': 0}

            for idx, clause in enumerate(clauses, 1):
                # v18.4: 检查取消
                if self._cancelled:
                    self.log_signal.emit("⛔ 用户取消了比对操作", "warning")
                    self.finished_signal.emit(False, "用户取消")
                    return

                self.progress_signal.emit(idx, len(clauses))

                # 翻译
                original_title = clause.title
                translated_title, was_translated = logic.translate_title(clause.title)
                if was_translated:
                    clause.title = translated_title
                    clause.original_title = original_title

                # 检查用户自定义映射
                user_library_name = None
                if HAS_MAPPING_MANAGER:
                    mapping_mgr = get_mapping_manager()
                    # 按原标题或翻译后标题查找
                    user_library_name = mapping_mgr.get_library_name(original_title)
                    if not user_library_name and was_translated:
                        user_library_name = mapping_mgr.get_library_name(translated_title)

                # v17.1: 根据是否有用户映射决定匹配策略
                # v18.5: 使用 create_user_mapping_result 方法减少重复代码
                match_results = []
                if user_library_name:
                    # 有用户映射，只返回映射的那一条
                    lib_entry = logic.find_library_entry_by_name(user_library_name, index)
                    mapped_result = logic.create_user_mapping_result(lib_entry, user_library_name)
                    match_results = [mapped_result]
                else:
                    # 无用户映射，使用多结果匹配（最多3条）
                    match_results = logic.match_clause_multiple(clause, index, is_title_only, max_results=3)

                # 统计使用第一个匹配结果
                primary_match = match_results[0] if match_results else MatchResult()
                if primary_match.match_level == MatchLevel.EXACT:
                    stats['exact'] += 1
                elif primary_match.match_level == MatchLevel.SEMANTIC:
                    stats['semantic'] += 1
                elif primary_match.match_level == MatchLevel.KEYWORD:
                    stats['keyword'] += 1
                elif primary_match.match_level == MatchLevel.FUZZY:
                    stats['fuzzy'] += 1
                else:
                    stats['none'] += 1

                # v18.15: 提取限额信息
                limit_info = ''
                # 优先从客户条款原名提取
                _, limit_info = logic.extract_limit_info(original_title)
                # 如果客户条款没有，再从匹配结果提取
                if not limit_info and match_results and match_results[0].matched_name:
                    _, limit_info = logic.extract_limit_info(match_results[0].matched_name)

                # v17.1: 构建多结果行
                row = {
                    ExcelColumns.SEQ: idx,
                    ExcelColumns.CLIENT_ORIG: original_title,
                    ExcelColumns.CLIENT_TRANS: translated_title if was_translated else "",
                    ExcelColumns.CLIENT_CONTENT: clause.content[:500] if clause.content else "",
                    ExcelColumns.LIMIT_INFO: limit_info,  # v18.15: 约定的限额
                }

                # 填充最多3条匹配结果
                for match_num in range(1, 4):
                    if match_num <= len(match_results):
                        mr = match_results[match_num - 1]
                        # v18.15: 显示时去掉限额后缀
                        display_name, _ = logic.extract_limit_info(mr.matched_name or "")
                        row[f'匹配{match_num}_条款名称'] = display_name
                        row[f'匹配{match_num}_注册号'] = logic.clean_reg_number(mr.matched_reg)
                        row[f'匹配{match_num}_条款内容'] = mr.matched_content[:500] if mr.matched_content else ""
                        row[f'匹配{match_num}_匹配度'] = round(mr.score, 3)
                        row[f'匹配{match_num}_匹配级别'] = mr.match_level.value
                    else:
                        row[f'匹配{match_num}_条款名称'] = ""
                        row[f'匹配{match_num}_注册号'] = ""
                        row[f'匹配{match_num}_条款内容'] = ""
                        row[f'匹配{match_num}_匹配度'] = ""
                        row[f'匹配{match_num}_匹配级别'] = ""

                results.append(row)

            # 保存结果
            df_res = pd.DataFrame(results)
            df_res.to_excel(self.output_path, index=False)
            ExcelStyler.apply_styles(self.output_path)

            # 输出统计
            self.log_signal.emit(f"📊 匹配统计:", "info")
            self.log_signal.emit(f"   精确匹配: {stats['exact']}", "success")
            self.log_signal.emit(f"   语义匹配: {stats['semantic']}", "success")
            self.log_signal.emit(f"   关键词匹配: {stats['keyword']}", "info")
            self.log_signal.emit(f"   模糊匹配: {stats['fuzzy']}", "warning")
            self.log_signal.emit(f"   无匹配: {stats['none']}", "error")

            self.log_signal.emit(f"🎉 完成！", "success")
            self.log_signal.emit(f"💡 提示: 报告中每个客户条款最多显示3条匹配结果供您选择", "info")
            self.finished_signal.emit(True, self.output_path)

        except Exception as e:
            logger.exception("匹配过程出错")
            self.log_signal.emit(f"❌ 错误: {str(e)}", "error")
            self.finished_signal.emit(False, str(e))


class BatchMatchWorker(QThread):
    """批量匹配工作线程"""
    log_signal = pyqtSignal(str, str)
    progress_signal = pyqtSignal(int, int)
    batch_progress_signal = pyqtSignal(int, int, str)  # 当前文件, 总数, 文件名
    finished_signal = pyqtSignal(bool, str, int, int)  # 成功, 消息, 成功数, 总数

    def __init__(self, doc_paths: List[str], excel_path: str, output_dir: str, sheet_name: str = None,
                 match_mode: str = "auto", precise_mode: bool = False):
        super().__init__()
        self.doc_paths = doc_paths
        self.excel_path = excel_path
        self.output_dir = output_dir
        self.sheet_name = sheet_name  # 指定的Sheet名称
        self.match_mode = match_mode  # v18.3: 匹配模式 (auto/title/content)
        self.precise_mode = precise_mode  # v18.9: 精准识别模式（仅蓝色文字）
        self._cancelled = False  # v18.4: 取消标志

    def cancel(self):
        """v18.4: 取消批量处理"""
        self._cancelled = True

    def run(self):
        try:
            logic = ClauseMatcherLogic()

            # 加载条款库（只需一次）
            sheet_info = f" [{self.sheet_name}]" if self.sheet_name else ""
            self.log_signal.emit(f"📚 加载条款库{sheet_info}...", "info")
            lib_data = LibraryLoader.load_excel(self.excel_path, sheet_name=self.sheet_name)
            self.log_signal.emit(f"✓ 条款库 {len(lib_data)} 条", "success")

            # v19.0: 设置险种上下文
            logic._current_category = logic.detect_category_from_sheet(self.sheet_name)
            if logic._current_category:
                self.log_signal.emit(f"🏷️ 检测到险种类别: {logic._current_category}", "info")

            # 构建索引（只需一次）
            self.log_signal.emit("🔧 构建索引...", "info")
            index = logic.build_index(lib_data)

            success_count = 0
            total = len(self.doc_paths)

            for file_idx, doc_path in enumerate(self.doc_paths, 1):
                # v18.4: 检查取消
                if self._cancelled:
                    self.log_signal.emit("⛔ 用户取消了批量处理", "warning")
                    self.finished_signal.emit(False, "用户取消", success_count, file_idx - 1)
                    return

                file_name = Path(doc_path).name
                self.batch_progress_signal.emit(file_idx, total, file_name)
                self.log_signal.emit(f"\n📄 [{file_idx}/{total}] {file_name}", "info")

                try:
                    # 解析文档
                    clauses, auto_detected_mode = logic.parse_docx(doc_path, precise_mode=self.precise_mode)

                    # v18.3: 根据用户选择的模式决定 is_title_only
                    if self.match_mode == "auto":
                        is_title_only = auto_detected_mode
                    elif self.match_mode == "title":
                        is_title_only = True
                    else:  # content
                        is_title_only = False

                    mode_hint = " (精准模式)" if self.precise_mode else ""
                    self.log_signal.emit(f"   提取 {len(clauses)} 条款{mode_hint}", "info")

                    # 匹配 (v17.1 多结果匹配)
                    results = []
                    mapping_mgr = get_mapping_manager() if HAS_MAPPING_MANAGER else None

                    for idx, clause in enumerate(clauses, 1):
                        original_title = clause.title
                        translated_title, was_translated = logic.translate_title(clause.title)
                        if was_translated:
                            clause.title = translated_title
                            clause.original_title = original_title

                        # 检查用户自定义映射
                        user_library_name = None
                        if mapping_mgr:
                            user_library_name = mapping_mgr.get_library_name(original_title)
                            if not user_library_name and was_translated:
                                user_library_name = mapping_mgr.get_library_name(translated_title)

                        # v17.1: 根据是否有用户映射决定匹配策略
                        # v18.5: 使用辅助方法减少重复代码
                        match_results = []
                        if user_library_name:
                            # 有用户映射，只返回映射的那一条
                            lib_entry = logic.find_library_entry_by_name(user_library_name, index)
                            mapped_result = logic.create_user_mapping_result(lib_entry, user_library_name)
                            match_results = [mapped_result]
                        else:
                            # 无用户映射，使用多结果匹配（最多3条）
                            match_results = logic.match_clause_multiple(clause, index, is_title_only, max_results=3)

                        # v17.1: 构建多结果行
                        row = {
                            ExcelColumns.SEQ: idx,
                            ExcelColumns.CLIENT_ORIG: original_title,
                            ExcelColumns.CLIENT_TRANS: translated_title if was_translated else "",
                            ExcelColumns.CLIENT_CONTENT: clause.content[:500] if clause.content else "",
                        }

                        # 填充最多3条匹配结果
                        for match_num in range(1, 4):
                            if match_num <= len(match_results):
                                mr = match_results[match_num - 1]
                                row[f'匹配{match_num}_条款名称'] = mr.matched_name or ""
                                row[f'匹配{match_num}_注册号'] = logic.clean_reg_number(mr.matched_reg)
                                row[f'匹配{match_num}_条款内容'] = mr.matched_content[:500] if mr.matched_content else ""
                                row[f'匹配{match_num}_匹配度'] = round(mr.score, 3)
                                row[f'匹配{match_num}_匹配级别'] = mr.match_level.value
                            else:
                                row[f'匹配{match_num}_条款名称'] = ""
                                row[f'匹配{match_num}_注册号'] = ""
                                row[f'匹配{match_num}_条款内容'] = ""
                                row[f'匹配{match_num}_匹配度'] = ""
                                row[f'匹配{match_num}_匹配级别'] = ""

                        results.append(row)

                    # 保存
                    output_name = f"报告_{Path(doc_path).stem}.xlsx"
                    output_path = Path(self.output_dir) / output_name
                    df_res = pd.DataFrame(results)
                    df_res.to_excel(output_path, index=False)
                    ExcelStyler.apply_styles(str(output_path))

                    self.log_signal.emit(f"   ✓ 已保存: {output_name}", "success")
                    success_count += 1

                except Exception as e:
                    self.log_signal.emit(f"   ✗ 失败: {e}", "error")

            self.log_signal.emit(f"\n🎉 批量处理完成: {success_count}/{total}", "success")
            self.finished_signal.emit(True, self.output_dir, success_count, total)

        except Exception as e:
            logger.exception("批量处理出错")
            self.log_signal.emit(f"❌ 错误: {str(e)}", "error")
            self.finished_signal.emit(False, str(e), 0, 0)


# ==========================================
# UI组件 - Anthropic 风格
# ==========================================
class AnthropicCard(QFrame):
    """Anthropic 风格卡片组件"""
    def __init__(self, parent=None, variant="default"):
        super().__init__(parent)
        # 根据变体选择背景色
        if variant == "mint":
            bg_color = AnthropicColors.BG_MINT
        elif variant == "lavender":
            bg_color = AnthropicColors.BG_LAVENDER
        else:
            bg_color = AnthropicColors.BG_CARD

        self.setStyleSheet(f"""
            AnthropicCard {{
                background: {bg_color};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 12px;
            }}
        """)
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(20)
        shadow.setColor(QColor(0, 0, 0, 25))
        shadow.setOffset(0, 4)
        self.setGraphicsEffect(shadow)


# 保留旧名称以兼容
GlassCard = AnthropicCard


class AddMappingDialog(QDialog):
    """添加映射对话框 - Anthropic 风格"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("添加术语映射")
        self.setMinimumWidth(400)
        self.setStyleSheet(f"""
            QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; font-size: 14px; }}
            QLineEdit {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 10px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QLineEdit:focus {{ border-color: {AnthropicColors.ACCENT}; }}
            QPushButton {{
                background: {AnthropicColors.BG_DARK}; color: {AnthropicColors.TEXT_LIGHT}; border: none;
                border-radius: 8px; padding: 10px 20px; font-weight: bold;
            }}
            QPushButton:hover {{ background: {AnthropicColors.ACCENT}; }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        form = QFormLayout()
        self.eng_input = QLineEdit()
        self.eng_input.setPlaceholderText("例如: reinstatement value")
        form.addRow("英文术语:", self.eng_input)

        self.chn_input = QLineEdit()
        self.chn_input.setPlaceholderText("例如: 重置价值条款")
        form.addRow("中文翻译:", self.chn_input)
        layout.addLayout(form)

        btn_layout = QHBoxLayout()
        cancel_btn = QPushButton("取消")
        cancel_btn.setStyleSheet(f"background: transparent; color: {AnthropicColors.TEXT_PRIMARY}; border: 1px solid {AnthropicColors.BORDER};")
        cancel_btn.clicked.connect(self.reject)
        save_btn = QPushButton("保存")
        save_btn.clicked.connect(self.accept)
        btn_layout.addWidget(cancel_btn)
        btn_layout.addWidget(save_btn)
        layout.addLayout(btn_layout)

    def get_mapping(self) -> Tuple[str, str]:
        return self.eng_input.text().strip(), self.chn_input.text().strip()


# ==========================================
# 打赏对话框
# ==========================================
class DonateDialog(QDialog):
    """支持作者对话框 - 微信和支付宝双二维码"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('💝 支持作者')
        self.setMinimumSize(520, 520)
        self._setup_ui()

    def _get_qr_image_path(self, name: str) -> str:
        """获取二维码图片路径（支持PyInstaller打包）"""
        possible_paths = []

        # PyInstaller 打包后的路径
        if getattr(sys, 'frozen', False):
            bundle_dir = getattr(sys, '_MEIPASS', os.path.dirname(sys.executable))
            possible_paths.append(os.path.join(bundle_dir, name))
            # exe 同级目录
            possible_paths.append(os.path.join(os.path.dirname(sys.executable), name))

        # 常规开发路径
        possible_paths.extend([
            os.path.join(os.path.dirname(os.path.abspath(__file__)), name),
            os.path.join(os.getcwd(), name),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Resources', name),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets', name),
        ])

        for path in possible_paths:
            if os.path.exists(path):
                return path
        return ""

    def _setup_ui(self):
        self.setStyleSheet(f"""
            QDialog {{
                background: {AnthropicColors.BG_PRIMARY};
            }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; }}
            QPushButton {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT}; border: none; border-radius: 12px;
                padding: 12px 30px; font-weight: bold; font-size: 14px;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.ACCENT};
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(30, 25, 30, 25)

        # 标题区域 - 带动画效果
        title = QLabel('✨ 感谢您的支持！✨')
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet(f'''
            font-size: 22px; font-weight: bold;
            color: {AnthropicColors.ACCENT};
            padding: 5px;
        ''')
        layout.addWidget(title)

        desc = QLabel('如果这个工具对您有帮助，欢迎请作者喝杯咖啡 ☕')
        desc.setAlignment(Qt.AlignCenter)
        desc.setStyleSheet(f'color: {AnthropicColors.TEXT_SECONDARY}; font-size: 13px;')
        layout.addWidget(desc)

        # 打赏区域
        donate_container = QHBoxLayout()
        donate_container.setSpacing(25)

        # 微信支付
        wechat_widget = QWidget()
        wechat_layout = QVBoxLayout(wechat_widget)
        wechat_layout.setAlignment(Qt.AlignCenter)
        wechat_layout.setSpacing(8)

        wechat_label = QLabel('💚 微信支付')
        wechat_label.setAlignment(Qt.AlignCenter)
        wechat_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #07C160;')
        wechat_layout.addWidget(wechat_label)

        wechat_qr_label = QLabel()
        wechat_qr_label.setFixedSize(160, 160)
        wechat_qr_label.setAlignment(Qt.AlignCenter)

        from PyQt5.QtGui import QPixmap
        wx_path = self._get_qr_image_path('wx.jpg')
        if wx_path:
            pixmap = QPixmap(wx_path)
            if not pixmap.isNull():
                wechat_qr_label.setPixmap(pixmap.scaled(154, 154, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                wechat_qr_label.setStyleSheet('''
                    background-color: white; border-radius: 12px;
                    border: 3px solid #07C160; padding: 3px;
                ''')
            else:
                wechat_qr_label.setText('💚\n微信扫码')
                wechat_qr_label.setStyleSheet(f'''
                    font-size: 20px; background-color: {AnthropicColors.BG_CARD}; border-radius: 12px;
                    border: 3px solid #07C160; color: #07C160;
                ''')
        else:
            wechat_qr_label.setText('💚\n微信扫码')
            wechat_qr_label.setStyleSheet('''
                font-size: 20px; background-color: rgba(255,255,255,0.1); border-radius: 12px;
                border: 3px solid #07C160; color: #07C160;
            ''')
        wechat_layout.addWidget(wechat_qr_label, alignment=Qt.AlignCenter)

        wechat_hint = QLabel('微信扫一扫')
        wechat_hint.setAlignment(Qt.AlignCenter)
        wechat_hint.setStyleSheet('font-size: 12px; color: #07C160;')
        wechat_layout.addWidget(wechat_hint)
        donate_container.addWidget(wechat_widget)

        # 分隔线
        separator = QFrame()
        separator.setFrameShape(QFrame.VLine)
        separator.setStyleSheet(f'background-color: {AnthropicColors.BORDER};')
        donate_container.addWidget(separator)

        # 支付宝
        alipay_widget = QWidget()
        alipay_layout = QVBoxLayout(alipay_widget)
        alipay_layout.setAlignment(Qt.AlignCenter)
        alipay_layout.setSpacing(8)

        alipay_label = QLabel('💙 支付宝')
        alipay_label.setAlignment(Qt.AlignCenter)
        alipay_label.setStyleSheet('font-weight: bold; font-size: 14px; color: #1677FF;')
        alipay_layout.addWidget(alipay_label)

        alipay_qr_label = QLabel()
        alipay_qr_label.setFixedSize(160, 160)
        alipay_qr_label.setAlignment(Qt.AlignCenter)

        zfb_path = self._get_qr_image_path('zfb.jpg')
        if zfb_path:
            pixmap = QPixmap(zfb_path)
            if not pixmap.isNull():
                alipay_qr_label.setPixmap(pixmap.scaled(154, 154, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                alipay_qr_label.setStyleSheet('''
                    background-color: white; border-radius: 12px;
                    border: 3px solid #1677FF; padding: 3px;
                ''')
            else:
                alipay_qr_label.setText('💙\n支付宝扫码')
                alipay_qr_label.setStyleSheet(f'''
                    font-size: 20px; background-color: {AnthropicColors.BG_CARD}; border-radius: 12px;
                    border: 3px solid #1677FF; color: #1677FF;
                ''')
        else:
            alipay_qr_label.setText('💙\n支付宝扫码')
            alipay_qr_label.setStyleSheet('''
                font-size: 20px; background-color: rgba(255,255,255,0.1); border-radius: 12px;
                border: 3px solid #1677FF; color: #1677FF;
            ''')
        alipay_layout.addWidget(alipay_qr_label, alignment=Qt.AlignCenter)

        alipay_hint = QLabel('支付宝扫一扫')
        alipay_hint.setAlignment(Qt.AlignCenter)
        alipay_hint.setStyleSheet('font-size: 12px; color: #1677FF;')
        alipay_layout.addWidget(alipay_hint)
        donate_container.addWidget(alipay_widget)

        layout.addLayout(donate_container)

        # 感谢语
        thanks_label = QLabel('「大鑽戒基金會」へのご支援、誠にありがとうございます💎')
        thanks_label.setAlignment(Qt.AlignCenter)
        thanks_label.setStyleSheet(f'''
            font-size: 14px; font-weight: 500;
            color: {AnthropicColors.TEXT_PRIMARY}; padding: 15px 0 5px 0;
        ''')
        layout.addWidget(thanks_label)

        # 作者信息
        author_info = QLabel('Author: Dachi Yijin  |  智能条款比对工具')
        author_info.setAlignment(Qt.AlignCenter)
        author_info.setStyleSheet(f'color: {AnthropicColors.TEXT_SECONDARY}; font-size: 11px;')
        layout.addWidget(author_info)

        # 关闭按钮
        close_btn = QPushButton('关闭')
        close_btn.setFixedWidth(140)
        close_btn.setCursor(Qt.PointingHandCursor)
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn, alignment=Qt.AlignCenter)


class BatchSelectDialog(QDialog):
    """批量文件选择对话框 - Anthropic 风格"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("批量处理")
        self.setMinimumSize(500, 400)
        self.setStyleSheet(f"""
            QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; font-size: 14px; }}
            QListWidget {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item {{ padding: 10px 12px; border-radius: 6px; margin-bottom: 2px; }}
            QListWidget::item:hover {{ background: rgba(217, 119, 87, 0.08); }}
            QListWidget::item:selected {{ background: {AnthropicColors.BG_MINT}; color: {AnthropicColors.TEXT_PRIMARY}; }}
            QPushButton {{
                background: transparent;
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 10px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QPushButton:hover {{ background: {AnthropicColors.BG_CARD}; border-color: {AnthropicColors.ACCENT}; }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        layout.addWidget(QLabel("选择要批量处理的 Word 文件:"))

        self.file_list = QListWidget()
        layout.addWidget(self.file_list)

        btn_row = QHBoxLayout()
        add_btn = QPushButton("➕ 添加文件")
        add_btn.clicked.connect(self._add_files)
        clear_btn = QPushButton("🗑️ 清空")
        clear_btn.clicked.connect(self.file_list.clear)
        btn_row.addWidget(add_btn)
        btn_row.addWidget(clear_btn)
        layout.addLayout(btn_row)

        action_row = QHBoxLayout()
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        start_btn = QPushButton("开始批量处理")
        start_btn.setStyleSheet(f"background: {AnthropicColors.BG_DARK}; color: {AnthropicColors.TEXT_LIGHT};")
        start_btn.clicked.connect(self.accept)
        action_row.addWidget(cancel_btn)
        action_row.addWidget(start_btn)
        layout.addLayout(action_row)

        self.selected_files: List[str] = []

    def _add_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择Word文件", "", "Word Files (*.docx)")
        for f in files:
            if f not in self.selected_files:
                self.selected_files.append(f)
                self.file_list.addItem(Path(f).name)

    def get_files(self) -> List[str]:
        return self.selected_files


class ClauseQueryDialog(QDialog):
    """v17.1: 条款查询对话框 - 仅查询条款标题"""
    def __init__(self, parent=None, library_index=None, logic=None, mapping_mgr=None):
        super().__init__(parent)
        self.setWindowTitle("🔍 条款智能查询")
        self.setMinimumSize(600, 500)
        self.library_index = library_index
        self.logic = logic or ClauseMatcherLogic()
        self.mapping_mgr = mapping_mgr
        self._setup_ui()

    def _setup_ui(self):
        self.setStyleSheet(f"""
            QDialog {{ background: {AnthropicColors.BG_PRIMARY}; }}
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; font-size: 14px; }}
            QLineEdit, QTextEdit {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 10px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QLineEdit:focus, QTextEdit:focus {{ border-color: {AnthropicColors.ACCENT}; }}
            QPushButton {{
                background: transparent;
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 10px 20px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QPushButton:hover {{ background: {AnthropicColors.BG_CARD}; border-color: {AnthropicColors.ACCENT}; }}
            QListWidget {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item {{ padding: 10px 12px; border-radius: 6px; margin-bottom: 2px; }}
            QListWidget::item:hover {{ background: rgba(217, 119, 87, 0.08); }}
            QListWidget::item:selected {{ background: {AnthropicColors.BG_MINT}; color: {AnthropicColors.TEXT_PRIMARY}; }}
        """)

        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # 说明
        hint = QLabel("输入条款名称或关键词，系统将自动匹配最相近的条款库条款（仅匹配标题）")
        hint.setStyleSheet(f"color: {AnthropicColors.TEXT_MUTED}; font-size: 12px;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # 输入行
        input_row = QHBoxLayout()
        self.query_input = QLineEdit()
        self.query_input.setPlaceholderText("例如: 自动升值 或 REINSTATEMENT VALUE...")
        self.query_input.returnPressed.connect(self._do_search)
        self.search_btn = QPushButton("🔍 搜索")
        self.search_btn.setStyleSheet(f"background: {AnthropicColors.ACCENT}; color: {AnthropicColors.TEXT_LIGHT};")
        self.search_btn.clicked.connect(self._do_search)
        input_row.addWidget(self.query_input, 4)
        input_row.addWidget(self.search_btn, 1)
        layout.addLayout(input_row)

        # 结果列表
        layout.addWidget(QLabel("查询结果（最多5条）:"))
        self.result_list = QListWidget()
        self.result_list.itemDoubleClicked.connect(self._show_detail)
        layout.addWidget(self.result_list, 1)

        # 详情区
        layout.addWidget(QLabel("选中条款详情:"))
        self.detail_text = QTextEdit()
        self.detail_text.setReadOnly(True)
        self.detail_text.setMaximumHeight(120)
        layout.addWidget(self.detail_text)

        # 关闭按钮
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn, alignment=Qt.AlignRight)

        # 存储结果数据
        self._search_results = []

    def _do_search(self):
        """执行查询"""
        query = self.query_input.text().strip()
        if not query:
            return

        self.result_list.clear()
        self.detail_text.clear()
        self._search_results = []

        if not self.library_index or not self.library_index.data:
            self.result_list.addItem("⚠️ 请先选择条款库文件")
            return

        # 检查是否有用户映射
        if self.mapping_mgr:
            mapped_name = self.mapping_mgr.get_library_name(query)
            if mapped_name:
                # 有映射，直接返回映射的那一条
                lib_entry = self.logic.find_library_entry_by_name(mapped_name, self.library_index)
                if lib_entry:
                    self._search_results = [{
                        'name': lib_entry.get('条款名称', mapped_name),
                        'content': lib_entry.get('条款内容', ''),
                        'reg': lib_entry.get('产品注册号', lib_entry.get('注册号', '')),
                        'score': 1.0,
                        'matchType': 'mapping'
                    }]
                else:
                    self._search_results = [{
                        'name': mapped_name,
                        'content': '(用户映射条款，未在库中找到)',
                        'reg': '',
                        'score': 1.0,
                        'matchType': 'mapping'
                    }]
                self._display_results()
                return

        # 使用search_library_titles进行查询
        results = self.logic.search_library_titles(query, self.library_index, max_results=5)
        self._search_results = results
        self._display_results()

    def _display_results(self):
        """显示查询结果"""
        if not self._search_results:
            self.result_list.addItem("未找到匹配的条款")
            return

        for i, r in enumerate(self._search_results):
            match_type = r.get('matchType', '')
            if match_type == 'mapping':
                type_str = "[用户映射]"
            elif match_type == 'exact':
                type_str = "[精确匹配]"
            elif match_type == 'contain':
                type_str = "[包含匹配]"
            else:
                type_str = f"[模糊 {r.get('score', 0):.2f}]"

            self.result_list.addItem(f"{i+1}. {type_str} {r.get('name', '')}")

        # 自动选择第一条
        if self.result_list.count() > 0:
            self.result_list.setCurrentRow(0)
            self._show_detail(self.result_list.item(0))

    def _show_detail(self, item):
        """显示条款详情"""
        row = self.result_list.row(item)
        if 0 <= row < len(self._search_results):
            r = self._search_results[row]
            detail = f"【条款名称】{r.get('name', '')}\n\n"
            detail += f"【产品注册号】{r.get('reg', '无')}\n\n"
            detail += f"【条款内容】\n{r.get('content', '无内容')[:500]}..."
            self.detail_text.setText(detail)


# ==========================================
# 条款提取Tab - V18.0新增
# ==========================================
class ClauseExtractorTab(QWidget):
    """条款提取Tab - 支持文件夹分类和条款提取"""

    # 文件分类信号
    extraction_log = pyqtSignal(str, str)  # message, level

    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_files = []
        self.classified_files = {'fujia': [], 'feilv': [], 'zhu': []}
        self.doc_files = []  # .doc文件列表（需要转换）
        self.extracted_data = []
        self.categories = set()
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(8)
        layout.setContentsMargins(15, 10, 15, 10)

        # 滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setStyleSheet(get_anthropic_scrollbar_style())
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setSpacing(8)
        scroll_layout.setContentsMargins(0, 0, 0, 0)

        # 紧凑型统计面板（水平对齐）- 初始隐藏，有数据时显示
        self.stats_frame = QFrame()
        self.stats_frame.setMinimumHeight(40)
        self.stats_frame.setStyleSheet(f"""
            QFrame {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
            }}
        """)
        self.stats_frame.setVisible(False)  # 初始隐藏
        stats_layout = QHBoxLayout(self.stats_frame)
        stats_layout.setContentsMargins(20, 0, 20, 0)
        stats_layout.setSpacing(0)

        # 使用固定宽度的统计项确保对齐
        stat_style = "font-size: 14px; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;"

        self.stat_total_label = QLabel("待处理: 0")
        self.stat_total_label.setMinimumWidth(120)
        self.stat_total_label.setAlignment(Qt.AlignCenter)
        self.stat_total_label.setStyleSheet(f"color: {AnthropicColors.ACCENT}; {stat_style} font-weight: 600;")

        sep1 = QLabel("|")
        sep1.setFixedWidth(20)
        sep1.setAlignment(Qt.AlignCenter)
        sep1.setStyleSheet(f"color: {AnthropicColors.BORDER}; font-size: 14px;")

        self.stat_extracted_label = QLabel("已提取: 0")
        self.stat_extracted_label.setMinimumWidth(100)
        self.stat_extracted_label.setAlignment(Qt.AlignCenter)
        self.stat_extracted_label.setStyleSheet(f"color: {AnthropicColors.SUCCESS}; {stat_style} font-weight: 600;")

        sep2 = QLabel("|")
        sep2.setFixedWidth(20)
        sep2.setAlignment(Qt.AlignCenter)
        sep2.setStyleSheet(f"color: {AnthropicColors.BORDER}; font-size: 14px;")

        self.stat_categories_label = QLabel("分类数: 0")
        self.stat_categories_label.setMinimumWidth(100)
        self.stat_categories_label.setAlignment(Qt.AlignCenter)
        self.stat_categories_label.setStyleSheet(f"color: {AnthropicColors.INFO}; {stat_style} font-weight: 600;")

        sep3 = QLabel("|")
        sep3.setFixedWidth(20)
        sep3.setAlignment(Qt.AlignCenter)
        sep3.setStyleSheet(f"color: {AnthropicColors.BORDER}; font-size: 14px;")

        self.stat_skipped_label = QLabel("已跳过: 0")
        self.stat_skipped_label.setMinimumWidth(100)
        self.stat_skipped_label.setAlignment(Qt.AlignCenter)
        self.stat_skipped_label.setStyleSheet(f"color: {AnthropicColors.WARNING}; {stat_style} font-weight: 600;")

        stats_layout.addStretch()
        stats_layout.addWidget(self.stat_total_label)
        stats_layout.addWidget(sep1)
        stats_layout.addWidget(self.stat_extracted_label)
        stats_layout.addWidget(sep2)
        stats_layout.addWidget(self.stat_categories_label)
        stats_layout.addWidget(sep3)
        stats_layout.addWidget(self.stat_skipped_label)
        stats_layout.addStretch()

        scroll_layout.addWidget(self.stats_frame)

        # 文件选择卡片
        file_card = GlassCard()
        file_card_layout = QVBoxLayout(file_card)
        file_card_layout.setSpacing(12)

        card_title = QLabel("📂 选择条款文件或文件夹")
        card_title.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-weight: 600; font-size: 14px;")
        file_card_layout.addWidget(card_title)

        # 模式切换按钮
        mode_layout = QHBoxLayout()
        self.mode_files_btn = QPushButton("📄 选择文件")
        self.mode_folder_btn = QPushButton("📁 选择文件夹")

        mode_btn_style = f"""
            QPushButton {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 10px 20px;
                color: {AnthropicColors.TEXT_MUTED};
                font-weight: 500;
            }}
            QPushButton:hover {{
                border-color: {AnthropicColors.ACCENT};
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QPushButton:checked {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
                border-color: {AnthropicColors.BG_DARK};
            }}
        """
        self.mode_files_btn.setStyleSheet(mode_btn_style)
        self.mode_folder_btn.setStyleSheet(mode_btn_style)
        self.mode_files_btn.setCheckable(True)
        self.mode_folder_btn.setCheckable(True)
        self.mode_files_btn.setChecked(True)
        self.mode_files_btn.clicked.connect(lambda: self._switch_mode('files'))
        self.mode_folder_btn.clicked.connect(lambda: self._switch_mode('folder'))

        mode_layout.addWidget(self.mode_files_btn)
        mode_layout.addWidget(self.mode_folder_btn)
        mode_layout.addStretch()
        file_card_layout.addLayout(mode_layout)

        # 文件选择区域
        self.file_select_btn = QPushButton("点击选择文件 (.docx / .pdf)")
        self.file_select_btn.setMinimumHeight(60)
        self.file_select_btn.setCursor(Qt.PointingHandCursor)
        self.file_select_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 2px dashed {AnthropicColors.BORDER};
                border-radius: 12px;
                color: {AnthropicColors.TEXT_MUTED};
                font-size: 14px;
            }}
            QPushButton:hover {{
                border-color: {AnthropicColors.ACCENT};
                color: {AnthropicColors.TEXT_PRIMARY};
                background: rgba(217, 119, 87, 0.05);
            }}
        """)
        self.file_select_btn.clicked.connect(self._select_files)
        file_card_layout.addWidget(self.file_select_btn)

        # 文件列表
        self.file_list = QListWidget()
        self.file_list.setMaximumHeight(140)
        self.file_list.setStyleSheet(f"""
            QListWidget {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 10px;
                font-family: 'Cascadia Code', 'Consolas', monospace;
                font-size: 13px;
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item {{
                padding: 8px 12px;
                border-radius: 6px;
                color: {AnthropicColors.TEXT_PRIMARY};
                margin-bottom: 2px;
            }}
            QListWidget::item:hover {{
                background: rgba(217, 119, 87, 0.1);
            }}
            QListWidget::item:selected {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
            }}
        """)
        self.file_list.setVisible(False)
        file_card_layout.addWidget(self.file_list)

        # 分类预览区域（文件夹模式）
        self.classify_preview = QWidget()
        classify_layout = QHBoxLayout(self.classify_preview)
        classify_layout.setContentsMargins(0, 10, 0, 0)

        self.preview_fujia = self._create_classify_box("📗 附加条款", "0", "#d97757")
        self.preview_feilv = self._create_classify_box("📘 费率表", "0", "#6a9bcc")
        self.preview_zhu = self._create_classify_box("📙 主条款", "0", "#788c5d")

        classify_layout.addWidget(self.preview_fujia)
        classify_layout.addWidget(self.preview_feilv)
        classify_layout.addWidget(self.preview_zhu)
        self.classify_preview.setVisible(False)
        file_card_layout.addWidget(self.classify_preview)

        scroll_layout.addWidget(file_card)

        # 操作按钮行
        btn_layout = QHBoxLayout()

        self.extract_btn = QPushButton("🚀 开始提取")
        self.extract_btn.setMinimumHeight(40)
        self.extract_btn.setCursor(Qt.PointingHandCursor)
        self.extract_btn.setEnabled(False)
        self.extract_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
                border: none;
                border-radius: 8px;
                font-size: 15px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background: #2a2a28;
            }}
            QPushButton:disabled {{
                background: {AnthropicColors.BORDER};
                color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)
        self.extract_btn.clicked.connect(self._start_extraction)

        self.download_zip_btn = QPushButton("📦 进行分类ZIP打包")
        self.download_zip_btn.setMinimumHeight(40)
        self.download_zip_btn.setCursor(Qt.PointingHandCursor)
        self.download_zip_btn.setVisible(False)
        self.download_zip_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BG_DARK};
                border-radius: 8px;
                font-size: 14px;
                font-weight: 500;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
            }}
        """)
        self.download_zip_btn.clicked.connect(self._download_classified_zip)

        self.download_excel_btn = QPushButton("📊 下载Excel报告")
        self.download_excel_btn.setMinimumHeight(40)
        self.download_excel_btn.setCursor(Qt.PointingHandCursor)
        self.download_excel_btn.setVisible(False)
        self.download_excel_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                color: {AnthropicColors.SUCCESS};
                border: 1px solid {AnthropicColors.SUCCESS};
                border-radius: 8px;
                font-size: 14px;
                font-weight: 500;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.SUCCESS};
                color: white;
            }}
        """)
        self.download_excel_btn.clicked.connect(self._download_excel_report)

        self.clear_btn = QPushButton("🗑 清空")
        self.clear_btn.setMinimumHeight(40)
        self.clear_btn.setCursor(Qt.PointingHandCursor)
        self.clear_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                font-size: 14px;
            }}
            QPushButton:hover {{
                border-color: {AnthropicColors.ERROR};
                color: {AnthropicColors.ERROR};
            }}
        """)
        self.clear_btn.clicked.connect(self._clear_all)

        btn_layout.addWidget(self.extract_btn, 3)
        btn_layout.addWidget(self.download_zip_btn, 1)
        btn_layout.addWidget(self.download_excel_btn, 1)
        btn_layout.addWidget(self.clear_btn, 1)
        scroll_layout.addLayout(btn_layout)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.setStyleSheet(f"""
            QProgressBar {{ background: {AnthropicColors.BORDER}; border-radius: 3px; }}
            QProgressBar::chunk {{
                background: {AnthropicColors.ACCENT};
                border-radius: 3px;
            }}
        """)
        scroll_layout.addWidget(self.progress_bar)

        # 完成滚动区域设置
        scroll_area.setWidget(scroll_widget)
        layout.addWidget(scroll_area, 1)

        # 日志区域
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background: {AnthropicColors.BG_DARK};
                border: none;
                border-radius: 8px;
                color: {AnthropicColors.TEXT_LIGHT};
                padding: 15px;
                font-family: 'Consolas', 'Cascadia Code', monospace;
                font-size: 13px;
            }}
        """)
        layout.addWidget(self.log_text, 1)

        # 初始日志
        self._log("📊 条款提取工具已就绪", "info")
        self._log("   支持格式: .docx / .pdf", "info")
        self._log("   文件夹模式可自动分类：附加条款、费率表、主条款", "info")

    def _create_classify_box(self, title: str, count: str, color: str) -> QFrame:
        """创建分类预览框"""
        frame = QFrame()
        frame.setMinimumWidth(160)
        frame.setMinimumHeight(90)
        frame.setStyleSheet(f"""
            QFrame {{
                background: {AnthropicColors.BG_CARD};
                border: 2px solid {color};
                border-radius: 12px;
            }}
        """)
        layout = QVBoxLayout(frame)
        layout.setSpacing(8)
        layout.setContentsMargins(20, 15, 20, 15)

        title_label = QLabel(title)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet(f"""
            QLabel {{
                color: {AnthropicColors.TEXT_PRIMARY};
                background: transparent;
                border: none;
                font-size: 14px;
                font-weight: 500;
            }}
        """)

        count_label = QLabel(count)
        count_label.setAlignment(Qt.AlignCenter)
        count_label.setStyleSheet(f"""
            QLabel {{
                color: {color};
                background: transparent;
                border: none;
                font-size: 24px;
                font-weight: bold;
                font-family: 'Consolas', 'Cascadia Code', monospace;
            }}
        """)
        count_label.setObjectName("count")

        layout.addWidget(title_label)
        layout.addWidget(count_label)
        return frame

    def _switch_mode(self, mode: str):
        """切换文件/文件夹模式"""
        self.mode_files_btn.setChecked(mode == 'files')
        self.mode_folder_btn.setChecked(mode == 'folder')

        if mode == 'files':
            self.file_select_btn.setText("点击选择文件 (.docx / .pdf)")
        else:
            self.file_select_btn.setText("点击选择文件夹")

        self._clear_all()
        self._log(f"📋 切换到{'文件模式' if mode == 'files' else '文件夹模式'}", "info")

    def _select_files(self):
        """选择文件或文件夹"""
        if self.mode_files_btn.isChecked():
            files, _ = QFileDialog.getOpenFileNames(
                self, "选择条款文件", "",
                "文档文件 (*.docx *.pdf);;Word文档 (*.docx);;PDF文档 (*.pdf)"
            )
            if files:
                self._handle_files(files)
        else:
            # macOS 需要 ShowDirsOnly 选项才能正确选择文件夹
            folder = QFileDialog.getExistingDirectory(
                self, "选择文件夹", "",
                QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
            )
            if folder:
                self._handle_folder(folder)

    def _handle_files(self, file_paths: list):
        """处理选择的文件"""
        self.selected_files = []
        self.file_list.clear()

        for fp in file_paths:
            fname = os.path.basename(fp)
            ext = fname.split('.')[-1].lower()

            # 检查格式
            if ext not in ['docx', 'pdf']:
                self._log(f"⚠️ 跳过不支持的格式: {fname}", "warning")
                continue

            # 检查是否包含"附加"
            if '附加' not in fname:
                self._log(f"⚠️ 跳过不含「附加」的文件: {fname}", "warning")
                continue

            # 排除费率文件
            if '费率' in fname:
                self._log(f"⚠️ 跳过费率文件: {fname}", "warning")
                continue

            self.selected_files.append(fp)
            self.file_list.addItem(f"📄 {fname}")

        if self.selected_files:
            self.file_list.setVisible(True)
            self.extract_btn.setEnabled(True)
            self._update_stats()
            self._log(f"✓ 已选择 {len(self.selected_files)} 个文件", "success")

    def _handle_folder(self, folder_path: str):
        """处理文件夹 - 自动分类（支持多层子目录穿透）"""
        self.classified_files = {'fujia': [], 'feilv': [], 'zhu': []}
        self.selected_files = []
        self.doc_files = []  # 需要转换的.doc文件
        self.file_list.clear()

        # 使用os.walk递归遍历所有子目录
        for root, dirs, files in os.walk(folder_path):
            # 跳过隐藏目录
            dirs[:] = [d for d in dirs if not d.startswith('.')]

            for fname in files:
                if fname.startswith('.') or fname.startswith('~'):
                    continue

                ext = fname.split('.')[-1].lower()
                if ext not in ['doc', 'docx', 'pdf']:
                    continue

                full_path = os.path.join(root, fname)
                category = self._classify_file(fname)
                self.classified_files[category].append(full_path)

                # 记录.doc文件（需要手动转换）
                if ext == 'doc':
                    self.doc_files.append(fname)

                # 只有附加条款的 docx/pdf 才能提取
                if category == 'fujia' and ext in ['docx', 'pdf']:
                    self.selected_files.append(full_path)

        # 更新分类预览
        self.classify_preview.setVisible(True)
        self.preview_fujia.findChild(QLabel, "count").setText(str(len(self.classified_files['fujia'])))
        self.preview_feilv.findChild(QLabel, "count").setText(str(len(self.classified_files['feilv'])))
        self.preview_zhu.findChild(QLabel, "count").setText(str(len(self.classified_files['zhu'])))

        # 显示文件列表
        self.file_list.clear()
        category_icons = {'fujia': '📗', 'feilv': '📘', 'zhu': '📙'}
        for cat in ['fujia', 'feilv', 'zhu']:
            for fp in self.classified_files[cat]:
                fname = os.path.basename(fp)
                ext = fname.split('.')[-1].lower()
                # 标记.doc文件
                suffix = " ⚠️" if ext == 'doc' else ""
                self.file_list.addItem(f"{category_icons[cat]} {fname}{suffix}")
        self.file_list.setVisible(True)

        total = sum(len(v) for v in self.classified_files.values())
        self._log(f"📁 文件夹加载完成，共 {total} 个文件", "info")
        self._log(f"   📗 附加条款: {len(self.classified_files['fujia'])} 个", "info")
        self._log(f"   📘 费率表: {len(self.classified_files['feilv'])} 个", "info")
        self._log(f"   📙 主条款: {len(self.classified_files['zhu'])} 个", "info")

        # 警告.doc文件 - 弹出对话框
        if self.doc_files:
            self._log(f"⚠️ 发现 {len(self.doc_files)} 个 .doc 文件需要先转换为 .docx 格式:", "warning")
            for df in self.doc_files[:5]:
                self._log(f"   • {df}", "warning")
            if len(self.doc_files) > 5:
                self._log(f"   ... 还有 {len(self.doc_files) - 5} 个文件", "warning")
            self._log("💡 请使用 Microsoft Word 或 LibreOffice 打开后另存为 .docx 格式", "info")

            # 显示警告对话框
            self._show_doc_warning_dialog()

        # 启用提取按钮 - 只要有可提取的文件就启用
        if self.selected_files:
            self.extract_btn.setEnabled(True)
            self._log(f"✓ 将提取 {len(self.selected_files)} 个附加条款(.docx/.pdf)", "success")
        else:
            # 如果没有可提取文件但有附加条款的.doc文件，也提示
            fujia_doc_count = sum(1 for f in self.classified_files['fujia'] if f.endswith('.doc'))
            if fujia_doc_count > 0:
                self._log(f"⚠️ 有 {fujia_doc_count} 个附加条款为.doc格式，转换后即可提取", "warning")
                self.extract_btn.setEnabled(False)
            else:
                self._log("ℹ️ 未找到可提取的附加条款文件", "info")
                self.extract_btn.setEnabled(False)

        # 显示ZIP下载按钮
        if total > 0:
            self.download_zip_btn.setVisible(True)

        self._update_stats()

    def _classify_file(self, filename: str) -> str:
        """文件分类"""
        if '费率表' in filename or '费率方案' in filename:
            return 'feilv'

        # 匹配"附加xxx保险"或"附加xxx条款"
        fujia_pattern = r'附加.{1,20}(保险|条款)'
        if re.search(fujia_pattern, filename):
            return 'fujia'

        return 'zhu'

    def _show_doc_warning_dialog(self):
        """显示.doc文件警告对话框 - 使用自定义Dialog确保按钮可见"""
        doc_count = len(self.doc_files)
        fujia_doc_count = sum(1 for f in self.classified_files.get('fujia', []) if f.endswith('.doc'))

        # 创建自定义对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("发现旧版Word文档")
        dialog.setMinimumWidth(500)
        dialog.setStyleSheet(f"background: {AnthropicColors.BG_PRIMARY};")

        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        layout.setContentsMargins(25, 25, 25, 20)

        # 标题
        title = QLabel(f"发现 {doc_count} 个 .doc 格式文件")
        title.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 16px; font-weight: bold;")
        layout.addWidget(title)

        subtitle = QLabel("是否自动转换为 .docx 格式？")
        subtitle.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 14px;")
        layout.addWidget(subtitle)

        # 文件列表
        file_list = QLabel()
        detail_text = "检测到以下 .doc 文件:\n"
        for df in self.doc_files[:8]:
            detail_text += f"  • {df}\n"
        if len(self.doc_files) > 8:
            detail_text += f"  ... 还有 {len(self.doc_files) - 8} 个文件"
        file_list.setText(detail_text)
        file_list.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 12px; padding: 10px; background: {AnthropicColors.BG_CARD}; border-radius: 8px;")
        file_list.setWordWrap(True)
        layout.addWidget(file_list)

        # 提示信息
        if fujia_doc_count > 0:
            hint = QLabel(f"💡 其中 {fujia_doc_count} 个附加条款转换后可立即提取")
            hint.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-size: 13px; font-weight: 500;")
            layout.addWidget(hint)

        layout.addSpacing(10)

        # 按钮行
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        skip_btn = QPushButton("跳过")
        skip_btn.setMinimumSize(100, 44)
        skip_btn.setCursor(Qt.PointingHandCursor)
        skip_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                font-size: 14px;
                font-weight: 500;
                padding: 10px 25px;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.BG_PRIMARY};
                border-color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)
        skip_btn.clicked.connect(dialog.reject)

        convert_btn = QPushButton("✓ 自动转换")
        convert_btn.setMinimumSize(120, 44)
        convert_btn.setCursor(Qt.PointingHandCursor)
        convert_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.ACCENT};
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 14px;
                font-weight: 600;
                padding: 10px 25px;
            }}
            QPushButton:hover {{
                background: #c96747;
            }}
        """)
        convert_btn.clicked.connect(dialog.accept)

        btn_layout.addWidget(skip_btn)
        btn_layout.addWidget(convert_btn)
        layout.addLayout(btn_layout)

        # 显示对话框
        if dialog.exec_() == QDialog.Accepted:
            self._convert_doc_files()

    def _convert_doc_files(self):
        """批量转换.doc文件为.docx格式"""
        import subprocess
        import platform

        self._log(f"🔄 开始转换 {len(self.doc_files)} 个 .doc 文件...", "info")
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        converted = 0
        failed = 0
        converted_paths = []

        for i, doc_name in enumerate(self.doc_files):
            progress = int((i + 1) / len(self.doc_files) * 100)
            self.progress_bar.setValue(progress)
            QApplication.processEvents()

            # 查找完整路径
            doc_path = None
            for cat in ['fujia', 'feilv', 'zhu']:
                for fp in self.classified_files[cat]:
                    if os.path.basename(fp) == doc_name:
                        doc_path = fp
                        break
                if doc_path:
                    break

            if not doc_path:
                self._log(f"  ✗ 未找到文件路径: {doc_name}", "error")
                failed += 1
                continue

            docx_path = doc_path.rsplit('.', 1)[0] + '.docx'

            try:
                # Windows: 使用 LibreOffice 转换 .doc → .docx
                soffice_paths = [
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                    'soffice.exe',
                ]
                converted_ok = False
                for soffice in soffice_paths:
                    if soffice == 'soffice.exe' or os.path.exists(soffice):
                        output_dir = os.path.dirname(doc_path)
                        try:
                            result = subprocess.run(
                                [soffice, '--headless', '--convert-to', 'docx', '--outdir', output_dir, doc_path],
                                capture_output=True, text=True, timeout=120
                            )
                            if result.returncode == 0 and os.path.exists(docx_path):
                                converted_ok = True
                                break
                        except FileNotFoundError:
                            continue

                if not converted_ok:
                    self._log(f"  ✗ 需安装 LibreOffice 才能转换 .doc 文件", "error")

                if os.path.exists(docx_path):
                    converted += 1
                    converted_paths.append(docx_path)
                    self._log(f"  ✓ {doc_name} → .docx", "success")

                    # 更新分类列表
                    for cat in ['fujia', 'feilv', 'zhu']:
                        if doc_path in self.classified_files[cat]:
                            self.classified_files[cat].remove(doc_path)
                            self.classified_files[cat].append(docx_path)
                            # 如果是附加条款，添加到待提取列表
                            if cat == 'fujia':
                                self.selected_files.append(docx_path)
                            break
                else:
                    self._log(f"  ✗ 转换失败: {doc_name}", "error")
                    failed += 1

            except subprocess.TimeoutExpired:
                self._log(f"  ✗ 转换超时: {doc_name}", "error")
                failed += 1
            except Exception as e:
                self._log(f"  ✗ 转换错误: {doc_name} - {str(e)}", "error")
                failed += 1

        self.progress_bar.setValue(100)
        self._log(f"🎉 转换完成! 成功: {converted}, 失败: {failed}", "success" if failed == 0 else "warning")

        # 更新UI
        if converted > 0:
            self._refresh_file_list()
            self._update_stats()
            if self.selected_files:
                self.extract_btn.setEnabled(True)
                self._log(f"✓ 现在可以提取 {len(self.selected_files)} 个附加条款", "success")

    def _refresh_file_list(self):
        """刷新文件列表显示"""
        self.file_list.clear()
        category_icons = {'fujia': '📗', 'feilv': '📘', 'zhu': '📙'}
        for cat in ['fujia', 'feilv', 'zhu']:
            for fp in self.classified_files[cat]:
                fname = os.path.basename(fp)
                ext = fname.split('.')[-1].lower()
                suffix = " ⚠️" if ext == 'doc' else ""
                self.file_list.addItem(f"{category_icons[cat]} {fname}{suffix}")

        # 更新分类预览
        self.preview_fujia.findChild(QLabel, "count").setText(str(len(self.classified_files['fujia'])))
        self.preview_feilv.findChild(QLabel, "count").setText(str(len(self.classified_files['feilv'])))
        self.preview_zhu.findChild(QLabel, "count").setText(str(len(self.classified_files['zhu'])))

    def _start_extraction(self):
        """开始提取条款"""
        if not self.selected_files:
            self._log("⚠️ 请先选择文件", "warning")
            return

        self.extracted_data = []
        self.categories = set()
        self.extract_btn.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        self._log(f"🚀 开始处理 {len(self.selected_files)} 个文件...", "info")

        for i, fp in enumerate(self.selected_files):
            fname = os.path.basename(fp)
            progress = int((i + 1) / len(self.selected_files) * 100)
            self.progress_bar.setValue(progress)
            QApplication.processEvents()

            try:
                results = self._extract_clause(fp)
                for result in results:
                    self.extracted_data.append(result)
                    self.categories.add(result['Category'])
                    if result.get('Error'):
                        self._log(f"✗ {result['ClauseName']}: {result['Error']}", "error")
                    else:
                        self._log(f"✓ {result['ClauseName']} → {result['Category']}", "success")
            except Exception as e:
                self._log(f"✗ {fname}: {sanitize_error_message(e)}", "error")

        self.progress_bar.setValue(100)
        self._update_stats()

        success_count = len([d for d in self.extracted_data if not d.get('Error')])
        self._log(f"🎉 处理完成! 新增: {success_count} 条，共 {len(self.categories)} 个分类", "success")

        self.extract_btn.setEnabled(True)
        if self.extracted_data:
            self.download_excel_btn.setVisible(True)

    def _extract_clause(self, file_path: str) -> list:
        """提取单个文件的条款"""
        fname = os.path.basename(file_path)
        clause_name = os.path.splitext(fname)[0]
        today = datetime.now().strftime('%Y-%m-%d')
        ext = fname.split('.')[-1].lower()

        result = {
            'FileName': fname,
            'ClauseName': clause_name,
            'RegistrationNo': '',
            'Content': '',
            'Category': self._get_category(fname, clause_name),
            'AddDate': today,
            'Error': ''
        }

        try:
            if ext == 'pdf':
                paragraphs = self._parse_pdf(file_path)
            else:
                paragraphs = self._parse_docx(file_path)

            if not paragraphs:
                result['Error'] = '文档内容为空'
                return [result]

            # 提取注册号 - v18.16: 智能判断第三行是否为注册号
            reg_no_idx = -1  # 记录注册号所在的段落索引
            for i, para in enumerate(paragraphs[:8]):
                if '注册号' in para or re.search(r'[A-Z]\d{10,}', para):
                    match = re.search(r'[（\(]([^）\)]+)[）\)]', para)
                    if match:
                        result['RegistrationNo'] = match.group(1)
                    else:
                        result['RegistrationNo'] = re.sub(r'(产品)?注册号[:：]?', '', para).strip()
                    reg_no_idx = i
                    break

            # 提取正文 - v18.16: 智能确定起始位置
            # 只跳过：公司名（段落0）、条款名（段落1）、注册号（如果存在）
            content_lines = []

            # 判断应该从哪里开始提取内容
            if len(paragraphs) >= 3:
                # 检查段落2是否为注册号（已在上面识别）
                if reg_no_idx == 2:
                    # 段落2是注册号，从段落3开始
                    start_idx = 3
                else:
                    # 段落2不是注册号，从段落2开始（它是正文内容）
                    start_idx = 2
            else:
                start_idx = 0

            for para in paragraphs[start_idx:]:
                clean = para.strip()
                if clean and clean != clause_name and not self._is_noise_line(clean):
                    content_lines.append(clean)

            result['Content'] = '\n'.join(content_lines)
            return [result]

        except Exception as e:
            result['Error'] = f'解析出错: {str(e)}'
            return [result]

    def _parse_docx(self, file_path: str) -> list:
        """解析Word文档 - v18.17: 支持表格提取

        使用 <b>...</b> 标记加粗文本，便于后续导出时保留格式
        优化：合并同一段落中连续的加粗run，避免产生</b><b>
        新增：提取文档中的表格，转换为可读的文本格式
        """
        doc = Document(file_path)
        paragraphs = []

        # 构建元素顺序映射，以便按文档顺序处理段落和表格
        # Word文档的body中包含段落(p)和表格(tbl)，需要按顺序处理
        from docx.oxml.ns import qn

        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for child in body:
            if child.tag == qn('w:p'):
                # 这是一个段落
                if para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para_idx += 1

                    # 检查段落是否有内容
                    if not para.text.strip():
                        continue

                    # 构建带格式标记的文本，合并连续的加粗run
                    formatted_parts = []
                    current_bold_text = []

                    for run in para.runs:
                        text = run.text
                        if not text:
                            continue

                        if run.bold:
                            current_bold_text.append(text)
                        else:
                            if current_bold_text:
                                formatted_parts.append(f'<b>{"".join(current_bold_text)}</b>')
                                current_bold_text = []
                            formatted_parts.append(text)

                    if current_bold_text:
                        formatted_parts.append(f'<b>{"".join(current_bold_text)}</b>')

                    formatted_text = ''.join(formatted_parts).strip()
                    if formatted_text:
                        paragraphs.append(formatted_text)

            elif child.tag == qn('w:tbl'):
                # 这是一个表格
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_idx += 1

                    # 将表格转换为文本格式
                    table_text = self._table_to_text(table)
                    if table_text:
                        paragraphs.append(table_text)

        return paragraphs

    def _table_to_text(self, table) -> str:
        """将Word表格转换为可读的文本格式 - v18.17

        使用管道符和横线创建类似Markdown的表格格式
        """
        if not table.rows:
            return ''

        rows_data = []
        col_widths = []

        # 收集所有单元格数据并计算列宽
        for row in table.rows:
            row_cells = []
            for idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace('\n', ' ')
                row_cells.append(cell_text)
                # 更新列宽（考虑中文字符宽度）
                text_width = sum(2 if ord(c) > 127 else 1 for c in cell_text)
                if idx >= len(col_widths):
                    col_widths.append(text_width)
                else:
                    col_widths[idx] = max(col_widths[idx], text_width)
            rows_data.append(row_cells)

        # 确保最小列宽
        col_widths = [max(w, 4) for w in col_widths]

        # 生成表格文本
        lines = []
        lines.append('<table>')  # 表格开始标记

        for row_idx, row_cells in enumerate(rows_data):
            # 格式化每个单元格
            formatted_cells = []
            for idx, cell_text in enumerate(row_cells):
                width = col_widths[idx] if idx < len(col_widths) else 10
                # 计算实际需要的填充（考虑中文）
                text_width = sum(2 if ord(c) > 127 else 1 for c in cell_text)
                padding = width - text_width
                formatted_cells.append(cell_text + ' ' * max(0, padding))

            line = '| ' + ' | '.join(formatted_cells) + ' |'
            lines.append(line)

            # 在表头后添加分隔线
            if row_idx == 0:
                separator = '|' + '|'.join(['-' * (w + 2) for w in col_widths[:len(row_cells)]]) + '|'
                lines.append(separator)

        lines.append('</table>')  # 表格结束标记
        return '\n'.join(lines)

    def _parse_pdf(self, file_path: str) -> list:
        """解析PDF文档"""
        paragraphs = []

        if HAS_PDFPLUMBER:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        paragraphs.extend([l.strip() for l in lines if l.strip()])
        elif HAS_PYPDF2:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        paragraphs.extend([l.strip() for l in lines if l.strip()])
        else:
            raise ImportError("未安装PDF解析库 (pdfplumber 或 PyPDF2)")

        return paragraphs

    def _is_noise_line(self, text: str) -> bool:
        """判断是否为噪声行（页码、网址等明显非内容行）"""
        # 先清理文本（去除<b>标记以便正确匹配）
        clean_text = re.sub(r'</?b>', '', text).strip()
        if not clean_text:
            return True

        # 噪声正则模式（只过滤明显的非内容行）
        noise_patterns = [
            # 页码格式
            r'^第?\s*\d+\s*页\s*$',                      # 第1页
            r'^Page\s*\d+\s*$',                          # Page 1
            r'^第\s*\d+\s*页\s*共\s*\d+\s*页\s*$',       # 第1页共10页
            r'^\d+\s*/\s*\d+\s*$',                       # 1/10 页码格式
            r'^[-—]\s*\d+\s*[-—]\s*$',                   # -1- 页码格式
            # Word域代码（textutil转换产生）
            r'PAGE\s*\\?\*?\s*MERGEFORMAT',              # PAGE \* MERGEFORMAT
            r'NUMPAGES',                                  # NUMPAGES 域代码
            # 日期格式
            r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}\s*$',
            # 网址
            r'^www\.',
            r'^http',
        ]

        for pattern in noise_patterns:
            if re.search(pattern, clean_text, re.IGNORECASE):
                return True

        return False

    def _get_category(self, filename: str, title: str) -> str:
        """获取条款分类"""
        text = title or filename
        if '附加' in text:
            parts = text.split('附加')
            prefix = parts[0].replace('条款', '').strip()
            if prefix and len(prefix) > 2:
                return prefix + '附加条款'
        return '通用附加条款'

    def _download_classified_zip(self):
        """下载分类后的ZIP文件"""
        if not any(self.classified_files.values()):
            self._log("⚠️ 没有可下载的文件", "warning")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "保存分类ZIP",
            f"条款分类_{datetime.now():%Y%m%d_%H%M}.zip",
            "ZIP文件 (*.zip)"
        )
        if not save_path:
            return

        self._log("📦 正在生成分类ZIP文件...", "info")

        try:
            with zipfile.ZipFile(save_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                folder_names = {'fujia': '附加条款', 'feilv': '费率表', 'zhu': '主条款'}
                for cat, files in self.classified_files.items():
                    folder_name = folder_names[cat]
                    for fp in files:
                        fname = os.path.basename(fp)
                        zf.write(fp, f"{folder_name}/{fname}")

            total = sum(len(v) for v in self.classified_files.values())
            self._log(f"✅ 分类ZIP已保存: {os.path.basename(save_path)}", "success")
            self._log(f"   包含 {len(self.classified_files['fujia'])} 附加条款 + {len(self.classified_files['feilv'])} 费率表 + {len(self.classified_files['zhu'])} 主条款", "info")
        except Exception as e:
            self._log(f"❌ ZIP生成失败: {sanitize_error_message(e)}", "error")

    def _parse_rich_text(self, text: str):
        """v18.9: 解析带 <b> 标记的文本，返回富文本对象

        返回: CellRichText 对象（如果有格式标记）或普通字符串

        优化：将换行符合并到相邻文本块中，避免Excel不渲染单独的换行块
        """
        if not text or '<b>' not in text:
            return text

        try:
            from openpyxl.cell.rich_text import CellRichText, TextBlock
            from openpyxl.cell.text import InlineFont

            rich_text = CellRichText()
            pattern = re.compile(r'<b>(.*?)</b>', re.DOTALL)
            last_end = 0
            pending_whitespace = ''  # 待处理的空白/换行

            for match in pattern.finditer(text):
                # 处理当前匹配之前的非加粗部分
                if match.start() > last_end:
                    normal_text = text[last_end:match.start()]
                    if normal_text:
                        # 如果只是空白/换行，先保存待后续处理
                        if normal_text.strip() == '':
                            pending_whitespace = normal_text
                        else:
                            # 有实际内容的非加粗文本
                            rich_text.append(pending_whitespace + normal_text)
                            pending_whitespace = ''

                # 添加加粗部分（包含前置的换行）
                bold_text = match.group(1)
                if bold_text:
                    # 将待处理的空白合并到加粗文本开头
                    full_bold = pending_whitespace + bold_text
                    rich_text.append(TextBlock(InlineFont(b=True), full_bold))
                    pending_whitespace = ''

                last_end = match.end()

            # 添加最后的非加粗部分
            if last_end < len(text):
                remaining = text[last_end:]
                if remaining:
                    rich_text.append(pending_whitespace + remaining)
                    pending_whitespace = ''

            # 处理末尾剩余的空白
            if pending_whitespace:
                rich_text.append(pending_whitespace)

            return rich_text if rich_text else text

        except ImportError:
            # 如果不支持富文本，返回去除标记的纯文本
            return re.sub(r'</?b>', '', text)

    def _download_excel_report(self):
        """下载Excel报告 - Anthropic风格，v18.15支持保留加粗格式"""
        if not self.extracted_data:
            self._log("⚠️ 没有可导出的数据", "warning")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "保存Excel报告",
            f"附加条款提取_{datetime.now():%Y%m%d_%H%M}.xlsx",
            "Excel文件 (*.xlsx)"
        )
        if not save_path:
            return

        try:
            # 按分类分组
            grouped = defaultdict(list)
            for item in self.extracted_data:
                cat = item.get('Category', '其他附加条款')
                grouped[cat].append(item)

            # 创建工作簿
            wb = openpyxl.Workbook()
            wb.remove(wb.active)

            # Anthropic 风格颜色
            header_fill = PatternFill(start_color="141413", end_color="141413", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=11)
            accent_fill = PatternFill(start_color="FAF9F5", end_color="FAF9F5", fill_type="solid")
            success_font = Font(color="5A9A7A")
            error_font = Font(color="C75050")
            border_style = Border(
                left=Side(style='thin', color='E0DED5'),
                right=Side(style='thin', color='E0DED5'),
                top=Side(style='thin', color='E0DED5'),
                bottom=Side(style='thin', color='E0DED5')
            )

            for sheet_name, items in grouped.items():
                safe_name = sheet_name[:30].replace('/', ' ').replace('\\', ' ')
                ws = wb.create_sheet(title=safe_name)

                # 表头
                headers = ['条款名称', '注册号', '条款内容', '原文件名', '添加日期', '状态']
                ws.append(headers)

                # 表头样式
                for col in range(1, 7):
                    cell = ws.cell(row=1, column=col)
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal='center', vertical='center')

                # 数据
                for row_idx, item in enumerate(items, start=2):
                    # 先添加基本数据（除内容列外）
                    ws.cell(row=row_idx, column=1, value=item['ClauseName'])
                    ws.cell(row=row_idx, column=2, value=item['RegistrationNo'])

                    # v18.15: 内容列使用富文本保留加粗格式
                    content = item['Content'][:30000] if item['Content'] else ''
                    content_cell = ws.cell(row=row_idx, column=3)
                    rich_content = self._parse_rich_text(content)
                    content_cell.value = rich_content

                    ws.cell(row=row_idx, column=4, value=item['FileName'])
                    ws.cell(row=row_idx, column=5, value=item['AddDate'])
                    ws.cell(row=row_idx, column=6, value=f"失败: {item['Error']}" if item.get('Error') else '成功')

                    # 数据行样式
                    for col in range(1, 7):
                        cell = ws.cell(row=row_idx, column=col)
                        cell.border = border_style
                        cell.alignment = Alignment(vertical='center', wrap_text=(col == 3))
                        if row_idx % 2 == 0:
                            cell.fill = accent_fill

                    # 状态列颜色
                    status_cell = ws.cell(row=row_idx, column=6)
                    if item.get('Error'):
                        status_cell.font = error_font
                    else:
                        status_cell.font = success_font

                # 列宽
                ws.column_dimensions['A'].width = 40
                ws.column_dimensions['B'].width = 25
                ws.column_dimensions['C'].width = 100
                ws.column_dimensions['D'].width = 45
                ws.column_dimensions['E'].width = 12
                ws.column_dimensions['F'].width = 12

                # 冻结首行
                ws.freeze_panes = 'A2'

            wb.save(save_path)
            self._log(f"✅ Excel报告已保存: {os.path.basename(save_path)}", "success")

            # 打开文件所在目录
            try:
                os.startfile(os.path.dirname(save_path))
            except Exception:
                pass

        except Exception as e:
            self._log(f"❌ Excel导出失败: {sanitize_error_message(e)}", "error")

    def _clear_all(self):
        """清空所有"""
        self.selected_files = []
        self.classified_files = {'fujia': [], 'feilv': [], 'zhu': []}
        self.doc_files = []
        self.extracted_data = []
        self.categories = set()

        self.file_list.clear()
        self.file_list.setVisible(False)
        self.classify_preview.setVisible(False)
        self.download_zip_btn.setVisible(False)
        self.download_excel_btn.setVisible(False)
        self.extract_btn.setEnabled(False)
        self.progress_bar.setVisible(False)
        self.progress_bar.setValue(0)

        self._update_stats()
        self._log("🗑 已清空所有文件", "info")

    def _update_stats(self):
        """更新统计 - 有数据时显示统计面板"""
        total = len(self.selected_files)
        extracted = len(self.extracted_data)
        categories = len(self.categories)
        skipped = len([d for d in self.extracted_data if d.get('Error')])

        self.stat_total_label.setText(f"待处理: {total}")
        self.stat_extracted_label.setText(f"已提取: {extracted}")
        self.stat_categories_label.setText(f"分类数: {categories}")
        self.stat_skipped_label.setText(f"已跳过: {skipped}")

        # 有任何数据时显示统计面板，否则隐藏
        has_data = total > 0 or extracted > 0
        self.stats_frame.setVisible(has_data)

    def _log(self, message: str, level: str = "info"):
        """添加日志"""
        colors = {
            'info': '#e0e0e0',
            'success': '#7ec9a0',
            'warning': '#e5c07b',
            'error': '#e06c75'
        }
        color = colors.get(level, '#e0e0e0')
        self.log_text.append(f'<span style="color: {color}">{message}</span>')


# ==========================================
# 条款输出Tab - V18.0 完整功能
# ==========================================
class ClauseOutputTab(QWidget):
    """条款输出Tab - Excel/提取结果转Word文档"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_window = parent
        self.report_data = []  # 存储读取的条款数据
        self.selected_clauses = []  # 用户选中的条款
        self.source_excel_path = None  # v18.15: 源Excel路径，用于录单增强模式
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(6)
        layout.setContentsMargins(15, 10, 15, 10)

        # 滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setStyleSheet(get_anthropic_scrollbar_style())
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setSpacing(6)
        scroll_layout.setContentsMargins(0, 0, 0, 0)

        # 紧凑型标题栏
        header = QHBoxLayout()
        title = QLabel("📝 条款输出")
        title.setStyleSheet(f"""
            color: {AnthropicColors.TEXT_PRIMARY};
            font-size: 18px;
            font-weight: bold;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
        """)
        header.addWidget(title)
        header.addStretch()

        # 输出模式选择
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["按条款逐个输出", "按分类合并输出", "全部合并为一个文档"])
        self.mode_combo.setStyleSheet(f"""
            QComboBox {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 6px;
                padding: 10px 14px;
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
                min-width: 180px;
            }}
            QComboBox:hover {{ border-color: {AnthropicColors.ACCENT}; }}
            QComboBox::drop-down {{
                border: none;
                width: 20px;
            }}
            QComboBox QAbstractItemView {{
                background: {AnthropicColors.BG_PRIMARY};
                color: {AnthropicColors.TEXT_PRIMARY};
                selection-background-color: {AnthropicColors.ACCENT};
                selection-color: white;
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 6px;
                padding: 5px;
            }}
        """)
        mode_label = QLabel("输出模式:")
        mode_label.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-weight: 500;")
        header.addWidget(mode_label)
        header.addWidget(self.mode_combo)
        scroll_layout.addLayout(header)

        # 数据源选择卡片
        source_card = GlassCard()
        source_layout = QVBoxLayout(source_card)
        source_layout.setSpacing(12)

        source_title = QLabel("📊 选择数据源")
        source_title.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-weight: 600; font-size: 15px;")
        source_layout.addWidget(source_title)

        # 数据源按钮行
        source_btn_layout = QHBoxLayout()

        self.from_extract_btn = QPushButton("📄 从条款提取获取")
        self.from_extract_btn.setCursor(Qt.PointingHandCursor)
        self.from_extract_btn.setMinimumWidth(180)
        self.from_extract_btn.setMinimumHeight(38)
        self.from_extract_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 12px 24px;
                color: {AnthropicColors.TEXT_PRIMARY};
                font-weight: 500;
                font-size: 15px;
                min-width: 180px;
            }}
            QPushButton:hover {{
                border-color: {AnthropicColors.ACCENT};
                background: rgba(217, 119, 87, 0.08);
            }}
        """)
        self.from_extract_btn.clicked.connect(self._load_from_extractor)

        self.from_file_btn = QPushButton("📁 从Excel文件加载")
        self.from_file_btn.setCursor(Qt.PointingHandCursor)
        self.from_file_btn.setMinimumWidth(180)
        self.from_file_btn.setMinimumHeight(38)
        self.from_file_btn.setStyleSheet(self.from_extract_btn.styleSheet())
        self.from_file_btn.clicked.connect(self._load_from_excel)

        source_btn_layout.addWidget(self.from_extract_btn)
        source_btn_layout.addWidget(self.from_file_btn)
        source_btn_layout.addStretch()
        source_layout.addLayout(source_btn_layout)

        # 文件路径显示
        self.source_label = QLabel("未选择数据源")
        self.source_label.setStyleSheet(f"color: {AnthropicColors.TEXT_MUTED}; font-size: 14px; padding: 5px 0;")
        source_layout.addWidget(self.source_label)

        scroll_layout.addWidget(source_card)

        # 条款预览列表
        preview_card = GlassCard()
        preview_layout = QVBoxLayout(preview_card)
        preview_layout.setSpacing(10)

        preview_header = QHBoxLayout()
        preview_title = QLabel("📋 条款预览")
        preview_title.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-weight: 600; font-size: 15px;")
        preview_header.addWidget(preview_title)

        self.clause_count_label = QLabel("共 0 条")
        self.clause_count_label.setStyleSheet(f"color: {AnthropicColors.TEXT_MUTED}; font-size: 14px;")
        preview_header.addWidget(self.clause_count_label)
        preview_header.addStretch()

        # 全选/取消按钮
        self.select_all_btn = QPushButton("全选")
        self.select_all_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                border: none;
                color: {AnthropicColors.ACCENT};
                font-size: 14px;
                padding: 4px 8px;
            }}
            QPushButton:hover {{ text-decoration: underline; }}
        """)
        self.select_all_btn.clicked.connect(self._toggle_select_all)
        preview_header.addWidget(self.select_all_btn)

        preview_layout.addLayout(preview_header)

        # 条款列表
        self.clause_list = QListWidget()
        self.clause_list.setMinimumHeight(150)
        self.clause_list.setStyleSheet(f"""
            QListWidget {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 8px;
                font-size: 15px;
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
            QListWidget::item {{
                padding: 10px 12px;
                border-radius: 6px;
                margin-bottom: 2px;
            }}
            QListWidget::item:hover {{
                background: rgba(217, 119, 87, 0.08);
            }}
            QListWidget::item:selected {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
            }}
        """)
        preview_layout.addWidget(self.clause_list)

        scroll_layout.addWidget(preview_card)

        # v18.17: Word样式设置 + 录单增强模式 合并为一个卡片，左右并排
        settings_card = GlassCard()
        settings_main_layout = QHBoxLayout(settings_card)
        settings_main_layout.setSpacing(12)

        # 左侧: Word样式设置
        style_section = QWidget()
        style_layout = QVBoxLayout(style_section)
        style_layout.setContentsMargins(0, 0, 0, 0)
        style_layout.setSpacing(10)

        style_title = QLabel("🎨 Word样式设置")
        style_title.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-weight: 600; font-size: 15px;")
        style_layout.addWidget(style_title)

        style_grid = QHBoxLayout()
        style_grid.setSpacing(15)

        # 标签通用样式
        label_style = f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 15px; font-weight: 500;"
        spin_style = f"""
            QSpinBox, QDoubleSpinBox {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 6px;
                padding: 8px;
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
                min-width: 70px;
                min-height: 32px;
            }}
            QSpinBox:focus, QDoubleSpinBox:focus {{
                border-color: {AnthropicColors.ACCENT};
            }}
        """

        # 标题字号
        title_size_layout = QVBoxLayout()
        title_label = QLabel("标题字号")
        title_label.setStyleSheet(label_style)
        title_size_layout.addWidget(title_label)
        self.title_size_spin = QSpinBox()
        self.title_size_spin.setRange(12, 36)
        self.title_size_spin.setValue(16)
        self.title_size_spin.setStyleSheet(spin_style)
        title_size_layout.addWidget(self.title_size_spin)
        style_grid.addLayout(title_size_layout)

        # 正文字号
        body_size_layout = QVBoxLayout()
        body_label = QLabel("正文字号")
        body_label.setStyleSheet(label_style)
        body_size_layout.addWidget(body_label)
        self.body_size_spin = QSpinBox()
        self.body_size_spin.setRange(9, 18)
        self.body_size_spin.setValue(12)
        self.body_size_spin.setStyleSheet(spin_style)
        body_size_layout.addWidget(self.body_size_spin)
        style_grid.addLayout(body_size_layout)

        # 行距
        line_spacing_layout = QVBoxLayout()
        spacing_label = QLabel("行距")
        spacing_label.setStyleSheet(label_style)
        line_spacing_layout.addWidget(spacing_label)
        self.line_spacing_spin = QDoubleSpinBox()
        self.line_spacing_spin.setRange(1.0, 3.0)
        self.line_spacing_spin.setValue(1.5)
        self.line_spacing_spin.setSingleStep(0.25)
        self.line_spacing_spin.setStyleSheet(spin_style)
        line_spacing_layout.addWidget(self.line_spacing_spin)
        style_grid.addLayout(line_spacing_layout)

        # 包含注册号
        include_reg_layout = QVBoxLayout()
        reg_label = QLabel("注册号")
        reg_label.setStyleSheet(label_style)
        include_reg_layout.addWidget(reg_label)
        self.include_reg_check = QCheckBox("显示")
        self.include_reg_check.setChecked(True)
        self.include_reg_check.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 15px;")
        include_reg_layout.addWidget(self.include_reg_check)
        style_grid.addLayout(include_reg_layout)

        style_layout.addLayout(style_grid)
        style_layout.addStretch()
        settings_main_layout.addWidget(style_section)

        # 分隔线
        separator = QFrame()
        separator.setFrameShape(QFrame.VLine)
        separator.setStyleSheet(f"background-color: {AnthropicColors.BORDER};")
        separator.setFixedWidth(1)
        settings_main_layout.addWidget(separator)

        # 右侧: 录单增强模式
        enhanced_section = QWidget()
        enhanced_layout = QVBoxLayout(enhanced_section)
        enhanced_layout.setContentsMargins(0, 0, 0, 0)
        enhanced_layout.setSpacing(8)

        self.enhanced_mode_check = QCheckBox("📋 录单增强模式")
        self.enhanced_mode_check.setStyleSheet(f"color: {AnthropicColors.ACCENT}; font-size: 15px; font-weight: 600;")
        self.enhanced_mode_check.toggled.connect(self._toggle_enhanced_mode)
        enhanced_layout.addWidget(self.enhanced_mode_check)

        enhanced_desc = QLabel("在条款名称前添加保险公司名称")
        enhanced_desc.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 14px;")
        enhanced_layout.addWidget(enhanced_desc)

        self.enhanced_options = QWidget()
        enhanced_options_layout = QVBoxLayout(self.enhanced_options)
        enhanced_options_layout.setContentsMargins(0, 5, 0, 0)
        enhanced_options_layout.setSpacing(5)

        self.company_prefix_edit = QLineEdit()
        self.company_prefix_edit.setPlaceholderText("输入公司名称前缀...")
        self.company_prefix_edit.setStyleSheet(f"""
            QLineEdit {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 6px;
                padding: 10px 12px;
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
            }}
            QLineEdit:focus {{ border-color: {AnthropicColors.ACCENT}; }}
        """)
        enhanced_options_layout.addWidget(self.company_prefix_edit)

        hint_label = QLabel("💡 输出: Excel + Word")
        hint_label.setStyleSheet(f"color: {AnthropicColors.TEXT_MUTED}; font-size: 13px;")
        enhanced_options_layout.addWidget(hint_label)

        self.enhanced_options.setVisible(False)
        enhanced_layout.addWidget(self.enhanced_options)
        enhanced_layout.addStretch()

        settings_main_layout.addWidget(enhanced_section)

        # 设置左右两侧的比例 (左侧稍宽)
        settings_main_layout.setStretch(0, 3)  # Word样式
        settings_main_layout.setStretch(1, 0)  # 分隔线
        settings_main_layout.setStretch(2, 2)  # 录单增强

        scroll_layout.addWidget(settings_card)

        # 操作按钮行
        btn_layout = QHBoxLayout()

        self.generate_btn = QPushButton("📄 生成Word文档")
        self.generate_btn.setMinimumHeight(44)
        self.generate_btn.setCursor(Qt.PointingHandCursor)
        self.generate_btn.setEnabled(False)
        self.generate_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
                border: none;
                border-radius: 8px;
                font-size: 16px;
                font-weight: 600;
            }}
            QPushButton:hover {{ background: #2a2a28; }}
            QPushButton:disabled {{
                background: {AnthropicColors.BORDER};
                color: {AnthropicColors.TEXT_SECONDARY};
            }}
        """)
        self.generate_btn.clicked.connect(self._generate_word)

        self.preview_btn = QPushButton("👁 预览")
        self.preview_btn.setMinimumHeight(44)
        self.preview_btn.setCursor(Qt.PointingHandCursor)
        self.preview_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent;
                color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                font-size: 15px;
                padding: 0 25px;
            }}
            QPushButton:hover {{
                border-color: {AnthropicColors.ACCENT};
                background: rgba(217, 119, 87, 0.05);
            }}
        """)
        self.preview_btn.clicked.connect(self._preview_output)

        btn_layout.addWidget(self.generate_btn, 3)
        btn_layout.addWidget(self.preview_btn, 1)
        scroll_layout.addLayout(btn_layout)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.setStyleSheet(f"""
            QProgressBar {{ background: {AnthropicColors.BORDER}; border-radius: 3px; }}
            QProgressBar::chunk {{ background: {AnthropicColors.ACCENT}; border-radius: 3px; }}
        """)
        scroll_layout.addWidget(self.progress_bar)

        # 完成滚动区域设置
        scroll_area.setWidget(scroll_widget)
        layout.addWidget(scroll_area, 1)

        # 日志区域
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(120)
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background: {AnthropicColors.BG_DARK};
                border: none;
                border-radius: 8px;
                color: {AnthropicColors.TEXT_LIGHT};
                padding: 12px;
                font-family: 'Cascadia Code', 'Consolas', monospace;
                font-size: 13px;
            }}
        """)
        layout.addWidget(self.log_text)

        self._log("📝 条款输出工具已就绪", "info")
        self._log("   支持从条款提取结果或Excel文件加载数据", "info")

    def _load_from_extractor(self):
        """从条款提取Tab获取数据"""
        if not self.parent_window:
            self._log("❌ 无法获取父窗口引用", "error")
            return

        try:
            extractor_tab = self.parent_window.extractor_tab
            if not extractor_tab.extracted_data:
                self._log("⚠️ 条款提取Tab中没有已提取的数据", "warning")
                self._log("   请先在「条款提取」页面提取条款", "info")
                return

            self.report_data = []
            for item in extractor_tab.extracted_data:
                if not item.get('Error'):
                    self.report_data.append({
                        'name': item.get('ClauseName', ''),
                        'regNo': item.get('RegistrationNo', ''),
                        'content': item.get('Content', ''),
                        'category': item.get('Category', '其他'),
                        'filename': item.get('FileName', '')
                    })

            if self.report_data:
                self._update_clause_list()
                self.source_label.setText(f"✓ 已从条款提取加载 {len(self.report_data)} 条数据")
                self._log(f"✓ 从条款提取Tab加载了 {len(self.report_data)} 条条款", "success")
                self.generate_btn.setEnabled(True)
            else:
                self._log("⚠️ 没有成功提取的条款数据", "warning")

        except Exception as e:
            self._log(f"❌ 加载失败: {sanitize_error_message(e)}", "error")

    def _load_from_excel(self):
        """从Excel文件加载数据 - v18.9: 支持富文本格式（加粗）"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择Excel文件", "",
            "Excel文件 (*.xlsx);;所有文件 (*.*)"
        )
        if not file_path:
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(10)
        self.source_excel_path = file_path  # v18.15: 保存源文件路径

        try:
            self._log(f"📖 读取文件: {os.path.basename(file_path)}", "info")

            # v18.9: 使用 rich_text=True 以保留加粗格式
            wb = openpyxl.load_workbook(file_path, rich_text=True)
            self.report_data = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows())
                if not rows:
                    continue

                headers = [str(cell.value) if cell.value else '' for cell in rows[0]]

                # 智能识别列
                col_map = self._detect_columns(headers)

                if not col_map.get('name'):
                    continue

                self.progress_bar.setValue(30)

                for row in rows[1:]:
                    if not row or not any(cell.value for cell in row):
                        continue

                    name = self._extract_cell_text(row, col_map.get('name'))
                    if not name:
                        continue

                    self.report_data.append({
                        'name': name,
                        'regNo': self._extract_cell_text(row, col_map.get('regNo')),
                        'content': self._extract_cell_text(row, col_map.get('content')),
                        'category': sheet_name if sheet_name != 'Sheet' else '条款',
                        'filename': self._extract_cell_text(row, col_map.get('filename'))
                    })

            wb.close()
            self.progress_bar.setValue(80)

            # 去重
            seen = set()
            unique_data = []
            for item in self.report_data:
                if item['name'] not in seen:
                    seen.add(item['name'])
                    unique_data.append(item)
            self.report_data = unique_data

            if self.report_data:
                self._update_clause_list()
                self.source_label.setText(f"✓ {os.path.basename(file_path)} ({len(self.report_data)} 条)")
                self._log(f"✓ 加载了 {len(self.report_data)} 条不重复条款", "success")
                self.generate_btn.setEnabled(True)
            else:
                self._log("⚠️ 文件中未找到有效条款数据", "warning")

        except Exception as e:
            self._log(f"❌ 读取Excel失败: {sanitize_error_message(e)}", "error")
        finally:
            self.progress_bar.setVisible(False)

    def _detect_columns(self, headers: list) -> dict:
        """智能识别Excel列 - 优先匹配「匹配1_」前缀的列（E/F/G）"""
        col_map = {}

        # 第一优先级：查找「匹配1_」前缀的列
        for i, h in enumerate(headers):
            if not h:
                continue
            h_str = str(h)
            if '匹配1_条款名称' in h_str or h_str == '匹配1_条款名称':
                col_map['name'] = i
            elif '匹配1_注册号' in h_str or '匹配1_产品注册号' in h_str:
                col_map['regNo'] = i
            elif '匹配1_条款内容' in h_str:
                col_map['content'] = i

        # 如果找到了匹配1_列，记录日志
        if col_map.get('name'):
            self._log(f"✓ 识别到匹配列: E={col_map.get('name')}, F={col_map.get('regNo')}, G={col_map.get('content')}", "success")
            return col_map

        # 第二优先级：直接使用E/F/G列（索引4/5/6）
        col_map['name'] = 4      # E列 = 匹配1_条款名称
        col_map['regNo'] = 5     # F列 = 匹配1_注册号
        col_map['content'] = 6   # G列 = 匹配1_条款内容
        self._log("ℹ️ 使用默认列: E=条款名称, F=注册号, G=条款内容", "info")

        return col_map

    def _safe_get(self, row: tuple, index: int) -> str:
        """安全获取行数据"""
        if index is None or index >= len(row):
            return ''
        return str(row[index] or '').strip()

    def _extract_cell_text(self, row, index: int) -> str:
        """v18.9: 从单元格提取文本，保留加粗格式（使用<b>标记）"""
        if index is None or index >= len(row):
            return ''

        cell = row[index]
        if cell.value is None:
            return ''

        # 检查是否为富文本
        try:
            from openpyxl.cell.rich_text import CellRichText
            if isinstance(cell.value, CellRichText):
                result = []
                for block in cell.value:
                    if hasattr(block, 'font') and block.font and block.font.b:
                        # 加粗文本
                        result.append(f'<b>{block.text}</b>')
                    elif hasattr(block, 'text'):
                        result.append(block.text)
                    else:
                        result.append(str(block))
                return ''.join(result).strip()
        except (ImportError, AttributeError):
            pass

        return str(cell.value).strip() if cell.value else ''

    def _update_clause_list(self):
        """更新条款列表显示"""
        self.clause_list.clear()

        for item in self.report_data:
            list_item = QListWidgetItem()
            list_item.setCheckState(Qt.Checked)

            # 显示格式：条款名称 (分类)
            display_text = item['name']
            if item.get('category'):
                display_text += f"  [{item['category']}]"
            if item.get('regNo'):
                display_text += f"  ({item['regNo'][:20]}...)" if len(item.get('regNo', '')) > 20 else f"  ({item['regNo']})"

            list_item.setText(display_text)
            list_item.setData(Qt.UserRole, item)
            self.clause_list.addItem(list_item)

        self.clause_count_label.setText(f"共 {len(self.report_data)} 条")

    def _toggle_select_all(self):
        """切换全选/取消"""
        # 检查当前状态
        all_checked = all(
            self.clause_list.item(i).checkState() == Qt.Checked
            for i in range(self.clause_list.count())
        )

        new_state = Qt.Unchecked if all_checked else Qt.Checked
        for i in range(self.clause_list.count()):
            self.clause_list.item(i).setCheckState(new_state)

        self.select_all_btn.setText("取消全选" if not all_checked else "全选")

    def _get_selected_clauses(self) -> list:
        """获取选中的条款"""
        selected = []
        for i in range(self.clause_list.count()):
            item = self.clause_list.item(i)
            if item.checkState() == Qt.Checked:
                selected.append(item.data(Qt.UserRole))
        return selected

    def _preview_output(self):
        """预览输出"""
        selected = self._get_selected_clauses()
        if not selected:
            self._log("⚠️ 请至少选择一条条款", "warning")
            return

        preview_text = f"将输出 {len(selected)} 条条款:\n\n"
        for i, clause in enumerate(selected[:10], 1):
            preview_text += f"{i}. {clause['name']}\n"
        if len(selected) > 10:
            preview_text += f"... 还有 {len(selected) - 10} 条\n"

        preview_text += f"\n输出模式: {self.mode_combo.currentText()}"
        preview_text += f"\n标题字号: {self.title_size_spin.value()}pt"
        preview_text += f"\n正文字号: {self.body_size_spin.value()}pt"

        QMessageBox.information(self, "输出预览", preview_text)


    def _toggle_enhanced_mode(self, checked):
        """v18.15: 切换录单增强模式"""
        self.enhanced_options.setVisible(checked)
        if checked:
            self.generate_btn.setText("📄 生成文档（Excel + Word）")
        else:
            self.generate_btn.setText("📄 生成Word文档")


    def _generate_enhanced_documents(self, clauses: list, company_prefix: str):
        """v18.16: 录单增强模式 - 复制原Excel并在F列(匹配1_条款名称)添加前缀，保留富文本"""
        from datetime import datetime
        from copy import copy
        from openpyxl.cell.rich_text import CellRichText, TextBlock
        from openpyxl.cell.text import InlineFont

        # 检查源文件
        if not self.source_excel_path or not os.path.exists(self.source_excel_path):
            self._log("⚠️ 请先从Excel文件加载数据（录单增强模式需要原始Excel文件）", "warning")
            return

        # 选择保存目录
        output_dir = QFileDialog.getExistingDirectory(
                self, "选择输出目录", "",
                QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
            )
        if not output_dir:
            return

        self.progress_bar.setVisible(True)
        self._log(f"🏢 录单增强模式: 添加前缀 '{company_prefix}'", "info")

        try:
            # 1. 复制原Excel并修改F列
            self.progress_bar.setValue(20)
            self._log("📊 生成增强版Excel（保持原格式+富文本）...", "info")

            # v18.16: 加载原文件（保留样式+富文本）
            wb = openpyxl.load_workbook(self.source_excel_path, rich_text=True)

            for ws in wb.worksheets:
                # 找到 匹配1_条款名称 列 (通常是F列)
                name_col = None
                headers = [cell.value for cell in ws[1]] if ws[1] else []

                for i, h in enumerate(headers):
                    h_str = str(h) if h else ''
                    if '匹配1_条款名称' in h_str:
                        name_col = i + 1  # openpyxl 列从1开始
                        break

                if not name_col:
                    # 如果没找到，尝试找普通的条款名称列
                    for i, h in enumerate(headers):
                        h_str = str(h) if h else ''
                        if '条款名称' in h_str or 'MATCH' in h_str.upper():
                            name_col = i + 1
                            break

                if not name_col:
                    self._log(f"⚠️ 工作表 '{ws.title}' 未找到条款名称列，跳过", "warning")
                    continue

                self._log(f"   处理工作表 '{ws.title}': 条款名称在第 {name_col} 列", "info")

                # v18.16: 在条款名称前添加前缀（跳过表头），保留富文本
                for row_idx in range(2, ws.max_row + 1):
                    cell = ws.cell(row=row_idx, column=name_col)
                    if cell.value:
                        # 处理富文本 - 在最前面添加前缀
                        if isinstance(cell.value, CellRichText):
                            # 创建新的富文本，前缀不加粗
                            new_blocks = [TextBlock(InlineFont(), company_prefix)]
                            for block in cell.value:
                                new_blocks.append(block)
                            cell.value = CellRichText(*new_blocks)
                        else:
                            cell.value = company_prefix + str(cell.value)

            self.progress_bar.setValue(50)

            # 保存增强版Excel
            date_str = datetime.now().strftime("%Y%m%d_%H%M")
            excel_path = os.path.join(output_dir, f"条款比对报告_录单版_{date_str}.xlsx")
            wb.save(excel_path)
            wb.close()

            self._log(f"✓ Excel已保存: {os.path.basename(excel_path)}", "success")
            self.progress_bar.setValue(70)

            # 2. 生成Word文档（使用增强后的名称）
            self._log("📄 生成Word文档...", "info")

            enhanced_clauses = []
            for clause in clauses:
                enhanced_clauses.append({
                    'name': company_prefix + clause.get('name', ''),
                    'regNo': clause.get('regNo', ''),
                    'content': clause.get('content', ''),
                    'category': clause.get('category', '其他'),
                })

            word_path = os.path.join(output_dir, f"条款清单_录单版_{date_str}.docx")
            self._generate_combined_doc(enhanced_clauses, word_path)

            self.progress_bar.setValue(100)
            self._log(f"✓ Word已保存: {os.path.basename(word_path)}", "success")
            self._log(f"🎉 录单增强模式完成！已生成 Excel + Word 两个文件", "success")

            # 打开输出目录
            try:
                os.startfile(output_dir)
            except Exception:
                pass

        except Exception as e:
            logger.exception("录单增强模式生成失败")
            self._log(f"❌ 生成失败: {str(e)}", "error")
        finally:
            self.progress_bar.setVisible(False)

    def _generate_word(self):
        """生成Word文档"""
        selected = self._get_selected_clauses()
        if not selected:
            self._log("⚠️ 请至少选择一条条款", "warning")
            return

        # v18.15: 检查录单增强模式
        if self.enhanced_mode_check.isChecked():
            company_prefix = self.company_prefix_edit.text().strip()
            if not company_prefix:
                self._log("⚠️ 请输入保险公司名称前缀", "warning")
                self.company_prefix_edit.setFocus()
                return
            self._generate_enhanced_documents(selected, company_prefix)
            return

        output_mode = self.mode_combo.currentIndex()

        if output_mode == 0:
            # 按条款逐个输出 - 选择输出目录
            output_dir = QFileDialog.getExistingDirectory(
                self, "选择输出目录", "",
                QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
            )
            if not output_dir:
                return
            self._generate_individual_docs(selected, output_dir)

        elif output_mode == 1:
            # 按分类合并输出 - 选择输出目录
            output_dir = QFileDialog.getExistingDirectory(
                self, "选择输出目录", "",
                QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
            )
            if not output_dir:
                return
            self._generate_category_docs(selected, output_dir)

        else:
            # 全部合并为一个文档
            save_path, _ = QFileDialog.getSaveFileName(
                self, "保存Word文档",
                f"条款汇总_{datetime.now():%Y%m%d_%H%M}.docx",
                "Word文档 (*.docx)"
            )
            if not save_path:
                return
            self._generate_combined_doc(selected, save_path)

    def _generate_individual_docs(self, clauses: list, output_dir: str):
        """按条款逐个生成Word文档"""
        self.progress_bar.setVisible(True)
        self._log(f"📄 开始生成 {len(clauses)} 个独立文档...", "info")

        success_count = 0
        for i, clause in enumerate(clauses):
            progress = int((i + 1) / len(clauses) * 100)
            self.progress_bar.setValue(progress)
            QApplication.processEvents()

            try:
                # 清理文件名
                safe_name = re.sub(r'[\\/*?:"<>|]', '_', clause['name'])[:50]
                file_path = os.path.join(output_dir, f"{safe_name}.docx")

                doc = self._create_clause_document(clause)
                doc.save(file_path)
                success_count += 1

            except Exception as e:
                self._log(f"  ✗ {clause['name']}: {str(e)}", "error")

        self.progress_bar.setValue(100)
        self._log(f"✅ 完成! 成功生成 {success_count}/{len(clauses)} 个文档", "success")
        self._log(f"   输出目录: {output_dir}", "info")
        self.progress_bar.setVisible(False)

        # 打开输出目录
        try:
            os.startfile(output_dir)
        except Exception:
            pass

    def _generate_category_docs(self, clauses: list, output_dir: str):
        """按分类生成Word文档"""
        self.progress_bar.setVisible(True)

        # 按分类分组
        categorized = defaultdict(list)
        for clause in clauses:
            cat = clause.get('category', '其他') or '其他'
            categorized[cat].append(clause)

        self._log(f"📄 按 {len(categorized)} 个分类生成文档...", "info")

        total = len(categorized)
        for i, (category, cat_clauses) in enumerate(categorized.items()):
            progress = int((i + 1) / total * 100)
            self.progress_bar.setValue(progress)
            QApplication.processEvents()

            try:
                safe_cat = re.sub(r'[\\/*?:"<>|]', '_', category)[:30]
                file_path = os.path.join(output_dir, f"{safe_cat}_条款汇总.docx")

                doc = self._create_category_document(category, cat_clauses)
                doc.save(file_path)
                self._log(f"  ✓ {category}: {len(cat_clauses)} 条条款", "success")

            except Exception as e:
                self._log(f"  ✗ {category}: {str(e)}", "error")

        self.progress_bar.setValue(100)
        self._log(f"✅ 完成! 输出目录: {output_dir}", "success")
        self.progress_bar.setVisible(False)

        # 打开输出目录
        try:
            os.startfile(output_dir)
        except Exception:
            pass

    def _set_run_font(self, run, size_pt: float, bold: bool = False, color_rgb=None):
        """设置run的字体：宋体(中文) + Times New Roman(英文)

        v18.15: 支持float类型字号（如10.5pt = 5号字）
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        run.font.size = Pt(size_pt)
        run.font.name = 'Times New Roman'  # 英文字体
        run.bold = bold

        # 设置中文字体为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        if color_rgb:
            run.font.color.rgb = color_rgb

    def _add_formatted_text_to_paragraph(self, para, text: str, size_pt: float, base_bold: bool = False, color_rgb=None):
        """v18.16: 向段落添加带格式的文本，解析<b>标记应用加粗

        支持处理不完整的<b>标签（当加粗文本跨越换行被分割时）:
        - 文本以 </b> 开头：开头部分应加粗
        - 文本以 <b> 结尾无配对 </b>：结尾部分应加粗

        Args:
            para: Word段落对象
            text: 可能包含<b>标记的文本
            size_pt: 字号
            base_bold: 基础是否加粗（<b>标记外的文本）
            color_rgb: 颜色
        """
        if not text:
            return

        # v18.16: 处理不完整的标签
        # 步骤1: 检查是否以 </b> 开头（orphan closing tag）
        orphan_close_match = re.match(r'^(.*?)</b>', text, re.DOTALL)
        orphan_close_text = None
        if orphan_close_match:
            # 确保这不是完整标签对的一部分
            before_close = text[:orphan_close_match.end()]
            if '<b>' not in before_close:
                orphan_close_text = orphan_close_match.group(1)
                text = text[orphan_close_match.end():]

        # 步骤2: 检查是否以 <b> 结尾无配对（orphan opening tag）
        orphan_open_match = re.search(r'<b>([^<]*)$', text, re.DOTALL)
        orphan_open_text = None
        if orphan_open_match:
            orphan_open_text = orphan_open_match.group(1)
            text = text[:orphan_open_match.start()]

        # 步骤3: 添加开头的孤立加粗部分
        if orphan_close_text:
            run = para.add_run(orphan_close_text)
            self._set_run_font(run, size_pt, bold=True, color_rgb=color_rgb)

        # 步骤4: 处理中间的完整 <b>...</b> 标记
        pattern = re.compile(r'<b>(.*?)</b>', re.DOTALL)
        last_end = 0

        for match in pattern.finditer(text):
            # 添加非加粗部分
            if match.start() > last_end:
                normal_text = text[last_end:match.start()]
                if normal_text:
                    run = para.add_run(normal_text)
                    self._set_run_font(run, size_pt, bold=base_bold, color_rgb=color_rgb)

            # 添加加粗部分
            bold_text = match.group(1)
            if bold_text:
                run = para.add_run(bold_text)
                self._set_run_font(run, size_pt, bold=True, color_rgb=color_rgb)

            last_end = match.end()

        # 添加最后的非加粗部分（在孤立开头标签之前）
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                run = para.add_run(remaining)
                self._set_run_font(run, size_pt, bold=base_bold, color_rgb=color_rgb)

        # 步骤5: 添加结尾的孤立加粗部分
        if orphan_open_text:
            run = para.add_run(orphan_open_text)
            self._set_run_font(run, size_pt, bold=True, color_rgb=color_rgb)

    def _generate_combined_doc(self, clauses: list, save_path: str):
        """生成合并的Word文档 - v18.15格式：宋体+Times New Roman, 5号字, 两端对齐, 单倍行距, 段后0.5行"""
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(20)

        try:
            self._log(f"📄 生成合并文档，共 {len(clauses)} 条条款...", "info")

            doc = Document()

            from docx.shared import Pt, RGBColor, Twips
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.oxml.ns import qn

            # v18.15: 固定格式参数 - 5号字=10.5pt, 两端对齐, 单倍行距, 段后0.5行≈120twips
            BODY_SIZE = 10.5  # 5号字
            TITLE_SIZE = 10.5  # 标题也用5号字
            SPACE_AFTER_HALF_LINE = 120  # 0.5行 (twips)

            # 文档标题
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run('条款汇总清单')
            self._set_run_font(title_run, 16, bold=True)  # 三号字
            title_para.paragraph_format.space_after = Twips(400)

            # 生成日期
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f"生成日期: {datetime.now():%Y年%m月%d日}")
            self._set_run_font(date_run, 10, color_rgb=RGBColor(128, 128, 128))
            date_para.paragraph_format.space_after = Twips(200)

            self.progress_bar.setValue(40)

            # 按分类组织
            categorized = defaultdict(list)
            for clause in clauses:
                cat = clause.get('category', '其他') or '其他'
                categorized[cat].append(clause)

            clause_num = 1
            for category, cat_clauses in categorized.items():
                # 分类标题
                cat_para = doc.add_paragraph()
                cat_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                cat_run = cat_para.add_run(f"【{category}】")
                self._set_run_font(cat_run, TITLE_SIZE, bold=True, color_rgb=RGBColor(217, 119, 87))
                cat_para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE)

                for clause in cat_clauses:
                    # 条款前空行（除第一条外）
                    if clause_num > 1:
                        blank_para = doc.add_paragraph()
                        blank_para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE)

                    # 条款名称（加粗，无下划线）
                    name_para = doc.add_paragraph()
                    name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    name_run = name_para.add_run(f"{clause_num}. {clause['name']}")
                    self._set_run_font(name_run, TITLE_SIZE, bold=True)
                    name_para.paragraph_format.space_after = Twips(60)
                    name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                    # 注册号 - v18.15: 直接输出，不添加额外前缀（数据已包含"注册号"或"产品注册号"）
                    if self.include_reg_check.isChecked() and clause.get('regNo'):
                        reg_para = doc.add_paragraph()
                        reg_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        reg_run = reg_para.add_run(clause['regNo'])  # 直接输出，不加前缀
                        self._set_run_font(reg_run, BODY_SIZE)
                        reg_para.paragraph_format.space_after = Twips(60)
                        reg_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                    # 条款内容 - v18.15: 支持<b>标记保留加粗格式
                    if clause.get('content'):
                        content_lines = clause['content'].split('\n')
                        for i, para_text in enumerate(content_lines):
                            para_text = para_text.strip()
                            if para_text:
                                content_para = doc.add_paragraph()
                                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                # 使用格式化方法处理可能包含<b>标记的文本
                                self._add_formatted_text_to_paragraph(content_para, para_text, BODY_SIZE)
                                # 最后一行段后0.5行，其他行无段后
                                is_last_line = (i == len(content_lines) - 1)
                                content_para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE if is_last_line else 0)
                                content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                    clause_num += 1

            self.progress_bar.setValue(80)

            doc.save(save_path)

            self.progress_bar.setValue(100)
            self._log(f"✅ Word文档已生成: {os.path.basename(save_path)}", "success")
            self._log(f"   共导出 {len(clauses)} 条条款，{len(categorized)} 个分类", "info")

            # 打开生成的文档
            try:
                os.startfile(save_path)
            except Exception:
                pass

        except Exception as e:
            self._log(f"❌ 生成失败: {sanitize_error_message(e)}", "error")
            logger.error(f"生成Word文档失败: {e}")  # 完整错误记录到日志
        finally:
            self.progress_bar.setVisible(False)

    def _create_clause_document(self, clause: dict) -> Document:
        """创建单个条款的Word文档 - v18.15格式：宋体+Times New Roman, 5号字, 两端对齐, 单倍行距"""
        from docx.shared import RGBColor, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

        doc = Document()

        # v18.15: 固定格式参数
        BODY_SIZE = 10.5  # 5号字
        TITLE_SIZE = 10.5
        SPACE_AFTER_HALF_LINE = 120  # 0.5行

        # 条款名称（居中，加粗）
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(clause['name'])
        self._set_run_font(title_run, TITLE_SIZE, bold=True)
        title_para.paragraph_format.space_after = Twips(60)
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 注册号 - 直接输出，不添加额外前缀
        if self.include_reg_check.isChecked() and clause.get('regNo'):
            reg_para = doc.add_paragraph()
            reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            reg_run = reg_para.add_run(clause['regNo'])
            self._set_run_font(reg_run, BODY_SIZE)
            reg_para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE)
            reg_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 条款内容 - v18.15: 支持<b>标记保留加粗格式
        if clause.get('content'):
            content_lines = clause['content'].split('\n')
            for i, line in enumerate(content_lines):
                line = line.strip()
                if line:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # 使用格式化方法处理可能包含<b>标记的文本
                    self._add_formatted_text_to_paragraph(para, line, BODY_SIZE)
                    is_last_line = (i == len(content_lines) - 1)
                    para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE if is_last_line else 0)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        return doc

    def _create_category_document(self, category: str, clauses: list) -> Document:
        """创建分类条款文档 - v18.15格式：宋体+Times New Roman, 5号字, 两端对齐, 单倍行距"""
        from docx.shared import RGBColor, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

        doc = Document()

        # v18.15: 固定格式参数
        BODY_SIZE = 10.5  # 5号字
        TITLE_SIZE = 10.5
        SPACE_AFTER_HALF_LINE = 120  # 0.5行

        # 分类标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"【{category}】条款汇总")
        self._set_run_font(title_run, 14, bold=True)  # 标题稍大
        title_para.paragraph_format.space_after = Twips(200)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"共 {len(clauses)} 条 · {datetime.now():%Y-%m-%d}")
        self._set_run_font(date_run, 10, color_rgb=RGBColor(128, 128, 128))
        date_para.paragraph_format.space_after = Twips(200)

        for i, clause in enumerate(clauses, 1):
            # 条款前空行（除第一条外）
            if i > 1:
                blank_para = doc.add_paragraph()
                blank_para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE)

            # 条款名称（加粗，无下划线）
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            name_run = name_para.add_run(f"{i}. {clause['name']}")
            self._set_run_font(name_run, TITLE_SIZE, bold=True)
            name_para.paragraph_format.space_after = Twips(60)
            name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 注册号 - 直接输出，不添加额外前缀
            if self.include_reg_check.isChecked() and clause.get('regNo'):
                reg_para = doc.add_paragraph()
                reg_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                reg_run = reg_para.add_run(clause['regNo'])
                self._set_run_font(reg_run, BODY_SIZE)
                reg_para.paragraph_format.space_after = Twips(60)
                reg_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # 条款内容 - v18.15: 支持<b>标记保留加粗格式
            if clause.get('content'):
                content_lines = clause['content'].split('\n')
                for j, line in enumerate(content_lines):
                    line = line.strip()
                    if line:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # 使用格式化方法处理可能包含<b>标记的文本
                        self._add_formatted_text_to_paragraph(para, line, BODY_SIZE)
                        is_last_line = (j == len(content_lines) - 1)
                        para.paragraph_format.space_after = Twips(SPACE_AFTER_HALF_LINE if is_last_line else 0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        return doc

    def _log(self, message: str, level: str = "info"):
        """添加日志"""
        colors = {
            'info': '#e0e0e0',
            'success': '#7ec9a0',
            'warning': '#e5c07b',
            'error': '#e06c75'
        }
        color = colors.get(level, '#e0e0e0')
        self.log_text.append(f'<span style="color: {color}">{message}</span>')


class ClauseComparisonAssistant(QMainWindow):
    """主界面 - Anthropic 风格 - V18.0 Tab版"""
    def __init__(self):
        super().__init__()
        self.setWindowTitle("智能条款工具箱 V18.0")
        self.setMinimumSize(860, 700)
        # Anthropic 风格：温暖的奶油白背景
        self.setStyleSheet(f"""
            QMainWindow {{
                background: {AnthropicColors.BG_PRIMARY};
            }}
        """)

        if HAS_CONFIG_MANAGER:
            self._config = get_config()
        else:
            self._config = None

        # 初始化映射管理器
        if HAS_MAPPING_MANAGER:
            self._mapping_manager = get_mapping_manager()
            self._mapping_manager.load()
            # 应用用户映射到配置
            if self._config:
                self._mapping_manager.apply_to_config(self._config)
        else:
            self._mapping_manager = None

        self._setup_ui()

    def _setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 12, 20, 12)

        # 标题行 - Anthropic 风格
        header_layout = QHBoxLayout()

        title = QLabel("🔧 智能条款工具箱")
        title.setStyleSheet(f"color: {AnthropicColors.TEXT_PRIMARY}; font-size: 22px; font-weight: bold;")
        header_layout.addWidget(title)

        header_layout.addStretch()

        # 版本信息
        subtitle = QLabel("V18.9 · 条款提取 · 条款比对 · 条款输出 · 保险计算")
        subtitle.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 13px;")
        header_layout.addWidget(subtitle)

        # 支持作者按钮 - Anthropic 强调色风格
        self.donate_btn = QPushButton("💝 支持作者")
        self.donate_btn.setCursor(Qt.PointingHandCursor)
        self.donate_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.ACCENT};
                color: {AnthropicColors.TEXT_LIGHT};
                border: none;
                border-radius: 15px;
                padding: 6px 16px;
                font-size: 12px;
                font-weight: 500;
                margin-left: 15px;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.ACCENT_DARK};
            }}
        """)
        self.donate_btn.clicked.connect(self._show_donate_dialog)

        # 为支持作者按钮添加柔和阴影
        donate_shadow = QGraphicsDropShadowEffect()
        donate_shadow.setBlurRadius(12)
        donate_shadow.setColor(QColor(217, 119, 87, 80))  # Anthropic ACCENT 色
        donate_shadow.setOffset(0, 2)
        self.donate_btn.setGraphicsEffect(donate_shadow)

        # 呼吸动画定时器
        self._donate_glow_step = 0
        self._donate_timer = QTimer(self)
        self._donate_timer.timeout.connect(self._animate_donate_button)
        self._donate_timer.start(50)  # 50ms间隔

        header_layout.addWidget(self.donate_btn)
        layout.addLayout(header_layout)

        # ==========================================
        # 主Tab区域 - Anthropic风格
        # ==========================================
        self.main_tabs = QTabWidget()
        self.main_tabs.setStyleSheet(f"""
            QTabWidget::pane {{
                border: none;
                background: {AnthropicColors.BG_PRIMARY};
            }}
            QTabBar::tab {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_SECONDARY};
                border: none;
                padding: 10px 28px;
                margin-right: 8px;
                border-radius: 8px 8px 0 0;
                font-size: 14px;
                font-weight: 600;
                min-width: 110px;
            }}
            QTabBar::tab:selected {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT};
            }}
            QTabBar::tab:hover:!selected {{
                background: {AnthropicColors.BG_CARD};
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
        """)

        # Tab 1: 条款提取
        self.extractor_tab = ClauseExtractorTab(self)
        self.main_tabs.addTab(self.extractor_tab, "📄 条款提取")

        # Tab 2: 条款比对
        self.comparison_tab = self._create_comparison_tab()
        self.main_tabs.addTab(self.comparison_tab, "🔍 条款比对")

        # Tab 3: 条款输出
        self.output_tab = ClauseOutputTab(self)
        self.main_tabs.addTab(self.output_tab, "📝 条款输出")

        # Tab 4 & 5: 保险计算器（如已安装）
        if HAS_INSURANCE_CALC:
            self.main_insurance_tab = MainInsuranceTab(self)
            self.main_tabs.addTab(self.main_insurance_tab, "🧮 主险计算")

            self.addon_insurance_tab = AddonInsuranceTab(self)
            self.main_tabs.addTab(self.addon_insurance_tab, "📋 附加险计算")

            # 连接信号：主险计算结果 → 附加险
            self.main_insurance_tab.premium_calculated.connect(
                self.addon_insurance_tab.receive_main_premium
            )
            self.main_insurance_tab.full_result_calculated.connect(
                self.addon_insurance_tab.receive_full_data
            )

        layout.addWidget(self.main_tabs, 1)

        # 版本信息
        version = QLabel("V18.9 Insurance Calculator Edition · Made with ❤️ by Dachi Yijin")
        version.setAlignment(Qt.AlignCenter)
        version.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 11px;")
        layout.addWidget(version)

    def _create_comparison_tab(self) -> QWidget:
        """创建条款比对Tab"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(8)
        layout.setContentsMargins(15, 10, 15, 10)

        # 滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setStyleSheet(get_anthropic_scrollbar_style())
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setSpacing(8)
        scroll_layout.setContentsMargins(0, 0, 0, 0)

        # 配置统计
        if self._config:
            stats = self._config.get_stats()
            user_mappings = self._mapping_manager.get_mapping_count() if self._mapping_manager else 0
            stats_text = f"📊 {stats['client_mappings']} 映射 | {user_mappings} 自定义 | {stats['semantic_aliases']} 别名"
        else:
            stats_text = "📊 使用DCYJIN智能AI配置"
        self.stats_label = QLabel(stats_text)
        self.stats_label.setAlignment(Qt.AlignCenter)
        self.stats_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 13px;")
        scroll_layout.addWidget(self.stats_label)

        # 输入卡片 - Anthropic 风格
        card = GlassCard()
        card_layout = QVBoxLayout(card)
        card_layout.setSpacing(12)
        card_layout.setContentsMargins(20, 20, 20, 20)

        # Anthropic 风格的输入框样式
        style = f"""
            QLabel {{ color: {AnthropicColors.TEXT_PRIMARY}; font-weight: 500; }}
            QLineEdit {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 14px 16px;
                color: {AnthropicColors.TEXT_PRIMARY}; font-size: 15px;
            }}
            QLineEdit:focus {{ border-color: {AnthropicColors.ACCENT}; }}
        """
        card.setStyleSheet(card.styleSheet() + style)

        btn_style = f"""
            QPushButton {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px; padding: 14px 20px;
                color: {AnthropicColors.TEXT_PRIMARY}; font-weight: 500; font-size: 15px;
            }}
            QPushButton:hover {{ background: {AnthropicColors.BG_CARD}; border-color: {AnthropicColors.ACCENT}; }}
        """

        self.doc_input = self._create_file_row(card_layout, "📂 客户文档",
            "Word 条款清单 (.docx)", "Word Files (*.docx)", btn_style)
        self.lib_input = self._create_file_row(card_layout, "📚 条款库",
            "Excel 条款库 (.xlsx)", "Excel Files (*.xlsx)", btn_style)

        # 添加Sheet选择行
        sheet_row = QHBoxLayout()
        sheet_label = QLabel("📋 险种Sheet")
        sheet_label.setMinimumWidth(90)
        self.sheet_combo = QComboBox()
        self.sheet_combo.setMinimumHeight(38)
        self.sheet_combo.setStyleSheet(f"""
            QComboBox {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                padding: 10px 14px;
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
            }}
            QComboBox:hover {{ border-color: {AnthropicColors.ACCENT}; }}
            QComboBox::drop-down {{
                border: none;
                width: 30px;
            }}
            QComboBox::down-arrow {{
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid {AnthropicColors.ACCENT};
            }}
            QComboBox QAbstractItemView {{
                background: {AnthropicColors.BG_PRIMARY};
                border: 1px solid {AnthropicColors.ACCENT};
                selection-background-color: {AnthropicColors.ACCENT};
                color: {AnthropicColors.TEXT_PRIMARY};
            }}
        """)
        self.sheet_combo.addItem("自动选择第一个Sheet")
        self.sheet_combo.setToolTip("选择条款库中的险种Sheet（如财产险、责任险等）")
        # 当条款库文件改变时更新Sheet列表
        self.lib_input.textChanged.connect(self._update_sheet_list)
        sheet_row.addWidget(sheet_label)
        sheet_row.addWidget(self.sheet_combo, 1)
        card_layout.addLayout(sheet_row)

        line = QFrame()
        line.setFixedHeight(2)
        line.setStyleSheet(f"background: {AnthropicColors.BORDER};")
        card_layout.addWidget(line)

        row3 = QHBoxLayout()
        label3 = QLabel("💾 保存路径")
        label3.setMinimumWidth(90)
        self.out_input = QLineEdit()
        self.out_input.setPlaceholderText("报告保存位置...")
        btn3 = QPushButton("选择")
        btn3.setCursor(Qt.PointingHandCursor)
        btn3.setStyleSheet(btn_style)
        btn3.clicked.connect(self._browse_save)
        row3.addWidget(label3)
        row3.addWidget(self.out_input, 1)
        row3.addWidget(btn3)
        card_layout.addLayout(row3)

        scroll_layout.addWidget(card)

        # v18.3: 匹配模式选择
        mode_layout = QHBoxLayout()
        mode_layout.setSpacing(12)

        mode_label = QLabel("匹配模式：")
        mode_label.setStyleSheet(f"color: {AnthropicColors.TEXT_SECONDARY}; font-size: 15px;")

        self.match_mode_combo = QComboBox()
        self.match_mode_combo.addItems(["🔄 自动检测（推荐）", "📝 纯标题模式", "📄 完整内容模式"])
        self.match_mode_combo.setMinimumHeight(36)
        self.match_mode_combo.setMinimumWidth(220)
        self.match_mode_combo.setCursor(Qt.PointingHandCursor)
        self.match_mode_combo.setStyleSheet(f"""
            QComboBox {{
                padding: 10px 14px;
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 8px;
                background: {AnthropicColors.BG_PRIMARY};
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
            }}
            QComboBox:hover {{
                border-color: {AnthropicColors.ACCENT};
            }}
            QComboBox::drop-down {{
                border: none;
                width: 24px;
            }}
            QComboBox::down-arrow {{
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 6px solid {AnthropicColors.TEXT_MUTED};
                margin-right: 8px;
            }}
            QComboBox QAbstractItemView {{
                background: {AnthropicColors.BG_PRIMARY};
                color: {AnthropicColors.TEXT_PRIMARY};
                selection-background-color: {AnthropicColors.BG_CARD};
                selection-color: {AnthropicColors.TEXT_PRIMARY};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 4px;
                padding: 4px;
            }}
        """)

        self.mode_hint_label = QLabel("")
        self.mode_hint_label.setStyleSheet(f"color: {AnthropicColors.TEXT_MUTED}; font-size: 14px;")

        # v18.9: 精准识别模式勾选框
        self.precise_mode_checkbox = QCheckBox("🎯 精准识别")
        self.precise_mode_checkbox.setToolTip("仅提取蓝色字体的条款\n适用于干扰项较多的文档")
        self.precise_mode_checkbox.setCursor(Qt.PointingHandCursor)
        self.precise_mode_checkbox.setStyleSheet(f"""
            QCheckBox {{
                color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px;
                spacing: 8px;
            }}
            QCheckBox::indicator {{
                width: 20px;
                height: 20px;
                border: 2px solid {AnthropicColors.BORDER};
                border-radius: 4px;
                background: {AnthropicColors.BG_PRIMARY};
            }}
            QCheckBox::indicator:hover {{
                border-color: {AnthropicColors.ACCENT};
            }}
            QCheckBox::indicator:checked {{
                background: {AnthropicColors.ACCENT};
                border-color: {AnthropicColors.ACCENT};
            }}
            QCheckBox::indicator:checked::after {{
                content: "✓";
            }}
        """)

        mode_layout.addWidget(mode_label)
        mode_layout.addWidget(self.match_mode_combo)
        mode_layout.addWidget(self.mode_hint_label)
        mode_layout.addSpacing(20)
        mode_layout.addWidget(self.precise_mode_checkbox)
        mode_layout.addStretch()
        scroll_layout.addLayout(mode_layout)

        # 按钮行
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        self.start_btn = QPushButton("🚀 开始比对")
        self.start_btn.setCursor(Qt.PointingHandCursor)
        self.start_btn.setMinimumHeight(44)
        self.start_btn.setStyleSheet(f"""
            QPushButton {{
                background: {AnthropicColors.BG_DARK};
                color: {AnthropicColors.TEXT_LIGHT}; font-size: 16px; font-weight: bold;
                border-radius: 12px; border: none;
            }}
            QPushButton:hover {{
                background: {AnthropicColors.ACCENT};
            }}
            QPushButton:disabled {{ background: {AnthropicColors.BORDER_DARK}; color: {AnthropicColors.TEXT_MUTED}; }}
        """)
        self.start_btn.clicked.connect(self._start_process)

        # v18.4: 取消比对按钮（替代原批量处理按钮）
        self.cancel_btn = QPushButton("⛔ 取消比对")
        self.cancel_btn.setCursor(Qt.PointingHandCursor)
        self.cancel_btn.setMinimumHeight(44)
        self.cancel_btn.setEnabled(False)  # 默认禁用
        self.cancel_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent; color: {AnthropicColors.ERROR};
                font-size: 15px; font-weight: 500;
                border-radius: 8px; border: 1px solid {AnthropicColors.ERROR};
            }}
            QPushButton:hover {{ background: {AnthropicColors.ERROR}; color: {AnthropicColors.TEXT_LIGHT}; }}
            QPushButton:disabled {{ color: {AnthropicColors.BORDER}; border-color: {AnthropicColors.BORDER}; }}
        """)
        self.cancel_btn.clicked.connect(self._cancel_process)

        # 普通按钮样式
        normal_btn_style = f"""
            QPushButton {{
                background: transparent; color: {AnthropicColors.TEXT_PRIMARY};
                font-size: 15px; font-weight: 500;
                border-radius: 8px; border: 1px solid {AnthropicColors.BG_DARK};
            }}
            QPushButton:hover {{ background: {AnthropicColors.BG_DARK}; color: {AnthropicColors.TEXT_LIGHT}; }}
        """

        self.add_btn = QPushButton("🔧 映射设置")
        self.add_btn.setCursor(Qt.PointingHandCursor)
        self.add_btn.setMinimumHeight(44)
        self.add_btn.setStyleSheet(normal_btn_style)
        self.add_btn.clicked.connect(self._show_add_mapping_dialog)

        # v17.1: 条款查询按钮
        self.query_btn = QPushButton("🔍 条款查询")
        self.query_btn.setCursor(Qt.PointingHandCursor)
        self.query_btn.setMinimumHeight(44)
        self.query_btn.setStyleSheet(normal_btn_style)
        self.query_btn.clicked.connect(self._show_query_dialog)

        self.open_btn = QPushButton("📂 打开目录")
        self.open_btn.setCursor(Qt.PointingHandCursor)
        self.open_btn.setMinimumHeight(44)
        self.open_btn.setEnabled(False)
        self.open_btn.setStyleSheet(f"""
            QPushButton {{
                background: transparent; color: {AnthropicColors.TEXT_SECONDARY};
                font-size: 15px; font-weight: 500;
                border-radius: 8px; border: 1px solid {AnthropicColors.BORDER};
            }}
            QPushButton:hover {{ border-color: {AnthropicColors.ACCENT}; color: {AnthropicColors.ACCENT}; }}
            QPushButton:disabled {{ color: {AnthropicColors.BORDER}; border-color: {AnthropicColors.BORDER}; }}
        """)
        self.open_btn.clicked.connect(self._open_output_folder)

        btn_layout.addWidget(self.start_btn, 3)
        btn_layout.addWidget(self.cancel_btn, 1)
        btn_layout.addWidget(self.add_btn, 1)
        btn_layout.addWidget(self.query_btn, 1)  # v17.1: 条款查询
        btn_layout.addWidget(self.open_btn, 1)
        scroll_layout.addLayout(btn_layout)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.setStyleSheet(f"""
            QProgressBar {{ background: {AnthropicColors.BORDER}; border-radius: 3px; }}
            QProgressBar::chunk {{
                background: {AnthropicColors.ACCENT};
                border-radius: 3px;
            }}
        """)
        scroll_layout.addWidget(self.progress_bar)

        # 完成滚动区域设置
        scroll_area.setWidget(scroll_widget)
        layout.addWidget(scroll_area, 1)

        # 日志
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background: {AnthropicColors.BG_CARD};
                border: 1px solid {AnthropicColors.BORDER};
                border-radius: 12px; color: {AnthropicColors.TEXT_PRIMARY};
                padding: 15px;
                font-family: 'JetBrains Mono', 'Consolas', 'Cascadia Code', monospace;
                font-size: 13px;
            }}
        """)
        layout.addWidget(self.log_text, 1)

        return tab

    def _create_file_row(self, layout, label_text: str, placeholder: str,
                         filter_str: str, btn_style: str) -> QLineEdit:
        row = QHBoxLayout()
        label = QLabel(label_text)
        label.setMinimumWidth(90)
        line_edit = QLineEdit()
        line_edit.setPlaceholderText(placeholder)
        btn = QPushButton("浏览")
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet(btn_style)
        btn.clicked.connect(lambda: self._browse_file(line_edit, filter_str))
        row.addWidget(label)
        row.addWidget(line_edit, 1)
        row.addWidget(btn)
        layout.addLayout(row)
        return line_edit

    def _browse_file(self, line_edit: QLineEdit, filter_str: str):
        f, _ = QFileDialog.getOpenFileName(self, "选择文件", "", filter_str)
        if f:
            line_edit.setText(f)
            if line_edit == self.doc_input and not self.out_input.text():
                # v18.9: 文件名加上日期，避免覆盖
                date_str = datetime.now().strftime("%Y%m%d")
                self.out_input.setText(os.path.join(os.path.dirname(f), f"条款比对报告_{date_str}.xlsx"))

    def _browse_save(self):
        # v18.9: 文件名加上日期，避免覆盖
        date_str = datetime.now().strftime("%Y%m%d")
        f, _ = QFileDialog.getSaveFileName(self, "保存结果", f"条款比对报告_{date_str}.xlsx", "Excel Files (*.xlsx)")
        if f:
            self.out_input.setText(f)

    def changeEvent(self, event):
        """窗口状态变化时暂停/恢复动画定时器"""
        if event.type() == event.WindowStateChange:
            if self.windowState() & Qt.WindowMinimized:
                self._donate_timer.stop()
            else:
                self._donate_timer.start(50)
        super().changeEvent(event)

    def _animate_donate_button(self):
        """支持作者按钮的呼吸发光动画"""
        import math
        self._donate_glow_step = (self._donate_glow_step + 1) % 120

        # 使用正弦函数创建平滑的呼吸效果
        glow_intensity = int(80 + 70 * math.sin(self._donate_glow_step * math.pi / 60))
        blur_radius = int(12 + 8 * math.sin(self._donate_glow_step * math.pi / 60))

        effect = self.donate_btn.graphicsEffect()
        if effect and isinstance(effect, QGraphicsDropShadowEffect):
            effect.setBlurRadius(blur_radius)
            effect.setColor(QColor(217, 119, 87, glow_intensity))  # Anthropic accent color

    def _show_donate_dialog(self):
        """显示支持作者对话框"""
        dialog = DonateDialog(self)
        dialog.exec_()

    def _show_add_mapping_dialog(self):
        """打开条款映射管理对话框"""
        if HAS_MAPPING_MANAGER:
            # 获取当前条款库中的条款名称列表（用于下拉提示）
            library_clauses = self._get_library_clauses()

            dialog = ClauseMappingDialog(self, library_clauses=library_clauses)
            dialog.mappings_changed.connect(self._on_mappings_changed)
            dialog.exec_()
        elif self._config:
            # 兼容旧版：使用简单的添加对话框
            dialog = AddMappingDialog(self)
            if dialog.exec_() == QDialog.Accepted:
                eng, chn = dialog.get_mapping()
                if eng and chn:
                    self._config.add_client_mapping(eng, chn)
                    self._config.save()
                    self._append_log(f"✓ 已添加映射: '{eng}' -> '{chn}'", "success")
        else:
            QMessageBox.warning(self, "提示", "映射管理功能不可用")

    def _show_query_dialog(self):
        """v17.1: 打开条款查询对话框"""
        library_path = self.lib_input.text().strip()
        if not library_path or not os.path.exists(library_path):
            QMessageBox.warning(self, "提示", "请先选择条款库文件！")
            return

        try:
            # 加载条款库并构建索引
            logic = ClauseMatcherLogic()
            sheet_name = self._get_selected_sheet()
            lib_data = LibraryLoader.load_excel(library_path, sheet_name=sheet_name)
            library_index = logic.build_index(lib_data)

            # 获取映射管理器
            mapping_mgr = get_mapping_manager() if HAS_MAPPING_MANAGER else None

            # 打开查询对话框
            dialog = ClauseQueryDialog(
                parent=self,
                library_index=library_index,
                logic=logic,
                mapping_mgr=mapping_mgr
            )
            dialog.exec_()
        except Exception as e:
            QMessageBox.warning(self, "错误", f"加载条款库失败: {sanitize_error_message(e)}")

    def _get_library_clauses(self) -> List[str]:
        """从当前条款库获取条款名称列表"""
        library_path = self.lib_input.text().strip()
        if not library_path or not os.path.exists(library_path):
            return []

        try:
            clauses = []
            wb = openpyxl.load_workbook(library_path, read_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(max_col=1, values_only=True):
                    if row[0] and isinstance(row[0], str):
                        name = row[0].strip()
                        if name and len(name) > 3 and name not in clauses:
                            clauses.append(name)
            wb.close()
            return clauses[:500]  # 限制数量防止内存问题
        except Exception as e:
            logger.warning(f"读取条款库失败: {e}")
            return []

    def _on_mappings_changed(self):
        """映射变更回调：更新配置"""
        if HAS_MAPPING_MANAGER and self._config:
            mapping_manager = get_mapping_manager()
            count = mapping_manager.apply_to_config(self._config)
            self._append_log(f"✓ 已应用 {count} 条用户映射", "success")

            # 更新统计显示
            stats = self._config.get_stats()
            user_mappings = mapping_manager.get_mapping_count()
            self.stats_label.setText(f"📊 {stats['client_mappings']} 映射 | {user_mappings} 自定义 | {stats['semantic_aliases']} 别名")

    def _show_batch_dialog(self):
        if not self.lib_input.text():
            QMessageBox.warning(self, "提示", "请先选择条款库")
            return

        dialog = BatchSelectDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            files = dialog.get_files()
            if not files:
                return

            output_dir = QFileDialog.getExistingDirectory(
                self, "选择输出目录", "",
                QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
            )
            if not output_dir:
                return

            self._start_batch_process(files, output_dir)

    def _append_log(self, msg: str, level: str):
        colors = {
            "info": AnthropicColors.TEXT_SECONDARY,
            "success": AnthropicColors.SUCCESS,
            "error": AnthropicColors.ERROR,
            "warning": AnthropicColors.WARNING
        }
        self.log_text.append(f'<span style="color:{colors.get(level, AnthropicColors.TEXT_PRIMARY)}">{msg}</span>')
        self.log_text.moveCursor(QTextCursor.End)

    def _start_process(self):
        doc = self.doc_input.text().strip()
        excel = self.lib_input.text().strip()
        out = self.out_input.text().strip()

        if not all([doc, excel, out]):
            QMessageBox.warning(self, "提示", "请完善所有文件路径！")
            return

        self._set_ui_state(False)
        self.log_text.clear()

        # 获取选择的Sheet名称
        sheet_name = self._get_selected_sheet()

        # v18.3: 获取选择的匹配模式
        match_mode = self._get_match_mode()

        # v18.9: 获取精准识别模式
        precise_mode = self.precise_mode_checkbox.isChecked()

        self.worker = MatchWorker(doc, excel, out, sheet_name, match_mode, precise_mode)
        self.worker.log_signal.connect(self._append_log)
        self.worker.progress_signal.connect(lambda c, t: self.progress_bar.setValue(int(c/t*100)))
        self.worker.finished_signal.connect(self._on_finished)
        self.worker.start()

    def _start_batch_process(self, files: List[str], output_dir: str):
        self._set_ui_state(False)
        self.log_text.clear()

        # 获取选择的Sheet名称
        sheet_name = self._get_selected_sheet()

        # v18.3: 获取选择的匹配模式
        match_mode = self._get_match_mode()

        # v18.9: 获取精准识别模式
        precise_mode = self.precise_mode_checkbox.isChecked()

        self.batch_worker = BatchMatchWorker(files, self.lib_input.text(), output_dir, sheet_name, match_mode, precise_mode)
        self.batch_worker.log_signal.connect(self._append_log)
        self.batch_worker.batch_progress_signal.connect(
            lambda c, t, n: self.progress_bar.setValue(int(c/t*100))
        )
        self.batch_worker.finished_signal.connect(self._on_batch_finished)
        self.batch_worker.start()

    def _get_selected_sheet(self) -> Optional[str]:
        """获取选择的Sheet名称"""
        if self.sheet_combo.currentIndex() == 0:  # "自动选择第一个Sheet"
            return None
        return self.sheet_combo.currentText()

    def _get_match_mode(self) -> str:
        """v18.3: 获取选择的匹配模式"""
        idx = self.match_mode_combo.currentIndex()
        if idx == 0:
            return "auto"
        elif idx == 1:
            return "title"
        else:
            return "content"

    def _update_sheet_list(self, excel_path: str):
        """当条款库文件改变时更新Sheet列表"""
        self.sheet_combo.clear()
        self.sheet_combo.addItem("自动选择第一个Sheet")

        if not excel_path or not os.path.exists(excel_path):
            return

        try:
            sheets = LibraryLoader.get_sheet_names(excel_path)
            if sheets:
                for sheet in sheets:
                    self.sheet_combo.addItem(sheet)
                # 如果只有一个Sheet，保持默认选择
                if len(sheets) > 1:
                    self._append_log(f"📋 检测到 {len(sheets)} 个Sheet: {', '.join(sheets)}", "info")
        except Exception as e:
            logger.warning(f"读取Sheet列表失败: {e}")

    def _set_ui_state(self, enabled: bool):
        self.start_btn.setEnabled(enabled)
        self.start_btn.setText("🚀 开始比对" if enabled else "⏳ 处理中...")
        self.progress_bar.setVisible(not enabled)

        # v18.4: 取消按钮 - 空闲时禁用，处理中时启用
        self.cancel_btn.setEnabled(not enabled)
        if not enabled:
            self.progress_bar.setValue(0)

    def _cancel_process(self):
        """v18.4: 取消比对操作"""
        if hasattr(self, 'worker') and self.worker and self.worker.isRunning():
            self.worker.cancel()
            self._append_log("⏳ 正在取消...", "warning")
        elif hasattr(self, 'batch_worker') and self.batch_worker and self.batch_worker.isRunning():
            self.batch_worker.cancel()
            self._append_log("⏳ 正在取消批量处理...", "warning")

    def _on_finished(self, success: bool, msg: str):
        self._set_ui_state(True)
        if success:
            self.open_btn.setEnabled(True)
            self.open_btn.setStyleSheet("""
                QPushButton {
                    background: transparent; color: #2ecc71;
                    font-size: 14px; font-weight: 500;
                    border-radius: 26px; border: 2px solid #2ecc71;
                }
                QPushButton:hover { background: #2ecc71; color: white; }
            """)
            QMessageBox.information(self, "完成", f"比对完成！\n{msg}")

    def _on_batch_finished(self, success: bool, msg: str, ok_count: int, total: int):
        self._set_ui_state(True)
        if success:
            self.open_btn.setEnabled(True)
            QMessageBox.information(self, "完成", f"批量处理完成！\n成功: {ok_count}/{total}\n输出目录: {msg}")

    def _open_output_folder(self):
        path = self.out_input.text().strip()
        if path and os.path.exists(path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(os.path.dirname(path)))


def run_clause_test(doc_paths: List[Tuple[str, Optional[int], str]]):
    """
    运行条款识别回归测试
    doc_paths: [(文件路径, 期望条款数, 名称), ...]
    """
    print('=' * 60)
    print('条款识别回归测试 (主脚本)')
    print('=' * 60)

    # 使用ClauseMatcherLogic类来测试（parse_docx是它的方法）
    extractor = ClauseMatcherLogic()

    all_passed = True

    for doc_path, expected, name in doc_paths:
        if not os.path.exists(doc_path):
            print(f'{name}: ⚠️ 文件不存在 - {doc_path}')
            continue

        try:
            clauses, _ = extractor.parse_docx(doc_path)
            count = len(clauses)

            if expected is not None:
                status = '✅ PASS' if count == expected else f'❌ FAIL (期望{expected})'
                if count != expected:
                    all_passed = False
            else:
                status = '📊 检查中'

            print(f'{name}: {count}条 {status}')

            # 如果测试失败或是检查文档，显示部分条款
            if (expected and count != expected) or expected is None:
                print(f'  前15条条款:')
                for i, c in enumerate(clauses[:15]):
                    title_display = c.title[:60] + '...' if len(c.title) > 60 else c.title
                    print(f'    {i+1}. {title_display}')
                if count > 15:
                    print(f'    ... 共{count}条')

        except Exception as e:
            print(f'{name}: ❌ 错误 - {e}')
            import traceback
            traceback.print_exc()
            all_passed = False

    print('=' * 60)
    if all_passed:
        print('✅ 所有测试通过!')
    else:
        print('❌ 部分测试失败，请检查')
    print('=' * 60)


def main():
    if hasattr(Qt, 'AA_EnableHighDpiScaling'):
        QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
        QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)

    app = QApplication(sys.argv)
    app.setFont(QFont("Microsoft YaHei", 10))

    window = ClauseComparisonAssistant()
    window.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    # 支持命令行测试模式
    if len(sys.argv) > 1 and sys.argv[1] == '--test':
        import glob as glob_module
        # 创建最小化的QApplication（某些类可能需要）
        app = QApplication(sys.argv)

        # 使用glob查找文件（处理特殊空格字符）
        def find_file(pattern):
            matches = glob_module.glob(pattern)
            return matches[0] if matches else pattern

        test_docs = [
            (find_file('/Volumes/4TB-Samsung/works/Trammo*Quote*.docx'), 27, 'Trammo'),
            (find_file('/Volumes/4TB-Samsung/works/广州天河林和君Hilton*.docx'), 128, 'Hilton'),
            (find_file('/Volumes/4TB-Samsung/works/马勒投资*PAR.docx'), 63, '马勒PAR'),
            (find_file('/Volumes/4TB-Samsung/works/*梅花*Quotation*.docx'), None, '梅花'),
        ]
        run_clause_test(test_docs)
    elif len(sys.argv) > 1 and sys.argv[1] == '--parse':
        # 解析指定文件
        import glob as glob_module
        app = QApplication(sys.argv)

        if len(sys.argv) < 3:
            print('用法: python Clause_Comparison_Assistant.py --parse <文件路径>')
            sys.exit(1)

        file_path = sys.argv[2]
        # 处理glob模式
        if '*' in file_path:
            matches = glob_module.glob(file_path)
            if matches:
                file_path = matches[0]

        matcher = ClauseMatcherLogic()
        try:
            clauses, is_title_only = matcher.parse_docx(file_path)
            print(f'\n文件: {file_path}')
            print(f'识别到 {len(clauses)} 条条款 (仅标题模式: {is_title_only})')
            print('-' * 60)
            for i, c in enumerate(clauses):
                print(f'{i+1}. {c.title}')
        except Exception as e:
            print(f'错误: {e}')
            import traceback
            traceback.print_exc()
    else:
        main()
